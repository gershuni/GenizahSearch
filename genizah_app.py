"""PyQt6 GUI for Genizah search and browsing."""

# genizah_app.py
import sys
import os
import re
import threading
import json
import requests
import urllib3
import csv
import openpyxl
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.cell.rich_text import TextBlock, CellRichText
from openpyxl.cell.text import InlineFont

from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
                             QLabel, QLineEdit, QPushButton, QTabWidget, QTableWidget,
                             QTableWidgetItem, QListWidgetItem, QHeaderView, QComboBox, QCheckBox,
                             QTextEdit, QMessageBox, QProgressBar, QSplitter, QDialog,
                             QTextBrowser, QFileDialog, QMenu, QGroupBox, QSpinBox, QDoubleSpinBox,
                             QTreeWidget, QTreeWidgetItem, QListWidget, QPlainTextEdit, QStyle, QFormLayout,
                             QGridLayout, QToolTip, QProgressDialog, QStackedLayout,
                             QScrollArea, QFrame, QSlider, QStyleOptionButton, QSizePolicy, QInputDialog,
                             QToolButton, QGraphicsView, QGraphicsScene, QGraphicsPixmapItem, QGraphicsSimpleTextItem,
                             QCompleter, QAbstractItemView)
from PyQt6.QtCore import (Qt, QTimer, QUrl, QSize, pyqtSignal, QThread, QEventLoop, QEvent, QRect, QRectF)
from PyQt6.QtGui import (QFont, QIcon, QDesktopServices, QPixmap, QImage, QFontMetrics, QTextDocument, QTransform, QPainter, QColor,
                         QStandardItemModel, QStandardItem, QPalette, QTextCursor, QTextCharFormat, QPen, QBrush, QPainterPath, QCursor)
from PyQt6 import sip

from version import APP_VERSION

from collections import defaultdict

_CORE_IMPORT_ERROR = None
try:
    from genizah_core import Config, MetadataManager, VariantManager, SearchEngine, LabEngine, Indexer, ListsManager, JoinsManager, tr, save_language, CURRENT_LANG, get_logger, natural_sort_key, load_app_config, save_app_config, get_library_display, normalize_shelfmark, generate_tabular_syntax
except ImportError as import_error:
    _CORE_IMPORT_ERROR = import_error

if _CORE_IMPORT_ERROR:
    def _show_core_import_error(err):
        app = QApplication.instance() or QApplication(sys.argv)
        QMessageBox.critical(None, "Missing dependency", str(err))

    if __name__ == "__main__":
        _show_core_import_error(_CORE_IMPORT_ERROR)
        sys.exit(1)
    else:
        raise _CORE_IMPORT_ERROR
from gui_threads import SearchThread, LabSearchThread, IndexerThread, ShelfmarkLoaderThread, CompositionThread, LabCompositionThread, GroupingThread, StartupThread, EnrichMetadataThread, ExternalResourceThread, UpdateCheckerThread, PGPSourceWorker, ReadingDeskWorker, PGPBadgeWorker, PGPTagsWorker, PGPTagSearchWorker, SidecarUpdateThread, SidecarDownloadThread
from filter_text_dialog import FilterTextDialog
from column_filter_dialog import ColumnFilterDialog
from list_filter_dialog import ListFilterDialog
from shared_export_utils import sanitize_text_for_excel as shared_sanitize_excel
from shared.reading_desk_model import ReadingDeskEntry, ReadingDeskState

# NLI crossref service for folio labels and source indicators (Phase 31)
try:
    from shared.nli_crossref_service import parse_folio_label, get_nli_crossref_service
    _HAS_NLI_CROSSREF = True
except ImportError:
    _HAS_NLI_CROSSREF = False

# Community features - corrections, comments, discoveries
from corrections_client import get_corrections_client
from corrections_ui import (
    LoginDialog, RegisterDialog,
    CorrectionSubmitDialog, CorrectionsViewerDialog, CorrectionDetailDialog,
    MyCorrectionsDialog, AllCorrectionsDialog,
    CommentDialog, CommentsViewerDialog, MyCommentsDialog,
    DiscoveriesDialog, CreateDiscoveryDialog, DiscoveryDetailDialog,
    TextEditorDialog, JoinsDialog
)

logger = get_logger(__name__)

# Global exception handler to log crashes to file
def _setup_crash_handler():
    import traceback
    from datetime import datetime

    def exception_hook(exc_type, exc_value, exc_tb):
        # Log to file
        try:
            crash_log = os.path.join(os.path.dirname(__file__), 'crash_log.txt')
            with open(crash_log, 'a', encoding='utf-8') as f:
                f.write(f"\n{'='*60}\n")
                f.write(f"Crash at {datetime.now().isoformat()}\n")
                f.write(''.join(traceback.format_exception(exc_type, exc_value, exc_tb)))
        except:
            pass
        # Also print to console
        traceback.print_exception(exc_type, exc_value, exc_tb)
        # Call default handler
        sys.__excepthook__(exc_type, exc_value, exc_tb)

    sys.excepthook = exception_hook

_setup_crash_handler()

def apply_find_highlight(text_browser, query):
    if not text_browser:
        return
    if not query:
        text_browser.setExtraSelections([])
        return
    doc = text_browser.document()
    cursor = QTextCursor(doc)
    highlight_format = QTextCharFormat()
    highlight_format.setBackground(QColor("#fff59d"))
    selections = []
    while True:
        cursor = doc.find(query, cursor)
        if cursor.isNull():
            break
        selection = QTextEdit.ExtraSelection()
        selection.cursor = cursor
        selection.format = highlight_format
        selections.append(selection)
    text_browser.setExtraSelections(selections)

def _generate_oxford_dynamic_url(oxford_part_id, folio_num, side='a'):
    """Generate dynamic Oxford image URL for a folio not in the database.

    Args:
        oxford_part_id: Part ID like "MS. Heb. f. 21/1"
        folio_num: Folio number (e.g., 21)
        side: 'a' for recto, 'b' for verso

    Returns URL string or None if generation not possible.
    """
    if not oxford_part_id:
        return None

    match = re.match(r'^MS\.?\s*Heb\.?\s*([a-z])\.?\s*(\d+)', oxford_part_id, re.IGNORECASE)
    if not match:
        return None

    letter, volume = match.groups()
    return f"https://hebrew.bodleian.ox.ac.uk/fragments/full/MS_HEB_{letter}_{volume}_{folio_num}{side}.jpg"


def _get_initial_image_index(meta, page_num):
    if page_num is None:
        return 0
    try:
        p_num = int(page_num)
    except (TypeError, ValueError):
        return 0

    images = (meta or {}).get('images_ext') or (meta or {}).get('images') or []
    folio_entries = []
    for idx, img in enumerate(images):
        folio_num = img.get('folio_num')
        if folio_num is None:
            continue
        try:
            folio_entries.append((idx, int(folio_num)))
        except (TypeError, ValueError):
            continue

    if not folio_entries:
        return max(p_num - 1, 0)

    for idx, folio_num in folio_entries:
        if folio_num == p_num:
            return idx

    prior = [(idx, folio_num) for idx, folio_num in folio_entries if folio_num <= p_num]
    if prior:
        return max(prior, key=lambda pair: pair[1])[0]

    return min(folio_entries, key=lambda pair: pair[1])[0]

def _get_folio_number_from_shelfmark(shelfmark):
    """Extract folio number from Oxford-style shelfmarks only.

    Oxford shelfmarks like "MS. Heb. a. 1/1" or "Bodl. Or. 12/3" contain
    actual folio numbers after the slash. Other libraries (Cambridge, NLI, etc.)
    use classmarks where trailing numbers are not folio references.
    """
    if not shelfmark:
        return None
    upper = shelfmark.upper()
    # Only extract folio from Oxford-style shelfmarks (MS. Heb., Bodl., etc.)
    is_oxford = (
        'MS. HEB' in upper or
        'MS HEB' in upper or
        upper.startswith('BODL') or
        'BODLEIAN' in upper
    )
    if not is_oxford:
        return None
    match = re.search(r'[/.](\d+)\s*$', shelfmark)
    if not match:
        return None
    try:
        return int(match.group(1))
    except (TypeError, ValueError):
        return None

def _get_folio_image_index(meta, folio_num, side_offset=0):
    base_idx = _get_initial_image_index(meta, folio_num)
    if side_offset <= 0:
        return base_idx

    images = (meta or {}).get('images_ext') or (meta or {}).get('images') or []
    if not images or base_idx >= len(images):
        return base_idx

    target_folio = images[base_idx].get('folio_num')
    if target_folio is None:
        return base_idx

    label = str(images[base_idx].get('label', '')).lower()
    if label.endswith('b'):
        return base_idx

    next_idx = base_idx + 1
    if next_idx < len(images) and images[next_idx].get('folio_num') == target_folio:
        return next_idx

    for idx, img in enumerate(images):
        if img.get('folio_num') == target_folio and str(img.get('label', '')).lower().endswith('b'):
            return idx

    return base_idx

class UpdateNotificationBar(QFrame):
    """A narrow notification bar at the top of the screen."""

    dismissed = pyqtSignal(str) # Emits version string on dismiss
    update_requested = pyqtSignal(str, str, str)  # version, html_url, installer_url

    def __init__(self, parent=None):
        super().__init__(parent)
        self.html_url = ""
        self.installer_url = ""
        self.version_tag = ""

        self.setFrameShape(QFrame.Shape.StyledPanel)
        self.setStyleSheet("background-color: #d1ecf1; color: #0c5460; border-bottom: 1px solid #bee5eb;")
        self.setFixedHeight(40)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(10, 0, 10, 0)

        self.lbl_msg = QLabel()
        self.lbl_msg.setStyleSheet("font-weight: bold; font-size: 13px; border: none; background: transparent;")

        self.btn_download = QPushButton(tr("Update Now"))
        self.btn_download.setStyleSheet("background-color: #17a2b8; color: white; font-weight: bold; border-radius: 4px; padding: 4px 8px;")
        self.btn_download.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_download.clicked.connect(self.on_download)

        self.btn_dismiss = QPushButton("✕")
        self.btn_dismiss.setToolTip(tr("Dismiss until next version"))
        self.btn_dismiss.setStyleSheet("""
            QPushButton { background: transparent; color: #0c5460; font-weight: bold; border: none; font-size: 16px; }
            QPushButton:hover { color: #dc3545; }
        """)
        self.btn_dismiss.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_dismiss.clicked.connect(self.on_dismiss)

        layout.addWidget(self.lbl_msg)
        layout.addStretch()
        layout.addWidget(self.btn_download)
        layout.addSpacing(10)
        layout.addWidget(self.btn_dismiss)

        self.hide()

    def show_update(self, version: str, html_url: str, installer_url: str = ""):
        self.version_tag = version
        self.html_url = html_url
        self.installer_url = installer_url
        self.lbl_msg.setText(tr("New version available: {}").format(version))
        self.show()

    def on_download(self):
        self.update_requested.emit(self.version_tag, self.html_url, self.installer_url)

    def on_dismiss(self):
        self.hide()
        self.dismissed.emit(self.version_tag)


class WhatsNewBar(QFrame):
    """A notification bar showing new features after a version update."""

    dismissed = pyqtSignal()
    learn_more = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFrameShape(QFrame.Shape.StyledPanel)
        self.setStyleSheet("background-color: #d1fae5; color: #065f46; border-bottom: 1px solid #a7f3d0;")
        self.setFixedHeight(40)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(10, 0, 10, 0)

        self.lbl_msg = QLabel()
        self.lbl_msg.setStyleSheet("font-weight: bold; font-size: 13px; border: none; background: transparent;")

        self.btn_learn_more = QPushButton(tr("Learn More"))
        self.btn_learn_more.setStyleSheet("background-color: #10b981; color: white; font-weight: bold; border-radius: 4px; padding: 4px 8px;")
        self.btn_learn_more.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_learn_more.clicked.connect(self.on_learn_more)

        self.btn_dismiss = QPushButton("\u2715")
        self.btn_dismiss.setToolTip(tr("Dismiss"))
        self.btn_dismiss.setStyleSheet("""
            QPushButton { background: transparent; color: #065f46; font-weight: bold; border: none; font-size: 16px; }
            QPushButton:hover { color: #dc3545; }
        """)
        self.btn_dismiss.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_dismiss.clicked.connect(self.on_dismiss)

        layout.addWidget(self.lbl_msg)
        layout.addStretch()
        layout.addWidget(self.btn_learn_more)
        layout.addSpacing(10)
        layout.addWidget(self.btn_dismiss)

        self.hide()

    def show_whats_new(self, version: str):
        self.lbl_msg.setText(tr("New: Responsa-Project style search, PGP corpus (35K documents), and FJMS scholarly metadata"))
        self.show()

    def on_learn_more(self):
        self.learn_more.emit()

    def on_dismiss(self):
        self.hide()
        self.dismissed.emit()


class WhatsNewDialog(QDialog):
    """Dialog showing detailed What's New information."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle(tr("New Features!"))
        self.setModal(True)
        self.setFixedSize(500, 380)
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(16)

        title = QLabel(tr("New Features!"))
        title.setStyleSheet("font-size: 18px; font-weight: bold; color: #065f46;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)

        features_html = (
            "<ul dir='rtl' style='font-size: 14px; line-height: 1.8; text-align: right;'>"
            f"<li><b>{tr('Responsa-Project style Search: advanced syntax parsing with grammatical expansion, Judeo-Arabic support, and tabular query builder')}</b></li>"
            f"<li><b>{tr('PGP Integration: 35,000 documents from the Princeton Geniza Project with editions, translations, and more information')}</b></li>"
            f"<li><b>{tr('FJMS scholarly metadata: 390K domain classifications, 48K scientific joins, and 500K catalog records from the Friedberg Genizah Project')}</b></li>"
            "</ul>"
        )
        features_label = QLabel(features_html)
        features_label.setWordWrap(True)
        features_label.setTextFormat(Qt.TextFormat.RichText)
        features_label.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
        layout.addWidget(features_label)

        layout.addStretch()

        btn_ok = QPushButton(tr("Got it!"))
        btn_ok.setStyleSheet("background-color: #10b981; color: white; font-weight: bold; border-radius: 4px; padding: 8px 24px; font-size: 14px;")
        btn_ok.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_ok.clicked.connect(self.accept)
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        btn_layout.addWidget(btn_ok)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)


class UpdateProgressDialog(QDialog):
    """Shows download progress and handles update installation."""

    def __init__(self, parent, version: str, installer_url: str, html_url: str):
        super().__init__(parent)
        self.version = version
        self.installer_url = installer_url
        self.html_url = html_url
        self.download_thread = None
        self.downloaded_path = None

        self.setWindowTitle(tr("Updating GenizahSearch"))
        self.setModal(True)
        self.setFixedSize(420, 180)
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(12)

        # Title
        title_label = QLabel(tr("Updating to version {}").format(version))
        title_label.setStyleSheet("font-size: 14px; font-weight: bold;")
        layout.addWidget(title_label)

        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setMinimum(0)
        self.progress_bar.setMaximum(100)
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(True)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 1px solid #ccc;
                border-radius: 4px;
                text-align: center;
                height: 24px;
            }
            QProgressBar::chunk {
                background-color: #17a2b8;
                border-radius: 3px;
            }
        """)
        layout.addWidget(self.progress_bar)

        # Status label
        self.status_label = QLabel(tr("Preparing download..."))
        self.status_label.setStyleSheet("color: #666;")
        layout.addWidget(self.status_label)

        layout.addStretch()

        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()

        self.btn_cancel = QPushButton(tr("Cancel"))
        self.btn_cancel.clicked.connect(self.on_cancel)
        btn_layout.addWidget(self.btn_cancel)

        layout.addLayout(btn_layout)

    def start_download(self):
        """Start the download process."""
        import tempfile
        import os

        if not self.installer_url:
            QMessageBox.warning(
                self, tr("Download Error"),
                tr("No direct download available. Opening browser instead...")
            )
            QDesktopServices.openUrl(QUrl(self.html_url))
            self.reject()
            return

        # Determine download path
        temp_dir = tempfile.gettempdir()
        safe_version = self.version.replace('/', '_').replace('\\', '_')
        target_path = os.path.join(temp_dir, f"GenizahSearchPro_{safe_version}_Setup.exe")

        # Import and start download thread
        from gui_threads import UpdateDownloaderThread
        self.download_thread = UpdateDownloaderThread(self.installer_url, target_path)
        self.download_thread.progress_signal.connect(self.on_progress)
        self.download_thread.finished_signal.connect(self.on_download_finished)
        self.download_thread.start()

        self.status_label.setText(tr("Downloading..."))

    def on_progress(self, downloaded: int, total: int):
        """Update progress bar with download progress."""
        if total > 0:
            percent = int((downloaded / total) * 100)
            self.progress_bar.setValue(percent)

            # Format size display
            downloaded_mb = downloaded / (1024 * 1024)
            total_mb = total / (1024 * 1024)
            self.status_label.setText(
                tr("Downloading: {:.1f} MB / {:.1f} MB").format(downloaded_mb, total_mb)
            )
        else:
            # Unknown total size - show indeterminate
            self.progress_bar.setMaximum(0)
            downloaded_mb = downloaded / (1024 * 1024)
            self.status_label.setText(tr("Downloaded: {:.1f} MB").format(downloaded_mb))

    def on_download_finished(self, success: bool, result: str):
        """Handle download completion."""
        if success:
            self.downloaded_path = result
            self.progress_bar.setValue(100)
            self.status_label.setText(tr("Download complete. Installing update..."))
            self.btn_cancel.setEnabled(False)

            # Execute the update
            QTimer.singleShot(500, self.execute_update)
        else:
            # Download failed
            self.progress_bar.setValue(0)
            self.status_label.setText(tr("Download failed"))

            reply = QMessageBox.question(
                self, tr("Download Failed"),
                tr("Failed to download update: {}").format(result) + "\n\n" +
                tr("Would you like to download manually from the website?"),
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                QDesktopServices.openUrl(QUrl(self.html_url))

            self.reject()

    def execute_update(self):
        """Run the installer in silent mode (Windows only)."""
        import subprocess
        import sys
        import os

        # Check platform
        if sys.platform != 'win32':
            QMessageBox.information(
                self, tr("Update Ready"),
                tr("The update has been downloaded to:") + f"\n{self.downloaded_path}\n\n" +
                tr("Please run the installer manually.")
            )
            self.accept()
            return

        # Check if running as compiled executable
        if not getattr(sys, 'frozen', False):
            # Running from Python (development mode)
            QMessageBox.information(
                self, tr("Development Mode"),
                tr("Auto-update is not available in development mode.") + "\n" +
                tr("The installer has been downloaded to:") + f"\n{self.downloaded_path}"
            )
            self.accept()
            return

        # Update status
        self.status_label.setText(tr("Launching installer..."))

        # Run the installer with silent mode
        # The installer will:
        # 1. Close this running app (CloseApplications=force)
        # 2. Install the update
        # 3. Restart the app (RestartApplications=yes)
        try:
            subprocess.Popen(
                [self.downloaded_path, '/VERYSILENT', '/RESTARTAPPLICATIONS'],
                creationflags=subprocess.DETACHED_PROCESS
            )
        except Exception as e:
            QMessageBox.critical(
                self, tr("Update Error"),
                tr("Could not start installer: {}").format(str(e))
            )
            self.reject()
            return

        # Close the dialog and quit the application
        # The installer will handle closing us, but we quit gracefully to speed things up
        self.accept()
        QApplication.quit()

    def on_cancel(self):
        """Handle cancel button click."""
        if self.download_thread and self.download_thread.isRunning():
            reply = QMessageBox.question(
                self, tr("Cancel Download"),
                tr("Are you sure you want to cancel the update?"),
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                self.download_thread.cancel()
                self.status_label.setText(tr("Cancelling..."))
                self.btn_cancel.setEnabled(False)
        else:
            self.reject()

    def closeEvent(self, event):
        """Handle dialog close event."""
        if self.download_thread and self.download_thread.isRunning():
            reply = QMessageBox.question(
                self, tr("Cancel Download"),
                tr("Download is in progress. Cancel and close?"),
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                self.download_thread.cancel()
                self.download_thread.wait(3000)  # Wait up to 3 seconds
                event.accept()
            else:
                event.ignore()
        else:
            event.accept()


BATCH_SIZE = 500

class LabScoringDialog(QDialog):
    """Configuration for Lab Mode Scoring (Advanced)."""
    def __init__(self, parent, lab_engine):
        super().__init__(parent)
        self.setWindowTitle(tr("Advanced Scoring"))
        self.resize(500, 500)
        self.lab_engine = lab_engine
        self.settings = lab_engine.settings
        if CURRENT_LANG == 'he':
            self.setLayoutDirection(Qt.LayoutDirection.RightToLeft)

        layout = QVBoxLayout()
        layout.addWidget(QLabel(tr("Adjust how the algorithm prioritizes results.")))
        
        grid = QGridLayout()
        
        # Order Bonus
        self.spin_order_bonus = QDoubleSpinBox(); self.spin_order_bonus.setRange(0.0, 100.0); self.spin_order_bonus.setSingleStep(1.0); self.spin_order_bonus.setValue(getattr(self.settings, 'order_bonus', 10.0))
        lbl_order = QLabel(tr("Sequential Order Bonus:")); lbl_order.setStyleSheet("color: #2980b9; font-weight: bold;")
        grid.addWidget(lbl_order, 0, 0); grid.addWidget(self.spin_order_bonus, 0, 1)

        # Coverage
        self.spin_coverage_power = QDoubleSpinBox(); self.spin_coverage_power.setRange(1.0, 10.0); self.spin_coverage_power.setValue(self.settings.coverage_power)
        grid.addWidget(QLabel(tr("Coverage Penalty Power:")), 1, 0); grid.addWidget(self.spin_coverage_power, 1, 1)

        # Noise Suppression
        lbl_noise = QLabel(tr("Stop-Word Suppression:")); lbl_noise.setStyleSheet("font-weight: bold; margin-top: 10px;")
        grid.addWidget(lbl_noise, 2, 0, 1, 2)

        # Short Word Score
        self.spin_stop_score = QDoubleSpinBox(); self.spin_stop_score.setRange(0.0, 50.0); self.spin_stop_score.setSingleStep(0.5); self.spin_stop_score.setValue(getattr(self.settings, 'stop_word_score', 1.0))
        self.spin_stop_score.setToolTip(tr("Points given for very short words (<3 letters). Keep low to reduce noise."))
        grid.addWidget(QLabel(tr("Score for Short Words (<3):")), 3, 0); grid.addWidget(self.spin_stop_score, 3, 1)

        # Common 3-Char Score
        self.spin_common3_score = QDoubleSpinBox(); self.spin_common3_score.setRange(0.0, 50.0); self.spin_common3_score.setSingleStep(0.5); self.spin_common3_score.setValue(getattr(self.settings, 'common_3char_score', 2.0))
        self.spin_common3_score.setToolTip(tr("Points for common 3-letter words (e.g. 'ליה', 'הכי')."))
        grid.addWidget(QLabel(tr("Score for Common 3-Letter:")), 4, 0); grid.addWidget(self.spin_common3_score, 4, 1)

        # Other Weights
        lbl_other = QLabel(tr("Standard Weights:")); lbl_other.setStyleSheet("font-weight: bold; margin-top: 10px;")
        grid.addWidget(lbl_other, 5, 0, 1, 2)

        self.spin_len_bonus = QDoubleSpinBox(); self.spin_len_bonus.setRange(1.0, 10.0); self.spin_len_bonus.setValue(self.settings.length_bonus_factor)
        grid.addWidget(QLabel(tr("Long Word Bonus:")), 6, 0); grid.addWidget(self.spin_len_bonus, 6, 1)

        self.spin_unique_base = QSpinBox(); self.spin_unique_base.setRange(10, 1000); self.spin_unique_base.setValue(self.settings.unique_bonus_base)
        grid.addWidget(QLabel(tr("Unique Match Base Score:")), 7, 0); grid.addWidget(self.spin_unique_base, 7, 1)

        self.spin_density = QDoubleSpinBox(); self.spin_density.setRange(0.0, 5.0); self.spin_density.setValue(self.settings.density_penalty)
        grid.addWidget(QLabel(tr("Distance Penalty:")), 8, 0); grid.addWidget(self.spin_density, 8, 1)

        self.spin_common_factor = QDoubleSpinBox(); self.spin_common_factor.setRange(0.0, 1.0); self.spin_common_factor.setValue(self.settings.common_penalty_factor)
        grid.addWidget(QLabel(tr("Repeated Word Factor:")), 9, 0); grid.addWidget(self.spin_common_factor, 9, 1)

        # Display Limit
        self.spin_display_limit = QSpinBox(); self.spin_display_limit.setRange(50, 1000); self.spin_display_limit.setValue(getattr(self.settings, 'lab_display_limit', 500))
        self.spin_display_limit.setToolTip(tr("Lower values prevent the app from freezing. All results are still exported."))
        grid.addWidget(QLabel(tr("Max Results to Display:")), 10, 0); grid.addWidget(self.spin_display_limit, 10, 1)

        layout.addLayout(grid)
        layout.addStretch()
        
        btn_box = QHBoxLayout()
        # Help Button
        btn_help = QPushButton("?")
        btn_help.setFixedWidth(30)
        btn_help.setStyleSheet("background-color: #f39c12; color: white; font-weight: bold; border-radius: 15px;")
        # Find main window to call open_help_center
        def open_help():
            main = parent
            while main and not hasattr(main, 'open_help_center'):
                main = main.parent()
            if main: main.open_help_center(anchor="lab")

        btn_help.clicked.connect(open_help)
        btn_box.addWidget(btn_help)

        btn_box.addStretch()
        self.btn_save = QPushButton(tr("Save & Close")); self.btn_save.clicked.connect(self.save_and_close)
        self.btn_cancel = QPushButton(tr("Cancel")); self.btn_cancel.clicked.connect(self.reject)
        btn_box.addStretch(); btn_box.addWidget(self.btn_cancel); btn_box.addWidget(self.btn_save)
        layout.addLayout(btn_box)
        self.setLayout(layout)

    def save_and_close(self):
        self.settings.coverage_power = self.spin_coverage_power.value()
        self.settings.length_bonus_factor = self.spin_len_bonus.value()
        self.settings.common_penalty_factor = self.spin_common_factor.value()
        self.settings.density_penalty = self.spin_density.value()
        self.settings.unique_bonus_base = self.spin_unique_base.value()
        if hasattr(self.settings, 'order_bonus'): self.settings.order_bonus = self.spin_order_bonus.value()
        if hasattr(self.settings, 'stop_word_score'):
            self.settings.stop_word_score = self.spin_stop_score.value()
            self.settings.common_3char_score = self.spin_common3_score.value()

        self.settings.lab_display_limit = self.spin_display_limit.value()
        self.settings.save()
        self.accept()


class SearchSettingsDialog(QDialog):
    """Settings for Standard Search - Variant configuration and custom pairs."""
    def __init__(self, parent, settings):
        super().__init__(parent)
        self.setWindowTitle(tr("Search Settings"))
        self.resize(450, 400)
        self.settings = settings
        if CURRENT_LANG == 'he':
            self.setLayoutDirection(Qt.LayoutDirection.RightToLeft)

        layout = QVBoxLayout()
        layout.addWidget(QLabel(tr("Configure variant search behavior for Standard Search modes.")))

        grid = QGridLayout()

        # --- Variant Limits Section ---
        lbl_variant = QLabel(tr("Variant Search Limits:"))
        lbl_variant.setStyleSheet("font-weight: bold; margin-top: 10px; color: #8e44ad;")
        grid.addWidget(lbl_variant, 0, 0, 1, 2)

        # Min Word Length for limiting changes
        self.spin_variant_min_len = QSpinBox()
        self.spin_variant_min_len.setRange(1, 5)
        self.spin_variant_min_len.setValue(getattr(self.settings, 'variant_min_word_len', 2))
        self.spin_variant_min_len.setToolTip(tr("Words with this length or less get only 1 character change. Increase to be more conservative."))
        grid.addWidget(QLabel(tr("Limit Short Words (≤N chars):")), 1, 0)
        grid.addWidget(self.spin_variant_min_len, 1, 1)

        # Max Changes
        self.spin_variant_max_changes = QSpinBox()
        self.spin_variant_max_changes.setRange(1, 3)
        self.spin_variant_max_changes.setValue(getattr(self.settings, 'variant_max_changes', 2))
        self.spin_variant_max_changes.setToolTip(tr("Maximum character substitutions per word. Higher = more results but slower."))
        grid.addWidget(QLabel(tr("Max Changes per Word:")), 2, 0)
        grid.addWidget(self.spin_variant_max_changes, 2, 1)

        # Aggressive Mode
        self.chk_variant_aggressive = QCheckBox(tr("Aggressive Mode (ignore word length limits)"))
        self.chk_variant_aggressive.setChecked(getattr(self.settings, 'variant_aggressive', False))
        self.chk_variant_aggressive.setToolTip(tr("Like old behavior: apply max changes to all words regardless of length. More results, more noise."))
        grid.addWidget(self.chk_variant_aggressive, 3, 0, 1, 2)

        # Use slider instead of presets
        self.chk_use_slider = QCheckBox(tr("Use slider instead of preset buttons (Basic, Extended, Maximum)"))
        self.chk_use_slider.setChecked(getattr(self.settings, 'variant_use_slider', False))
        self.chk_use_slider.setToolTip(tr("When enabled, shows a slider in the search bar instead of preset buttons"))
        grid.addWidget(self.chk_use_slider, 4, 0, 1, 2)

        # --- Variant Pairs Slider (shown only when slider mode is enabled) ---
        self.slider_container = QWidget()
        slider_container_layout = QVBoxLayout(self.slider_container)
        slider_container_layout.setContentsMargins(0, 0, 0, 0)

        lbl_pairs = QLabel(tr("Variant Pairs Level:"))
        lbl_pairs.setStyleSheet("font-weight: bold; margin-top: 10px; color: #2980b9;")
        slider_container_layout.addWidget(lbl_pairs)

        # Slider for number of variant pairs to use
        slider_layout = QHBoxLayout()
        self.slider_variant_pairs = QSlider(Qt.Orientation.Horizontal)
        self.slider_variant_pairs.setRange(10, 300)
        self.slider_variant_pairs.setValue(getattr(self.settings, 'variant_pairs_count', 70))
        self.slider_variant_pairs.setToolTip(tr("Number of variant pairs to use. Higher = more substitutions but slower search.\nBased on frequency: top pairs are most common HTR confusions."))

        self.lbl_pairs_value = QLabel(str(self.slider_variant_pairs.value()))
        self.lbl_pairs_value.setMinimumWidth(40)
        self.slider_variant_pairs.valueChanged.connect(
            lambda v: self.lbl_pairs_value.setText(str(v))
        )

        slider_layout.addWidget(QLabel(tr("10")))
        slider_layout.addWidget(self.slider_variant_pairs)
        slider_layout.addWidget(QLabel(tr("300")))
        slider_layout.addWidget(self.lbl_pairs_value)
        slider_container_layout.addLayout(slider_layout)

        lbl_pairs_help = QLabel(tr("Controls how many character substitution pairs to use. Higher values find more variants but are slower."))
        lbl_pairs_help.setStyleSheet("font-size: 10px; color: gray; font-style: italic;")
        lbl_pairs_help.setWordWrap(True)
        slider_container_layout.addWidget(lbl_pairs_help)

        grid.addWidget(self.slider_container, 5, 0, 1, 2)

        # Show/hide slider container based on checkbox
        self.slider_container.setVisible(self.chk_use_slider.isChecked())
        self.chk_use_slider.toggled.connect(self.slider_container.setVisible)

        layout.addLayout(grid)

        # --- Custom Variants Section ---
        lbl_custom = QLabel(tr("Custom Variant Pairs:"))
        lbl_custom.setStyleSheet("font-weight: bold; margin-top: 15px; color: #27ae60;")
        layout.addWidget(lbl_custom)

        lbl_custom_help = QLabel(tr("Add character pairs that should be treated as interchangeable (e.g. ק=א means ק↔א)."))
        lbl_custom_help.setStyleSheet("font-size: 10px; color: gray; font-style: italic;")
        lbl_custom_help.setWordWrap(True)
        layout.addWidget(lbl_custom_help)

        # Custom variants text edit
        self.txt_custom_variants = QTextEdit()
        self.txt_custom_variants.setPlaceholderText(tr("Enter one pair per line:\nק=א\nכו=מ\nב=פ"))
        self.txt_custom_variants.setMaximumHeight(120)

        # Load existing custom variants
        custom = getattr(self.settings, 'custom_variants', {})
        if custom:
            lines = [k for k in custom.keys()]
            self.txt_custom_variants.setPlainText('\n'.join(lines))

        layout.addWidget(self.txt_custom_variants)

        layout.addStretch()

        # Buttons
        btn_box = QHBoxLayout()
        btn_box.addStretch()
        self.btn_cancel = QPushButton(tr("Cancel"))
        self.btn_cancel.clicked.connect(self.reject)
        self.btn_save = QPushButton(tr("Save & Close"))
        self.btn_save.clicked.connect(self.save_and_close)
        btn_box.addWidget(self.btn_cancel)
        btn_box.addWidget(self.btn_save)
        layout.addLayout(btn_box)
        self.setLayout(layout)

    def save_and_close(self):
        # Save variant limits
        self.settings.variant_min_word_len = self.spin_variant_min_len.value()
        self.settings.variant_max_changes = self.spin_variant_max_changes.value()
        self.settings.variant_aggressive = self.chk_variant_aggressive.isChecked()
        self.settings.variant_pairs_count = self.slider_variant_pairs.value()
        self.settings.variant_use_slider = self.chk_use_slider.isChecked()

        # Parse custom variants
        text = self.txt_custom_variants.toPlainText().strip()
        custom = {}
        if text:
            for line in text.split('\n'):
                line = line.strip()
                if '=' in line:
                    custom[line] = True
        self.settings.custom_variants = custom

        self.settings.save()

        # Update VariantManager if available
        main = self.parent()
        while main and not hasattr(main, 'var_mgr'):
            main = main.parent()
        if main and main.var_mgr:
            main.var_mgr.set_settings(self.settings)

        self.accept()


class LabPanel(QFrame):
    def __init__(self, parent, mode):
        super().__init__(parent)
        self.mode = mode
        self.lab_engine = None
        self.settings = None
        self.setFrameShape(QFrame.Shape.StyledPanel)
        self.setFrameShadow(QFrame.Shadow.Raised)
        # Removed hardcoded background color to support both light and dark themes
        self.setStyleSheet("border-radius: 5px; margin-bottom: 5px;")

        self.init_ui()
        self.setVisible(False)

    def set_engine(self, engine):
        self.lab_engine = engine
        self.settings = engine.settings
        self.refresh_values()
        self.enable_controls(True)
        self._mark_rebuild_required()

    def enable_controls(self, enabled):
        self.setEnabled(enabled)

    def init_ui(self):
        self.layout = QHBoxLayout(self)
        self.layout.setContentsMargins(10, 5, 10, 5)

        # Shared: Rebuild
        self.btn_rebuild = QPushButton(tr("Rebuild Lab Index"))
        self.btn_rebuild.setStyleSheet("background-color: #d35400; color: white; font-weight: bold; border-radius: 4px; padding: 4px;")
        self.btn_rebuild.clicked.connect(self.run_rebuild)
        self.layout.addWidget(self.btn_rebuild)

        self.lbl_idx_status = QLabel("")
        self.layout.addWidget(self.lbl_idx_status)

        # Help Button (Shared)
        btn_help = QPushButton("?")
        btn_help.setFixedWidth(24)
        btn_help.setStyleSheet("background-color: #f39c12; color: white; font-weight: bold; border-radius: 12px;")

        def open_help_panel():
            main = self.parent()
            while main and not hasattr(main, 'open_help_center'):
                main = main.parent()
            if main: main.open_help_center(anchor="lab")

        btn_help.clicked.connect(open_help_panel)
        self.layout.addWidget(btn_help)

        # Spacer
        self.layout.addSpacing(20)
        self.layout.addWidget(QLabel("|"))
        self.layout.addSpacing(20)

        if self.mode == 'search':
            # Min Match
            self.layout.addWidget(QLabel(tr("Minimum Match %:")))
            self.spin_min = QSpinBox()
            self.spin_min.setToolTip(tr("Minimum percentage of query terms required to consider a result relevant."))
            self.spin_min.setRange(10, 100)
            self.spin_min.setSuffix("%")
            self.spin_min.valueChanged.connect(self.on_change)
            self.layout.addWidget(self.spin_min)

            # Candidate Limit
            self.layout.addWidget(QLabel(tr("Max Results to Process:")))
            self.spin_limit = QSpinBox()
            self.spin_limit.setToolTip(tr("Maximum number of raw candidates to fetch from the index before detailed scoring."))
            self.spin_limit.setRange(500, 1000000)
            self.spin_limit.setSingleStep(500)
            self.spin_limit.valueChanged.connect(self.on_change)
            self.layout.addWidget(self.spin_limit)

            # Deep Scan Limit
            self.layout.addWidget(QLabel(tr("Deep Limit:")))
            self.spin_scan_limit = QSpinBox()
            self.spin_scan_limit.setToolTip(tr("Maximum number of documents to scan in Deep Scan mode."))
            self.spin_scan_limit.setRange(10000, 1000000)
            self.spin_scan_limit.setSingleStep(10000)
            self.spin_scan_limit.valueChanged.connect(self.on_change)
            self.layout.addWidget(self.spin_scan_limit)

            # Dynamic Weights
            self.chk_dynamic = QCheckBox(tr("Use Dynamic Corpus Stats (HTR-Aware)"))
            self.chk_dynamic.setToolTip(tr("Analyzes the corpus to detect and penalize HTR errors (e.g., if 'Tet' appears 4x more than expected, it is treated as noise). Saves a report to Reports folder."))
            self.chk_dynamic.clicked.connect(self.on_change)
            self.layout.addWidget(self.chk_dynamic)

            self.layout.addStretch()

            # Advanced Scoring
            self.btn_scoring = QPushButton(tr("Advanced Scoring..."))
            self.btn_scoring.setStyleSheet("background-color: #7f8c8d; color: white; font-weight: bold; border-radius: 4px; padding: 4px;")
            self.btn_scoring.clicked.connect(self.open_scoring)
            self.layout.addWidget(self.btn_scoring)

        elif self.mode == 'comp':
            # Chunk Limit
            self.layout.addWidget(QLabel(tr("Max Candidates per Chunk:")))
            self.spin_chunk_limit = QSpinBox()
            self.spin_chunk_limit.setToolTip(tr("Maximum number of index hits to process per text chunk."))
            self.spin_chunk_limit.setRange(50, 5000)
            self.spin_chunk_limit.setSingleStep(50)
            self.spin_chunk_limit.valueChanged.connect(self.on_change)
            self.layout.addWidget(self.spin_chunk_limit)

            # Min Score
            self.layout.addWidget(QLabel(tr("Min Chunk Score:")))
            self.spin_min_score = QSpinBox()
            self.spin_min_score.setToolTip(tr("Minimum score required for a chunk to be considered a match."))
            self.spin_min_score.setRange(10, 500)
            self.spin_min_score.valueChanged.connect(self.on_change)
            self.layout.addWidget(self.spin_min_score)

            # Max Final
            self.layout.addWidget(QLabel(tr("Max Final Results:")))
            self.spin_max_final = QSpinBox()
            self.spin_max_final.setToolTip(tr("Maximum number of results to display in the tree (prevents freezing). All results are exported."))
            self.spin_max_final.setRange(10, 250)
            self.spin_max_final.setValue(200)
            self.spin_max_final.valueChanged.connect(self.on_change)
            self.layout.addWidget(self.spin_max_final)

            self.layout.addStretch()

            # Advanced Scoring
            self.btn_scoring = QPushButton(tr("Advanced Scoring..."))
            self.btn_scoring.setStyleSheet("background-color: #7f8c8d; color: white; font-weight: bold; border-radius: 4px; padding: 4px;")
            self.btn_scoring.clicked.connect(self.open_scoring)
            self.layout.addWidget(self.btn_scoring)

    def refresh_values(self):
        if not self.settings: return
        self.blockSignals(True)
        if self.mode == 'search':
            self.spin_min.setValue(self.settings.min_should_match)
            self.spin_limit.setValue(self.settings.candidate_limit)
            if hasattr(self, 'spin_scan_limit'):
                self.spin_scan_limit.setValue(getattr(self.settings, 'lab_scan_limit', 50000))
            if hasattr(self, 'chk_dynamic'):
                self.chk_dynamic.setChecked(getattr(self.settings, 'use_dynamic_weights', False))
        elif self.mode == 'comp':
            self.spin_chunk_limit.setValue(self.settings.comp_chunk_limit)
            self.spin_min_score.setValue(self.settings.comp_min_score)
            self.spin_max_final.setValue(self.settings.comp_max_final_results)
        self.blockSignals(False)

    def on_change(self):
        if not self.settings: return
        if self.mode == 'search':
            self.settings.min_should_match = self.spin_min.value()
            self.settings.candidate_limit = self.spin_limit.value()
            if hasattr(self, 'spin_scan_limit'):
                self.settings.lab_scan_limit = self.spin_scan_limit.value()
            if hasattr(self, 'chk_dynamic'):
                self.settings.use_dynamic_weights = self.chk_dynamic.isChecked()
        elif self.mode == 'comp':
            self.settings.comp_chunk_limit = self.spin_chunk_limit.value()
            self.settings.comp_min_score = self.spin_min_score.value()
            self.settings.comp_max_final_results = self.spin_max_final.value()
        self.settings.save()

    def open_scoring(self):
         if not self.lab_engine: return
         d = LabScoringDialog(self, self.lab_engine)
         d.exec()

    def _mark_rebuild_required(self):
        if self.lab_engine.lab_index_needs_rebuild:
            self.lbl_idx_status.setText(tr("Index missing or outdated."))
            self.lbl_idx_status.setStyleSheet("color: #c0392b;")
        else:
            self.lbl_idx_status.setText(tr("Index is ready."))
            self.lbl_idx_status.setStyleSheet("color: #27ae60;")

    def run_rebuild(self):
        self.btn_rebuild.setEnabled(False)
        self.lbl_idx_status.setText(tr("Starting..."))
        QApplication.processEvents()

        class RebuildThread(QThread):
            finished_sig = pyqtSignal(int)
            progress_sig = pyqtSignal(int, int)
            error_sig = pyqtSignal(str)

            def __init__(self, engine):
                super().__init__()
                self.engine = engine
            def run(self):
                try:
                    if not os.path.exists(Config.LAB_DIR):
                        os.makedirs(Config.LAB_DIR)
                    
                    def cb(curr, total):
                        self.progress_sig.emit(curr, total)

                    count = self.engine.rebuild_lab_index(progress_callback=cb)
                    self.finished_sig.emit(count)
                except Exception as e:
                    self.error_sig.emit(str(e))

        self.worker = RebuildThread(self.lab_engine)
        self.worker.progress_sig.connect(self.on_rebuild_progress)
        self.worker.finished_sig.connect(self.on_rebuild_finished)
        self.worker.error_sig.connect(self.on_rebuild_error)
        self.worker.start()

    def on_rebuild_progress(self, current, total):
        self.lbl_idx_status.setText(tr("Processing docs: {}").format(current))

    def on_rebuild_error(self, err):
        self.btn_rebuild.setEnabled(True)
        self.lbl_idx_status.setText(tr("Error"))
        QMessageBox.critical(self, tr("Error"), str(err))

    def on_rebuild_finished(self, count):
        self.lbl_idx_status.setText(tr("Done. {} docs.").format(count))
        self.lbl_idx_status.setStyleSheet("color: #27ae60; font-weight: bold;")
        self.btn_rebuild.setEnabled(True)
        QMessageBox.information(self, tr("Success"), tr("Lab Index rebuilt successfully."))
      
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

_TLS_NOTICE_LOGGED = False


def log_tls_relaxation_notice():
    """Log once that TLS verification is intentionally disabled for thumbnail fetches."""
    global _TLS_NOTICE_LOGGED
    if not _TLS_NOTICE_LOGGED:
        logger.info(
            "TLS verification is disabled for thumbnail downloads to accommodate legacy IIIF endpoints "
            "with outdated certificates; certificate validation is skipped for these image requests."
        )
        _TLS_NOTICE_LOGGED = True


class ShelfmarkCompleter(QCompleter):
    """
    Custom Completer that normalizes input before matching.
    Input "T-S" -> Normalized "ts" -> Matches model items where UserRole starts with "ts".
    """
    def __init__(self, model, parent=None, valid_keys=None):
        super().__init__(model, parent)
        self.valid_keys = valid_keys or set()

    @staticmethod
    def normalize(text):
        t = re.sub(r'^\s*m[\.\s]*s[\.\s]*\.?\s*', '', text, flags=re.IGNORECASE)
        return re.sub(r"[^\w\./]", "", t).lower()

    def splitPath(self, path):
        return [self.normalize(path)]

    def pathFromIndex(self, index):
        # Return the pretty display text when an item is selected
        return index.data(Qt.ItemDataRole.DisplayRole)

    def complete(self, rect=QRect()):
        # Hide popup if there is an exact match
        text = self.widget().text()
        norm = self.normalize(text)
        if norm in self.valid_keys:
            self.popup().hide()
            return
        super().complete(rect)

class ShelfmarkTableWidgetItem(QTableWidgetItem):
    """Custom item for sorting shelfmarks by ignoring 'Ms.' prefix and case."""
    def __lt__(self, other):
        text1 = self.text()
        text2 = other.text()
        return natural_sort_key(text1) < natural_sort_key(text2)

class CheckBoxHeader(QHeaderView):
    """Custom HeaderView that draws a checkbox in the first section."""
    toggled = pyqtSignal(bool)

    def __init__(self, parent=None, non_sortable_cols=None, filter_columns=None, filter_callback=None, star_columns=None, star_callback=None, desc_first_cols=None):
        super().__init__(Qt.Orientation.Horizontal, parent)
        self.isChecked = False
        self.setSectionsClickable(True)
        self.non_sortable_cols = non_sortable_cols if non_sortable_cols else []
        self.filter_columns = set(filter_columns or [])
        self.filter_callback = filter_callback
        self.filter_states = {}
        self.star_columns = set(star_columns or [])
        self.star_callback = star_callback
        self.star_states = {}
        self.desc_first_cols = set(desc_first_cols or [])

    def get_checkbox_rect(self, rect):
        box_size = 20
        padding = 4
        y = rect.top() + (rect.height() - box_size) // 2

        if self.layoutDirection() == Qt.LayoutDirection.RightToLeft:
            x = rect.right() - box_size - padding
        else:
            x = rect.left() + padding

        return QRect(x, y, box_size, box_size)

    def paintSection(self, painter, rect, logicalIndex):
        painter.save()
        super().paintSection(painter, rect, logicalIndex)
        painter.restore()

        if logicalIndex in self.filter_columns:
            # Filter is usually right-most (index 0 from edge)
            icon_rect = self._get_icon_rect(rect, 0)
            self._draw_filter_icon(painter, icon_rect, self.filter_states.get(logicalIndex, False))

        if logicalIndex in self.star_columns:
            # Star is next to filter (index 1 if filter exists, else 0)
            offset = 1 if logicalIndex in self.filter_columns else 0
            icon_rect = self._get_icon_rect(rect, offset)
            self._draw_star_icon(painter, icon_rect, self.star_states.get(logicalIndex, False))

        if logicalIndex == 0:
            option = QStyleOptionButton()
            option.rect = self.get_checkbox_rect(rect)
            option.state = QStyle.StateFlag.State_Enabled | QStyle.StateFlag.State_Active
            if self.isChecked:
                option.state |= QStyle.StateFlag.State_On
            else:
                option.state |= QStyle.StateFlag.State_Off

            self.style().drawControl(QStyle.ControlElement.CE_CheckBox, option, painter)

    def mousePressEvent(self, event):
        idx = self.logicalIndexAt(event.pos())

        # Handle Filter/Star clicks
        if (idx in self.filter_columns and self.filter_callback) or (idx in self.star_columns and self.star_callback):
            sec_pos = self.sectionViewportPosition(idx)
            sec_width = self.sectionSize(idx)
            sec_rect = QRect(sec_pos, 0, sec_width, self.height())

            if idx in self.filter_columns and self.filter_callback:
                if self._get_icon_rect(sec_rect, 0).contains(event.pos()):
                    self.filter_callback(idx)
                    return

            if idx in self.star_columns and self.star_callback:
                offset = 1 if idx in self.filter_columns else 0
                if self._get_icon_rect(sec_rect, offset).contains(event.pos()):
                    self.star_callback(idx)
                    return

        if idx == 0:
            sec_pos = self.sectionViewportPosition(0)
            sec_width = self.sectionSize(0)
            sec_rect = QRect(sec_pos, 0, sec_width, self.height())

            chk_rect = self.get_checkbox_rect(sec_rect)

            if chk_rect.contains(event.pos()):
                self.isChecked = not self.isChecked
                self.viewport().update()
                self.toggled.emit(self.isChecked)
                return # Consume event (Checkbox toggle)

            # If we clicked the header area but NOT the checkbox, check if sort should be blocked
            if 0 in self.non_sortable_cols:
                return # Prevent sort on col 0

        elif idx in self.non_sortable_cols:
            return # Prevent sort

        # For desc-first columns: if not currently sorted on this column,
        # pre-set indicator to ascending so the toggle goes to descending
        if idx in self.desc_first_cols and self.sortIndicatorSection() != idx:
            self.setSortIndicator(idx, Qt.SortOrder.AscendingOrder)

        super().mousePressEvent(event)

    def setChecked(self, checked):
        if self.isChecked != checked:
            self.isChecked = checked
            self.viewport().update()

    def _get_icon_rect(self, rect, offset_index=0):
        icon_size = 12
        padding = 6
        spacing = 4

        total_offset = padding + (offset_index * (icon_size + spacing))

        y = rect.top() + (rect.height() - icon_size) // 2

        if self.layoutDirection() == Qt.LayoutDirection.RightToLeft:
            x = rect.left() + total_offset
        else:
            x = rect.right() - icon_size - total_offset

        return QRect(x, y, icon_size, icon_size)

    def _draw_filter_icon(self, painter, rect, active):
        painter.save()
        color = self.palette().color(QPalette.ColorRole.Highlight if active else QPalette.ColorRole.Mid)
        pen = QPen(color)
        pen.setWidth(1)
        painter.setPen(pen)
        painter.setBrush(QBrush(color if active else Qt.BrushStyle.NoBrush))

        x = rect.x()
        y = rect.y()
        w = rect.width()
        h = rect.height()
        top_h = int(h * 0.55)
        stem_w = max(2, int(w * 0.3))
        mid_x = x + w // 2
        stem_left = mid_x - stem_w // 2
        stem_right = stem_left + stem_w

        path = QPainterPath()
        path.moveTo(x, y)
        path.lineTo(x + w, y)
        path.lineTo(stem_right, y + top_h)
        path.lineTo(stem_right, y + h)
        path.lineTo(stem_left, y + h)
        path.lineTo(stem_left, y + top_h)
        path.closeSubpath()
        painter.drawPath(path)
        painter.restore()

    def _draw_star_icon(self, painter, rect, active):
        painter.save()
        # Star color: Gold if active, Gray if inactive
        if active:
            color = QColor("#f1c40f") # Gold
            brush = QBrush(color)
        else:
            color = self.palette().color(QPalette.ColorRole.Mid)
            brush = Qt.BrushStyle.NoBrush

        pen = QPen(color)
        pen.setWidth(1)
        painter.setPen(pen)
        painter.setBrush(brush)

        # Draw Star using a simple path
        center = rect.center()
        radius = rect.width() / 2.0

        path = QPainterPath()
        import math
        points = []
        for i in range(5):
            # Outer point
            angle_deg = -90 + i * 72 # Start from top (rotated -90)
            angle_rad = math.radians(angle_deg)
            ox = center.x() + radius * math.cos(angle_rad)
            oy = center.y() + radius * math.sin(angle_rad)
            points.append((ox, oy))

            # Inner point
            angle_deg = -90 + i * 72 + 36
            angle_rad = math.radians(angle_deg)
            ix = center.x() + (radius * 0.4) * math.cos(angle_rad)
            iy = center.y() + (radius * 0.4) * math.sin(angle_rad)
            points.append((ix, iy))

        path.moveTo(points[0][0], points[0][1])
        for x, y in points[1:]:
            path.lineTo(x, y)
        path.closeSubpath()

        painter.drawPath(path)
        painter.restore()

    def set_filter_active(self, column, active):
        if active:
            self.filter_states[column] = True
        else:
            self.filter_states.pop(column, None)
        self.viewport().update()

    def set_star_active(self, column, active):
        if active:
            self.star_states[column] = True
        else:
            self.star_states.pop(column, None)
        self.viewport().update()

    def event(self, event):
        if event.type() == QEvent.Type.ToolTip:
            pos = event.pos()
            idx = self.logicalIndexAt(pos)

            if idx in self.star_columns or idx in self.filter_columns:
                sec_pos = self.sectionViewportPosition(idx)
                sec_width = self.sectionSize(idx)
                sec_rect = QRect(sec_pos, 0, sec_width, self.height())

                # Check Star
                if idx in self.star_columns:
                    offset = 1 if idx in self.filter_columns else 0
                    if self._get_icon_rect(sec_rect, offset).contains(pos):
                        QToolTip.showText(event.globalPos(), tr("Show entries in selected lists") if self.star_states.get(idx) else tr("Filter by List (Click to enable)"))
                        return True

                # Check Filter
                if idx in self.filter_columns:
                    if self._get_icon_rect(sec_rect, 0).contains(pos):
                        QToolTip.showText(event.globalPos(), tr("Filter configuration"))
                        return True

        return super().event(event)

class ZoomableScrollArea(QGraphicsView):
    """A GraphicsView that supports hand-panning and wheel-zooming."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.scene = QGraphicsScene(self)
        self.setScene(self.scene)

        self.setRenderHint(QPainter.RenderHint.Antialiasing)
        self.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform)
        self.setDragMode(QGraphicsView.DragMode.ScrollHandDrag)
        self.setTransformationAnchor(QGraphicsView.ViewportAnchor.AnchorUnderMouse)
        self.setResizeAnchor(QGraphicsView.ViewportAnchor.AnchorUnderMouse)
        self.setStyleSheet("background: #222; border: none;")

        # Hide scrollbars but keep functionality
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

        self._pixmap_item = QGraphicsPixmapItem()
        self._pixmap_item.setTransformationMode(Qt.TransformationMode.SmoothTransformation)
        self.scene.addItem(self._pixmap_item)

        self._msg_item = QGraphicsSimpleTextItem()
        self._msg_item.setBrush(QColor("white"))
        font = QFont("Arial", 16)
        self._msg_item.setFont(font)
        self.scene.addItem(self._msg_item)

        self._pixmap = None
        self._rotation = 0
        self._auto_fit_enabled = False

    def set_image(self, pixmap):
        self._pixmap = pixmap
        self._rotation = 0
        self._auto_fit_enabled = bool(pixmap)

        # Guard against destroyed graphics items (async callback after widget close)
        if sip.isdeleted(self._msg_item) or sip.isdeleted(self._pixmap_item):
            return

        if not pixmap or pixmap.isNull():
            self._pixmap_item.setVisible(False)
            self.set_status_message(tr("No Image"))
            return

        self._msg_item.setVisible(False)
        self._pixmap_item.setPixmap(pixmap)
        self._pixmap_item.setVisible(True)

        # Reset transform
        self.resetTransform()

        # Center item transform origin
        rect = QRectF(pixmap.rect())
        # Let scene grow automatically to fit rotated items
        self.scene.setSceneRect(QRectF())
        self._pixmap_item.setPos(0, 0)
        self._pixmap_item.setTransformOriginPoint(rect.center())

        if not self._apply_fit_to_viewport():
             self.centerOn(self._pixmap_item)

    def set_status_message(self, text):
        if sip.isdeleted(self._msg_item) or sip.isdeleted(self._pixmap_item):
            return
        self._pixmap_item.setVisible(False)
        self._msg_item.setText(text)
        self._msg_item.setVisible(True)
        self._update_text_pos()

    def _update_text_pos(self):
        if sip.isdeleted(self._msg_item):
            return
        if not self._msg_item.isVisible(): return

        # Simple center in view
        # We need to map viewport center to scene
        center = self.mapToScene(self.viewport().rect().center())
        brect = self._msg_item.boundingRect()
        self._msg_item.setPos(center.x() - brect.width()/2, center.y() - brect.height()/2)

    def set_rotation(self, angle: float):
        """Set absolute rotation (degrees clockwise) and update view."""
        self._rotation = angle % 360 if angle is not None else 0
        self._pixmap_item.setRotation(self._rotation)

    def rotate_view(self, degrees):
        """Add degrees to current rotation and update."""
        self._rotation = (self._rotation + degrees) % 360
        self._pixmap_item.setRotation(self._rotation)

    def wheelEvent(self, event):
        if event.modifiers() & Qt.KeyboardModifier.ControlModifier:
            # Ctrl+wheel: zoom
            self._auto_fit_enabled = False
            delta = event.angleDelta().y()
            factor = 1.1 if delta > 0 else 0.9
            self._apply_zoom(factor)
            event.accept()
        else:
            # Plain wheel: propagate to parent for normal scrolling
            event.ignore()

    def zoom_in(self):
        self._auto_fit_enabled = False
        self._apply_zoom(1.2)

    def zoom_out(self):
        self._auto_fit_enabled = False
        self._apply_zoom(0.8)

    def _apply_zoom(self, factor):
        current_scale = self.transform().m11()
        new_scale = current_scale * factor

        # Clamp zoom level (0.1 to 5.0)
        if new_scale < 0.1:
            factor = 0.1 / current_scale
        elif new_scale > 5.0:
            factor = 5.0 / current_scale

        self.scale(factor, factor)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        if self._auto_fit_enabled:
            self._apply_fit_to_viewport()
        else:
            self._update_text_pos()

    def _apply_fit_to_viewport(self):
        if not self._pixmap or self._pixmap.isNull():
            return False

        self.fitInView(self._pixmap_item, Qt.AspectRatioMode.KeepAspectRatio)
        # Scale down a bit to have margins
        self.scale(0.95, 0.95)
        return True

class ManuscriptViewerWidget(QWidget):
    """Reusable widget for displaying manuscript images with navigation."""
    _thumbnail_ready = pyqtSignal(QPixmap, int)  # pixmap, page_index

    def __init__(self, parent=None):
        super().__init__(parent)
        self.images_nli = []
        self.images_ext = []
        self.active_list = []
        self.current_idx = 0
        self.loader_thread = None
        self.preload_worker = None
        self.external_provider = None
        self._closing = False
        self._thumbnail_ready.connect(self._on_thumbnail_ready)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Top Bar (Source + Zoom)
        top_bar = QHBoxLayout()
        top_bar.setContentsMargins(5, 5, 5, 5)

        self.combo_source = QComboBox()
        self.combo_source.addItem("NLI")
        self.combo_source.addItem("External (Cambridge/Other)")
        self.combo_source.setVisible(False)
        self.combo_source.currentIndexChanged.connect(self._on_source_changed)

        btn_zoom_out = QPushButton("-")
        btn_zoom_out.setFixedWidth(30)
        btn_zoom_out.clicked.connect(lambda: self.scroll_area.zoom_out())

        btn_zoom_in = QPushButton("+")
        btn_zoom_in.setFixedWidth(30)
        btn_zoom_in.clicked.connect(lambda: self.scroll_area.zoom_in())

        # Rotation controls
        self.slider_rotation = QSlider(Qt.Orientation.Horizontal)
        self.slider_rotation.setRange(0, 360)
        self.slider_rotation.setValue(0)
        self.slider_rotation.setFixedWidth(160)
        self.slider_rotation.setToolTip(tr("Rotate image (0-360°)"))
        self.slider_rotation.valueChanged.connect(lambda val: self.scroll_area.set_rotation(val))

        btn_rot_left = QPushButton("↺")
        btn_rot_left.setToolTip(tr("Rotate Left 90°"))
        btn_rot_left.setFixedWidth(30)
        btn_rot_left.clicked.connect(lambda: self.adjust_rotation(-90))

        btn_rot_right = QPushButton("↻")
        btn_rot_right.setToolTip(tr("Rotate Right 90°"))
        btn_rot_right.setFixedWidth(30)
        btn_rot_right.clicked.connect(lambda: self.adjust_rotation(90))

        btn_rot_reset = QPushButton(tr("Reset"))
        btn_rot_reset.setToolTip(tr("Reset rotation"))
        btn_rot_reset.setFixedWidth(50)
        btn_rot_reset.clicked.connect(lambda: self.slider_rotation.setValue(0))

        self.btn_external = QPushButton(tr("External Website"))
        self.btn_external.setVisible(False)
        self.btn_external.clicked.connect(self.open_external)

        # KTIV / NLI Viewer button (Phase 31)
        self.btn_ktiv = QPushButton(tr("View on Ktiv"))
        self.btn_ktiv.setToolTip(tr("View image in Ktiv"))
        self.btn_ktiv.setVisible(False)
        self.btn_ktiv.clicked.connect(self._open_ktiv_viewer)
        self._ktiv_sys_id = None

        top_bar.addWidget(self.combo_source)
        top_bar.addStretch()
        top_bar.addWidget(self.btn_ktiv)
        top_bar.addWidget(self.btn_external)
        top_bar.addWidget(btn_rot_left)
        top_bar.addWidget(self.slider_rotation)
        top_bar.addWidget(btn_rot_right)
        top_bar.addWidget(btn_rot_reset)
        top_bar.addSpacing(10)
        top_bar.addWidget(btn_zoom_out)
        top_bar.addWidget(btn_zoom_in)

        layout.addLayout(top_bar)

        # Attribution
        self.lbl_attribution = QLabel("")
        self.lbl_attribution.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lbl_attribution.setWordWrap(True)
        self.lbl_attribution.setStyleSheet("font-size: 10px; color: #7f8c8d; background: transparent; margin: 0px;")
        self.lbl_attribution.setVisible(False)
        layout.addWidget(self.lbl_attribution)

        # Image Area
        self.scroll_area = ZoomableScrollArea()
        layout.addWidget(self.scroll_area, 1)

    def _detect_external_provider(self, meta):
        # Check explicit provider from enrich_metadata (set by Plan 03 for Manchester/JTS)
        if meta.get('external_provider'):
            return meta['external_provider']

        marc = meta.get('marc', {}) if meta else {}
        url = (marc.get('external_iiif_link') or "").lower()
        if "cudl.lib.cam.ac.uk" in url:
            return "cambridge"

        for img in meta.get('images_ext', []) or []:
            img_url = (img.get('url', '').lower())
            if "cudl.lib.cam.ac.uk" in img_url:
                return "cambridge"
            if "bodleian.ox.ac.uk" in img_url or "hebrew.bodleian" in img_url:
                return "oxford"
            if "luna.manchester.ac.uk" in img_url:
                return "manchester"
            if "iiif-cloud.princeton.edu" in img_url or "figgy.princeton.edu" in img_url:
                return "jts"

        # Check for Oxford Part ID
        if meta.get('oxford_part_id'):
            return "oxford"

        return None

    def set_image_by_fl_id(self, fl_id):
        digits = re.sub(r"\D", "", str(fl_id or ""))
        if not digits:
            return False

        fallback_url = f"{Config.NLI_IIIF_BASE}/FL{digits}/full/2000,/0/default.jpg"
        self.images_nli = [{'label': f"FL{digits}", 'url': fallback_url, 'fl_id': digits}]
        self.images_ext = []
        self.active_list = self.images_nli
        self.current_source = "nli"
        self.combo_source.clear()
        self.combo_source.addItem(f"NLI (1 page)", "nli")
        self.combo_source.setVisible(False)
        self.btn_ktiv.setVisible(False)
        self._ktiv_sys_id = None
        return True

    def load_images(self, meta, initial_idx=0, target_folio=None):
        self.external_provider = self._detect_external_provider(meta)
        self._current_meta = meta  # Store for dynamic image generation

        # Attribution
        attr = meta.get('attribution')
        if attr:
            self.lbl_attribution.setText(attr)
            self.lbl_attribution.setVisible(True)
        else:
            self.lbl_attribution.setVisible(False)

        # meta contains 'images_nli' and 'images_ext'
        # Make copies to avoid modifying the cached meta
        self.images_nli = list(meta.get('images_nli', []))
        self.images_ext = list(meta.get('images_ext', []))

        # For Oxford: check if target_folio is missing and add dynamic images
        if target_folio is not None and self.images_ext:
            folio_in_list = any(img.get('folio_num') == target_folio for img in self.images_ext)
            if not folio_in_list:
                oxford_part_id = meta.get('oxford_part_id', '')
                oxford_part_meta = meta.get('oxford_part_metadata', {})
                folio_range = oxford_part_meta.get('folio_range', [])
                if oxford_part_id and len(folio_range) >= 2 and folio_range[0] <= target_folio <= folio_range[1]:
                    # Generate dynamic URLs for this folio
                    for side in ['a', 'b']:
                        url = _generate_oxford_dynamic_url(oxford_part_id, target_folio, side)
                        if url:
                            self.images_ext.append({
                                'label': f'{target_folio}{side}',
                                'url': url,
                                'folio_num': target_folio,
                                '_dynamic': True
                            })
                    # Update initial_idx to point to the new recto image
                    initial_idx = len(self.images_ext) - 2  # Index of recto (a)

        # Determine default source
        self.combo_source.blockSignals(True)
        self.combo_source.clear()

        if self.images_ext:
            if self.external_provider == "cambridge":
                ext_label = "Cambridge"
            elif self.external_provider == "oxford":
                ext_label = "Oxford"
            elif self.external_provider == "manchester":
                ext_label = "Manchester"
            elif self.external_provider == "jts":
                ext_label = "JTS/Princeton"
            else:
                ext_label = "External"
            self.combo_source.addItem(f"{ext_label} ({len(self.images_ext)} pages)", "ext")
            if self.images_nli:
                self.combo_source.addItem(f"NLI ({len(self.images_nli)} pages)", "nli")
            self.active_list = self.images_ext
            self.current_source = "ext"
        elif self.images_nli:
            self.combo_source.addItem(f"NLI ({len(self.images_nli)} pages)", "nli")
            self.active_list = self.images_nli
            self.current_source = "nli"
        else:
            self.active_list = []
            self.current_source = None

        self.combo_source.setVisible(len(self.images_nli) > 0 and len(self.images_ext) > 0)
        self.combo_source.blockSignals(False)

        if not self.active_list:
            fl_ids = meta.get('fl_ids') if meta else []
            if isinstance(fl_ids, str):
                fl_ids = [fl_ids]
            for fl in fl_ids or []:
                if self.set_image_by_fl_id(fl):
                    break

        # External Link
        marc = meta.get('marc', {})
        self.external_url = meta.get('external_url') or marc.get('external_iiif_link')

        # Prefer library_viewer_url when available (detail page > generic manifest URL)
        lib_viewer = meta.get('library_viewer_url')
        if lib_viewer and lib_viewer.get('url'):
            if self.external_provider in ('manchester', 'jts'):
                self.external_url = lib_viewer['url']

        if self.external_url:
            if self.external_provider == "cambridge":
                btn_label = tr("Cambridge")
            elif self.external_provider == "oxford":
                btn_label = tr("Oxford")
            elif self.external_provider == "manchester":
                btn_label = "Manchester LUNA"
            elif self.external_provider == "jts":
                btn_label = "Princeton Digital Library"
            else:
                btn_label = tr("External Website")
            self.btn_external.setText(btn_label)
        self.btn_external.setVisible(bool(self.external_url))

        # KTIV button: show when NLI FGP images available (Phase 31)
        image_source_info = meta.get('image_source_info', {})
        if image_source_info.get('nli_fgp'):
            # Use sys_id from meta or try fl_id-based detection
            sys_id = meta.get('sys_id', '')
            if not sys_id:
                # Try to get sys_id from the fl_ids
                fl_ids = meta.get('fl_ids', [])
                if isinstance(fl_ids, str):
                    fl_ids = [fl_ids]
            self._ktiv_sys_id = sys_id
            self.btn_ktiv.setVisible(bool(sys_id))
        else:
            self._ktiv_sys_id = None
            self.btn_ktiv.setVisible(False)

        # Set Page
        self.set_page(initial_idx)

    def _on_source_changed(self):
        data = self.combo_source.currentData()
        if data == "nli":
            self.active_list = self.images_nli
            self.current_source = "nli"
        else:
            self.active_list = self.images_ext
            self.current_source = "ext"

        # Try to keep index within bounds
        if self.current_idx >= len(self.active_list):
            self.current_idx = 0

        self.set_page(self.current_idx)

    def _resolve_url(self, base_url):
        if not base_url: return None
        if base_url.endswith('.jpg'): return base_url
        return f"{base_url}/full/2000,/0/default.jpg"

    def _preload(self, index):
        if index < 0 or index >= len(self.active_list): return
        url = self.active_list[index]['url']
        final = self._resolve_url(url)

        # Spawn thread without connecting signals (just for cache)
        # Store ref to prevent GC
        self.preload_worker = ImageLoaderThread(final)
        self.preload_worker.start()

    def stop_threads(self):
        """Stop all running image loading threads. Call before destroying widget."""
        self._closing = True
        if self.loader_thread and self.loader_thread.isRunning():
            self.loader_thread.cancel()
            self.loader_thread.wait(2000)
        if self.preload_worker and self.preload_worker.isRunning():
            self.preload_worker.cancel()
            self.preload_worker.wait(1000)

    def _on_thumbnail_ready(self, pix, page_idx):
        """Handle thumbnail loaded signal - only display if still on same page."""
        if self._closing:
            return
        if self.current_idx == page_idx and pix and not pix.isNull():
            self.scroll_area.set_image(pix)

    def _load_thumbnail_async(self, thumb_url):
        """Load thumbnail asynchronously for quick display while full image loads."""
        current_idx = self.current_idx  # Capture current state
        signal = self._thumbnail_ready  # Capture signal reference for thread
        closing_ref = lambda: self._closing  # Capture closing flag check

        def fetch_and_emit():
            try:
                if closing_ref():
                    return
                import urllib.request
                req = urllib.request.Request(thumb_url, headers=Config.HTTP_HEADERS)
                with urllib.request.urlopen(req, timeout=3) as response:
                    data = response.read()
                    if data and not closing_ref():
                        image = QImage()
                        if image.loadFromData(data):
                            pix = QPixmap.fromImage(image)
                            signal.emit(pix, current_idx)
            except Exception:
                pass  # Thumbnail load failed, full image will replace it

        # Run in background thread to avoid blocking UI
        threading.Thread(target=fetch_and_emit, daemon=True).start()

    def set_page(self, index):
        if not self.active_list:
            self.scroll_area.set_image(None)
            self.scroll_area.set_status_message(tr("No images available"))
            return

        # Bounds check
        if index < 0: index = 0
        if index >= len(self.active_list): index = len(self.active_list) - 1

        self.current_idx = index
        img_data = self.active_list[index]
        base_url = img_data['url']

        # Check for thumbnail URL (Oxford images have this)
        thumb_url = img_data.get('thumb_url', '')
        # For NLI IIIF images, auto-generate a fast preview URL (400px)
        if not thumb_url and 'iiif.nli.org.il' in base_url:
            thumb_url = f"{base_url}/full/400,/0/default.jpg"

        self.scroll_area.set_status_message(tr("Loading..."))

        if self.loader_thread and self.loader_thread.isRunning():
            self.loader_thread.cancel()
            # Use short timeout to avoid blocking UI - thread will finish in background
            self.loader_thread.wait(500)

        # Load low-res preview first for instant display, then high-res replaces it
        if thumb_url:
            self._load_thumbnail_async(thumb_url)

        final_url = self._resolve_url(base_url)

        self.loader_thread = ImageLoaderThread(final_url)
        self.loader_thread.image_loaded.connect(self.display_image)
        self.loader_thread.load_failed.connect(lambda: None if self._closing else self.scroll_area.set_status_message(tr("No Image")))
        self.loader_thread.start()

        # Preload next image
        self._preload(index + 1)

    def display_image(self, image):
        if self._closing:
            return
        pix = QPixmap.fromImage(image)
        self.scroll_area.set_image(pix)
        self.slider_rotation.setValue(0)

    def open_external(self):
        if self.external_url:
            url = self.external_url
            # Transform CUDL IIIF manifest URL to viewer URL
            if "cudl.lib.cam.ac.uk/iiif/" in url:
                url = url.replace("/iiif/", "/view/")
            QDesktopServices.openUrl(QUrl(url))

    def _open_ktiv_viewer(self):
        """Open the NLI KTIV manuscript viewer at the current page."""
        if self._ktiv_sys_id:
            # Use docid query param (not hash fragment) — hash-based URLs fail on direct navigation
            docid = f"PNX_MANUSCRIPTS{self._ktiv_sys_id}-1"
            # Append FL ID to navigate to the current page
            if self.active_list and self.current_source == "nli" and 0 <= self.current_idx < len(self.active_list):
                fl_id = self.active_list[self.current_idx].get('fl_id')
                if fl_id:
                    docid += f",FL{fl_id}"
            url = f"https://www.nli.org.il/he/discover/manuscripts/hebrew-manuscripts/viewerpage?vid=MANUSCRIPTS&docid={docid}"
            QDesktopServices.openUrl(QUrl(url))

    def adjust_rotation(self, delta):
        """Adjust rotation via slider to keep controls in sync."""
        new_val = (self.slider_rotation.value() + delta) % 360
        self.slider_rotation.setValue(int(new_val))

class HiddenScrollArea(QScrollArea):
    def __init__(self, text_with_markers="", anchor_text=None, parent=None):
        super().__init__(parent)
        self._raw_text = text_with_markers
        self._anchor_text = anchor_text
        
        # Hide scrollbars
        self.setWidgetResizable(True)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setFrameShape(QFrame.Shape.NoFrame)
        
        # Add horizontal margins to create a gap between columns (5px on each side = 10px total gap)
        # We also ensure the background is transparent to show the row selection color
        self.setStyleSheet("QScrollArea { background: transparent; margin-left: 5px; margin-right: 5px; }")
        
        # Keep height strictly slim
        self.setFixedHeight(self.fontMetrics().lineSpacing() + 4)
        
        self.label = QLabel()
        self.label.setTextFormat(Qt.TextFormat.RichText)
        self.label.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        self.label.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        # Ensure label background is transparent
        self.label.setStyleSheet("background: transparent;")
        self.setWidget(self.label)
        
        self._update_content()

    def _update_content(self):
        if not self._raw_text:
            self.label.setText(""); return

        # Apply coloring to markers
        processed = re.sub(r'\*(.*?)\*', r"<b style='color:#c0392b;'>\1</b>", self._raw_text)
        processed = re.sub(r'\*([^*]+)$', r"<b style='color:#c0392b;'>\1</b>", processed)
        processed = re.sub(r'^([^*]+)\*', r"<b style='color:#c0392b;'>\1</b>", processed)
        final_html = processed.replace("*", "")
        
        # Enforce non-breaking text
        self.label.setText(f"<div dir='rtl' style='white-space:nowrap; padding: 0 5px;'>{final_html}</div>")
        self.setToolTip(self._raw_text.replace("*", ""))
        
        # Position highlight in view initially
        QTimer.singleShot(10, self._center_on_match)

    def _center_on_match(self):
        target_pos = -1
        if self._anchor_text:
            target_pos = self._raw_text.find(f"*{self._anchor_text}*")
        if target_pos == -1:
            target_pos = self._raw_text.find('*')
            
        if target_pos != -1:
            bar = self.horizontalScrollBar()
            max_val = bar.maximum()
            if max_val > 0:
                # Calculate center ratio for RTL scrollbar
                ratio = (len(self._raw_text) - target_pos) / len(self._raw_text)
                bar.setValue(int(max_val * ratio))

    def wheelEvent(self, event):
        # Convert vertical wheel movement to horizontal scroll
        if event.angleDelta().y() != 0:
            bar = self.horizontalScrollBar()
            # Sensitivity adjustment
            bar.setValue(bar.value() - event.angleDelta().y())
            event.accept()
        else:
            super().wheelEvent(event)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        # Maintain highlight focus when column width changes
        QTimer.singleShot(10, self._center_on_match)
        
class ImageLoaderThread(QThread):
    """
    Smart Image Loader:
    1. Checks Local Disk Cache first.
    2. If missing, Downloads from IIIF (with Rosetta fallback).
    3. Saves successful downloads to Disk Cache.
    """

    image_loaded = pyqtSignal(QImage)
    load_failed = pyqtSignal()

    def __init__(self, url):
        super().__init__()
        self.url = url
        self._cancelled = False
        
        # Ensure cache directory exists
        if not os.path.exists(Config.IMAGE_CACHE_DIR):
            try:
                os.makedirs(Config.IMAGE_CACHE_DIR)
            except Exception as e:
                logger.warning(
                    "Could not create image cache directory at %s: %s; image caching disabled for this session.",
                    Config.IMAGE_CACHE_DIR,
                    e,
                )

    def cancel(self):
        self._cancelled = True

    def run(self):
        if not self.url:
            self.load_failed.emit()
            return

        # 1. Determine cache filename: FL ID for NLI, URL hash for external
        fl_match = re.search(r'FL(\d+)', self.url)
        local_path = None

        if fl_match:
            fl_id = fl_match.group(1)
            # v2 cache: high resolution (2000px). Old v1 cache was 600px.
            local_path = os.path.join(Config.IMAGE_CACHE_DIR, f"FL{fl_id}_v2.jpg")
        else:
            # Cache external images (Cambridge, Manchester, Oxford, JTS) by URL hash
            import hashlib
            url_hash = hashlib.md5(self.url.encode('utf-8')).hexdigest()[:16]
            local_path = os.path.join(Config.IMAGE_CACHE_DIR, f"ext_{url_hash}.jpg")

        # --- CHECK LOCAL CACHE ---
        if local_path and os.path.exists(local_path) and os.path.getsize(local_path) > 0:
            img = QImage(local_path)
            if not img.isNull():
                self.image_loaded.emit(img)
                return
            else:
                # Corrupt file? Delete it so we re-download
                try:
                    os.remove(local_path)
                except Exception as e:
                    logger.warning("Failed to remove corrupt cache file %s: %s", local_path, e)

        # 2. Download from Network (if not in cache)
        headers = dict(Config.HTTP_HEADERS)
        headers["Referer"] = "https://www.nli.org.il/"

        data = None
        
        # Attempt A: Original URL
        data = self._download_bytes(self.url, headers)
        
        # Attempt B: Fallback to Rosetta stream if IIIF failed (full-res TIFF)
        if data is None and fl_match and not self._cancelled:
            fl_digits = fl_match.group(1)
            logger.info("IIIF failed for FL%s. Trying Rosetta stream fallback...", fl_digits)
            fallback_url = MetadataManager.get_rosetta_fallback_url(fl_digits)
            if fallback_url:
                data = self._download_bytes(fallback_url, headers)

        # Attempt C: Rosetta thumbnail if stream also failed (e.g. 401 for some libraries)
        if data is None and fl_match and not self._cancelled:
            fl_digits = fl_match.group(1)
            logger.info("Rosetta stream failed for FL%s. Trying thumbnail fallback...", fl_digits)
            thumb_url = f"https://rosetta.nli.org.il/delivery/DeliveryManagerServlet?dps_func=thumbnail&dps_pid=FL{fl_digits}"
            data = self._download_bytes(thumb_url, headers)

        # 3. Process Result
        if data:
            img = QImage.fromData(data)
            if not img.isNull():
                self.image_loaded.emit(img)

                # --- SAVE TO LOCAL CACHE (always as JPEG for compact storage) ---
                if local_path and not self._cancelled:
                    try:
                        img.save(local_path, "JPEG", 85)
                        logger.debug("Saved image cache to %s", local_path)
                    except Exception as e:
                        logger.warning(
                            "Failed to write image cache for %s: %s; future loads will re-download.",
                            local_path,
                            e,
                        )
            else:
                self.load_failed.emit()
        else:
            self.load_failed.emit()

    def _download_bytes(self, target_url, headers):
        """Helper to download bytes safely."""
        try:
            # Rosetta stream returns large TIFF files (7-15MB) — allow longer timeout
            timeout = 30 if 'rosetta.nli.org.il' in target_url else 10
            resp = requests.get(target_url, headers=headers, timeout=timeout, stream=True, verify=False)
            if self._cancelled: return None
            if resp.status_code == 200:
                return resp.content
            return None
        except Exception as e:
            logger.warning("Image download failed for %s: %s", target_url, e)
            return None
                
class HelpDialog(QDialog):
    """Display HTML help content from the bundled Help.html file with graceful fallback."""
    def __init__(self, parent, title, source_path=None, anchor=None, fallback_html="", lang="en"):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.setWindowIcon(QIcon(os.path.join(Config.BASE_DIR, "icon.ico")))
        self.resize(900, 700)
        layout = QVBoxLayout()
        self.text = QTextBrowser()
        self.text.setOpenExternalLinks(True)
        layout.addWidget(self.text)

        self._load_content(source_path, anchor, fallback_html, lang)

        btn = QPushButton(tr("Close"))
        btn.clicked.connect(self.close)
        layout.addWidget(btn)
        self.setLayout(layout)

    def _load_content(self, source_path, anchor, fallback_html, lang):
        if source_path and os.path.exists(source_path):
            try:
                with open(source_path, "r", encoding="utf-8") as f:
                    content = f.read()

                # --- 1. LANGUAGE FILTERING (Content Stripping) ---
                # Since QTextBrowser ignores "display: none", we must remove the unused language block manually.
                # We rely on markers added to Help.html.
                if lang == 'he':
                    # Keep Hebrew -> Remove English
                    start_marker = "<!-- START_LANG_EN -->"
                    end_marker = "<!-- END_LANG_EN -->"
                else:
                    # Keep English -> Remove Hebrew
                    start_marker = "<!-- START_LANG_HE -->"
                    end_marker = "<!-- END_LANG_HE -->"

                s_idx = content.find(start_marker)
                e_idx = content.find(end_marker)

                if s_idx != -1 and e_idx != -1:
                    # Remove the block including markers
                    content = content[:s_idx] + content[e_idx + len(end_marker):]

                # Removed explicit Dark Mode CSS injection to allow "native" palette behavior
                # as requested by the user ("do what the previous version did").
                # By removing explicit background colors in HTML, QTextBrowser uses QPalette.

                self.text.setHtml(content)
                if anchor:
                    QTimer.singleShot(0, lambda: self.text.scrollToAnchor(anchor))
                return
            except Exception as e:
                logger.warning("Failed to load help file %s: %s", source_path, e)
        # Fallback: prefer clean content without warning if we have a fallback HTML snippet
        if fallback_html:
            self.text.setHtml(fallback_html)
        else:
            notice = "<p style='color:#c0392b;'><b>Help file is missing or could not be loaded.</b></p>"
            self.text.setHtml(notice)
        if anchor:
            QTimer.singleShot(0, lambda: self.text.scrollToAnchor(anchor))

class ExcludeDialog(QDialog):
    """Collect system IDs or shelfmarks that should be excluded from searches."""
    def __init__(self, parent, existing_entries=None):
        super().__init__(parent)
        self.setWindowTitle(tr("Exclude Manuscripts"))
        self.resize(500, 400)
        layout = QVBoxLayout()

        help_lbl = QLabel(tr("Enter system IDs or shelfmarks to exclude (one per line). Matching values are filled automatically."))
        help_lbl.setWordWrap(True)
        layout.addWidget(help_lbl)

        self._syncing = False
        self._shelf_to_sys = None
        self._last_edited = None
        self._full_titles = []
        self._display_titles = []
        self.meta_mgr = getattr(parent, "meta_mgr", None)

        grid = QGridLayout()
        grid.addWidget(QLabel(tr("System IDs")), 0, 0)
        grid.addWidget(QLabel(tr("Shelfmarks")), 0, 1)
        grid.addWidget(QLabel(tr("Title")), 0, 2)

        self.sys_text_area = QPlainTextEdit()
        self.sys_text_area.setPlaceholderText("990051564290205171\n990053963680205171")
        self.sys_text_area.textChanged.connect(self._on_sys_text_changed)

        self.shelf_text_area = QPlainTextEdit()
        self.shelf_text_area.setPlaceholderText("T-S NS 192.21\nMS heb. e.34/30\nMs. EVR II B 1011\nMs. Kaufmann GEN 227/A")
        self.shelf_text_area.textChanged.connect(self._on_shelf_text_changed)

        self.title_text_area = QPlainTextEdit()
        self.title_text_area.setPlaceholderText(tr("Title"))
        self.title_text_area.setReadOnly(True)
        self.title_text_area.setLineWrapMode(QPlainTextEdit.LineWrapMode.NoWrap)

        self.sys_text_area.installEventFilter(self)
        self.shelf_text_area.installEventFilter(self)
        self.title_text_area.installEventFilter(self)

        grid.addWidget(self.sys_text_area, 1, 0)
        grid.addWidget(self.shelf_text_area, 1, 1)
        grid.addWidget(self.title_text_area, 1, 2)
        layout.addLayout(grid)

        if existing_entries:
            sys_entries, shelf_entries = self._split_existing_entries(existing_entries)
            if sys_entries:
                self.sys_text_area.setPlainText("\n".join(sys_entries))
            if shelf_entries:
                self.shelf_text_area.setPlainText("\n".join(shelf_entries))
            if sys_entries and not shelf_entries:
                self._last_edited = "sys"
                self._sync_from_sys()
            elif shelf_entries and not sys_entries:
                self._last_edited = "shelf"
                self._sync_from_shelf()
            elif sys_entries:
                self._set_titles(self._resolve_titles_from_sys(sys_entries))

        btn_row = QHBoxLayout()
        self.btn_load = QPushButton(tr("Load from File"))
        self.btn_load.clicked.connect(self.load_file)
        btn_row.addWidget(self.btn_load)

        btn_row.addStretch()
        btn_apply = QPushButton(tr("Apply"))
        btn_apply.clicked.connect(self.accept)
        btn_cancel = QPushButton(tr("Cancel"))
        btn_cancel.clicked.connect(self.reject)
        btn_row.addWidget(btn_cancel)
        btn_row.addWidget(btn_apply)
        layout.addLayout(btn_row)

        self.setLayout(layout)

    def eventFilter(self, obj, event):
        if event.type() == QEvent.Type.FocusIn:
            if obj is self.sys_text_area:
                self._last_edited = "sys"
            elif obj is self.shelf_text_area:
                self._last_edited = "shelf"
        if obj is self.title_text_area and event.type() == QEvent.Type.ToolTip:
            cursor = self.title_text_area.cursorForPosition(event.pos())
            line_idx = cursor.blockNumber()
            if 0 <= line_idx < len(self._full_titles):
                full_title = self._full_titles[line_idx]
                display_title = self._display_titles[line_idx]
                if full_title and full_title != display_title:
                    QToolTip.showText(event.globalPos(), full_title, self.title_text_area)
                    return True
            QToolTip.hideText()
            return True
        return super().eventFilter(obj, event)

    def _split_existing_entries(self, entries):
        sys_entries = []
        shelf_entries = []
        for entry in entries:
            cleaned = re.sub(r"\s+", "", entry or "")
            digits_only = re.sub(r"\D", "", cleaned)
            if digits_only and digits_only == cleaned:
                sys_entries.append(cleaned)
            else:
                stripped = (entry or "").strip()
                if stripped:
                    shelf_entries.append(stripped)
        return sys_entries, shelf_entries

    def _on_sys_text_changed(self):
        if self._syncing or self._last_edited != "sys":
            return
        self._sync_from_sys()

    def _on_shelf_text_changed(self):
        if self._syncing or self._last_edited != "shelf":
            return
        self._sync_from_shelf()

    def _sync_from_sys(self):
        self._syncing = True
        sys_lines = self._get_lines(self.sys_text_area.toPlainText())
        shelves = self._resolve_shelves_from_sys(sys_lines)
        titles = self._resolve_titles_from_sys(sys_lines)
        self.shelf_text_area.setPlainText("\n".join(shelves))
        self._set_titles(titles)
        self._syncing = False

    def _sync_from_shelf(self):
        self._syncing = True
        shelf_lines = self._get_lines(self.shelf_text_area.toPlainText())
        sys_ids = self._resolve_sys_from_shelves(shelf_lines)
        self.sys_text_area.setPlainText("\n".join(sys_ids))
        titles = self._resolve_titles_from_sys(sys_ids)
        self._set_titles(titles)
        self._syncing = False

    def resizeEvent(self, event):
        super().resizeEvent(event)
        if self._full_titles:
            self._refresh_title_display()

    def _get_lines(self, text):
        return text.splitlines()

    def _set_titles(self, titles):
        self._full_titles = titles
        self._refresh_title_display()

    def _refresh_title_display(self):
        metrics = QFontMetrics(self.title_text_area.font())
        width = max(self.title_text_area.viewport().width() - 6, 20)
        self._display_titles = [
            metrics.elidedText(title, Qt.TextElideMode.ElideRight, width) if title else ""
            for title in self._full_titles
        ]
        self.title_text_area.setPlainText("\n".join(self._display_titles))

    def _resolve_shelves_from_sys(self, sys_lines):
        shelves = []
        for line in sys_lines:
            cleaned = re.sub(r"\D", "", line or "")
            if not cleaned or not self.meta_mgr:
                shelves.append("")
                continue
            shelf, _ = self.meta_mgr.get_meta_for_id(cleaned)
            if shelf == "Unknown" and cleaned not in self.meta_mgr.nli_cache:
                self.meta_mgr.fetch_nli_data(cleaned)
                shelf, _ = self.meta_mgr.get_meta_for_id(cleaned)
            shelves.append("" if shelf == "Unknown" else shelf)
        return shelves

    def _resolve_titles_from_sys(self, sys_lines):
        titles = []
        for line in sys_lines:
            cleaned = re.sub(r"\D", "", line or "")
            if not cleaned or not self.meta_mgr:
                titles.append("")
                continue
            _, title = self.meta_mgr.get_meta_for_id(cleaned)
            if not title and cleaned not in self.meta_mgr.nli_cache:
                self.meta_mgr.fetch_nli_data(cleaned)
                _, title = self.meta_mgr.get_meta_for_id(cleaned)
            titles.append(title or "")
        return titles

    def _ensure_shelf_map(self):
        if self._shelf_to_sys is not None:
            return
        self._shelf_to_sys = {}
        if not self.meta_mgr:
            return
        for sys_id, meta in self.meta_mgr.csv_bank.items():
            self._add_shelf_map(meta.get("shelfmark"), sys_id)
        for sys_id, meta in self.meta_mgr.nli_cache.items():
            self._add_shelf_map(meta.get("shelfmark"), sys_id)

    def _add_shelf_map(self, shelf, sys_id):
        norm = self._normalize_shelfmark(shelf)
        if norm and norm not in self._shelf_to_sys:
            self._shelf_to_sys[norm] = sys_id

    def _resolve_sys_from_shelves(self, shelf_lines):
        self._ensure_shelf_map()
        sys_ids = []
        for line in shelf_lines:
            norm = self._normalize_shelfmark(line)
            sys_ids.append(self._shelf_to_sys.get(norm, "") if norm else "")
        return sys_ids

    def _normalize_shelfmark(self, shelf):
        if not shelf:
            return ""
        without_prefix = re.sub(r"^\s*m[\.\s]*s[\.\s]*\.?\s*", "", shelf, flags=re.IGNORECASE)
        cleaned = re.sub(r"[^\w]", "", without_prefix).lower()
        if cleaned.startswith("ms"):
            cleaned = cleaned[2:]
        return cleaned

    def load_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Load", "", "Text (*.txt)")
        if path:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            entries = [line for line in content.splitlines() if line.strip()]
            sys_entries, shelf_entries = self._split_existing_entries(entries)
            self._syncing = True
            self.sys_text_area.setPlainText("\n".join(sys_entries))
            self.shelf_text_area.setPlainText("\n".join(shelf_entries))
            self._syncing = False
            if sys_entries and not shelf_entries:
                self._last_edited = "sys"
                self._sync_from_sys()
            elif shelf_entries and not sys_entries:
                self._last_edited = "shelf"
                self._sync_from_shelf()
            elif sys_entries:
                self._set_titles(self._resolve_titles_from_sys(sys_entries))

    def get_entries_text(self):
        entries = []
        seen = set()

        sys_lines = self._get_lines(self.sys_text_area.toPlainText())
        for line in sys_lines:
            cleaned = re.sub(r"\D", "", line or "")
            if cleaned and cleaned not in seen:
                entries.append(cleaned)
                seen.add(cleaned)

        shelf_lines = self._get_lines(self.shelf_text_area.toPlainText())
        for line in shelf_lines:
            stripped = (line or "").strip()
            if stripped and stripped not in seen:
                entries.append(stripped)
                seen.add(stripped)

        return "\n".join(entries)

class ResultDialog(QDialog):
    """Allow browsing a single search result and its surrounding pages."""

    metadata_loaded = pyqtSignal(int, dict)
    thumb_resolved = pyqtSignal(str, object)

    def __init__(self, parent, all_results, current_index, meta_mgr, searcher):
        super().__init__(parent)
        
        self.all_results = all_results
        self.current_result_idx = current_index
        self.meta_mgr = meta_mgr
        self.searcher = searcher
        self.thumb_resolved.connect(self._on_thumb_resolved)
        
        # State for internal browsing
        self.current_sys_id = None
        self.current_p_num = None
        self.current_fl_id = None
        self.current_page_text = None
        self.current_page_uid = None
        self.current_internal_idx = None
        
        self.current_meta_request = 0
        self.extended_info_visible = False
        self.external_url = None

        # External Viewer State
        self.ext_data = None
        self.ext_canvases = []

        self.init_ui()
        self.metadata_loaded.connect(self.on_metadata_loaded)
        self.load_result_by_index(self.current_result_idx)

    def init_ui(self):
        self.setWindowTitle(tr("Manuscript Viewer"))
        self.resize(1300, 850) # Wider for split view
        
        main_layout = QVBoxLayout()
        
        # --- Top Bar (Result Nav) ---
        top_bar = QHBoxLayout()
        self.btn_res_prev = QPushButton(tr("◀ Prev Result")); self.btn_res_prev.clicked.connect(lambda: self.navigate_results(-1))
        self.lbl_res_count = QLabel(); self.lbl_res_count.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.btn_compact_toggle = QPushButton("⏶")
        self.btn_compact_toggle.setToolTip(tr("Compact"))
        self.btn_compact_toggle.setCheckable(True)
        self.btn_compact_toggle.setChecked(False)
        self.btn_compact_toggle.setFixedWidth(36)
        self.btn_compact_toggle.clicked.connect(lambda checked: self._toggle_compact_mode(checked))
        self.btn_res_next = QPushButton(tr("Next Result ▶")); self.btn_res_next.clicked.connect(lambda: self.navigate_results(1))
        top_bar.addWidget(self.btn_res_prev); top_bar.addWidget(self.lbl_res_count, 1); top_bar.addWidget(self.btn_compact_toggle); top_bar.addWidget(self.btn_res_next)
        main_layout.addLayout(top_bar)
        main_layout.addWidget(QSplitter(Qt.Orientation.Horizontal))

        # --- Compact Bar (initially hidden, shown in compact mode) ---
        self.compact_bar = QWidget()
        self.compact_bar.setVisible(False)
        compact_layout = QHBoxLayout(self.compact_bar)
        compact_layout.setContentsMargins(4, 2, 4, 2)
        compact_layout.setSpacing(6)

        # Shelfmark (compact)
        self.lbl_compact_shelf = QLabel()
        self.lbl_compact_shelf.setFont(QFont("Arial", 13, QFont.Weight.Bold))
        self.lbl_compact_shelf.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        compact_layout.addWidget(self.lbl_compact_shelf)

        compact_layout.addWidget(QLabel(" | "))

        # Image navigation (compact)
        compact_layout.addWidget(QLabel(tr("Image:")))
        self.btn_compact_pg_prev = QPushButton("<")
        self.btn_compact_pg_prev.setFixedWidth(25)
        self.btn_compact_pg_prev.clicked.connect(lambda: self.load_page(offset=-1))
        compact_layout.addWidget(self.btn_compact_pg_prev)

        self.lbl_compact_page = QLabel("1 / ?")
        self.lbl_compact_page.setMinimumWidth(50)
        self.lbl_compact_page.setAlignment(Qt.AlignmentFlag.AlignCenter)
        compact_layout.addWidget(self.lbl_compact_page)

        self.btn_compact_pg_next = QPushButton(">")
        self.btn_compact_pg_next.setFixedWidth(25)
        self.btn_compact_pg_next.clicked.connect(lambda: self.load_page(offset=1))
        compact_layout.addWidget(self.btn_compact_pg_next)

        compact_layout.addWidget(QLabel(" | "))

        # Add to List (compact)
        self.btn_compact_add_list = QPushButton(_format_add_to_list_label(False))
        self.btn_compact_add_list.clicked.connect(self.add_current_to_list)
        compact_layout.addWidget(self.btn_compact_add_list)

        # Extended Info (compact)
        self.btn_compact_ext_info = QPushButton(tr("Show Extended Info"))
        self.btn_compact_ext_info.setCheckable(True)
        self.btn_compact_ext_info.setVisible(False)  # shown when extended info available
        self.btn_compact_ext_info.toggled.connect(self.toggle_extended_info)
        compact_layout.addWidget(self.btn_compact_ext_info)

        # Bib buttons (compact)
        self.btn_compact_bib_fjms = QPushButton()
        self.btn_compact_bib_fjms.setVisible(False)
        self.btn_compact_bib_fjms.clicked.connect(self._show_rd_fjms_bib)
        compact_layout.addWidget(self.btn_compact_bib_fjms)
        self.btn_compact_bib_nli = QPushButton()
        self.btn_compact_bib_nli.setVisible(False)
        self.btn_compact_bib_nli.clicked.connect(self._show_rd_nli_bib)
        compact_layout.addWidget(self.btn_compact_bib_nli)

        # Catalog Records (compact)
        self.btn_compact_catalog = QPushButton()
        self.btn_compact_catalog.setVisible(False)
        self.btn_compact_catalog.clicked.connect(self._show_rd_catalog)
        compact_layout.addWidget(self.btn_compact_catalog)

        # Joins (compact) - chain icon like normal mode
        self.btn_compact_joins = QToolButton()
        self.btn_compact_joins.setText("🔗")
        self.btn_compact_joins.setToolTip(tr("View joined fragments"))
        self.btn_compact_joins.setFixedSize(40, 32)
        self.btn_compact_joins.setStyleSheet("background-color: #95a5a6; color: white; border-radius: 4px;")
        self.btn_compact_joins.setPopupMode(QToolButton.ToolButtonPopupMode.MenuButtonPopup)
        self.btn_compact_joins.clicked.connect(self._rd_view_joins)
        compact_layout.addWidget(self.btn_compact_joins)

        compact_layout.addStretch()

        main_layout.addWidget(self.compact_bar)

        # --- Header ---
        header_widget = QWidget()
        header_layout = QHBoxLayout(header_widget); header_layout.setContentsMargins(0, 5, 0, 10)
        
        # Left: Meta + Controls
        meta_col = QVBoxLayout(); meta_col.setAlignment(Qt.AlignmentFlag.AlignTop); meta_col.setSpacing(4)
        
        self.lbl_shelf = QLabel(); self.lbl_shelf.setFont(QFont("Arial", 16, QFont.Weight.Bold)); self.lbl_shelf.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.lbl_title = QLabel(); self.lbl_title.setFont(QFont("Arial", 14)); self.lbl_title.setAlignment(Qt.AlignmentFlag.AlignLeft); self.lbl_title.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)

        # Controls Row
        info_row = QHBoxLayout()
        self.btn_img = QPushButton(tr("Go to Ktiv")); self.btn_img.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_DialogHelpButton)); self.btn_img.clicked.connect(self.open_catalog); self.btn_img.setFixedWidth(100)
        self.btn_external_link = QPushButton(tr("External Website"))
        self.btn_external_link.setVisible(False)
        self.btn_external_link.clicked.connect(self.open_external_link)
        self.lbl_info = QLabel(); self.lbl_info.setStyleSheet("font-size: 11px;"); self.lbl_info.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self.lbl_meta_loading = QLabel(tr("Loading...")); self.lbl_meta_loading.setStyleSheet("color: orange; font-size: 11px;"); self.lbl_meta_loading.setVisible(False)

        # Domain info (inlined on info_row)
        self.lbl_rd_domains = QLabel("")
        self.lbl_rd_domains.setStyleSheet("color: #8e44ad; font-size: 11px;")
        self.lbl_rd_domains.setVisible(False)

        info_row.addWidget(self.btn_img); info_row.addWidget(self.btn_external_link); info_row.addWidget(self.lbl_info); info_row.addWidget(self.lbl_rd_domains); info_row.addWidget(self.lbl_meta_loading); info_row.addStretch()

        # Nav Row (Inside Header)
        nav_row = QHBoxLayout()

        # Arrows logic (Standard: Prev <, Next > regardless of RTL)
        prev_arrow = "<"
        next_arrow = ">"

        btn_pg_prev = QPushButton(prev_arrow); btn_pg_prev.setFixedWidth(30); btn_pg_prev.clicked.connect(lambda: self.load_page(offset=-1))
        self.spin_page = QSpinBox(); self.spin_page.setRange(1, 9999); self.spin_page.setFixedWidth(80); self.spin_page.editingFinished.connect(lambda: self.load_page(target=self.spin_page.value()))
        btn_pg_next = QPushButton(next_arrow); btn_pg_next.setFixedWidth(30); btn_pg_next.clicked.connect(lambda: self.load_page(offset=1))
        self.lbl_total = QLabel("/ ?")

        self.lbl_img_label = QLabel("")
        self.lbl_img_label.setStyleSheet("color: #2980b9; font-weight: bold; margin-left: 10px;")

        nav_row.addWidget(QLabel(tr("Image:"))); nav_row.addWidget(btn_pg_prev); nav_row.addWidget(self.spin_page);
        nav_row.addWidget(self.lbl_total); nav_row.addWidget(btn_pg_next); nav_row.addWidget(self.lbl_img_label); nav_row.addStretch()

        action_row = QHBoxLayout()
        self.btn_view_transcription = QPushButton(tr("Browse manuscript")) # Renamed
        self.btn_view_transcription.clicked.connect(self.open_full_transcription)
        self.btn_search_parallels = QPushButton(tr("Search for parallels"))
        self.btn_search_parallels.clicked.connect(self.search_for_parallels)

        # Add to List button
        self.btn_add_to_list = QPushButton(_format_add_to_list_label(False))
        self.btn_add_to_list.clicked.connect(self.add_current_to_list)

        self.btn_ext_info = QPushButton(tr("Show Extended Info"))
        self.btn_ext_info.setCheckable(True)
        self.btn_ext_info.toggled.connect(self.toggle_extended_info)
        self.btn_ext_info.setVisible(False)

        # Toggle Image Button
        self.btn_toggle_image = QPushButton("🖼️")
        self.btn_toggle_image.setCheckable(True)
        self.btn_toggle_image.setChecked(True) # Default open
        self.btn_toggle_image.clicked.connect(self.toggle_external_viewer)
        self.btn_toggle_image.setVisible(False) # Hidden until images avail

        # Deprecated: btn_external_view replaced/merged logic
        self.btn_external_view = self.btn_toggle_image

        self.btn_rd_bib_fjms = QPushButton()
        self.btn_rd_bib_fjms.setVisible(False)
        self.btn_rd_bib_fjms.clicked.connect(self._show_rd_fjms_bib)
        self.btn_rd_bib_nli = QPushButton()
        self.btn_rd_bib_nli.setVisible(False)
        self.btn_rd_bib_nli.clicked.connect(self._show_rd_nli_bib)
        self.btn_rd_catalog = QPushButton(f"{tr('Catalog Records')} (0)")
        self.btn_rd_catalog.setEnabled(False)
        self.btn_rd_catalog.setVisible(False)
        self.btn_rd_catalog.clicked.connect(self._show_rd_catalog)
        self._rd_fjms_bib = []
        self._rd_marc_bib = []
        self._rd_catalog_detail = None

        action_row.addWidget(self.btn_view_transcription)
        action_row.addWidget(self.btn_search_parallels)
        action_row.addWidget(self.btn_add_to_list)
        action_row.addWidget(self.btn_ext_info)
        action_row.addWidget(self.btn_rd_bib_fjms)
        action_row.addWidget(self.btn_rd_bib_nli)
        action_row.addWidget(self.btn_rd_catalog)
        action_row.addWidget(self.btn_toggle_image)

        action_row.addStretch()

        # --- Second row: Community features (Edit, Version, Comment) ---
        community_row = QHBoxLayout()

        # Version selector
        community_row.addWidget(QLabel(tr("Version:")))
        self.rd_version_combo = QComboBox()
        self.rd_version_combo.addItem("V0.8", {"source": "original"})
        self.rd_version_combo.setFixedWidth(240)  # Wider for PGP scholar names
        self.rd_version_combo.setEnabled(False)
        self.rd_version_combo.currentIndexChanged.connect(self._rd_change_version)
        community_row.addWidget(self.rd_version_combo)
        self._rd_versions_cache = {}

        community_row.addWidget(QLabel(" | "))

        # Edit button
        self.btn_rd_edit = QPushButton(tr("✏️ Edit"))
        self.btn_rd_edit.setToolTip(tr("Enable edit mode to make corrections"))
        self.btn_rd_edit.clicked.connect(self._rd_toggle_edit_mode)
        community_row.addWidget(self.btn_rd_edit)

        # Edit action buttons (hidden by default, shown in edit mode)
        self.btn_rd_save_draft = QPushButton(f"💾 {tr('Save')}")
        self.btn_rd_save_draft.clicked.connect(lambda: self._rd_save_correction(submit=False))
        self.btn_rd_save_draft.setEnabled(False)
        self.btn_rd_save_draft.setVisible(False)
        community_row.addWidget(self.btn_rd_save_draft)

        self.btn_rd_submit = QPushButton(f"📤 {tr('Submit')}")
        self.btn_rd_submit.clicked.connect(lambda: self._rd_save_correction(submit=True))
        self.btn_rd_submit.setEnabled(False)
        self.btn_rd_submit.setVisible(False)
        community_row.addWidget(self.btn_rd_submit)

        self.btn_rd_cancel_edit = QPushButton(tr("Cancel"))
        self.btn_rd_cancel_edit.clicked.connect(self._rd_cancel_edit)
        self.btn_rd_cancel_edit.setVisible(False)
        community_row.addWidget(self.btn_rd_cancel_edit)

        # Edit status label (hidden by default)
        self.rd_edit_status = QLabel("")
        self.rd_edit_status.setVisible(False)
        community_row.addWidget(self.rd_edit_status)

        community_row.addWidget(QLabel(" | "))

        # Comment button
        self.btn_comment = QPushButton(tr("💬 Comment"))
        self.btn_comment.clicked.connect(self.add_comment)
        community_row.addWidget(self.btn_comment)

        # View Corrections button
        self.btn_view_corrections = QPushButton(tr("View Corrections"))
        self.btn_view_corrections.clicked.connect(self.view_corrections)
        community_row.addWidget(self.btn_view_corrections)

        # View Comments button (icon, visible when comments exist)
        self.btn_view_comments = QPushButton("💬")
        self.btn_view_comments.setToolTip(tr("View Comments"))
        self.btn_view_comments.setFixedSize(32, 32)
        self.btn_view_comments.setVisible(False)
        self.btn_view_comments.clicked.connect(self.view_comments)
        community_row.addWidget(self.btn_view_comments)

        # Joins button with dropdown
        self.btn_joins = QToolButton()
        self.btn_joins.setText("🔗")
        self.btn_joins.setToolTip(tr("View joined fragments"))
        self.btn_joins.setFixedSize(40, 32)
        self.btn_joins.setStyleSheet("background-color: #95a5a6; color: white; border-radius: 4px;")
        self.btn_joins.setPopupMode(QToolButton.ToolButtonPopupMode.MenuButtonPopup)
        self.btn_joins.clicked.connect(self._rd_view_joins)
        self.rd_joins_menu = QMenu(self)
        self.rd_joins_menu.aboutToShow.connect(self._rd_on_joins_menu_show)
        self.btn_joins.setMenu(self.rd_joins_menu)
        community_row.addWidget(self.btn_joins)

        community_row.addStretch()

        self.txt_extended_info = QTextBrowser()
        self.txt_extended_info.setVisible(False)
        self.txt_extended_info.setMaximumHeight(200)
        # Use standard palette (transparent background allowed) to support dark mode
        self.txt_extended_info.setStyleSheet("border: 1px solid #ccc; padding: 5px;")
        self.txt_extended_info.setOpenLinks(False)
        self.txt_extended_info.anchorClicked.connect(self._on_rd_ext_link_clicked)

        meta_col.addWidget(self.lbl_shelf); meta_col.addWidget(self.lbl_title); meta_col.addLayout(info_row); meta_col.addLayout(nav_row); meta_col.addLayout(action_row); meta_col.addLayout(community_row)

        # Thumbnail (kept as hidden dummy for compatibility with existing methods)
        self.lbl_thumb = QLabel()
        self.lbl_thumb.setVisible(False)

        header_layout.addLayout(meta_col, 1)
        self.header_widget = header_widget
        main_layout.addWidget(header_widget)

        # Extended info (moved outside header to remain visible in compact mode)
        main_layout.addWidget(self.txt_extended_info)

        # Set compact joins button menu (now that rd_joins_menu is created)
        self.btn_compact_joins.setMenu(self.rd_joins_menu)
        
        # --- SPLIT VIEW (Manuscript | Source | External) ---
        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)

        # 1. Manuscript View (Left)
        ms_widget = QWidget()
        ms_layout = QVBoxLayout(ms_widget); ms_layout.setContentsMargins(0,0,0,0)
        ms_text_widget = QWidget()
        ms_text_layout = QVBoxLayout(ms_text_widget); ms_text_layout.setContentsMargins(0,0,0,0)
        ms_text_layout.addWidget(QLabel("<b>" + tr("Manuscript Text") + "</b>"))
        ms_find_row = QHBoxLayout()
        ms_find_row.addWidget(QLabel(tr("Find:")))
        self.find_ms_input = QLineEdit()
        self.find_ms_input.setPlaceholderText(tr("Find in text..."))
        self.find_ms_input.textChanged.connect(lambda text: apply_find_highlight(self.text_ms, text.strip()))
        ms_find_row.addWidget(self.find_ms_input)
        ms_text_layout.addLayout(ms_find_row)
        self.text_ms = QTextBrowser(); self.text_ms.setFont(QFont("SBL Hebrew", 16)); self.text_ms.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
        ms_text_layout.addWidget(self.text_ms)

        # 2. Source Context (Below Manuscript Text)
        self.src_widget = QWidget() # Container to hide/show easily
        src_layout = QVBoxLayout(self.src_widget); src_layout.setContentsMargins(0,0,0,0)
        src_layout.addWidget(QLabel("<b>" + tr("Match Context (Source)") + "</b>"))
        self.text_src = QTextBrowser()
        self.text_src.setFont(QFont("SBL Hebrew", 16))
        self.text_src.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
        line_height = self.text_src.fontMetrics().lineSpacing()
        self.text_src.setMinimumHeight(line_height * 3 + 12)
        src_layout.addWidget(self.text_src)

        self.ms_text_splitter = QSplitter(Qt.Orientation.Vertical)
        self.ms_text_splitter.addWidget(ms_text_widget)
        self.ms_text_splitter.addWidget(self.src_widget)
        self.ms_text_splitter.setStretchFactor(0, 5)
        self.ms_text_splitter.setStretchFactor(1, 1)
        self.ms_text_splitter.setSizes([600, line_height * 3 + 12])
        ms_layout.addWidget(self.ms_text_splitter)

        self.main_splitter.addWidget(ms_widget)

        # 3. External Viewer Pane (Initially Hidden)
        self.external_pane = QWidget()
        self.external_pane.setVisible(False)
        ext_layout = QVBoxLayout(self.external_pane); ext_layout.setContentsMargins(0,0,0,0)

        self.lbl_ext_attr = QLabel(tr("External Viewer"))
        self.lbl_ext_attr.setStyleSheet("font-weight: bold; padding: 5px; background: #ecf0f1;")
        self.lbl_ext_attr.setWordWrap(True)

        self.txt_ext_meta = QTextBrowser()
        self.txt_ext_meta.setMaximumHeight(100)
        self.txt_ext_meta.setStyleSheet("font-size: 11px;")

        # New: Reusable Viewer Widget
        self.ms_viewer = ManuscriptViewerWidget()

        ext_layout.addWidget(self.lbl_ext_attr)
        ext_layout.addWidget(self.txt_ext_meta)
        ext_layout.addWidget(self.ms_viewer, 1)

        self.main_splitter.addWidget(self.external_pane)
        self.main_splitter.setStretchFactor(0, 1)
        self.main_splitter.setStretchFactor(1, 1)
        self.main_splitter.setSizes([650, 650])

        main_layout.addWidget(self.main_splitter, 1)
        
        # Footer
        btn_close = QPushButton("Close"); btn_close.clicked.connect(self.close); main_layout.addWidget(btn_close)
        self.setLayout(main_layout)

    def _toggle_compact_mode(self, compact):
        """Toggle between compact and full header mode."""
        self.compact_bar.setVisible(compact)
        self.header_widget.setVisible(not compact)
        self.btn_compact_toggle.setChecked(compact)
        self.btn_compact_toggle.setText("⏷" if compact else "⏶")
        self.btn_compact_toggle.setToolTip(tr("Expand") if compact else tr("Compact"))

        if compact:
            # Sync compact bar state from full header
            self.lbl_compact_shelf.setText(self.lbl_shelf.text())
            page_num = self.spin_page.value()
            total_text = self.lbl_total.text()  # "/ N"
            self.lbl_compact_page.setText(f"{page_num} {total_text}")

            # Sync extended info button state
            self.btn_compact_ext_info.setVisible(self.btn_ext_info.isVisible())
            self.btn_compact_ext_info.blockSignals(True)
            self.btn_compact_ext_info.setChecked(self.btn_ext_info.isChecked())
            self.btn_compact_ext_info.blockSignals(False)
            self.btn_compact_ext_info.setText(self.btn_ext_info.text())

    def navigate_results(self, direction):
        new_idx = self.current_result_idx + direction
        if 0 <= new_idx < len(self.all_results):
            self.current_result_idx = new_idx
            self.load_result_by_index(new_idx)

    def open_full_transcription(self):
        parent = self.parent()
        if parent and hasattr(parent, "open_result_in_browse"):
            parent.open_result_in_browse(
                self.data,
                shelfmark=self.lbl_shelf.text(),
                title=self.lbl_title.text(),
                fl_id=self.current_fl_id,
            )
            self.close()

    def search_for_parallels(self):
        parent = self.parent()
        if parent and hasattr(parent, "send_result_to_composition"):
            # Trim title to first 6 words and append ... if longer
            full_title = self.lbl_title.text() or ""
            words = full_title.split()
            if len(words) > 6:
                short_title = " ".join(words[:6]) + "..."
            else:
                short_title = full_title

            parent.send_result_to_composition(
                self.data,
                source_text=self.current_page_text,
                title=short_title,
            )
            self.close()

    def add_current_to_list(self):
        """Add the current manuscript to a list."""
        parent = self.parent()
        if not parent or not hasattr(parent, 'lists_mgr') or not parent.lists_mgr:
            return

        # Get system ID from current result
        sys_id = None
        if self.data:
            display = self.data.get('display', {})
            sys_id = display.get('id')

        if sys_id:
            fl_id = parent._normalize_fl_id(self.current_fl_id)
            img = self.current_p_num
            parent.show_add_to_list_menu(
                [{'sys_id': sys_id, 'fl_id': fl_id, 'img': img}],
                source=tr("from browse"),
                anchor_widget=self.btn_add_to_list
            )
            # Also add to recently viewed
            parent.lists_mgr.add_to_recent(sys_id, fl_id=fl_id, img=img)
            self._update_add_to_list_button()

    def _update_add_to_list_button(self):
        parent = self.parent()
        if not parent or not hasattr(parent, 'lists_mgr') or not parent.lists_mgr:
            return
        if not self.current_sys_id:
            return
        in_list = parent._is_item_in_non_recent_list(
            self.current_sys_id,
            img=self.current_p_num,
            fl_id=parent._normalize_fl_id(self.current_fl_id),
        )
        label = _format_add_to_list_label(in_list)
        self.btn_add_to_list.setText(label)
        if hasattr(self, 'btn_compact_add_list'):
            self.btn_compact_add_list.setText(label)

    def add_comment(self):
        """Open comment dialog for current document."""
        parent = self.parent()
        if not parent or not hasattr(parent, 'corrections_client'):
            return
        if not parent.corrections_client.is_logged_in():
            QMessageBox.warning(self, tr("Login Required"), tr("Please login to add a comment."))
            return
        dialog = CommentDialog(
            self, parent.corrections_client,
            document_id=self.current_sys_id,
            shelfmark=self.lbl_shelf.text(),
            page_number=self.current_p_num
        )
        dialog.exec()

    def view_corrections(self):
        """View corrections for current document."""
        parent = self.parent()
        if not parent or not hasattr(parent, 'corrections_client'):
            return
        dialog = CorrectionsViewerDialog(
            self, parent.corrections_client,
            document_id=self.current_sys_id,
            shelfmark=self.lbl_shelf.text(),
            on_view_result=lambda s: parent._open_document_result_dialog(shelfmark=s) if hasattr(parent, '_open_document_result_dialog') else None,
            on_browse=lambda s: parent._browse_document_by_shelfmark(s) if hasattr(parent, '_browse_document_by_shelfmark') else None
        )
        dialog.exec()

    def view_comments(self):
        """View comments for current document."""
        parent = self.parent()
        if not parent or not hasattr(parent, 'corrections_client'):
            return
        dialog = CommentsViewerDialog(
            self, parent.corrections_client,
            document_id=self.current_sys_id,
            shelfmark=self.lbl_shelf.text()
        )
        dialog.exec()

    def _rd_view_joins(self):
        """View joined fragments for current document."""
        parent = self.parent()
        if not parent or not hasattr(parent, 'corrections_client'):
            return

        shelfmark = self.lbl_shelf.text()
        if not shelfmark:
            return

        def navigate_to_shelfmark(target_shelfmark):
            """Navigate to a shelfmark within the same results dialog."""
            # Note: JoinsDialog already closes itself before calling this callback
            # Load the document in the same ResultDialog
            self.load_by_shelfmark(target_shelfmark)

        dialog = JoinsDialog(
            self, parent.corrections_client,
            document_id=self.current_sys_id,
            shelfmark=shelfmark,
            on_browse=navigate_to_shelfmark,
            shelf_model=getattr(parent, 'shelf_model', None),
            joins_mgr=getattr(parent, 'joins_mgr', None),
            shelf_completer=getattr(parent, 'shelf_completer', None),
            lists_mgr=getattr(parent, 'lists_mgr', None),
            meta_mgr=getattr(parent, 'meta_mgr', None)
        )
        dialog.exec()

    def _rd_update_joins_menu(self):
        """Update the joins dropdown menu with connected fragments."""
        self.rd_joins_menu.clear()
        parent = self.parent()

        # Use document_id (sys_id) for lookup - this is the reliable key
        document_id = self.current_sys_id
        display_shelfmark = self.lbl_shelf.text()  # For display purposes only

        if not document_id:
            action = self.rd_joins_menu.addAction(tr("No document ID"))
            action.setEnabled(False)
            return

        # Get joins from JoinsManager using document_id (offline-first)
        connected = None
        plain_shelfmark = display_shelfmark.split(' | ')[-1] if ' | ' in display_shelfmark else display_shelfmark

        if parent and hasattr(parent, 'joins_mgr') and parent.joins_mgr:
            # Debug: show what's in the indexes
            joins_mgr = parent.joins_mgr
            by_doc_id = joins_mgr.data.get('by_document_id', {})
            by_normalized = joins_mgr.data.get('by_normalized', {})
            total_joins = len(joins_mgr.data.get('joins', {}))
            logger.debug("ResultDialog joins: total_joins=%s, by_document_id=%s, by_normalized=%s", total_joins, len(by_doc_id), len(by_normalized))
            logger.debug("Looking for doc_id='%s', plain_shelfmark='%s'", document_id, plain_shelfmark)

            # First try document_id lookup
            if document_id in by_doc_id:
                logger.debug("Found in by_document_id with join_ids: %s", by_doc_id[document_id])
            connected = joins_mgr.get_connected_fragments_by_id(document_id)

            # If no results by document_id, try shelfmark
            if not connected or connected.get('total_fragments', 0) <= 1:
                normalized = joins_mgr._normalize_shelfmark(plain_shelfmark)
                logger.debug("Not found by doc_id, trying normalized shelfmark: '%s'", normalized)
                if normalized in by_normalized:
                    logger.debug("Found in by_normalized with join_ids: %s", by_normalized[normalized])
                connected = joins_mgr.get_connected_fragments(plain_shelfmark)

            logger.debug("Final connected result: fragments=%s", connected.get('fragments', []) if connected else 'None')

        if not connected or connected.get('total_fragments', 0) <= 1:
            # Check PGP multi-fragment joins as fallback
            try:
                from shared.document_service import get_document_for_fragment, get_fragments_for_document
                pgp_doc = get_document_for_fragment(self.current_sys_id)
                if pgp_doc:
                    pgp_frags = get_fragments_for_document(pgp_doc.get('pgpid'))
                    if pgp_frags and len(pgp_frags) > 1:
                        self.btn_joins.setStyleSheet("background-color: #27ae60; color: white; border-radius: 4px;")
                        if hasattr(self, 'btn_compact_joins'):
                            self.btn_compact_joins.setStyleSheet("background-color: #27ae60; color: white; border-radius: 4px;")
                        header_action = self.rd_joins_menu.addAction(
                            tr("{} connected fragments").format(len(pgp_frags)) + " [PGP]"
                        )
                        header_action.setEnabled(False)
                        self.rd_joins_menu.addSeparator()
                        for frag in pgp_frags:
                            frag_sid = frag.get('sys_id', '')
                            frag_shelf = frag.get('shelfmark', frag_sid)
                            if frag_sid == self.current_sys_id:
                                continue
                            action = self.rd_joins_menu.addAction(f"[PGP] {frag_shelf}")
                            action.triggered.connect(lambda checked, sh=frag_shelf: self._rd_navigate_to_joined_fragment(sh))
                        return
            except Exception as e:
                logger.debug("PGP joins RD dropdown fallback error: %s", e)

            action = self.rd_joins_menu.addAction(tr("No joined fragments"))
            action.setEnabled(False)
            self.btn_joins.setStyleSheet("background-color: #95a5a6; color: white; border-radius: 4px;")
            if hasattr(self, 'btn_compact_joins'):
                self.btn_compact_joins.setStyleSheet("background-color: #95a5a6; color: white; border-radius: 4px;")
            return

        # Has joins - update button style
        self.btn_joins.setStyleSheet("background-color: #27ae60; color: white; border-radius: 4px;")
        if hasattr(self, 'btn_compact_joins'):
            self.btn_compact_joins.setStyleSheet("background-color: #27ae60; color: white; border-radius: 4px;")

        header_action = self.rd_joins_menu.addAction(
            tr("{} connected fragments").format(connected.get('total_fragments', 0))
        )
        header_action.setEnabled(False)
        self.rd_joins_menu.addSeparator()

        fragments_list = connected.get('fragments', []) if connected else []
        joins_list = connected.get('joins', []) if connected else []
        fragment_details = connected.get('fragment_details', []) if connected else []

        # Extract plain shelfmark for comparison
        plain_shelfmark = display_shelfmark.split(' | ')[-1] if ' | ' in display_shelfmark else display_shelfmark

        # Build set of directly connected fragments
        direct_fragments = set()
        for join in joins_list:
            frag_a = join.get('fragment_a', '') if isinstance(join, dict) else getattr(join, 'fragment_a', '')
            frag_b = join.get('fragment_b', '') if isinstance(join, dict) else getattr(join, 'fragment_b', '')
            if frag_a.upper() == plain_shelfmark.upper():
                direct_fragments.add(frag_b.upper())
            elif frag_b.upper() == plain_shelfmark.upper():
                direct_fragments.add(frag_a.upper())

        # Build map of shelfmark -> document_id from fragment_details for title lookup
        shelfmark_to_docid = {}
        for fd in fragment_details:
            shelf = fd.get('shelfmark', '') if isinstance(fd, dict) else getattr(fd, 'shelfmark', '')
            doc_id = fd.get('document_id') if isinstance(fd, dict) else getattr(fd, 'document_id', None)
            if shelf and doc_id:
                shelfmark_to_docid[shelf.upper()] = doc_id

        logger.debug("_rd_update_joins_menu: doc_id='%s', plain_shelfmark='%s', direct=%s", document_id, plain_shelfmark, direct_fragments)
        for frag in fragments_list:
            # Compare with plain shelfmark (joins store plain shelfmarks)
            is_current = frag.upper() == plain_shelfmark.upper()
            is_direct = frag.upper() in direct_fragments

            # Get title for display
            title_preview = ""
            frag_doc_id = shelfmark_to_docid.get(frag.upper())

            # Fallback: use parent's _shelf_to_sys map from csv_bank
            if not frag_doc_id and parent and hasattr(parent, '_shelf_to_sys') and parent._shelf_to_sys:
                norm = parent._normalize_shelfmark(frag) if hasattr(parent, '_normalize_shelfmark') else None
                if norm:
                    frag_doc_id = parent._shelf_to_sys.get(norm)

            if frag_doc_id and parent and hasattr(parent, 'meta_mgr') and parent.meta_mgr:
                try:
                    _, title = parent.meta_mgr.get_meta_for_id(frag_doc_id)
                    if title:
                        words = title.split()[:4]
                        title_preview = ' '.join(words)
                        if len(title.split()) > 4:
                            title_preview += "..."
                except:
                    pass

            if is_current:
                label = f"• {frag}"
                if title_preview:
                    label += f" - {title_preview}"
                label += f" ({tr('current')})"
                action = self.rd_joins_menu.addAction(label)
                action.setEnabled(False)
            else:
                label = f"→ {frag}"
                if title_preview:
                    label += f" - {title_preview}"
                if is_direct:
                    label += f" ({tr('direct')})"
                action = self.rd_joins_menu.addAction(label)
                action.setData(frag)
                action.triggered.connect(lambda checked, f=frag: self._rd_navigate_to_joined_fragment(f))

        self.rd_joins_menu.addSeparator()
        view_all = self.rd_joins_menu.addAction(tr("View all joins..."))
        view_all.triggered.connect(self._rd_view_joins)

    def _rd_on_joins_menu_show(self):
        """Called when joins menu is about to show - trigger sync and update."""
        parent = self.parent()
        # Trigger a background sync to get latest joins from server
        if parent and hasattr(parent, 'joins_mgr') and parent.joins_mgr:
            import threading
            def sync_and_update():
                parent.joins_mgr.sync_with_server()
            threading.Thread(target=sync_and_update, daemon=True).start()
        # Update menu with current data
        self._rd_update_joins_menu()

    def _rd_navigate_to_joined_fragment(self, shelfmark: str):
        """Navigate to a joined fragment within the same results dialog."""
        # Load the document in the same ResultDialog instead of switching to browse tab
        self.load_by_shelfmark(shelfmark)

    def _rd_load_versions(self):
        """Load versions for current document page."""
        parent = self.parent()
        if not parent or not hasattr(parent, 'corrections_client'):
            return

        doc_id = self.current_sys_id
        page_num = self.current_p_num or 1
        client = parent.corrections_client

        # Store original text
        original_text = self.text_ms.toPlainText()
        self._rd_original_text = original_text
        self._rd_versions_cache = {'original': original_text}

        # Force fresh server availability check (500ms timeout) to prevent UI freeze
        if not client.is_server_available(force_check=True):
            # Server is down - skip API calls, hide version-related UI
            self.btn_view_comments.setVisible(False)
            return

        # Check for comments
        try:
            comments = client.get_comments_for_document(doc_id, page_size=1)
            if comments and len(comments) > 0:
                self.btn_view_comments.setVisible(True)
            else:
                self.btn_view_comments.setVisible(False)
        except:
            self.btn_view_comments.setVisible(False)

        # Fetch versions and corrections using shared method
        self._rd_refresh_versions(select_latest=True)

    def _rd_change_version(self, index):
        """Handle version change in ResultDialog."""
        version_data = self.rd_version_combo.currentData()
        if version_data:
            self._rd_load_version_content(version_data)

    def _rd_refresh_versions(self, select_latest=False):
        """Refresh version list. If select_latest=True, select and load the latest version."""
        parent = self.parent()
        if not parent or not hasattr(parent, 'corrections_client'):
            return

        doc_id = self.current_sys_id
        page_num = self.current_p_num or 1
        client = parent.corrections_client

        # Quick server availability check (500ms timeout) to prevent UI freeze
        if not client.is_server_available():
            # Server is down - skip API calls
            return

        # Remember current selection
        current_data = self.rd_version_combo.currentData()

        # Reset version combo
        self.rd_version_combo.blockSignals(True)
        self.rd_version_combo.clear()
        self.rd_version_combo.addItem("V0.8", {"source": "original"})

        new_user_idx = -1  # Track user's own correction/version
        users_with_versions = set()  # Track users who have versions (to avoid duplicate corrections)

        try:
            versions_data = client.get_page_versions(doc_id, page_num)
            all_versions = versions_data.get('all_versions', [])
            logger.debug("_rd_refresh_versions: doc_id=%s, page=%s, versions=%s", doc_id, page_num, len(all_versions))
            for v in all_versions:
                logger.debug("version: source=%s, user=%s, id=%s", v.get('source'), v.get('user_name'), v.get('id'))

            # Filter to only latest version per user (use user_name as key for consistent deduplication)
            user_versions = [v for v in all_versions if v.get('source') == 'user']
            latest_by_user = {}
            for ver in user_versions:
                # Use user_name as key for deduplication (more reliable than user_id which may vary)
                user_key = ver.get('user_name', 'unknown')
                if user_key not in latest_by_user:
                    latest_by_user[user_key] = ver
                else:
                    existing = latest_by_user[user_key]
                    # Keep the one with the later created_at date
                    if ver.get('created_at', '') > existing.get('created_at', ''):
                        latest_by_user[user_key] = ver

            logger.debug("After dedup: %s unique users from %s versions", len(latest_by_user), len(user_versions))

            # Add V0.7 if available
            for ver in all_versions:
                if ver.get('source') == 'V0.7':
                    ver_id = ver.get('id')
                    is_default = ver.get('is_current_default', False)
                    label = 'V0.7'
                    if is_default:
                        label += f" ({tr('Default')})"
                    self.rd_version_combo.addItem(label, {
                        "source": "V0.7", "version_id": ver_id, "is_default": is_default
                    })

            # Add unique user versions (latest per user)
            for ver in latest_by_user.values():
                ver_id = ver.get('id')
                user_name = ver.get('user_name') or 'User'
                created_at = ver.get('created_at', '')[:10] if ver.get('created_at') else ''
                is_default = ver.get('is_current_default', False)

                label = f"{tr('by')} {user_name}"
                if created_at:
                    label += f" ({created_at})"
                if is_default:
                    label += " ✓"

                self.rd_version_combo.addItem(label, {
                    "source": "user", "version_id": ver_id, "user_name": user_name, "is_default": is_default
                })

                # Check if this is the current user's version (by username or full name)
                is_current_user = False
                if client.current_user:
                    if user_name == client.current_user.username:
                        is_current_user = True
                    elif client.current_user.full_name and user_name == client.current_user.full_name:
                        is_current_user = True

                if is_current_user:
                    new_user_idx = self.rd_version_combo.count() - 1
                    # Also add the username to tracked set (for matching with corrections)
                    users_with_versions.add(client.current_user.username)

                # Track users who already have versions (to avoid duplicates with corrections)
                users_with_versions.add(user_name)

        except Exception as e:
            logger.debug("Error refreshing versions: %s", e)

        # Also fetch corrections from corrections API (separate from versions)
        # Only add corrections for users who don't already have a version entry
        try:
            corrections = client.get_corrections_for_document(doc_id, include_drafts=True)
            # Filter corrections by page number
            page_corrections = [c for c in corrections if c.page_number == page_num or c.page_number is None]
            logger.debug("_rd_refresh_versions: corrections=%s, page_corrections=%s", len(corrections), len(page_corrections))
            for c in corrections:
                logger.debug("corr id=%s, status=%s, author=%s, page=%s", c.id, c.status, c.author_username, c.page_number)

            # Group by user, keep latest per user
            corrections_by_user = {}
            for corr in page_corrections:
                user_key = corr.author_username or f"user_{corr.author_id}"
                if user_key not in corrections_by_user:
                    corrections_by_user[user_key] = corr
                else:
                    existing = corrections_by_user[user_key]
                    if (corr.created_at or '') > (existing.created_at or ''):
                        corrections_by_user[user_key] = corr

            # Determine user permissions for viewing corrections
            current_username = client.current_user.username if client.current_user else None
            is_reviewer_or_admin = client.current_user and client.current_user.role in ('reviewer', 'editor', 'admin')

            for corr in corrections_by_user.values():
                user_name = corr.author_username or 'User'
                status = corr.status

                # Skip if user already has a version entry (avoid duplicates)
                if user_name in users_with_versions:
                    logger.debug("correction: user=%s SKIPPED (has version)", user_name)
                    continue

                # Filter based on status and user permissions:
                # - Authors can see their own corrections (any status)
                # - Reviewers/admins can see all corrections
                # - Regular users can only see approved corrections from others
                is_own_correction = current_username and user_name == current_username

                if status == 'rejected':
                    # Rejected corrections: only visible to author or admin
                    if not is_own_correction and not is_reviewer_or_admin:
                        logger.debug("correction: user=%s SKIPPED (rejected, not authorized)", user_name)
                        continue
                elif status in ('draft', 'pending'):
                    # Draft/Pending: only visible to author or reviewer/admin
                    if not is_own_correction and not is_reviewer_or_admin:
                        logger.debug("correction: user=%s SKIPPED (%s, not authorized)", user_name, status)
                        continue

                created_at = corr.created_at[:10] if corr.created_at else ''

                # Status indicators and label
                if status == 'draft':
                    # For drafts, just show "📝 Draft" without username
                    label = f"📝 {tr('Draft')}"
                elif status == 'pending':
                    label = f"⏳ {tr('Pending')} - {user_name}"
                elif status == 'approved':
                    label = f"✅ {tr('by')} {user_name}"
                    if created_at:
                        label += f" ({created_at})"
                elif status == 'rejected':
                    label = f"❌ {tr('Rejected')} - {user_name}"
                else:
                    label = f"{tr('by')} {user_name}"
                    if created_at:
                        label += f" ({created_at})"

                logger.debug("correction: user=%s, status=%s, id=%s", user_name, status, corr.id)

                self.rd_version_combo.addItem(label, {
                    "source": "correction",
                    "correction_id": corr.id,
                    "user_name": user_name,
                    "status": status,
                    "corrected_text": corr.corrected_text
                })

                # Check if this is the current user's correction
                if client.current_user and user_name == client.current_user.username:
                    new_user_idx = self.rd_version_combo.count() - 1

        except Exception as e:
            logger.debug("Error fetching corrections: %s", e)

        # Cache corrections/versions for re-appending after PGP combo rebuild
        self._rd_cached_corrections = []
        for i in range(self.rd_version_combo.count()):
            item_data = self.rd_version_combo.itemData(i)
            if item_data and item_data.get('source') not in ('original', 'header', None):
                self._rd_cached_corrections.append(
                    (self.rd_version_combo.itemText(i), item_data))

        # Enable combo if we have versions/corrections
        if self.rd_version_combo.count() > 1:
            self.rd_version_combo.setEnabled(True)

            if select_latest:
                # Select and load the latest (last) version as default
                latest_idx = self.rd_version_combo.count() - 1
                self.rd_version_combo.setCurrentIndex(latest_idx)
                self.rd_version_combo.blockSignals(False)
                data = self.rd_version_combo.itemData(latest_idx)
                if data and data.get('source') != 'original':
                    self._rd_load_version_content(data)
                return
            elif new_user_idx >= 0:
                # Select user's own version/correction if just saved
                self.rd_version_combo.setCurrentIndex(new_user_idx)
        else:
            self.rd_version_combo.setEnabled(False)

        self.rd_version_combo.blockSignals(False)

    def _rd_load_version_content(self, version_data):
        """Load and display version content."""
        source = version_data.get('source')
        version_id = version_data.get('version_id')
        correction_id = version_data.get('correction_id')
        source_id = version_data.get('source_id')

        # Build cache key
        if source in ('pgp_edition', 'pgp_translation'):
            cache_key = f"pgp_{source_id}" if source_id else source
        else:
            cache_key = f"{source}_{version_id or correction_id}" if (version_id or correction_id) else source

        if cache_key in self._rd_versions_cache:
            content = self._rd_versions_cache[cache_key]
            if source == 'pgp_translation':
                language = version_data.get('language', '')
                is_rtl = language != 'English'
                self._rd_display_pgp_text(content, is_rtl=is_rtl)
            elif source == 'pgp_edition':
                self._rd_display_pgp_text(content, is_rtl=True)
            else:
                self._rd_display_text(content)
            return

        if source == "original":
            # Restore RTL direction for V0.8 text
            self.text_ms.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
            if hasattr(self, '_rd_original_text'):
                self._rd_display_text(self._rd_original_text)
        elif source == "pgp_edition":
            # PGP edition content is stored directly in version_data
            content = version_data.get('content', '')
            if content:
                if source_id:
                    self._rd_versions_cache[f"pgp_{source_id}"] = content
                self._rd_display_pgp_text(content, is_rtl=True)
        elif source == "pgp_translation":
            # PGP translation content is stored directly in version_data
            content = version_data.get('content', '')
            language = version_data.get('language', '')
            if content:
                if source_id:
                    self._rd_versions_cache[f"pgp_{source_id}"] = content
                # English translations are LTR, everything else RTL
                is_rtl = language != 'English'
                self._rd_display_pgp_text(content, is_rtl=is_rtl)
        elif source == "correction":
            # Correction text is included directly in version_data
            # Restore RTL for corrections (Hebrew text)
            self.text_ms.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
            content = version_data.get('corrected_text', '')
            if content:
                self._rd_versions_cache[cache_key] = content
                self._rd_display_text(content)
        elif version_id:
            parent = self.parent()
            if parent and hasattr(parent, 'corrections_client'):
                # Restore RTL for user versions (Hebrew text)
                self.text_ms.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
                # Quick server availability check (500ms timeout) to prevent UI freeze
                if not parent.corrections_client.is_server_available():
                    return
                try:
                    ver_data = parent.corrections_client.get_version_content(version_id)
                    content = ver_data.get('content', '')
                    if content:
                        self._rd_versions_cache[cache_key] = content
                        self._rd_display_text(content)
                except Exception as e:
                    logger.debug("Error loading version: %s", e)

    def _rd_display_text(self, text):
        """Display text in the manuscript viewer."""
        if text:
            self.text_ms.setHtml(self._htmlify(text))

    def _rd_display_pgp_text(self, text, is_rtl=True):
        """Display PGP edition/translation text with proper directionality."""
        if not text:
            return
        direction = 'rtl' if is_rtl else 'ltr'
        layout_dir = Qt.LayoutDirection.RightToLeft if is_rtl else Qt.LayoutDirection.LeftToRight
        self.text_ms.setLayoutDirection(layout_dir)
        html_text = text.replace('\n', '<br>')
        self.text_ms.setHtml(f"<div dir='{direction}'>{html_text}</div>")
        self._refresh_find_highlights()

    def _on_rd_pgp_loaded(self, sys_id, sources, pgp_doc):
        """Handle PGP sources loaded from background thread."""
        # Stale-request guard: user may have navigated to a different result
        if sys_id != self.current_sys_id:
            return

        # Store PGP data
        self._rd_pgp_sources = sources
        self._rd_pgp_doc = pgp_doc

        # Handle PGP extended info display:
        # Case 1: Enriched data already built HTML -> append PGP section
        # Case 2: Enriched data ran but had nothing (early return) -> build PGP-only
        # Case 3: Enriched data hasn't arrived yet -> PGP included when it runs
        if pgp_doc:
            if getattr(self, '_rd_enriched_data_loaded', False):
                self._rd_update_extended_info_with_pgp()
            elif not self.btn_ext_info.isVisible():
                parent_win = self.parent()
                if parent_win and hasattr(parent_win, '_build_pgp_extended_info_html'):
                    pal = self.txt_extended_info.palette()
                    tc = pal.color(QPalette.ColorRole.Text).name()
                    bc = pal.color(QPalette.ColorRole.Base).name()
                    ph = parent_win._build_pgp_extended_info_html(pgp_doc, palette=pal)
                    if ph:
                        h = f"<div style='font-family:Arial; color:{tc}; background-color:{bc};'>{ph}</div>"
                        self.txt_extended_info.setHtml(h)
                        self.btn_ext_info.setVisible(True)
                        if hasattr(self, 'btn_compact_ext_info'):
                            self.btn_compact_ext_info.setVisible(True)

        if not sources:
            return

        parent = self.parent()
        if not parent:
            return

        # Populate combo with PGP items (clears and rebuilds: PGP Editions > Translations > V0.8)
        has_pgp = parent._populate_pgp_combo(self.rd_version_combo, sources, pgp_doc)

        if has_pgp:
            # Re-add cached corrections/versions after V0.8
            cached = getattr(self, '_rd_cached_corrections', [])
            if cached:
                self.rd_version_combo.blockSignals(True)
                self.rd_version_combo.insertSeparator(self.rd_version_combo.count())
                for label, data in cached:
                    self.rd_version_combo.addItem(label, data)
                self.rd_version_combo.blockSignals(False)

            # Store original V0.8 text (always refresh from current display)
            current_text = self.text_ms.toPlainText()
            if current_text:
                self._rd_original_text = current_text

            # Auto-select first PGP edition and display it
            edition_data = parent._auto_select_pgp_edition(self.rd_version_combo)
            if edition_data:
                content = edition_data.get('content', '')
                if content:
                    self._rd_display_pgp_text(content, is_rtl=True)

            self.rd_version_combo.setEnabled(True)

    def _on_rd_pgp_error(self, sys_id, error_msg):
        """Handle PGP source fetch error -- silently fall back to existing behavior."""
        logger.debug("PGP fetch error for %s: %s", sys_id, error_msg)

    def _rd_update_extended_info_with_pgp(self):
        """Append PGP metadata to existing extended info HTML (race condition handler).

        Called when PGP data arrives after on_enriched_data_loaded() already built
        the extended info HTML. Rebuilds the HTML with PGP section appended.
        """
        pgp_doc = getattr(self, '_rd_pgp_doc', None)
        if not pgp_doc:
            return
        parent = self.parent()
        if not parent or not hasattr(parent, '_build_pgp_extended_info_html'):
            return
        palette = self.txt_extended_info.palette()
        pgp_html = parent._build_pgp_extended_info_html(pgp_doc, palette=palette)
        if not pgp_html:
            return
        # Append PGP section to existing HTML content
        current_html = self.txt_extended_info.toHtml()
        # Insert PGP HTML before the final closing </div> of the wrapper
        # The wrapper div is: <div style='font-family:Arial; ...'>...</div>
        # We inject the PGP section before the last </div>
        close_idx = current_html.rfind('</div>')
        if close_idx >= 0:
            updated_html = current_html[:close_idx] + pgp_html + current_html[close_idx:]
            self.txt_extended_info.setHtml(updated_html)
        else:
            # Fallback: just append
            self.txt_extended_info.setHtml(current_html + pgp_html)
        self.btn_ext_info.setVisible(True)
        if hasattr(self, 'btn_compact_ext_info'):
            self.btn_compact_ext_info.setVisible(True)

    def _rd_toggle_edit_mode(self):
        """Toggle edit mode in ResultDialog."""
        parent = self.parent()
        if not parent or not hasattr(parent, 'corrections_client'):
            return
        if not parent.corrections_client.is_logged_in():
            QMessageBox.warning(self, tr("Login Required"), tr("Please login to edit."))
            return

        if not hasattr(self, '_rd_edit_mode'):
            self._rd_edit_mode = False

        self._rd_edit_mode = not self._rd_edit_mode

        if self._rd_edit_mode:
            # Enter edit mode - reset draft tracking
            self._rd_draft_correction_id = None
            self._rd_original_edit_text = self.text_ms.toPlainText()
            self.text_ms.setReadOnly(False)
            # Use palette-aware colors for dark mode support
            palette = self.palette()
            is_dark = palette.color(QPalette.ColorRole.Window).lightness() < 128
            if is_dark:
                edit_bg = "#3d3522"  # Dark yellowish for dark mode
            else:
                edit_bg = "#fffacd"  # Light lemon for light mode
            self.text_ms.setStyleSheet(f"background-color: {edit_bg}; border: 2px solid #f39c12;")
            # Show edit action buttons
            self.btn_rd_save_draft.setVisible(True)
            self.btn_rd_submit.setVisible(True)
            self.btn_rd_cancel_edit.setVisible(True)
            self.rd_edit_status.setVisible(True)
            self.btn_rd_edit.setText(tr("✏️ Editing..."))
            self.btn_rd_edit.setStyleSheet("background-color: #f39c12; color: white;")
            self.text_ms.textChanged.connect(self._rd_on_text_changed)
        else:
            self._rd_exit_edit_mode()

    def _rd_exit_edit_mode(self):
        """Exit edit mode."""
        self._rd_edit_mode = False
        self._rd_draft_correction_id = None
        try:
            self.text_ms.textChanged.disconnect(self._rd_on_text_changed)
        except:
            pass
        self.text_ms.setReadOnly(True)
        self.text_ms.setStyleSheet("")
        # Hide edit action buttons
        self.btn_rd_save_draft.setVisible(False)
        self.btn_rd_submit.setVisible(False)
        self.btn_rd_cancel_edit.setVisible(False)
        self.rd_edit_status.setVisible(False)
        self.rd_edit_status.setText("")
        self.btn_rd_edit.setText(tr("✏️ Edit"))
        self.btn_rd_edit.setStyleSheet("")

    def _rd_on_text_changed(self):
        """Handle text changes in edit mode."""
        current = self.text_ms.toPlainText()
        has_changes = current != getattr(self, '_rd_original_edit_text', '')
        draft_id = getattr(self, '_rd_draft_correction_id', None)
        self.btn_rd_save_draft.setEnabled(has_changes)
        # Enable submit if has changes OR has saved draft
        self.btn_rd_submit.setEnabled(has_changes or draft_id is not None)

        # Get palette-aware background color
        palette = self.palette()
        is_dark = palette.color(QPalette.ColorRole.Window).lightness() < 128
        edit_bg = "#3d3522" if is_dark else "#fffacd"

        if has_changes:
            self.rd_edit_status.setText(tr("Modified"))
            self.rd_edit_status.setStyleSheet("color: #e67e22;")
            # Orange border for unsaved changes
            self.text_ms.setStyleSheet(f"background-color: {edit_bg}; border: 2px solid #f39c12;")
        elif draft_id:
            self.rd_edit_status.setText(f"✓ {tr('Saved')}")
            self.rd_edit_status.setStyleSheet("color: #27ae60; font-weight: bold;")
            # Green border for saved draft
            self.text_ms.setStyleSheet(f"background-color: {edit_bg}; border: 2px solid #27ae60;")
        else:
            self.rd_edit_status.setText("")
            # Orange border (default edit mode)
            self.text_ms.setStyleSheet(f"background-color: {edit_bg}; border: 2px solid #f39c12;")

    def _rd_cancel_edit(self):
        """Cancel edit mode and restore original text."""
        if hasattr(self, '_rd_original_edit_text'):
            self._rd_display_text(self._rd_original_edit_text)
        self._rd_exit_edit_mode()

    def _rd_save_correction(self, submit=False):
        """Save correction from ResultDialog."""
        parent = self.parent()
        if not parent or not hasattr(parent, 'corrections_client'):
            return

        new_text = self.text_ms.toPlainText()
        original = getattr(self, '_rd_original_edit_text', new_text)

        # Check if there are changes OR if we have a saved draft to submit
        draft_correction_id = getattr(self, '_rd_draft_correction_id', None)
        has_changes = new_text != original

        if not has_changes and not draft_correction_id:
            if submit:
                QMessageBox.information(self, tr("No Changes"), tr("No changes were made to the text."))
            return

        notes = None
        if submit:
            notes, ok = QInputDialog.getMultiLineText(
                self, tr("Correction Notes"),
                tr("Please provide a brief explanation for your correction (optional):"), ""
            )
            if not ok:
                return
            notes = notes if notes else None

        try:
            # If submitting an existing draft, try submit_correction API first
            if submit and draft_correction_id:
                success, message = parent.corrections_client.submit_correction(draft_correction_id, notes)
                if success or 'approved' in message.lower():
                    # Success, or already approved (which means it succeeded earlier)
                    QMessageBox.information(self, tr("Correction Submitted"),
                        tr("Your correction has been submitted for review. Thank you for your contribution!"))
                    self._rd_exit_edit_mode()
                    self._rd_original_edit_text = new_text
                    self._rd_draft_correction_id = None
                    # Refresh versions to show the submitted correction
                    self._rd_refresh_versions()
                else:
                    # Submit failed, try creating a new pending correction instead
                    correction, create_msg = parent.corrections_client.create_correction(
                        document_id=self.current_sys_id,
                        original_text=original if original else new_text,
                        corrected_text=new_text,
                        correction_type="text_correction",
                        page_number=self.current_p_num,
                        notes=notes,
                        shelfmark=self.lbl_shelf.text(),
                        system_id=self.current_sys_id,
                        status='pending'
                    )
                    if correction:
                        QMessageBox.information(self, tr("Correction Submitted"),
                            tr("Your correction has been submitted for review. Thank you for your contribution!"))
                        self._rd_exit_edit_mode()
                        self._rd_original_edit_text = new_text
                        self._rd_draft_correction_id = None
                        self._rd_refresh_versions()
                    else:
                        QMessageBox.warning(self, tr("Error"), f"{tr('Failed to submit correction')}: {create_msg}")
            else:
                # Create new correction (draft or direct submit)
                correction, message = parent.corrections_client.create_correction(
                    document_id=self.current_sys_id,
                    original_text=original if original else new_text,
                    corrected_text=new_text,
                    correction_type="text_correction",
                    page_number=self.current_p_num,
                    notes=notes,
                    shelfmark=self.lbl_shelf.text(),
                    system_id=self.current_sys_id,
                    status='pending' if submit else 'draft',
                    save_as_draft=not submit  # Don't auto-submit when saving as draft
                )
                if correction:
                    if submit:
                        QMessageBox.information(self, tr("Correction Submitted"),
                            tr("Your correction has been submitted for review. Thank you for your contribution!"))
                        self._rd_exit_edit_mode()
                        self._rd_original_edit_text = new_text
                        self._rd_draft_correction_id = None
                        # Refresh versions to show the submitted correction
                        self._rd_refresh_versions()
                    else:
                        self.rd_edit_status.setText(f"✓ {tr('Saved')}")
                        self.rd_edit_status.setStyleSheet("color: #27ae60; font-weight: bold;")
                        self._rd_draft_correction_id = correction.id  # Store draft ID for later submit
                        self._rd_original_edit_text = new_text  # Update original to mark as saved
                        # Update border to green (saved)
                        palette = self.palette()
                        is_dark = palette.color(QPalette.ColorRole.Window).lightness() < 128
                        edit_bg = "#3d3522" if is_dark else "#fffacd"
                        self.text_ms.setStyleSheet(f"background-color: {edit_bg}; border: 2px solid #27ae60;")
                        # Keep submit button enabled after saving draft
                        self.btn_rd_submit.setEnabled(True)
                        self.btn_rd_save_draft.setEnabled(False)  # Disable save since no changes
                        # Refresh versions to show the draft
                        self._rd_refresh_versions()
                else:
                    QMessageBox.warning(self, tr("Error"), f"{tr('Failed to save correction')}: {message}")
        except Exception as e:
            QMessageBox.warning(self, tr("Error"), f"{tr('Failed to save correction')}: {str(e)}")

    def _refresh_find_highlights(self):
        apply_find_highlight(self.text_ms, self.find_ms_input.text().strip())

    def _apply_source_highlights(self, text, pattern_str):
        if not text:
            return ""
        if pattern_str and '*' not in text:
            try:
                regex = re.compile(pattern_str, re.IGNORECASE)
                text = regex.sub(r'*\g<0>*', text)
            except:
                pass
        return text

    def open_external_link(self):
        if self.external_url:
            url = self.external_url
            # Transform CUDL IIIF manifest URL to viewer URL
            if "cudl.lib.cam.ac.uk/iiif/" in url:
                url = url.replace("/iiif/", "/view/")
            QDesktopServices.openUrl(QUrl(url))

    def _htmlify(self, text):
        if not text: return ""
        t = text.replace("\n", "<br>")
        t = re.sub(r'\*(.*?)\*', r"<b style='color:red;'>\1</b>", t)
        return f"<div dir='rtl'>{t}</div>"

    def _apply_manual_highlights_to_text(self, text, uid):
        if not text or not uid:
            return text
        spans = []
        for ph in self.data.get('page_highlights', []) if self.data else []:
            if ph.get('uid') == uid:
                span = ph.get('span')
                if span and len(span) == 2:
                    spans.append(span)
        if not spans:
            return text
        # Apply in reverse order to keep indices stable
        spans.sort(key=lambda s: s[0], reverse=True)
        for s, e in spans:
            if s is None or e is None:
                continue
            if s < 0 or e > len(text) or s >= e:
                continue
            text = text[:e] + "*" + text[e:]
            text = text[:s] + "*" + text[s:]
        return text

    def load_result_by_index(self, idx):
        data = self.all_results[idx]
        if not data.get('full_text'):
            uid = data.get('uid')
            if uid:
                data['full_text'] = self.searcher.get_full_text_by_id(uid) or data.get('text', '')
            else:
                # Tag search results: get full text by sys_id from display dict
                sid = data.get('display', {}).get('id', '')
                if sid and self.searcher:
                    pages = self.searcher.get_full_manuscript(sid)
                    data['full_text'] = '\n'.join(p['text'] for p in pages if p.get('text')) if pages else data.get('text', '')
                else:
                    data['full_text'] = data.get('text', '')
        self.data = data
        
        # Nav UI Updates
        self.lbl_res_count.setText(tr("Result {} of {}").format(idx + 1, len(self.all_results)))
        self.btn_res_prev.setEnabled(idx > 0)
        self.btn_res_next.setEnabled(idx < len(self.all_results) - 1)
        
        # Parse Meta
        ids = self.meta_mgr.parse_full_id_components(data.get('raw_header', ''))
        self.current_sys_id = ids['sys_id']
        if not self.current_sys_id:
            # Fallback for tag search results: get sys_id from display dict
            self.current_sys_id = data.get('display', {}).get('id', '')
        try: p = int(ids['p_num'])
        except: p = 1

        # Add to Recently Viewed
        parent = self.parent()
        if parent and hasattr(parent, 'lists_mgr') and parent.lists_mgr and self.current_sys_id:
            fl_id = parent._normalize_fl_id(ids.get('fl_id'))
            parent.lists_mgr.add_to_recent(self.current_sys_id, fl_id=fl_id, img=ids.get('p_num'))

        # --- Prepare Text Content ---
        # 1. Manuscript Text (Apply Pattern!)
        ms_raw = data.get('full_text', '') or data.get('text', '')
        pattern_str = data.get('highlight_pattern') # Get regex pattern
        
        if pattern_str:
            try:
                # Apply Regex to clean full-text to verify highlighting on load
                regex = re.compile(pattern_str, re.IGNORECASE)
                ms_raw = regex.sub(r'*\g<0>*', ms_raw)
            except:
                pass

        self.text_ms.setHtml(self._htmlify(ms_raw))
        self._refresh_find_highlights()
        
        # 2. Source Context
        source_text = ""
        if 'source_ctx' in data:
            parent = self.parent()
            if parent and hasattr(parent, "comp_text_area"):
                source_text = parent.comp_text_area.toPlainText().strip()
            if not source_text:
                source_text = data.get('source_ctx', '')
        source_text = self._apply_source_highlights(source_text, pattern_str)
        if source_text:
            self.src_widget.setVisible(True)
            self.text_src.setHtml(self._htmlify(source_text))
        else:
            self.src_widget.setVisible(False)
            self.text_src.clear()
        
        # Load Page & Metadata
        self.load_page(target=p)

        # Preload next result
        self._preload_next_result(idx + 1)

    def _preload_next_result(self, next_idx):
        if next_idx >= len(self.all_results): return
        res = self.all_results[next_idx]

        # Extract SID logic from load_result
        meta = res.get('display', {})
        parsed = self.meta_mgr.parse_full_id_components(res.get('raw_header', ''))
        sid = parsed['sys_id'] or meta.get('id')

        if not sid: return

        # Trigger Enrich Fetch (caches metadata)
        # We don't connect signals, just run it
        self.preload_meta_worker = EnrichMetadataThread(self.meta_mgr, sid)
        self.preload_meta_worker.start()

    def load_by_shelfmark(self, shelfmark: str, page_num: int = 1):
        """Load a document by shelfmark within the same dialog."""
        try:
            parent = self.parent()
            if not parent:
                return False

            # Look up sys_id from shelfmark
            if hasattr(parent, '_ensure_shelf_map'):
                parent._ensure_shelf_map()
            if hasattr(parent, '_normalize_shelfmark') and hasattr(parent, '_shelf_to_sys'):
                norm = parent._normalize_shelfmark(shelfmark)
                sys_id = parent._shelf_to_sys.get(norm) if norm else None
            else:
                return False

            if not sys_id:
                QMessageBox.warning(self, tr("Error"), tr("Document not found: {}").format(shelfmark))
                return False

            # Get page data
            page_data = self.searcher.get_browse_page(sys_id, p_num=page_num)
            if not page_data:
                QMessageBox.warning(self, tr("View Error"), tr("Could not load manuscript data."))
                return False

            try:
                shelfmark_display, title = self.meta_mgr.get_meta_for_id(sys_id)
            except:
                shelfmark_display = shelfmark
                title = ""

            # Create result dict
            result = {
                'uid': page_data.get('uid', ''),
                'raw_header': page_data.get('full_header', ''),
                'full_header': page_data.get('full_header', ''),
                'text': page_data.get('text', ''),
                'full_text': page_data.get('text', ''),
                'display': {
                    'id': sys_id,
                    'shelfmark': shelfmark_display,
                    'title': title,
                    'img': str(page_num),
                    'source': ''
                }
            }

            # Add to results and navigate
            self.all_results.append(result)
            new_idx = len(self.all_results) - 1
            self.current_result_idx = new_idx
            self.load_result_by_index(new_idx)
            return True
        except Exception as e:
            print(f"[ERROR] load_by_shelfmark failed: {e}", flush=True)
            import traceback
            traceback.print_exc()
            return False

    def load_page(self, offset=0, target=None):
        if not self.current_sys_id: return
        self.cancel_image_thread()
        
        # Determine strict navigation source
        # If target (Spinbox jump) is set -> Use p_num logic (target)
        # If offset (Next/Prev) is set -> Use internal_index logic (prevents loops)
        
        page_data = None
        
        if target is not None:
            # Jump by number (user typed in box)
            try: p = int(target)
            except: p = 1
            page_data = self.searcher.get_browse_page(self.current_sys_id, p_num=p, next_prev=0, allow_cross=True)
        else:
            # Relative Navigation (Next/Prev)
            # Use internal index if we have it, otherwise rely on p_num
            idx_arg = self.current_internal_idx 
            p_arg = int(self.current_p_num) if self.current_p_num is not None else None
            
            page_data = self.searcher.get_browse_page(
                self.current_sys_id, 
                p_num=p_arg, 
                next_prev=offset,
                absolute_index=idx_arg, # <--- THIS FIXES THE BUG
                allow_cross=True
            )
            
        if not page_data: return

        # --- UPDATE STATE ---
        new_sys = page_data.get('sys_id', self.current_sys_id)
        if new_sys and new_sys != self.current_sys_id:
            self.current_sys_id = new_sys

        self.current_p_num = page_data['p_num']
        self.current_internal_idx = page_data['internal_index'] # <--- SAVE IT
        
        parsed_new = self.meta_mgr.parse_full_id_components(page_data['full_header'])
        self.current_fl_id = parsed_new['fl_id']
        self.current_full_header = page_data.get('full_header', '')
        self.current_page_text = page_data.get('text', '')
        self.current_page_uid = page_data.get('uid')
        self._update_add_to_list_button()

        # Keep the dialog's data object aligned with the currently displayed folio
        if self.data is not None:
            self.data['raw_header'] = page_data.get('full_header', self.data.get('raw_header', ''))
            self.data['uid'] = page_data.get('uid', self.data.get('uid'))
            self.data['full_text'] = page_data.get('text', self.data.get('full_text', ''))
            display_block = self.data.get('display', {})
            display_block['id'] = self.current_sys_id
            self.data['display'] = display_block

        # Update Info Label
        info_html = f"<b>{tr('Sys')}:</b> {self.current_sys_id} | <b>{tr('FL')}:</b> {self.current_fl_id or '?'}"
        self.lbl_info.setText(info_html)

        # Update Page Controls
        self.spin_page.blockSignals(True); self.spin_page.setValue(self.current_p_num); self.spin_page.blockSignals(False)
        self.lbl_total.setText(f"/ {page_data['total_pages']}")

        # Sync compact bar page label
        if hasattr(self, 'lbl_compact_page') and self.compact_bar.isVisible():
            self.lbl_compact_page.setText(f"{self.current_p_num} {self.lbl_total.text()}")

        # 2. Sync Image (Non-Blocking)
        if self.btn_external_view.isChecked():
            QTimer.singleShot(0, self.sync_external_view)

        # --- Render Text ---
        raw_text = page_data['text']
        raw_text = self._apply_manual_highlights_to_text(raw_text, self.current_page_uid)
        pattern_str = self.data.get('highlight_pattern')
        
        if pattern_str:
            try:
                regex = re.compile(pattern_str, re.IGNORECASE)
                highlighted_text = regex.sub(r'*\g<0>*', raw_text)
                raw_text = highlighted_text
            except: pass
        
        self.text_ms.setHtml(self._htmlify(raw_text))
        self._refresh_find_highlights()

        # Load versions for this page
        self._rd_load_versions()

        # Start PGP source fetch for this page (runs in background)
        # Reset PGP and enriched data flags for new result
        self._rd_pgp_doc = None
        self._rd_pgp_sources = []
        self._rd_enriched_data_loaded = False
        self._rd_fjms_bib = []
        self._rd_marc_bib = []
        self._rd_catalog_detail = None
        self.btn_rd_bib_fjms.setVisible(False)
        self.btn_rd_bib_nli.setVisible(False)
        self.btn_rd_catalog.setVisible(False)
        self.btn_rd_catalog.setEnabled(False)
        if hasattr(self, 'btn_compact_bib_fjms'):
            self.btn_compact_bib_fjms.setVisible(False)
            self.btn_compact_bib_nli.setVisible(False)
        if hasattr(self, 'btn_compact_catalog'):
            self.btn_compact_catalog.setVisible(False)
        parent = self.parent()
        if parent:
            # Disconnect old worker signals first to prevent stale results
            if hasattr(self, '_rd_pgp_worker') and self._rd_pgp_worker is not None:
                try:
                    self._rd_pgp_worker.finished_signal.disconnect(self._on_rd_pgp_loaded)
                    self._rd_pgp_worker.error_signal.disconnect(self._on_rd_pgp_error)
                except (TypeError, RuntimeError):
                    pass
            self._rd_pgp_worker = PGPSourceWorker(self.current_sys_id, self.current_p_num or 1)
            self._rd_pgp_worker.finished_signal.connect(self._on_rd_pgp_loaded)
            self._rd_pgp_worker.error_signal.connect(self._on_rd_pgp_error)
            self._rd_pgp_worker.start()

        # Update joins menu
        self._rd_update_joins_menu()

        # Update Domain info + start enrichment (AFTER reset so buttons aren't wiped)
        self._update_rd_domain_label()

    def _update_rd_domain_label(self):
        """Update domain info label for the current result in ResultDialog."""
        parent = self.parent()
        if not parent or not hasattr(parent, '_result_domain_map'):
            self.lbl_rd_domains.setVisible(False)
            return

        domain_names = parent._result_domain_map.get(self.current_sys_id, [])
        if domain_names:
            display_names = [parent._domain_display_name(d) for d in domain_names] if hasattr(parent, '_domain_display_name') else domain_names
            self.lbl_rd_domains.setText(" | " + tr("Domain") + ": " + ", ".join(display_names))
            self.lbl_rd_domains.setVisible(True)
        else:
            self.lbl_rd_domains.setVisible(False)

        self.lbl_meta_loading.setVisible(False)
        self.lbl_title.setText('')
        self.lbl_img_label.setText("")
        
        if self.ext_data and self.current_sys_id not in self.meta_mgr.nli_cache:
             self.ext_data = None
             self.ext_canvases = []
             self.btn_external_view.setVisible(False)
             self.external_pane.setVisible(False)

        cached_meta = self.meta_mgr.nli_cache.get(self.current_sys_id)
        if cached_meta:
            self.apply_metadata(cached_meta)
        else:
            self.lbl_meta_loading.setVisible(True)
            self.current_meta_request += 1
            request_id = self.current_meta_request
            def worker():
                meta = self.meta_mgr.fetch_nli_data(self.current_sys_id)
                self.metadata_loaded.emit(request_id, meta or {})
            threading.Thread(target=worker, daemon=True).start()

        if not cached_meta or 'marc' not in cached_meta:
            # Disconnect old enrich worker to prevent stale signals and GC crash
            if hasattr(self, 'enrich_worker') and self.enrich_worker is not None:
                try:
                    self.enrich_worker.finished_signal.disconnect(self.on_enriched_data_loaded)
                except (TypeError, RuntimeError):
                    pass
            self.enrich_worker = EnrichMetadataThread(self.meta_mgr, self.current_sys_id)
            self.enrich_worker.finished_signal.connect(self.on_enriched_data_loaded)
            self.enrich_worker.start()
        else:
            self.on_enriched_data_loaded(self.current_sys_id, cached_meta)

    def apply_metadata(self, meta):
        # 1. Update Text Labels
        shelf = self.meta_mgr.get_shelfmark_from_header(self.current_full_header) or meta.get('shelfmark', 'Unknown Shelf')
        # Add library name prefix
        library_code = self.meta_mgr.get_library_for_id(self.current_sys_id)
        if library_code:
            library = get_library_display(library_code, short=False)
            shelf = f"{library} | {shelf}"
        self.lbl_shelf.setText(shelf)
        if hasattr(self, 'lbl_compact_shelf'):
            self.lbl_compact_shelf.setText(shelf)
        _set_label_with_tooltip(self.lbl_title, meta.get('title', ''))
        self.lbl_meta_loading.setVisible(False)

        # 2. Trigger Image Fetch using the FRESH metadata
        # (This meta object now contains 'thumb_url' from the XML 907 $d field)
        self.fetch_image(self.current_sys_id, meta)

    def toggle_extended_info(self, checked):
        self.extended_info_visible = checked
        self.txt_extended_info.setVisible(checked)
        label = tr("Hide Extended Info") if checked else tr("Show Extended Info")
        self.btn_ext_info.setText(label)
        if hasattr(self, 'btn_compact_ext_info'):
            self.btn_compact_ext_info.blockSignals(True)
            self.btn_compact_ext_info.setChecked(checked)
            self.btn_compact_ext_info.setText(label)
            self.btn_compact_ext_info.blockSignals(False)

    def _on_rd_ext_link_clicked(self, url):
        """Handle clicks on links in ResultDialog extended info."""
        url_str = url.toString()
        if url_str.startswith('tag:'):
            tag = url_str[4:]
            parent = self.parent()
            if parent and hasattr(parent, '_search_by_pgp_tag'):
                self.close()
                parent._search_by_pgp_tag(tag)
        elif url_str.startswith('http'):
            QDesktopServices.openUrl(url)

    def _show_rd_fjms_bib(self):
        """Open FJMS bibliography dialog from ResultDialog."""
        if not self._rd_fjms_bib:
            return
        shelf = self.meta_mgr.get_meta_for_id(self.current_sys_id)[0] if self.current_sys_id else ''
        dlg = FjmsBibliographyDialog(
            self._rd_fjms_bib,
            sys_id=self.current_sys_id or '',
            shelfmark=shelf,
            parent=self,
        )
        dlg.exec()

    def _show_rd_nli_bib(self):
        """Open NLI bibliography dialog from ResultDialog."""
        if not self._rd_marc_bib:
            return
        shelf = self.meta_mgr.get_meta_for_id(self.current_sys_id)[0] if self.current_sys_id else ''
        dlg = NliBibliographyDialog(
            self._rd_marc_bib,
            sys_id=self.current_sys_id or '',
            shelfmark=shelf,
            parent=self,
        )
        dlg.exec()

    def _show_rd_catalog(self):
        """Open FJMS catalog records dialog from reading desk (lazy fetch)."""
        # Lazy fetch: load catalog detail on first click if not yet loaded
        if self._rd_catalog_detail is None and self.current_sys_id:
            try:
                from shared.fjms_service import get_fjms_service
                fjms_svc = get_fjms_service()
                if fjms_svc.is_available():
                    self._rd_catalog_detail = fjms_svc.get_catalog_detail(self.current_sys_id)
            except Exception:
                pass

        if not self._rd_catalog_detail:
            return
        shelf = self.meta_mgr.get_meta_for_id(self.current_sys_id)[0] if self.current_sys_id else ''
        dlg = FjmsCatalogDialog(
            self._rd_catalog_detail,
            sys_id=self.current_sys_id or '',
            shelfmark=shelf,
            parent=self,
        )
        dlg.exec()

    def toggle_external_viewer(self, checked):
        self.external_pane.setVisible(checked)
        if checked:
            QTimer.singleShot(0, self.sync_external_view)

    def on_enriched_data_loaded(self, sid, meta):
        if not meta: return
        # Verify this data is for the currently displayed result to prevent race conditions
        if sid != self.current_sys_id:
            return
        if self.current_sys_id not in self.meta_mgr.nli_cache: return

        # 1. Update Image Label
        fl_digits = re.sub(r"\D", "", str(self.current_fl_id or ""))
        canvas_map = meta.get('canvas_map', {})
        label = canvas_map.get(fl_digits)
        self.lbl_img_label.setText(f"({label})" if label else "")

        # Check for Oxford Part metadata
        oxford_part_id = meta.get('oxford_part_id')
        part_meta = None
        if oxford_part_id:
            part_meta = self.meta_mgr.get_part_metadata(oxford_part_id)
        elif self.current_sys_id:
            # Check if this folio belongs to a Part
            part_id = self.meta_mgr.get_part_for_folio(self.current_sys_id)
            if part_id:
                oxford_part_id = part_id
                part_meta = self.meta_mgr.get_part_metadata(part_id)

        # 2. Populate External / Image Viewer
        has_images = bool(meta.get('images_nli') or meta.get('images_ext'))

        self.btn_toggle_image.setVisible(has_images)

        if has_images:
            # Show viewer by default
            self.external_pane.setVisible(True)
            self.btn_toggle_image.setChecked(True)

            self.lbl_ext_attr.setVisible(False)
            self.txt_ext_meta.setHtml("")
            self.txt_ext_meta.setVisible(False)

            # Load images into widget
            shelfmark = meta.get('shelfmark') or self.meta_mgr.get_meta_for_id(self.current_sys_id)[0]
            folio_num = _get_folio_number_from_shelfmark(shelfmark)
            side_offset = 1 if (self.current_internal_idx or 0) % 2 == 1 else 0
            initial_idx = _get_folio_image_index(
                meta,
                folio_num if folio_num is not None else self.current_p_num,
                side_offset=side_offset
            )
            self.ms_viewer.load_images(meta, initial_idx, target_folio=folio_num)
        else:
            self.external_pane.setVisible(False)
            self.btn_toggle_image.setChecked(False)

        self.external_url = meta.get('external_url') or meta.get('marc', {}).get('external_iiif_link')
        if self.external_url:
            provider = meta.get('external_provider', '')
            if oxford_part_id or provider == 'oxford':
                btn_label = tr("Oxford")
            elif provider == 'cambridge' or "cudl.lib.cam.ac.uk" in (self.external_url or "").lower():
                btn_label = tr("Cambridge")
            elif provider == 'manchester':
                btn_label = "Manchester LUNA"
            elif provider == 'jts':
                btn_label = "Princeton Digital Library"
            else:
                btn_label = tr("External Website")
            self.btn_external_link.setText(btn_label)
            self.btn_external_link.setVisible(True)
        else:
            self.btn_external_link.setVisible(False)

        # 3. Populate bibliography buttons (before early-return guard)
        marc = meta.get('marc', {})
        fjms_bib = meta.get('bibliography', [])
        marc_bib = marc.get('bibliography', [])
        if fjms_bib:
            self._rd_fjms_bib = fjms_bib
            lbl = f"{tr('Bibliography FJMS')} ({len(fjms_bib)})"
            self.btn_rd_bib_fjms.setText(lbl)
            self.btn_rd_bib_fjms.setVisible(True)
            if hasattr(self, 'btn_compact_bib_fjms'):
                self.btn_compact_bib_fjms.setText(lbl)
                self.btn_compact_bib_fjms.setVisible(True)
        if marc_bib:
            self._rd_marc_bib = marc_bib
            lbl = f"{tr('Bibliography Ktiv')} ({len(marc_bib)})"
            self.btn_rd_bib_nli.setText(lbl)
            self.btn_rd_bib_nli.setVisible(True)
            if hasattr(self, 'btn_compact_bib_nli'):
                self.btn_compact_bib_nli.setText(lbl)
                self.btn_compact_bib_nli.setVisible(True)

        # Catalog Records button
        # Detail is fetched lazily on button click, not during page load
        self._rd_catalog_detail = None
        try:
            from shared.fjms_service import get_fjms_service
            fjms_svc = get_fjms_service()
            if fjms_svc.is_available():
                source_names = fjms_svc.get_source_names(self.current_sys_id)
                catalog_count = len(source_names)
                self.btn_rd_catalog.setText(f"{tr('Catalog Records')} ({catalog_count})")
                self.btn_rd_catalog.setEnabled(catalog_count > 0)
                self.btn_rd_catalog.setVisible(True)
                if hasattr(self, 'btn_compact_catalog'):
                    self.btn_compact_catalog.setText(f"{tr('Catalog Records')} ({catalog_count})")
                    self.btn_compact_catalog.setEnabled(catalog_count > 0)
                    self.btn_compact_catalog.setVisible(True)
        except Exception:
            self.btn_rd_catalog.setVisible(False)
            if hasattr(self, 'btn_compact_catalog'):
                self.btn_compact_catalog.setVisible(False)

        # 4. Build Extended Info HTML (Text)
        external_meta = meta.get('external_meta', {})
        has_pgp = bool(getattr(self, '_rd_pgp_doc', None))
        if not marc and not meta.get('physical_desc') and not part_meta and not external_meta and not has_pgp:
            self.btn_ext_info.setVisible(False)
            if hasattr(self, 'btn_compact_ext_info'):
                self.btn_compact_ext_info.setVisible(False)
            return

        palette = self.txt_extended_info.palette()
        text_color = palette.color(QPalette.ColorRole.Text).name()
        base_color = palette.color(QPalette.ColorRole.Base).name()
        part_bg = QColor(base_color).lighter(115).name()

        kti_html = ""
        date_val = marc.get('date')
        if date_val:
            kti_html += f"<p><b>{tr('Date')}:</b> {date_val}</p>"

        dims = marc.get('dimensions'); phys = meta.get('physical_desc')
        if dims or phys:
            kti_html += f"<p><b>{tr('Physical Description')}:</b> {phys or ''} {dims or ''}</p>"

        eng_title = marc.get('english_title')
        if eng_title:
            kti_html += f"<p><b>{tr('English Title')}:</b> {eng_title}</p>"

        subjects = marc.get('subjects', [])
        if subjects:
            kti_html += f"<p><b>{tr('Subjects')}:</b> {'; '.join(subjects)}</p>"

        notes = marc.get('notes', [])
        if notes:
            kti_html += f"<p><b>{tr('Notes')}:</b><ul>"
            for n in notes:
                kti_html += f"<li>{n}</li>"
            kti_html += "</ul></p>"

        people = marc.get('people', [])
        if people:
            kti_html += f"<p><b>{tr('People')}:</b> {'; '.join(people)}</p>"

        external_html = ""
        if part_meta:
            part_display = self.meta_mgr.codico_mgr.get_part_display_name(oxford_part_id)
            external_html += (
                f"<div style='background-color: {part_bg}; color:{text_color}; padding: 10px; "
                "margin-bottom: 10px; border-left: 3px solid #3498db; text-align: left;' dir='ltr'>"
            )
            external_html += f"<p><b>📖 {tr('Codicological Part')}:</b> {part_display}</p>"

            folio_range = part_meta.get('folio_range', [])
            if len(folio_range) == 2:
                if folio_range[0] == folio_range[1]:
                    external_html += f"<p><b>{tr('Folio')}:</b> {folio_range[0]}</p>"
                else:
                    external_html += f"<p><b>{tr('Folio Range')}:</b> {folio_range[0]} - {folio_range[1]}</p>"

            part_title = part_meta.get('title', '')
            if part_title:
                external_html += f"<p><b>{tr('Oxford Title')}:</b> {part_title}</p>"

            part_contents = part_meta.get('contents', '')
            if part_contents:
                external_html += f"<p><b>{tr('Contents')}:</b> {part_contents}</p>"

            external_html += "</div>"

        if external_meta:
            external_html += f"<div style='margin-bottom: 10px; text-align: left;' dir='ltr'><ul>"
            for k, v in external_meta.items():
                external_html += f"<li><b>{k}:</b> {v}</li>"
            external_html += "</ul></div>"

        is_rtl = self.layoutDirection() == Qt.LayoutDirection.RightToLeft
        dir_attr = "rtl" if is_rtl else "ltr"
        header_align = "right" if is_rtl else "left"
        kti_header = tr("Ktiv Info")
        if oxford_part_id:
            external_header = tr("Oxford Info")
        else:
            external_header = tr("Cambridge Info")

        html = f"<div style='font-family:Arial; color:{text_color}; background-color:{base_color};'>"
        if external_html:
            if is_rtl:
                first_title, first_html = kti_header, kti_html
                second_title, second_html = external_header, external_html
            else:
                first_title, first_html = external_header, external_html
                second_title, second_html = kti_header, kti_html

            html += (
                f"<table style='width:100%; border-collapse:collapse;' dir='{dir_attr}'>"
                f"<tr>"
                f"<th style='text-align:{header_align}; padding:4px; border-bottom:1px solid #ccc;'>{first_title}</th>"
                f"<th style='text-align:{header_align}; padding:4px; border-bottom:1px solid #ccc;'>{second_title}</th>"
                f"</tr>"
                f"<tr>"
                f"<td style='vertical-align:top; padding:6px;'>{first_html}</td>"
                f"<td style='vertical-align:top; padding:6px;'>{second_html}</td>"
                f"</tr></table>"
            )
        else:
            html += kti_html

        # Append FJMS catalog section
        parent = self.parent()
        if parent and hasattr(parent, '_build_fjms_catalog_html'):
            fjms_catalog = parent._build_fjms_catalog_html(self.current_sys_id, text_color)
            if fjms_catalog:
                html += fjms_catalog

        # Append PGP metadata section if available
        pgp_doc = getattr(self, '_rd_pgp_doc', None)
        if pgp_doc:
            if parent and hasattr(parent, '_build_pgp_extended_info_html'):
                pgp_html = parent._build_pgp_extended_info_html(pgp_doc, palette=palette)
                if pgp_html:
                    html += pgp_html

        html += "</div>"
        self.txt_extended_info.setHtml(html)
        self.btn_ext_info.setVisible(True)
        if hasattr(self, 'btn_compact_ext_info'):
            self.btn_compact_ext_info.setVisible(True)
        # Store flag so PGP late-arrival handler knows enriched data was processed
        self._rd_enriched_data_loaded = True

        _set_label_with_tooltip(self.lbl_title, meta.get('title', ''))
        shelf = meta.get('shelfmark')
        if shelf and shelf != "Unknown":
            # Try CSV library_code first, then MARC as fallback
            library_code = self.meta_mgr.get_library_for_id(self.current_sys_id)
            if library_code:
                library = get_library_display(library_code, short=False)
            else:
                library = marc.get('current_owner', '')
            if library:
                shelf = f"{library} | {shelf}"
            # Add Part info to shelfmark if available
            if oxford_part_id:
                part_label = self.meta_mgr.codico_mgr.get_part_label(oxford_part_id)
                if part_label:
                    shelf = f"{shelf} [{part_label}]"
            self.lbl_shelf.setText(shelf)
            if hasattr(self, 'lbl_compact_shelf'):
                self.lbl_compact_shelf.setText(shelf)

        thumb_url = meta.get('thumb_url')
        if thumb_url and thumb_url != getattr(self, 'current_thumb_url', None):
            self.fetch_image(self.current_sys_id, meta)

    def sync_external_view(self):
        meta = self.meta_mgr.nli_cache.get(self.current_sys_id, {})
        if not meta:
            return
        shelfmark = meta.get('shelfmark') or self.meta_mgr.get_meta_for_id(self.current_sys_id)[0]
        folio_num = _get_folio_number_from_shelfmark(shelfmark)
        side_offset = 1 if (self.current_internal_idx or 0) % 2 == 1 else 0

        # Use viewer's images (may include dynamic images added by load_images)
        viewer_images = getattr(self.ms_viewer, 'images_ext', None)
        if viewer_images:
            idx = _get_folio_image_index(
                {'images_ext': viewer_images},
                folio_num if folio_num is not None else self.current_p_num,
                side_offset=side_offset
            )
            self.ms_viewer.set_page(idx)

    def on_metadata_loaded(self, request_id, meta):
        if request_id != self.current_meta_request:
            return
        self.apply_metadata(meta or {})

    def cancel_image_thread(self):
        img_thread = getattr(self, 'img_thread', None)
        if img_thread and img_thread.isRunning():
            img_thread.cancel()
            # Use short timeout to avoid blocking UI - thread will finish in background
            img_thread.wait(500)

        if getattr(self, 'ext_img_thread', None) and self.ext_img_thread.isRunning():
            self.ext_img_thread.cancel()
            # Use short timeout to avoid blocking UI - thread will finish in background
            self.ext_img_thread.wait(500)

    def fetch_image(self, sys_id, meta=None):
        self.cancel_image_thread()
        self.lbl_thumb.setText(tr("Loading..."))
        self.lbl_thumb.setPixmap(QPixmap())

        # Ensure we look at the global cache which acts as the "Source of Truth"
        if not meta:
            meta = self.meta_mgr.nli_cache.get(sys_id)

        # Retrieve the URL that MetadataManager logic (XML 907 $d) has determined
        thumb_url = meta.get('thumb_url') if meta else None

        if thumb_url:
            self.start_download(sys_id, thumb_url)
        else:
            # If meta exists but no thumb_url, it means no representative image found
            if meta:
                self.lbl_thumb.setText(tr("No Preview"))
            else:
                self.lbl_thumb.setText(tr("Waiting..."))

        def worker(target_sid=sys_id):
            url = self.meta_mgr.get_thumbnail(target_sid)
            self.thumb_resolved.emit(target_sid, url)

        threading.Thread(target=worker, daemon=True).start()

    def _on_thumb_resolved(self, sid, thumb_url):
        if sid != self.current_sys_id:
            return
        if thumb_url:
            self.start_download(sid, thumb_url)
        else:
            self.on_img_failed()

    def start_download(self, sid, thumb_url):
        if sid != self.current_sys_id:
            return

        self.current_thumb_url = thumb_url
        self.cancel_image_thread()

        if not thumb_url:
            self.on_img_failed()
            return

        self.img_thread = ImageLoaderThread(thumb_url)
        self.img_thread.image_loaded.connect(self.on_img_loaded)
        self.img_thread.load_failed.connect(self.on_img_failed)
        self.img_thread.start()
        
    def start_browse_download(self, sid, thumb_url):
        if sid != self.current_browse_sid:
            return

        logger.debug("Starting browse image download for SID=%s, URL=%s", sid, thumb_url)

        self.browse_thumb_url = thumb_url
        self.cancel_browse_image_thread()

        if not thumb_url:
            self.on_browse_img_failed()
            return

        # Create and start thread
        self.browse_img_thread = ImageLoaderThread(thumb_url)
        self.browse_img_thread.image_loaded.connect(self.on_browse_img_loaded)
        self.browse_img_thread.load_failed.connect(self.on_browse_img_failed)
        self.browse_img_thread.start()

    def on_img_loaded(self, image):
        pix = QPixmap.fromImage(image)
        scaled = pix.scaled(self.lbl_thumb.size(), Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
        self.lbl_thumb.setPixmap(scaled)
        self.lbl_thumb.setText("")

    def on_img_failed(self):
        self.lbl_thumb.setPixmap(QPixmap())
        self.lbl_thumb.setText(tr("No Preview"))

    def closeEvent(self, event):
        try:
            if hasattr(self, 'meta_mgr'):
                self.meta_mgr.save_caches()
                logger.info("Metadata caches flushed to disk on exit.")
        except Exception as e:
            logger.error("Failed to save metadata caches on exit: %s", e)

        # 2. Stop worker threads safely
        try:
            if getattr(self, 'meta_loader', None) and self.meta_loader.isRunning():
                self.meta_loader.request_cancel()
                self.meta_loader.wait()

            if getattr(self, 'search_thread', None) and self.search_thread.isRunning():
                self.search_thread.requestInterruption()
                self.search_thread.wait(2000)
                if self.search_thread.isRunning():
                    self.search_thread.terminate()
                    self.search_thread.wait()

            if getattr(self, 'comp_thread', None) and self.comp_thread.isRunning():
                self.comp_thread.requestInterruption()
                self.comp_thread.wait(2000)
                if self.comp_thread.isRunning():
                    self.comp_thread.terminate()
                    self.comp_thread.wait()

            if getattr(self, 'group_thread', None) and self.group_thread.isRunning():
                self.group_thread.requestInterruption()
                self.group_thread.wait(2000)
                if self.group_thread.isRunning():
                    self.group_thread.terminate()
                    self.group_thread.wait()
                    
            if getattr(self, 'browse_img_thread', None) and self.browse_img_thread.isRunning():
                self.browse_img_thread.cancel()
                self.browse_img_thread.wait()

            # Stop manuscript viewer image threads
            if getattr(self, 'ms_viewer', None):
                self.ms_viewer.stop_threads()

        finally:
            super().closeEvent(event)

    def open_catalog(self):
        if self.current_sys_id: QDesktopServices.openUrl(QUrl(f"https://www.nli.org.il/he/discover/manuscripts/hebrew-manuscripts/itempage?vid=KTIV&scope=KTIV&docId=PNX_MANUSCRIPTS{self.current_sys_id}"))

    def open_viewer(self):
        if self.current_sys_id:
            # Use docid query param (not hash fragment) — hash-based URLs fail on direct navigation
            docid = f"PNX_MANUSCRIPTS{self.current_sys_id}-1"
            if self.current_fl_id:
                docid += f",FL{self.current_fl_id}"
            QDesktopServices.openUrl(QUrl(f"https://www.nli.org.il/he/discover/manuscripts/hebrew-manuscripts/viewerpage?vid=MANUSCRIPTS&docid={docid}"))

class ActionsHoverWidget(QWidget):
    def __init__(self, parent=None, alignment=Qt.AlignmentFlag.AlignCenter):
        super().__init__(parent)
        layout = QHBoxLayout(self)
        layout.setContentsMargins(2, 2, 2, 2)
        layout.setSpacing(4)
        layout.setAlignment(alignment)
        self.buttons = []
        self.always_visible_buttons = set()

    def add_btn(self, btn, always_visible=False):
        self.layout().addWidget(btn)
        self.buttons.append(btn)
        if always_visible:
            self.always_visible_buttons.add(btn)
            btn.setVisible(True)
        else:
            btn.setVisible(False)

    def set_buttons_visible(self, visible):
        for b in self.buttons:
            if b in self.always_visible_buttons:
                b.setVisible(True)
            else:
                b.setVisible(visible)


class ListsTreeWidget(QTreeWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.main_window = parent  # שומרים את ההפניה ל-GenizahGUI
        
        self.setDragEnabled(True)
        self.setAcceptDrops(True)
        self.setDropIndicatorShown(True)
        self.setDefaultDropAction(Qt.DropAction.MoveAction)
        self.setDragDropMode(QAbstractItemView.DragDropMode.InternalMove)

    def dropEvent(self, event):
        super().dropEvent(event)
        # קריאה לפונקציה בחלון הראשי לעדכון הצבעים והסדר
        if self.main_window and hasattr(self.main_window, 'lists_handle_tree_reorder'):
            self.main_window.lists_handle_tree_reorder()


def _format_add_to_list_label(in_list=False):
    star = "⭐" if in_list else "☆"
    return f"{star} {tr('Add to List')}"


def _format_list_star(in_list=False):
    return "⭐" if in_list else "☆"


def _truncate_title(text, max_chars=100):
    """Truncate long title text with ellipsis. Returns (truncated_text, tooltip_or_None)."""
    if not text:
        return text, None
    if len(text) <= max_chars:
        return text, None
    return text[:max_chars].rstrip() + "...", text


def _set_label_with_tooltip(label, text, max_chars=100):
    """Set label text with truncation and tooltip for full text."""
    truncated, full = _truncate_title(text, max_chars)
    label.setText(truncated or '')
    label.setToolTip(full or '')


class DomainFilterDialog(QDialog):
    """Hierarchical domain filter dialog with checkboxes and type-ahead search.

    Post-search dynamic filter: Shows only domains from current results,
    all checked by default. Unchecking excludes domains.
    """

    def __init__(self, parent=None, result_domains: dict = None, excluded_domains: set = None, uncategorized_count: int = 0):
        super().__init__(parent)
        self.setWindowTitle(tr("Filter by Subject Domain"))
        self.setMinimumSize(550, 650)
        self.result_domains = result_domains or {}  # domain_name -> count
        self.excluded_domains = excluded_domains or set()
        self.uncategorized_count = uncategorized_count
        self._updating_checks = False  # Guard for programmatic checkbox changes

        layout = QVBoxLayout(self)

        # Search input for type-ahead filtering
        search_label = QLabel(tr("Search domains:"))
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText(tr("Type to filter..."))
        self.search_input.textChanged.connect(self._filter_tree)
        layout.addWidget(search_label)
        layout.addWidget(self.search_input)

        # Tree widget with checkboxes
        self.tree = QTreeWidget()
        self.tree.setHeaderLabels([tr("Domain"), tr("Manuscripts")])
        self.tree.setColumnWidth(0, 380)
        self.tree.itemChanged.connect(self._handle_item_changed)
        layout.addWidget(self.tree)

        # Selection summary
        self.summary_label = QLabel(tr("Showing all domains"))
        layout.addWidget(self.summary_label)

        # Buttons
        btn_layout = QHBoxLayout()
        check_all_btn = QPushButton(tr("Select All"))
        check_all_btn.clicked.connect(self._check_all)
        btn_layout.addWidget(check_all_btn)
        uncheck_all_btn = QPushButton(tr("Select None"))
        uncheck_all_btn.clicked.connect(self._uncheck_all)
        btn_layout.addWidget(uncheck_all_btn)
        btn_layout.addStretch()

        ok_btn = QPushButton(tr("OK"))
        ok_btn.setDefault(True)
        ok_btn.clicked.connect(self.accept)
        cancel_btn = QPushButton(tr("Cancel"))
        cancel_btn.clicked.connect(self.reject)
        btn_layout.addWidget(ok_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)

        self._populate_tree()
        self._restore_exclusions()
        self._update_summary()

    def _populate_tree(self):
        """Populate tree with domains from current search results only."""
        from shared.fjms_service import get_fjms_service, qualify_domain_name, AMBIGUOUS_CHILD_DOMAINS
        fjms = get_fjms_service()
        if not fjms.is_available() or not self.result_domains:
            return

        # Get full hierarchy to maintain parent/child structure
        hierarchy = fjms.get_domain_hierarchy()
        self.tree.blockSignals(True)

        # Only show domains that appear in result_domains
        for parent_name, info in hierarchy.items():
            # Show parent if it or any of its children are in result_domains
            parent_in_results = parent_name in self.result_domains
            # Check both qualified and bare names for ambiguous domains
            children_in_results = []
            for child in info.get('children', []):
                qname = qualify_domain_name(child['domain'], parent_name)
                if qname in self.result_domains:
                    children_in_results.append((child, qname))
                elif child['domain'] in self.result_domains and child['domain'] not in AMBIGUOUS_CHILD_DOMAINS:
                    children_in_results.append((child, child['domain']))

            if not parent_in_results and not children_in_results:
                continue  # Skip this parent entirely

            # Add parent item (display Hebrew name when available, store English as UserRole key)
            parent_count = self.result_domains.get(parent_name, 0)
            # If parent count is 0 but has children in results, sum their counts
            if children_in_results and parent_count == 0:
                parent_count = sum(self.result_domains.get(domain_key, 0) for _, domain_key in children_in_results)
            parent_display = info.get('parent_domain_heb', parent_name) if CURRENT_LANG == 'he' else parent_name
            parent_item = QTreeWidgetItem([parent_display, f"{parent_count:,}"])
            parent_item.setFlags(parent_item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
            parent_item.setCheckState(0, Qt.CheckState.Checked)  # All checked by default
            parent_item.setData(0, Qt.ItemDataRole.UserRole, parent_name)
            self.tree.addTopLevelItem(parent_item)

            # Add children that are in results (each entry is (child_dict, domain_key))
            for child, domain_key in children_in_results:
                child_count = self.result_domains.get(domain_key, 0)
                child_display = child.get('domain_heb', child['domain']) if CURRENT_LANG == 'he' else domain_key
                child_item = QTreeWidgetItem([child_display, f"{child_count:,}"])
                child_item.setFlags(child_item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
                child_item.setCheckState(0, Qt.CheckState.Checked)  # All checked by default
                child_item.setData(0, Qt.ItemDataRole.UserRole, domain_key)
                parent_item.addChild(child_item)

        # Add "Uncategorized" node for results without domain data
        if self.uncategorized_count > 0:
            uncat_display = tr("Uncategorized")
            uncat_item = QTreeWidgetItem([uncat_display, f"{self.uncategorized_count:,}"])
            uncat_item.setFlags(uncat_item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
            uncat_item.setCheckState(0, Qt.CheckState.Checked)
            uncat_item.setData(0, Qt.ItemDataRole.UserRole, "Uncategorized")
            self.tree.addTopLevelItem(uncat_item)

        self.tree.blockSignals(False)
        self.tree.expandAll()

    def _filter_tree(self):
        """Filter tree items by search text."""
        search_text = self.search_input.text().lower()
        root = self.tree.invisibleRootItem()
        for i in range(root.childCount()):
            parent_item = root.child(i)
            parent_text = parent_item.text(0).lower()
            parent_match = search_text in parent_text
            any_child_match = False

            for j in range(parent_item.childCount()):
                child_item = parent_item.child(j)
                child_text = child_item.text(0).lower()
                child_match = search_text in child_text
                child_item.setHidden(not child_match and search_text)
                if child_match:
                    any_child_match = True

            parent_item.setHidden(not (parent_match or any_child_match) and search_text)

    def _handle_item_changed(self, item, column):
        """Handle checkbox state changes with parent-child propagation."""
        if self._updating_checks or column != 0:
            return

        self._updating_checks = True
        check_state = item.checkState(0)

        # If parent changed, update all visible children
        if item.parent() is None:
            for i in range(item.childCount()):
                child = item.child(i)
                if not child.isHidden():
                    child.setCheckState(0, check_state)

        self._updating_checks = False
        self._update_summary()

    def _check_all(self):
        """Check all items (no filtering)."""
        self.tree.blockSignals(True)
        root = self.tree.invisibleRootItem()
        for i in range(root.childCount()):
            parent_item = root.child(i)
            parent_item.setCheckState(0, Qt.CheckState.Checked)
            for j in range(parent_item.childCount()):
                parent_item.child(j).setCheckState(0, Qt.CheckState.Checked)
        self.tree.blockSignals(False)
        self._update_summary()

    def _uncheck_all(self):
        """Uncheck all items (exclude all domains)."""
        self.tree.blockSignals(True)
        root = self.tree.invisibleRootItem()
        for i in range(root.childCount()):
            parent_item = root.child(i)
            parent_item.setCheckState(0, Qt.CheckState.Unchecked)
            for j in range(parent_item.childCount()):
                parent_item.child(j).setCheckState(0, Qt.CheckState.Unchecked)
        self.tree.blockSignals(False)
        self._update_summary()

    def _restore_exclusions(self):
        """Restore previously excluded domains by unchecking them."""
        if not self.excluded_domains:
            return

        self.tree.blockSignals(True)
        root = self.tree.invisibleRootItem()
        for i in range(root.childCount()):
            parent_item = root.child(i)
            parent_domain = parent_item.data(0, Qt.ItemDataRole.UserRole)
            if parent_domain in self.excluded_domains:
                parent_item.setCheckState(0, Qt.CheckState.Unchecked)

            for j in range(parent_item.childCount()):
                child_item = parent_item.child(j)
                child_domain = child_item.data(0, Qt.ItemDataRole.UserRole)
                if child_domain in self.excluded_domains:
                    child_item.setCheckState(0, Qt.CheckState.Unchecked)

        self.tree.blockSignals(False)

    def get_excluded_domains(self):
        """Return set of excluded (unchecked) domain names."""
        excluded = set()
        root = self.tree.invisibleRootItem()
        for i in range(root.childCount()):
            parent_item = root.child(i)
            if parent_item.checkState(0) == Qt.CheckState.Unchecked:
                parent_domain = parent_item.data(0, Qt.ItemDataRole.UserRole)
                excluded.add(parent_domain)

            for j in range(parent_item.childCount()):
                child_item = parent_item.child(j)
                if child_item.checkState(0) == Qt.CheckState.Unchecked:
                    child_domain = child_item.data(0, Qt.ItemDataRole.UserRole)
                    excluded.add(child_domain)

        return excluded

    def _update_summary(self):
        """Update exclusion summary label."""
        excluded = self.get_excluded_domains()
        count = len(excluded)
        if count == 0:
            self.summary_label.setText(tr("Showing all domains"))
        elif count == 1:
            domain_name = next(iter(excluded))
            self.summary_label.setText(f"{tr('Excluding')}: {domain_name}")
        else:
            self.summary_label.setText(f"{tr('Excluding')} {count} {tr('domains')}")


class FjmsBibliographyDialog(QDialog):
    """FJMS bibliography dialog with structured table."""

    def __init__(self, fjms_entries, sys_id='', shelfmark='', parent=None):
        super().__init__(parent)
        from shared.fjms_service import format_page_ref, _ts_symbol
        self.entries = fjms_entries
        self.sys_id = sys_id
        self._format_page_ref = format_page_ref
        self._ts_symbol = _ts_symbol
        self.setWindowTitle(f"{tr('Bibliography FJMS')} \u2014 {shelfmark}" if shelfmark else tr('Bibliography FJMS'))
        self.setMinimumSize(900, 500)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)

        # Filter row 1: text + type
        filter_row = QHBoxLayout()
        self.text_filter = QLineEdit()
        self.text_filter.setPlaceholderText(tr('Filter by author, title...'))
        self.text_filter.textChanged.connect(self._filter_rows)
        filter_row.addWidget(self.text_filter, 1)
        self.type_combo = QComboBox()
        for label, val in [(tr('All'), 'All'), (tr('Discussion'), 'Discussion'),
                           (tr('Mentioned'), 'Mentioned'), (tr('Index'), 'Index')]:
            self.type_combo.addItem(label, val)
        self.type_combo.currentIndexChanged.connect(lambda _: self._filter_rows())
        filter_row.addWidget(QLabel(tr('Type') + ':'))
        filter_row.addWidget(self.type_combo)
        layout.addLayout(filter_row)

        # Filter row 2: checkboxes
        check_row = QHBoxLayout()
        self.chk_transcription = QCheckBox(tr('Has Transcription'))
        self.chk_transcription.toggled.connect(self._filter_rows)
        check_row.addWidget(self.chk_transcription)
        self.chk_translation = QCheckBox(tr('Has Translation'))
        self.chk_translation.toggled.connect(self._filter_rows)
        check_row.addWidget(self.chk_translation)
        check_row.addStretch()
        layout.addLayout(check_row)

        # Table: Author, Article/Title, Year, Vol., Pages, Type, T, S
        headers = [tr('Author'), tr('Article/Title'), tr('Year'), tr('Vol.'),
                    tr('Pages'), tr('Type'), tr('col_T'), tr('col_S')]
        self.table = QTableWidget(len(fjms_entries), len(headers))
        self.table.setHorizontalHeaderLabels(headers)
        self.table.horizontalHeader().model().setHeaderData(6, Qt.Orientation.Horizontal, tr('Transcription'), Qt.ItemDataRole.ToolTipRole)
        self.table.horizontalHeader().model().setHeaderData(7, Qt.Orientation.Horizontal, tr('Translation'), Qt.ItemDataRole.ToolTipRole)
        self.table.setSortingEnabled(False)
        self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.table.verticalHeader().setVisible(False)
        self.table.setAlternatingRowColors(True)
        for col_idx in (6, 7):
            self.table.setColumnWidth(col_idx, 36)

        for row, e in enumerate(fjms_entries):
            author = (e.get('article_author_eng') or e.get('article_author_heb') or '').strip()
            item0 = QTableWidgetItem(author)
            item0.setData(Qt.ItemDataRole.UserRole, row)
            self.table.setItem(row, 0, item0)
            article_name = (e.get('article_name') or '').strip()
            running_title = (e.get('running_title') or e.get('title_acronym') or '').strip()
            self.table.setItem(row, 1, QTableWidgetItem(article_name if article_name else running_title))
            year = str(e.get('title_year') or '').strip()
            self.table.setItem(row, 2, QTableWidgetItem(year if year and year != 'None' else ''))
            vol = str(e.get('volume') or '').strip()
            self.table.setItem(row, 3, QTableWidgetItem(vol if vol and vol != 'None' else ''))
            self.table.setItem(row, 4, QTableWidgetItem(format_page_ref(e)))
            mt = (e.get('mention_type') or '').strip()
            self.table.setItem(row, 5, QTableWidgetItem(tr(mt) if mt and mt != 'None' else ''))
            self.table.setItem(row, 6, QTableWidgetItem(_ts_symbol(e.get('transcription_type'))))
            self.table.setItem(row, 7, QTableWidgetItem(_ts_symbol(e.get('translation_type'))))

        self.table.resizeColumnsToContents()
        for col_idx in (6, 7):
            self.table.setColumnWidth(col_idx, 36)
        self.table.setSortingEnabled(True)
        self.table.currentCellChanged.connect(self._on_row_selected)
        layout.addWidget(self.table, 1)

        # Detail panel
        self.detail_panel = QTextBrowser()
        self.detail_panel.setMaximumHeight(80)
        self.detail_panel.setVisible(False)
        self.detail_panel.setStyleSheet("border: 1px solid #ccc; padding: 4px; font-size: 12px;")
        layout.addWidget(self.detail_panel)

        # Bottom row
        bottom_row = QHBoxLayout()
        if sys_id:
            ktiv_url = f"https://www.nli.org.il/he/discover/manuscripts/hebrew-manuscripts/itempage?vid=KTIV&scope=KTIV&docId=PNX_MANUSCRIPTS{sys_id}"
            btn_ktiv = QPushButton(tr('Open in KTIV'))
            btn_ktiv.clicked.connect(lambda: QDesktopServices.openUrl(QUrl(ktiv_url)))
            bottom_row.addWidget(btn_ktiv)
        bottom_row.addStretch()
        btn_close = QPushButton(tr('Close'))
        btn_close.clicked.connect(self.close)
        bottom_row.addWidget(btn_close)
        layout.addLayout(bottom_row)

    def _filter_rows(self):
        text_val = self.text_filter.text().strip().lower()
        type_val = self.type_combo.currentData() or 'All'
        need_trans = self.chk_transcription.isChecked()
        need_transl = self.chk_translation.isChecked()
        skip_vals = ('', 'None', 'Unknown')

        for row in range(self.table.rowCount()):
            item = self.table.item(row, 0)
            orig_idx = item.data(Qt.ItemDataRole.UserRole) if item else -1
            if not isinstance(orig_idx, int) or orig_idx < 0 or orig_idx >= len(self.entries):
                continue
            e = self.entries[orig_idx]
            show = True
            mt = (e.get('mention_type') or '').strip()
            if type_val != 'All' and mt != type_val:
                show = False
            if show and need_trans:
                tt = (e.get('transcription_type') or '').strip()
                if not tt or tt in skip_vals:
                    show = False
            if show and need_transl:
                tl = (e.get('translation_type') or '').strip()
                if not tl or tl in skip_vals:
                    show = False
            if show and text_val:
                searchable = ' '.join([
                    e.get('article_author_eng') or '', e.get('article_author_heb') or '',
                    e.get('article_name') or '', e.get('running_title') or '',
                    e.get('title_acronym') or '',
                ]).lower()
                if text_val not in searchable:
                    show = False
            self.table.setRowHidden(row, not show)

    def _on_row_selected(self, row, col, prev_row, prev_col):
        item = self.table.item(row, 0)
        orig_idx = item.data(Qt.ItemDataRole.UserRole) if item else -1
        if isinstance(orig_idx, int) and 0 <= orig_idx < len(self.entries):
            e = self.entries[orig_idx]
            parts = []
            article = (e.get('article_name') or '').strip()
            if article:
                parts.append(f"{tr('Article')}: {article}")
            author_heb = (e.get('article_author_heb') or '').strip()
            if author_heb:
                parts.append(f"{tr('Author')}: {author_heb}")
            tt = (e.get('transcription_type') or '').strip()
            if tt and tt not in ('', 'None'):
                parts.append(f"{tr('Transcription')}: {tr(tt)}")
            tl = (e.get('translation_type') or '').strip()
            if tl and tl not in ('', 'None'):
                parts.append(f"{tr('Translation')}: {tr(tl)}")
            cat = (e.get('catalog_acronym') or '').strip()
            if cat and cat != 'None':
                parts.append(f"{tr('Catalog')}: {cat}")
            if parts:
                self.detail_panel.setPlainText('\n'.join(parts))
                self.detail_panel.setVisible(True)
            else:
                self.detail_panel.setVisible(False)
        else:
            self.detail_panel.setVisible(False)


class FjmsCatalogDialog(QDialog):
    """Dialog showing FJMS catalog records with multi-team scholarly descriptions.

    Mirrors the FIST web interface "Cataloging Data Details" view:
    teams as columns, fields as rows, grouped into 5 labeled sections:
    1. Shelfmark Description  2. Content Description  3. Script Description
    4. Format Description  5. Miscellaneous
    """

    def __init__(self, detail: dict, sys_id: str = '', shelfmark: str = '', parent=None):
        super().__init__(parent)
        self.setWindowTitle(f'{tr("Catalog Records")} \u2014 {shelfmark}' if shelfmark else tr('Catalog Records'))
        self.setMinimumSize(800, 500)
        self.resize(900, 650)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)

        # Header — use palette text color so it works in dark mode
        palette = QApplication.palette()
        is_dark = palette.color(QPalette.ColorRole.Window).lightness() < 128
        header_color = '#bb86fc' if is_dark else '#6c3483'
        header = QLabel(f'<h3 style="color: {header_color};">{tr("Catalog Records")} \u2014 {shelfmark}</h3>')
        layout.addWidget(header)

        # Content browser — rely on app-level RTL layout direction for Hebrew.
        # Qt's QTextBrowser inherits RTL from the application, so we use plain
        # LTR HTML (no text-align or column reversal) and let Qt handle alignment.
        self.text_browser = QTextBrowser()
        self.text_browser.setOpenExternalLinks(True)
        self.text_browser.setHtml(self._build_html(detail, shelfmark=shelfmark or ''))
        layout.addWidget(self.text_browser)

        # Close button
        btn_row = QHBoxLayout()
        btn_row.addStretch()
        close_btn = QPushButton(tr("Close"))
        close_btn.clicked.connect(self.accept)
        btn_row.addWidget(close_btn)
        layout.addLayout(btn_row)

    def _build_html(self, detail: dict, shelfmark: str = '') -> str:
        """Build HTML table mirroring FIST Cataloging Data Details view."""
        from shared.fjms_service import parse_textual_frame, split_textual_frames, get_team_display_name, get_team_header_name, is_team_source, GENERIC_SOURCE_NAMES

        records = detail.get("records", [])
        running_titles = detail.get("running_titles", {})
        sizes = detail.get("sizes", {})
        fields = detail.get("fields", {})
        free_descriptions = detail.get("free_descriptions", [])
        full_texts = detail.get("full_texts", [])
        textual_frames = detail.get("textual_frames", {})
        mentions = detail.get("mentions", {})

        is_heb = CURRENT_LANG == 'he'

        # Dark mode detection — define color palette for HTML
        palette = QApplication.palette()
        text_color = palette.color(QPalette.ColorRole.Text).name()
        base_color = palette.color(QPalette.ColorRole.Base).name()
        is_dark = palette.color(QPalette.ColorRole.Window).lightness() < 128
        c = {
            'text': text_color,
            'base': base_color,
            'muted': '#777' if is_dark else '#999',
            'border': '#444' if is_dark else '#eee',
            'section_bg': '#2d1f3d' if is_dark else '#f3e5f5',
            'section_text': '#bb86fc' if is_dark else '#6c3483',
            'label': '#aaa' if is_dark else '#555',
            'header_border': '#9b59b6',
            'full_text_bg': '#2a2a2a' if is_dark else '#fafafa',
            'author_muted': '#888' if is_dark else 'gray',
        }
        # Store for use in helper methods
        self._colors = c

        # Group records by source_name to get team columns, skipping generic sources
        teams = []
        team_map = {}
        for rec in records:
            sn = rec.get("source_name") or tr("Unknown")
            if sn in GENERIC_SOURCE_NAMES:
                continue
            if sn not in team_map:
                team_map[sn] = len(teams)
                teams.append({
                    "source_name": sn,
                    "source_name_heb": rec.get("source_name_heb") or sn,
                    "records": [],
                })
            teams[team_map[sn]]["records"].append(rec)

        num_teams = len(teams)
        total_cols = num_teams + 1  # label column + team columns

        if num_teams == 0 and not free_descriptions and not full_texts:
            return f'<p style="color: {c["muted"]};">No catalog data available</p>'

        # Calculate column widths for table-layout:fixed
        label_width = 130
        team_col_width = max(150, (700 - label_width) // max(num_teams, 1)) if num_teams > 0 else 150

        html_parts = []
        html_parts.append(
            f'<table style="width:100%; border-collapse:collapse; table-layout:fixed; '
            f'font-family:Arial; font-size:13px; color:{c["text"]};">'
        )
        # Column width definitions — RTL: reverse column order (team cols first,
        # label last) so Hebrew readers see labels on the right.  Qt handles
        # text alignment within cells automatically via app-level RTL.
        if num_teams > 0:
            html_parts.append('<colgroup>')
            if is_heb:
                for _ in teams:
                    html_parts.append(f'<col style="width:{team_col_width}px;"/>')
                html_parts.append(f'<col style="width:{label_width}px;"/>')
            else:
                html_parts.append(f'<col style="width:{label_width}px;"/>')
                for _ in teams:
                    html_parts.append(f'<col style="width:{team_col_width}px;"/>')
            html_parts.append('</colgroup>')

        if num_teams > 0:
            # === Team header row ===
            team_ths = []
            for team in teams:
                header_name = get_team_header_name(team["source_name"], is_heb=is_heb)
                team_ths.append(
                    f'<th style="padding:8px; border-bottom:2px solid {c["header_border"]}; '
                    f'overflow:hidden; text-overflow:ellipsis; white-space:nowrap;" title="{header_name}">'
                    f'{header_name}</th>'
                )
            empty_th = f'<th style="padding:8px;"></th>'
            html_parts.append('<tr>')
            if is_heb:
                html_parts.extend(team_ths)
                html_parts.append(empty_th)
            else:
                html_parts.append(empty_th)
                html_parts.extend(team_ths)
            html_parts.append('</tr>')

            # === Section 1: Shelfmark Description ===
            html_parts.append(self._section_row(tr('Shelfmark Description'), total_cols))

            # Shelfmark
            if shelfmark:
                sm_vals = [shelfmark] * num_teams
                html_parts.append(self._field_row(tr('Shelfmark'), sm_vals, is_heb))

            # Source — "{Author}, Head of {Team}" for teams, raw name for catalogs
            source_vals = []
            for team in teams:
                sn = team["source_name"]
                first_rec = team["records"][0] if team["records"] else None
                author = ""
                if first_rec:
                    a = first_rec.get("author_text")
                    if a and str(a).strip():
                        author = str(a).strip()
                if is_team_source(sn) and author:
                    header = get_team_header_name(sn, is_heb=is_heb)
                    if is_heb:
                        source_vals.append(f"{author}, ראש {header}")
                    else:
                        source_vals.append(f"{author}, Head of {header}")
                elif is_team_source(sn):
                    source_vals.append(get_team_display_name(sn, is_heb=is_heb))
                else:
                    sn_display = team.get("source_name_heb", sn) if is_heb else sn
                    source_vals.append(sn_display or sn)
            html_parts.append(self._field_row(tr('Source'), source_vals, is_heb))

            # Number of Folios
            folio_vals = []
            for team in teams:
                folios = [self._fmt_int(r.get("num_folio")) for r in team["records"]
                          if r.get("num_folio") and str(r["num_folio"]).strip() and str(r["num_folio"]).strip() != '0']
                folio_vals.append(', '.join(folios) if folios else '')
            html_parts.append(self._field_row(tr('Number of Folios'), folio_vals, is_heb))

            # Number of Bifolios
            bifolio_vals = []
            for team in teams:
                bifolios = [self._fmt_int(r.get("num_bifolio")) for r in team["records"]
                            if r.get("num_bifolio") and str(r["num_bifolio"]).strip() and str(r["num_bifolio"]).strip() != '0']
                bifolio_vals.append(', '.join(bifolios) if bifolios else '')
            html_parts.append(self._field_row(tr('Number of Bifolios'), bifolio_vals, is_heb))

            # === Section 2: Content Description ===
            html_parts.append(self._section_row(tr('Content Description'), total_cols))

            # Domain
            domain_vals = []
            for team in teams:
                categories = []
                for rec in team["records"]:
                    tf_eng = rec.get("textual_frame_eng") or ""
                    tf_heb = rec.get("textual_frame_heb") or ""
                    tf = tf_heb if is_heb and tf_heb else tf_eng
                    if tf:
                        parts = split_textual_frames(tf)
                        if not parts and tf.strip():
                            parts = [tf.strip()]
                        for part in parts:
                            cat, content = parse_textual_frame(part)
                            display_parts = []
                            if cat:
                                display_parts.append(f"[{cat}]")
                            if content:
                                display_parts.append(content)
                            if display_parts:
                                categories.append(' '.join(display_parts))
                domain_vals.append('; '.join(categories) if categories else '')
            html_parts.append(self._field_row(tr('Domain'), domain_vals, is_heb))

            # Running Title
            rt_vals = []
            for team in teams:
                titles = []
                for rec in team["records"]:
                    rec_id = rec.get("unit_catalog_rec_id")
                    if rec_id and rec_id in running_titles:
                        for rt in running_titles[rec_id]:
                            rt_text = rt.get("running_title", "")
                            if rt_text and str(rt_text).strip():
                                titles.append(str(rt_text).strip())
                rt_vals.append('; '.join(titles) if titles else '')
            html_parts.append(self._field_row(tr('Running Title'), rt_vals, is_heb))

            # Detailed Content (from catalog_textual_frames)
            if textual_frames:
                dc_vals = []
                for team in teams:
                    frames = []
                    for rec in team["records"]:
                        rec_id = rec.get("unit_catalog_rec_id")
                        if rec_id and rec_id in textual_frames:
                            for tf in textual_frames[rec_id]:
                                text = tf.get("heb") if is_heb and tf.get("heb") else tf.get("eng")
                                if text and str(text).strip():
                                    frames.append(str(text).strip())
                    dc_vals.append('; '.join(frames) if frames else '')
                html_parts.append(self._field_row(tr('Detailed Content'), dc_vals, is_heb))

            # GenizahTitle
            gt_vals = []
            for team in teams:
                titles = []
                for rec in team["records"]:
                    gt_org = rec.get("genizah_title_org")
                    gt_eng = rec.get("genizah_title_eng")
                    gt = gt_org if gt_org and str(gt_org).strip() else gt_eng
                    if gt and str(gt).strip():
                        titles.append(str(gt).strip())
                gt_vals.append('; '.join(titles) if titles else '')
            html_parts.append(self._field_row(tr('Title'), gt_vals, is_heb))

            # === Section 3: Mentions ===
            if mentions:
                html_parts.append(self._section_row(tr('Mentions'), total_cols))
                mention_types_ordered = ['Personalities', 'Places', 'Creations', 'Dates', 'Groups']
                all_types = set()
                for rec_id, items in mentions.items():
                    for item in items:
                        mt = item.get("mention_type")
                        if mt:
                            all_types.add(mt)
                extra_types = sorted(all_types - set(mention_types_ordered))
                type_order = [t for t in mention_types_ordered if t in all_types] + extra_types
                for mention_type in type_order:
                    mn_vals = []
                    for team in teams:
                        names = []
                        for rec in team["records"]:
                            rec_id = rec.get("unit_catalog_rec_id")
                            if rec_id and rec_id in mentions:
                                for m in mentions[rec_id]:
                                    if m.get("mention_type") == mention_type:
                                        name = m.get("mention", "")
                                        if name and str(name).strip():
                                            names.append(str(name).strip())
                        mn_vals.append(', '.join(names) if names else '')
                    html_parts.append(self._field_row(tr(mention_type), mn_vals, is_heb))

            # === Section 4: Script Description ===
            html_parts.append(self._section_row(tr('Script Description'), total_cols))

            html_parts.append(self._field_category_row('GenizahLanguages', tr('Language'), teams, fields, is_heb))
            html_parts.append(self._field_category_row('TypeOfScript', tr('Script Type'), teams, fields, is_heb))
            html_parts.append(self._field_category_row('TypeOfScriptStyle', tr('Script Style'), teams, fields, is_heb))
            html_parts.append(self._field_category_row('TypeOfScriptPlace', tr('Script Place'), teams, fields, is_heb))
            html_parts.append(self._field_category_row('TypeOfVocalization', tr('Vocalization'), teams, fields, is_heb))

            # === Section 5: Format Description ===
            html_parts.append(self._section_row(tr('Format Description'), total_cols))

            # No. of Rows
            row_vals = []
            for team in teams:
                rows = [str(r.get("num_row", "")).strip() for r in team["records"]
                        if r.get("num_row") and str(r["num_row"]).strip() and str(r["num_row"]).strip() != '0']
                row_vals.append(', '.join(rows) if rows else '')
            html_parts.append(self._field_row(tr('Number of Lines'), row_vals, is_heb))

            # No. of Columns
            col_vals = []
            for team in teams:
                cols = [str(r.get("num_column", "")).strip() for r in team["records"]
                        if r.get("num_column") and str(r["num_column"]).strip() and str(r["num_column"]).strip() != '0']
                col_vals.append(', '.join(cols) if cols else '')
            html_parts.append(self._field_row(tr('Number of Columns'), col_vals, is_heb))

            # Material
            html_parts.append(self._field_category_row('FragmentMaterial', tr('Material'), teams, fields, is_heb))

            # Physical Status
            html_parts.append(self._field_category_row('FragmentStatus', tr('Physical Status'), teams, fields, is_heb))

            # Sizes
            size_vals = []
            for team in teams:
                size_parts = []
                for rec in team["records"]:
                    rec_id = rec.get("unit_catalog_rec_id")
                    if rec_id and rec_id in sizes:
                        for sz in sizes[rec_id]:
                            sx = sz.get("size_x")
                            sy = sz.get("size_y")
                            isx = sz.get("inner_size_x")
                            isy = sz.get("inner_size_y")
                            if sx and sy:
                                dim = f"{self._fmt_num(sx)} \u00d7 {self._fmt_num(sy)}"
                                if isx and isy:
                                    dim += f" ({tr('Inner Size')}: {self._fmt_num(isx)} \u00d7 {self._fmt_num(isy)})"
                                dim += " mm"
                                size_parts.append(dim)
                size_vals.append('; '.join(size_parts) if size_parts else '')
            html_parts.append(self._field_row(tr('Size'), size_vals, is_heb))

        # === Section 6: Miscellaneous ===
        if free_descriptions or full_texts:
            html_parts.append(self._section_row(tr('Miscellaneous'), total_cols if num_teams > 0 else 2))

            col_span = total_cols if num_teams > 0 else 2
            for desc in free_descriptions:
                text = desc.get("text", "")
                if text and str(text).strip():
                    eng_source = desc.get("source_name")
                    source = get_team_display_name(eng_source, is_heb=is_heb) if eng_source else None
                    source_html = f'<div style="font-weight:bold; font-size:11px; color:{c["section_text"]}; margin-bottom:2px;">{source}</div>' if source else ''
                    html_parts.append(
                        f'<tr><td colspan="{col_span}" '
                        f'style="padding:8px; border-bottom:1px solid {c["border"]};"'
                        f'>{source_html}{str(text).strip()}</td></tr>'
                    )

            # Full texts (scholarly descriptions) with distinct styling
            if full_texts:
                html_parts.append(
                    f'<tr><td colspan="{col_span}" style="padding:6px 8px; font-weight:bold; '
                    f'color:{c["section_text"]}; font-size:12px;">{tr("Scholarly Description")}</td></tr>'
                )
                for ft in full_texts:
                    text = ft.get("text", "")
                    if text and str(text).strip():
                        html_parts.append(
                            f'<tr><td colspan="{col_span}" '
                            f'style="padding:8px; border-bottom:1px solid {c["border"]}; background:{c["full_text_bg"]};"'
                            f'>{str(text).strip()}</td></tr>'
                        )

        html_parts.append('</table>')
        return '\n'.join(html_parts)

    def _section_row(self, title: str, colspan: int) -> str:
        """Build a section header row."""
        c = self._colors
        return (
            f'<tr><td colspan="{colspan}" style="background:{c["section_bg"]}; font-weight:bold; '
            f'padding:8px; color:{c["section_text"]}; font-size:13px;">{title}</td></tr>'
        )

    def _field_row(self, label: str, values: list, is_heb: bool) -> str:
        """Build a field row: label + value columns. RTL: values first, label last.
        Qt handles text alignment via app-level layout direction. Returns '' if all values empty."""
        if not any(v for v in values):
            return ''
        c = self._colors
        label_cell = f'<td style="padding:6px 8px; font-weight:bold; color:{c["label"]}; vertical-align:top; word-wrap:break-word; overflow-wrap:break-word;">{label}</td>'
        value_cells = []
        for val in values:
            display = str(val).strip() if val else '\u2014'
            style = f'padding:6px 8px; border-bottom:1px solid {c["border"]}; vertical-align:top; word-wrap:break-word; overflow-wrap:break-word;'
            if not val:
                style += f' color:{c["muted"]};'
            value_cells.append(f'<td style="{style}">{display}</td>')
        if is_heb:
            return '<tr>' + ''.join(value_cells) + label_cell + '</tr>'
        return '<tr>' + label_cell + ''.join(value_cells) + '</tr>'

    def _field_category_row(self, category: str, label: str, teams: list, fields: dict, is_heb: bool) -> str:
        """Build a row for a specific FieldCategory from catalog_fields."""
        vals = []
        for team in teams:
            field_vals = []
            for rec in team["records"]:
                rec_id = rec.get("unit_catalog_rec_id")
                if rec_id and rec_id in fields:
                    cat_fields = fields[rec_id].get(category, [])
                    for fv in cat_fields:
                        val = fv.get("value_heb") if is_heb else fv.get("value")
                        if not val or not str(val).strip():
                            val = fv.get("value") or fv.get("value_heb")
                        if val and str(val).strip():
                            field_vals.append(str(val).strip())
            vals.append('; '.join(field_vals) if field_vals else '')
        return self._field_row(label, vals, is_heb)

    @staticmethod
    def _fmt_num(val) -> str:
        """Format a numeric value for size display, removing trailing .0."""
        if val is None:
            return ""
        s = str(val)
        if s.endswith('.0'):
            return s[:-2]
        return s

    @staticmethod
    def _fmt_int(val) -> str:
        """Format a numeric value as integer (2.0 → '2')."""
        if val is None:
            return ""
        s = str(val).strip()
        if s.endswith('.0'):
            return s[:-2]
        return s


class NliBibliographyDialog(QDialog):
    """NLI bibliography dialog with MARC 581 reference strings."""

    def __init__(self, marc_strings, sys_id='', shelfmark='', parent=None):
        super().__init__(parent)
        from shared.fjms_service import _parse_marc_annotations, strip_marc_annotation_suffix, _ts_symbol
        self.marc_strings = marc_strings
        self.sys_id = sys_id
        self._ts_symbol = _ts_symbol
        self.setWindowTitle(f"{tr('Bibliography Ktiv')} \u2014 {shelfmark}" if shelfmark else tr('Bibliography Ktiv'))
        self.setMinimumSize(900, 500)

        # Pre-parse all MARC strings
        self.parsed = []
        for ms in marc_strings:
            ann = _parse_marc_annotations(ms)
            ref = strip_marc_annotation_suffix(ms)
            self.parsed.append({
                'reference': ref,
                'raw': ms,
                'mention_type': ann.get('mention_type', ''),
                'has_image': ann.get('has_image', False),
                'transcription': ann.get('transcription', ''),
                'translation': ann.get('translation', ''),
            })

        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)

        # Filter row
        filter_row = QHBoxLayout()
        self.text_filter = QLineEdit()
        self.text_filter.setPlaceholderText(tr('Filter references...'))
        self.text_filter.textChanged.connect(self._filter_rows)
        filter_row.addWidget(self.text_filter, 1)
        self.type_combo = QComboBox()
        for label, val in [(tr('All'), 'All'), (tr('Discussion'), 'Discussion'),
                           (tr('Mentioned'), 'Mentioned'), (tr('Index'), 'Index')]:
            self.type_combo.addItem(label, val)
        self.type_combo.currentIndexChanged.connect(lambda _: self._filter_rows())
        filter_row.addWidget(QLabel(tr('Type') + ':'))
        filter_row.addWidget(self.type_combo)
        layout.addLayout(filter_row)

        check_row = QHBoxLayout()
        self.chk_transcription = QCheckBox(tr('Has Transcription'))
        self.chk_transcription.toggled.connect(self._filter_rows)
        check_row.addWidget(self.chk_transcription)
        self.chk_translation = QCheckBox(tr('Has Translation'))
        self.chk_translation.toggled.connect(self._filter_rows)
        check_row.addWidget(self.chk_translation)
        self.chk_image = QCheckBox(tr('Has Image'))
        self.chk_image.toggled.connect(self._filter_rows)
        check_row.addWidget(self.chk_image)
        check_row.addStretch()
        layout.addLayout(check_row)

        # Table: Reference, D, T, S, I
        headers = [tr('Reference'), tr('col_D'), tr('col_T'), tr('col_S'), tr('col_I')]
        self.table = QTableWidget(len(self.parsed), len(headers))
        self.table.setHorizontalHeaderLabels(headers)
        hdr_model = self.table.horizontalHeader().model()
        for col_idx, tooltip in [(1, tr('Discussion')), (2, tr('Transcription')),
                                  (3, tr('Translation')), (4, tr('Image'))]:
            hdr_model.setHeaderData(col_idx, Qt.Orientation.Horizontal, tooltip, Qt.ItemDataRole.ToolTipRole)
        self.table.setSortingEnabled(False)
        self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.table.verticalHeader().setVisible(False)
        self.table.setAlternatingRowColors(True)
        for col_idx in (1, 2, 3, 4):
            self.table.setColumnWidth(col_idx, 36)

        for row, pe in enumerate(self.parsed):
            item0 = QTableWidgetItem(pe['reference'])
            item0.setData(Qt.ItemDataRole.UserRole, row)
            self.table.setItem(row, 0, item0)
            self.table.setItem(row, 1, QTableWidgetItem('\u2713' if pe['mention_type'] == 'Discussion' else ''))
            self.table.setItem(row, 2, QTableWidgetItem(_ts_symbol(pe['transcription'])))
            self.table.setItem(row, 3, QTableWidgetItem(_ts_symbol(pe['translation'])))
            self.table.setItem(row, 4, QTableWidgetItem('\u2713' if pe['has_image'] else ''))

        self.table.resizeColumnsToContents()
        for col_idx in (1, 2, 3, 4):
            self.table.setColumnWidth(col_idx, 36)
        self.table.setSortingEnabled(True)
        self.table.currentCellChanged.connect(self._on_row_selected)
        layout.addWidget(self.table, 1)

        # Detail panel
        self.detail_panel = QTextBrowser()
        self.detail_panel.setMaximumHeight(80)
        self.detail_panel.setVisible(False)
        self.detail_panel.setStyleSheet("border: 1px solid #ccc; padding: 4px; font-size: 12px;")
        layout.addWidget(self.detail_panel)

        # Bottom row
        bottom_row = QHBoxLayout()
        if sys_id:
            ktiv_url = f"https://www.nli.org.il/he/discover/manuscripts/hebrew-manuscripts/itempage?vid=KTIV&scope=KTIV&docId=PNX_MANUSCRIPTS{sys_id}"
            btn_ktiv = QPushButton(tr('Open in KTIV'))
            btn_ktiv.clicked.connect(lambda: QDesktopServices.openUrl(QUrl(ktiv_url)))
            bottom_row.addWidget(btn_ktiv)
        bottom_row.addStretch()
        btn_close = QPushButton(tr('Close'))
        btn_close.clicked.connect(self.close)
        bottom_row.addWidget(btn_close)
        layout.addLayout(bottom_row)

    def _filter_rows(self):
        text_val = self.text_filter.text().strip().lower()
        type_val = self.type_combo.currentData() or 'All'
        need_trans = self.chk_transcription.isChecked()
        need_transl = self.chk_translation.isChecked()
        need_image = self.chk_image.isChecked()
        skip_vals = ('', 'None', 'Unknown')

        for row in range(self.table.rowCount()):
            item = self.table.item(row, 0)
            orig_idx = item.data(Qt.ItemDataRole.UserRole) if item else -1
            if not isinstance(orig_idx, int) or orig_idx < 0 or orig_idx >= len(self.parsed):
                continue
            pe = self.parsed[orig_idx]
            show = True
            if type_val != 'All' and pe['mention_type'] != type_val:
                show = False
            if show and need_trans:
                if not pe['transcription'] or pe['transcription'] in skip_vals:
                    show = False
            if show and need_transl:
                if not pe['translation'] or pe['translation'] in skip_vals:
                    show = False
            if show and need_image:
                if not pe['has_image']:
                    show = False
            if show and text_val:
                if text_val not in pe['reference'].lower() and text_val not in pe['raw'].lower():
                    show = False
            self.table.setRowHidden(row, not show)

    def _on_row_selected(self, row, col, prev_row, prev_col):
        item = self.table.item(row, 0)
        orig_idx = item.data(Qt.ItemDataRole.UserRole) if item else -1
        if isinstance(orig_idx, int) and 0 <= orig_idx < len(self.parsed):
            pe = self.parsed[orig_idx]
            parts = [pe['raw']]
            details = []
            if pe['mention_type']:
                details.append(tr(pe['mention_type']))
            if pe['transcription']:
                details.append(f"{tr('Transcription')}: {tr(pe['transcription'])}")
            if pe['translation']:
                details.append(f"{tr('Translation')}: {tr(pe['translation'])}")
            if pe['has_image']:
                details.append(tr('Has Image'))
            if details:
                parts.append(', '.join(details))
            self.detail_panel.setPlainText('\n'.join(parts))
            self.detail_panel.setVisible(True)
        else:
            self.detail_panel.setVisible(False)


class TabularQueryBuilderDialog(QDialog):
    """Tabular Query Builder for Responsa syntax composition.

    Provides a visual interface for composing Responsa queries using
    2-4 component columns with per-word modifiers and distance controls.
    """

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle(tr("Tabular Search"))
        self.setMinimumSize(750, 500)
        self.resize(800, 550)
        self.setLayoutDirection(Qt.LayoutDirection.RightToLeft)

        self._syntax = ''
        self._negated_words = []
        self._active_word = None  # (comp_idx, word_idx)
        self._updating_modifiers = False
        self._max_components = 4
        self._max_words_per_component = 4
        self._initial_words_visible = 2
        self._initial_components = 2

        # Internal state
        self._component_data = []  # List of component state dicts
        self._distance_spinners = []  # QSpinBox list
        self._component_widgets = []  # List of component UI widget groups
        self._distance_containers = []  # Container widgets for distance spinners

        # Dark mode detection
        palette = self.palette()
        self._is_dark = palette.color(palette.ColorRole.Window).lightness() < 128

        self._setup_ui()
        self._initialize_components()

    def _setup_ui(self):
        """Build the complete dialog UI."""
        main_layout = QVBoxLayout(self)
        main_layout.setSpacing(8)

        # --- Scope Row ---
        scope_row = QHBoxLayout()
        scope_row.addWidget(QLabel(tr("Scope") + ":"))
        from PyQt6.QtWidgets import QRadioButton, QButtonGroup
        self._scope_group = QButtonGroup(self)
        self._rb_word_range = QRadioButton(tr("Word Range"))
        self._rb_word_range.setChecked(True)
        self._rb_within_doc = QRadioButton(tr("Within Document"))
        self._scope_group.addButton(self._rb_word_range, 0)
        self._scope_group.addButton(self._rb_within_doc, 1)
        scope_row.addWidget(self._rb_word_range)
        scope_row.addWidget(self._rb_within_doc)
        scope_row.addStretch()
        self._scope_group.idToggled.connect(self._on_scope_changed)
        main_layout.addLayout(scope_row)

        # --- Components Area (scrollable) ---
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setMinimumHeight(250)

        self._components_container = QWidget()
        self._components_layout = QHBoxLayout(self._components_container)
        self._components_layout.setSpacing(6)
        self._components_layout.setContentsMargins(4, 4, 4, 4)
        scroll.setWidget(self._components_container)
        main_layout.addWidget(scroll)

        # --- Add Component Button ---
        self._btn_add_component = QPushButton("+ " + tr("Add Component"))
        self._btn_add_component.setFixedWidth(160)
        self._btn_add_component.clicked.connect(self._add_component)
        add_comp_row = QHBoxLayout()
        add_comp_row.addWidget(self._btn_add_component)
        add_comp_row.addStretch()
        main_layout.addLayout(add_comp_row)

        # --- Modifiers Row ---
        mod_row = QHBoxLayout()
        mod_row.addWidget(QLabel(tr("Modifiers") + ":"))

        self.chk_prefix = QCheckBox(tr("Prefixes #_"))
        self.chk_prefix.setToolTip(tr("Grammatical prefixes tooltip"))
        mod_row.addWidget(self.chk_prefix)

        self.chk_suffix = QCheckBox(tr("Suffixes _#"))
        self.chk_suffix.setToolTip(tr("Grammatical suffixes tooltip"))
        mod_row.addWidget(self.chk_suffix)

        self.chk_wild_start = QCheckBox(tr("Wildcard *_"))
        self.chk_wild_start.setToolTip(tr("Words ending with..."))
        mod_row.addWidget(self.chk_wild_start)

        self.chk_wild_end = QCheckBox(tr("Wildcard _*"))
        self.chk_wild_end.setToolTip(tr("Words starting with..."))
        mod_row.addWidget(self.chk_wild_end)

        self.chk_plene = QCheckBox(tr("Plene/Defective %"))
        self.chk_plene.setToolTip(tr("Plene/defective spelling tooltip"))
        mod_row.addWidget(self.chk_plene)

        self.chk_negation = QCheckBox(tr("Negation −"))
        self.chk_negation.setToolTip(tr("Negation tooltip"))
        mod_row.addWidget(self.chk_negation)

        mod_row.addStretch()
        main_layout.addLayout(mod_row)

        # Connect modifier checkboxes
        for chk in [self.chk_prefix, self.chk_suffix, self.chk_wild_start,
                     self.chk_wild_end, self.chk_plene, self.chk_negation]:
            chk.stateChanged.connect(self._on_modifier_changed)

        # --- Search Options Row ---
        opts_row = QHBoxLayout()
        opts_row.addWidget(QLabel(tr("Search Options") + ":"))
        self.chk_opt_variants = QCheckBox(tr("Variants"))
        self.chk_opt_ja = QCheckBox(tr("Judeo-Arabic"))
        self.chk_opt_flex = QCheckBox(tr("Flex Spacing"))
        self.chk_opt_bidir = QCheckBox(tr("Bidirectional"))
        opts_row.addWidget(self.chk_opt_variants)
        opts_row.addWidget(self.chk_opt_ja)
        opts_row.addWidget(self.chk_opt_flex)
        opts_row.addWidget(self.chk_opt_bidir)
        opts_row.addStretch()
        main_layout.addLayout(opts_row)

        # --- Preview Row ---
        preview_row = QHBoxLayout()
        preview_row.addWidget(QLabel(tr("Preview") + ":"))
        self._preview_label = QLabel("")
        if self._is_dark:
            preview_bg = '#2d2d2d'
            preview_border = '#555'
        else:
            preview_bg = '#f8f9fa'
            preview_border = '#dee2e6'
        self._preview_label.setStyleSheet(
            f"font-family: 'Consolas', 'Courier New', monospace; font-size: 13px; "
            f"padding: 4px 8px; background: {preview_bg}; border: 1px solid {preview_border}; border-radius: 4px; "
            f"min-height: 22px;"
        )
        self._preview_label.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        self._preview_label.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
        preview_row.addWidget(self._preview_label, 1)
        main_layout.addLayout(preview_row)

        # --- Buttons Row ---
        btn_row = QHBoxLayout()
        btn_clear = QPushButton(tr("Clear All"))
        btn_clear.clicked.connect(self._clear_all)
        btn_row.addWidget(btn_clear)
        btn_row.addStretch()
        btn_cancel = QPushButton(tr("Cancel"))
        btn_cancel.clicked.connect(self.reject)
        btn_row.addWidget(btn_cancel)
        btn_search = QPushButton(tr("Search"))
        btn_search.setStyleSheet("background-color: #27ae60; color: white; font-weight: bold; padding: 6px 20px;")
        btn_search.clicked.connect(self._apply)
        btn_row.addWidget(btn_search)
        main_layout.addLayout(btn_row)

    def _initialize_components(self):
        """Create the initial 2 components with distance spinner between them."""
        for i in range(self._initial_components):
            self._create_component(i)
            if i < self._initial_components - 1:
                self._create_distance_spinner(i)
        self._update_add_component_visibility()

    def _create_component(self, index):
        """Create a component card (QFrame with word inputs)."""
        # Data
        comp_data = {
            'words': [{'text': '', 'mods': {}} for _ in range(self._max_words_per_component)]
        }
        self._component_data.append(comp_data)

        # UI
        frame = QFrame()
        frame.setFrameStyle(QFrame.Shape.Box | QFrame.Shadow.Plain)
        if self._is_dark:
            frame_bg = '#2a2a2a'
            frame_border = '#555'
        else:
            frame_bg = '#fafafa'
            frame_border = '#bdc3c7'
        frame.setStyleSheet(
            f"QFrame {{ border: 1px solid {frame_border}; border-radius: 6px; background: {frame_bg}; }}"
        )
        frame_layout = QVBoxLayout(frame)
        frame_layout.setSpacing(4)
        frame_layout.setContentsMargins(8, 6, 8, 6)

        # Title
        title_label = QLabel(tr("Component") + f" {index + 1}")
        title_color = '#ddd' if self._is_dark else '#333'
        title_label.setStyleSheet(f"font-weight: bold; font-size: 12px; border: none; background: transparent; color: {title_color};")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        frame_layout.addWidget(title_label)

        # Word inputs and modifier indicators
        inputs = []
        indicators = []
        ind_color = '#7cabd4' if self._is_dark else '#2980b9'
        for wi in range(self._max_words_per_component):
            inp = QLineEdit()
            inp.setPlaceholderText(tr("Word") + f" {wi + 1}")
            inp.setMinimumWidth(120)
            if self._is_dark:
                inp_bg = '#3a3a3a'
                inp_border = '#666'
                inp_color = '#eee'
            else:
                inp_bg = 'white'
                inp_border = '#ccc'
                inp_color = '#333'
            inp.setStyleSheet(f"border: 1px solid {inp_border}; border-radius: 3px; padding: 3px; background: {inp_bg}; color: {inp_color};")
            inp.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
            inp.installEventFilter(self)
            inp.textChanged.connect(self._on_word_text_changed)
            frame_layout.addWidget(inp)
            # Modifier indicator label below input
            mod_ind = QLabel("")
            mod_ind.setStyleSheet(f"font-size: 9px; color: {ind_color}; border: none; background: transparent; margin-top: -2px;")
            mod_ind.setVisible(False)
            frame_layout.addWidget(mod_ind)
            indicators.append(mod_ind)
            inputs.append(inp)
            # Hide extra word slots
            if wi >= self._initial_words_visible:
                inp.setVisible(False)
                mod_ind.setVisible(False)

        # Add word button
        btn_add_word = QPushButton("+ " + tr("Add Word"))
        add_word_color = '#5dade2' if self._is_dark else '#2980b9'
        btn_add_word.setStyleSheet(f"font-size: 10px; border: none; color: {add_word_color}; background: transparent;")
        btn_add_word.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        frame_layout.addWidget(btn_add_word)

        # Remove button (only for components 3+)
        btn_remove = QPushButton(tr("Remove"))
        remove_color = '#e74c3c' if self._is_dark else '#c0392b'
        btn_remove.setStyleSheet(f"font-size: 10px; color: {remove_color}; border: none; background: transparent;")
        btn_remove.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        btn_remove.setVisible(index >= self._initial_components)
        frame_layout.addWidget(btn_remove)

        frame_layout.addStretch()

        comp_widget = {
            'frame': frame,
            'inputs': inputs,
            'indicators': indicators,
            'btn_add_word': btn_add_word,
            'btn_remove': btn_remove,
            'title_label': title_label,
            'visible_words': self._initial_words_visible,
        }
        self._component_widgets.append(comp_widget)

        # Connect buttons with closure over index
        ci = len(self._component_widgets) - 1
        btn_add_word.clicked.connect(lambda checked=False, idx=ci: self._show_next_word(idx))
        btn_remove.clicked.connect(lambda checked=False, idx=ci: self._remove_component(idx))

        self._components_layout.addWidget(frame)
        self._update_add_word_visibility(ci)

    def _create_distance_spinner(self, pair_index):
        """Create a distance spinner between components pair_index and pair_index+1."""
        container = QWidget()
        container_layout = QVBoxLayout(container)
        container_layout.setContentsMargins(2, 0, 2, 0)
        container_layout.setSpacing(2)
        container_layout.addStretch()

        dist_label = QLabel(tr("Distance"))
        dist_color = '#aab' if self._is_dark else '#7f8c8d'
        dist_label.setStyleSheet(f"font-size: 10px; color: {dist_color};")
        dist_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        container_layout.addWidget(dist_label)

        spinner = QSpinBox()
        spinner.setRange(0, 50)
        spinner.setValue(0)
        spinner.setFixedWidth(60)
        spinner.setAlignment(Qt.AlignmentFlag.AlignCenter)
        spinner.valueChanged.connect(lambda v: self._update_preview())
        container_layout.addWidget(spinner)

        words_label = QLabel(tr("words"))
        words_sub_color = '#999' if self._is_dark else '#95a5a6'
        words_label.setStyleSheet(f"font-size: 9px; color: {words_sub_color};")
        words_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        container_layout.addWidget(words_label)

        container_layout.addStretch()

        self._distance_spinners.append(spinner)
        self._distance_containers.append(container)

        # Insert in layout before the next component
        # The layout has: comp0, dist0, comp1, dist1, comp2, ...
        # We insert at position 2*pair_index + 1
        insert_pos = 2 * pair_index + 1
        self._components_layout.insertWidget(insert_pos, container)

        # Hide if scope is Within Document
        if self._rb_within_doc.isChecked():
            container.setVisible(False)

    def _update_add_word_visibility(self, comp_idx):
        """Show/hide the + button based on how many word slots are visible."""
        if comp_idx >= len(self._component_widgets):
            return
        cw = self._component_widgets[comp_idx]
        visible_count = cw['visible_words']
        cw['btn_add_word'].setVisible(visible_count < self._max_words_per_component)

    def _update_add_component_visibility(self):
        """Show/hide the + Component button based on current count."""
        active_count = len(self._component_widgets)
        self._btn_add_component.setVisible(active_count < self._max_components)

    def _show_next_word(self, comp_idx):
        """Reveal the next hidden word input in the given component."""
        if comp_idx >= len(self._component_widgets):
            return
        cw = self._component_widgets[comp_idx]
        visible = cw['visible_words']
        if visible < self._max_words_per_component:
            cw['inputs'][visible].setVisible(True)
            cw['visible_words'] = visible + 1
            self._update_add_word_visibility(comp_idx)

    def _add_component(self):
        """Add a new component (up to max 4)."""
        current_count = len(self._component_widgets)
        if current_count >= self._max_components:
            return
        # Add distance spinner before the new component
        self._create_distance_spinner(current_count - 1)
        self._create_component(current_count)
        self._update_add_component_visibility()
        self._update_preview()

    def _remove_component(self, comp_idx):
        """Remove a component (cannot go below 2)."""
        if len(self._component_widgets) <= self._initial_components:
            return
        if comp_idx < self._initial_components:
            return

        # Remove the component widget
        cw = self._component_widgets.pop(comp_idx)
        cw['frame'].setParent(None)
        cw['frame'].deleteLater()

        # Remove component data
        self._component_data.pop(comp_idx)

        # Remove the distance spinner before this component
        dist_idx = comp_idx - 1
        if dist_idx >= 0 and dist_idx < len(self._distance_spinners):
            self._distance_spinners.pop(dist_idx)
            container = self._distance_containers.pop(dist_idx)
            container.setParent(None)
            container.deleteLater()

        # Reset active word if it was in the removed component
        if self._active_word and self._active_word[0] >= len(self._component_widgets):
            self._active_word = None

        # Renumber component titles
        for i, cw in enumerate(self._component_widgets):
            cw['title_label'].setText(tr("Component") + f" {i + 1}")
            cw['btn_remove'].setVisible(i >= self._initial_components)

        # Reconnect button lambdas (re-bind indices)
        for i, cw in enumerate(self._component_widgets):
            try:
                cw['btn_add_word'].clicked.disconnect()
            except TypeError:
                pass
            try:
                cw['btn_remove'].clicked.disconnect()
            except TypeError:
                pass
            cw['btn_add_word'].clicked.connect(lambda checked=False, idx=i: self._show_next_word(idx))
            cw['btn_remove'].clicked.connect(lambda checked=False, idx=i: self._remove_component(idx))

        self._update_add_component_visibility()
        self._update_preview()

    def _on_scope_changed(self, button_id, checked):
        """Toggle distance spinner visibility based on scope."""
        if not checked:
            return
        show_distances = (button_id == 0)  # 0 = Word Range
        for container in self._distance_containers:
            container.setVisible(show_distances)
        self._update_preview()

    def _on_word_focus(self, comp_idx, word_idx):
        """Handle focus on a word input -- update modifier checkboxes."""
        self._active_word = (comp_idx, word_idx)
        self._updating_modifiers = True
        try:
            mods = self._component_data[comp_idx]['words'][word_idx].get('mods', {})
            self.chk_prefix.setChecked(mods.get('prefix', False))
            self.chk_suffix.setChecked(mods.get('suffix', False))
            self.chk_wild_start.setChecked(mods.get('wildcard_prefix', False))
            self.chk_wild_end.setChecked(mods.get('wildcard_suffix', False))
            self.chk_plene.setChecked(mods.get('plene', False))
            self.chk_negation.setChecked(mods.get('negation', False))
        finally:
            self._updating_modifiers = False

    _MOD_DISPLAY = {
        'prefix': '#_', 'suffix': '_#',
        'wildcard_prefix': '*_', 'wildcard_suffix': '_*',
        'plene': '%', 'negation': '−',
    }

    def _update_mod_indicator(self, ci, wi):
        """Update the modifier indicator label for a specific word."""
        if ci < len(self._component_widgets) and wi < len(self._component_widgets[ci].get('indicators', [])):
            mods = self._component_data[ci]['words'][wi]['mods']
            parts = [v for k, v in self._MOD_DISPLAY.items() if mods.get(k)]
            text = ' '.join(parts)
            ind = self._component_widgets[ci]['indicators'][wi]
            ind.setText(text)
            ind.setVisible(bool(text))

    def _on_modifier_changed(self):
        """Save modifier state to the active word's data."""
        if self._updating_modifiers or self._active_word is None:
            return
        ci, wi = self._active_word
        if ci >= len(self._component_data):
            return
        mods = {
            'prefix': self.chk_prefix.isChecked(),
            'suffix': self.chk_suffix.isChecked(),
            'wildcard_prefix': self.chk_wild_start.isChecked(),
            'wildcard_suffix': self.chk_wild_end.isChecked(),
            'plene': self.chk_plene.isChecked(),
            'negation': self.chk_negation.isChecked(),
        }
        self._component_data[ci]['words'][wi]['mods'] = mods
        self._update_mod_indicator(ci, wi)
        self._update_preview()

    def _on_word_text_changed(self, text):
        """Sync QLineEdit text back to component data and update preview."""
        sender = self.sender()
        if sender is None:
            return
        for ci, cw in enumerate(self._component_widgets):
            for wi, inp in enumerate(cw['inputs']):
                if inp is sender:
                    self._component_data[ci]['words'][wi]['text'] = text
                    self._update_preview()
                    return

    def _update_preview(self):
        """Regenerate syntax from current state and update preview label."""
        # Build components list in generate_tabular_syntax format
        components = []
        for ci, comp in enumerate(self._component_data):
            words = []
            for wi, word_data in enumerate(comp['words']):
                # Only include words from visible slots
                if ci < len(self._component_widgets) and wi < self._component_widgets[ci]['visible_words']:
                    words.append({
                        'text': word_data.get('text', ''),
                        'mods': word_data.get('mods', {}),
                    })
            components.append({'words': words})

        # Build distances list
        distances = [s.value() for s in self._distance_spinners]

        # Get scope
        scope = 'word_range' if self._rb_word_range.isChecked() else 'within_document'

        try:
            syntax, negated = generate_tabular_syntax(components, distances, scope)
            self._syntax = syntax
            self._negated_words = negated
        except Exception:
            self._syntax = ''
            self._negated_words = []

        self._preview_label.setText(self._syntax if self._syntax else "")

    def _clear_all(self):
        """Reset all inputs, modifiers, spinners, and components to initial state."""
        # Remove extra components (keep only initial 2)
        while len(self._component_widgets) > self._initial_components:
            idx = len(self._component_widgets) - 1
            cw = self._component_widgets.pop(idx)
            cw['frame'].setParent(None)
            cw['frame'].deleteLater()
            self._component_data.pop(idx)

        # Remove extra distance spinners
        while len(self._distance_spinners) > self._initial_components - 1:
            self._distance_spinners.pop()
            container = self._distance_containers.pop()
            container.setParent(None)
            container.deleteLater()

        # Reset remaining components
        for ci, cw in enumerate(self._component_widgets):
            for wi, inp in enumerate(cw['inputs']):
                inp.blockSignals(True)
                inp.clear()
                inp.blockSignals(False)
                inp.setVisible(wi < self._initial_words_visible)
            cw['visible_words'] = self._initial_words_visible
            self._update_add_word_visibility(ci)
            # Reset data
            self._component_data[ci] = {
                'words': [{'text': '', 'mods': {}} for _ in range(self._max_words_per_component)]
            }

        # Reset spinners
        for spinner in self._distance_spinners:
            spinner.blockSignals(True)
            spinner.setValue(0)
            spinner.blockSignals(False)

        # Reset modifiers
        self._active_word = None
        self._updating_modifiers = True
        for chk in [self.chk_prefix, self.chk_suffix, self.chk_wild_start,
                     self.chk_wild_end, self.chk_plene, self.chk_negation]:
            chk.setChecked(False)
        self._updating_modifiers = False

        # Reset scope
        self._rb_word_range.setChecked(True)

        self._update_add_component_visibility()
        self._update_preview()

    def _apply(self):
        """Generate final syntax and accept the dialog."""
        self._update_preview()
        self.accept()

    def get_syntax(self) -> str:
        """Return the generated Responsa syntax string."""
        self._update_preview()
        return self._syntax

    def get_negated_words(self) -> list:
        """Return list of words marked for exclusion."""
        self._update_preview()
        return self._negated_words

    def eventFilter(self, obj, event):
        """Catch focus events on word inputs to update modifier checkboxes."""
        if event.type() == QEvent.Type.FocusIn:
            for ci, comp in enumerate(self._component_widgets):
                for wi, inp in enumerate(comp['inputs']):
                    if inp is obj:
                        self._on_word_focus(ci, wi)
                        return super().eventFilter(obj, event)
        return super().eventFilter(obj, event)


class GenizahGUI(QMainWindow):
    """Main application window orchestrating search, browsing, and indexing."""
    browse_thumb_resolved = pyqtSignal(str, object)
    
    def __init__(self):
        super().__init__()
        self.comp_col_library = 1  # Library before Shelfmark
        self.comp_col_shelfmark = 2
        self.comp_col_title = 3
        self.comp_col_sysid = 4
        self.comp_col_context = 5
        self.comp_col_ms_context = 6
        self.setWindowTitle(tr(f"Genizah Search Pro V{APP_VERSION}"))
        # Initial size - will be overridden by showMaximized() at startup
        self.setMinimumSize(1200, 700)
        log_tls_relaxation_notice()

        self.meta_mgr = None
        self.var_mgr = None
        self.searcher = None
        self.indexer = None
        self.lab_engine = None
        self.lists_mgr = None
        self.joins_mgr = None

        # Community features - corrections client
        self.corrections_client = get_corrections_client()

        self.last_results = []
        self.last_search_query = ""
        self.result_row_by_sys_id = {}
        self.comp_main = []
        self.comp_appendix = {}
        self.comp_summary = {}
        self.comp_filtered_main = []
        self.comp_filtered_appendix = {}
        self.comp_filtered_summary = {}
        self.comp_raw_items = []
        self.comp_raw_filtered = []
        self.comp_grouped_main = []
        self.comp_grouped_appendix = {}
        self.comp_grouped_summary = {}
        self.comp_sort_mode = "score"
        self.comp_sort_reverse = True
        self.comp_grouped_filtered_main = []
        self.comp_grouped_filtered_appendix = {}
        self.comp_grouped_filtered_summary = {}
        self.comp_has_grouped_results = False
        self.comp_known = []
        self.pending_recursive_search = False
        self.excluded_raw_entries = []
        self.excluded_sys_ids = set()
        self.excluded_shelfmarks = set()
        self.filter_sources = {}  # dict of {ref: cleaned_text}
        self.filter_enabled_sources = set()  # set of enabled source refs
        self.results_filters = {}
        self.list_filter_state = {'active': False, 'mode': 'in', 'lists': 'all'}
        self._pgp_transcription_sys_ids = set()
        self._pgp_badge_worker = None
        self._domain_worker = None
        self._pgp_tags_worker = None
        self._pgp_tag_search_worker = None
        self._pgp_tags = []
        self.comp_filters = {}
        self.group_thread = None
        self.is_searching = False
        self.is_comp_running = False
        self.last_browse_field = None
        self.current_browse_sid = None
        self.current_browse_p = None
        self.current_browse_internal_idx = None
        self.browse_highlight_data = []
        self.browse_highlight_pattern = None
        # Codicological Parts (Neubauer) browsing state
        self.current_browse_part_id = None
        self.current_browse_part_folios = []
        self.current_browse_part_folio_idx = 0
        self.meta_loader = None
        self.meta_cached_count = 0
        self.meta_to_fetch_count = 0
        self.meta_progress_current = 0
        self.browse_thumb_url = None
        self.browse_img_thread = None
        self.shelfmark_items_by_sid = {}
        self.title_items_by_sid = {}
        self.comp_thread = None 
        self.comp_worker = None
        self.hovered_row = -1
        self.lists_hovered_row = -1
        self.results_loaded = 0
        self.snippet_queue = []

        # Shelfmark to sys_id mapping for community features
        self._shelf_to_sys = None

        self.init_ui()

        # Step 2: Start heavy initialization in background
        self.status_label.setText(tr("Initializing components... Please wait."))
        self.set_results_loading(True)
        QTimer.singleShot(100, self.start_background_init)

    def start_background_init(self):
        try:
            self.startup_thread = StartupThread()
            self.startup_thread.finished_signal.connect(self.on_startup_finished)
            self.startup_thread.error_signal.connect(lambda e: QMessageBox.critical(self, tr("Fatal Error"), tr("Failed to initialize:\n{}").format(e)))
            self.startup_thread.start()
        except Exception as e:
            QMessageBox.critical(self, tr("Fatal Error"), tr("Failed to start initialization:\n{}").format(e))

    def on_startup_finished(self, meta_mgr, var_mgr, searcher, indexer):
        try:
            self.meta_mgr = meta_mgr
            self.var_mgr = var_mgr
            self.searcher = searcher
            self.indexer = indexer

            # Init Lists Manager
            self.lists_mgr = ListsManager(self.meta_mgr)

            # Init Joins Manager (offline-first fragment connections)
            self.joins_mgr = JoinsManager(self.corrections_client)
            self.joins_mgr.start_background_sync()

            # Init Lab Engine (lightweight init)
            self.lab_engine = LabEngine(self.meta_mgr, self.var_mgr)

            # Connect VariantManager to Lab settings for variant search configuration
            if self.var_mgr and self.lab_engine:
                self.var_mgr.set_settings(self.lab_engine.settings)

            # Setup Panels (guaranteed to exist as init_ui runs before startup thread)
            if hasattr(self, 'lab_panel_search'):
                self.lab_panel_search.set_engine(self.lab_engine)
            else:
                logger.warning("lab_panel_search not found during startup finish")

            if hasattr(self, 'lab_panel_comp'):
                self.lab_panel_comp.set_engine(self.lab_engine)
            else:
                logger.warning("lab_panel_comp not found during startup finish")

            os.makedirs(Config.REPORTS_DIR, exist_ok=True)
            self.browse_thumb_resolved.connect(self._on_browse_thumb_resolved)

            # Enable UI interactions
            self.btn_search.setEnabled(True)
            self.btn_comp_run.setEnabled(True)
            self.btn_browse_go.setEnabled(True)
            self.btn_build_index.setEnabled(True)
            
            self.status_label.setText(tr("Components loaded. Ready."))
            self.set_results_loading(False)

            # Initialize Lists Tab UI
            self.lists_refresh_all()

            db_path = os.path.join(Config.INDEX_DIR, "tantivy_db")
            index_exists = os.path.exists(db_path) and os.listdir(db_path)
            
            if not index_exists:
                msg = tr("Index not found.\nWould you like to build it now?\n(Requires 'Transcriptions.txt' next to this app)")
                reply = QMessageBox.question(self, tr("Index Missing"), msg,
                                             QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                if reply == QMessageBox.StandardButton.Yes:
                    self.tabs.setCurrentIndex(3) 
                    self.run_indexing()

            # Start checking for CSV bank readiness to init Autocomplete
            self.shelf_init_timer = QTimer(self)
            self.shelf_init_timer.timeout.connect(self._check_shelfmark_completer_ready)
            self.shelf_init_timer.start(500)
            # Automatic Update Check
            self.check_updates_auto()

            # Show What's New bar if version is new
            cfg = load_app_config()
            if cfg.get('whats_new_seen') != APP_VERSION:
                self.whats_new_bar.show_whats_new(APP_VERSION)

        except Exception as e:
            QMessageBox.critical(self, tr("Fatal Error"), tr("Failed to finalize initialization:\n{}").format(e))

    def _check_shelfmark_completer_ready(self):
        if self.meta_mgr and len(self.meta_mgr.csv_bank) > 0:
            if self.setup_shelfmark_completer():
                self.shelf_init_timer.stop()
                self.shelf_init_timer = None
        # Add a timeout/limit? For now, we assume it eventually loads or stays empty (if file missing).
        # If libraries.csv is missing, csv_bank remains empty, so this never setups. That's acceptable.

    def setup_shelfmark_completer(self):
        """Initialize the shelfmark autocomplete with data from csv_bank and Parts."""
        if not self.meta_mgr: return False

        # 1. Extract unique shelfmarks (Protected against background updates)
        try:
            shelfmarks = sorted(list({v['shelfmark'] for v in self.meta_mgr.csv_bank.values() if v.get('shelfmark')}))
        except RuntimeError:
            # Dictionary changed size during iteration (background loader still running)
            return False

        if not shelfmarks: return False

        # 2. Setup Models (Optimized with QStandardItemModel + UserRole)
        self.shelf_model = QStandardItemModel()
        self.valid_shelf_keys = set()

        # Add regular shelfmarks
        for s in shelfmarks:
            item = QStandardItem(s)
            norm = ShelfmarkCompleter.normalize(s)
            self.valid_shelf_keys.add(norm)
            item.setData(norm, Qt.ItemDataRole.UserRole)
            self.shelf_model.appendRow(item)

        # 2b. Add Codicological Parts (Neubauer) with (neubauer) suffix
        try:
            part_list = self.meta_mgr.get_part_autocomplete_list()
            for part_info in part_list:
                display = part_info['display']  # e.g., "MS. Heb. d. 29/2 (neubauer)"
                normalized = part_info['normalized']

                item = QStandardItem(display)
                item.setData(normalized, Qt.ItemDataRole.UserRole)
                # Store part_id for later retrieval
                item.setData(part_info['part_id'], Qt.ItemDataRole.UserRole + 1)
                self.valid_shelf_keys.add(normalized)
                self.shelf_model.appendRow(item)
        except Exception as e:
            # Parts might not be loaded yet, that's OK
            pass

        # 3. Setup Completer
        # Note: We use ShelfmarkCompleter which overrides splitPath to normalize input
        self.shelf_completer = ShelfmarkCompleter(self.shelf_model, self, valid_keys=self.valid_shelf_keys)
        self.shelf_completer.setCompletionRole(Qt.ItemDataRole.UserRole)
        self.shelf_completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        self.shelf_completer.setCompletionMode(QCompleter.CompletionMode.PopupCompletion)
        self.shelf_completer.setFilterMode(Qt.MatchFlag.MatchStartsWith)

        # 4. Attach to Input
        if hasattr(self, 'browse_shelf_input'):
            self.browse_shelf_input.setCompleter(self.shelf_completer)

        return True
             
    def init_ui(self):
        if CURRENT_LANG == 'he':
            QApplication.instance().setLayoutDirection(Qt.LayoutDirection.RightToLeft)

        self.tabs = QTabWidget()
        self.search_tab = self.create_search_tab()
        self.composition_tab = self.create_composition_tab()
        self.browse_tab = self.create_browse_tab()
        self.catalog_browse_tab = self.create_catalog_browse_tab()
        self.lists_tab = self.create_lists_tab()
        self.community_tab = self.create_community_tab()
        self.settings_tab = self.create_settings_tab()
        self.tabs.addTab(self.search_tab, tr("Search"))
        self.tabs.addTab(self.composition_tab, tr("Composition Search"))
        self.tabs.addTab(self.browse_tab, tr("Browse by Shelfmark"))
        self.tabs.addTab(self.catalog_browse_tab, tr("Browse by Identification"))
        self.tabs.addTab(self.lists_tab, tr("Personal Lists"))
        self.tabs.addTab(self.community_tab, tr("Community"))
        self.tabs.addTab(self.settings_tab, tr("Settings & About"))

        # Corner widget with Website, Version, Login button and Language toggle
        corner_widget = QWidget()
        corner_layout = QHBoxLayout(corner_widget)
        corner_layout.setContentsMargins(5, 0, 5, 0)
        corner_layout.setSpacing(10)

        # Website button (highlighted)
        self.corner_website_btn = QPushButton("🌐 GenizahSearch.com")
        self.corner_website_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.corner_website_btn.setToolTip(tr("Visit our website"))
        self.corner_website_btn.clicked.connect(lambda: QDesktopServices.openUrl(QUrl("https://genizahsearch.com")))
        self.corner_website_btn.setStyleSheet("""
            QPushButton {
                background-color: #e8f4fc;
                color: #1a73e8;
                border: 1px solid #1a73e8;
                border-radius: 4px;
                padding: 3px 8px;
                font-weight: bold;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: #1a73e8;
                color: white;
            }
        """)
        corner_layout.addWidget(self.corner_website_btn)

        # Separator
        sep0 = QLabel("|")
        sep0.setStyleSheet("color: gray;")
        corner_layout.addWidget(sep0)

        # Version button
        self.corner_version_btn = QPushButton(f"v{APP_VERSION}")
        self.corner_version_btn.setFlat(True)
        self.corner_version_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.corner_version_btn.setToolTip(tr("Check for updates"))
        self.corner_version_btn.clicked.connect(self.check_updates_manual)
        self.corner_version_btn.setStyleSheet("color: #7f8c8d; font-size: 11px;")
        corner_layout.addWidget(self.corner_version_btn)

        # Separator
        sep1 = QLabel("|")
        sep1.setStyleSheet("color: gray;")
        corner_layout.addWidget(sep1)

        # Login/Logout button
        self.corner_login_btn = QPushButton(tr("Login"))
        self.corner_login_btn.setFlat(True)
        self.corner_login_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.corner_login_btn.clicked.connect(self._corner_login_clicked)
        corner_layout.addWidget(self.corner_login_btn)

        # Separator
        sep2 = QLabel("|")
        sep2.setStyleSheet("color: gray;")
        corner_layout.addWidget(sep2)

        # Language toggle
        self.lang_btn = QPushButton("English" if CURRENT_LANG == 'he' else "עברית")
        self.lang_btn.setFlat(True)
        self.lang_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.lang_btn.clicked.connect(self.toggle_language)
        corner_layout.addWidget(self.lang_btn)

        self.tabs.setCornerWidget(corner_widget, Qt.Corner.TopRightCorner if CURRENT_LANG == 'en' else Qt.Corner.TopLeftCorner)
        self._update_corner_login_state()

        # Connect tab change to refresh community data when needed
        self.tabs.currentChanged.connect(self._on_tab_changed)
        self._community_data_loaded = False

        # Main Layout wrapper for Notification Bar
        central_widget = QWidget()
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Update Bar (Hidden by default)
        self.update_bar = UpdateNotificationBar()
        self.update_bar.dismissed.connect(self.on_update_dismissed)
        self.update_bar.update_requested.connect(self.start_in_app_update)
        main_layout.addWidget(self.update_bar)

        # What's New Bar (Hidden by default)
        self.whats_new_bar = WhatsNewBar()
        self.whats_new_bar.dismissed.connect(self.on_whats_new_dismissed)
        self.whats_new_bar.learn_more.connect(self.show_whats_new_dialog)
        main_layout.addWidget(self.whats_new_bar)

        main_layout.addWidget(self.tabs)
        self.setCentralWidget(central_widget)

    def _update_corner_login_state(self):
        """Update the corner login button based on login state."""
        if self.corrections_client.is_logged_in():
            user = self.corrections_client.current_user
            if user:
                self.corner_login_btn.setText(f"{user.username} ({tr('Logout')})")
            else:
                self.corner_login_btn.setText(tr("Logout"))
        else:
            self.corner_login_btn.setText(tr("Login"))

    def _on_tab_changed(self, index):
        """Handle tab change events."""
        import sys
        logger.debug("_on_tab_changed called with index=%s", index)
        try:
            current_widget = self.tabs.widget(index)
            logger.debug("current_widget=%s", current_widget)
            if hasattr(self, 'community_tab') and current_widget == self.community_tab:
                logger.debug("Matched community_tab, _community_data_loaded=%s", getattr(self, '_community_data_loaded', False))
                # Load community data when tab is first shown or refresh if needed
                if not getattr(self, '_community_data_loaded', False):
                    logger.debug("About to call _refresh_community_panels")
                    self._refresh_community_panels()
                    logger.debug("_refresh_community_panels completed")
                    self._community_data_loaded = True
            # Lazy-load catalog browse tree on first tab activation
            if hasattr(self, 'catalog_browse_tab') and current_widget == self.catalog_browse_tab:
                if not getattr(self, '_catalog_tree_loaded', False):
                    self._catalog_populate_tree()
        except Exception as e:
            import traceback
            print(f"Error in _on_tab_changed: {e}", flush=True)
            traceback.print_exc()
            sys.stdout.flush()

    def _corner_login_clicked(self):
        """Handle corner login button click."""
        if self.corrections_client.is_logged_in():
            self._do_logout()
        else:
            self._show_login_dialog()

    def _show_login_dialog(self):
        dialog = LoginDialog(self, self.corrections_client)
        if dialog.exec():
            self._update_corner_login_state()
            self._refresh_community_panels()
            # Enable cloud sync for lists after successful login
            self._enable_lists_cloud_sync()

    def _show_register_dialog(self):
        dialog = RegisterDialog(self, self.corrections_client)
        if dialog.exec():
            self._update_corner_login_state()
            self._refresh_community_panels()

    def _do_logout(self):
        # Disable cloud sync before logout
        self._disable_lists_cloud_sync()
        self.corrections_client.logout()
        self._update_corner_login_state()
        self._refresh_community_panels()
        QMessageBox.information(self, tr("Logged Out"), tr("You have been logged out."))

    def _enable_lists_cloud_sync(self):
        """Enable cloud sync for user lists after login - shows sync dialog."""
        try:
            user = self.corrections_client.current_user
            logger.debug(f"Cloud sync: user={user}")

            # Get user UUID - try multiple approaches
            user_uuid = None
            if user:
                # Try _uuid attribute (supabase_corrections_client)
                if hasattr(user, '_uuid') and user._uuid:
                    user_uuid = user._uuid
                    logger.debug(f"Cloud sync: Got UUID from _uuid: {user_uuid[:8]}...")
                # Try getting from supabase auth session directly
                elif hasattr(self.corrections_client, '_client') and self.corrections_client._client:
                    try:
                        session = self.corrections_client._client.auth.get_session()
                        if session and session.user:
                            user_uuid = str(session.user.id)
                            logger.debug(f"Cloud sync: Got UUID from session: {user_uuid[:8]}...")
                    except Exception as e:
                        logger.debug(f"Could not get UUID from session: {e}")

            if not user_uuid:
                logger.warning("Cloud sync: No user UUID available - cannot sync")
                return

            # Enable cloud sync connection (but don't sync yet)
            # Pass the authenticated Supabase client for RLS to work
            logger.info(f"Enabling cloud sync for user UUID: {user_uuid}")
            supabase_client = None
            if hasattr(self.corrections_client, '_client'):
                supabase_client = self.corrections_client._client
                logger.debug("Using authenticated client from corrections system")
            self.lists_mgr.enable_cloud_sync(user_uuid, supabase_client=supabase_client)

            # Get preview of what's in cloud vs local
            cloud_preview = self.lists_mgr.get_cloud_lists_preview()
            logger.info(f"Cloud preview returned: {cloud_preview}")
            local_lists = self.lists_mgr.get_local_lists_summary()

            logger.debug(f"Cloud preview result: success={cloud_preview.get('success')}, "
                        f"lists_count={len(cloud_preview.get('lists', []))}, "
                        f"error={cloud_preview.get('error')}")

            cloud_lists = cloud_preview.get('lists', []) if cloud_preview.get('success') else []

            # If preview failed, show error but still allow upload
            cloud_error = None
            if not cloud_preview.get('success'):
                cloud_error = cloud_preview.get('error', 'Unknown error')
                logger.warning(f"Cloud preview failed: {cloud_error}")

            # If both are empty (and no error), nothing to sync
            if not cloud_lists and not local_lists and not cloud_error:
                logger.info("Cloud sync: No lists to sync (both empty)")
                return

            # Check if already in sync (same list names with cloud_ids set)
            if not cloud_error and cloud_lists and local_lists:
                # Filter out "Recently Viewed" - it's local-only and not synced
                local_names = {lst.get('name') for lst in local_lists}
                cloud_names = {lst.get('name') for lst in cloud_lists
                              if lst.get('name') != 'Recently Viewed'}
                # Check if local lists have cloud_ids (already synced)
                local_with_cloud_ids = sum(1 for lst in self.lists_mgr.data.get('lists', {}).values()
                                           if lst.get('cloud_id'))
                # Consider synced if: same user lists AND has some cloud_ids
                if local_names == cloud_names and local_with_cloud_ids > 0:
                    logger.info("Cloud sync: Already in sync, skipping dialog")
                    # Don't sync - lists are already in sync
                    return

            # Show sync dialog
            self._show_lists_sync_dialog(local_lists, cloud_lists, cloud_error)

        except Exception as e:
            logger.warning(f"Cloud sync dialog error: {e}")
            import traceback
            traceback.print_exc()

    def _show_lists_sync_dialog(self, local_lists, cloud_lists, cloud_error=None):
        """Show dialog to let user choose how to sync lists."""
        dialog = QDialog(self)
        dialog.setWindowTitle(tr("Sync Your Lists"))
        dialog.setMinimumWidth(500)
        layout = QVBoxLayout(dialog)

        # Header
        if cloud_error:
            header = QLabel(tr("Could not load cloud lists. You can still upload your local lists."))
        elif cloud_lists:
            header = QLabel(tr("Your account has lists in the cloud. How would you like to sync?"))
        else:
            header = QLabel(tr("Upload your local lists to the cloud?"))
        header.setWordWrap(True)
        header.setStyleSheet("font-size: 14px; margin-bottom: 10px;")
        layout.addWidget(header)

        # Show error if any
        if cloud_error:
            error_label = QLabel(f"Error: {cloud_error}")
            error_label.setStyleSheet("color: #f44336; font-size: 12px; margin-bottom: 10px;")
            error_label.setWordWrap(True)
            layout.addWidget(error_label)

        # Two columns: Local vs Cloud
        columns = QHBoxLayout()

        # Local lists column
        local_group = QGroupBox(tr("Local Lists (this device)"))
        local_layout = QVBoxLayout(local_group)
        if local_lists:
            for lst in local_lists[:10]:  # Show max 10
                item_label = QLabel(f"• {lst['name']} ({lst['item_count']} items)")
                local_layout.addWidget(item_label)
            if len(local_lists) > 10:
                local_layout.addWidget(QLabel(f"... and {len(local_lists) - 10} more"))
        else:
            local_layout.addWidget(QLabel(tr("No local lists")))
        local_layout.addStretch()
        columns.addWidget(local_group)

        # Cloud lists column
        cloud_group = QGroupBox(tr("Cloud Lists (your account)"))
        cloud_layout = QVBoxLayout(cloud_group)
        if cloud_lists:
            for lst in cloud_lists[:10]:  # Show max 10
                item_label = QLabel(f"• {lst['name']} ({lst['item_count']} items)")
                cloud_layout.addWidget(item_label)
            if len(cloud_lists) > 10:
                cloud_layout.addWidget(QLabel(f"... and {len(cloud_lists) - 10} more"))
        else:
            cloud_layout.addWidget(QLabel(tr("No cloud lists")))
        cloud_layout.addStretch()
        columns.addWidget(cloud_group)

        layout.addLayout(columns)

        # Action buttons
        btn_layout = QHBoxLayout()

        # Download from cloud button
        if cloud_lists:
            download_btn = QPushButton(tr("Download from Cloud"))
            download_btn.setToolTip(tr("Add cloud lists to this device (keeps both)"))
            download_btn.setStyleSheet("background-color: #4CAF50; color: white; padding: 8px 16px;")
            download_btn.clicked.connect(lambda: self._do_sync_action(dialog, 'download'))
            btn_layout.addWidget(download_btn)

        # Upload to cloud button
        if local_lists:
            upload_btn = QPushButton(tr("Upload to Cloud"))
            upload_btn.setToolTip(tr("Push local lists to your account"))
            upload_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 8px 16px;")
            upload_btn.clicked.connect(lambda: self._do_sync_action(dialog, 'upload'))
            btn_layout.addWidget(upload_btn)

        # Merge both (if both have lists)
        if cloud_lists and local_lists:
            merge_btn = QPushButton(tr("Merge Both"))
            merge_btn.setToolTip(tr("Download cloud AND upload local (combine everything)"))
            merge_btn.setStyleSheet("background-color: #9C27B0; color: white; padding: 8px 16px;")
            merge_btn.clicked.connect(lambda: self._do_sync_action(dialog, 'merge'))
            btn_layout.addWidget(merge_btn)

        # Skip button
        skip_btn = QPushButton(tr("Skip"))
        skip_btn.setToolTip(tr("Don't sync now - you can sync later from Settings"))
        skip_btn.clicked.connect(dialog.reject)
        btn_layout.addWidget(skip_btn)

        layout.addLayout(btn_layout)

        dialog.exec()

    def _do_sync_action(self, dialog, action):
        """Execute the chosen sync action."""
        dialog.accept()

        progress = QProgressDialog(tr("Syncing lists..."), None, 0, 0, self)
        progress.setWindowTitle(tr("Sync"))
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.show()
        QApplication.processEvents()

        try:
            if action == 'download':
                # Download cloud lists to local (merge)
                result = self.lists_mgr.sync_from_cloud()
                if result.get('success'):
                    added = result.get('lists_added', 0)
                    items = result.get('items_added', 0)
                    QMessageBox.information(
                        self, tr("Sync Complete"),
                        tr("Downloaded {lists} lists and {items} items from cloud.").format(
                            lists=added, items=items
                        )
                    )
                else:
                    QMessageBox.warning(self, tr("Sync Error"), result.get('error', 'Unknown error'))

            elif action == 'upload':
                # Upload local lists to cloud
                result = self.lists_mgr.sync_to_cloud()
                if result.get('success'):
                    pushed = result.get('lists_pushed', 0)
                    items = result.get('items_pushed', 0)
                    QMessageBox.information(
                        self, tr("Sync Complete"),
                        tr("Uploaded {lists} lists and {items} items to cloud.").format(
                            lists=pushed, items=items
                        )
                    )
                else:
                    QMessageBox.warning(self, tr("Sync Error"), result.get('error', 'Unknown error'))

            elif action == 'merge':
                # Both directions
                download_result = self.lists_mgr.sync_from_cloud()
                upload_result = self.lists_mgr.sync_to_cloud()

                if download_result.get('success') and upload_result.get('success'):
                    QMessageBox.information(
                        self, tr("Sync Complete"),
                        tr("Lists merged successfully! Downloaded {dl} lists, uploaded {ul} lists.").format(
                            dl=download_result.get('lists_added', 0),
                            ul=upload_result.get('lists_pushed', 0)
                        )
                    )
                else:
                    errors = []
                    if not download_result.get('success'):
                        errors.append(f"Download: {download_result.get('error')}")
                    if not upload_result.get('success'):
                        errors.append(f"Upload: {upload_result.get('error')}")
                    QMessageBox.warning(self, tr("Sync Error"), "\n".join(errors))

            # Refresh the lists UI if it exists
            if hasattr(self, 'lists_tree'):
                self.lists_refresh_all()

        except Exception as e:
            QMessageBox.critical(self, tr("Sync Error"), str(e))
        finally:
            progress.close()

    def _disable_lists_cloud_sync(self):
        """Disable cloud sync on logout."""
        try:
            # Skip sync on logout if recently synced (auto-sync handles it)
            import time
            if hasattr(self.lists_mgr, '_last_sync') and self.lists_mgr._last_sync:
                if time.time() - self.lists_mgr._last_sync < 60:  # Synced in last minute
                    logger.debug("Skipping logout sync - recently synced")
                    self.lists_mgr.disable_cloud_sync()
                    return

            # Quick sync with 10-second timeout
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(self.lists_mgr.sync_to_cloud)
                try:
                    future.result(timeout=10)
                except concurrent.futures.TimeoutError:
                    logger.debug("Logout sync timed out after 10s - continuing")

            self.lists_mgr.disable_cloud_sync()
        except Exception as e:
            logger.debug(f"Cloud sync disable: {e}")

    def _show_discoveries_dialog(self):
        dialog = DiscoveriesDialog(
            self, self.corrections_client,
            on_view_result=lambda s: self._open_document_result_dialog(shelfmark=s),
            on_browse=lambda s: self._browse_document_by_shelfmark(s)
        )
        dialog.exec()

    def _show_create_discovery_dialog(self):
        if not self.corrections_client.is_logged_in():
            QMessageBox.warning(self, tr("Login Required"), tr("Please login to share a discovery."))
            return
        # Get current document context if in browse tab
        doc_id = self.current_browse_sid
        shelfmark = None
        page_number = None
        if doc_id and self.meta_mgr:
            try:
                shelfmark, _ = self.meta_mgr.get_meta_for_id(doc_id)
                page_number = self.current_browse_p
            except:
                pass
        dialog = CreateDiscoveryDialog(
            self, self.corrections_client,
            document_id=doc_id, shelfmark=shelfmark, page_number=page_number,
            lists_mgr=self.lists_mgr,
            shelf_completer=getattr(self, 'shelf_completer', None),
            meta_mgr=self.meta_mgr
        )
        if dialog.exec():
            # Refresh discoveries panel after successful creation
            self._refresh_discoveries_panel(use_cache_first=False)

    def _show_all_corrections_dialog(self):
        dialog = AllCorrectionsDialog(self, self.corrections_client)
        dialog.exec()

    def _show_my_corrections_dialog(self):
        if not self.corrections_client.is_logged_in():
            QMessageBox.warning(self, tr("Login Required"), tr("Please login to view your corrections."))
            return
        dialog = MyCorrectionsDialog(self, self.corrections_client)
        dialog.exec()

    def _show_my_comments_dialog(self):
        if not self.corrections_client.is_logged_in():
            QMessageBox.warning(self, tr("Login Required"), tr("Please login to view your comments."))
            return
        dialog = MyCommentsDialog(
            self, self.corrections_client,
            on_view_result=lambda s: self._open_document_result_dialog(shelfmark=s),
            on_browse=lambda s: self._browse_document_by_shelfmark(s)
        )
        dialog.exec()

    # === Browse Tab Community Button Handlers ===

    def _browse_toggle_edit_mode(self):
        """Toggle edit mode for inline corrections."""
        if not self.corrections_client.is_logged_in():
            QMessageBox.warning(self, tr("Login Required"), tr("Please login to make corrections."))
            return
        if not self.current_browse_sid:
            QMessageBox.warning(self, tr("No Document"), tr("Please load a document first."))
            return

        if self.browse_edit_mode:
            # Exit edit mode
            self._browse_cancel_edit()
        else:
            # Enter edit mode - reset draft tracking
            self.browse_edit_mode = True
            self._browse_draft_correction_id = None
            self.browse_original_edit_text = self.browse_text.toPlainText()
            self.browse_text.setReadOnly(False)
            # Set initial orange border (no changes yet)
            palette = self.palette()
            is_dark = palette.color(QPalette.ColorRole.Window).lightness() < 128
            edit_bg = "#3d3522" if is_dark else "#fffacd"
            self.browse_text.setStyleSheet(f"background-color: {edit_bg}; border: 2px solid #f39c12;")
            self.browse_text.textChanged.connect(self._browse_on_text_changed)
            self.browse_edit_bar.show()
            self.browse_edit_status.setText("")
            self.btn_b_save_draft.setEnabled(False)
            self.btn_b_save_correction.setEnabled(False)
            self.btn_b_edit.setText(tr("❌ Exit Edit"))
            self.btn_b_edit.setStyleSheet("background-color: #e74c3c; color: white;")

    def _browse_on_text_changed(self):
        """Handle text changes in edit mode."""
        if not self.browse_edit_mode:
            return
        current_text = self.browse_text.toPlainText()
        original = getattr(self, 'browse_original_edit_text', self.browse_original_text)
        has_changes = current_text != original
        self.btn_b_save_draft.setEnabled(has_changes)
        # Enable submit if has changes OR has saved draft
        draft_id = getattr(self, '_browse_draft_correction_id', None)
        self.btn_b_save_correction.setEnabled(has_changes or draft_id is not None)

        # Get palette-aware background color
        palette = self.palette()
        is_dark = palette.color(QPalette.ColorRole.Window).lightness() < 128
        edit_bg = "#3d3522" if is_dark else "#fffacd"

        if has_changes:
            self.browse_edit_status.setText(f"● {tr('Unsaved changes')}")
            self.browse_edit_status.setStyleSheet("color: #e67e22;")
            # Orange border for unsaved changes
            self.browse_text.setStyleSheet(f"background-color: {edit_bg}; border: 2px solid #f39c12;")
        elif draft_id:
            self.browse_edit_status.setText(f"✓ {tr('Saved')}")
            self.browse_edit_status.setStyleSheet("color: #27ae60; font-weight: bold;")
            # Green border for saved draft
            self.browse_text.setStyleSheet(f"background-color: {edit_bg}; border: 2px solid #27ae60;")
        else:
            self.browse_edit_status.setText("")
            # Orange border (default edit mode)
            self.browse_text.setStyleSheet(f"background-color: {edit_bg}; border: 2px solid #f39c12;")

    def _browse_cancel_edit(self):
        """Cancel edit mode and restore original text."""
        self.browse_edit_mode = False
        self._browse_draft_correction_id = None
        try:
            self.browse_text.textChanged.disconnect(self._browse_on_text_changed)
        except:
            pass
        # Restore to original text before editing
        original = getattr(self, 'browse_original_edit_text', self.browse_original_text)
        self.browse_text.setPlainText(original)
        self.browse_text.setReadOnly(True)
        self.browse_text.setStyleSheet("")
        self.browse_edit_bar.hide()
        self.browse_edit_status.setText("")
        self.btn_b_edit.setText(tr("✏️ Edit"))
        self.btn_b_edit.setStyleSheet("")

    def _browse_exit_edit_mode(self):
        """Exit edit mode without restoring text (after successful submit)."""
        self.browse_edit_mode = False
        self._browse_draft_correction_id = None
        try:
            self.browse_text.textChanged.disconnect(self._browse_on_text_changed)
        except:
            pass
        self.browse_text.setReadOnly(True)
        self.browse_text.setStyleSheet("")
        self.browse_edit_bar.hide()
        self.browse_edit_status.setText("")
        self.btn_b_edit.setText(tr("✏️ Edit"))
        self.btn_b_edit.setStyleSheet("")

    def _browse_save_correction(self, submit=False):
        """Save the inline correction.

        Args:
            submit: If True, submit for review. If False, save as draft (silent).
        """
        if not self.browse_edit_mode:
            return

        new_text = self.browse_text.toPlainText()
        original = getattr(self, 'browse_original_edit_text', self.browse_original_text)
        draft_correction_id = getattr(self, '_browse_draft_correction_id', None)
        has_changes = new_text != original

        if not has_changes and not draft_correction_id:
            if submit:
                QMessageBox.information(self, tr("No Changes"), tr("No changes were made to the text."))
            return

        # Get document context
        doc_id = self.current_browse_sid
        shelfmark = None
        if self.meta_mgr:
            try:
                shelfmark, _ = self.meta_mgr.get_meta_for_id(doc_id)
            except:
                pass

        # Ask for notes/reason (only for submit)
        notes = None
        if submit:
            notes, ok = QInputDialog.getMultiLineText(
                self,
                tr("Correction Notes"),
                tr("Please provide a brief explanation for your correction (optional):"),
                ""
            )
            if not ok:
                return
            notes = notes if notes else None

        # Submit or save the correction
        try:
            # If submitting an existing draft, try submit_correction API first
            if submit and draft_correction_id:
                success, message = self.corrections_client.submit_correction(draft_correction_id, notes)
                if success or 'approved' in message.lower():
                    # Success, or already approved (which means it succeeded earlier)
                    QMessageBox.information(
                        self,
                        tr("Correction Submitted"),
                        tr("Your correction has been submitted for review. Thank you for your contribution!")
                    )
                    # Exit edit mode after submit
                    self._browse_exit_edit_mode()
                    self._browse_draft_correction_id = None
                    # Refresh versions
                    self._check_document_community_status()
                else:
                    # Submit failed, try creating a new pending correction instead
                    correction, create_msg = self.corrections_client.create_correction(
                        document_id=doc_id,
                        original_text=original if original else new_text,
                        corrected_text=new_text,
                        correction_type="text_correction",
                        page_number=self.current_browse_p,
                        notes=notes,
                        shelfmark=shelfmark,
                        system_id=doc_id,
                        status='pending'
                    )
                    if correction:
                        QMessageBox.information(
                            self,
                            tr("Correction Submitted"),
                            tr("Your correction has been submitted for review. Thank you for your contribution!")
                        )
                        self._browse_exit_edit_mode()
                        self._browse_draft_correction_id = None
                        self._check_document_community_status()
                    else:
                        QMessageBox.warning(self, tr("Error"), f"{tr('Failed to submit correction')}: {create_msg}")
            else:
                # Create new correction (draft or direct submit)
                correction, message = self.corrections_client.create_correction(
                    document_id=doc_id,
                    original_text=original if original else new_text,
                    corrected_text=new_text,
                    correction_type="text_correction",
                    page_number=self.current_browse_p,
                    notes=notes,
                    shelfmark=shelfmark,
                    system_id=doc_id,
                    status='pending' if submit else 'draft',
                    save_as_draft=not submit  # Don't auto-submit when saving as draft
                )
                if correction:
                    if submit:
                        QMessageBox.information(
                            self,
                            tr("Correction Submitted"),
                            tr("Your correction has been submitted for review. Thank you for your contribution!")
                        )
                        # Exit edit mode after submit
                        self._browse_exit_edit_mode()
                        self._browse_draft_correction_id = None
                        # Refresh versions
                        self._check_document_community_status()
                    else:
                        # Silent save - just show checkmark in status, keep editing
                        self.browse_edit_status.setText(f"✓ {tr('Saved')}")
                        self.browse_edit_status.setStyleSheet("color: #27ae60; font-weight: bold;")
                        self._browse_draft_correction_id = correction.id
                        self.browse_original_edit_text = new_text  # Update original to mark as saved
                        # Update border to green (saved)
                        palette = self.palette()
                        is_dark = palette.color(QPalette.ColorRole.Window).lightness() < 128
                        edit_bg = "#3d3522" if is_dark else "#fffacd"
                        self.browse_text.setStyleSheet(f"background-color: {edit_bg}; border: 2px solid #27ae60;")
                        # Disable save draft button since no changes
                        self.btn_b_save_draft.setEnabled(False)
                        # Enable submit button
                        self.btn_b_save_correction.setEnabled(True)
                        # Refresh versions
                        self._check_document_community_status()
                else:
                    QMessageBox.warning(self, tr("Error"), f"{tr('Failed to save correction')}: {message}")
        except Exception as e:
            QMessageBox.warning(self, tr("Error"), f"{tr('Failed to save correction')}: {str(e)}")

    def _browse_change_version(self, index):
        """Change between text versions."""
        if not hasattr(self, 'current_browse_sid') or not self.current_browse_sid:
            return
        if not hasattr(self, 'browse_version_combo'):
            return

        version_data = self.browse_version_combo.currentData()
        if not version_data:
            return
        # Skip non-selectable header items
        if version_data.get('source') == 'header':
            return

        self._browse_load_version(version_data)

    def _browse_load_version(self, version_data):
        """Load and display a specific version."""
        if not version_data:
            return

        source = version_data.get('source')
        version_id = version_data.get('version_id')
        source_id = version_data.get('source_id')

        # Check cache first
        if source in ('pgp_edition', 'pgp_translation'):
            cache_key = f"pgp_{source_id}" if source_id else source
        else:
            cache_key = f"{source}_{version_id}" if version_id else source

        if cache_key in self._browse_versions_cache:
            content = self._browse_versions_cache[cache_key]
            if source == 'pgp_translation':
                language = version_data.get('language', '')
                is_rtl = language != 'English'
                self._browse_display_pgp_text(content, is_rtl=is_rtl)
            elif source == 'pgp_edition':
                self._browse_display_pgp_text(content, is_rtl=True)
            else:
                self._browse_display_version_text(content)
            return

        if source == "pgp_edition":
            # PGP edition content is stored directly in version_data
            content = version_data.get('content', '')
            if content:
                if source_id:
                    self._browse_versions_cache[f"pgp_{source_id}"] = content
                self._browse_display_pgp_text(content, is_rtl=True)
        elif source == "pgp_translation":
            # PGP translation content is stored directly in version_data
            content = version_data.get('content', '')
            language = version_data.get('language', '')
            if content:
                if source_id:
                    self._browse_versions_cache[f"pgp_{source_id}"] = content
                # English translations are LTR, everything else RTL
                is_rtl = language != 'English'
                self._browse_display_pgp_text(content, is_rtl=is_rtl)
        elif source == "original":
            # Show original V0.8 text and restore RTL direction
            self.browse_text.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
            if hasattr(self, 'browse_original_page_text') and self.browse_original_page_text:
                self._browse_display_version_text(self.browse_original_page_text)
        elif source == "correction":
            # Correction content is stored directly in version_data
            content = version_data.get('corrected_text', '')
            if content:
                correction_id = version_data.get('correction_id')
                cache_key = f"correction_{correction_id}"
                self._browse_versions_cache[cache_key] = content
                # Restore RTL for corrections (Hebrew text)
                self.browse_text.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
                self._browse_display_version_text(content)
            else:
                if hasattr(self, 'browse_original_page_text'):
                    self._browse_display_version_text(self.browse_original_page_text)
        elif version_id:
            # Quick server availability check (500ms timeout) to prevent UI freeze
            if not self.corrections_client.is_server_available():
                return
            # Fetch version content from API
            try:
                ver_data = self.corrections_client.get_version_content(version_id)
                content = ver_data.get('content', '')
                if content:
                    self._browse_versions_cache[cache_key] = content
                    self._browse_display_version_text(content)
                else:
                    # Fall back to original if no content
                    if hasattr(self, 'browse_original_page_text'):
                        self._browse_display_version_text(self.browse_original_page_text)
            except Exception as e:
                logger.debug("Error loading version content: %s", e)
                if hasattr(self, 'browse_original_page_text'):
                    self._browse_display_version_text(self.browse_original_page_text)

    def _browse_display_version_text(self, text):
        """Display version text in the browse text area."""
        if not text:
            return
        # Apply RTL formatting like browse_render_page does
        browse_html_text = text.replace('\n', '<br>')
        self.browse_text.setHtml(f"<div dir='rtl'>{browse_html_text}</div>")
        apply_find_highlight(self.browse_text, self.browse_find_input.text().strip())

    # ── PGP Version Selector Helpers (shared by Browse tab and ResultDialog) ──

    def _populate_pgp_combo(self, combo, sources, pgp_doc):
        """Build combo items with PGP editions and translations grouped.

        Matches the web app grouping pattern:
        PGP Editions (header) -> edition items -> separator ->
        Translations (header) -> translation items -> separator -> V0.8

        Args:
            combo: QComboBox to populate
            sources: list of source dicts from PGPSourceWorker
            pgp_doc: dict of PGP document metadata

        Returns:
            True if any PGP editions/translations were added, False if only V0.8.
        """
        editions = [s for s in sources
                     if 'Edition' in (s.get('doc_relation') or '') and s.get('content')]
        # Exclude sources already classified as editions (compound doc_relation like
        # "Edition ; Translation ; Discussion" should only appear once, as edition)
        edition_ids = {id(s) for s in editions}
        translations = [s for s in sources
                         if 'Translation' in (s.get('doc_relation') or '')
                         and s.get('content')
                         and id(s) not in edition_ids]

        if not editions and not translations:
            return False

        combo.blockSignals(True)
        combo.clear()

        # === PGP Editions Group ===
        if editions:
            combo.addItem("-- PGP Editions --", {"source": "header"})
            combo.model().item(combo.count() - 1).setEnabled(False)

            for edition in editions:
                scholar = edition.get('source_scholar', 'Unknown')
                label = f"  {scholar}"
                combo.addItem(label, {
                    "source": "pgp_edition",
                    "content": edition.get('content', ''),
                    "scholar": scholar,
                    "pgpid": edition.get('pgpid'),
                    "source_id": edition.get('id')
                })

        # === Translations Group ===
        if translations:
            combo.addItem("─────────────", {"source": "header"})
            combo.model().item(combo.count() - 1).setEnabled(False)
            combo.addItem("-- Translations --", {"source": "header"})
            combo.model().item(combo.count() - 1).setEnabled(False)

            # Group translations by language (Hebrew first, English second, others last)
            # Matches web app grouping: web/components/version_selector.py:256-264
            hebrew_trans = [t for t in translations if t.get('language') == 'Hebrew']
            english_trans = [t for t in translations if t.get('language') == 'English']
            other_trans = [t for t in translations if t.get('language') not in ('Hebrew', 'English')]

            for trans_group in [hebrew_trans, english_trans, other_trans]:
                for trans in trans_group:
                    scholar = trans.get('source_scholar', 'Unknown')
                    language = trans.get('language', '')
                    label = f"  {language} - {scholar}" if language else f"  {scholar}"
                    combo.addItem(label, {
                        "source": "pgp_translation",
                        "content": trans.get('content', ''),
                        "scholar": scholar,
                        "language": language,
                        "pgpid": trans.get('pgpid'),
                        "source_id": trans.get('id')
                    })

        # === Visual divider before HTR ===
        combo.addItem("─────────────", {"source": "header"})
        combo.model().item(combo.count() - 1).setEnabled(False)

        # === HTR V0.8 (always present) ===
        combo.addItem("V0.8", {"source": "original"})

        combo.blockSignals(False)
        return True

    def _auto_select_pgp_edition(self, combo):
        """Find the first PGP edition item and set it as current.

        Returns:
            The item data dict if found, None otherwise.
        """
        for i in range(combo.count()):
            data = combo.itemData(i)
            if data and data.get('source') == 'pgp_edition':
                combo.setCurrentIndex(i)
                return data
        return None

    def _check_document_community_status(self):
        """Check if document has comments and load available versions."""
        if not self.current_browse_sid or not self.corrections_client:
            return

        doc_id = self.current_browse_sid
        page_num = self.current_browse_p or 1

        # Store original text for this document
        original_text = self.browse_text.toPlainText()
        self.browse_original_page_text = original_text

        # Reset version cache, PGP state, and combo for new document
        self._browse_versions_cache = {'original': original_text}
        self._browse_pgp_sources = []
        self._browse_pgp_doc = {}
        self._browse_enriched_html = ''
        self.browse_version_combo.blockSignals(True)
        self.browse_version_combo.clear()
        self.browse_version_combo.addItem("V0.8", {"source": "original"})
        self.browse_version_combo.blockSignals(False)

        # Force fresh server availability check (500ms timeout) to prevent UI freeze
        if not self.corrections_client.is_server_available(force_check=True):
            # Server is down - skip API calls, hide community UI elements
            self.btn_b_view_comments.setVisible(False)
            return

        # Check for comments
        try:
            comments = self.corrections_client.get_comments_for_document(doc_id, page_size=1)
            if comments and len(comments) > 0:
                self.btn_b_view_comments.setVisible(True)
                self.btn_b_view_comments.setEnabled(True)
            else:
                self.btn_b_view_comments.setVisible(False)
        except:
            self.btn_b_view_comments.setVisible(False)

        # Fetch versions from API
        try:
            versions_data = self.corrections_client.get_page_versions(doc_id, page_num)
            all_versions = versions_data.get('all_versions', [])
            current_default = versions_data.get('current_default')

            # Add V0.7 if available
            v07_versions = [v for v in all_versions if v.get('source') == 'V0.7']
            for ver in v07_versions:
                ver_id = ver.get('id')
                is_default = ver.get('is_current_default', False)
                label = 'V0.7'
                if is_default:
                    label += f" ({tr('Default')})"
                self.browse_version_combo.addItem(label, {
                    "source": "V0.7",
                    "version_id": ver_id,
                    "is_default": is_default
                })

            # Add user versions
            users_with_versions = set()  # Track users who have versions (to avoid duplicate corrections)
            user_versions = [v for v in all_versions if v.get('source') == 'user']
            for ver in user_versions:
                ver_id = ver.get('id')
                user_name = ver.get('user_name') or 'User'
                created_at = ver.get('created_at', '')[:10] if ver.get('created_at') else ''
                is_default = ver.get('is_current_default', False)

                label = f"{tr('by')} {user_name}"
                if created_at:
                    label += f" ({created_at})"
                if is_default:
                    label += " ✓"

                self.browse_version_combo.addItem(label, {
                    "source": "user",
                    "version_id": ver_id,
                    "user_name": user_name,
                    "is_default": is_default
                })
                users_with_versions.add(user_name)

            # Also fetch corrections (including drafts) for current user
            try:
                corrections = self.corrections_client.get_corrections_for_document(doc_id, include_drafts=True)
                # Filter corrections by page number
                page_corrections = [c for c in corrections if c.page_number == page_num or c.page_number is None]

                # Group by user, keep latest per user
                corrections_by_user = {}
                for corr in page_corrections:
                    user_key = corr.author_username or f"user_{corr.author_id}"
                    if user_key not in corrections_by_user:
                        corrections_by_user[user_key] = corr
                    else:
                        existing = corrections_by_user[user_key]
                        if (corr.created_at or '') > (existing.created_at or ''):
                            corrections_by_user[user_key] = corr

                # Determine user permissions for viewing corrections
                current_username = self.corrections_client.current_user.username if self.corrections_client.current_user else None
                is_reviewer_or_admin = self.corrections_client.current_user and self.corrections_client.current_user.role in ('reviewer', 'editor', 'admin')

                for corr in corrections_by_user.values():
                    user_name = corr.author_username or 'User'
                    status = corr.status

                    # Skip if user already has a version entry (avoid duplicates)
                    if user_name in users_with_versions:
                        continue

                    # Filter based on status and user permissions:
                    # - Authors can see their own corrections (any status)
                    # - Reviewers/admins can see all corrections
                    # - Regular users can only see approved corrections from others
                    is_own_correction = current_username and user_name == current_username

                    if status == 'rejected':
                        if not is_own_correction and not is_reviewer_or_admin:
                            continue
                    elif status in ('draft', 'pending'):
                        if not is_own_correction and not is_reviewer_or_admin:
                            continue

                    created_at = corr.created_at[:10] if corr.created_at else ''

                    # Status indicators and label
                    if status == 'draft':
                        label = f"📝 {tr('Draft')}"
                    elif status == 'pending':
                        label = f"⏳ {tr('Pending')} - {user_name}"
                    elif status == 'approved':
                        label = f"✅ {tr('by')} {user_name}"
                        if created_at:
                            label += f" ({created_at})"
                    elif status == 'rejected':
                        label = f"❌ {tr('Rejected')} - {user_name}"
                    else:
                        label = f"{tr('by')} {user_name}"
                        if created_at:
                            label += f" ({created_at})"

                    self.browse_version_combo.addItem(label, {
                        "source": "correction",
                        "correction_id": corr.id,
                        "user_name": user_name,
                        "status": status,
                        "corrected_text": corr.corrected_text
                    })
            except Exception as e:
                logger.debug("Error fetching corrections for browse: %s", e)

            # Enable combo if we have more than just V0.8
            if self.browse_version_combo.count() > 1:
                self.browse_version_combo.setEnabled(True)

                # If there's a current default that's not V0.8, select it
                if current_default and current_default.get('source') != 'V0.8':
                    default_id = current_default.get('id')
                    for i in range(self.browse_version_combo.count()):
                        data = self.browse_version_combo.itemData(i)
                        if data and data.get('version_id') == default_id:
                            self.browse_version_combo.blockSignals(True)
                            self.browse_version_combo.setCurrentIndex(i)
                            self.browse_version_combo.blockSignals(False)
                            # Load and display this version
                            self._browse_load_version(data)
                            break
            else:
                self.browse_version_combo.setEnabled(False)

        except Exception as e:
            logger.debug("Error fetching versions: %s", e)
            self.browse_version_combo.setEnabled(False)

    def _browse_add_comment(self):
        """Open comment dialog for current document."""
        if not self.corrections_client.is_logged_in():
            QMessageBox.warning(self, tr("Login Required"), tr("Please login to add a comment."))
            return
        if not self.current_browse_sid:
            QMessageBox.warning(self, tr("No Document"), tr("Please load a document first."))
            return

        doc_id = self.current_browse_sid
        shelfmark = None
        if self.meta_mgr:
            try:
                shelfmark, _ = self.meta_mgr.get_meta_for_id(doc_id)
            except:
                pass

        dialog = CommentDialog(
            self, self.corrections_client,
            document_id=doc_id,
            shelfmark=shelfmark,
            page_number=self.current_browse_p
        )
        dialog.exec()

    def _browse_view_corrections(self):
        """View corrections for current document."""
        if not self.current_browse_sid:
            QMessageBox.warning(self, tr("No Document"), tr("Please load a document first."))
            return

        doc_id = self.current_browse_sid
        shelfmark = None
        if self.meta_mgr:
            try:
                shelfmark, _ = self.meta_mgr.get_meta_for_id(doc_id)
            except:
                pass

        dialog = CorrectionsViewerDialog(
            self, self.corrections_client,
            document_id=doc_id,
            shelfmark=shelfmark,
            on_view_result=lambda s: self._open_document_result_dialog(shelfmark=s),
            on_browse=lambda s: self._browse_document_by_shelfmark(s)
        )
        dialog.exec()

    def _browse_view_comments(self):
        """View comments for current document."""
        if not self.current_browse_sid:
            QMessageBox.warning(self, tr("No Document"), tr("Please load a document first."))
            return

        doc_id = self.current_browse_sid
        shelfmark = None
        if self.meta_mgr:
            try:
                shelfmark, _ = self.meta_mgr.get_meta_for_id(doc_id)
            except:
                pass

        dialog = CommentsViewerDialog(
            self, self.corrections_client,
            document_id=doc_id,
            shelfmark=shelfmark
        )
        dialog.exec()

    def _browse_view_joins(self):
        """View joined fragments for current document."""
        if not self.current_browse_sid:
            QMessageBox.warning(self, tr("No Document"), tr("Please load a document first."))
            return

        doc_id = self.current_browse_sid
        shelfmark = None
        if self.meta_mgr:
            try:
                shelfmark, _ = self.meta_mgr.get_meta_for_id(doc_id)
            except:
                pass

        if not shelfmark:
            QMessageBox.warning(self, tr("No Shelfmark"), tr("Could not determine shelfmark for this document."))
            return

        def browse_shelfmark(target_shelfmark):
            """Navigate to a shelfmark in the browse tab."""
            self.browse_shelf_input.setText(target_shelfmark)
            self._set_last_browse_field("shelf")
            self.browse_load()

        dialog = JoinsDialog(
            self, self.corrections_client,
            document_id=doc_id,
            shelfmark=shelfmark,
            on_browse=browse_shelfmark,
            shelf_model=getattr(self, 'shelf_model', None),
            joins_mgr=getattr(self, 'joins_mgr', None),
            shelf_completer=getattr(self, 'shelf_completer', None),
            lists_mgr=getattr(self, 'lists_mgr', None),
            meta_mgr=getattr(self, 'meta_mgr', None)
        )
        dialog.exec()

    def _update_joins_dropdown(self):
        """Update the joins dropdown menu with connected fragments."""
        self.joins_menu.clear()

        if not self.current_browse_sid:
            return

        # Use document_id (sys_id) for lookup - this is the reliable key
        document_id = self.current_browse_sid

        shelfmark = None
        if self.meta_mgr:
            try:
                shelfmark, _ = self.meta_mgr.get_meta_for_id(self.current_browse_sid)
            except:
                pass

        if not shelfmark:
            return

        # Get joins from JoinsManager using document_id (offline-first)
        connected = None
        if self.joins_mgr:
            connected = self.joins_mgr.get_connected_fragments_by_id(document_id)

        # Fall back to shelfmark-based lookup if no results with document_id
        if (not connected or connected.get('total_fragments', 0) <= 1) and self.joins_mgr:
            connected = self.joins_mgr.get_connected_fragments(shelfmark)

        if not connected or connected.get('total_fragments', 0) <= 1:
            # Check PGP multi-fragment joins as fallback (use cached data to avoid sync Supabase calls on UI thread)
            pgp_doc = getattr(self, '_browse_pgp_doc', {})
            if pgp_doc:
                try:
                    from shared.document_service import get_fragments_for_document
                    pgp_frags = get_fragments_for_document(pgp_doc.get('pgpid'))
                    # Filter to multi-fragment documents (>1 fragment)
                    if pgp_frags and len(pgp_frags) > 1:
                        self.btn_b_joins.setStyleSheet("background-color: #27ae60; color: white; border-radius: 4px;")
                        header_action = self.joins_menu.addAction(
                            tr("{} connected fragments").format(len(pgp_frags)) + " [PGP]"
                        )
                        header_action.setEnabled(False)
                        self.joins_menu.addSeparator()
                        for frag in pgp_frags:
                            frag_sid = frag.get('sys_id', '')
                            frag_shelf = frag.get('shelfmark', frag_sid)
                            if frag_sid == self.current_browse_sid:
                                continue  # Skip current manuscript
                            action = self.joins_menu.addAction(f"[PGP] {frag_shelf}")
                            action.triggered.connect(lambda checked, sh=frag_shelf: self._navigate_to_joined_fragment(sh))
                        self.joins_menu.addSeparator()
                        open_rd = self.joins_menu.addAction(tr("Open in Reading Desk"))
                        open_rd.triggered.connect(self._browse_open_pgp_joins_in_reading_desk)
                        return
                except Exception as e:
                    logger.debug("PGP joins dropdown fallback error: %s", e)

            # Check FJMS scholarly joins as additional fallback
            try:
                from shared.fjms_service import get_fjms_service
                fjms_svc = get_fjms_service()
                if fjms_svc.is_available():
                    fjms_members = fjms_svc.get_join_group(document_id)
                    if fjms_members:
                        # Filter to valid, non-self members
                        valid_members = []
                        for member in fjms_members:
                            alma_id = member.get('alma_id', '')
                            if not alma_id or alma_id == document_id:
                                continue
                            shelf = None
                            if self.meta_mgr:
                                try:
                                    shelf, _ = self.meta_mgr.get_meta_for_id(alma_id)
                                except Exception:
                                    pass
                            if not shelf or shelf == 'Unknown':
                                continue
                            valid_members.append((shelf, member))
                        if valid_members:
                            self.btn_b_joins.setStyleSheet("background-color: #27ae60; color: white; border-radius: 4px;")
                            header_action = self.joins_menu.addAction(
                                tr("{} connected fragments").format(len(valid_members) + 1) + " [FJMS]"
                            )
                            header_action.setEnabled(False)
                            self.joins_menu.addSeparator()
                            for shelf, member in valid_members:
                                join_type = ', '.join(member.get('join_types', []))
                                scholar_name = ', '.join(member.get('scholar_names', []))
                                label = f"[FJMS] {shelf}"
                                if join_type:
                                    label += f" \u2014 {join_type}"
                                if scholar_name:
                                    label += f" ({scholar_name})"
                                action = self.joins_menu.addAction(label)
                                action.triggered.connect(lambda checked, sh=shelf: self._navigate_to_joined_fragment(sh))
                            return
            except Exception as e:
                logger.debug("FJMS joins dropdown fallback error: %s", e)

            # No user, PGP, or FJMS joins
            action = self.joins_menu.addAction(tr("No joined fragments"))
            action.setEnabled(False)
            self.btn_b_joins.setStyleSheet("background-color: #95a5a6; color: white; border-radius: 4px;")
            return

        # Has joins - update button style and add fragment actions
        self.btn_b_joins.setStyleSheet("background-color: #27ae60; color: white; border-radius: 4px;")

        # Add header
        header_action = self.joins_menu.addAction(
            tr("{} connected fragments").format(connected.get('total_fragments', 0))
        )
        header_action.setEnabled(False)
        self.joins_menu.addSeparator()

        # Build set of directly connected fragments
        joins_list = connected.get('joins', [])
        fragment_details = connected.get('fragment_details', [])
        direct_fragments = set()
        for join in joins_list:
            frag_a = join.get('fragment_a', '') if isinstance(join, dict) else getattr(join, 'fragment_a', '')
            frag_b = join.get('fragment_b', '') if isinstance(join, dict) else getattr(join, 'fragment_b', '')
            if frag_a.upper() == shelfmark.upper():
                direct_fragments.add(frag_b.upper())
            elif frag_b.upper() == shelfmark.upper():
                direct_fragments.add(frag_a.upper())

        # Build map of shelfmark -> document_id from fragment_details for title lookup
        shelfmark_to_docid = {}
        for fd in fragment_details:
            shelf = fd.get('shelfmark', '') if isinstance(fd, dict) else getattr(fd, 'shelfmark', '')
            doc_id = fd.get('document_id') if isinstance(fd, dict) else getattr(fd, 'document_id', None)
            if shelf and doc_id:
                shelfmark_to_docid[shelf.upper()] = doc_id

        # Add each connected fragment
        for frag in connected.get('fragments', []):
            is_direct = frag.upper() in direct_fragments

            # Get title for display
            title_preview = ""
            frag_doc_id = shelfmark_to_docid.get(frag.upper())

            # Fallback: use _shelf_to_sys map from csv_bank
            if not frag_doc_id and self._shelf_to_sys:
                norm = self._normalize_shelfmark(frag) if hasattr(self, '_normalize_shelfmark') else None
                if norm:
                    frag_doc_id = self._shelf_to_sys.get(norm)

            if frag_doc_id and self.meta_mgr:
                try:
                    _, title = self.meta_mgr.get_meta_for_id(frag_doc_id)
                    if title:
                        words = title.split()[:4]
                        title_preview = ' '.join(words)
                        if len(title.split()) > 4:
                            title_preview += "..."
                except:
                    pass

            if frag.upper() == shelfmark.upper():
                # Current fragment - mark but don't make clickable
                label = f"• {frag}"
                if title_preview:
                    label += f" - {title_preview}"
                label += f" ({tr('current')})"
                action = self.joins_menu.addAction(label)
                action.setEnabled(False)
            else:
                label = f"→ {frag}"
                if title_preview:
                    label += f" - {title_preview}"
                if is_direct:
                    label += f" ({tr('direct')})"
                action = self.joins_menu.addAction(label)
                action.setData(frag)
                action.triggered.connect(lambda checked, f=frag: self._navigate_to_joined_fragment(f))

        # Merge FJMS scholarly joins into existing dropdown
        try:
            from shared.fjms_service import get_fjms_service
            fjms_svc = get_fjms_service()
            if fjms_svc.is_available():
                fjms_members = fjms_svc.get_join_group(document_id)
                fjms_valid = []
                existing_upper = set(f.upper() for f in connected.get('fragments', []))
                for member in fjms_members:
                    alma_id = member.get('alma_id', '')
                    if not alma_id or alma_id == document_id:
                        continue
                    shelf = None
                    if self.meta_mgr:
                        try:
                            shelf, _ = self.meta_mgr.get_meta_for_id(alma_id)
                        except Exception:
                            pass
                    if not shelf or shelf == 'Unknown':
                        continue
                    # Deduplicate against existing fragments
                    if shelf.upper() in existing_upper:
                        continue
                    existing_upper.add(shelf.upper())
                    fjms_valid.append((shelf, member))
                if fjms_valid:
                    self.joins_menu.addSeparator()
                    fjms_header = self.joins_menu.addAction("[FJMS Scholarly Joins]")
                    fjms_header.setEnabled(False)
                    for shelf, member in fjms_valid:
                        join_type = ', '.join(member.get('join_types', []))
                        scholar_name = ', '.join(member.get('scholar_names', []))
                        label = f"[FJMS] {shelf}"
                        if join_type:
                            label += f" \u2014 {join_type}"
                        if scholar_name:
                            label += f" ({scholar_name})"
                        action = self.joins_menu.addAction(label)
                        action.triggered.connect(lambda checked, sh=shelf: self._navigate_to_joined_fragment(sh))
        except Exception as e:
            logger.debug("FJMS joins merge error: %s", e)

        # Add separator and "View All" / "Open in Reading Desk" actions
        self.joins_menu.addSeparator()
        view_all = self.joins_menu.addAction(tr("View all joins..."))
        view_all.triggered.connect(self._browse_view_joins)
        open_rd = self.joins_menu.addAction(tr("Open in Reading Desk"))
        open_rd.triggered.connect(self._browse_open_joins_in_reading_desk)

    def _on_joins_menu_show(self):
        """Called when joins menu is about to show - trigger sync and update."""
        # Trigger a background sync to get latest joins from server
        if self.joins_mgr:
            import threading
            def sync_and_update():
                self.joins_mgr.sync_with_server()
            threading.Thread(target=sync_and_update, daemon=True).start()
        # Update menu with current data
        self._update_joins_dropdown()

    def _navigate_to_joined_fragment(self, shelfmark: str):
        """Navigate to a joined fragment in browse tab."""
        self.browse_shelf_input.setText(shelfmark)
        self._set_last_browse_field("shelf")
        self.browse_load()

    # === Search Results Context Menu ===

    def _show_results_context_menu(self, pos):
        """Show context menu for search results with community options."""
        row = self.results_table.rowAt(pos.y())
        if row < 0:
            return

        # Get the result data
        item = self.results_table.item(row, self.COL_SYS_ID)
        if not item:
            return
        res = item.data(Qt.ItemDataRole.UserRole)
        if not res:
            return

        sys_id = res.get('sys_id') or res.get('system_id', '')
        shelfmark = res.get('shelfmark', '')
        if not shelfmark and self.meta_mgr:
            try:
                shelfmark, _ = self.meta_mgr.get_meta_for_id(sys_id)
            except:
                pass

        menu = QMenu(self)

        # View action
        action_view = menu.addAction(tr("View Document"))
        action_view.triggered.connect(lambda: self._context_view_document(sys_id))

        menu.addSeparator()

        # Community actions
        action_correction = menu.addAction(tr("Submit Correction..."))
        action_correction.triggered.connect(lambda: self._context_submit_correction(sys_id, shelfmark))

        action_comment = menu.addAction(tr("Add Comment..."))
        action_comment.triggered.connect(lambda: self._context_add_comment(sys_id, shelfmark))

        menu.addSeparator()

        action_view_corrections = menu.addAction(tr("View Corrections..."))
        action_view_corrections.triggered.connect(lambda: self._context_view_corrections(sys_id, shelfmark))

        action_view_comments = menu.addAction(tr("View Comments..."))
        action_view_comments.triggered.connect(lambda: self._context_view_comments(sys_id, shelfmark))

        menu.addSeparator()

        action_discovery = menu.addAction(tr("Share Discovery..."))
        action_discovery.triggered.connect(lambda: self._context_share_discovery(sys_id, shelfmark))

        menu.exec(self.results_table.mapToGlobal(pos))

    def _context_view_document(self, sys_id):
        """Navigate to browse tab for this document."""
        self.browse_sys_input.setText(sys_id)
        self._set_last_browse_field("sys")
        self.tabs.setCurrentWidget(self.browse_tab)
        self.browse_load()

    def _context_submit_correction(self, doc_id, shelfmark):
        """Open correction dialog from context menu."""
        if not self.corrections_client.is_logged_in():
            QMessageBox.warning(self, tr("Login Required"), tr("Please login to submit a correction."))
            return
        dialog = CorrectionSubmitDialog(
            self, self.corrections_client,
            document_id=doc_id,
            shelfmark=shelfmark
        )
        dialog.exec()

    def _context_add_comment(self, doc_id, shelfmark):
        """Open comment dialog from context menu."""
        if not self.corrections_client.is_logged_in():
            QMessageBox.warning(self, tr("Login Required"), tr("Please login to add a comment."))
            return
        dialog = CommentDialog(
            self, self.corrections_client,
            document_id=doc_id,
            shelfmark=shelfmark
        )
        dialog.exec()

    def _context_view_corrections(self, doc_id, shelfmark):
        """View corrections from context menu."""
        dialog = CorrectionsViewerDialog(
            self, self.corrections_client,
            document_id=doc_id,
            shelfmark=shelfmark,
            on_view_result=lambda s: self._open_document_result_dialog(shelfmark=s),
            on_browse=lambda s: self._browse_document_by_shelfmark(s)
        )
        dialog.exec()

    def _context_view_comments(self, doc_id, shelfmark):
        """View comments from context menu."""
        dialog = CommentsViewerDialog(
            self, self.corrections_client,
            document_id=doc_id,
            shelfmark=shelfmark
        )
        dialog.exec()

    def _context_share_discovery(self, doc_id, shelfmark):
        """Share discovery from context menu."""
        if not self.corrections_client.is_logged_in():
            QMessageBox.warning(self, tr("Login Required"), tr("Please login to share a discovery."))
            return
        dialog = CreateDiscoveryDialog(
            self, self.corrections_client,
            document_id=doc_id,
            shelfmark=shelfmark,
            lists_mgr=self.lists_mgr,
            shelf_completer=getattr(self, 'shelf_completer', None),
            meta_mgr=self.meta_mgr
        )
        dialog.exec()

    def toggle_language(self):
        new_lang = 'en' if CURRENT_LANG == 'he' else 'he'
        save_language(new_lang)
        QMessageBox.information(self, tr("נדרש אתחול מחדש"), tr("אנא הפעילו מחדש את התוכנה כדי שהשינוי בשפה ייכנס לתוקף."))

    def create_search_tab(self):
        panel = QWidget(); layout = QVBoxLayout()

        # Top Container with Multi-Row Layout
        top_container = QWidget()
        top_layout = QVBoxLayout(top_container)
        top_layout.setContentsMargins(0, 0, 0, 0)

        # Row 1: Query & Search Buttons (wrapped in container for PGP Tags mode hide)
        self.search_row1_container = QWidget()
        row1 = QHBoxLayout(self.search_row1_container)
        row1.setContentsMargins(0, 0, 0, 0)
        self.query_input = QLineEdit(); self.query_input.setPlaceholderText(tr("Search terms, title or shelfmark..."))
        self.query_input.setToolTip(tr("Search Shortcuts:\n= = Exact match\n? = Variants (use buttons to select level)\n~ = Fuzzy search\n/ = Regex\n$ = Title search\n# = Shelfmark search\nR = Responsa mode\n\nExample: ?שלום"))
        self.query_input.returnPressed.connect(self.toggle_search)
        self.query_input.textChanged.connect(self._on_query_text_changed)

        self.btn_search = QPushButton(tr("Search")); self.btn_search.clicked.connect(self.toggle_search)
        self.btn_search.setStyleSheet("background-color: #27ae60; color: white; font-weight: bold; min-width: 80px;")
        self.btn_search.setEnabled(False)

        self.query_label = QLabel(tr("Query:"))
        row1.addWidget(self.query_label)
        row1.addWidget(self.query_input)
        row1.addWidget(self.btn_search)

        # Row 2: Search Parameters & Lab Mode
        row2 = QHBoxLayout()

        self.mode_combo = QComboBox()
        self.mode_combo.addItems([tr("Exact (=)"), tr("Variants (?)"), tr("Responsa (R)"), tr("Fuzzy (~)"), tr("Regex (/)"), tr("Title ($)"), tr("Shelfmark (#)"), tr("PGP Tags")])
        self.MODE_RESPONSA = 2  # Index of Responsa mode
        self.MODE_PGP_TAGS = 7  # Index of PGP Tags mode (shifted by Responsa insertion)

        # Feature discovery glow on mode combo (one-time hint)
        self._mode_glow_active = False
        cfg = load_app_config()
        if not cfg.get('hint_responsa_seen'):
            self._mode_glow_active = True
            self._mode_glow_on = True
            self.mode_combo.setToolTip(tr("Try the Responsa-project style search mode!"))
            self._mode_glow_timer = QTimer(self)
            self._mode_glow_timer.timeout.connect(self._pulse_mode_glow)
            self._mode_glow_timer.start(800)
            self._pulse_mode_glow()
            # Highlight the Responsa item inside the dropdown
            from PyQt6.QtGui import QColor
            self.mode_combo.model().item(self.MODE_RESPONSA).setBackground(QColor("#d1fae5"))
            self.mode_combo.setItemText(self.MODE_RESPONSA, "\u2728 " + tr("Responsa (R)"))
        # Tooltips
        self.mode_combo.setItemData(0, tr("Exact match"))
        self.mode_combo.setItemData(1, tr("Variant search with configurable intensity"))
        self.mode_combo.setItemData(2, tr("Responsa-Project style grammatical expansion for Hebrew search"))
        self.mode_combo.setItemData(3, tr("Fuzzy search: Levenshtein distance"))
        self.mode_combo.setItemData(4, tr("Regex: Advanced pattern matching"))
        self.mode_combo.setItemData(5, tr("Search in Title metadata"))
        self.mode_combo.setItemData(6, tr("Search in Shelfmark metadata"))
        self.mode_combo.setItemData(7, tr("Search by topic tags from the Princeton Geniza Project"))
        self.mode_combo.currentIndexChanged.connect(self._on_search_mode_changed)

        # Variant controls container (visible only in Variants mode)
        self.variant_controls_container = QWidget()
        variant_layout = QHBoxLayout(self.variant_controls_container)
        variant_layout.setContentsMargins(0, 0, 0, 0)
        variant_layout.setSpacing(4)

        # === Preset buttons (default mode) ===
        self.variant_presets_widget = QWidget()
        presets_layout = QHBoxLayout(self.variant_presets_widget)
        presets_layout.setContentsMargins(0, 0, 0, 0)
        presets_layout.setSpacing(4)

        self.btn_variant_basic = QPushButton("○ " + tr("Basic"))
        self.btn_variant_basic.setToolTip(tr("Basic variants (30 pairs)"))
        self.btn_variant_basic.setCheckable(True)
        self.btn_variant_basic.setStyleSheet("padding: 2px 6px;")
        self.btn_variant_basic.clicked.connect(lambda: self._set_variant_preset(30))

        self.btn_variant_extended = QPushButton("◐ " + tr("Extended"))
        self.btn_variant_extended.setToolTip(tr("Extended variants (70 pairs)"))
        self.btn_variant_extended.setCheckable(True)
        self.btn_variant_extended.setStyleSheet("padding: 2px 6px;")
        self.btn_variant_extended.clicked.connect(lambda: self._set_variant_preset(70))

        self.btn_variant_maximum = QPushButton("● " + tr("Maximum"))
        self.btn_variant_maximum.setToolTip(tr("Maximum variants (150 pairs) - slower"))
        self.btn_variant_maximum.setCheckable(True)
        self.btn_variant_maximum.setStyleSheet("padding: 2px 6px;")
        self.btn_variant_maximum.clicked.connect(lambda: self._set_variant_preset(150))

        # Set default selection
        self.btn_variant_extended.setChecked(True)
        self._current_variant_preset = 70

        presets_layout.addWidget(self.btn_variant_basic)
        presets_layout.addWidget(self.btn_variant_extended)
        presets_layout.addWidget(self.btn_variant_maximum)

        # === Slider (advanced mode) ===
        self.variant_slider_widget = QWidget()
        slider_layout = QHBoxLayout(self.variant_slider_widget)
        slider_layout.setContentsMargins(0, 0, 0, 0)
        slider_layout.setSpacing(4)

        self.variant_slider = QSlider(Qt.Orientation.Horizontal)
        self.variant_slider.setRange(10, 300)
        self.variant_slider.setValue(getattr(self.lab_engine.settings if hasattr(self, 'lab_engine') and self.lab_engine else None, 'variant_pairs_count', 70) if hasattr(self, 'lab_engine') else 70)
        self.variant_slider.setFixedWidth(120)
        self.variant_slider.setToolTip(tr("Variant intensity: more pairs = more results but slower"))

        self.variant_slider_label = QLabel("70")
        self.variant_slider_label.setFixedWidth(28)
        self.variant_slider_label.setStyleSheet("font-size: 11px;")

        self.variant_slider.valueChanged.connect(lambda v: (
            self.variant_slider_label.setText(str(v)),
            self._sync_variant_sliders(v, 'search') if hasattr(self, '_sync_variant_sliders') else None,
            self._update_variant_count_preview()
        ))

        slider_layout.addWidget(self.variant_slider)
        slider_layout.addWidget(self.variant_slider_label)

        # === Common controls (shown in both modes) ===
        # Dynamic label showing estimated variants for current query
        self.variant_count_label = QLabel("")
        self.variant_count_label.setFixedWidth(50)
        self.variant_count_label.setStyleSheet("font-size: 10px; color: #7f8c8d;")
        self.variant_count_label.setToolTip(tr("Total estimated variants for all words in query"))

        # Max changes spinbox
        self.spin_max_changes = QSpinBox()
        self.spin_max_changes.setRange(1, 3)
        self.spin_max_changes.setValue(getattr(self.lab_engine.settings if hasattr(self, 'lab_engine') and self.lab_engine else None, 'variant_max_changes', 2) if hasattr(self, 'lab_engine') else 2)
        self.spin_max_changes.setFixedWidth(40)
        self.spin_max_changes.setToolTip(tr("Max character changes per word (1-3)"))
        self.spin_max_changes.setPrefix("×")

        # Add widgets to main container
        variant_layout.addWidget(self.variant_presets_widget)
        variant_layout.addWidget(self.variant_slider_widget)
        variant_layout.addWidget(self.variant_count_label)
        variant_layout.addWidget(self.spin_max_changes)

        # Show presets or slider based on setting
        use_slider = getattr(self.lab_engine.settings if hasattr(self, 'lab_engine') and self.lab_engine else None, 'variant_use_slider', False) if hasattr(self, 'lab_engine') else False
        self.variant_presets_widget.setVisible(not use_slider)
        self.variant_slider_widget.setVisible(use_slider)

        self.variant_controls_container.setVisible(False)  # Hidden by default (Exact mode)

        # PGP Tag selector (hidden by default, shown in row2 when PGP Tags mode selected)
        self.tag_search_combo = QComboBox()
        self.tag_search_combo.setEditable(True)
        self.tag_search_combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.tag_search_combo.setPlaceholderText(tr("Select a tag..."))
        self.tag_search_combo.setFixedWidth(300)
        self.tag_search_combo.addItem("")  # empty placeholder item
        completer = self.tag_search_combo.completer()
        if completer:
            completer.setFilterMode(Qt.MatchFlag.MatchContains)
        self.tag_search_combo.activated.connect(lambda idx: self._execute_tag_search() if idx > 0 else None)
        self.tag_search_combo.setVisible(False)

        # Search params container (Gap, Exclude, settings, Lab, Deep) — hidden in PGP Tags mode
        self.search_params_container = QWidget()
        self.search_params_container.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Preferred)
        params_layout = QHBoxLayout(self.search_params_container)
        params_layout.setContentsMargins(0, 0, 0, 0)

        self.gap_input = QLineEdit(); self.gap_input.setPlaceholderText(tr("Gap")); self.gap_input.setFixedWidth(50)
        self.gap_input.setToolTip(tr("Maximum word distance (0 = Exact phrase)"))

        # Exclude Words Filter (New)
        self.exclude_input = QLineEdit(); self.exclude_input.setPlaceholderText(tr("Exclude Words"))
        self.exclude_input.setToolTip(tr("Results containing these words will be filtered out"))
        self.exclude_input.setFixedWidth(120)

        # Gear button for search settings
        self.btn_search_settings = QPushButton("⚙")
        self.btn_search_settings.setFixedWidth(30)
        self.btn_search_settings.setToolTip(tr("Variant search settings"))
        self.btn_search_settings.setStyleSheet("font-size: 14px;")
        self.btn_search_settings.clicked.connect(self.open_search_settings)

        self.btn_lab_mode_toggle = QPushButton(tr("Lab Mode"))
        self.btn_lab_mode_toggle.setCheckable(True)
        self.btn_lab_mode_toggle.setToolTip(tr("Experimental search mode using advanced proximity scoring. WARNING: Can freeze the program. Use with caution."))
        self.btn_lab_mode_toggle.toggled.connect(self.on_lab_mode_toggled_search)

        # Deep Scan Checkbox
        self.chk_lab_deep = QCheckBox(tr("Deep Scan"))
        self.chk_lab_deep.setToolTip(tr("Slower but checks deeper. Use for common phrases/quotes"))
        self.chk_lab_deep.setEnabled(False) # Enabled only in Lab Mode
        self.chk_lab_deep.toggled.connect(self.on_deep_scan_toggled_search)

        params_layout.addWidget(QLabel(tr("Gap:")))
        params_layout.addWidget(self.gap_input)
        params_layout.addWidget(QLabel(tr("Exclude:")))
        params_layout.addWidget(self.exclude_input)
        params_layout.addWidget(self.btn_search_settings)
        params_layout.addWidget(self.btn_lab_mode_toggle)
        params_layout.addWidget(self.chk_lab_deep)

        # Help Button
        btn_help = QPushButton("?")
        btn_help.setFixedWidth(30)
        btn_help.setStyleSheet("background-color: #f39c12; color: white; font-weight: bold; border-radius: 15px;")
        btn_help.clicked.connect(lambda: self.open_help_center(anchor=None))

        # Domain filter button and label
        self.btn_domain_filter = QPushButton(tr("Filter by domains"))
        self.btn_domain_filter.setToolTip(tr("Filter results by subject domain (post-search)"))
        self.btn_domain_filter.setStyleSheet("padding: 2px 8px;")
        self.btn_domain_filter.clicked.connect(self._open_domain_filter_dialog)
        self.btn_domain_filter.setEnabled(False)  # Enabled after search with domain data

        self.lbl_domain_filter = QLabel("")
        self.lbl_domain_filter.setStyleSheet("color: #9b59b6; font-size: 11px;")
        self.lbl_domain_filter.setVisible(False)

        # Store domain exclusions and result-specific domain data
        self._domain_exclusions = set()
        self._result_domain_counts = {}  # domain_name -> count in current results
        self._result_domain_map = {}  # sys_id -> list of domain names
        self._has_result_domains = False

        row2.addWidget(QLabel(tr("Mode:")))
        row2.addWidget(self.mode_combo)
        row2.addWidget(self.tag_search_combo)
        row2.addWidget(self.variant_controls_container)
        row2.addWidget(self.search_params_container)
        row2.addWidget(self.btn_domain_filter)
        row2.addWidget(self.lbl_domain_filter)

        row2.addStretch()
        row2.addWidget(btn_help)

        # --- Responsa Sub-Options Row (visible only when Responsa mode is selected) ---
        self.responsa_sub_row = QWidget()
        responsa_sub_layout = QHBoxLayout(self.responsa_sub_row)
        responsa_sub_layout.setContentsMargins(0, 2, 0, 2)
        responsa_sub_layout.setSpacing(12)

        self.chk_responsa_variants = QCheckBox(tr("Variants"))
        self.chk_responsa_variants.setToolTip(tr("Include spelling variant pairs"))
        responsa_sub_layout.addWidget(self.chk_responsa_variants)

        self.chk_responsa_ja = QCheckBox(tr("Judeo-Arabic"))
        self.chk_responsa_ja.setToolTip(tr("Expand with Judeo-Arabic article forms (al-)"))
        responsa_sub_layout.addWidget(self.chk_responsa_ja)

        self.chk_responsa_flex = QCheckBox(tr("Flex Spacing"))
        self.chk_responsa_flex.setToolTip(tr("Allow flexible spacing between characters (helps with OCR text)"))
        responsa_sub_layout.addWidget(self.chk_responsa_flex)

        self.chk_bidirectional = QCheckBox(tr("Bidirectional"))
        self.chk_bidirectional.setToolTip(tr("Search for words in either order"))
        responsa_sub_layout.addWidget(self.chk_bidirectional)

        # Syntax legend label
        syntax_legend = QLabel("  #מילה " + tr("prefix") + "  |  מילה# " + tr("suffix") + "  |  %מילה " + tr("plene") + "  |  *מילה " + tr("wildcard") + "  |  (א/ב) " + tr("OR") + "  |  -מילה " + tr("Exclude"))
        syntax_legend.setStyleSheet("font-size: 10px; color: #7f8c8d;")
        responsa_sub_layout.addWidget(syntax_legend)

        # Tabular Search button
        self.btn_query_builder = QPushButton(tr("Tabular Search"))
        self.btn_query_builder.setToolTip(tr("Open the tabular query builder"))
        self.btn_query_builder.setStyleSheet("font-size: 11px; padding: 2px 8px;")
        self.btn_query_builder.clicked.connect(self._open_query_builder)
        responsa_sub_layout.addWidget(self.btn_query_builder)

        # Feature discovery glow on tabular button (one-time hint)
        self._tabular_glow_active = False
        if not cfg.get('hint_tabular_seen'):
            self._tabular_glow_active = True
            self._tabular_glow_on = True
            self.btn_query_builder.setToolTip(tr("Try the Tabular Search!"))
            self._tabular_glow_timer = QTimer(self)
            self._tabular_glow_timer.timeout.connect(self._pulse_tabular_glow)
            self._tabular_glow_timer.start(800)
            self._pulse_tabular_glow()

        responsa_sub_layout.addStretch()

        # Initially hidden (shown when Responsa mode selected)
        self.responsa_sub_row.setVisible(False)

        # Lazily load PGP tags in background
        self._pgp_tags_worker = PGPTagsWorker()
        self._pgp_tags_worker.finished.connect(self._on_pgp_tags_loaded)
        self._pgp_tags_worker.start()

        top_layout.addWidget(self.search_row1_container)
        top_layout.addLayout(row2)
        top_layout.addWidget(self.responsa_sub_row)
        layout.addWidget(top_container)

        self.lab_panel_search = LabPanel(self, 'search')
        layout.addWidget(self.lab_panel_search)
        
        self.search_progress = QProgressBar(); self.search_progress.setVisible(False)
        layout.addWidget(self.search_progress)
        
        # Results Table Setup
        self.COL_CHECKBOX = 0
        self.COL_ACTIONS = 1
        self.COL_SYS_ID = 2
        self.COL_LIBRARY = 3  # Library column (before Shelfmark)
        self.COL_SHELF = 4
        self.COL_IMG = 5
        self.COL_TITLE = 6
        self.COL_SNIPPET = 7
        self.COL_SRC = 8
        self.COL_PGP = 9
        self.COL_DOMAIN = 10

        self.results_table = QTableWidget(); self.results_table.setColumnCount(11)
        self.results_table.setHorizontalHeaderLabels(["", "", tr("System ID"), tr("Library"), tr("Shelfmark"), tr("Img"), tr("Title"), tr("Snippet"), tr("Src"), tr("PGP"), tr("Domain")])
        # Tooltip for PGP column header
        self.results_table.horizontalHeaderItem(self.COL_PGP).setToolTip(tr("Scholarly transcriptions/data available from the Princeton Geniza Project"))

        # Custom Header
        # Disable sort for Checkbox (0), Actions (1), and Image (5)
        self.chk_search_header = CheckBoxHeader(
            self.results_table,
            non_sortable_cols=[0, 1, self.COL_IMG],
            filter_columns=[self.COL_ACTIONS, self.COL_SHELF, self.COL_LIBRARY, self.COL_TITLE, self.COL_SNIPPET, self.COL_DOMAIN],
            filter_callback=self._open_results_filter_dialog,
            star_columns=[self.COL_ACTIONS],
            star_callback=self.toggle_list_filter,
            desc_first_cols=[self.COL_PGP]
        )
        self.chk_search_header.toggled.connect(self.on_search_select_all_toggled)
        self.results_table.setHorizontalHeader(self.chk_search_header)
        self._update_results_filter_indicators()

        self.results_table.setColumnWidth(self.COL_CHECKBOX, 30) # Checkbox column
        self.results_table.setColumnWidth(self.COL_ACTIONS, 95)
        self.results_table.setColumnWidth(self.COL_SYS_ID, 135)
        self.results_table.setColumnWidth(self.COL_SHELF, 175)
        self.results_table.setColumnWidth(self.COL_LIBRARY, 90)  # Library column
        self.results_table.setColumnWidth(self.COL_PGP, 40)  # PGP badge column
        self.results_table.setColumnWidth(self.COL_DOMAIN, 130)  # Domain column
        self.results_table.horizontalHeader().setSectionResizeMode(self.COL_SNIPPET, QHeaderView.ResizeMode.Stretch)
        self.results_table.horizontalHeader().setSectionResizeMode(self.COL_PGP, QHeaderView.ResizeMode.Fixed)
        # Ensure column 0 is not sortable to avoid confusion with check action
        self.results_table.horizontalHeader().setSectionResizeMode(self.COL_CHECKBOX, QHeaderView.ResizeMode.Fixed)
        self.results_table.horizontalHeader().setSectionResizeMode(self.COL_ACTIONS, QHeaderView.ResizeMode.Fixed)

        self.results_table.setMouseTracking(True)
        self.results_table.cellEntered.connect(self.on_table_cell_entered)
        self.results_table.installEventFilter(self)
        self.results_table.viewport().installEventFilter(self)

        self.results_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.results_table.setSortingEnabled(True) # Enable sorting
        self.results_table.doubleClicked.connect(self.show_full_text)
        self.results_table.itemChanged.connect(self.on_search_result_item_changed)
        self.results_table.verticalScrollBar().valueChanged.connect(self.check_scroll_load)

        # Context menu for community features
        self.results_table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.results_table.customContextMenuRequested.connect(self._show_results_context_menu)

        self.results_placeholder = QLabel(tr("Please wait while components load..."))
        self.results_placeholder.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.results_placeholder.setWordWrap(True)
        self.results_placeholder.setStyleSheet("font-size: 16px; font-weight: bold; color: #c0392b;")

        # Container for table
        self.table_container = QWidget()
        table_layout = QVBoxLayout(self.table_container)
        table_layout.setContentsMargins(0, 0, 0, 0)

        table_layout.addWidget(self.results_table)

        self.results_stack = QStackedLayout()
        self.results_stack.addWidget(self.results_placeholder)
        self.results_stack.addWidget(self.table_container)

        results_container = QWidget()
        results_container.setLayout(self.results_stack)
        layout.addWidget(results_container)

        bot = QHBoxLayout()

        self.status_label = QLabel(tr("Ready."))
        self.lbl_search_export = QLabel(tr("Export Results") + ":")
        
        # Separate export buttons
        self.btn_exp_xlsx = QPushButton("XLSX")
        self.btn_exp_xlsx.clicked.connect(lambda: self.export_results('xlsx'))
        self.btn_exp_xlsx.setFixedWidth(50)
        
        self.btn_exp_csv = QPushButton("CSV")
        self.btn_exp_csv.clicked.connect(lambda: self.export_results('csv'))
        self.btn_exp_csv.setFixedWidth(50)
        
        self.btn_exp_txt = QPushButton("TXT")
        self.btn_exp_txt.clicked.connect(lambda: self.export_results('txt'))
        self.btn_exp_txt.setFixedWidth(50)

        self.btn_exp_docx = QPushButton("DOCX")
        self.btn_exp_docx.clicked.connect(lambda: self.export_results('docx'))
        self.btn_exp_docx.setFixedWidth(60)
        
        # Track export buttons for bulk enable/disable
        self.export_buttons = [self.btn_exp_xlsx, self.btn_exp_csv, self.btn_exp_txt, self.btn_exp_docx]
        for b in self.export_buttons: b.setEnabled(False)

        # Add controls to status row
        bot.addWidget(self.status_label, 1)

        # Add to List button
        self.btn_add_to_list = QPushButton(_format_add_to_list_label(False))
        self.btn_add_to_list.clicked.connect(self.search_add_selected_to_list)
        self.btn_add_to_list.setEnabled(False)
        bot.addWidget(self.btn_add_to_list)

        # Append export controls to the right
        bot.addWidget(QLabel("|"))
        bot.addWidget(self.lbl_search_export)
        bot.addWidget(self.btn_exp_xlsx)
        bot.addWidget(self.btn_exp_csv)
        bot.addWidget(self.btn_exp_txt)
        bot.addWidget(self.btn_exp_docx)
        
        layout.addLayout(bot)
        panel.setLayout(layout)
        return panel

    def set_results_loading(self, is_loading: bool):
        """Toggle the search results placeholder while components initialize."""
        if hasattr(self, "results_stack") and hasattr(self, "results_placeholder") and hasattr(self, "table_container"):
            target = self.results_placeholder if is_loading else self.table_container
            self.results_stack.setCurrentWidget(target)

    def create_composition_tab(self):
        panel = QWidget(); layout = QVBoxLayout(); splitter = QSplitter(Qt.Orientation.Vertical)
        
        inp_w = QWidget(); in_l = QVBoxLayout()
        top_row = QHBoxLayout()
        self.comp_title_input = QLineEdit(); self.comp_title_input.setPlaceholderText(tr("Composition Title"))
        top_row.addWidget(QLabel(tr("Title:"))); top_row.addWidget(self.comp_title_input)
        
        # Load Button Cmoved to top row
        btn_load = QPushButton(tr("Load Text File")); btn_load.clicked.connect(self.load_comp_file)
        top_row.addWidget(btn_load)

        # 1. Exclude & Filter (Moved to top row)
        btn_exclude = QPushButton(tr("Exclude Manuscripts")); btn_exclude.clicked.connect(self.open_exclude_dialog)
        btn_filter_text = QPushButton(tr("Filter Text")); btn_filter_text.clicked.connect(self.open_filter_dialog)
        self.lbl_exclude_status = QLabel(tr("Excluded: {}").format(0))
        self.lbl_exclude_status.setStyleSheet("color: #8e44ad; font-weight: bold;")
        self.lbl_comp_status = QLabel("")

        # Domain filter button for composition results
        self.btn_comp_domain_filter = QPushButton(tr("Filter by domains"))
        self.btn_comp_domain_filter.setToolTip(tr("Filter results by subject domain (post-search)"))
        self.btn_comp_domain_filter.setStyleSheet("padding: 2px 8px;")
        self.btn_comp_domain_filter.clicked.connect(self._open_comp_domain_filter_dialog)
        self.btn_comp_domain_filter.setEnabled(False)

        self.lbl_comp_domain_filter = QLabel("")
        self.lbl_comp_domain_filter.setStyleSheet("color: #9b59b6; font-size: 11px;")
        self.lbl_comp_domain_filter.setVisible(False)

        # Composition domain state
        self._comp_domain_exclusions = set()
        self._comp_result_domain_counts = {}
        self._comp_result_domain_map = {}
        self._comp_has_result_domains = False

        top_row.addWidget(btn_exclude); top_row.addWidget(btn_filter_text)
        top_row.addWidget(self.btn_comp_domain_filter)
        top_row.addWidget(self.lbl_comp_domain_filter)
        top_row.addWidget(self.lbl_exclude_status)
        top_row.addWidget(self.lbl_comp_status)

        # Help Button
        btn_help = QPushButton("?")
        btn_help.setFixedWidth(30)
        btn_help.setStyleSheet("background-color: #f39c12; color: white; font-weight: bold; border-radius: 15px;")
        btn_help.clicked.connect(lambda: self.open_help_center(anchor="composition"))
        top_row.addWidget(btn_help)
        
        in_l.addLayout(top_row)
        self.comp_text_area = QPlainTextEdit(); self.comp_text_area.setPlaceholderText(tr("Paste source text..."))
        if CURRENT_LANG == 'he': self.comp_text_area.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
        in_l.addWidget(self.comp_text_area)

        # Single Row for Controls
        cr = QHBoxLayout()

        # 2. Parameters
        self.spin_chunk = QSpinBox(); self.spin_chunk.setValue(5); self.spin_chunk.setPrefix(tr("Chunk: "))
        self.spin_chunk.setToolTip(tr("Words per search block (Rec: 5-7)"))
        self.spin_chunk.valueChanged.connect(self._update_boundary_stats)
        
        self.spin_freq = QSpinBox(); self.spin_freq.setValue(10); self.spin_freq.setRange(1,1000); self.spin_freq.setPrefix(tr("Max Freq: "))
        self.spin_freq.setToolTip(tr("Ignore phrases appearing > X times (filters common phrases)"))
        
        self.comp_mode_combo = QComboBox(); self.comp_mode_combo.addItems([tr("Exact"), tr("Variants"), tr("Fuzzy")])
        self.comp_mode_combo.setItemData(0, tr("Exact match"))
        self.comp_mode_combo.setItemData(1, tr("Variant search (uses slider level)"))
        self.comp_mode_combo.setItemData(2, tr("Fuzzy search"))
        self.comp_mode_combo.currentIndexChanged.connect(self._on_comp_mode_changed)

        # Variant Level Slider for Composition (visible only in Variants mode)
        self.comp_variant_slider_container = QWidget()
        comp_slider_layout = QHBoxLayout(self.comp_variant_slider_container)
        comp_slider_layout.setContentsMargins(0, 0, 0, 0)
        self.comp_variant_slider = QSlider(Qt.Orientation.Horizontal)
        self.comp_variant_slider.setRange(10, 500)
        self.comp_variant_slider.setValue(50)
        self.comp_variant_slider.setFixedWidth(80)
        self.comp_variant_slider.setToolTip(tr("Variant intensity"))
        self.comp_variant_slider_label = QLabel("50")
        self.comp_variant_slider_label.setFixedWidth(25)
        self.comp_variant_slider.valueChanged.connect(lambda v: (
            self.comp_variant_slider_label.setText(str(v)),
            self._sync_variant_sliders(v, 'comp') if hasattr(self, '_sync_variant_sliders') else None
        ))
        comp_slider_layout.addWidget(self.comp_variant_slider)
        comp_slider_layout.addWidget(self.comp_variant_slider_label)
        self.comp_variant_slider_container.setVisible(False)

        self.spin_filter = QSpinBox(); self.spin_filter.setValue(5); self.spin_filter.setPrefix(tr("Filter > "))
        self.spin_filter.setToolTip(tr("Move titles appearing > X times to Appendix"))

        # Shortened Text
        self.chk_comp_flat = QCheckBox(tr("Sort by shelfmark only"))
        self.chk_comp_flat.setToolTip(tr("Disable Main/Appendix grouping"))
        self.chk_comp_flat.toggled.connect(self.on_comp_display_mode_changed)

        cr.addWidget(self.spin_chunk); cr.addWidget(self.spin_freq)
        cr.addWidget(self.comp_mode_combo); cr.addWidget(self.comp_variant_slider_container)
        cr.addWidget(self.spin_filter); cr.addWidget(self.chk_comp_flat)

        # Boundary Search Controls Row
        boundary_row = QHBoxLayout()
        boundary_row.setContentsMargins(0, 0, 0, 0)

        # Boundary Mode ComboBox
        self.boundary_mode_combo = QComboBox()
        self.boundary_mode_combo.addItem(tr("Full search"), "full")
        self.boundary_mode_combo.addItem(tr("Cross-paragraph only"), "boundary")
        self.boundary_mode_combo.addItem(tr("Full + Cross-paragraph boost"), "combined")
        self.boundary_mode_combo.setToolTip(tr("Search all text chunks regardless of paragraph breaks"))
        self.boundary_mode_combo.currentIndexChanged.connect(self._on_boundary_mode_changed)

        # Delimiter ComboBox
        self.boundary_delimiter_combo = QComboBox()
        self.boundary_delimiter_combo.addItem(tr("Line break"), "\n")
        self.boundary_delimiter_combo.addItem(tr("Blank line (paragraph)"), "\n\n")
        self.boundary_delimiter_combo.addItem(tr("Period (.)"), ".")
        self.boundary_delimiter_combo.addItem(tr("Colon (:)"), ":")
        self.boundary_delimiter_combo.setToolTip(tr("Character or pattern that separates paragraphs in your text"))
        self.boundary_delimiter_combo.currentIndexChanged.connect(self._on_boundary_delimiter_changed)

        # Boundary Stats Label (colors chosen for both light and dark mode visibility)
        self.boundary_stats_label = QLabel("")
        self.boundary_stats_label.setStyleSheet("color: #5dade2; font-size: 11px;")

        # Advanced Settings Button
        self.btn_boundary_advanced = QPushButton("⚙")
        self.btn_boundary_advanced.setFixedWidth(30)
        self.btn_boundary_advanced.setToolTip(tr("Advanced cross-paragraph settings"))
        self.btn_boundary_advanced.clicked.connect(self._open_boundary_advanced_dialog)
        self.btn_boundary_advanced.setVisible(False)  # Only show in non-full modes

        boundary_row.addWidget(QLabel(tr("Paragraph search") + ":"))
        boundary_row.addWidget(self.boundary_mode_combo)
        boundary_row.addWidget(QLabel(tr("Paragraph separator") + ":"))
        boundary_row.addWidget(self.boundary_delimiter_combo)
        boundary_row.addWidget(self.btn_boundary_advanced)
        boundary_row.addWidget(self.boundary_stats_label)
        boundary_row.addStretch()

        in_l.addLayout(boundary_row)

        # Initialize boundary settings from LabSettings if available
        if hasattr(self, 'lab_engine') and self.lab_engine:
            settings = self.lab_engine.settings
            mode_index = {'full': 0, 'boundary': 1, 'combined': 2}.get(settings.boundary_mode, 0)
            self.boundary_mode_combo.setCurrentIndex(mode_index)
            delim_index = {'\n': 0, '\n\n': 1, '.': 2, ':': 3}.get(settings.boundary_delimiter, 0)
            self.boundary_delimiter_combo.setCurrentIndex(delim_index)

        # Connect text area changes to update stats
        self.comp_text_area.textChanged.connect(self._update_boundary_stats)

        # 3. Lab & Action
        self.btn_lab_mode_toggle_comp = QPushButton(tr("Lab Mode"))
        self.btn_lab_mode_toggle_comp.setCheckable(True)
        self.btn_lab_mode_toggle_comp.setToolTip(tr("Experimental search mode using advanced proximity scoring. WARNING: Can freeze the program. Use with caution."))
        self.btn_lab_mode_toggle_comp.toggled.connect(self.on_lab_mode_toggled_comp)

        self.chk_lab_deep_comp = QCheckBox(tr("Deep Scan"))
        self.chk_lab_deep_comp.setToolTip(tr("Slower but checks deeper. Use for common phrases/quotes"))
        self.chk_lab_deep_comp.setEnabled(False)
        self.chk_lab_deep_comp.toggled.connect(self.on_deep_scan_toggled_comp)

        # Shortened Text
        self.btn_comp_run = QPushButton(tr("Analyze")); self.btn_comp_run.clicked.connect(self.toggle_composition)
        self.btn_comp_run.setStyleSheet("background-color: #2980b9; color: white; font-weight: bold;")
        self.btn_comp_run.setEnabled(False)
        self.btn_comp_recursive = QPushButton(tr("Full Recursive Search")); self.btn_comp_recursive.clicked.connect(self.run_recursive_composition)
        self.btn_comp_recursive.setStyleSheet("background-color: #27ae60; color: white; font-weight: bold;")
        self.btn_comp_recursive.setEnabled(True)

        cr.addWidget(self.btn_lab_mode_toggle_comp)
        cr.addWidget(self.chk_lab_deep_comp)
        cr.addWidget(self.btn_comp_run)
        cr.addWidget(self.btn_comp_recursive)

        in_l.addLayout(cr)

        self.lab_panel_comp = LabPanel(self, 'comp')
        in_l.addWidget(self.lab_panel_comp)

        self.comp_progress = QProgressBar(); self.comp_progress.setVisible(False)
        in_l.addWidget(self.comp_progress)
        inp_w.setLayout(in_l); splitter.addWidget(inp_w)
        
        res_w = QWidget(); rl = QVBoxLayout()
        self.comp_tree = QTreeWidget(); self.comp_tree.setHeaderLabels([tr("Score"), tr("Library"), tr("Shelfmark"), tr("Title"), tr("System ID"), tr("Context"), tr("MS Context")])
        self.comp_tree.itemChanged.connect(self.on_comp_tree_item_changed)
        self.comp_tree.itemExpanded.connect(self.on_comp_tree_item_expanded)
        self.comp_tree.itemCollapsed.connect(self.on_comp_tree_item_collapsed)
        self.comp_tree_updating = False
        self.comp_tree.setStyleSheet(
            "QTreeWidget::indicator { width: 16px; height: 16px; }"
            "QTreeWidget::indicator:unchecked { border: 1px solid #9b9b9b; background: transparent; }"
            "QTreeWidget::indicator:checked { border: 1px solid #9b9b9b; background: rgba(255, 255, 255, 0.35); }"
            "QTreeWidget::indicator:indeterminate { border: 1px solid #9b9b9b; background: rgba(255, 255, 255, 0.18); }"
        )

        # Configure columns width
        header = self.comp_tree.header()
        header.setSectionsClickable(True)
        header.sectionClicked.connect(self.on_comp_header_clicked)
        header.sectionResized.connect(self._refresh_comp_tree_tooltips)
        header.setSortIndicatorShown(True)
        header.setSortIndicator(0, Qt.SortOrder.DescendingOrder)
        header.setSectionResizeMode(self.comp_col_context, QHeaderView.ResizeMode.Interactive)
        header.setSectionResizeMode(self.comp_col_ms_context, QHeaderView.ResizeMode.Stretch)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Interactive) # Shelfmark
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents) # System ID

        self.comp_tree.setColumnWidth(0, 160) # Score - widened

        # Title column (~25 chars)
        title_width = self.comp_tree.fontMetrics().averageCharWidth() * 25
        self.comp_tree.setColumnWidth(2, int(title_width))
        context_width = self.comp_tree.fontMetrics().averageCharWidth() * 35
        self.comp_tree.setColumnWidth(self.comp_col_context, int(context_width))
        self.comp_tree.setColumnWidth(self.comp_col_ms_context, int(context_width))
        header.setStretchLastSection(True)

        self.comp_tree.itemDoubleClicked.connect(self.on_comp_item_double_clicked)
        self.comp_tree.itemExpanded.connect(self._on_comp_item_expanded)
        self.comp_tree.itemCollapsed.connect(self._on_comp_item_collapsed)

        # Use CheckBoxHeader for tree
        self.chk_comp_header = CheckBoxHeader(
            self.comp_tree,
            filter_columns=[self.comp_col_library, self.comp_col_shelfmark, self.comp_col_title, self.comp_col_context, self.comp_col_ms_context],
            filter_callback=self._open_comp_filter_dialog,
        )
        self.chk_comp_header.toggled.connect(self.on_comp_header_toggled)
        self.comp_tree.setHeader(self.chk_comp_header)
        self._update_comp_filter_indicators()
        comp_header = self.comp_tree.header()
        comp_header.setSectionResizeMode(self.comp_col_context, QHeaderView.ResizeMode.Interactive)
        comp_header.setSectionResizeMode(self.comp_col_ms_context, QHeaderView.ResizeMode.Stretch)
        comp_header.setSectionResizeMode(self.comp_col_library, QHeaderView.ResizeMode.Interactive) # Library
        comp_header.setSectionResizeMode(self.comp_col_shelfmark, QHeaderView.ResizeMode.Interactive) # Shelfmark
        comp_header.setSectionResizeMode(self.comp_col_sysid, QHeaderView.ResizeMode.ResizeToContents) # System ID
        comp_header.setStretchLastSection(True)

        rl.addWidget(self.comp_tree)
        
        exp_layout = QHBoxLayout()
        self.btn_comp_add_to_list = QPushButton(_format_add_to_list_label(False))
        self.btn_comp_add_to_list.clicked.connect(self.comp_add_selected_to_list)
        self.btn_comp_add_to_list.setEnabled(False)
        exp_layout.addWidget(self.btn_comp_add_to_list)
        exp_layout.addStretch()
        self.lbl_comp_export = QLabel(tr("Save Report"))
        exp_layout.addWidget(self.lbl_comp_export)
        
        self.btn_comp_xlsx = QPushButton("XLSX")
        self.btn_comp_xlsx.clicked.connect(lambda: self.export_comp_report('xlsx'))
        
        self.btn_comp_csv = QPushButton("CSV")
        self.btn_comp_csv.clicked.connect(lambda: self.export_comp_report('csv'))
        
        self.btn_comp_txt = QPushButton("TXT")
        self.btn_comp_txt.clicked.connect(lambda: self.export_comp_report('txt'))

        self.btn_comp_docx = QPushButton("DOCX")
        self.btn_comp_docx.clicked.connect(lambda: self.export_comp_report('docx'))
        
        self.comp_export_buttons = [self.btn_comp_xlsx, self.btn_comp_csv, self.btn_comp_txt, self.btn_comp_docx]
        for b in self.comp_export_buttons:
            b.setEnabled(False) 
            
        exp_layout.addWidget(self.btn_comp_xlsx)
        exp_layout.addWidget(self.btn_comp_csv)
        exp_layout.addWidget(self.btn_comp_txt)
        exp_layout.addWidget(self.btn_comp_docx)
        
        rl.addLayout(exp_layout)
        
        res_w.setLayout(rl); splitter.addWidget(res_w)
        
        layout.addWidget(splitter); panel.setLayout(layout)
        return panel

    def create_browse_tab(self):
        panel = QWidget(); layout = QVBoxLayout(panel)
        
        # --- Top Area: Metadata (Gray Bar) ---
        top_container = QFrame();
        top_container.setFrameShape(QFrame.Shape.StyledPanel)
        
        top_layout = QVBoxLayout(top_container)
        top_layout.setContentsMargins(10, 5, 10, 5)

        # Row 1: Search Inputs
        row1 = QHBoxLayout()
        self.btn_browse_by_list = QPushButton(tr("Browse by List"))
        self.btn_browse_by_list.clicked.connect(self.browse_toggle_lists_panel)
        row1.addWidget(self.btn_browse_by_list)

        self.btn_prev_ms = QPushButton(tr("◀"))
        self.btn_prev_ms.setToolTip(tr("Previous Manuscript (File Order)"))
        self.btn_prev_ms.setFixedWidth(25)
        self.btn_prev_ms.clicked.connect(lambda: self.navigate_manuscript(-1))

        self.btn_next_ms = QPushButton(tr("▶"))
        self.btn_next_ms.setToolTip(tr("Next Manuscript (File Order)"))
        self.btn_next_ms.setFixedWidth(25)
        self.btn_next_ms.clicked.connect(lambda: self.navigate_manuscript(1))

        self.browse_sys_input = QLineEdit(); self.browse_sys_input.setPlaceholderText(tr("Enter System ID..."))
        self.browse_shelf_input = QLineEdit(); self.browse_shelf_input.setPlaceholderText(tr("Enter shelfmark..."))
        self.browse_fl_input = QLineEdit(); self.browse_fl_input.setPlaceholderText(tr("Enter FL ID..."))
        self.browse_fl_input.setFixedWidth(140)

        self.btn_browse_go = QPushButton(tr("Go")); self.btn_browse_go.setFixedWidth(50)
        self.btn_browse_go.clicked.connect(self.browse_load)
        self.btn_browse_go.setEnabled(False)
        
        self.browse_sys_input.returnPressed.connect(self.browse_load)
        self.browse_shelf_input.returnPressed.connect(self.browse_load)
        self.browse_fl_input.returnPressed.connect(self.browse_load)
        self.browse_sys_input.textEdited.connect(lambda _t: self._set_last_browse_field("sys"))
        self.browse_shelf_input.textEdited.connect(lambda _t: self._set_last_browse_field("shelf"))
        self.browse_fl_input.textEdited.connect(lambda _t: self._set_last_browse_field("fl"))
        
        # Find Parallels (Top Row)
        self.btn_find_parallels = QPushButton(tr("Find parallels"))
        self.btn_find_parallels.clicked.connect(self.browse_search_parallels)
        self.btn_find_parallels.setEnabled(False)

        self.btn_b_catalog = QPushButton(tr("View on Ktiv")); self.btn_b_catalog.setToolTip(tr("Open in Ktiv Website"))
        self.btn_b_catalog.clicked.connect(self.browse_open_catalog); self.btn_b_catalog.setEnabled(False)
        
        # View All and Save moved to Nav Bar, defined here as class members
        self.btn_b_save = QPushButton(tr("Save")); self.btn_b_save.setToolTip(tr("Save full manuscript to file"))
        self.btn_b_save.clicked.connect(self.browse_save_full); self.btn_b_save.setEnabled(False)
        
        self.btn_b_all = QPushButton(tr("View All"))
        self.btn_b_all.setCheckable(True)
        self.btn_b_all.clicked.connect(self.toggle_browse_view_all)
        self.btn_b_all.setEnabled(False)
        
        row1.addWidget(self.btn_prev_ms)
        row1.addWidget(QLabel(tr("System ID:")))
        row1.addWidget(self.browse_sys_input)
        row1.addWidget(QLabel(tr("Shelfmark:"))); row1.addWidget(self.browse_shelf_input)
        row1.addWidget(self.btn_next_ms)
        row1.addSpacing(10)
        row1.addWidget(QLabel(tr("FL:"))); row1.addWidget(self.browse_fl_input)
        row1.addWidget(self.btn_browse_go)

        # Add to View button for reading desk (immediately after Go for discoverability)
        self.btn_b_add_to_view = QPushButton(tr("Add to View"))
        self.btn_b_add_to_view.setToolTip(tr("Add current manuscript to Reading Desk"))
        self.btn_b_add_to_view.setEnabled(False)
        self.btn_b_add_to_view.clicked.connect(self._browse_add_to_view)
        row1.addWidget(self.btn_b_add_to_view)

        row1.addWidget(self.btn_find_parallels)

        # Add to List button for browse tab
        self.btn_browse_add_to_list = QPushButton(_format_add_to_list_label(False))
        self.btn_browse_add_to_list.clicked.connect(self.browse_add_to_list)
        self.btn_browse_add_to_list.setEnabled(False)
        row1.addWidget(self.btn_browse_add_to_list)

        row1.addSpacing(20)
        row1.addWidget(self.btn_b_catalog)
        row1.addStretch()

        # Help
        btn_browse_help = QPushButton("?")
        btn_browse_help.setFixedWidth(30)
        btn_browse_help.setStyleSheet("background-color: #f39c12; color: white; font-weight: bold; border-radius: 15px;")
        btn_browse_help.clicked.connect(lambda: self.open_help_center(anchor="browse"))
        row1.addWidget(btn_browse_help)

        top_layout.addLayout(row1)

        # Row 2: Metadata Display (Compact)
        self.browse_info_lbl = QLabel(tr("Enter ID to browse."))
        self.browse_info_lbl.setStyleSheet("font-size: 12px;")
        self.browse_info_lbl.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        top_layout.addWidget(self.browse_info_lbl)

        # Extended Info button + Bibliography buttons (shared row)
        ext_info_row = QHBoxLayout()
        self.btn_b_ext_info = QPushButton(tr("Show Extended Info"))
        self.btn_b_ext_info.setCheckable(True)
        self.btn_b_ext_info.toggled.connect(self._browse_toggle_extended_info)
        self.btn_b_ext_info.setVisible(False)
        ext_info_row.addWidget(self.btn_b_ext_info)
        ext_info_row.addStretch()
        self.btn_b_bibliography_fjms = QPushButton()
        self.btn_b_bibliography_fjms.setVisible(False)
        self.btn_b_bibliography_fjms.clicked.connect(self._show_fjms_bibliography_dialog)
        ext_info_row.addWidget(self.btn_b_bibliography_fjms)
        self.btn_b_bibliography_nli = QPushButton()
        self.btn_b_bibliography_nli.setVisible(False)
        self.btn_b_bibliography_nli.clicked.connect(self._show_nli_bibliography_dialog)
        ext_info_row.addWidget(self.btn_b_bibliography_nli)
        self.btn_b_catalog_records = QPushButton(f"{tr('Catalog Records')} (0)")
        self.btn_b_catalog_records.setEnabled(False)
        self.btn_b_catalog_records.setVisible(False)
        self.btn_b_catalog_records.clicked.connect(self._show_fjms_catalog_dialog)
        ext_info_row.addWidget(self.btn_b_catalog_records)
        top_layout.addLayout(ext_info_row)
        self._browse_fjms_bib = []
        self._browse_marc_bib = []
        self._browse_catalog_detail = None

        self.txt_b_extended_info = QTextBrowser()
        self.txt_b_extended_info.setVisible(False)
        self.txt_b_extended_info.setMaximumHeight(200)
        self.txt_b_extended_info.setStyleSheet("border: 1px solid #ccc; padding: 5px;")
        self.txt_b_extended_info.setOpenLinks(False)
        self.txt_b_extended_info.anchorClicked.connect(self._on_browse_ext_link_clicked)
        top_layout.addWidget(self.txt_b_extended_info)

        layout.addWidget(top_container)

        # --- Main Splitter (Left: List Panel, Center: Text, Right: Images) ---
        self.browse_splitter = QSplitter(Qt.Orientation.Horizontal)

        # Left: Lists Side Panel
        self.browse_lists_panel = QFrame()
        self.browse_lists_panel.setFrameShape(QFrame.Shape.StyledPanel)
        browse_lists_layout = QVBoxLayout(self.browse_lists_panel)
        browse_lists_layout.setContentsMargins(8, 8, 8, 8)

        browse_lists_header = QHBoxLayout()
        browse_lists_title = QLabel(tr("Browse by List"))
        browse_lists_title.setStyleSheet("font-weight: bold;")
        browse_lists_header.addWidget(browse_lists_title)
        browse_lists_header.addStretch()
        self.btn_close_browse_lists = QPushButton("✕")
        self.btn_close_browse_lists.setFixedWidth(28)
        self.btn_close_browse_lists.setToolTip(tr("Close"))
        self.btn_close_browse_lists.clicked.connect(lambda: self.browse_set_lists_panel_visible(False))
        browse_lists_header.addWidget(self.btn_close_browse_lists)
        browse_lists_layout.addLayout(browse_lists_header)

        self.browse_lists_tree = QTreeWidget()
        self.browse_lists_tree.setHeaderHidden(True)
        self.browse_lists_tree.itemClicked.connect(self.browse_on_list_selected)
        browse_lists_layout.addWidget(self.browse_lists_tree, 2)

        self.browse_list_items_label = QLabel(tr("Items"))
        self.browse_list_items_label.setStyleSheet("font-weight: bold;")
        browse_lists_layout.addWidget(self.browse_list_items_label)

        self.browse_list_items = QListWidget()
        self.browse_list_items.itemClicked.connect(self.browse_on_list_item_clicked)
        browse_lists_layout.addWidget(self.browse_list_items, 3)

        # Left: Text Browser
        text_widget = QWidget(); text_layout = QVBoxLayout(text_widget); text_layout.setContentsMargins(0,0,0,0)

        # Navigation Bar (Above Text)
        nav_bar = QHBoxLayout()
        self.btn_b_prev = QPushButton(tr("< Prev")); self.btn_b_prev.clicked.connect(lambda: self.browse_navigate(-1))
        self.btn_b_next = QPushButton(tr("Next >")); self.btn_b_next.clicked.connect(lambda: self.browse_navigate(1))
        self.btn_b_prev.setEnabled(False); self.btn_b_next.setEnabled(False)

        # Page Combo
        self.combo_browse_page = QComboBox()
        self.combo_browse_page.setEditable(False)
        self.combo_browse_page.setFixedWidth(100)
        self.combo_browse_page.currentIndexChanged.connect(self.on_browse_page_combo_changed)

        # Folio label (Phase 31: shows "Folio 1r" or empty when no folio data)
        self.lbl_browse_folio = QLabel("")
        self.lbl_browse_folio.setStyleSheet("font-weight: bold; font-size: 12px; margin-left: 4px;")
        self.lbl_browse_folio.setVisible(False)

        # Page count label (Phase 31: "of N pages")
        self.lbl_browse_page_count = QLabel("")
        self.lbl_browse_page_count.setStyleSheet("font-size: 11px; color: #7f8c8d; margin-left: 4px;")
        self.lbl_browse_page_count.setVisible(False)

        # Folio images cache for the current manuscript
        self._browse_folio_images = []

        # Image Toggle
        self.btn_b_toggle_img = QPushButton()
        self.btn_b_toggle_img.setText("🖼️")
        self.btn_b_toggle_img.setToolTip(tr("Show/Hide Image"))
        self.btn_b_toggle_img.setCheckable(True)
        self.btn_b_toggle_img.setChecked(True)
        self.btn_b_toggle_img.clicked.connect(self.toggle_browse_image)
        self.btn_b_toggle_img.setEnabled(False)

        # Layout: [< Prev] [Folio Label] [Page Combo] [of N pages] [Next >] [View All] [Save] [Image Toggle]
        nav_bar.addWidget(self.btn_b_prev)
        nav_bar.addWidget(self.lbl_browse_folio)
        nav_bar.addWidget(self.combo_browse_page)
        nav_bar.addWidget(self.lbl_browse_page_count)
        nav_bar.addWidget(self.btn_b_next)
        nav_bar.addWidget(self.btn_b_all)
        nav_bar.addWidget(self.btn_b_save)
        nav_bar.addWidget(self.btn_b_toggle_img)

        nav_bar.addStretch()
        text_layout.addLayout(nav_bar)

        # Second row: Community buttons and version selector
        community_bar = QHBoxLayout()
        community_bar.setContentsMargins(0, 2, 0, 2)

        # Version selector
        community_bar.addWidget(QLabel(tr("Version:")))
        self.browse_version_combo = QComboBox()
        self.browse_version_combo.blockSignals(True)
        self.browse_version_combo.addItem("V0.8", {"source": "original"})
        self.browse_version_combo.setFixedWidth(240)  # Wider for PGP scholar names
        self.browse_version_combo.setEnabled(False)
        self.browse_version_combo.blockSignals(False)
        self.browse_version_combo.currentIndexChanged.connect(self._browse_change_version)
        community_bar.addWidget(self.browse_version_combo)
        # Version data cache for current document
        self._browse_versions_cache = {}
        # PGP source state for current document
        self._browse_pgp_sources = []
        self._browse_pgp_doc = {}
        self._browse_enriched_html = ''
        self._browse_pgp_worker = None

        community_bar.addWidget(QLabel(" | "))

        # Edit Mode button (for inline corrections)
        self.btn_b_edit = QPushButton(tr("✏️ Edit"))
        self.btn_b_edit.setToolTip(tr("Enable edit mode to make corrections"))
        self.btn_b_edit.clicked.connect(self._browse_toggle_edit_mode)
        self.btn_b_edit.setEnabled(False)
        community_bar.addWidget(self.btn_b_edit)

        # Add Comment button
        self.btn_b_comment = QPushButton(tr("💬 Comment"))
        self.btn_b_comment.setToolTip(tr("Add a comment on this document"))
        self.btn_b_comment.clicked.connect(self._browse_add_comment)
        self.btn_b_comment.setEnabled(False)
        community_bar.addWidget(self.btn_b_comment)

        # View Corrections button
        self.btn_b_view_corrections = QPushButton(tr("View Corrections"))
        self.btn_b_view_corrections.setToolTip(tr("View community corrections for this document"))
        self.btn_b_view_corrections.clicked.connect(self._browse_view_corrections)
        self.btn_b_view_corrections.setEnabled(False)
        community_bar.addWidget(self.btn_b_view_corrections)

        # View Comments indicator (icon that appears when comments exist)
        self.btn_b_view_comments = QPushButton("💬")
        self.btn_b_view_comments.setToolTip(tr("View comments on this document"))
        self.btn_b_view_comments.clicked.connect(self._browse_view_comments)
        self.btn_b_view_comments.setEnabled(False)
        self.btn_b_view_comments.setVisible(False)  # Hidden by default, show when comments exist
        self.btn_b_view_comments.setFixedSize(32, 32)
        self.btn_b_view_comments.setStyleSheet("background-color: #f39c12; color: white; border-radius: 4px;")
        community_bar.addWidget(self.btn_b_view_comments)

        # Joins button with dropdown - show connected fragments
        self.btn_b_joins = QToolButton()
        self.btn_b_joins.setText("🔗")
        self.btn_b_joins.setToolTip(tr("View joined fragments"))
        self.btn_b_joins.setEnabled(False)
        self.btn_b_joins.setFixedSize(40, 32)
        self.btn_b_joins.setStyleSheet("background-color: #3498db; color: white; border-radius: 4px;")
        self.btn_b_joins.setPopupMode(QToolButton.ToolButtonPopupMode.MenuButtonPopup)
        self.btn_b_joins.clicked.connect(self._browse_view_joins)
        self.joins_menu = QMenu(self)
        self.joins_menu.aboutToShow.connect(self._on_joins_menu_show)
        self.btn_b_joins.setMenu(self.joins_menu)
        community_bar.addWidget(self.btn_b_joins)

        community_bar.addStretch()
        text_layout.addLayout(community_bar)

        # Edit mode action bar (hidden by default)
        self.browse_edit_bar = QWidget()
        self.browse_edit_bar.setStyleSheet("background-color: #2c3e50; border-radius: 5px;")
        edit_bar_layout = QHBoxLayout(self.browse_edit_bar)
        edit_bar_layout.setContentsMargins(10, 5, 10, 5)

        # Edit mode label
        edit_mode_label = QLabel(f"<b style='color: white;'>{tr('Edit Mode')}</b>")
        edit_bar_layout.addWidget(edit_mode_label)

        # Save status indicator
        self.browse_edit_status = QLabel()
        self.browse_edit_status.setStyleSheet("color: #f39c12; font-weight: bold;")
        edit_bar_layout.addWidget(self.browse_edit_status)

        edit_bar_layout.addStretch()

        # Save button (draft)
        self.btn_b_save_draft = QPushButton(f"💾 {tr('Save')}")
        self.btn_b_save_draft.setStyleSheet("background-color: #3498db; color: white; padding: 5px 15px; border-radius: 3px;")
        self.btn_b_save_draft.setToolTip(tr("Save as draft for later editing"))
        self.btn_b_save_draft.clicked.connect(lambda: self._browse_save_correction(submit=False))
        self.btn_b_save_draft.setEnabled(False)
        edit_bar_layout.addWidget(self.btn_b_save_draft)

        # Save and Submit button
        self.btn_b_save_correction = QPushButton(f"📤 {tr('Submit')}")
        self.btn_b_save_correction.setStyleSheet("background-color: #27ae60; color: white; padding: 5px 15px; border-radius: 3px;")
        self.btn_b_save_correction.setToolTip(tr("Submit correction for review"))
        self.btn_b_save_correction.clicked.connect(lambda: self._browse_save_correction(submit=True))
        self.btn_b_save_correction.setEnabled(False)
        edit_bar_layout.addWidget(self.btn_b_save_correction)

        self.btn_b_cancel_edit = QPushButton(tr("Cancel"))
        self.btn_b_cancel_edit.setStyleSheet("background-color: #95a5a6; color: white; padding: 5px 15px; border-radius: 3px;")
        self.btn_b_cancel_edit.clicked.connect(self._browse_cancel_edit)
        edit_bar_layout.addWidget(self.btn_b_cancel_edit)
        self.browse_edit_bar.hide()
        text_layout.addWidget(self.browse_edit_bar)

        browse_find_row = QHBoxLayout()
        browse_find_row.addWidget(QLabel(tr("Find:")))
        self.browse_find_input = QLineEdit()
        self.browse_find_input.setPlaceholderText(tr("Find in text..."))
        self.browse_find_input.textChanged.connect(lambda text: apply_find_highlight(self.browse_text, text.strip()))
        browse_find_row.addWidget(self.browse_find_input)
        text_layout.addLayout(browse_find_row)

        # Reading Desk toolbar (hidden by default, shown when reading desk is active)
        self.browse_rd_toolbar = QWidget()
        self.browse_rd_toolbar.setStyleSheet(
            "background-color: #2d6a4f; border-radius: 4px;"
        )
        rd_toolbar_layout = QHBoxLayout(self.browse_rd_toolbar)
        rd_toolbar_layout.setContentsMargins(10, 5, 10, 5)

        rd_label = QLabel(f"<b style='color: white;'>{tr('Reading Desk')}</b>")
        rd_toolbar_layout.addWidget(rd_label)

        self.browse_rd_count_label = QLabel()
        self.browse_rd_count_label.setStyleSheet("color: #a7d8c0; font-weight: bold;")
        rd_toolbar_layout.addWidget(self.browse_rd_count_label)

        rd_toolbar_layout.addSpacing(10)

        rd_shelf_label = QLabel(f"<span style='color: white;'>{tr('Shelfmark:')}</span>")
        rd_toolbar_layout.addWidget(rd_shelf_label)
        self.browse_rd_shelf_input = QLineEdit()
        self.browse_rd_shelf_input.setPlaceholderText(tr("Add shelfmark..."))
        self.browse_rd_shelf_input.setFixedWidth(180)
        self.browse_rd_shelf_input.returnPressed.connect(self._browse_rd_add_by_shelfmark)
        rd_toolbar_layout.addWidget(self.browse_rd_shelf_input)

        btn_rd_add = QPushButton(tr("Add to Desk"))
        btn_rd_add.setStyleSheet(
            "background-color: #40916c; color: white; padding: 3px 10px; border-radius: 3px;"
        )
        btn_rd_add.clicked.connect(self._browse_rd_add_by_shelfmark)
        rd_toolbar_layout.addWidget(btn_rd_add)

        rd_toolbar_layout.addSpacing(10)

        btn_rd_add_from_list = QPushButton(tr("Add from List"))
        btn_rd_add_from_list.setStyleSheet(
            "background-color: #52b788; color: white; padding: 3px 10px; border-radius: 3px;"
        )
        btn_rd_add_from_list.clicked.connect(self._browse_rd_add_from_list)
        rd_toolbar_layout.addWidget(btn_rd_add_from_list)

        rd_toolbar_layout.addStretch()

        btn_rd_exit = QPushButton(tr("Exit Reading Desk"))
        btn_rd_exit.setStyleSheet(
            "background-color: #c0392b; color: white; padding: 3px 10px; border-radius: 3px;"
        )
        btn_rd_exit.clicked.connect(self._browse_exit_reading_desk)
        rd_toolbar_layout.addWidget(btn_rd_exit)

        self.browse_rd_toolbar.hide()
        text_layout.addWidget(self.browse_rd_toolbar)

        # Text display/edit widget
        self.browse_text = QTextEdit()
        self.browse_text.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
        self.browse_text.setFont(QFont("SBL Hebrew", 16))
        self.browse_text.setReadOnly(True)  # Start in read-only mode
        self.browse_edit_mode = False
        self.browse_original_text = ""
        text_layout.addWidget(self.browse_text)

        # Install link click handler for genizah:// URLs in QTextEdit
        # (QTextEdit doesn't have anchorClicked signal like QTextBrowser,
        # so we intercept mouseReleaseEvent and use anchorAt to detect clicks)
        _original_mouse_release = self.browse_text.mouseReleaseEvent

        def _browse_text_mouse_release(event, orig=_original_mouse_release):
            if event.button() == Qt.MouseButton.LeftButton:
                pos = event.pos()
                anchor = self.browse_text.anchorAt(pos)
                if anchor and anchor.startswith("genizah://"):
                    self._on_browse_link_clicked(QUrl(anchor))
                    event.accept()
                    return
            orig(event)

        self.browse_text.mouseReleaseEvent = _browse_text_mouse_release
        
        # Right: Image Viewer
        self.browse_viewer = ManuscriptViewerWidget()

        self.browse_splitter.addWidget(self.browse_lists_panel)
        self.browse_splitter.addWidget(text_widget)
        self.browse_splitter.addWidget(self.browse_viewer)
        self.browse_splitter.setStretchFactor(0, 0)
        self.browse_splitter.setStretchFactor(1, 1)
        self.browse_splitter.setStretchFactor(2, 1)
        self.browse_lists_panel.setVisible(False)

        layout.addWidget(self.browse_splitter, 1)

        # Dummy placeholders
        self.browse_thumb = QLabel()
        self.btn_b_ext = QPushButton()
        self.browse_side_panel = QTextBrowser()
        self.browse_lists_panel_sizes = None
        self.browse_current_list_id = None

        # Reading Desk state
        self.browse_reading_desk_active = False
        self.browse_reading_desk_state = ReadingDeskState()
        self.browse_reading_desk_pgpid = None
        self._browse_rd_worker = None
        self._browse_rd_image_widgets = []  # list of (sys_id, ZoomableScrollArea, ImageLoaderThread)
        self._browse_rd_image_scroll = None  # QScrollArea for stacked images
        self._browse_rd_syncing = False  # prevents infinite scroll sync loop
        self._rd_text_sync_handler = None  # stored ref for targeted disconnect
        self._rd_image_sync_handler = None  # stored ref for targeted disconnect

        return panel

    def browse_toggle_lists_panel(self):
        """Toggle the browse lists side panel."""
        self.browse_set_lists_panel_visible(not self.browse_lists_panel.isVisible())

    def browse_set_lists_panel_visible(self, visible):
        """Show or hide the browse lists side panel."""
        if visible and not self.lists_mgr:
            QMessageBox.information(self, tr("Browse by List"), tr("Lists are not available."))
            return

        if self.browse_lists_panel.isVisible() == visible:
            return

        if visible:
            self.browse_lists_panel.setVisible(True)
            self.browse_refresh_lists_panel()
            sizes = self.browse_lists_panel_sizes
            rd_active = self.browse_reading_desk_active and self._browse_rd_image_scroll is not None
            if not sizes:
                total = sum(self.browse_splitter.sizes()) or 1000
                if rd_active:
                    # 4-widget layout: lists, text, (hidden viewer), image_scroll
                    sizes = [max(200, int(total * 0.20)), max(300, int(total * 0.35)), 0, max(300, int(total * 0.45))]
                else:
                    # 3-widget layout: lists, text, viewer
                    sizes = [max(250, int(total * 0.25)), max(350, int(total * 0.45)), max(350, int(total * 0.30))]
            elif rd_active and len(sizes) == 3:
                # Cached sizes from normal mode -- need to expand to 4 elements
                total = sum(sizes) or 1000
                sizes = [sizes[0], max(300, int(total * 0.35)), 0, max(300, int(total * 0.45))]
            self.browse_splitter.setSizes(sizes)
        else:
            self.browse_lists_panel_sizes = self.browse_splitter.sizes()
            self.browse_lists_panel.setVisible(False)
            sizes = self.browse_splitter.sizes()
            if sizes:
                sizes[0] = 0
                self.browse_splitter.setSizes(sizes)

    def browse_refresh_lists_panel(self):
        """Refresh the lists tree and items list in the browse panel."""
        self.browse_lists_tree.clear()
        self.browse_list_items.clear()
        self.browse_current_list_id = None

        if not self.lists_mgr:
            return

        lists = self.lists_mgr.get_all_lists(include_recent=True)
        projects = {proj['id']: proj for proj in self.lists_mgr.get_projects()}
        project_lists = {}
        top_level_lists = []

        for lst in lists:
            project_id = lst.get('project_id')
            if project_id and project_id in projects:
                project_lists.setdefault(project_id, []).append(lst)
            else:
                top_level_lists.append(lst)

        for lst in top_level_lists:
            item = QTreeWidgetItem([self._get_list_display_name(lst)])
            item.setData(0, Qt.ItemDataRole.UserRole, lst.get('id'))
            self.browse_lists_tree.addTopLevelItem(item)

        for project in self.lists_mgr.get_projects():
            proj_item = QTreeWidgetItem([project.get('name', tr("Project"))])
            proj_item.setFlags(proj_item.flags() & ~Qt.ItemFlag.ItemIsSelectable)
            self.browse_lists_tree.addTopLevelItem(proj_item)
            for lst in project_lists.get(project.get('id'), []):
                child = QTreeWidgetItem([self._get_list_display_name(lst)])
                child.setData(0, Qt.ItemDataRole.UserRole, lst.get('id'))
                proj_item.addChild(child)

        self.browse_lists_tree.expandAll()

    def browse_on_list_selected(self, item, column):
        """Handle selection of a list in the browse lists panel."""
        list_id = item.data(0, Qt.ItemDataRole.UserRole)
        if not list_id or not self.lists_mgr:
            return

        self.browse_current_list_id = list_id
        self.browse_list_items.clear()

        items = self.lists_mgr.get_items_sorted(list_id, sort_by='shelfmark')
        if not items:
            empty_item = QListWidgetItem(tr("No items in this list."))
            empty_item.setFlags(empty_item.flags() & ~Qt.ItemFlag.ItemIsSelectable)
            self.browse_list_items.addItem(empty_item)
            return

        for entry in items:
            sys_id = entry.get('sys_id')
            item_id = entry.get('item_id')
            if not sys_id or not item_id:
                continue

            shelfmark, title = self.meta_mgr.get_meta_for_id(sys_id)
            shelfmark = entry.get('shelfmark_override') or shelfmark or sys_id
            img = entry.get('img')
            fl_id = entry.get('fl_id')

            label_parts = [shelfmark]
            if img not in (None, ""):
                label_parts.append(tr("Img {}").format(self._format_image_display(img)))
            extra_bits = []  
            if title:
                extra_bits.append(title)                    
            if extra_bits:
                label_parts.append(f"({', '.join(extra_bits)})")

            list_item = QListWidgetItem(" • ".join(label_parts))
            list_item.setData(Qt.ItemDataRole.UserRole, item_id)
            self.browse_list_items.addItem(list_item)

    def browse_on_list_item_clicked(self, item):
        """Open a list item in the browse tab using FL/Image ID lookup.

        When reading desk is active, adds the item to the desk instead of navigating.
        """
        item_id = item.data(Qt.ItemDataRole.UserRole)
        if not item_id or not self.lists_mgr:
            return

        entry = self.lists_mgr.get_item(item_id)
        if not entry:
            return

        # Reading desk mode: add item to desk instead of navigating
        if self.browse_reading_desk_active:
            sys_id = entry.get('sys_id')
            if sys_id:
                shelfmark, _ = self.meta_mgr.get_meta_for_id(sys_id)
                if not shelfmark or shelfmark == "Unknown":
                    shelfmark = entry.get('shelfmark_override') or sys_id
                self._browse_rd_add_entry(sys_id, shelfmark)
            return

        fl_id = entry.get('fl_id')
        img_id = entry.get('img')
        sys_id = entry.get('sys_id')

        self.browse_sys_input.clear()
        self.browse_shelf_input.clear()
        self.browse_fl_input.clear()

        target_fl = fl_id if fl_id else img_id

        if target_fl:
            clean_fl = str(target_fl).replace('FL', '').strip()

            self.browse_fl_input.setText(clean_fl)
            self._set_last_browse_field("fl")
            self.browse_load()

        elif sys_id:
            self.browse_sys_input.setText(str(sys_id))
            self._set_last_browse_field("sys")
            self.browse_load()

    def on_browse_enriched_loaded(self, sid, meta):
        if not meta: return
        if sid != self.current_browse_sid: return
        if self.current_browse_sid not in self.meta_mgr.nli_cache: return

        # Reset extended info panel for new manuscript
        self.btn_b_ext_info.setVisible(False)
        self.btn_b_ext_info.setChecked(False)
        self.txt_b_extended_info.setVisible(False)
        self.txt_b_extended_info.setHtml("")
        self.btn_b_bibliography_fjms.setVisible(False)
        self.btn_b_bibliography_nli.setVisible(False)
        self.btn_b_catalog_records.setVisible(False)
        self.btn_b_catalog_records.setEnabled(False)
        self._browse_fjms_bib = []
        self._browse_marc_bib = []
        self._browse_catalog_detail = None

        # Store folio images for page combo labels (Phase 31)
        self._browse_folio_images = meta.get('folio_images', [])

        # 1. Update Info Label (Top Bar)
        marc = meta.get('marc', {})
        shelf = meta.get('shelfmark')
        title = meta.get('title')
        if shelf and shelf != "Unknown":
            # Try CSV library_code first, then MARC as fallback
            library_code = self.meta_mgr.get_library_for_id(sid)
            if library_code:
                library = get_library_display(library_code, short=False)
            else:
                library = marc.get('current_owner', '')
            if library:
                shelf = f"{library} | {shelf}"

        # Check for Oxford Part metadata - integrated into shelfmark
        part_id = self.current_browse_part_id
        if not part_id:
            part_id = self.meta_mgr.get_part_for_folio(sid)

        if part_id:
            part_meta = self.meta_mgr.get_part_metadata(part_id)
            oxford_title = part_meta.get('title', '') if part_meta else ''
            folio_range = part_meta.get('folio_range', []) if part_meta else []

            # Extract part number from part_id (e.g., "MS. Heb. b. 10/43" -> "43")
            part_num = part_id.split('/')[-1] if '/' in part_id else ''

            # Build combined: "MS heb. b. 10/79 (part 43: fols. 79-82)" or "fol. 14" for single
            shelf_with_part = f"{shelf or ''}"
            if part_num:
                shelf_with_part += f" (part {part_num}"
                if len(folio_range) == 2:
                    if folio_range[0] == folio_range[1]:
                        shelf_with_part += f": fol. {folio_range[0]}"
                    else:
                        shelf_with_part += f": fols. {folio_range[0]}–{folio_range[1]}"
                shelf_with_part += ")"

            label_text = f"<b>{shelf_with_part}</b>"
            tooltip_parts = []
            if oxford_title:
                truncated, full = _truncate_title(oxford_title)
                label_text += f"<br/><span style='font-size: 11px;'>{truncated}</span>"
                if full: tooltip_parts.append(full)
            if title and title != oxford_title:
                truncated, full = _truncate_title(title)
                label_text += f"<br/>{truncated}"
                if full: tooltip_parts.append(full)
            if meta.get('physical_desc'):
                label_text += f" | {meta['physical_desc']}"
        else:
            label_text = f"<b>{shelf or ''}</b>"
            tooltip_parts = []
            if title:
                truncated, full = _truncate_title(title)
                label_text += f" | {truncated}"
                if full: tooltip_parts.append(full)
            if meta.get('physical_desc'):
                label_text += f" | {meta['physical_desc']}"

        # Append Neubauer-Cowley catalog entry for Oxford manuscripts
        catalog_entry = meta.get('catalog_entry')
        if catalog_entry:
            label_text += f" | {catalog_entry}"

        self.browse_info_lbl.setText(label_text)
        self.browse_info_lbl.setToolTip('\n'.join(tooltip_parts) if tooltip_parts else '')

        # 2. Populate Image Viewer (using new logic)
        # Guard: skip when reading desk is active to preserve stacked image layout.
        # This handles both new requests and in-flight enrichment threads that
        # complete after reading desk activation.
        if not self.browse_reading_desk_active:
            folio_num = _get_folio_number_from_shelfmark(shelf)
            idx = _get_initial_image_index(meta, folio_num if folio_num is not None else self.current_browse_p)
            self.browse_viewer.load_images(meta, idx, target_folio=folio_num)

        # 3. Enable buttons
        self.btn_b_catalog.setEnabled(True)
        self.btn_b_save.setEnabled(True)
        self.btn_b_all.setEnabled(True)
        self.btn_find_parallels.setEnabled(True)
        self.btn_browse_add_to_list.setEnabled(True)
        self.btn_b_toggle_img.setEnabled(True)
        # Enable community buttons
        self.btn_b_edit.setEnabled(True)
        self.btn_b_comment.setEnabled(True)
        self.btn_b_view_corrections.setEnabled(True)
        self.btn_b_joins.setEnabled(True)
        self.browse_version_combo.setEnabled(True)
        self.btn_b_add_to_view.setEnabled(True)

        # Check for comments and corrections FIRST (resets PGP state and version combo)
        self._check_document_community_status()

        # Build enrichment extended info (KTI/Oxford/Cambridge)
        marc = meta.get('marc', {})
        part_id_for_ext = self.current_browse_part_id or self.meta_mgr.get_part_for_folio(sid)
        part_meta_ext = self.meta_mgr.get_part_metadata(part_id_for_ext) if part_id_for_ext else None
        external_meta = meta.get('external_meta', {})

        palette = self.txt_b_extended_info.palette()
        text_color = palette.color(QPalette.ColorRole.Text).name()
        base_color = palette.color(QPalette.ColorRole.Base).name()

        physical_metadata = meta.get('physical_metadata')
        library_viewer_url = meta.get('library_viewer_url')

        enriched_html = self._build_browse_enriched_html(
            marc, part_id_for_ext, part_meta_ext, external_meta, text_color, base_color,
            physical_metadata=physical_metadata, library_viewer_url=library_viewer_url
        )

        # Build FJMS domain and catalog HTML and prepend to enriched HTML
        fjms_catalog_html = self._build_fjms_catalog_html(sid, text_color)
        fjms_domain_html = self._build_fjms_domain_html(sid, text_color)

        # Build Phase 33 metadata HTML (catalog refs, secondary metadata)
        catalog_refs_html = self._build_catalog_refs_html(meta, text_color)
        secondary_meta_html = self._build_secondary_metadata_html(meta, text_color)

        # Separate bibliography buttons (FJMS / NLI)
        fjms_bib = meta.get('bibliography', [])
        marc_bib = marc.get('bibliography', [])
        if fjms_bib:
            self._browse_fjms_bib = fjms_bib
            self.btn_b_bibliography_fjms.setText(f"{tr('Bibliography FJMS')} ({len(fjms_bib)})")
            self.btn_b_bibliography_fjms.setVisible(True)
        else:
            self._browse_fjms_bib = []
            self.btn_b_bibliography_fjms.setVisible(False)
        if marc_bib:
            self._browse_marc_bib = marc_bib
            self.btn_b_bibliography_nli.setText(f"{tr('Bibliography Ktiv')} ({len(marc_bib)})")
            self.btn_b_bibliography_nli.setVisible(True)
        else:
            self._browse_marc_bib = []
            self.btn_b_bibliography_nli.setVisible(False)

        # Catalog Records button (FJMS catalog detail dialog)
        # Detail is fetched lazily on button click, not during page load
        self._browse_catalog_detail = None
        try:
            from shared.fjms_service import get_fjms_service
            fjms_svc = get_fjms_service()
            if fjms_svc.is_available():
                source_names = fjms_svc.get_source_names(sid)
                catalog_count = len(source_names)
                self.btn_b_catalog_records.setText(f"{tr('Catalog Records')} ({catalog_count})")
                self.btn_b_catalog_records.setEnabled(catalog_count > 0)
                self.btn_b_catalog_records.setVisible(True)
        except Exception:
            self.btn_b_catalog_records.setVisible(False)

        enriched_html = fjms_domain_html + fjms_catalog_html + catalog_refs_html + secondary_meta_html + enriched_html

        self._browse_enriched_html = enriched_html

        if enriched_html.strip():
            full_html = f"<div style='font-family:Arial; color:{text_color}; background-color:{base_color};'>{enriched_html}</div>"
            self.txt_b_extended_info.setHtml(full_html)
            self.btn_b_ext_info.setVisible(True)

        # Update joins dropdown menu (PGP joins will be added when PGP worker completes)
        self._update_joins_dropdown()

        # Start PGP source fetch (runs in background, populates combo when done)
        # Disconnect old worker signals first to prevent stale results
        if self._browse_pgp_worker is not None:
            try:
                self._browse_pgp_worker.finished_signal.disconnect(self._on_browse_pgp_loaded)
                self._browse_pgp_worker.error_signal.disconnect(self._on_browse_pgp_error)
            except (TypeError, RuntimeError):
                pass
        self._browse_pgp_worker = PGPSourceWorker(self.current_browse_sid, self.current_browse_p or 1)
        self._browse_pgp_worker.finished_signal.connect(self._on_browse_pgp_loaded)
        self._browse_pgp_worker.error_signal.connect(self._on_browse_pgp_error)
        self._browse_pgp_worker.start()

        # 4. Trigger Page Load to show text (skip if already rendered by arrow navigation)
        if getattr(self, '_browse_nav_rendered', False):
            self._browse_nav_rendered = False
        else:
            self.browse_load_page()

    def _on_browse_pgp_loaded(self, sys_id, sources, pgp_doc):
        """Handle PGP sources loaded from background thread."""
        # Stale-request guard: user may have navigated to a different manuscript
        if sys_id != self.current_browse_sid:
            return

        # Store PGP data for later use (page changes)
        self._browse_pgp_sources = sources
        self._browse_pgp_doc = pgp_doc

        # Update extended info panel with PGP metadata combined with enrichment
        pgp_html = self._build_pgp_extended_info_html(pgp_doc) or ''
        enriched_html = getattr(self, '_browse_enriched_html', '') or ''

        combined = enriched_html + pgp_html
        if combined.strip():
            palette = self.txt_b_extended_info.palette()
            text_color = palette.color(QPalette.ColorRole.Text).name()
            base_color = palette.color(QPalette.ColorRole.Base).name()
            full_html = f"<div style='font-family:Arial; color:{text_color}; background-color:{base_color};'>{combined}</div>"
            self.txt_b_extended_info.setHtml(full_html)
            self.btn_b_ext_info.setVisible(True)
        else:
            self.btn_b_ext_info.setVisible(False)
            self.btn_b_ext_info.setChecked(False)
            self.txt_b_extended_info.setVisible(False)

        # Refresh joins dropdown now that PGP data is available
        self._update_joins_dropdown()

        if not sources:
            return

        # Save existing corrections/versions from the combo before rebuilding.
        # Only save non-PGP, non-original, non-header items (corrections, V0.7, user versions).
        saved_corrections = []
        for i in range(self.browse_version_combo.count()):
            data = self.browse_version_combo.itemData(i)
            if data and data.get('source') not in ('original', 'header', 'pgp_edition', 'pgp_translation', None):
                saved_corrections.append((self.browse_version_combo.itemText(i), data))

        # Populate combo with PGP items (clears and rebuilds: PGP Editions > Translations > V0.8)
        has_pgp = self._populate_pgp_combo(self.browse_version_combo, sources, pgp_doc)

        if has_pgp:
            # Re-add saved corrections/versions after V0.8
            if saved_corrections:
                self.browse_version_combo.blockSignals(True)
                self.browse_version_combo.insertSeparator(self.browse_version_combo.count())
                for label, data in saved_corrections:
                    self.browse_version_combo.addItem(label, data)
                self.browse_version_combo.blockSignals(False)

            # Store original V0.8 text from browse_render_page output
            current_text = self.browse_text.toPlainText()
            if current_text:
                self.browse_original_page_text = current_text

            # Auto-select first PGP edition and display it
            edition_data = self._auto_select_pgp_edition(self.browse_version_combo)
            if edition_data:
                content = edition_data.get('content', '')
                if content:
                    self._browse_display_pgp_text(content, is_rtl=True)

            self.browse_version_combo.setEnabled(True)

    def _on_browse_pgp_error(self, sys_id, error_message):
        """Handle PGP source fetch error -- silently fall back to existing behavior."""
        logger.debug("PGP source fetch error for %s: %s", sys_id, error_message)

    def _build_pgp_extended_info_html(self, pgp_doc, palette=None):
        """Build HTML for PGP metadata section in extended info panels.

        Args:
            pgp_doc: dict from PGPSourceWorker (document metadata)
            palette: optional QPalette for text color; defaults to app palette

        Returns:
            HTML string for the PGP section, or empty string if no PGP data.
        """
        if not pgp_doc:
            return ""

        if palette is None:
            palette = self.palette()
        text_color = palette.color(QPalette.ColorRole.Text).name()

        pgp_html = (
            f"<div style='background-color: transparent; color:{text_color}; "
            "padding: 10px; margin-bottom: 10px; "
            "border-left: 3px solid #27ae60; text-align: left;' dir='ltr'>"
        )
        pgp_html += f"<p style='margin-top:0;'><b>Princeton Geniza Project</b></p>"

        doc_type = pgp_doc.get('document_type')
        if doc_type:
            pgp_html += f"<p><b>{tr('Document Type')}:</b> {doc_type}</p>"

        tags = pgp_doc.get('tags', [])
        if tags:
            tag_links = []
            for tag in tags:
                tag_links.append(
                    f"<a href='tag:{tag}' style='color: #27ae60; text-decoration: underline;'>{tag}</a>"
                )
            pgp_html += f"<p><b>{tr('Tags')}:</b> {', '.join(tag_links)}</p>"

        description = pgp_doc.get('description')
        if description:
            pgp_html += f"<p><b>{tr('Description')}:</b> {description}</p>"

        date = pgp_doc.get('inferred_date_display') or pgp_doc.get('doc_date_standard')
        if date:
            pgp_html += f"<p><b>{tr('Date')}:</b> {date}</p>"

        pgp_url = pgp_doc.get('pgp_url')
        if pgp_url:
            pgp_html += (
                f"<p><a href='{pgp_url}' style='color: #27ae60; text-decoration: underline;'>"
                f"{tr('View on PGP')}</a></p>"
            )

        pgp_html += "</div>"
        return pgp_html

    def _build_fjms_domain_html(self, sys_id, text_color):
        """Build HTML for FJMS domain classifications in extended info."""
        from shared.fjms_service import get_fjms_service
        fjms = get_fjms_service()
        if not fjms.is_available():
            return ""

        domains = fjms.get_domains(sys_id)
        if not domains:
            return ""

        # Deduplicate: skip parent if child already shown
        all_domain_names = {d['domain'] for d in domains}
        filtered = []
        for dom in domains:
            parent = dom.get('parent_domain')
            if parent and parent in all_domain_names and parent != dom['domain']:
                continue
            filtered.append(dom)

        if not filtered:
            return ""

        html = (
            f"<div style='color:{text_color}; padding: 5px 0; margin-bottom: 8px; "
            "border-left: 3px solid #9b59b6;' dir='ltr'>"
            f"<p style='margin: 0 0 4px 10px;'><b>{tr('Subject Domains')}</b></p>"
            "<p style='margin: 0 0 0 10px;'>"
        )

        links = []
        for dom in filtered:
            domain_name = dom['domain']
            display_name = dom.get('domain_heb', domain_name) if CURRENT_LANG == 'he' else domain_name
            # Clickable link that navigates to search with domain filter
            links.append(
                f"<a href='domain:{domain_name}' style='color: #9b59b6; "
                f"text-decoration: underline;'>{display_name}</a>"
            )
        html += ", ".join(links)
        html += "</p></div>"
        return html

    def _build_fjms_catalog_html(self, sys_id, text_color):
        """Build HTML for FJMS catalog metadata in extended info."""
        from shared.fjms_service import get_fjms_service, merge_catalog_records, parse_textual_frame
        fjms = get_fjms_service()
        if not fjms.is_available():
            return ""

        records = fjms.get_catalog_records(sys_id)
        if not records:
            return ""

        merged = merge_catalog_records(records)

        html = (
            f"<div style='color:{text_color}; padding: 10px; margin-bottom: 10px; "
            "border-left: 3px solid #9b59b6; text-align: left;' dir='ltr'>"
            f"<p style='margin-top:0;'><b>{tr('FJMS Catalog')}</b></p>"
        )

        # Title (language-aware)
        title = merged.get('title_heb') if CURRENT_LANG == 'he' else merged.get('title')
        if title and title.strip():
            html += f"<p><b>{tr('Title')}:</b> {title}</p>"

        # Author
        if merged.get('author_text') and merged['author_text'].strip():
            html += f"<p><b>{tr('Author')}:</b> {merged['author_text']}</p>"

        # Date and Place
        date = merged.get('copy_date')
        place = merged.get('copy_place')
        if date or place:
            parts = []
            if date:
                parts.append(f"<b>{tr('Copy Date')}:</b> {date}")
            if place:
                parts.append(f"<b>{tr('Place')}:</b> {place}")
            html += f"<p>{' &nbsp;|&nbsp; '.join(parts)}</p>"

        # Content Identifications (TextualFrames)
        frames = merged.get('textual_frames', [])
        if frames:
            html += f"<p style='margin-bottom:4px;'><b>{tr('Content Identification')}:</b></p>"

            html += "<ul style='margin-top:2px; padding-left:20px;'>"
            for frame in frames:
                text = frame.get('heb') if CURRENT_LANG == 'he' else frame.get('eng')
                if not text or not text.strip():
                    text = frame.get('eng') if CURRENT_LANG == 'he' else frame.get('heb')
                if text and text.strip():
                    category, content = parse_textual_frame(text)
                    source = frame.get('source_name_heb') if CURRENT_LANG == 'he' else frame.get('source_name')
                    if category:
                        li = f"<b style='color:#9b59b6;'>{category}:</b> {content}"
                    else:
                        li = text
                    if source and source.strip():
                        li += f" <span style='color:gray; font-size:0.85em;'>({source})</span>"
                    html += f"<li>{li}</li>"
            html += "</ul>"

        html += "</div>"
        return html

    def _build_catalog_refs_html(self, meta, text_color):
        """Build HTML for FIST catalog cross-references in extended info."""
        cat_refs = meta.get('catalog_refs', [])
        if not cat_refs:
            return ""

        html = (
            f"<div style='color:{text_color}; padding: 10px; margin-bottom: 10px; "
            "border-left: 3px solid #1abc9c; text-align: left;' dir='ltr'>"
            f"<p style='margin-top:0;'><b>{tr('Catalog References')}</b></p>"
        )

        html += "<ul style='margin-top:2px; padding-left:20px;'>"
        for ref in cat_refs:
            acronym = ref.get('cat_acronym', '')
            entry = ref.get('catalog_entry', '')
            display = f"{acronym} #{entry}" if entry else acronym
            html += f"<li>{display}</li>"
        html += "</ul></div>"
        return html

    def _build_secondary_metadata_html(self, meta, text_color):
        """Build HTML for secondary metadata (source names, collection, storage)."""
        source_names = meta.get('source_names', [])
        coll_storage = meta.get('collection_storage')

        if not source_names and not coll_storage:
            return ""

        html = (
            f"<div style='color:{text_color}; padding: 5px 0; margin-bottom: 8px; "
            "border-left: 3px solid #bdc3c7; text-align: left;' dir='ltr'>"
        )

        if source_names:
            html += f"<p style='margin: 0 0 4px 10px;'><b>{tr('Scholarly Sources')}:</b> "
            html += ', '.join(source_names) + "</p>"

        if coll_storage:
            parts = []
            coll = coll_storage.get('collection_name', '')
            if coll:
                parts.append(coll)
            storage = []
            box = coll_storage.get('ob_box', '')
            vol = coll_storage.get('ob_volume', '')
            folio = coll_storage.get('ob_folio', '')
            if box:
                storage.append(f'Box {box}')
            if vol:
                storage.append(f'Vol. {vol}')
            if folio:
                storage.append(f'Fol. {folio}')
            if storage:
                parts.append(', '.join(storage))
            if parts:
                html += f"<p style='margin: 0 0 0 10px;'><b>{tr('Collection & Storage')}:</b> {' — '.join(parts)}</p>"

        html += "</div>"
        return html

    def _build_browse_enriched_html(self, marc, part_id, part_meta, external_meta, text_color, base_color,
                                     physical_metadata=None, library_viewer_url=None):
        """Build HTML for KTI/Oxford/Cambridge enrichment data in Browse extended info.

        Mirrors the enrichment HTML builder in ResultDialog.on_enriched_data_loaded.

        Args:
            marc: dict from MARC metadata
            part_id: Oxford part ID (if applicable)
            part_meta: Oxford part metadata dict
            external_meta: External metadata dict (Cambridge, etc.)
            text_color: CSS text color
            base_color: CSS background color
            physical_metadata: dict with material, num_folio, num_bifolio, size from NLI crossref
            library_viewer_url: dict with url, label for library digital collection link

        Returns:
            HTML string for enrichment section, or empty string.
        """
        if not marc and not part_meta and not external_meta and not physical_metadata and not library_viewer_url:
            return ""

        part_bg = QColor(base_color).lighter(115).name()

        # Physical metadata from NLI crossref (Phase 32: META-01, META-02)
        phys_html = ""
        if physical_metadata:
            material = physical_metadata.get('material', '')
            num_folio = physical_metadata.get('num_folio', '')
            num_bifolio = physical_metadata.get('num_bifolio', '')
            size = physical_metadata.get('size', '')

            if material or num_folio or num_bifolio or size:
                phys_html += f"<p style='margin-bottom:4px;'><b>{tr('Physical Description')}:</b></p>"
                if material:
                    phys_html += f"<p style='margin-left:12px;'><b>{tr('Material')}:</b> {tr(material)}</p>"
                folio_parts = []
                if num_folio and num_folio != '0':
                    folio_parts.append(f"{num_folio} {tr('Folios')}")
                if num_bifolio and num_bifolio != '0':
                    folio_parts.append(f"{num_bifolio} {tr('Bifolios')}")
                if folio_parts:
                    phys_html += f"<p style='margin-left:12px;'><b>{tr('Folios')}:</b> {' + '.join(folio_parts)}</p>"
                if size:
                    phys_html += f"<p style='margin-left:12px;'><b>{tr('Size')}:</b> {size}</p>"

        # Library viewer link (Phase 32: META-04)
        if library_viewer_url and library_viewer_url.get('url'):
            lib_label = library_viewer_url.get('label', tr('View in Library Catalog'))
            lib_url = library_viewer_url['url']
            phys_html += f"<p style='margin-left:12px;'><a href='{lib_url}' style='color:#1976d2;'>{lib_label}</a></p>"

        kti_html = ""
        date_val = marc.get('date')
        if date_val:
            kti_html += f"<p><b>{tr('Date')}:</b> {date_val}</p>"

        dims = marc.get('dimensions')
        phys = marc.get('physical_desc') or marc.get('physical_description')
        if dims or phys:
            kti_html += f"<p><b>{tr('Physical Description')}:</b> {phys or ''} {dims or ''}</p>"

        eng_title = marc.get('english_title')
        if eng_title:
            kti_html += f"<p><b>{tr('English Title')}:</b> {eng_title}</p>"

        subjects = marc.get('subjects', [])
        if subjects:
            kti_html += f"<p><b>{tr('Subjects')}:</b> {'; '.join(subjects)}</p>"

        notes = marc.get('notes', [])
        if notes:
            kti_html += f"<p><b>{tr('Notes')}:</b><ul>"
            for n in notes:
                kti_html += f"<li>{n}</li>"
            kti_html += "</ul></p>"

        people = marc.get('people', [])
        if people:
            kti_html += f"<p><b>{tr('People')}:</b> {'; '.join(people)}</p>"

        external_html = ""
        if part_meta:
            part_display = self.meta_mgr.codico_mgr.get_part_display_name(part_id) if hasattr(self.meta_mgr, 'codico_mgr') else part_id
            external_html += (
                f"<div style='background-color: {part_bg}; color:{text_color}; padding: 10px; "
                "margin-bottom: 10px; border-left: 3px solid #3498db; text-align: left;' dir='ltr'>"
            )
            external_html += f"<p><b>{tr('Codicological Part')}:</b> {part_display}</p>"

            folio_range = part_meta.get('folio_range', [])
            if len(folio_range) == 2:
                if folio_range[0] == folio_range[1]:
                    external_html += f"<p><b>{tr('Folio')}:</b> {folio_range[0]}</p>"
                else:
                    external_html += f"<p><b>{tr('Folio Range')}:</b> {folio_range[0]} - {folio_range[1]}</p>"

            part_title = part_meta.get('title', '')
            if part_title:
                external_html += f"<p><b>{tr('Oxford Title')}:</b> {part_title}</p>"

            part_contents = part_meta.get('contents', '')
            if part_contents:
                external_html += f"<p><b>{tr('Contents')}:</b> {part_contents}</p>"

            external_html += "</div>"

        if external_meta:
            external_html += f"<div style='margin-bottom: 10px; text-align: left;' dir='ltr'><ul>"
            for k, v in external_meta.items():
                external_html += f"<li><b>{k}:</b> {v}</li>"
            external_html += "</ul></div>"

        is_rtl = self.layoutDirection() == Qt.LayoutDirection.RightToLeft
        dir_attr = "rtl" if is_rtl else "ltr"
        header_align = "right" if is_rtl else "left"
        kti_header = tr("Ktiv Info")
        if part_id:
            external_header = tr("Oxford Info")
        else:
            external_header = tr("Cambridge Info")

        html = ""
        if external_html:
            if is_rtl:
                first_title, first_html = kti_header, kti_html
                second_title, second_html = external_header, external_html
            else:
                first_title, first_html = external_header, external_html
                second_title, second_html = kti_header, kti_html

            html += (
                f"<table style='width:100%; border-collapse:collapse;' dir='{dir_attr}'>"
                f"<tr>"
                f"<th style='text-align:{header_align}; padding:4px; border-bottom:1px solid #ccc;'>{first_title}</th>"
                f"<th style='text-align:{header_align}; padding:4px; border-bottom:1px solid #ccc;'>{second_title}</th>"
                f"</tr>"
                f"<tr>"
                f"<td style='vertical-align:top; padding:6px;'>{first_html}</td>"
                f"<td style='vertical-align:top; padding:6px;'>{second_html}</td>"
                f"</tr></table>"
            )
        elif kti_html:
            html += kti_html

        return phys_html + html

    def _browse_toggle_extended_info(self, checked):
        """Toggle browse tab extended info panel visibility."""
        self.txt_b_extended_info.setVisible(checked)
        self.btn_b_ext_info.setText(
            tr("Hide Extended Info") if checked else tr("Show Extended Info")
        )

    def _show_fjms_bibliography_dialog(self):
        """Open the FJMS bibliography dialog."""
        if not self._browse_fjms_bib:
            return
        shelf = self.meta_mgr.get_meta_for_id(self.current_browse_sid)[0] if self.current_browse_sid else ''
        dlg = FjmsBibliographyDialog(
            self._browse_fjms_bib,
            sys_id=self.current_browse_sid or '',
            shelfmark=shelf,
            parent=self,
        )
        dlg.exec()

    def _show_nli_bibliography_dialog(self):
        """Open the NLI bibliography dialog."""
        if not self._browse_marc_bib:
            return
        shelf = self.meta_mgr.get_meta_for_id(self.current_browse_sid)[0] if self.current_browse_sid else ''
        dlg = NliBibliographyDialog(
            self._browse_marc_bib,
            sys_id=self.current_browse_sid or '',
            shelfmark=shelf,
            parent=self,
        )
        dlg.exec()

    def _show_fjms_catalog_dialog(self):
        """Open the FJMS catalog records dialog from Browse tab (lazy fetch)."""
        # Lazy fetch: load catalog detail on first click if not yet loaded
        if self._browse_catalog_detail is None and self.current_browse_sid:
            self.statusBar().showMessage(tr("Loading catalog data..."), 3000)
            try:
                from shared.fjms_service import get_fjms_service
                fjms_svc = get_fjms_service()
                if fjms_svc.is_available():
                    self._browse_catalog_detail = fjms_svc.get_catalog_detail(self.current_browse_sid)
            except Exception:
                pass
            self.statusBar().clearMessage()

        if not self._browse_catalog_detail or not self._browse_catalog_detail.get("records"):
            # Still try to show if we have free_descriptions or other data
            if not self._browse_catalog_detail:
                return
        shelf = self.meta_mgr.get_meta_for_id(self.current_browse_sid)[0] if self.current_browse_sid else ''
        dlg = FjmsCatalogDialog(
            self._browse_catalog_detail,
            sys_id=self.current_browse_sid or '',
            shelfmark=shelf,
            parent=self,
        )
        dlg.exec()

    def _on_browse_ext_link_clicked(self, url):
        """Handle clicks on links in browse tab extended info."""
        url_str = url.toString()
        if url_str.startswith('tag:'):
            tag = url_str[4:]
            self._search_by_pgp_tag(tag)
        elif url_str.startswith('domain:'):
            domain = url_str[7:]
            self._navigate_to_search_with_domain(domain)
        elif url_str.startswith('http'):
            QDesktopServices.openUrl(url)

    def _browse_display_pgp_text(self, text, is_rtl=True):
        """Display PGP edition/translation text with proper directionality."""
        if not text:
            return
        direction = 'rtl' if is_rtl else 'ltr'
        layout_dir = Qt.LayoutDirection.RightToLeft if is_rtl else Qt.LayoutDirection.LeftToRight
        self.browse_text.setLayoutDirection(layout_dir)
        browse_html_text = text.replace('\n', '<br>')
        self.browse_text.setHtml(f"<div dir='{direction}'>{browse_html_text}</div>")
        apply_find_highlight(self.browse_text, self.browse_find_input.text().strip())

    def _browse_refresh_pgp_for_page(self):
        """Re-fetch PGP sources for current page (called on page change within same manuscript)."""
        if not hasattr(self, '_browse_pgp_sources') or not self._browse_pgp_sources:
            return  # No PGP data was loaded for this manuscript
        if not self.current_browse_sid:
            return
        # Disconnect old worker signals first
        if self._browse_pgp_worker is not None:
            try:
                self._browse_pgp_worker.finished_signal.disconnect(self._on_browse_pgp_loaded)
                self._browse_pgp_worker.error_signal.disconnect(self._on_browse_pgp_error)
            except (TypeError, RuntimeError):
                pass
        # Start a new PGP worker for the current page
        self._browse_pgp_worker = PGPSourceWorker(self.current_browse_sid, self.current_browse_p or 1)
        self._browse_pgp_worker.finished_signal.connect(self._on_browse_pgp_loaded)
        self._browse_pgp_worker.error_signal.connect(self._on_browse_pgp_error)
        self._browse_pgp_worker.start()

    def _on_browse_link_clicked(self, url):
        """Handle clicks on internal links in browse text (View All and Reading Desk modes)."""
        url_str = url.toString()

        # Reading Desk links
        if url_str.startswith("genizah://rd-navigate/"):
            sid = url_str.replace("genizah://rd-navigate/", "")
            if sid:
                self._browse_exit_reading_desk()
                self.browse_sys_input.setText(sid)
                shelf, _ = self.meta_mgr.get_meta_for_id(sid)
                if shelf and shelf != "Unknown":
                    self.browse_shelf_input.setText(shelf)
                self._set_last_browse_field("sys")
                self.browse_load()
            return

        if url_str.startswith("genizah://rd-remove/"):
            sid = url_str.replace("genizah://rd-remove/", "")
            if sid:
                self._browse_rd_remove_entry(sid)
            return

        if url_str.startswith("genizah://rd-version/"):
            parts = url_str.replace("genizah://rd-version/", "").split("/")
            if len(parts) >= 2:
                self._browse_rd_show_version_dialog(parts[0], int(parts[1]))
            return

        # View All mode links
        if url_str.startswith("genizah://load/"):
            # Extract system ID from URL
            sid = url_str.replace("genizah://load/", "")
            if sid:
                # Exit View All mode
                self.btn_b_all.setChecked(False)
                self.browse_viewer.setVisible(self.btn_b_toggle_img.isChecked())
                self.btn_b_toggle_img.setEnabled(True)

                # Clear Part state and load the specific folio
                self.current_browse_part_id = None
                self.current_browse_part_folios = []
                self.current_browse_part_folio_idx = 0

                # Load the folio
                self.browse_sys_input.setText(sid)
                shelf, _ = self.meta_mgr.get_meta_for_id(sid)
                if shelf and shelf != "Unknown":
                    self.browse_shelf_input.setText(shelf)
                self._set_last_browse_field("sys")
                self.browse_load()

    # -------------------------------------------------------------------------
    # Reading Desk -- Desktop dual-pane rendering
    # -------------------------------------------------------------------------

    def _browse_enter_reading_desk(self, fragments_info, pgpid=None):
        """Enter reading desk mode with the given fragments.

        Args:
            fragments_info: list of dicts with keys: sys_id, shelfmark, sequence_order
            pgpid: optional PGP document ID that groups these fragments
        """
        if not fragments_info:
            return

        self.browse_reading_desk_active = True
        self.browse_reading_desk_pgpid = pgpid

        # Build entries from fragments info
        state = ReadingDeskState()
        state.pgpid = pgpid
        for frag in fragments_info:
            sid = frag.get('sys_id', '')
            shelfmark = frag.get('shelfmark', '')
            if not shelfmark or shelfmark == "Unknown":
                shelfmark, _ = self.meta_mgr.get_meta_for_id(sid)
                if not shelfmark or shelfmark == "Unknown":
                    shelfmark = sid

            # Get pages from searcher
            pages = self.searcher.get_full_manuscript(sid)
            page_list = []
            if pages:
                for p in pages:
                    page_list.append({
                        'p_num': p.get('p_num', 1),
                        'text': p.get('text', ''),
                        'full_header': p.get('full_header', ''),
                        'fl_id': p.get('fl_id', '')
                    })

            entry = ReadingDeskEntry(
                sys_id=sid,
                shelfmark=shelfmark,
                pages=page_list,
                sequence_order=frag.get('sequence_order', 0)
            )
            state.entries.append(entry)

        # Sort by sequence order
        state.entries.sort(key=lambda e: e.sequence_order)
        self.browse_reading_desk_state = state

        # Disable normal navigation
        self.btn_b_prev.setEnabled(False)
        self.btn_b_next.setEnabled(False)
        self.combo_browse_page.setEnabled(False)
        self.btn_b_all.setEnabled(False)

        # Show reading desk toolbar
        self.browse_rd_toolbar.show()

        # Create image scroll area ONCE (will be repopulated on each render)
        if self._browse_rd_image_scroll is None:
            self._browse_rd_image_scroll = QScrollArea()
            self._browse_rd_image_scroll.setWidgetResizable(True)
            self._browse_rd_image_scroll.setStyleSheet("background: #1a1a2e;")
            self.browse_splitter.addWidget(self._browse_rd_image_scroll)

        # Initial render with V0.8 text
        self._browse_rd_render()

        # Launch background worker to fetch PGP sources
        sys_ids = [e.sys_id for e in state.entries]
        if sys_ids:
            if self._browse_rd_worker is not None:
                try:
                    self._browse_rd_worker.finished.disconnect()
                    self._browse_rd_worker.error.disconnect()
                except (TypeError, RuntimeError):
                    pass
            self._browse_rd_worker = ReadingDeskWorker(sys_ids)
            self._browse_rd_worker.finished.connect(self._browse_rd_on_sources_loaded)
            self._browse_rd_worker.error.connect(
                lambda msg: logger.debug("ReadingDeskWorker error: %s", msg)
            )
            self._browse_rd_worker.start()

    def _browse_rd_on_sources_loaded(self, results):
        """Handle PGP sources loaded from ReadingDeskWorker."""
        if not self.browse_reading_desk_active:
            return

        # Update entries with loaded sources
        results_map = {sid: (sources, pgp_doc) for sid, sources, pgp_doc in results}
        for entry in self.browse_reading_desk_state.entries:
            if entry.sys_id in results_map:
                sources, pgp_doc = results_map[entry.sys_id]
                entry.sources = sources
                entry.pgp_doc = pgp_doc

        # Re-render with PGP data now available
        self._browse_rd_render()

    def _browse_exit_reading_desk(self):
        """Exit reading desk mode and restore normal browse view."""
        if not self.browse_reading_desk_active:
            return

        self.browse_reading_desk_active = False
        self.browse_reading_desk_state = ReadingDeskState()
        self.browse_reading_desk_pgpid = None

        # Stop worker if running
        if self._browse_rd_worker is not None:
            try:
                self._browse_rd_worker.finished.disconnect()
                self._browse_rd_worker.error.disconnect()
            except (TypeError, RuntimeError):
                pass
            self._browse_rd_worker = None

        # Clean up image widgets
        self._browse_rd_image_widgets = []

        # Restore normal view
        self._browse_rd_restore_normal_view()

        # Hide reading desk toolbar
        self.browse_rd_toolbar.hide()

        # Re-enable navigation
        self.btn_b_prev.setEnabled(True)
        self.btn_b_next.setEnabled(True)
        self.combo_browse_page.setEnabled(True)
        self.btn_b_all.setEnabled(True)

        # Reload current page normally
        if self.current_browse_sid:
            self.browse_load_page()

    def _browse_add_to_view(self):
        """Handle 'Add to View' button click -- enter reading desk or add current manuscript."""
        if not self.current_browse_sid:
            return

        sid = self.current_browse_sid
        shelfmark, _ = self.meta_mgr.get_meta_for_id(sid)
        if not shelfmark or shelfmark == "Unknown":
            shelfmark = sid

        if not self.browse_reading_desk_active:
            # Start reading desk with current manuscript
            frag_info = [{'sys_id': sid, 'shelfmark': shelfmark, 'sequence_order': 0}]
            self._browse_enter_reading_desk(frag_info)
        else:
            # Add current manuscript to existing reading desk
            self._browse_rd_add_entry(sid, shelfmark)

    def _browse_rd_add_entry(self, sys_id, shelfmark, sequence_order=None):
        """Add a single manuscript entry to the reading desk (duplicate-safe).

        Args:
            sys_id: System ID of the manuscript
            shelfmark: Display shelfmark
            sequence_order: Optional sort order (default: after last entry)
        """
        if not self.browse_reading_desk_active:
            return

        state = self.browse_reading_desk_state

        # Check for duplicates
        existing_sids = {e.sys_id for e in state.entries}
        if sys_id in existing_sids:
            return

        # Get pages from searcher
        pages = self.searcher.get_full_manuscript(sys_id)
        page_list = []
        if pages:
            for p in pages:
                page_list.append({
                    'p_num': p.get('p_num', 1),
                    'text': p.get('text', ''),
                    'full_header': p.get('full_header', ''),
                    'fl_id': p.get('fl_id', '')
                })

        if sequence_order is None:
            sequence_order = max((e.sequence_order for e in state.entries), default=0) + 1

        entry = ReadingDeskEntry(
            sys_id=sys_id,
            shelfmark=shelfmark,
            pages=page_list,
            sequence_order=sequence_order
        )
        state.entries.append(entry)
        state.entries.sort(key=lambda e: e.sequence_order)

        # Launch ReadingDeskWorker for the new entry's PGP sources
        if self._browse_rd_worker is not None:
            try:
                self._browse_rd_worker.finished.disconnect()
                self._browse_rd_worker.error.disconnect()
            except (TypeError, RuntimeError):
                pass
        self._browse_rd_worker = ReadingDeskWorker([sys_id])
        self._browse_rd_worker.finished.connect(self._browse_rd_on_sources_loaded)
        self._browse_rd_worker.error.connect(
            lambda msg: logger.debug("ReadingDeskWorker error: %s", msg)
        )
        self._browse_rd_worker.start()

        # Re-render immediately with V0.8 text (PGP will update when worker finishes)
        self._browse_rd_render()

    def _browse_rd_add_by_shelfmark(self):
        """Add a manuscript to the reading desk by shelfmark (toolbar input)."""
        text = self.browse_rd_shelf_input.text().strip()
        if not text:
            return

        # Resolve shelfmark to sys_id
        shelf_res = self.meta_mgr.resolve_system_by_shelfmark(text)
        sid = shelf_res.get('sys_id')

        if not sid and shelf_res.get('options'):
            options = shelf_res['options']
            if len(options) == 1:
                sid = options[0]['sys_id']
                text = options[0].get('shelfmark', text)
            else:
                display_options = []
                for idx, opt in enumerate(options):
                    base = opt['shelfmark']
                    title = (opt.get('title') or "").strip()
                    if title:
                        base = f"{base} | {title}"
                    label = f"{idx + 1}. {base}"
                    if len(label) > 60:
                        label = label[:57] + "..."
                    display_options.append(label)
                choice, ok = QInputDialog.getItem(
                    self, tr("Shelfmark"), tr("Multiple shelfmarks found. Select one:"),
                    display_options, 0, False
                )
                if not ok:
                    return
                if choice in display_options:
                    chosen_idx = display_options.index(choice)
                    sid = options[chosen_idx]['sys_id']
                    text = options[chosen_idx].get('shelfmark', text)

        if not sid:
            QMessageBox.warning(self, tr("Not Found"), tr("Shelfmark not found: {}").format(text))
            return

        shelfmark, _ = self.meta_mgr.get_meta_for_id(sid)
        if not shelfmark or shelfmark == "Unknown":
            shelfmark = text

        self._browse_rd_add_entry(sid, shelfmark)
        self.browse_rd_shelf_input.clear()

    def _browse_rd_add_from_list(self):
        """Show the browse lists panel so items can be added to reading desk."""
        self.browse_set_lists_panel_visible(True)

    def _browse_open_joins_in_reading_desk(self):
        """Open all joined fragments in the reading desk."""
        if not self.current_browse_sid:
            return

        document_id = self.current_browse_sid
        shelfmark = None
        if self.meta_mgr:
            try:
                shelfmark, _ = self.meta_mgr.get_meta_for_id(document_id)
            except Exception:
                pass

        if not shelfmark:
            return

        # Get joins from JoinsManager
        connected = None
        if self.joins_mgr:
            connected = self.joins_mgr.get_connected_fragments_by_id(document_id)
        if (not connected or connected.get('total_fragments', 0) <= 1) and self.joins_mgr:
            connected = self.joins_mgr.get_connected_fragments(shelfmark)

        if not connected or connected.get('total_fragments', 0) <= 1:
            QMessageBox.information(
                self, tr("Reading Desk"), tr("No joined fragments found.")
            )
            return

        # Build fragments_info from connected fragments
        fragments = connected.get('fragments', [])
        fragment_details = connected.get('fragment_details', [])
        shelfmark_to_docid = {}
        for fd in fragment_details:
            shelf = fd.get('shelfmark', '') if isinstance(fd, dict) else getattr(fd, 'shelfmark', '')
            doc_id = fd.get('document_id') if isinstance(fd, dict) else getattr(fd, 'document_id', None)
            sid = fd.get('sys_id') if isinstance(fd, dict) else getattr(fd, 'sys_id', None)
            if shelf:
                if sid:
                    shelfmark_to_docid[shelf.upper()] = sid
                elif doc_id:
                    shelfmark_to_docid[shelf.upper()] = str(doc_id)

        fragments_info = []
        for idx, frag in enumerate(fragments):
            frag_sid = shelfmark_to_docid.get(frag.upper())
            # Fallback: use _shelf_to_sys map
            if not frag_sid and self._shelf_to_sys:
                norm = self._normalize_shelfmark(frag) if hasattr(self, '_normalize_shelfmark') else None
                if norm:
                    frag_sid = self._shelf_to_sys.get(norm)
            if frag_sid:
                fragments_info.append({
                    'sys_id': frag_sid,
                    'shelfmark': frag,
                    'sequence_order': idx
                })

        if not fragments_info:
            QMessageBox.information(
                self, tr("Reading Desk"), tr("Could not resolve fragment identifiers.")
            )
            return

        # Check for PGP document context
        pgpid = None
        try:
            from shared.document_service import get_document_for_fragment
            doc_data = get_document_for_fragment(document_id)
            if doc_data:
                pgpid = doc_data.get('pgpid')
        except Exception:
            pass

        self._browse_enter_reading_desk(fragments_info, pgpid=pgpid)

    def _browse_open_pgp_joins_in_reading_desk(self):
        """Open PGP multi-fragment joined document in reading desk."""
        pgp_doc = getattr(self, '_browse_pgp_doc', {})
        if not pgp_doc:
            return

        try:
            from shared.document_service import get_fragments_for_document
            pgp_frags = get_fragments_for_document(pgp_doc.get('pgpid'))
            if not pgp_frags or len(pgp_frags) <= 1:
                return

            fragments_info = []
            for idx, frag in enumerate(pgp_frags):
                frag_sid = frag.get('sys_id', '')
                if frag_sid:
                    fragments_info.append({
                        'sys_id': frag_sid,
                        'shelfmark': frag.get('shelfmark', frag_sid),
                        'sequence_order': idx
                    })

            if fragments_info:
                self._browse_enter_reading_desk(fragments_info, pgpid=pgp_doc.get('pgpid'))
        except Exception as e:
            logger.debug("Failed to open PGP joins in reading desk: %s", e)

    def _browse_rd_render(self):
        """Render reading desk: stacked texts in text pane, stacked images in viewer pane.

        This is the key v3 method that produces the dual-pane synchronized reading desk.
        """
        state = self.browse_reading_desk_state
        if not state.entries:
            return

        # Update toolbar count label
        count = len(state.entries)
        self.browse_rd_count_label.setText(
            f" ({count} fragment{'s' if count != 1 else ''})"
        )

        # === LEFT PANE: Stacked Texts in browse_text ===
        html_parts = []

        # Header bar with exit button
        html_parts.append(
            "<div style='background: #2c3e50; color: white; padding: 8px 12px; "
            "margin-bottom: 10px; border-radius: 4px;'>"
            "<b>Reading Desk</b>"
        )
        if state.pgpid:
            html_parts.append(f" &mdash; PGP #{state.pgpid}")
        html_parts.append(
            f" ({len(state.entries)} fragment{'s' if len(state.entries) != 1 else ''})"
            "</div>"
        )

        for idx, entry in enumerate(state.entries):
            sid = entry.sys_id
            shelfmark = entry.shelfmark

            # Fragment header (clickable to navigate to single view)
            is_current = (sid == self.current_browse_sid)
            current_badge = " <span style='background: #27ae60; color: white; padding: 1px 6px; border-radius: 3px; font-size: 10px;'>Current</span>" if is_current else ""
            html_parts.append(
                f"<div id='rd-text-frag-{idx}' style='background: #ecf0f1; padding: 6px 10px; "
                f"margin-top: {'0' if idx == 0 else '20'}px; border-bottom: 2px solid #3498db;'>"
                f"<a href='genizah://rd-navigate/{sid}' style='color: #2980b9; text-decoration: none; font-weight: bold;'>"
                f"{shelfmark}</a>{current_badge}"
                f" <a href='genizah://rd-remove/{sid}' style='color: #e74c3c; text-decoration: none; font-size: 11px; margin-left: 10px;'>[remove]</a>"
            )

            # Version selector link
            if entry.sources:
                html_parts.append(
                    f" <a href='genizah://rd-version/{sid}/{idx}' style='color: #8e44ad; text-decoration: none; font-size: 11px; margin-left: 10px;'>[change version]</a>"
                )

            html_parts.append("</div>")

            # Determine text to show: prefer PGP edition, else V0.8
            display_text = ""
            text_direction = "rtl"

            if entry.sources:
                # Find first edition
                for src in entry.sources:
                    if 'Edition' in (src.get('doc_relation') or '') and src.get('content'):
                        display_text = src['content']
                        break
                # If no edition, try first translation
                if not display_text:
                    for src in entry.sources:
                        if src.get('content'):
                            display_text = src['content']
                            lang = src.get('language', '')
                            if lang == 'English':
                                text_direction = 'ltr'
                            break

            # Fallback to V0.8 text from pages
            if not display_text and entry.pages:
                page_texts = [p.get('text', '') for p in entry.pages if p.get('text')]
                display_text = '\n\n'.join(page_texts)

            if display_text:
                html_text = display_text.replace('\n', '<br>')
                html_parts.append(
                    f"<div dir='{text_direction}' style='padding: 10px; font-family: SBL Hebrew, serif; "
                    f"font-size: 16px; line-height: 1.6;'>"
                    f"{html_text}</div>"
                )
            else:
                html_parts.append(
                    "<div style='padding: 10px; color: #95a5a6; font-style: italic;'>"
                    "No text available</div>"
                )

        full_html = "\n".join(html_parts)
        self.browse_text.setHtml(full_html)

        # === RIGHT PANE: Stacked Images in viewer pane ===
        self._browse_rd_render_images()

    def _browse_rd_render_images(self):
        """Render stacked images in the viewer pane (right side of browse splitter)."""
        state = self.browse_reading_desk_state

        # Hide normal viewer
        self.browse_viewer.setVisible(False)

        # Ensure scroll area exists (created in _browse_enter_reading_desk)
        if self._browse_rd_image_scroll is None:
            return

        # Disconnect existing sync handlers before re-render
        self._browse_rd_disconnect_sync()

        # Clear existing content by replacing the container widget
        # QScrollArea.setWidget() with a new widget causes the old one to be cleaned up
        container = QWidget()
        container_layout = QVBoxLayout(container)
        container_layout.setContentsMargins(0, 0, 0, 0)
        container_layout.setSpacing(5)

        self._browse_rd_image_widgets = []

        for idx, entry in enumerate(state.entries):
            sid = entry.sys_id
            meta = self.meta_mgr.nli_cache.get(sid, {})

            # Fragment header label
            header = QLabel(f"  {entry.shelfmark}")
            header.setStyleSheet(
                "background: #34495e; color: white; padding: 6px 10px; "
                "font-weight: bold; font-size: 13px;"
            )
            header.setObjectName(f"rd-img-frag-{idx}")
            container_layout.addWidget(header)

            # Get image URLs
            images_nli = meta.get('images_nli', [])
            images_ext = meta.get('images_ext', [])
            image_list = images_ext if images_ext else images_nli

            if not image_list:
                # Try FL ID fallback
                fl_ids = meta.get('fl_ids', [])
                if isinstance(fl_ids, str):
                    fl_ids = [fl_ids]
                for fl in fl_ids or []:
                    digits = re.sub(r"\D", "", str(fl or ""))
                    if digits:
                        fallback_url = f"{Config.NLI_IIIF_BASE}/FL{digits}/full/2000,/0/default.jpg"
                        image_list = [{'label': f"FL{digits}", 'url': fallback_url}]
                        break

            if not image_list:
                no_img = QLabel(tr("No images available"))
                no_img.setStyleSheet("color: #95a5a6; padding: 20px; font-style: italic;")
                no_img.setAlignment(Qt.AlignmentFlag.AlignCenter)
                container_layout.addWidget(no_img)
                continue

            # Create a ZoomableScrollArea + controls for each image page
            for img_idx, img_data in enumerate(image_list):
                base_url = img_data.get('url', '')
                label_text = img_data.get('label', f"Page {img_idx + 1}")

                # Controls row: zoom/rotate per image
                controls = QWidget()
                controls_layout = QHBoxLayout(controls)
                controls_layout.setContentsMargins(4, 2, 4, 2)
                controls_layout.setSpacing(4)

                img_label = QLabel(f"  {label_text}")
                img_label.setStyleSheet("color: #bdc3c7; font-size: 11px;")
                controls_layout.addWidget(img_label)
                controls_layout.addStretch()

                viewer = ZoomableScrollArea()
                viewer.setMinimumHeight(400)
                viewer.setMaximumHeight(600)

                btn_zoom_out = QPushButton("-")
                btn_zoom_out.setFixedWidth(28)
                btn_zoom_out.setStyleSheet("color: white; background: #555;")
                btn_zoom_out.clicked.connect(viewer.zoom_out)
                controls_layout.addWidget(btn_zoom_out)

                btn_zoom_in = QPushButton("+")
                btn_zoom_in.setFixedWidth(28)
                btn_zoom_in.setStyleSheet("color: white; background: #555;")
                btn_zoom_in.clicked.connect(viewer.zoom_in)
                controls_layout.addWidget(btn_zoom_in)

                btn_rot_left = QPushButton("\u21BA")  # ↺
                btn_rot_left.setFixedWidth(28)
                btn_rot_left.setStyleSheet("color: white; background: #555;")
                btn_rot_left.setToolTip(tr("Rotate Left 90"))
                btn_rot_left.clicked.connect(lambda checked, v=viewer: v.rotate_view(-90))
                controls_layout.addWidget(btn_rot_left)

                btn_rot_right = QPushButton("\u21BB")  # ↻
                btn_rot_right.setFixedWidth(28)
                btn_rot_right.setStyleSheet("color: white; background: #555;")
                btn_rot_right.setToolTip(tr("Rotate Right 90"))
                btn_rot_right.clicked.connect(lambda checked, v=viewer: v.rotate_view(90))
                controls_layout.addWidget(btn_rot_right)

                container_layout.addWidget(controls)
                container_layout.addWidget(viewer)

                # Load image
                viewer.set_status_message(tr("Loading..."))
                final_url = base_url
                if final_url and not final_url.endswith('.jpg'):
                    final_url = f"{final_url}/full/2000,/0/default.jpg"

                loader = ImageLoaderThread(final_url)
                loader.image_loaded.connect(
                    lambda img, v=viewer: v.set_image(QPixmap.fromImage(img))
                )
                loader.load_failed.connect(
                    lambda v=viewer: v.set_status_message(tr("No Image"))
                )
                loader.start()

                self._browse_rd_image_widgets.append((sid, viewer, loader))

        container_layout.addStretch()
        self._browse_rd_image_scroll.setWidget(container)
        self._browse_rd_image_scroll.setVisible(True)

        # Set up synchronized scrolling between text and image panes
        self._browse_rd_setup_sync_scroll()
        # Re-establish sync after images finish loading (scroll maximums change)
        QTimer.singleShot(500, self._browse_rd_setup_sync_scroll)

    def _browse_rd_disconnect_sync(self):
        """Disconnect sync scroll handlers without affecting other signal connections."""
        if self._rd_text_sync_handler is not None:
            try:
                self.browse_text.verticalScrollBar().valueChanged.disconnect(self._rd_text_sync_handler)
            except (TypeError, RuntimeError):
                pass
            self._rd_text_sync_handler = None

        if self._rd_image_sync_handler is not None:
            try:
                if self._browse_rd_image_scroll:
                    self._browse_rd_image_scroll.verticalScrollBar().valueChanged.disconnect(self._rd_image_sync_handler)
            except (TypeError, RuntimeError):
                pass
            self._rd_image_sync_handler = None

    def _browse_rd_setup_sync_scroll(self):
        """Set up proportional scroll synchronization between text and image panes."""
        if not self._browse_rd_image_scroll:
            return

        text_bar = self.browse_text.verticalScrollBar()
        image_bar = self._browse_rd_image_scroll.verticalScrollBar()

        # Disconnect only OUR sync handlers (not all valueChanged connections)
        self._browse_rd_disconnect_sync()

        def sync_text_to_image(value):
            if self._browse_rd_syncing:
                return
            self._browse_rd_syncing = True
            try:
                if sip.isdeleted(text_bar) or sip.isdeleted(image_bar):
                    return
                text_max = text_bar.maximum()
                image_max = image_bar.maximum()
                if text_max > 0 and image_max > 0:
                    ratio = value / text_max
                    image_bar.setValue(int(ratio * image_max))
            finally:
                self._browse_rd_syncing = False

        def sync_image_to_text(value):
            if self._browse_rd_syncing:
                return
            self._browse_rd_syncing = True
            try:
                if sip.isdeleted(text_bar) or sip.isdeleted(image_bar):
                    return
                text_max = text_bar.maximum()
                image_max = image_bar.maximum()
                if text_max > 0 and image_max > 0:
                    ratio = value / image_max
                    text_bar.setValue(int(ratio * text_max))
            finally:
                self._browse_rd_syncing = False

        # Store references for targeted disconnect
        self._rd_text_sync_handler = sync_text_to_image
        self._rd_image_sync_handler = sync_image_to_text

        text_bar.valueChanged.connect(self._rd_text_sync_handler)
        image_bar.valueChanged.connect(self._rd_image_sync_handler)

    def _browse_rd_restore_normal_view(self):
        """Hide reading desk image scroll and restore normal viewer."""
        if self._browse_rd_image_scroll is not None:
            self._browse_rd_disconnect_sync()
            # Remove from splitter and destroy
            self._browse_rd_image_scroll.setParent(None)
            self._browse_rd_image_scroll.deleteLater()
            self._browse_rd_image_scroll = None

        # Show normal viewer
        self.browse_viewer.setVisible(self.btn_b_toggle_img.isChecked())

    def _browse_rd_remove_entry(self, sys_id):
        """Remove a fragment entry from the reading desk and re-render or exit."""
        state = self.browse_reading_desk_state
        state.entries = [e for e in state.entries if e.sys_id != sys_id]

        if not state.entries:
            self._browse_exit_reading_desk()
        else:
            self._browse_rd_render()

    def _browse_rd_show_version_dialog(self, sys_id, entry_idx):
        """Show a dialog to select PGP version source for a specific fragment."""
        state = self.browse_reading_desk_state
        if entry_idx < 0 or entry_idx >= len(state.entries):
            return

        entry = state.entries[entry_idx]
        if not entry.sources:
            QMessageBox.information(
                self, tr("Version Selector"),
                tr("No PGP sources available for this fragment.")
            )
            return

        # Build options list
        options = []
        for src in entry.sources:
            relation = src.get('doc_relation', '')
            scholar = src.get('source_scholar', 'Unknown')
            language = src.get('language', '')
            if 'Edition' in relation:
                label = f"Edition: {scholar}"
            elif 'Translation' in relation:
                label = f"Translation ({language}): {scholar}" if language else f"Translation: {scholar}"
            else:
                label = f"{relation}: {scholar}"
            options.append(label)

        # Add V0.8 fallback
        options.append("V0.8 (HTR)")

        choice, ok = QInputDialog.getItem(
            self,
            tr("Select Version"),
            tr("Choose a text version for {}:").format(entry.shelfmark),
            options, 0, False
        )
        if not ok:
            return

        selected_idx = options.index(choice) if choice in options else -1
        if selected_idx < 0:
            return

        # Update the entry's displayed text
        if selected_idx < len(entry.sources):
            # Selected a PGP source -- reorder sources so selected is first
            selected_src = entry.sources[selected_idx]
            entry.sources.remove(selected_src)
            entry.sources.insert(0, selected_src)
        else:
            # V0.8 selected -- clear sources so fallback text is used
            entry.sources = []

        # Re-render
        self._browse_rd_render()

    def toggle_browse_view_all(self, checked):
        if checked:
            self.browse_viewer.setVisible(False)
            self.btn_b_toggle_img.setEnabled(False)
            self.browse_load_all()
        else:
            self.btn_b_toggle_img.setEnabled(True)
            self.browse_viewer.setVisible(self.btn_b_toggle_img.isChecked())
            self.browse_load_page()

    def on_browse_page_combo_changed(self, index):
        if index < 0: return
        if not self.current_browse_sid: return

        # We assume index in combo maps 1:1 to absolute_index
        page_data = self.searcher.get_browse_page(
            self.current_browse_sid,
            absolute_index=index,
            next_prev=0
        )

        if page_data:
            self.browse_render_page(page_data)
            # Re-fetch PGP sources for the new page (recto/verso may differ)
            self._browse_refresh_pgp_for_page()

    def toggle_browse_image(self):
        visible = self.btn_b_toggle_img.isChecked()
        self.browse_viewer.setVisible(visible)

    def browse_search_parallels(self):
        if not self.current_browse_sid: return

        # Get Title
        meta = self.meta_mgr.nli_cache.get(self.current_browse_sid, {})
        full_title = meta.get('title') or ""
        # Truncate title
        words = full_title.split()
        if len(words) > 6:
            short_title = " ".join(words[:6]) + "..."
        else:
            short_title = full_title

        text_content = ""

        if self.btn_b_all.isChecked():
             # Full Text
             pages = self.searcher.get_full_manuscript(self.current_browse_sid)
             if pages:
                 text_content = "\n\n".join([p['text'] for p in pages])
        else:
             # Current Page
             if self.current_browse_internal_idx is not None:
                 pd = self.searcher.get_browse_page(self.current_browse_sid, absolute_index=self.current_browse_internal_idx)
                 if pd: text_content = pd['text']

        if text_content:
            self.send_result_to_composition(
                {'display': {'id': self.current_browse_sid}},
                source_text=text_content,
                title=short_title
            )

    def browse_add_to_list(self):
        """Add current manuscript to a list."""
        if not self.current_browse_sid or not self.lists_mgr:
            return

        fl_id = self._normalize_fl_id(self.browse_fl_input.text().strip())
        img = self.current_browse_p
        self.show_add_to_list_menu(
            [{'sys_id': self.current_browse_sid, 'fl_id': fl_id, 'img': img}],
            source=tr("from browse"),
            anchor_widget=self.btn_browse_add_to_list
        )
        self._update_browse_add_to_list_button()

    def _set_last_browse_field(self, field):
        self.last_browse_field = field

    def browse_load_page(self):
        """Load single page text and sync viewer."""
        if not self.current_browse_sid: return

        p = self.current_browse_p or 1
        page_data = self.searcher.get_browse_page(self.current_browse_sid, p_num=p)

        if not page_data:
            self.browse_text.setText(tr("Page not found."))
            return

        self.browse_render_page(page_data)

    def _apply_browse_highlights(self, text, uid):
        if not text or not uid:
            return text
        spans = []
        for ph in self.browse_highlight_data or []:
            if ph.get('uid') == uid and ph.get('span'):
                span = ph.get('span')
                if span and len(span) == 2:
                    spans.append(tuple(span))
        if not spans:
            return text
        spans.sort(key=lambda s: s[0], reverse=True)
        for s, e in spans:
            if s < 0 or e > len(text) or s >= e:
                continue
            text = text[:e] + "*" + text[e:]
            text = text[:s] + "*" + text[s:]
        return text

    def browse_load_all(self):
        """Load all pages into the text browser for continuous scrolling."""
        if not self.current_browse_sid: return

        self.browse_text.setText(tr("Loading full manuscript..."))
        QApplication.processEvents() # Refresh UI

        html_content = []

        # If browsing a Part, load all folios in the Part
        if self.current_browse_part_id and self.current_browse_part_folios:
            # Get Oxford images for image labels
            part_images = self.meta_mgr.get_part_images(self.current_browse_part_id)
            image_idx = 0
            part_display = self.meta_mgr.codico_mgr.get_part_display_name(self.current_browse_part_id)

            for folio_idx, sid in enumerate(self.current_browse_part_folios):
                pages = self.searcher.get_full_manuscript(sid)
                if not pages:
                    continue

                # Get shelfmark for linking
                shelf, _ = self.meta_mgr.get_meta_for_id(sid)
                if not shelf or shelf == "Unknown":
                    shelf = sid
                shelf_display = shelf
                if part_display:
                    shelf_display = f"{part_display} | {shelf}"

                for p in pages:
                    # Get image label from Oxford images
                    img_label = ""
                    if part_images and image_idx < len(part_images):
                        img_label = part_images[image_idx].get('label', '')
                        image_idx += 1

                    # Anchor for scrolling
                    anchor = f'<a name="img_{image_idx}"></a>'

                    # Create clickable link separator using image label
                    # The link will load this specific shelfmark when clicked
                    link_text = f"image {img_label}" if img_label else f"image {image_idx}"
                    # Use custom URL scheme for internal navigation
                    link_href = f"genizah://load/{sid}"

                    separator = f"""
                    <div style='background-color: #f5f5f5; color: #333; padding: 6px 10px; margin-top: 20px; border-bottom: 1px solid #ddd;'>
                        <a href="{link_href}" style='color: #2980b9; text-decoration: none; font-weight: bold;'>
                            {link_text}
                        </a>
                        <span style='font-size: 0.85em; color: #777; margin-left: 10px;'>{shelf_display}</span>
                    </div>
                    """

                    # Content with line breaks preserved
                    content = p['text'].replace("\n", "<br>")

                    html_content.append(anchor + separator + f"<div dir='rtl'>{content}</div>")
        else:
            # Single folio mode (original behavior)
            pages = self.searcher.get_full_manuscript(self.current_browse_sid)
            if not pages:
                QMessageBox.warning(self, tr("Error"), tr("Could not load full text."))
                return

            # Get enriched map if available
            meta = self.meta_mgr.nli_cache.get(self.current_browse_sid, {})
            canvas_map = meta.get('canvas_map', {})

            for p in pages:
                # Anchor for scrolling
                anchor = f'<a name="page_{p["p_num"]}"></a>'

                # Visual Separator
                img_lbl = tr("Image")
                fl_id = p.get('fl_id')

                # Resolve Label
                fl_digits = re.sub(r"\D", "", str(fl_id or ""))
                label = canvas_map.get(fl_digits, "")

                fl_suffix = f" ({tr('FL')}: {fl_id})" if fl_id else ""
                label_suffix = f" - {label}" if label else ""

                separator = f"""
                <div style='background-color: #f0f0f0; color: #555; padding: 5px; margin-top: 20px; border-bottom: 2px solid #ccc;'>
                    <b>{img_lbl}: {p['p_num']}{label_suffix}</b> <span style='font-size:0.8em'>{fl_suffix}</span>
                </div>
                """

                # Content with line breaks preserved
                content = p['text'].replace("\n", "<br>")

                html_content.append(anchor + separator + f"<div dir='rtl'>{content}</div>")

        if not html_content:
            QMessageBox.warning(self, tr("Error"), tr("Could not load full text."))
            return

        full_html = "".join(html_content)
        self.browse_text.setHtml(full_html)
        apply_find_highlight(self.browse_text, self.browse_find_input.text().strip())

        # Update info label with Oxford Part info in View All mode
        part_id = self.current_browse_part_id
        if not part_id and self.current_browse_sid:
            part_id = self.meta_mgr.get_part_for_folio(self.current_browse_sid)

        if part_id:
            part_meta = self.meta_mgr.get_part_metadata(part_id)
            oxford_title = part_meta.get('title', '') if part_meta else ''
            folio_range = part_meta.get('folio_range', []) if part_meta else []

            # Get shelfmark for first folio in part
            shelf = ""
            first_sid = None
            if self.current_browse_part_folios:
                first_sid = self.current_browse_part_folios[0]
                shelf, _ = self.meta_mgr.get_meta_for_id(first_sid)
            if not shelf or shelf == "Unknown":
                shelf = part_id

            # Add library prefix
            sid_for_lib = first_sid or self.current_browse_sid
            if sid_for_lib:
                library_code = self.meta_mgr.get_library_for_id(sid_for_lib)
                if library_code:
                    library = get_library_display(library_code, short=False)
                    shelf = f"{library} | {shelf}"

            # Extract part number from part_id (e.g., "MS. Heb. b. 10/43" -> "43")
            part_num = part_id.split('/')[-1] if '/' in part_id else ''

            # Build combined: "MS heb. b. 10/79 (part 43: fols. 79-82) - View All"
            shelf_with_part = f"{shelf}"
            if part_num:
                shelf_with_part += f" (part {part_num}"
                if len(folio_range) == 2:
                    if folio_range[0] == folio_range[1]:
                        shelf_with_part += f": fol. {folio_range[0]}"
                    else:
                        shelf_with_part += f": fols. {folio_range[0]}–{folio_range[1]}"
                shelf_with_part += ")"

            info_text = f"<b>{shelf_with_part}</b> - View All"
            tooltip_text = ''
            if oxford_title:
                truncated, full = _truncate_title(oxford_title)
                info_text += f"<br/><span style='font-size: 11px;'>{truncated}</span>"
                if full: tooltip_text = full
            self.browse_info_lbl.setText(info_text)
            self.browse_info_lbl.setToolTip(tooltip_text)

        # Disable paging buttons since we are showing everything
        self.btn_b_prev.setEnabled(False)
        self.btn_b_next.setEnabled(False)
        self.combo_browse_page.setEnabled(False)

        # Scroll to the page we were looking at
        if self.current_browse_p:
            self.browse_text.scrollToAnchor(f"page_{self.current_browse_p}")

    def browse_save_full(self):
        if not self.current_browse_sid: return
        
        # Determine default filename from shelfmark if available
        meta = self.meta_mgr.nli_cache.get(self.current_browse_sid, {})
        shelfmark = meta.get('shelfmark')

        if shelfmark and shelfmark != "Unknown":
            # Logic: Remove "Ms."/"Ms" prefix unless followed only by a number
            # 1. Match Ms prefix
            ms_match = re.match(r'^\s*ms\.?\s*(.*)', shelfmark, re.IGNORECASE)
            if ms_match:
                remainder = ms_match.group(1)
                # If remainder is NOT just digits (e.g. "T-S ...", "Or. ..."), use remainder
                if not re.fullmatch(r'\d+', remainder.strip()):
                    shelfmark = remainder.strip()

            # Sanitize filename: remove illegal chars, preserve dots, convert spaces to underscores
            safe_shelf = re.sub(r'[<>:"/\\|?*]', '', shelfmark)
            safe_shelf = re.sub(r'\s+', '_', safe_shelf).strip('_')
            default_name = f"{safe_shelf}.txt"
        else:
            default_name = f"Manuscript_{self.current_browse_sid}.txt"

        path, _ = QFileDialog.getSaveFileName(self, tr("Save Manuscript"),
                                            os.path.join(Config.REPORTS_DIR, default_name), 
                                            "Text (*.txt)")
        if not path: return
        
        pages = self.searcher.get_full_manuscript(self.current_browse_sid)
        if not pages: return
        
        with open(path, 'w', encoding='utf-8') as f:
            # Header
            f.write(self._get_credit_header())
            f.write(f"System ID: {self.current_browse_sid}\n")
            # Add Library field
            library_code = self.meta_mgr.get_library_for_id(self.current_browse_sid)
            library_name = get_library_display(library_code, short=False) if library_code else ''
            f.write(f"Library: {library_name}\n")
            f.write(f"Shelfmark: {meta.get('shelfmark', 'Unknown')}\n")
            f.write(f"Title: {meta.get('title', 'Unknown')}\n")
            f.write("="*50 + "\n\n")
            
            for p in pages:
                f.write(f"--- Page {p['p_num']} ---\n")
                f.write(p['text'])
                f.write("\n\n")
            lab_config = self._get_lab_config_block()
            if lab_config:
                f.write(lab_config)
        
        QMessageBox.information(self, tr("Saved"), tr("Manuscript saved to:\n{}").format(path))

    # ==========================================================================
    #  CATALOG BROWSE TAB (Browse by Identification)
    # ==========================================================================

    def create_catalog_browse_tab(self):
        """Create the 'Browse by Identification' tab with domain tree, author/work search, and results table."""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(5, 5, 5, 5)

        # --- Internal state ---
        self._catalog_current_domain = None
        self._catalog_current_author = None
        self._catalog_current_work = None
        self._catalog_current_page = 0
        self._catalog_authors_cache = []
        self._catalog_works_cache = []
        self._catalog_tree_loaded = False
        self._CATALOG_PAGE_SIZE = 50

        # --- Top Controls Row: Active Filters + Result Count ---
        top_row = QHBoxLayout()
        self._catalog_chips_layout = QHBoxLayout()
        self._catalog_chips_layout.setSpacing(4)
        top_row.addLayout(self._catalog_chips_layout)

        self._catalog_clear_all_btn = QPushButton(tr("Clear All"))
        self._catalog_clear_all_btn.setFixedHeight(26)
        self._catalog_clear_all_btn.setStyleSheet(
            "QPushButton { background: #e74c3c; color: white; border-radius: 3px; padding: 2px 8px; font-size: 11px; }"
            "QPushButton:hover { background: #c0392b; }"
        )
        self._catalog_clear_all_btn.clicked.connect(lambda: self._catalog_remove_filter("all"))
        self._catalog_clear_all_btn.setVisible(False)
        top_row.addWidget(self._catalog_clear_all_btn)

        top_row.addStretch()

        self._catalog_count_label = QLabel("")
        self._catalog_count_label.setStyleSheet("color: #555; font-size: 12px;")
        top_row.addWidget(self._catalog_count_label)

        layout.addLayout(top_row)

        # --- Main Splitter ---
        splitter = QSplitter(Qt.Orientation.Horizontal)

        # LEFT PANEL: Domain Tree + Author Search + Work Search
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        left_layout.setContentsMargins(0, 0, 5, 0)

        # a) Domain Tree
        domain_label = QLabel(tr("Domain"))
        domain_label.setStyleSheet("font-weight: bold; font-size: 13px; margin-bottom: 2px;")
        left_layout.addWidget(domain_label)

        self.catalog_domain_tree = QTreeWidget()
        self.catalog_domain_tree.setHeaderLabels([tr("Domain"), "#"])
        self.catalog_domain_tree.setColumnWidth(0, 220)
        self.catalog_domain_tree.setAlternatingRowColors(True)
        self.catalog_domain_tree.setSelectionMode(QTreeWidget.SelectionMode.SingleSelection)
        self.catalog_domain_tree.itemClicked.connect(self._catalog_on_domain_select)
        self.catalog_domain_tree.setStyleSheet(
            "QTreeWidget { font-size: 12px; }"
            "QTreeWidget::item { padding: 2px 0; }"
            "QTreeWidget::item:selected { background-color: #3498db; color: white; }"
        )
        left_layout.addWidget(self.catalog_domain_tree, 3)

        # b) Author Search
        author_label = QLabel(tr("Author"))
        author_label.setStyleSheet("font-weight: bold; font-size: 13px; margin-top: 6px; margin-bottom: 2px;")
        left_layout.addWidget(author_label)

        self.catalog_author_input = QLineEdit()
        self.catalog_author_input.setPlaceholderText(tr("Search authors..."))
        left_layout.addWidget(self.catalog_author_input)

        self.catalog_author_list = QListWidget()
        self.catalog_author_list.setAlternatingRowColors(True)
        self.catalog_author_list.setStyleSheet("QListWidget { font-size: 12px; }")
        self.catalog_author_list.itemClicked.connect(self._catalog_on_author_select)
        left_layout.addWidget(self.catalog_author_list, 2)

        # Debounce timer for author search
        self._catalog_author_timer = QTimer()
        self._catalog_author_timer.setSingleShot(True)
        self._catalog_author_timer.setInterval(300)
        self._catalog_author_timer.timeout.connect(self._catalog_filter_authors)
        self.catalog_author_input.textChanged.connect(lambda _: self._catalog_author_timer.start())

        # c) Work Search
        work_label = QLabel(tr("Work / Title"))
        work_label.setStyleSheet("font-weight: bold; font-size: 13px; margin-top: 6px; margin-bottom: 2px;")
        left_layout.addWidget(work_label)

        self.catalog_work_input = QLineEdit()
        self.catalog_work_input.setPlaceholderText(tr("Search works..."))
        left_layout.addWidget(self.catalog_work_input)

        self.catalog_work_list = QListWidget()
        self.catalog_work_list.setAlternatingRowColors(True)
        self.catalog_work_list.setStyleSheet("QListWidget { font-size: 12px; }")
        self.catalog_work_list.itemClicked.connect(self._catalog_on_work_select)
        left_layout.addWidget(self.catalog_work_list, 2)

        # Debounce timer for work search
        self._catalog_work_timer = QTimer()
        self._catalog_work_timer.setSingleShot(True)
        self._catalog_work_timer.setInterval(300)
        self._catalog_work_timer.timeout.connect(self._catalog_filter_works)
        self.catalog_work_input.textChanged.connect(lambda _: self._catalog_work_timer.start())

        left_panel.setMinimumWidth(280)
        left_panel.setMaximumWidth(400)
        splitter.addWidget(left_panel)

        # RIGHT PANEL: Results Table + Pagination
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(5, 0, 0, 0)

        # d) Results Table
        self.catalog_results_table = QTableWidget()
        self.catalog_results_table.setColumnCount(5)
        headers = [tr("Shelfmark"), tr("Library"), tr("Domain"), tr("Identification"), tr("Date")]
        self.catalog_results_table.setHorizontalHeaderLabels(headers)
        self.catalog_results_table.horizontalHeader().setStretchLastSection(True)
        self.catalog_results_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Interactive)
        self.catalog_results_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.catalog_results_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Interactive)
        self.catalog_results_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)
        self.catalog_results_table.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeMode.ResizeToContents)
        self.catalog_results_table.setColumnWidth(0, 180)
        self.catalog_results_table.setColumnWidth(2, 150)
        self.catalog_results_table.setAlternatingRowColors(True)
        self.catalog_results_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.catalog_results_table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.catalog_results_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.catalog_results_table.verticalHeader().setVisible(False)
        self.catalog_results_table.doubleClicked.connect(self._catalog_open_manuscript)
        self.catalog_results_table.setStyleSheet(
            "QTableWidget { font-size: 12px; }"
            "QTableWidget::item { padding: 3px; }"
        )
        right_layout.addWidget(self.catalog_results_table)

        # e) Pagination
        pagination_layout = QHBoxLayout()
        pagination_layout.addStretch()

        self._catalog_prev_btn = QPushButton(tr("Previous"))
        self._catalog_prev_btn.setEnabled(False)
        self._catalog_prev_btn.clicked.connect(self._catalog_prev_page)
        pagination_layout.addWidget(self._catalog_prev_btn)

        self._catalog_page_label = QLabel("")
        self._catalog_page_label.setStyleSheet("margin: 0 10px; font-size: 12px;")
        pagination_layout.addWidget(self._catalog_page_label)

        self._catalog_next_btn = QPushButton(tr("Next"))
        self._catalog_next_btn.setEnabled(False)
        self._catalog_next_btn.clicked.connect(self._catalog_next_page)
        pagination_layout.addWidget(self._catalog_next_btn)

        pagination_layout.addStretch()
        right_layout.addLayout(pagination_layout)

        splitter.addWidget(right_panel)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 3)

        layout.addWidget(splitter)
        return panel

    # --- Catalog Browse: Refresh & Data Methods ---

    def _catalog_refresh(self):
        """Main refresh: re-fetch results with current filters + pagination, update UI."""
        from shared.fjms_service import get_fjms_service
        fjms = get_fjms_service()
        if not fjms.is_available():
            return

        offset = self._catalog_current_page * self._CATALOG_PAGE_SIZE
        data = fjms.get_browse_results(
            domain=self._catalog_current_domain,
            author=self._catalog_current_author,
            work=self._catalog_current_work,
            offset=offset,
            limit=self._CATALOG_PAGE_SIZE,
        )

        results = data.get("results", [])
        total = data.get("total", 0)

        # Update results table
        self.catalog_results_table.setRowCount(len(results))
        for row_idx, r in enumerate(results):
            sys_id = r.get("sys_id", "")
            # Resolve shelfmark/library from metadata manager
            shelfmark = "Unknown"
            library = ""
            if self.meta_mgr:
                shelfmark, _ = self.meta_mgr.get_meta_for_id(sys_id)
                library = self.meta_mgr.get_library_for_id(sys_id)

            # Domain: pick Hebrew/English based on current language
            domains = r.get("domains_heb", []) if CURRENT_LANG == 'he' else r.get("domains", [])
            domain_str = ", ".join(domains) if domains else ""

            # Identification: Author - Title
            author = r.get("author", "")
            title = r.get("title_heb", "") if CURRENT_LANG == 'he' else r.get("title", "")
            if not title:
                title = r.get("title", "") or r.get("title_heb", "")
            ident = f"{author} - {title}" if author and title else (author or title or "")

            date_str = r.get("copy_date", "")

            item_shelf = QTableWidgetItem(shelfmark)
            item_shelf.setData(Qt.ItemDataRole.UserRole, sys_id)  # Store sys_id for navigation
            self.catalog_results_table.setItem(row_idx, 0, item_shelf)
            self.catalog_results_table.setItem(row_idx, 1, QTableWidgetItem(library))
            self.catalog_results_table.setItem(row_idx, 2, QTableWidgetItem(domain_str))
            self.catalog_results_table.setItem(row_idx, 3, QTableWidgetItem(ident))
            self.catalog_results_table.setItem(row_idx, 4, QTableWidgetItem(date_str))

        # Update count label
        if total > 0:
            start = offset + 1
            end = min(offset + self._CATALOG_PAGE_SIZE, total)
            self._catalog_count_label.setText(
                tr("Showing {start}-{end} of {total} manuscripts").format(start=start, end=end, total=f"{total:,}")
            )
        else:
            self._catalog_count_label.setText(tr("No results"))

        # Update pagination
        total_pages = max(1, (total + self._CATALOG_PAGE_SIZE - 1) // self._CATALOG_PAGE_SIZE)
        current_page_display = self._catalog_current_page + 1
        self._catalog_page_label.setText(
            tr("Page {current} of {total}").format(current=current_page_display, total=total_pages)
        )
        self._catalog_prev_btn.setEnabled(self._catalog_current_page > 0)
        self._catalog_next_btn.setEnabled(current_page_display < total_pages)

        # Update chips bar
        self._catalog_update_chips()

    def _catalog_refresh_authors(self):
        """Fetch authors scoped to current domain, update author list widget."""
        from shared.fjms_service import get_fjms_service
        fjms = get_fjms_service()
        if not fjms.is_available():
            return
        self._catalog_authors_cache = fjms.get_browse_authors(domain=self._catalog_current_domain)
        self._catalog_filter_authors()

    def _catalog_filter_authors(self):
        """Filter author list widget based on current text input."""
        search_text = self.catalog_author_input.text().strip().lower()
        self.catalog_author_list.clear()
        shown = 0
        for entry in self._catalog_authors_cache:
            author = entry["author"]
            count = entry["count"]
            if search_text and search_text not in author.lower():
                continue
            self.catalog_author_list.addItem(f"{author}  ({count:,})")
            shown += 1
            if shown >= 50:
                break

    def _catalog_refresh_works(self):
        """Fetch works scoped to current domain + author, update works list widget."""
        from shared.fjms_service import get_fjms_service
        fjms = get_fjms_service()
        if not fjms.is_available():
            return
        self._catalog_works_cache = fjms.get_browse_works(
            domain=self._catalog_current_domain,
            author=self._catalog_current_author,
        )
        self._catalog_filter_works()

    def _catalog_filter_works(self):
        """Filter works list widget based on current text input."""
        search_text = self.catalog_work_input.text().strip().lower()
        self.catalog_work_list.clear()
        shown = 0
        for entry in self._catalog_works_cache:
            title = entry.get("title", "")
            title_heb = entry.get("title_heb", "")
            count = entry["count"]
            display = title_heb if CURRENT_LANG == 'he' and title_heb else title
            if not display:
                display = title or title_heb or "?"
            if search_text and search_text not in title.lower() and search_text not in title_heb.lower():
                continue
            item = QListWidgetItem(f"{display}  ({count:,})")
            item.setData(Qt.ItemDataRole.UserRole, title)  # Store English title for query
            self.catalog_work_list.addItem(item)
            shown += 1
            if shown >= 50:
                break

    # --- Catalog Browse: Selection Handlers ---

    def _catalog_on_domain_select(self, item):
        """Handle domain tree item click."""
        domain_key = item.data(0, Qt.ItemDataRole.UserRole)
        if domain_key == "__unclassified__":
            # Unclassified bucket - not a real domain filter for now
            # (Would need special query logic; skip for initial release)
            return
        self._catalog_current_domain = domain_key
        self._catalog_current_author = None
        self._catalog_current_work = None
        self._catalog_current_page = 0
        self.catalog_author_input.clear()
        self.catalog_work_input.clear()
        self._catalog_refresh_authors()
        self._catalog_refresh_works()
        self._catalog_refresh()

    def _catalog_on_author_select(self, item):
        """Handle author list item click."""
        text = item.text()
        # Parse author name from "Author Name  (123)" format
        author = text.rsplit("  (", 1)[0].strip()
        self._catalog_current_author = author
        self._catalog_current_work = None
        self._catalog_current_page = 0
        self.catalog_work_input.clear()
        self._catalog_refresh_works()
        self._catalog_refresh()

    def _catalog_on_work_select(self, item):
        """Handle work list item click."""
        # Get the English title stored in UserRole for the query
        work_title = item.data(Qt.ItemDataRole.UserRole)
        if not work_title:
            text = item.text()
            work_title = text.rsplit("  (", 1)[0].strip()
        self._catalog_current_work = work_title
        self._catalog_current_page = 0
        self._catalog_refresh()

    def _catalog_remove_filter(self, filter_type):
        """Remove a specific filter (or all) and refresh."""
        if filter_type == "all":
            self._catalog_current_domain = None
            self._catalog_current_author = None
            self._catalog_current_work = None
            self.catalog_domain_tree.clearSelection()
            self.catalog_author_input.clear()
            self.catalog_work_input.clear()
            self._catalog_refresh_authors()
            self._catalog_refresh_works()
        elif filter_type == "domain":
            self._catalog_current_domain = None
            self._catalog_current_author = None
            self._catalog_current_work = None
            self.catalog_domain_tree.clearSelection()
            self.catalog_author_input.clear()
            self.catalog_work_input.clear()
            self._catalog_refresh_authors()
            self._catalog_refresh_works()
        elif filter_type == "author":
            self._catalog_current_author = None
            self._catalog_current_work = None
            self.catalog_work_input.clear()
            self._catalog_refresh_works()
        elif filter_type == "work":
            self._catalog_current_work = None
        self._catalog_current_page = 0
        self._catalog_refresh()

    def _catalog_update_chips(self):
        """Update the active filter chips bar."""
        # Clear existing chips
        while self._catalog_chips_layout.count():
            child = self._catalog_chips_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        has_any = False
        chip_style = (
            "QPushButton { background: #3498db; color: white; border-radius: 12px; "
            "padding: 3px 10px; font-size: 11px; border: none; }"
            "QPushButton:hover { background: #2980b9; }"
        )

        if self._catalog_current_domain:
            has_any = True
            btn = QPushButton(f"{tr('Domain')}: {self._catalog_current_domain}  \u00d7")
            btn.setStyleSheet(chip_style)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            btn.setFixedHeight(24)
            btn.clicked.connect(lambda: self._catalog_remove_filter("domain"))
            self._catalog_chips_layout.addWidget(btn)

        if self._catalog_current_author:
            has_any = True
            display = self._catalog_current_author
            if len(display) > 30:
                display = display[:27] + "..."
            btn = QPushButton(f"{tr('Author')}: {display}  \u00d7")
            btn.setStyleSheet(chip_style)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            btn.setFixedHeight(24)
            btn.clicked.connect(lambda: self._catalog_remove_filter("author"))
            self._catalog_chips_layout.addWidget(btn)

        if self._catalog_current_work:
            has_any = True
            display = self._catalog_current_work
            if len(display) > 30:
                display = display[:27] + "..."
            btn = QPushButton(f"{tr('Work / Title')}: {display}  \u00d7")
            btn.setStyleSheet(chip_style)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            btn.setFixedHeight(24)
            btn.clicked.connect(lambda: self._catalog_remove_filter("work"))
            self._catalog_chips_layout.addWidget(btn)

        self._catalog_clear_all_btn.setVisible(has_any)

    def _catalog_open_manuscript(self, index):
        """Double-click result row: navigate to Browse by Shelfmark tab."""
        row = index.row()
        item = self.catalog_results_table.item(row, 0)
        if not item:
            return
        sys_id = item.data(Qt.ItemDataRole.UserRole)
        if sys_id:
            self.tabs.setCurrentWidget(self.browse_tab)
            self.browse_sys_input.setText(str(sys_id))
            self.browse_load()

    def _catalog_next_page(self):
        """Go to next page of results."""
        self._catalog_current_page += 1
        self._catalog_refresh()

    def _catalog_prev_page(self):
        """Go to previous page of results."""
        if self._catalog_current_page > 0:
            self._catalog_current_page -= 1
            self._catalog_refresh()

    # --- Catalog Browse: Tree Population (lazy) ---

    def _catalog_populate_tree(self):
        """Populate the domain tree from FjmsService hierarchy. Called on first tab activation."""
        from shared.fjms_service import get_fjms_service
        fjms = get_fjms_service()
        if not fjms.is_available():
            return

        hierarchy = fjms.get_domain_hierarchy()
        self.catalog_domain_tree.clear()

        # Sort parents by count descending
        sorted_parents = sorted(hierarchy.items(), key=lambda x: x[1].get('count', 0), reverse=True)

        expanded_count = 0
        for parent_name, info in sorted_parents:
            parent_heb = info.get('parent_domain_heb', parent_name)
            parent_count = info.get('count', 0)
            display_name = parent_heb if CURRENT_LANG == 'he' else parent_name

            parent_item = QTreeWidgetItem([display_name, f"{parent_count:,}"])
            parent_item.setData(0, Qt.ItemDataRole.UserRole, parent_name)
            self.catalog_domain_tree.addTopLevelItem(parent_item)

            # Sort children by count descending
            children = sorted(info.get('children', []), key=lambda c: c.get('count', 0), reverse=True)
            for child in children:
                child_name = child.get('domain', '')
                child_heb = child.get('domain_heb', child_name)
                child_count = child.get('count', 0)
                child_display = child_heb if CURRENT_LANG == 'he' else child_name

                child_item = QTreeWidgetItem([child_display, f"{child_count:,}"])
                child_item.setData(0, Qt.ItemDataRole.UserRole, child_name)
                parent_item.addChild(child_item)

            # Expand top 3 parent domains by default
            if expanded_count < 3:
                parent_item.setExpanded(True)
                expanded_count += 1

        # Add "Unclassified" item at bottom
        try:
            unclassified_count = fjms.get_unclassified_count()
        except Exception:
            unclassified_count = 0

        if unclassified_count > 0:
            uncat_display = tr("Unclassified")
            uncat_item = QTreeWidgetItem([uncat_display, f"{unclassified_count:,}"])
            uncat_item.setData(0, Qt.ItemDataRole.UserRole, "__unclassified__")
            self.catalog_domain_tree.addTopLevelItem(uncat_item)

        self._catalog_tree_loaded = True

        # Also pre-load author and work lists (unfiltered)
        self._catalog_refresh_authors()
        self._catalog_refresh_works()

    # ==========================================================================
    #  PERSONAL LISTS TAB
    # ==========================================================================

    def create_lists_tab(self):
        """Create the Personal Lists tab for managing starred manuscripts."""
        panel = QWidget()
        layout = QHBoxLayout(panel)
        layout.setContentsMargins(5, 5, 5, 5)

        # Main splitter for preview panel and content
        main_splitter = QSplitter(Qt.Orientation.Horizontal)

        # --- Preview Panel (Left/Optional) ---
        self.lists_preview_panel = QWidget()
        preview_layout = QVBoxLayout(self.lists_preview_panel)
        preview_layout.setContentsMargins(5, 5, 5, 5)
        preview_layout.setSpacing(5)

        # Preview header with collapse button
        preview_header = QHBoxLayout()
        self.lists_preview_title = QLabel(f"<b>{tr('Preview')}</b>")
        preview_header.addWidget(self.lists_preview_title)
        preview_header.addStretch()
        self.btn_toggle_preview = QPushButton(tr("◀"))
        self.btn_toggle_preview.setFixedSize(24, 24)
        self.btn_toggle_preview.setToolTip(tr("Toggle Preview Panel"))
        self.btn_toggle_preview.clicked.connect(self.lists_toggle_preview)
        preview_header.addWidget(self.btn_toggle_preview)
        preview_layout.addLayout(preview_header)

        self.lists_preview_contents = QWidget()
        preview_contents_layout = QVBoxLayout(self.lists_preview_contents)
        preview_contents_layout.setContentsMargins(0, 0, 0, 0)
        preview_contents_layout.setSpacing(5)

        # Text preview area
        self.lists_preview_text = QTextEdit()
        self.lists_preview_text.setReadOnly(True)
        self.lists_preview_text.setPlaceholderText(tr("Select an item to preview"))
        preview_contents_layout.addWidget(self.lists_preview_text, 2)

        # Image area
        self.lists_preview_image = QLabel()
        self.lists_preview_image.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lists_preview_image.setMinimumHeight(150)
        self.lists_preview_image.setStyleSheet("background-color: #1a1a1a; border: 1px solid #333;")
        self.lists_preview_image.setText(tr("No image"))
        preview_contents_layout.addWidget(self.lists_preview_image, 1)

        # Details section
        details_group = QGroupBox(tr("Item Details"))
        details_layout = QFormLayout(details_group)
        details_layout.setContentsMargins(5, 10, 5, 5)

        self.lists_detail_library = QLabel()
        self.lists_detail_library.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        details_layout.addRow(tr("Library:"), self.lists_detail_library)

        self.lists_detail_shelfmark = QLabel()
        self.lists_detail_shelfmark.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        details_layout.addRow(tr("Shelfmark:"), self.lists_detail_shelfmark)

        self.lists_detail_image = QLabel()
        self.lists_detail_image.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        details_layout.addRow(tr("Image:"), self.lists_detail_image)

        self.lists_detail_title = QLabel()
        self.lists_detail_title.setWordWrap(True)
        self.lists_detail_title.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        details_layout.addRow(tr("Title:"), self.lists_detail_title)

        self.lists_detail_lists = QLabel()
        details_layout.addRow(tr("Lists:"), self.lists_detail_lists)

        self.lists_detail_sys_id = QLabel()
        self.lists_detail_sys_id.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        details_layout.addRow(tr("System ID:"), self.lists_detail_sys_id)

        # Tags with add button
        tags_widget = QWidget()
        self.lists_detail_tags_container = QHBoxLayout(tags_widget)
        self.lists_detail_tags_container.setContentsMargins(0, 0, 0, 0)
        self.lists_detail_tags_container.setSpacing(3)
        details_layout.addRow(tr("Tags:"), tags_widget)

        # Editable note
        self.lists_detail_note = QLineEdit()
        self.lists_detail_note.setPlaceholderText(tr("Add a note..."))
        details_layout.addRow(tr("Note:"), self.lists_detail_note)

        self.lists_detail_source = QLabel()
        details_layout.addRow(tr("Source:"), self.lists_detail_source)

        self.lists_detail_added = QLabel()
        details_layout.addRow(tr("Added:"), self.lists_detail_added)

        # Save button for note
        self.btn_detail_save = QPushButton(tr("Save Changes"))
        self.btn_detail_save.clicked.connect(self.lists_save_item_details)
        details_layout.addRow("", self.btn_detail_save)

        preview_contents_layout.addWidget(details_group)
        preview_layout.addWidget(self.lists_preview_contents)

        # --- Center Content Area ---
        center_widget = QWidget()
        center_layout = QVBoxLayout(center_widget)
        center_layout.setContentsMargins(0, 0, 0, 0)

        # Splitter for sidebar and main area
        content_splitter = QSplitter(Qt.Orientation.Horizontal)

        # --- Main Area ---
        main_area = QWidget()
        main_layout = QVBoxLayout(main_area)
        main_layout.setContentsMargins(0, 0, 0, 0)

        # List header bar
        list_header = QHBoxLayout()

        self.lists_current_label = QLabel()
        self.lists_current_label.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        list_header.addWidget(self.lists_current_label)

        list_header.addStretch()

        # Edit/Delete buttons for current list
        self.btn_edit_list = QPushButton("✏️")
        self.btn_edit_list.setFixedSize(28, 28)
        self.btn_edit_list.setToolTip(tr("Edit List"))
        self.btn_edit_list.clicked.connect(self.lists_edit_current_list)
        list_header.addWidget(self.btn_edit_list)

        self.btn_delete_list = QPushButton("🗑️")
        self.btn_delete_list.setFixedSize(28, 28)
        self.btn_delete_list.setToolTip(tr("Delete List"))
        self.btn_delete_list.clicked.connect(self.lists_delete_current_list)
        list_header.addWidget(self.btn_delete_list)

        main_layout.addLayout(list_header)

        # Filter and sort bar
        filter_bar = QHBoxLayout()

        self.lists_filter_input = QLineEdit()
        self.lists_filter_input.setPlaceholderText(tr("Filter items..."))
        self.lists_filter_input.textChanged.connect(self.lists_apply_filter)
        filter_bar.addWidget(self.lists_filter_input)

        filter_bar.addWidget(QLabel(tr("Sort by:")))
        self.lists_sort_combo = QComboBox()
        self.lists_sort_combo.addItems([tr("Shelfmark"), tr("Title"), tr("Date Added")])
        self.lists_sort_combo.currentIndexChanged.connect(self.lists_refresh_items)
        filter_bar.addWidget(self.lists_sort_combo)

        main_layout.addLayout(filter_bar)

        # Items table
        self.lists_items_table = QTableWidget()
        self.lists_items_table.setColumnCount(7)
        self.lists_items_table.setHorizontalHeaderLabels(["", tr("Library"), tr("Shelfmark"), tr("Image"), tr("Title"), tr("Tags"), tr("Actions")])
        self.lists_items_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Fixed)
        self.lists_items_table.setColumnWidth(0, 30)  # Checkbox
        self.lists_items_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Interactive)
        self.lists_items_table.setColumnWidth(1, 100)  # Library
        self.lists_items_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Interactive)
        self.lists_items_table.setColumnWidth(2, 150)  # Shelfmark
        self.lists_items_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.Interactive)
        self.lists_items_table.setColumnWidth(3, 90)  # Image
        self.lists_items_table.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeMode.Stretch)  # Title
        self.lists_items_table.horizontalHeader().setSectionResizeMode(5, QHeaderView.ResizeMode.Interactive)
        self.lists_items_table.setColumnWidth(5, 120)  # Tags
        self.lists_items_table.horizontalHeader().setSectionResizeMode(6, QHeaderView.ResizeMode.Fixed)
        self.lists_items_table.setColumnWidth(6, 120)  # Actions
        self.lists_items_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.lists_items_table.itemClicked.connect(self.lists_on_item_clicked)
        self.lists_items_table.itemChanged.connect(self.lists_on_item_checkbox_changed)
        self.lists_items_table.setMouseTracking(True)
        self.lists_items_table.cellEntered.connect(self.on_lists_table_cell_entered)
        self.lists_items_table.installEventFilter(self)
        main_layout.addWidget(self.lists_items_table, 1)

        # Single action bar
        action_bar = QHBoxLayout()

        self.chk_lists_select_all = QCheckBox(tr("Select All"))
        self.chk_lists_select_all.setTristate(True)
        self.chk_lists_select_all.toggled.connect(self.lists_on_select_all_toggled)
        action_bar.addWidget(self.chk_lists_select_all)

        self.lists_selection_label = QLabel("")
        action_bar.addWidget(self.lists_selection_label)

        btn_move = QPushButton(tr("Move to List..."))
        btn_move.clicked.connect(self.lists_move_selected_items)
        action_bar.addWidget(btn_move)

        btn_add_tag = QPushButton(tr("Add Tag..."))
        btn_add_tag.clicked.connect(self.lists_add_tag_to_selected)
        action_bar.addWidget(btn_add_tag)

        btn_remove_selected = QPushButton(tr("Remove Selected"))
        btn_remove_selected.clicked.connect(self.lists_remove_selected_items)
        action_bar.addWidget(btn_remove_selected)

        action_bar.addStretch()

        btn_export = QPushButton(tr("Export List..."))
        btn_export.clicked.connect(self.lists_export_current_list)
        action_bar.addWidget(btn_export)

        btn_import = QPushButton(tr("Import List..."))
        btn_import.clicked.connect(self.lists_import_list)
        action_bar.addWidget(btn_import)

        main_layout.addLayout(action_bar)

        # --- Sidebar: Lists ---
        sidebar = QWidget()
        sidebar.setMinimumWidth(120)
        sidebar_layout = QVBoxLayout(sidebar)
        sidebar_layout.setContentsMargins(0, 0, 0, 0)

        # Lists header
        lists_header = QHBoxLayout()
        lists_header.setContentsMargins(5, 5, 5, 5)
        lists_header.addWidget(QLabel(f"<b>{tr('Personal Lists')}</b>"))
        lists_header.addStretch()

        # New list button
        btn_new_list = QPushButton("+")
        btn_new_list.setFixedSize(24, 24)
        btn_new_list.setToolTip(tr("New List..."))
        btn_new_list.clicked.connect(self.lists_create_new_list)
        lists_header.addWidget(btn_new_list)

        sidebar_layout.addLayout(lists_header)

        # Lists tree
        self.lists_tree = ListsTreeWidget(self)
        self.lists_tree.setHeaderHidden(True)
        self.lists_tree.setIndentation(10)
        self.lists_tree.setRootIsDecorated(True)
        self.lists_tree.itemClicked.connect(self.lists_on_list_selected)
        self.lists_tree.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.lists_tree.customContextMenuRequested.connect(self.lists_show_list_context_menu)
        sidebar_layout.addWidget(self.lists_tree, 1)

        # Sidebar action buttons
        sidebar_actions = QVBoxLayout()
        sidebar_actions.setSpacing(2)
        sidebar_actions.setContentsMargins(5, 0, 5, 5)

        btn_add_list = QPushButton(tr("Add list..."))
        btn_add_list.clicked.connect(self.lists_create_new_list)
        sidebar_actions.addWidget(btn_add_list)

        btn_add_project = QPushButton(tr("Add project..."))
        btn_add_project.clicked.connect(self.lists_create_new_project)
        sidebar_actions.addWidget(btn_add_project)

        btn_duplicate = QPushButton(tr("Duplicate List"))
        btn_duplicate.clicked.connect(self.lists_duplicate_selected_list)
        sidebar_actions.addWidget(btn_duplicate)

        btn_merge = QPushButton(tr("Merge Lists"))
        btn_merge.clicked.connect(self.lists_merge_lists)
        sidebar_actions.addWidget(btn_merge)

        btn_cleanup = QPushButton(tr("Fix Duplicates"))
        btn_cleanup.setToolTip(tr("Merge duplicate lists created by sync issues"))
        btn_cleanup.clicked.connect(self.lists_cleanup_duplicates)
        sidebar_actions.addWidget(btn_cleanup)

        btn_trash = QPushButton(tr("Trash"))
        btn_trash.setToolTip(tr("View and restore deleted lists"))
        btn_trash.clicked.connect(self.lists_show_trash)
        sidebar_actions.addWidget(btn_trash)

        sidebar_layout.addLayout(sidebar_actions)

        is_rtl = self.layoutDirection() == Qt.LayoutDirection.RightToLeft

        content_splitter.addWidget(sidebar)     # איבר 0: רשימות
        content_splitter.addWidget(main_area)   # איבר 1: טבלה ראשית
        
        content_splitter.setStretchFactor(0, 0)
        content_splitter.setStretchFactor(1, 1)
        content_splitter.setSizes([250, 800])

        center_layout.addWidget(content_splitter)

        main_splitter.addWidget(center_widget)            # המרכז
        main_splitter.addWidget(self.lists_preview_panel) # תצוגה מקדימה

        main_splitter.setStretchFactor(0, 1)
        main_splitter.setStretchFactor(1, 0)
        main_splitter.setSizes([800, 300])

        
        self.lists_main_splitter = main_splitter
        self.lists_preview_index = 1
        self.lists_preview_last_sizes = None  

        layout.addWidget(main_splitter)

        self.lists_current_list_id = 'default'
        self.lists_current_item_id = None
        self.lists_preview_visible = True
        
        self.lists_set_preview_visible(True, auto=True)

        return panel

    def lists_toggle_preview(self):
        """Toggle the preview panel visibility."""
        self.lists_set_preview_visible(not self.lists_preview_visible, auto=False)

    def _normalize_fl_id(self, fl_id):
        digits = re.sub(r"\D", "", str(fl_id or ""))
        return digits or None

    def _format_image_display(self, img):
        return str(img) if img not in (None, "") else ""

    def _get_list_display_name(self, lst):
        if CURRENT_LANG == 'en' and lst.get('name_en'):
            return lst.get('name_en')
        name = lst.get('name', lst.get('name_en', 'List'))
        if lst.get('is_system') or lst.get('is_default'):
            return tr(name)
        return name

    def _get_list_display_color(self, lst, projects=None):
        if lst.get('is_system') and lst.get('id') == 'recent':
            return lst.get('color', '#9E9E9E')
        project_id = lst.get('project_id')
        if project_id and projects and project_id in projects:
            return projects[project_id].get('color', '#FFD700')
        if lst.get('is_default'):
            return lst.get('color', '#FFD700')
        return self.lists_mgr.data.get('lists', {}).get('default', {}).get('color', '#FFD700')

    def lists_set_preview_visible(self, visible, auto=False):
        """Show/hide preview panel with a slim collapsed bar."""
        if self.lists_preview_visible == visible and not auto:
            return
            
        self.lists_preview_visible = visible
        
        self.lists_preview_contents.setVisible(visible)
        self.lists_preview_title.setVisible(visible)
        self.btn_toggle_preview.setText(tr("▶") if visible else tr("◀"))

        if not self.lists_main_splitter:
            return

        collapsed_width = 32 

        if visible:
            self.lists_preview_panel.setMaximumWidth(16777215) # QWIDGETSIZE_MAX
            
            self.lists_preview_panel.setMinimumWidth(250)
            
            self.lists_preview_panel.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Expanding)
            
            if self.lists_preview_last_sizes:
                self.lists_main_splitter.setSizes(self.lists_preview_last_sizes)
            
            self.lists_main_splitter.setCollapsible(self.lists_preview_index, False)

        else:
            
            self.lists_preview_last_sizes = self.lists_main_splitter.sizes()
            
            self.lists_preview_panel.setMinimumWidth(collapsed_width)
            self.lists_preview_panel.setMaximumWidth(collapsed_width)
            
            self.lists_preview_panel.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Expanding)
            
            sizes = self.lists_main_splitter.sizes()
            total = sum(sizes)
            
            new_sizes = [0] * len(sizes)
            for i in range(len(sizes)):
                if i == self.lists_preview_index:
                    new_sizes[i] = collapsed_width
                else:
                    new_sizes[i] = total - collapsed_width
            
            self.lists_main_splitter.setSizes(new_sizes)

    def lists_refresh_all(self):
        """Refresh the lists sidebar and current items view."""
        self.lists_refresh_sidebar()
        self.lists_refresh_items()
        self._update_search_action_stars()
        self._update_browse_add_to_list_button()

    _auto_sync_pending = False
    _auto_sync_last = 0

    def _lists_auto_sync(self):
        """Auto-sync to cloud after local changes (if logged in).

        Features:
        - Runs in background thread (won't freeze UI)
        - Quick network check before syncing
        - Debounced (max once per 2 seconds)
        - 30-second timeout
        """
        if not self.lists_mgr:
            logger.debug("Auto-sync: no lists_mgr")
            return
        if not hasattr(self.lists_mgr, 'is_sync_available') or not self.lists_mgr.is_sync_available():
            logger.debug("Auto-sync: sync not available")
            return

        # Debounce: skip if synced recently
        import time
        now = time.time()
        if now - self.__class__._auto_sync_last < 2:
            return

        # Skip if sync already pending
        if self.__class__._auto_sync_pending:
            return

        self.__class__._auto_sync_pending = True
        self.__class__._auto_sync_last = now

        try:
            import threading

            def sync_task():
                try:
                    # Quick network check (try to resolve Supabase host)
                    import socket
                    socket.setdefaulttimeout(5)
                    try:
                        socket.gethostbyname('ylcpglwxompwjcufdemz.supabase.co')
                    except socket.gaierror:
                        logger.debug("Auto-sync skipped: no network")
                        return
                    finally:
                        socket.setdefaulttimeout(None)

                    logger.debug("Auto-sync starting...")
                    # Run sync with timeout
                    import concurrent.futures
                    import time as _time
                    start = _time.time()
                    with concurrent.futures.ThreadPoolExecutor() as executor:
                        future = executor.submit(self.lists_mgr.sync_to_cloud)
                        try:
                            result = future.result(timeout=30)  # 30 second timeout
                            logger.debug(f"Auto-sync completed in {_time.time()-start:.1f}s: {result}")
                        except concurrent.futures.TimeoutError:
                            logger.debug("Auto-sync timed out after 30s")
                except Exception as e:
                    logger.debug(f"Auto-sync failed: {e}")
                finally:
                    self.__class__._auto_sync_pending = False

            threading.Thread(target=sync_task, daemon=True).start()
        except Exception as e:
            self.__class__._auto_sync_pending = False
            logger.debug(f"Auto-sync error: {e}")

    def lists_refresh_sidebar(self):
        """Refresh the lists tree in the sidebar."""
        self.lists_tree.clear()

        if not self.lists_mgr:
            return

        lists = self.lists_mgr.get_all_lists(include_recent=True)
        projects = {proj['id']: proj for proj in self.lists_mgr.get_projects()}
        project_items = {}
        list_items = {}

        for lst in lists:
            project_id = lst.get('project_id')
            parent = None
            if project_id and project_id in projects:
                parent = project_items.get(project_id)
                if not parent:
                    parent = QTreeWidgetItem()
                    parent.setText(0, f"📁 {projects[project_id]['name']}")
                    parent.setData(0, Qt.ItemDataRole.UserRole, None)
                    parent.setData(0, Qt.ItemDataRole.UserRole + 1, project_id)
                    parent.setForeground(0, QColor(projects[project_id].get('color', '#FFD700')))
                    parent.setFlags(
                        Qt.ItemFlag.ItemIsEnabled
                        | Qt.ItemFlag.ItemIsSelectable
                        | Qt.ItemFlag.ItemIsDropEnabled
                        | Qt.ItemFlag.ItemIsDragEnabled
                    )
                    self.lists_tree.addTopLevelItem(parent)
                    project_items[project_id] = parent

            item = QTreeWidgetItem()

            # Create colored dot
            color = self._get_list_display_color(lst, projects)
            name = self._get_list_display_name(lst)
            count = lst.get('count', 0)

            if lst.get('is_system') and lst['id'] == 'recent':
                display_text = f"🕐 {name} ({count})"
            else:
                display_text = f"● {name} ({count})"
                item.setForeground(0, QColor(color))

            item.setText(0, display_text)
            item.setData(0, Qt.ItemDataRole.UserRole, lst['id'])
            item.setFlags(
                Qt.ItemFlag.ItemIsEnabled
                | Qt.ItemFlag.ItemIsSelectable
                | Qt.ItemFlag.ItemIsDragEnabled
            )

            # Bold for default list
            if lst.get('is_default'):
                font = item.font(0)
                font.setBold(True)
                item.setFont(0, font)

            if parent:
                parent.addChild(item)
            else:
                self.lists_tree.addTopLevelItem(item)
            list_items[lst['id']] = item

        for project in self.lists_mgr.get_projects():
            if project['id'] not in project_items:
                parent = QTreeWidgetItem()
                parent.setText(0, f"📁 {project['name']}")
                parent.setData(0, Qt.ItemDataRole.UserRole, None)
                parent.setData(0, Qt.ItemDataRole.UserRole + 1, project['id'])
                parent.setForeground(0, QColor(project.get('color', '#FFD700')))
                parent.setFlags(
                    Qt.ItemFlag.ItemIsEnabled
                    | Qt.ItemFlag.ItemIsSelectable
                    | Qt.ItemFlag.ItemIsDropEnabled
                    | Qt.ItemFlag.ItemIsDragEnabled
                )
                self.lists_tree.addTopLevelItem(parent)
                project_items[project['id']] = parent

        self.lists_tree.expandAll()

        # Select current list
        current_item = list_items.get(self.lists_current_list_id)
        if current_item:
            self.lists_tree.setCurrentItem(current_item)

    def lists_handle_tree_reorder(self):
        """Apply drag-and-drop changes to list/project order and assignment."""
        if not self.lists_mgr:
            return

        list_project_map = {}
        list_order = []
        project_order = []

        for i in range(self.lists_tree.topLevelItemCount()):
            top_item = self.lists_tree.topLevelItem(i)
            list_id = top_item.data(0, Qt.ItemDataRole.UserRole)
            project_id = top_item.data(0, Qt.ItemDataRole.UserRole + 1)

            if project_id and not list_id:
                project_order.append(project_id)
                for j in range(top_item.childCount()):
                    child = top_item.child(j)
                    child_list_id = child.data(0, Qt.ItemDataRole.UserRole)
                    if child_list_id:
                        list_project_map[child_list_id] = project_id
                        list_order.append(child_list_id)
            elif list_id:
                list_project_map[list_id] = None
                list_order.append(list_id)

        for list_id, list_data in self.lists_mgr.data.get('lists', {}).items():
            if list_id not in list_order:
                list_order.append(list_id)
                list_project_map.setdefault(list_id, list_data.get('project_id'))

        for project_id in self.lists_mgr.data.get('projects', {}):
            if project_id not in project_order:
                project_order.append(project_id)

        self.lists_mgr.apply_list_layout(list_project_map, list_order, project_order)
        self.lists_refresh_sidebar()

    def lists_refresh_items(self):
        """Refresh the items table for the current list."""
        self.lists_items_table.setRowCount(0)

        if not self.lists_mgr:
            return

        # Get current list info
        lists = self.lists_mgr.get_all_lists()
        projects = {proj['id']: proj for proj in self.lists_mgr.get_projects()}
        current_list = None
        for lst in lists:
            if lst['id'] == self.lists_current_list_id:
                current_list = lst
                break

        if current_list:
            color = self._get_list_display_color(current_list, projects)
            name = self._get_list_display_name(current_list)
            self.lists_current_label.setText(f"<span style='color:{color}'>●</span> {name}")

            # Show/hide edit/delete buttons for system lists
            is_system = current_list.get('is_system', False)
            is_default = current_list.get('is_default', False)
            self.btn_edit_list.setEnabled(not is_system)
            self.btn_delete_list.setEnabled(not is_system and not is_default)

        # Get sort order
        sort_map = {0: 'shelfmark', 1: 'title', 2: 'added'}
        sort_by = sort_map.get(self.lists_sort_combo.currentIndex(), 'shelfmark')

        items = self.lists_mgr.get_items_sorted(self.lists_current_list_id, sort_by=sort_by)

        # Apply filter if any
        filter_text = self.lists_filter_input.text().strip().lower()

        self.lists_items_table.blockSignals(True)
        visible_item_ids = set()

        for item in items:
            shelfmark = item.get('shelfmark', 'Unknown')
            title = item.get('title', '')
            tags = item.get('tags', [])
            item_id = item.get('item_id') or item.get('sys_id')
            img = item.get('img')
            sys_id = item.get('sys_id')
            is_unidentified = item.get('shelfmark_override') is not None

            # Apply filter
            if filter_text:
                searchable = f"{shelfmark} {title} {' '.join(tags)} {img or ''}".lower()
                if filter_text not in searchable:
                    continue

            row = self.lists_items_table.rowCount()
            self.lists_items_table.insertRow(row)
            if item_id:
                visible_item_ids.add(item_id)

            # Checkbox
            chk_item = QTableWidgetItem()
            chk_item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
            chk_item.setCheckState(Qt.CheckState.Unchecked)
            chk_item.setData(Qt.ItemDataRole.UserRole, item_id)
            self.lists_items_table.setItem(row, 0, chk_item)

            # Library
            library_code = self.meta_mgr.get_library_for_id(sys_id) if sys_id else ''
            library_item = QTableWidgetItem(library_code)
            if library_code:
                library_item.setToolTip(get_library_display(library_code, short=False))
            self.lists_items_table.setItem(row, 1, library_item)

            # Shelfmark
            shelf_item = QTableWidgetItem(shelfmark)
            if is_unidentified:
                shelf_item.setForeground(QColor("#e74c3c"))
                shelf_item.setToolTip(tr("(unidentified)"))
            self.lists_items_table.setItem(row, 2, shelf_item)

            # Image
            self.lists_items_table.setItem(row, 3, QTableWidgetItem(str(img or "")))

            # Title
            self.lists_items_table.setItem(row, 4, QTableWidgetItem(title))

            # Tags
            self.lists_items_table.setItem(row, 5, QTableWidgetItem(", ".join(tags)))

            # Action buttons
            actions_widget = ActionsHoverWidget()

            btn_view = self._create_action_button("👁️", tr("Quick View"), lambda _, iid=item_id: self.lists_quick_view_by_id(iid), parent=self.lists_items_table)
            actions_widget.add_btn(btn_view)

            btn_browse = self._create_action_button("📖", tr("Browse"), lambda _, iid=item_id: self.lists_browse_by_id(iid), parent=self.lists_items_table)
            actions_widget.add_btn(btn_browse)

            btn_copy = self._create_action_button("📋", tr("Copy Info"), lambda _, iid=item_id: self.lists_copy_info_by_id(iid), parent=self.lists_items_table)
            actions_widget.add_btn(btn_copy)

            btn_remove = self._create_action_button("🗑️", tr("Remove from List"), lambda _, iid=item_id: self.lists_remove_item_by_id(iid), parent=self.lists_items_table)
            actions_widget.add_btn(btn_remove)

            self.lists_items_table.setCellWidget(row, 6, actions_widget)

        self.lists_items_table.blockSignals(False)
        self.lists_update_selection_label()
        self.chk_lists_select_all.setEnabled(self.lists_items_table.rowCount() > 0)
        self.lists_sync_select_all_checkbox()
        if not self.lists_current_item_id or self.lists_current_item_id not in visible_item_ids:
            self.lists_current_item_id = None
            self.lists_clear_details()

    def lists_on_list_selected(self, item, column):
        """Handle list selection in the sidebar."""
        list_id = item.data(0, Qt.ItemDataRole.UserRole)
        if not list_id:
            return
        self.lists_current_list_id = list_id
        self.lists_current_item_id = None
        self.lists_refresh_items()
        self.lists_clear_details()

    def lists_on_item_clicked(self, item):
        """Handle item click in the table."""
        if item.column() == 0:  # Checkbox column
            return

        row = item.row()
        chk_item = self.lists_items_table.item(row, 0)
        if chk_item:
            item_id = chk_item.data(Qt.ItemDataRole.UserRole)
            if item_id:
                self.lists_current_item_id = item_id
                self.lists_show_item_details(item_id)

    def lists_on_item_checkbox_changed(self, item):
        """Handle checkbox state change."""
        if item.column() != 0:
            return
        self.lists_update_selection_label()

    def lists_update_selection_label(self):
        """Update the selection count label."""
        count = 0
        for row in range(self.lists_items_table.rowCount()):
            chk = self.lists_items_table.item(row, 0)
            if chk and chk.checkState() == Qt.CheckState.Checked:
                count += 1

        if count > 0:
            self.lists_selection_label.setText(tr("{} selected").format(count))
        else:
            self.lists_selection_label.setText("")
        self.lists_sync_select_all_checkbox()

    def lists_on_select_all_toggled(self, checked):
        """Toggle all checkboxes in the list table."""
        if not hasattr(self, "_lists_select_all_guard"):
            self._lists_select_all_guard = False
        if self._lists_select_all_guard:
            return
        self.lists_items_table.blockSignals(True)
        for row in range(self.lists_items_table.rowCount()):
            chk = self.lists_items_table.item(row, 0)
            if chk:
                chk.setCheckState(Qt.CheckState.Checked if checked else Qt.CheckState.Unchecked)
        self.lists_items_table.blockSignals(False)
        self.lists_update_selection_label()

    def lists_sync_select_all_checkbox(self):
        """Sync 'Select All' checkbox state with row selections."""
        if not hasattr(self, "_lists_select_all_guard"):
            self._lists_select_all_guard = False
        total = self.lists_items_table.rowCount()
        if total == 0:
            self._lists_select_all_guard = True
            self.chk_lists_select_all.setCheckState(Qt.CheckState.Unchecked)
            self._lists_select_all_guard = False
            return
        checked_count = 0
        for row in range(total):
            chk = self.lists_items_table.item(row, 0)
            if chk and chk.checkState() == Qt.CheckState.Checked:
                checked_count += 1
        self._lists_select_all_guard = True
        if checked_count == 0:
            self.chk_lists_select_all.setCheckState(Qt.CheckState.Unchecked)
        elif checked_count == total:
            self.chk_lists_select_all.setCheckState(Qt.CheckState.Checked)
        else:
            self.chk_lists_select_all.setCheckState(Qt.CheckState.PartiallyChecked)
        self._lists_select_all_guard = False

    def lists_get_selected_item_ids(self):
        """Get list of selected item ids."""
        selected = []
        for row in range(self.lists_items_table.rowCount()):
            chk = self.lists_items_table.item(row, 0)
            if chk and chk.checkState() == Qt.CheckState.Checked:
                item_id = chk.data(Qt.ItemDataRole.UserRole)
                if item_id:
                    selected.append(item_id)
        return selected

    def lists_show_item_details(self, item_id):
        """Show details for a specific item."""
        if not self.lists_mgr:
            return

        item = self.lists_mgr.get_item(item_id)
        if not item:
            return
        sys_id = item.get('sys_id')
        img = item.get('img')

        # Get metadata
        shelfmark = 'Unknown'
        title = ''
        if item.get('shelfmark_override'):
            shelfmark = item['shelfmark_override']
        elif self.meta_mgr:
            shelfmark, title = self.meta_mgr.get_meta_for_id(sys_id)

        self.lists_detail_shelfmark.setText(shelfmark)
        # Library
        library_code = self.meta_mgr.get_library_for_id(sys_id) if sys_id and self.meta_mgr else ''
        library_name = get_library_display(library_code, short=False) if library_code else '-'
        self.lists_detail_library.setText(library_name)
        self.lists_detail_image.setText(str(img) if img not in (None, "") else "-")
        self.lists_detail_title.setText(title)
        self.lists_detail_sys_id.setText(sys_id or '')

        # Lists
        list_names = []
        for list_id in item.get('lists', []):
            lst_info = self.lists_mgr.data['lists'].get(list_id)
            if lst_info:
                list_names.append(self._get_list_display_name(lst_info))
        self.lists_detail_lists.setText(", ".join(list_names) if list_names else "-")

        # Tags
        # Clear previous tags
        while self.lists_detail_tags_container.count():
            child = self.lists_detail_tags_container.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        for tag in item.get('tags', []):
            tag_label = QLabel(f"[{tag}]")
            tag_label.setStyleSheet("color: #3498db; font-weight: bold;")
            self.lists_detail_tags_container.addWidget(tag_label)

        # Add tag button
        btn_add_tag = QPushButton("+")
        btn_add_tag.setFixedSize(20, 20)
        btn_add_tag.clicked.connect(lambda: self.lists_add_tag_to_item(item_id))
        self.lists_detail_tags_container.addWidget(btn_add_tag)
        self.lists_detail_tags_container.addStretch()

        # Note
        self.lists_detail_note.setText(item.get('note', ''))

        # Source
        self.lists_detail_source.setText(item.get('source', '-'))

        # Added date
        added_ts = item.get('added', 0)
        if added_ts:
            from datetime import datetime
            added_str = datetime.fromtimestamp(added_ts).strftime('%d.%m.%Y %H:%M')
            self.lists_detail_added.setText(added_str)
        else:
            self.lists_detail_added.setText('-')

        # Load text preview and image
        if sys_id:
            self._lists_load_preview(sys_id, img=img)
        if not self.lists_preview_visible:
            self.lists_set_preview_visible(True, auto=True)

    def _lists_load_preview(self, sys_id, img=None):
        """Load text and image preview for an item."""
        # Clear previous preview
        self.lists_preview_text.clear()
        self.lists_preview_image.clear()
        self.lists_preview_image.setText(tr("No image"))

        if not self.searcher:
            return

        # Get page data
        p_num = 1
        if img not in (None, ""):
            try:
                p_num = int(img)
            except ValueError:
                p_num = 1
        page_data = self.searcher.get_browse_page(sys_id, p_num=p_num)
        if not page_data:
            self.lists_preview_text.setPlainText(tr("Could not load text"))
            return

        # Load text
        text = page_data.get('text', '')
        if text:
            # Truncate for preview (first ~500 chars)
            preview_text = text[:1000] + ('...' if len(text) > 1000 else '')
            self.lists_preview_text.setPlainText(preview_text)
        else:
            self.lists_preview_text.setPlainText(tr("No text available"))

        # Load image (async would be better, but for now sync)
        self._lists_load_preview_image(sys_id, page_data)

    def _lists_load_preview_image(self, sys_id, page_data):
        """Load image for preview panel."""
        if not self.meta_mgr:
            return

        try:
            thumb_url = None
            meta = self.meta_mgr.nli_cache.get(sys_id)
            if meta:
                thumb_url = meta.get('thumb_url')

            if not thumb_url:
                thumb_url = self.meta_mgr.get_thumbnail(sys_id)

            if thumb_url:
                self._lists_start_preview_download(thumb_url)
            else:
                self.lists_preview_image.setText(tr("No image"))
        except Exception as e:
            self.lists_preview_image.setText(tr("No image"))

    def _lists_start_preview_download(self, thumb_url):
        """Download and display preview image for lists panel."""
        self._lists_cancel_preview_image_thread()
        self.lists_preview_thumb_url = thumb_url
        if not thumb_url:
            self.lists_preview_image.setText(tr("No image"))
            return

        self.lists_preview_image.setText(tr("Loading..."))
        self.lists_preview_image.setPixmap(QPixmap())

        self.lists_preview_img_thread = ImageLoaderThread(thumb_url)
        self.lists_preview_img_thread.image_loaded.connect(
            lambda image, url=thumb_url: self._lists_on_preview_image_loaded(image, url)
        )
        self.lists_preview_img_thread.load_failed.connect(
            lambda url=thumb_url: self._lists_on_preview_image_failed(url)
        )
        self.lists_preview_img_thread.start()

    def _lists_on_preview_image_loaded(self, image, thumb_url):
        """Handle preview image loaded for lists panel."""
        if thumb_url != getattr(self, 'lists_preview_thumb_url', None):
            return
        pix = QPixmap.fromImage(image)
        scaled = pix.scaled(
            self.lists_preview_image.size(),
            Qt.AspectRatioMode.KeepAspectRatio,
            Qt.TransformationMode.SmoothTransformation
        )
        self.lists_preview_image.setPixmap(scaled)
        self.lists_preview_image.setText("")

    def _lists_on_preview_image_failed(self, thumb_url):
        """Handle preview image load failure for lists panel."""
        if thumb_url != getattr(self, 'lists_preview_thumb_url', None):
            return
        self.lists_preview_image.setPixmap(QPixmap())
        self.lists_preview_image.setText(tr("No image"))

    def _lists_cancel_preview_image_thread(self):
        preview_thread = getattr(self, 'lists_preview_img_thread', None)
        if preview_thread and preview_thread.isRunning():
            preview_thread.cancel()
            preview_thread.wait(500)

    def lists_clear_details(self):
        """Clear the details panel and preview."""
        self.lists_detail_shelfmark.setText('')
        self.lists_detail_library.setText('')
        self.lists_detail_title.setText('')
        self.lists_detail_lists.setText('')
        self.lists_detail_sys_id.setText('')
        self.lists_detail_image.setText('')
        self.lists_detail_note.setText('')
        self.lists_detail_source.setText('')
        self.lists_detail_added.setText('')

        # Clear tags
        while self.lists_detail_tags_container.count():
            child = self.lists_detail_tags_container.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # Clear preview
        self.lists_preview_text.clear()
        self.lists_preview_image.clear()
        self.lists_preview_image.setText(tr("No image"))
        self._lists_cancel_preview_image_thread()
        if self.lists_preview_visible:
            self.lists_set_preview_visible(False, auto=True)

    def lists_save_item_details(self):
        """Save changes to the current item."""
        if not self.lists_mgr or not self.lists_current_item_id:
            return

        note = self.lists_detail_note.text()
        self.lists_mgr.update_item(self.lists_current_item_id, note=note)

    def lists_create_new_list(self):
        """Create a new list."""
        name, ok = QInputDialog.getText(self, tr("Create New List"), tr("List Name:"))
        if ok and name.strip():
            if self.lists_mgr:
                self.lists_mgr.create_list(name.strip())
                self.lists_refresh_sidebar()
                self._lists_auto_sync()

    def lists_create_new_project(self):
        """Create a new project."""
        name, ok = QInputDialog.getText(self, tr("Create New Project"), tr("Project Name:"))
        if ok and name.strip():
            if self.lists_mgr:
                self.lists_mgr.create_project(name.strip())
                self.lists_refresh_sidebar()
                self._lists_auto_sync()

    def lists_edit_current_list(self):
        """Edit the current list name/color."""
        if not self.lists_mgr or self.lists_current_list_id in ['default', 'recent']:
            return

        lst = self.lists_mgr.data['lists'].get(self.lists_current_list_id)
        if not lst or lst.get('is_system'):
            return

        name, ok = QInputDialog.getText(self, tr("Rename List"), tr("List Name:"), text=lst.get('name', ''))
        if ok and name.strip():
            self.lists_mgr.update_list(self.lists_current_list_id, name=name.strip())
            self.lists_refresh_all()
            self._lists_auto_sync()

    def lists_delete_current_list(self):
        """Delete the current list."""
        if not self.lists_mgr or self.lists_current_list_id in ['default', 'recent']:
            return

        lst = self.lists_mgr.data['lists'].get(self.lists_current_list_id)
        if not lst:
            return

        name = lst.get('name', 'List')
        reply = QMessageBox.question(
            self, tr("Delete List?"),
            tr("Move '{}' to trash?\nYou can restore it later from the Trash.").format(name),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            self.lists_mgr.delete_list(self.lists_current_list_id)
            self.lists_current_list_id = 'default'
            self.lists_refresh_all()
            self._lists_auto_sync()

    def lists_duplicate_selected_list(self):
        """Duplicate the current list."""
        if not self.lists_mgr or self.lists_current_list_id == 'recent':
            return

        self.lists_mgr.duplicate_list(self.lists_current_list_id)
        self.lists_refresh_sidebar()
        self._lists_auto_sync()

    def lists_merge_lists(self):
        """Show dialog to merge lists."""
        if not self.lists_mgr:
            return

        # Get non-system lists
        lists = [l for l in self.lists_mgr.get_all_lists(include_recent=False)
                 if not l.get('is_system')]

        if len(lists) < 2:
            QMessageBox.information(self, tr("Merge Lists"), tr("Need at least two lists to merge."))
            return

        # Simple dialog - select target
        items = [self._get_list_display_name(l) for l in lists]
        target_name, ok = QInputDialog.getItem(
            self, tr("Merge Lists"),
            tr("Merge '{}' into:").format(
                self.lists_mgr.data['lists'].get(self.lists_current_list_id, {}).get('name', '')
            ),
            items, 0, False
        )

        if ok and target_name:
            target_list = None
            for l in lists:
                if self._get_list_display_name(l) == target_name:
                    target_list = l
                    break

            if target_list and target_list['id'] != self.lists_current_list_id:
                self.lists_mgr.merge_lists(self.lists_current_list_id, target_list['id'])
                self.lists_current_list_id = target_list['id']
                self.lists_refresh_all()

    def lists_cleanup_duplicates(self):
        """Clean up duplicate lists created by sync bugs."""
        if not self.lists_mgr:
            return

        duplicate_groups = self.lists_mgr.find_duplicate_lists()

        if not duplicate_groups:
            QMessageBox.information(self, tr("Fix Duplicates"), tr("No duplicate lists found."))
            return

        auto_groups = [g for g in duplicate_groups if not g['has_conflict']]
        conflict_groups = [g for g in duplicate_groups if g['has_conflict']]

        total_merged = 0
        total_deleted = 0
        details = []

        for group in auto_groups:
            result = self.lists_mgr.auto_merge_duplicate_group(group)
            total_merged += result['merged_items']
            total_deleted += result['deleted_count']
            if result['merged_items'] > 0:
                details.append((group['name'], result['merged_items']))

        for group in conflict_groups:
            result = self._show_duplicate_conflict_dialog(group)
            if result:
                total_merged += result['merged_items']
                total_deleted += result['deleted_count']
                if result['merged_items'] > 0:
                    details.append((group['name'], result['merged_items']))

        hierarchy_result = self.lists_mgr.restore_project_hierarchy()
        restored = hierarchy_result.get('restored_count', 0)

        if total_deleted > 0 or restored > 0:
            msg = tr("Cleanup complete:\n")
            if total_deleted > 0:
                msg += tr("- Removed {} duplicate lists\n").format(total_deleted)
            if total_merged > 0:
                msg += tr("- Merged {} items total\n").format(total_merged)
            for list_name, count in details:
                msg += f"  * '{list_name}': {count} " + tr("items merged") + "\n"
            if restored > 0:
                msg += tr("- Restored {} lists to their projects\n").format(restored)
            QMessageBox.information(self, tr("Fix Duplicates"), msg)
            self.lists_refresh_all()
        else:
            QMessageBox.information(self, tr("Fix Duplicates"), tr("No changes made."))

    def _show_duplicate_conflict_dialog(self, group):
        """Show dialog for user to resolve a duplicate list conflict."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QRadioButton, QButtonGroup, QDialogButtonBox

        dialog = QDialog(self)
        dialog.setWindowTitle(tr("Resolve Duplicate: {}").format(group['name']))
        dialog.setMinimumWidth(400)

        layout = QVBoxLayout(dialog)

        header = QLabel(tr("Found {} lists named '{}' with different projects.\nChoose which to keep:").format(
            len(group['lists']), group['name']
        ))
        header.setWordWrap(True)
        layout.addWidget(header)

        button_group = QButtonGroup(dialog)

        for i, lst in enumerate(group['lists']):
            project_name = lst['project_name'] or tr("(no project)")
            item_count = lst['item_count']
            label = tr("Under '{}' ({} items)").format(project_name, item_count)
            if lst['id'] in ['default', 'recent']:
                label += tr(" [System]")

            radio = QRadioButton(label)
            radio.setProperty('list_info', lst)
            button_group.addButton(radio, i)
            layout.addWidget(radio)
            if i == 0:
                radio.setChecked(True)

        layout.addSpacing(10)
        info = QLabel(tr("Items from other lists will be merged into the selected one."))
        info.setStyleSheet("color: gray; font-size: 11px;")
        info.setWordWrap(True)
        layout.addWidget(info)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)

        if dialog.exec() != QDialog.DialogCode.Accepted:
            return None

        selected_id = button_group.checkedId()
        if selected_id < 0:
            return None

        selected_lst = group['lists'][selected_id]
        keep_id = selected_lst['id']
        target_project_id = selected_lst['project_id']
        duplicate_ids = [l['id'] for l in group['lists'] if l['id'] != keep_id]

        return self.lists_mgr.merge_duplicate_group(keep_id, duplicate_ids, target_project_id)

    def lists_show_trash(self):
        """Show dialog with deleted lists (trash)."""
        if not self.lists_mgr:
            return

        deleted_lists = self.lists_mgr.get_deleted_lists()

        if not deleted_lists:
            QMessageBox.information(self, tr("Trash"), tr("Trash is empty."))
            return

        from PyQt6.QtWidgets import (QDialog, QVBoxLayout, QHBoxLayout, QLabel,
                                      QListWidget, QListWidgetItem, QPushButton,
                                      QDialogButtonBox)
        from datetime import datetime

        dialog = QDialog(self)
        dialog.setWindowTitle(tr("Trash"))
        dialog.setMinimumWidth(400)
        dialog.setMinimumHeight(300)

        layout = QVBoxLayout(dialog)

        header = QLabel(tr("{} deleted lists").format(len(deleted_lists)))
        layout.addWidget(header)

        list_widget = QListWidget()
        for lst in deleted_lists:
            deleted_at = lst.get('deleted_at', 0)
            if deleted_at:
                deleted_str = datetime.fromtimestamp(deleted_at).strftime('%Y-%m-%d %H:%M')
            else:
                deleted_str = tr("Unknown")

            item = QListWidgetItem(f"{lst['name']} ({lst['count']} {tr('items')}) - {tr('Deleted')}: {deleted_str}")
            item.setData(Qt.ItemDataRole.UserRole, lst['id'])
            list_widget.addItem(item)

        layout.addWidget(list_widget)

        btn_layout = QHBoxLayout()

        btn_restore = QPushButton(tr("Restore"))
        btn_restore.clicked.connect(lambda: self._trash_restore(dialog, list_widget))
        btn_layout.addWidget(btn_restore)

        btn_delete_perm = QPushButton(tr("Delete Permanently"))
        btn_delete_perm.clicked.connect(lambda: self._trash_delete_permanently(dialog, list_widget))
        btn_layout.addWidget(btn_delete_perm)

        btn_empty = QPushButton(tr("Empty Trash"))
        btn_empty.clicked.connect(lambda: self._trash_empty(dialog))
        btn_layout.addWidget(btn_empty)

        layout.addLayout(btn_layout)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)

        dialog.exec()

    def _trash_restore(self, dialog, list_widget):
        """Restore selected list from trash."""
        item = list_widget.currentItem()
        if not item:
            QMessageBox.warning(self, tr("Restore"), tr("Please select a list to restore."))
            return

        list_id = item.data(Qt.ItemDataRole.UserRole)
        if self.lists_mgr.restore_list(list_id):
            list_widget.takeItem(list_widget.row(item))
            self.lists_refresh_all()
            self._lists_auto_sync()
            if list_widget.count() == 0:
                dialog.accept()

    def _trash_delete_permanently(self, dialog, list_widget):
        """Permanently delete selected list from trash."""
        item = list_widget.currentItem()
        if not item:
            QMessageBox.warning(self, tr("Delete Permanently"), tr("Please select a list to delete."))
            return

        list_id = item.data(Qt.ItemDataRole.UserRole)
        reply = QMessageBox.question(
            self, tr("Delete Permanently"),
            tr("Are you sure you want to permanently delete this list?\nThis cannot be undone."),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            if self.lists_mgr.permanently_delete_list(list_id):
                list_widget.takeItem(list_widget.row(item))
                self._lists_auto_sync()
                if list_widget.count() == 0:
                    dialog.accept()

    def _trash_empty(self, dialog):
        """Empty all trash."""
        reply = QMessageBox.question(
            self, tr("Empty Trash"),
            tr("Are you sure you want to permanently delete all lists in trash?\nThis cannot be undone."),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            count = self.lists_mgr.empty_trash()
            QMessageBox.information(self, tr("Empty Trash"), tr("Deleted {} lists permanently.").format(count))
            self.lists_refresh_all()
            self._lists_auto_sync()
            dialog.accept()

    def lists_move_selected_items(self):
        """Move selected items to another list."""
        if not self.lists_mgr:
            return

        selected = self.lists_get_selected_item_ids()
        if not selected:
            return

        lists = [l for l in self.lists_mgr.get_all_lists(include_recent=False)]
        items = [self._get_list_display_name(l) for l in lists]

        target_name, ok = QInputDialog.getItem(
            self, tr("Move to List..."),
            tr("Select list:"),
            items, 0, False
        )

        if ok and target_name:
            target_list = None
            for l in lists:
                if self._get_list_display_name(l) == target_name:
                    target_list = l
                    break

            if target_list:
                self.lists_mgr.move_items_to_list(selected, self.lists_current_list_id, target_list['id'])
                self.lists_refresh_all()

    def lists_add_tag_to_selected(self):
        """Add a tag to selected items."""
        if not self.lists_mgr:
            return

        selected = self.lists_get_selected_item_ids()
        if not selected:
            return

        tag, ok = QInputDialog.getText(self, tr("Add Tag"), tr("Enter tag:"))
        if ok and tag.strip():
            self.lists_mgr.add_tag_to_items(selected, tag.strip())
            self.lists_refresh_items()

    def lists_add_tag_to_item(self, item_id):
        """Add a tag to a specific item."""
        if not self.lists_mgr:
            return

        tag, ok = QInputDialog.getText(self, tr("Add Tag"), tr("Enter tag:"))
        if ok and tag.strip():
            self.lists_mgr.add_tag_to_items([item_id], tag.strip())
            self.lists_show_item_details(item_id)
            self.lists_refresh_items()

    def lists_remove_selected_items(self):
        """Remove selected items from current list."""
        if not self.lists_mgr:
            return

        selected = self.lists_get_selected_item_ids()
        if not selected:
            return

        reply = QMessageBox.question(
            self, tr("Delete Items?"),
            tr("Remove {} items from this list?").format(len(selected)),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            for item_id in selected:
                self.lists_mgr.remove_item_from_list(item_id, self.lists_current_list_id)
            self.lists_refresh_all()

    def lists_remove_item_by_id(self, item_id):
        """Remove a specific item from current list."""
        if not self.lists_mgr:
            return

        self.lists_mgr.remove_item_from_list(item_id, self.lists_current_list_id)
        self.lists_refresh_all()

    def lists_quick_view_item(self):
        """Quick view the current item."""
        if self.lists_current_item_id:
            self.lists_quick_view_by_id(self.lists_current_item_id)

    def lists_quick_view_by_id(self, item_id):
        """Open quick view dialog for an item."""
        if not self.searcher or not self.meta_mgr:
            return
        item = self.lists_mgr.get_item(item_id)
        if not item:
            return
        sys_id = item.get('sys_id')
        img = item.get('img')
        if not sys_id:
            return

        # Get page data from browse index to have all required fields
        p_num = 1
        if img not in (None, ""):
            try:
                p_num = int(img)
            except ValueError:
                p_num = 1
        page_data = self.searcher.get_browse_page(sys_id, p_num=p_num)
        if not page_data:
            QMessageBox.warning(self, tr("View Error"), tr("Could not load manuscript data."))
            return

        shelfmark, title = self.meta_mgr.get_meta_for_id(sys_id)

        # Create a complete result dict for ResultDialog
        result = {
            'uid': page_data.get('uid', ''),
            'raw_header': page_data.get('full_header', ''),
            'full_header': page_data.get('full_header', ''),
            'text': page_data.get('text', ''),
            'full_text': page_data.get('text', ''),
            'display': {
                'id': sys_id,
                'shelfmark': shelfmark,
                'title': title,
                'img': '',
                'source': ''
            }
        }

        ResultDialog(self, [result], 0, self.meta_mgr, self.searcher).exec()

    def lists_browse_item(self):
        """Browse the current item in the Browse tab."""
        if self.lists_current_item_id:
            self.lists_browse_by_id(self.lists_current_item_id)

    def lists_browse_by_id(self, item_id):
        """Open an item in the Browse tab."""
        item = self.lists_mgr.get_item(item_id)
        if not item:
            return
        sys_id = item.get('sys_id')
        img = item.get('img')
        if not sys_id:
            return

        # Switch to browse tab and load the manuscript
        self.tabs.setCurrentWidget(self.browse_tab)
        self.browse_sys_input.setText(sys_id)
        self._set_last_browse_field("sys")
        self.browse_load()

    def _open_document_result_dialog(self, shelfmark=None, sys_id=None, page_num=1):
        """Open ResultDialog for a document by shelfmark or sys_id."""
        logger.debug("_open_document_result_dialog called: shelfmark=%s, sys_id=%s", shelfmark, sys_id)
        try:
            if not self.searcher or not self.meta_mgr:
                logger.debug("searcher or meta_mgr not available")
                return

            # Get sys_id from shelfmark if needed
            if not sys_id and shelfmark:
                logger.debug("Looking up sys_id from shelfmark")
                self._ensure_shelf_map()
                norm = self._normalize_shelfmark(shelfmark)
                sys_id = self._shelf_to_sys.get(norm) if norm else None
                logger.debug("Normalized: %s, sys_id: %s", norm, sys_id)

            if not sys_id:
                logger.debug("No sys_id found")
                QMessageBox.warning(self, tr("Error"), tr("Document not found"))
                return

            # Get page data
            logger.debug("Getting page data for sys_id=%s", sys_id)
            page_data = self.searcher.get_browse_page(sys_id, p_num=page_num)
            if not page_data:
                logger.debug("No page data found")
                QMessageBox.warning(self, tr("View Error"), tr("Could not load manuscript data."))
                return

            shelfmark_display, title = self.meta_mgr.get_meta_for_id(sys_id)
            logger.debug("shelfmark_display=%s, title=%s", shelfmark_display, title)

            # Create result dict for ResultDialog
            result = {
                'uid': page_data.get('uid', ''),
                'raw_header': page_data.get('full_header', ''),
                'full_header': page_data.get('full_header', ''),
                'text': page_data.get('text', ''),
                'full_text': page_data.get('text', ''),
                'display': {
                    'id': sys_id,
                    'shelfmark': shelfmark_display,
                    'title': title,
                    'img': str(page_num),
                    'source': ''
                }
            }

            logger.debug("Opening ResultDialog")
            ResultDialog(self, [result], 0, self.meta_mgr, self.searcher).exec()
            logger.debug("ResultDialog closed")
        except Exception as e:
            logger.debug("Error in _open_document_result_dialog: %s", e)
            import traceback
            traceback.print_exc()

    def _browse_document_by_shelfmark(self, shelfmark, page_num=1):
        """Browse a document by shelfmark in the Browse tab."""
        self.tabs.setCurrentWidget(self.browse_tab)
        self.browse_shelf_input.setText(shelfmark)
        self._set_last_browse_field("shelf")
        self.browse_load()

    def lists_copy_item_info(self):
        """Copy current item info to clipboard."""
        if self.lists_current_item_id:
            self.lists_copy_info_by_id(self.lists_current_item_id)

    def lists_copy_info_by_id(self, item_id):
        """Copy item info to clipboard with format options."""
        if not self.lists_mgr:
            return

        menu = QMenu(self)

        action_compact = menu.addAction(tr("Compact"))
        action_compact.triggered.connect(lambda: self._do_copy_info(item_id, 'compact'))

        action_detailed = menu.addAction(tr("Detailed"))
        action_detailed.triggered.connect(lambda: self._do_copy_info(item_id, 'detailed'))

        action_link = menu.addAction(tr("With Link"))
        action_link.triggered.connect(lambda: self._do_copy_info(item_id, 'with_link'))

        menu.exec(QCursor.pos())

    def _do_copy_info(self, item_id, format_type):
        """Actually copy the info to clipboard."""
        if not self.lists_mgr:
            return

        # Get info even if item is not in a list
        item = self.lists_mgr.get_item(item_id)

        if item:
            sys_id = item.get('sys_id')
            text = self.lists_mgr.get_item_copy_text(item_id, format_type)
        else:
            # Item not in lists, generate text from metadata
            sys_id = item_id
            shelfmark, title = self.meta_mgr.get_meta_for_id(sys_id) if self.meta_mgr else ('Unknown', '')

            if format_type == 'compact':
                text = f"{shelfmark} - {title}" if title else shelfmark
            elif format_type == 'detailed':
                lines = [f"מספר מדף: {shelfmark}"]
                if title:
                    lines.append(f"כותרת: {title}")
                lines.append(f"מספר מערכת: {sys_id}")
                text = '\n'.join(lines)
            else:  # with_link
                lines = [f"מספר מדף: {shelfmark}"]
                if title:
                    lines.append(f"כותרת: {title}")
                lines.append(f"מספר מערכת: {sys_id}")
                ktiv_url = f"https://www.nli.org.il/he/discover/manuscripts/hebrew-manuscripts/itempage?vid=KTIV&scope=KTIV&docId=PNX_MANUSCRIPTS{sys_id}"
                lines.append(f"קישור: {ktiv_url}")
                text = '\n'.join(lines)

        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        self.status_label.setText(tr("Info copied to clipboard."))

    def lists_export_current_list(self):
        """Export the current list."""
        if not self.lists_mgr or not self.lists_current_list_id:
            return

        # Export format selection menu
        menu = QMenu(self)

        action_text = menu.addAction(tr("Text (plain)"))
        action_text.triggered.connect(lambda: self._export_list_format(self.lists_current_list_id, 'text'))

        action_json = menu.addAction(tr("JSON"))
        action_json.triggered.connect(lambda: self._export_list_format(self.lists_current_list_id, 'json'))

        menu.addSeparator()

        action_excel = menu.addAction(tr("Excel (.xlsx)"))
        action_excel.triggered.connect(lambda: self._export_list_format(self.lists_current_list_id, 'excel'))

        action_word = menu.addAction(tr("Word (.docx)"))
        action_word.triggered.connect(lambda: self._export_list_format(self.lists_current_list_id, 'word'))

        menu.exec(QCursor.pos())

    def lists_import_list(self):
        """Import a list from file."""
        if not self.lists_mgr:
            return

        path, _ = QFileDialog.getOpenFileName(
            self, tr("Select List File"),
            Config.REPORTS_DIR,
            tr("JSON Files") + " (*.json)"
        )

        if not path:
            return

        try:
            with open(path, 'r', encoding='utf-8') as f:
                import_data = json.load(f)

            list_id, imported, unidentified = self.lists_mgr.import_list(import_data)

            if list_id:
                self.lists_current_list_id = list_id
                self.lists_refresh_all()
                QMessageBox.information(
                    self, tr("Import List"),
                    tr("Imported {} items ({} unidentified).").format(imported, unidentified)
                )
            else:
                QMessageBox.warning(self, tr("Import Error"), tr("Invalid file format."))

        except Exception as e:
            QMessageBox.warning(self, tr("Import Error"), str(e))

    def lists_show_list_context_menu(self, pos):
        """Show context menu for list items in sidebar."""
        item = self.lists_tree.itemAt(pos)
        if not item:
            return

        list_id = item.data(0, Qt.ItemDataRole.UserRole)
        project_id = item.data(0, Qt.ItemDataRole.UserRole + 1)
        if not self.lists_mgr:
            return

        if not list_id and project_id:
            menu = QMenu(self)
            action_rename = menu.addAction(tr("Rename Project"))
            action_delete_keep = menu.addAction(tr("Delete Project (Keep Lists)"))
            action_delete_lists = menu.addAction(tr("Delete Project and Lists"))

            action = menu.exec(self.lists_tree.mapToGlobal(pos))
            if action == action_rename:
                name, ok = QInputDialog.getText(
                    self,
                    tr("Rename Project"),
                    tr("Project Name:"),
                    text=self.lists_mgr.data.get('projects', {}).get(project_id, {}).get('name', '')
                )
                if ok and name.strip():
                    self.lists_mgr.update_project(project_id, name=name.strip())
                    self.lists_refresh_sidebar()
            elif action == action_delete_keep:
                self.lists_mgr.delete_project(project_id, delete_lists=False)
                if self.lists_current_list_id not in self.lists_mgr.data.get('lists', {}):
                    self.lists_current_list_id = 'default'
                self.lists_refresh_all()
            elif action == action_delete_lists:
                self.lists_mgr.delete_project(project_id, delete_lists=True)
                if self.lists_current_list_id not in self.lists_mgr.data.get('lists', {}):
                    self.lists_current_list_id = 'default'
                self.lists_refresh_all()
            return

        if not list_id:
            return

        lst = self.lists_mgr.data['lists'].get(list_id)
        if not lst:
            return

        menu = QMenu(self)

        action_new_category = None
        category_actions = None

        if not lst.get('is_system'):
            action_rename = menu.addAction(tr("Rename List"))
            action_rename.triggered.connect(lambda: self._rename_list(list_id))

            if not lst.get('is_default'):
                action_delete = menu.addAction(tr("Delete List"))
                action_delete.triggered.connect(lambda: self._delete_list(list_id))

            if not lst.get('is_default'):
                category_menu = menu.addMenu(tr("Add to project..."))
                category_actions = {}

                action_clear = category_menu.addAction(tr("No project"))
                category_actions[action_clear] = None

                for project in self.lists_mgr.get_projects():
                    action_project = category_menu.addAction(project['name'])
                    category_actions[action_project] = project['id']

                category_menu.addSeparator()
                action_new_category = category_menu.addAction(tr("Add new..."))

            menu.addSeparator()

        action_duplicate = menu.addAction(tr("Duplicate List"))
        action_duplicate.triggered.connect(lambda: self._duplicate_list(list_id))

        # Export submenu
        export_menu = menu.addMenu(tr("Export List"))

        action_text = export_menu.addAction(tr("Text (plain)"))
        action_text.triggered.connect(lambda: self._export_list_format(list_id, 'text'))

        action_json = export_menu.addAction(tr("JSON"))
        action_json.triggered.connect(lambda: self._export_list_format(list_id, 'json'))

        export_menu.addSeparator()

        action_excel = export_menu.addAction(tr("Excel (.xlsx)"))
        action_excel.triggered.connect(lambda: self._export_list_format(list_id, 'excel'))

        action_word = export_menu.addAction(tr("Word (.docx)"))
        action_word.triggered.connect(lambda: self._export_list_format(list_id, 'word'))

        action = menu.exec(self.lists_tree.mapToGlobal(pos))
        if lst.get('is_system') or lst.get('is_default'):
            return

        if action:
            if category_actions and action in category_actions:
                self.lists_mgr.update_list_project(list_id, category_actions[action])
                self.lists_refresh_sidebar()
                self._lists_auto_sync()
            elif action_new_category and action == action_new_category:
                name, ok = QInputDialog.getText(self, tr("Create New Project"), tr("Project Name:"))
                if ok and name.strip():
                    project_id = self.lists_mgr.create_project(name.strip())
                    self.lists_mgr.update_list_project(list_id, project_id)
                    self.lists_refresh_sidebar()
                    self._lists_auto_sync()

    def _rename_list(self, list_id):
        """Rename a specific list."""
        if not self.lists_mgr:
            return

        lst = self.lists_mgr.data['lists'].get(list_id)
        if not lst:
            return

        name, ok = QInputDialog.getText(self, tr("Rename List"), tr("List Name:"), text=lst.get('name', ''))
        if ok and name.strip():
            self.lists_mgr.update_list(list_id, name=name.strip())
            self.lists_refresh_all()
            self._lists_auto_sync()

    def _delete_list(self, list_id):
        """Delete a specific list."""
        if not self.lists_mgr:
            return

        lst = self.lists_mgr.data['lists'].get(list_id)
        if not lst:
            return

        name = lst.get('name', 'List')
        reply = QMessageBox.question(
            self, tr("Delete List?"),
            tr("Move '{}' to trash?\nYou can restore it later from the Trash.").format(name),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            self.lists_mgr.delete_list(list_id)
            if self.lists_current_list_id == list_id:
                self.lists_current_list_id = 'default'
            self.lists_refresh_all()
            self._lists_auto_sync()

    def _duplicate_list(self, list_id):
        """Duplicate a specific list."""
        if not self.lists_mgr:
            return

        self.lists_mgr.duplicate_list(list_id)
        self.lists_refresh_sidebar()
        self._lists_auto_sync()

    def _export_list(self, list_id):
        """Export a specific list (opens format menu)."""
        self._export_list_format(list_id, None)  # Will show format menu

    def _export_list_format(self, list_id, format_type):
        """Export a specific list in the given format."""
        if not self.lists_mgr:
            return

        lst = self.lists_mgr.data['lists'].get(list_id)
        if not lst:
            return

        list_name = lst.get('name', 'list')
        items = self.lists_mgr.get_items_sorted(list_id, sort_by='shelfmark')

        if not items:
            QMessageBox.information(self, tr("Export List"), tr("List is empty."))
            return

        if format_type == 'text':
            self._export_as_text(list_id, list_name, items)
        elif format_type == 'json':
            self._export_as_json(list_id, list_name, items)
        elif format_type == 'excel':
            self._export_as_excel(list_id, list_name, items)
        elif format_type == 'word':
            self._export_as_word(list_id, list_name, items)
        else:
            # Show format selection menu
            menu = QMenu(self)
            menu.addAction(tr("Text (plain)")).triggered.connect(lambda: self._export_list_format(list_id, 'text'))
            menu.addAction(tr("JSON")).triggered.connect(lambda: self._export_list_format(list_id, 'json'))
            menu.addSeparator()
            menu.addAction(tr("Excel (.xlsx)")).triggered.connect(lambda: self._export_list_format(list_id, 'excel'))
            menu.addAction(tr("Word (.docx)")).triggered.connect(lambda: self._export_list_format(list_id, 'word'))
            menu.exec(QCursor.pos())

    def _format_item_text(self, item, include_notes=True):
        """Format a single item for text export."""
        sys_id = item.get('sys_id', 'Unknown')
        shelfmark = item.get('shelfmark', '')
        title = item.get('title', '')
        source = item.get('source', '')
        notes = item.get('notes', '')
        tags = item.get('tags', [])
        img = item.get('img')

        lines = []
        if shelfmark:
            lines.append(shelfmark)
        if img not in (None, ""):
            lines.append(f"  {tr('Image')}: {self._format_image_display(img)}")
        if title:
            lines.append(f"  {title}")
        lines.append(f"  ID: {sys_id}")
        if source:
            lines.append(f"  {tr('Source')}: {source}")
        if tags:
            lines.append(f"  {tr('Tags')}: {', '.join(tags)}")
        if include_notes and notes:
            lines.append(f"  {tr('Notes')}: {notes}")

        return '\n'.join(lines)

    def _export_as_text(self, list_id, list_name, items):
        """Export list as plain text."""
        lines = [f"=== {list_name} ===", f"{tr('Total items')}: {len(items)}", ""]

        for i, item in enumerate(items, 1):
            lines.append(f"{i}. {self._format_item_text(item)}")
            lines.append("")

        text = '\n'.join(lines)

        # Ask destination
        menu = QMenu(self)
        menu.addAction(tr("Save to File")).triggered.connect(lambda: self._save_text_to_file(text, list_name))
        menu.addAction(tr("Copy to Clipboard")).triggered.connect(lambda: self._copy_to_clipboard(text))
        menu.addAction(tr("Send by Email")).triggered.connect(lambda: self._send_by_email(text, list_name))
        menu.exec(QCursor.pos())

    def _export_as_json(self, list_id, list_name, items):
        """Export list as JSON."""
        from datetime import datetime
        export_data = {
            'list_name': list_name,
            'exported': datetime.now().isoformat(),
            'items': items
        }
        json_text = json.dumps(export_data, ensure_ascii=False, indent=2)

        default_name = f"{list_name}.json"
        path, _ = QFileDialog.getSaveFileName(
            self, tr("Export List"),
            os.path.join(Config.REPORTS_DIR, default_name),
            tr("JSON Files") + " (*.json)"
        )

        if path:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(json_text)
            self.status_label.setText(tr("List exported successfully."))

    def _export_as_excel(self, list_id, list_name, items):
        """Export list as Excel file."""
        try:
            import openpyxl
        except ImportError:
            QMessageBox.warning(self, tr("Export Error"),
                                tr("Excel export requires the 'openpyxl' package. Install it with: pip install openpyxl"))
            return

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = list_name[:31]  # Excel max sheet name length

        # Headers
        headers = [tr('Shelfmark'), tr('Library'), tr('Image'), tr('Title'), 'ID', tr('Source'), tr('Tags'), tr('Notes')]
        for col, header in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=header)

        # Data - use shared_sanitize_excel to prevent formula injection and control chars
        for row, item in enumerate(items, 2):
            sys_id = item.get('sys_id', '')
            library_code = self.meta_mgr.get_library_for_id(sys_id) if sys_id else ''
            library_name = get_library_display(library_code, short=False) if library_code else ''
            ws.cell(row=row, column=1, value=shared_sanitize_excel(item.get('shelfmark', '')))
            ws.cell(row=row, column=2, value=shared_sanitize_excel(library_name))
            ws.cell(row=row, column=3, value=shared_sanitize_excel(self._format_image_display(item.get('img'))))
            ws.cell(row=row, column=4, value=shared_sanitize_excel(item.get('title', '')))
            ws.cell(row=row, column=5, value=shared_sanitize_excel(sys_id))
            ws.cell(row=row, column=6, value=shared_sanitize_excel(item.get('source', '')))
            ws.cell(row=row, column=7, value=shared_sanitize_excel(', '.join(item.get('tags', []))))
            ws.cell(row=row, column=8, value=shared_sanitize_excel(item.get('notes', '')))

        default_name = f"{list_name}.xlsx"
        path, _ = QFileDialog.getSaveFileName(
            self, tr("Export List"),
            os.path.join(Config.REPORTS_DIR, default_name),
            tr("Excel Files") + " (*.xlsx)"
        )

        if path:
            wb.save(path)
            self.status_label.setText(tr("List exported successfully."))

    def _export_as_word(self, list_id, list_name, items):
        """Export list as Word document."""
        try:
            from docx import Document
        except ImportError:
            QMessageBox.warning(self, tr("Export Error"),
                                tr("Word export requires the 'python-docx' package. Install it with: pip install python-docx"))
            return

        doc = Document()
        doc.add_heading(list_name, 0)
        doc.add_paragraph(f"{tr('Total items')}: {len(items)}")

        for i, item in enumerate(items, 1):
            shelfmark = item.get('shelfmark', 'Unknown')
            title = item.get('title', '')
            sys_id = item.get('sys_id', '')
            img = item.get('img')
            source = item.get('source', '')
            tags = item.get('tags', [])
            notes = item.get('notes', '')
            # Get library info
            library_code = self.meta_mgr.get_library_for_id(sys_id) if sys_id else ''
            library_name = get_library_display(library_code, short=False) if library_code else ''

            doc.add_heading(f"{i}. {shelfmark}", level=2)
            if library_name:
                doc.add_paragraph(f"{tr('Library')}: {library_name}")
            if img not in (None, ""):
                doc.add_paragraph(f"{tr('Image')}: {self._format_image_display(img)}")
            if title:
                doc.add_paragraph(title)
            doc.add_paragraph(f"ID: {sys_id}")
            if source:
                doc.add_paragraph(f"{tr('Source')}: {source}")
            if tags:
                doc.add_paragraph(f"{tr('Tags')}: {', '.join(tags)}")
            if notes:
                doc.add_paragraph(f"{tr('Notes')}: {notes}")

        default_name = f"{list_name}.docx"
        path, _ = QFileDialog.getSaveFileName(
            self, tr("Export List"),
            os.path.join(Config.REPORTS_DIR, default_name),
            tr("Word Files") + " (*.docx)"
        )

        if path:
            doc.save(path)
            self.status_label.setText(tr("List exported successfully."))

    def _save_text_to_file(self, text, list_name):
        """Save text to file."""
        default_name = f"{list_name}.txt"
        path, _ = QFileDialog.getSaveFileName(
            self, tr("Export List"),
            os.path.join(Config.REPORTS_DIR, default_name),
            tr("Text Files") + " (*.txt)"
        )

        if path:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(text)
            self.status_label.setText(tr("List exported successfully."))

    def _copy_to_clipboard(self, text):
        """Copy text to clipboard."""
        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        self.status_label.setText(tr("Copied to clipboard."))

    def _send_by_email(self, text, subject):
        """
        Robust Email: Copies text to clipboard and opens empty email draft.
        Directly injecting body text into mailto links fails with Hebrew/Long text.
        """
        import urllib.parse
        
        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        
        QMessageBox.information(
            self, 
            tr("Email"), 
            tr("The items have been copied to your clipboard.\n\n"
               "Your email client will now open.\n"
               "Please paste (Ctrl+V) the text into the message body.")
        )
        
        subject_encoded = urllib.parse.quote(f"GenizahSearch - {subject}")
        
        QDesktopServices.openUrl(QUrl(f"mailto:?subject={subject_encoded}"))

    def lists_apply_filter(self, text):
        """Apply filter to items table."""
        self.lists_refresh_items()

    # --- Add to List from other places ---

    def show_add_to_list_menu(self, items, source='', anchor_widget=None):
        """Show menu for adding items to a list."""
        if not self.lists_mgr or not items:
            return

        menu = QMenu(self)

        # Get all lists (excluding recent)
        lists = self.lists_mgr.get_all_lists(include_recent=False)
        projects = {proj['id']: proj['name'] for proj in self.lists_mgr.get_projects()}
        project_lists = {}
        top_level_lists = []

        for lst in lists:
            project_id = lst.get('project_id')
            if project_id and project_id in projects:
                project_lists.setdefault(project_id, []).append(lst)
            else:
                top_level_lists.append(lst)

        for lst in top_level_lists:
            action = menu.addAction(f"● {self._get_list_display_name(lst)}")
            action.setData(lst['id'])

        for project in self.lists_mgr.get_projects():
            proj_lists = project_lists.get(project['id'], [])
            if not proj_lists:
                continue
            submenu = menu.addMenu(project['name'])
            for lst in proj_lists:
                action = submenu.addAction(f"● {self._get_list_display_name(lst)}")
                action.setData(lst['id'])

        menu.addSeparator()
        action_new = menu.addAction(tr("New List..."))

        if anchor_widget:
            pos = anchor_widget.mapToGlobal(anchor_widget.rect().bottomLeft())
        else:
            pos = QCursor.pos()

        action = menu.exec(pos)

        if action:
            if action == action_new:
                name, ok = QInputDialog.getText(self, tr("Create New List"), tr("List Name:"))
                if ok and name.strip():
                    list_id = self.lists_mgr.create_list(name.strip())
                    self.lists_mgr.add_items_bulk(items, list_id, source=source)
                    self.status_label.setText(tr("Added to list."))
                    self.lists_refresh_all()  # Refresh to show new list
            else:
                list_id = action.data()
                if list_id:
                    added = self.lists_mgr.add_items_bulk(items, list_id, source=source)
                    if added > 0:
                        self.status_label.setText(tr("Added to list."))
                        self.lists_refresh_all()  # Refresh to show new items
                    else:
                        # Items already in list
                        QMessageBox.information(self, tr("Already in list"),
                                                tr("Items are already in this list."))

    # ==========================================================================
    #  COMMUNITY TAB
    # ==========================================================================

    def create_community_tab(self):
        """Create the Community tab with panels for discoveries, corrections, and comments."""
        panel = QWidget()
        main_layout = QVBoxLayout(panel)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(10)

        # Header with user info (wrap in widget to control height)
        header_widget = QWidget()
        header_widget.setFixedHeight(30)
        header_layout = QHBoxLayout(header_widget)
        header_layout.setContentsMargins(0, 0, 0, 0)
        self.community_user_label = QLabel()
        self.community_user_label.setStyleSheet("font-size: 14px;")
        header_layout.addWidget(self.community_user_label)
        header_layout.addStretch()

        # Register button (shown when not logged in)
        self.community_register_btn = QPushButton(tr("Register"))
        self.community_register_btn.clicked.connect(self._show_register_dialog)
        header_layout.addWidget(self.community_register_btn)

        main_layout.addWidget(header_widget)

        # Create horizontal splitter for the three panels
        splitter = QSplitter(Qt.Orientation.Horizontal)

        # --- Discoveries Panel ---
        discoveries_panel = QWidget()
        discoveries_layout = QVBoxLayout(discoveries_panel)
        discoveries_layout.setContentsMargins(5, 5, 5, 5)

        discoveries_header = QHBoxLayout()
        discoveries_header.addWidget(QLabel(f"<b>{tr('Discoveries')}</b>"))
        discoveries_header.addStretch()

        # Store discoveries data for filtering (must be before filter creation)
        self._discoveries_cache_data = []

        # Category filter dropdown
        self.discoveries_filter = QComboBox()
        self.discoveries_filter.blockSignals(True)  # Block during setup
        self.discoveries_filter.addItem(tr("All Types"), "all")
        self.discoveries_filter.addItem(f"📜 {tr('Discovery')}", "discovery")
        self.discoveries_filter.addItem(f"❓ {tr('Question')}", "question")
        self.discoveries_filter.addItem(f"🔍 {tr('Identification')}", "identification")
        self.discoveries_filter.addItem(f"📝 {tr('Note')}", "note")
        self.discoveries_filter.setFixedWidth(140)
        self.discoveries_filter.blockSignals(False)  # Unblock after setup
        self.discoveries_filter.currentIndexChanged.connect(self._filter_discoveries)
        discoveries_header.addWidget(self.discoveries_filter)

        btn_new_discovery = QPushButton("+")
        btn_new_discovery.setFixedSize(24, 24)
        btn_new_discovery.setToolTip(tr("Share a new discovery"))
        btn_new_discovery.clicked.connect(self._show_create_discovery_dialog)
        discoveries_header.addWidget(btn_new_discovery)
        btn_refresh_discoveries = QPushButton("↻")
        btn_refresh_discoveries.setFixedSize(24, 24)
        btn_refresh_discoveries.setToolTip(tr("Refresh discoveries"))
        btn_refresh_discoveries.clicked.connect(self._refresh_discoveries_panel)
        discoveries_header.addWidget(btn_refresh_discoveries)
        discoveries_layout.addLayout(discoveries_header)

        self.discoveries_list = QListWidget()
        self.discoveries_list.setAlternatingRowColors(True)
        self.discoveries_list.itemDoubleClicked.connect(self._on_discovery_clicked)
        self.discoveries_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.discoveries_list.customContextMenuRequested.connect(self._discoveries_context_menu)
        discoveries_layout.addWidget(self.discoveries_list)

        btn_view_all_discoveries = QPushButton(tr("View All Discoveries..."))
        btn_view_all_discoveries.clicked.connect(self._show_discoveries_dialog)
        discoveries_layout.addWidget(btn_view_all_discoveries)

        splitter.addWidget(discoveries_panel)

        # --- Corrections Panel ---
        corrections_panel = QWidget()
        corrections_layout = QVBoxLayout(corrections_panel)
        corrections_layout.setContentsMargins(5, 5, 5, 5)

        corrections_header = QHBoxLayout()
        corrections_header.addWidget(QLabel(f"<b>{tr('Corrections')}</b>"))
        corrections_header.addStretch()
        btn_refresh_corrections = QPushButton("↻")
        btn_refresh_corrections.setFixedSize(24, 24)
        btn_refresh_corrections.setToolTip(tr("Refresh corrections"))
        btn_refresh_corrections.clicked.connect(self._refresh_corrections_panel)
        corrections_header.addWidget(btn_refresh_corrections)
        corrections_layout.addLayout(corrections_header)

        # Sub-tabs for All Corrections vs My Corrections
        corrections_tabs = QTabWidget()
        corrections_tabs.setTabPosition(QTabWidget.TabPosition.South)

        # All Corrections list (shown first)
        self.all_corrections_list = QListWidget()
        self.all_corrections_list.setAlternatingRowColors(True)
        self.all_corrections_list.itemDoubleClicked.connect(self._on_correction_clicked)
        self.all_corrections_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.all_corrections_list.customContextMenuRequested.connect(lambda pos: self._corrections_context_menu(pos, self.all_corrections_list))
        corrections_tabs.addTab(self.all_corrections_list, tr("All Corrections"))

        # My Corrections list
        self.my_corrections_list = QListWidget()
        self.my_corrections_list.setAlternatingRowColors(True)
        self.my_corrections_list.itemDoubleClicked.connect(self._on_correction_clicked)
        self.my_corrections_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.my_corrections_list.customContextMenuRequested.connect(lambda pos: self._corrections_context_menu(pos, self.my_corrections_list))
        corrections_tabs.addTab(self.my_corrections_list, tr("My Corrections"))

        corrections_layout.addWidget(corrections_tabs)

        btn_browse_corrections = QPushButton(tr("Browse Corrections..."))
        btn_browse_corrections.clicked.connect(self._show_all_corrections_dialog)
        corrections_layout.addWidget(btn_browse_corrections)

        splitter.addWidget(corrections_panel)

        # --- Comments Panel ---
        comments_panel = QWidget()
        comments_layout = QVBoxLayout(comments_panel)
        comments_layout.setContentsMargins(5, 5, 5, 5)

        comments_header = QHBoxLayout()
        comments_header.addWidget(QLabel(f"<b>{tr('Comments')}</b>"))
        comments_header.addStretch()
        btn_refresh_comments = QPushButton("↻")
        btn_refresh_comments.setFixedSize(24, 24)
        btn_refresh_comments.setToolTip(tr("Refresh comments"))
        btn_refresh_comments.clicked.connect(self._refresh_comments_panel)
        comments_header.addWidget(btn_refresh_comments)
        comments_layout.addLayout(comments_header)

        # Sub-tabs for All Comments vs My Comments (like Corrections)
        comments_tabs = QTabWidget()
        comments_tabs.setTabPosition(QTabWidget.TabPosition.South)

        # All Comments list (shown first)
        self.all_comments_list = QListWidget()
        self.all_comments_list.setAlternatingRowColors(True)
        self.all_comments_list.itemDoubleClicked.connect(self._on_comment_clicked)
        self.all_comments_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.all_comments_list.customContextMenuRequested.connect(lambda pos: self._comments_context_menu(pos, self.all_comments_list))
        comments_tabs.addTab(self.all_comments_list, tr("All Comments"))

        # My Comments list
        self.my_comments_list = QListWidget()
        self.my_comments_list.setAlternatingRowColors(True)
        self.my_comments_list.itemDoubleClicked.connect(self._on_comment_clicked)
        self.my_comments_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.my_comments_list.customContextMenuRequested.connect(lambda pos: self._comments_context_menu(pos, self.my_comments_list))
        comments_tabs.addTab(self.my_comments_list, tr("My Comments"))

        comments_layout.addWidget(comments_tabs)

        btn_view_my_comments = QPushButton(tr("View All My Comments..."))
        btn_view_my_comments.clicked.connect(self._show_my_comments_dialog)
        comments_layout.addWidget(btn_view_my_comments)

        splitter.addWidget(comments_panel)

        # --- Joins Panel ---
        joins_panel = QWidget()
        joins_layout = QVBoxLayout(joins_panel)
        joins_layout.setContentsMargins(5, 5, 5, 5)

        joins_header = QHBoxLayout()
        joins_header.addWidget(QLabel(f"<b>{tr('Joins')}</b>"))
        joins_header.addStretch()
        btn_refresh_joins = QPushButton("↻")
        btn_refresh_joins.setFixedSize(24, 24)
        btn_refresh_joins.setToolTip(tr("Refresh joins"))
        btn_refresh_joins.clicked.connect(self._refresh_joins_panel)
        joins_header.addWidget(btn_refresh_joins)
        joins_layout.addLayout(joins_header)

        # Sub-tabs for All Joins vs My Joins
        joins_tabs = QTabWidget()
        joins_tabs.setTabPosition(QTabWidget.TabPosition.South)

        # All Joins list (shown first)
        self.all_joins_list = QListWidget()
        self.all_joins_list.setAlternatingRowColors(True)
        self.all_joins_list.itemDoubleClicked.connect(self._on_join_clicked)
        self.all_joins_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.all_joins_list.customContextMenuRequested.connect(lambda pos: self._joins_context_menu(pos, self.all_joins_list))
        joins_tabs.addTab(self.all_joins_list, tr("All Joins"))

        # My Joins list
        self.my_joins_list = QListWidget()
        self.my_joins_list.setAlternatingRowColors(True)
        self.my_joins_list.itemDoubleClicked.connect(self._on_join_clicked)
        self.my_joins_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.my_joins_list.customContextMenuRequested.connect(lambda pos: self._joins_context_menu(pos, self.my_joins_list))
        joins_tabs.addTab(self.my_joins_list, tr("My Joins"))

        joins_layout.addWidget(joins_tabs)

        btn_browse_joins = QPushButton(tr("Browse Joins..."))
        btn_browse_joins.clicked.connect(self._show_joins_feed_dialog)
        joins_layout.addWidget(btn_browse_joins)

        splitter.addWidget(joins_panel)

        # Set equal sizes for all panels
        splitter.setSizes([250, 250, 250, 250])

        main_layout.addWidget(splitter)

        return panel

    def _refresh_community_panels(self, use_cache_first=True):
        """Refresh all community panels and update UI state.

        Args:
            use_cache_first: If True, display cached data first for instant response,
                           then fetch fresh data in background.
        """
        logger.debug("_refresh_community_panels started")

        # Quick connectivity check to avoid long timeouts when offline
        server_available = self.corrections_client.is_server_available()
        logger.debug("Server available: %s", server_available)

        # If offline, only use cached data - skip all API calls
        skip_api_calls = not server_available

        try:
            logger.debug("Calling _update_community_header...")
            self._update_community_header()
            logger.debug("_update_community_header completed")
        except Exception as e:
            print(f"Error in _update_community_header: {e}", flush=True)
        try:
            logger.debug("Calling _refresh_discoveries_panel...")
            self._refresh_discoveries_panel(use_cache_first, skip_api_calls=skip_api_calls)
            logger.debug("_refresh_discoveries_panel completed")
        except Exception as e:
            print(f"Error in _refresh_discoveries_panel: {e}", flush=True)
            import traceback
            traceback.print_exc()
        try:
            logger.debug("Calling _refresh_corrections_panel...")
            self._refresh_corrections_panel(use_cache_first, skip_api_calls=skip_api_calls)
            logger.debug("_refresh_corrections_panel completed")
        except Exception as e:
            print(f"Error in _refresh_corrections_panel: {e}", flush=True)
        try:
            logger.debug("Calling _refresh_comments_panel...")
            self._refresh_comments_panel(use_cache_first, skip_api_calls=skip_api_calls)
            logger.debug("_refresh_comments_panel completed")
        except Exception as e:
            print(f"Error in _refresh_comments_panel: {e}", flush=True)
        try:
            logger.debug("Calling _refresh_joins_panel...")
            self._refresh_joins_panel(use_cache_first, skip_api_calls=skip_api_calls)
            logger.debug("_refresh_joins_panel completed")
        except Exception as e:
            print(f"Error in _refresh_joins_panel: {e}", flush=True)
        logger.debug("_refresh_community_panels finished")

    def _update_community_header(self):
        """Update the community header with user info."""
        if self.corrections_client.is_logged_in() and self.corrections_client.current_user:
            user = self.corrections_client.current_user
            self.community_user_label.setText(
                f"👤 {tr('Logged in as')} <b>{user.username}</b> | ⭐ {tr('Reputation')}: {user.reputation_score}"
            )
            self.community_register_btn.hide()
        else:
            self.community_user_label.setText(f"👤 {tr('Not logged in')} - {tr('Login to participate in the community')}")
            self.community_register_btn.show()

    def _refresh_discoveries_panel(self, use_cache_first=True, skip_api_calls=False):
        """Refresh the discoveries list panel."""
        self.discoveries_list.clear()

        # Try to display cached data first for instant response
        cached = self.corrections_client.get_cached_data('discoveries') if use_cache_first else None
        if cached:
            self._discoveries_cache_data = cached
            self._populate_discoveries_list(cached)

        # Skip API calls if offline
        if skip_api_calls:
            if not cached:
                item = QListWidgetItem(f"ℹ️ {tr('Offline - no cached data available')}")
                self.discoveries_list.addItem(item)
            return

        # Fetch fresh data from API
        try:
            # Admins can see hidden discoveries
            is_admin = False
            if self.corrections_client.is_logged_in():
                current_user = self.corrections_client.current_user
                is_admin = current_user and current_user.role == 'admin'

            logger.debug("_refresh_discoveries_panel: is_admin=%s, include_hidden=%s", is_admin, is_admin)

            discoveries, total = self.corrections_client.get_discoveries(
                page_size=20,
                include_hidden=is_admin
            )
            logger.debug("Got %s discoveries from API, total=%s", len(discoveries), total)
            for d in discoveries:
                logger.debug("discovery id=%s, title=%s, is_hidden=%s", d.id, d.title[:30] if d.title else 'N/A', d.is_hidden)

            # Convert to cacheable format with discovery_type, is_pinned, is_hidden
            cache_data = [{
                'id': d.id,
                'title': d.title,
                'author_username': d.author_username,
                'discovery_type': d.discovery_type,
                'shelfmark': d.shelfmark,
                'document_id': d.document_id,
                'is_pinned': d.is_pinned,
                'is_hidden': d.is_hidden
            } for d in discoveries]
            self.corrections_client.set_cached_data('discoveries', cache_data)
            self._discoveries_cache_data = cache_data

            # Update display with fresh data
            self.discoveries_list.clear()
            self._filter_discoveries()  # Apply current filter
        except Exception as e:
            if not cached:
                item = QListWidgetItem(f"⚠️ {tr('Error loading discoveries')}: {str(e)[:30]}")
                self.discoveries_list.addItem(item)

    def _filter_discoveries(self):
        """Filter discoveries list by selected type."""
        if not hasattr(self, '_discoveries_cache_data'):
            self._discoveries_cache_data = []

        # Check if current user is admin
        is_admin = False
        if self.corrections_client.is_logged_in():
            current_user = self.corrections_client.current_user
            is_admin = current_user and current_user.role == 'admin'

        # Filter by type
        filter_type = self.discoveries_filter.currentData()
        if filter_type == "all":
            filtered = self._discoveries_cache_data
        else:
            filtered = [d for d in self._discoveries_cache_data if d.get('discovery_type') == filter_type]

        # Filter out hidden items for non-admins
        if not is_admin:
            filtered = [d for d in filtered if not d.get('is_hidden', False)]

        # Sort: pinned items first, then by id (newest first)
        filtered = sorted(filtered, key=lambda d: (not d.get('is_pinned', False), -d.get('id', 0)))

        self.discoveries_list.clear()
        self._populate_discoveries_list(filtered)

    def _populate_discoveries_list(self, discoveries_data):
        """Populate discoveries list from data."""
        if not discoveries_data:
            item = QListWidgetItem(f"ℹ️ {tr('No discoveries yet')}")
            self.discoveries_list.addItem(item)
            return

        # Type icons and colors
        type_config = {
            'discovery': ('📜', '#f39c12'),   # Orange
            'question': ('❓', '#3498db'),    # Blue
            'identification': ('🔍', '#9b59b6'),  # Purple
            'note': ('📝', '#27ae60')         # Green
        }

        for disc in discoveries_data:
            title = disc.get('title') or tr('Untitled')
            author = disc.get('author_username') or tr('Anonymous')
            disc_type = disc.get('discovery_type') or 'discovery'
            is_pinned = disc.get('is_pinned', False)
            is_hidden = disc.get('is_hidden', False)
            icon, color = type_config.get(disc_type, ('📜', '#f39c12'))

            # Add pin icon for pinned items
            prefix = "📌 " if is_pinned else ""
            # Add hidden indicator for admin view
            suffix = " [hidden]" if is_hidden else ""

            item = QListWidgetItem(f"{prefix}{icon} {title}{suffix}\n   {tr('by')} {author}")
            item.setData(Qt.ItemDataRole.UserRole, {
                'id': disc.get('id'),
                'title': disc.get('title'),
                'shelfmark': disc.get('shelfmark'),
                'document_id': disc.get('document_id')
            })
            # Set foreground color based on type (gray for hidden)
            if is_hidden:
                item.setForeground(QColor('#888888'))
            else:
                item.setForeground(QColor(color))
            self.discoveries_list.addItem(item)

    def _refresh_corrections_panel(self, use_cache_first=True, skip_api_calls=False):
        """Refresh the corrections list panels."""
        self.my_corrections_list.clear()
        self.all_corrections_list.clear()

        # My corrections (only if logged in)
        if self.corrections_client.is_logged_in():
            # Try cached data first
            cached_my = self.corrections_client.get_cached_data('my_corrections') if use_cache_first else None
            if cached_my:
                self._populate_my_corrections_list(cached_my)

            if not skip_api_calls:
                try:
                    corrections, total = self.corrections_client.get_my_corrections(page_size=20)
                    cache_data = [{'id': c.id, 'shelfmark': c.shelfmark, 'system_id': c.system_id, 'status': c.status, 'corrected_text': c.corrected_text, 'page_number': c.page_number} for c in corrections]
                    self.corrections_client.set_cached_data('my_corrections', cache_data)
                    self.my_corrections_list.clear()
                    self._populate_my_corrections_list(cache_data)
                except Exception as e:
                    if not cached_my:
                        item = QListWidgetItem(f"⚠️ {tr('Error')}: {str(e)[:30]}")
                        self.my_corrections_list.addItem(item)
            elif not cached_my:
                item = QListWidgetItem(f"ℹ️ {tr('Offline - no cached data available')}")
                self.my_corrections_list.addItem(item)

        # All corrections
        cached_all = self.corrections_client.get_cached_data('all_corrections') if use_cache_first else None
        if cached_all:
            self._populate_all_corrections_list(cached_all)

        if not skip_api_calls:
            try:
                corrections, total = self.corrections_client.get_all_corrections(page_size=20)
                cache_data = [{'id': c.id, 'shelfmark': c.shelfmark, 'system_id': c.system_id, 'author_username': c.author_username, 'status': c.status, 'page_number': c.page_number} for c in corrections]
                self.corrections_client.set_cached_data('all_corrections', cache_data)
                self.all_corrections_list.clear()
                self._populate_all_corrections_list(cache_data)
            except Exception as e:
                if not cached_all:
                    item = QListWidgetItem(f"⚠️ {tr('Error')}: {str(e)[:30]}")
                    self.all_corrections_list.addItem(item)
        elif not cached_all:
            item = QListWidgetItem(f"ℹ️ {tr('Offline - no cached data available')}")
            self.all_corrections_list.addItem(item)

    def _populate_my_corrections_list(self, corrections_data):
        """Populate my corrections list from data (only latest per document)."""
        if not corrections_data:
            item = QListWidgetItem(f"ℹ️ {tr('No corrections submitted yet')}")
            self.my_corrections_list.addItem(item)
            return

        # Deduplicate: keep only the latest correction per document (by system_id or shelfmark)
        latest_by_doc = {}
        for corr in corrections_data:
            doc_key = corr.get('system_id') or corr.get('shelfmark') or 'unknown'
            corr_id = corr.get('id', 0)
            if doc_key not in latest_by_doc or corr_id > latest_by_doc[doc_key].get('id', 0):
                latest_by_doc[doc_key] = corr

        for corr in latest_by_doc.values():
            shelfmark = corr.get('shelfmark') or tr('Unknown')
            page_number = corr.get('page_number')
            # Display shelfmark with page number if available
            display_shelfmark = shelfmark
            if page_number:
                display_shelfmark = f"{shelfmark}:{page_number}"
            status = corr.get('status') or 'pending'
            status_icon = {'pending': '🔵', 'approved': '✅', 'rejected': '❌', 'draft': '📝'}.get(status, '⚪')
            corrected_text = corr.get('corrected_text') or ''
            preview = (corrected_text[:30] + '...') if len(corrected_text) > 30 else corrected_text
            item = QListWidgetItem(f"{status_icon} {display_shelfmark}\n   {preview}")
            item.setData(Qt.ItemDataRole.UserRole, {
                'id': corr.get('id'),
                'shelfmark': shelfmark,
                'system_id': corr.get('system_id'),
                'page_number': page_number
            })
            self.my_corrections_list.addItem(item)

    def _populate_all_corrections_list(self, corrections_data):
        """Populate all corrections list from data (only latest per user per document)."""
        if not corrections_data:
            item = QListWidgetItem(f"ℹ️ {tr('No corrections yet')}")
            self.all_corrections_list.addItem(item)
            return

        # Deduplicate: keep only the latest correction per user per document
        latest_by_user_doc = {}
        for corr in corrections_data:
            doc_key = corr.get('system_id') or corr.get('shelfmark') or 'unknown'
            author = corr.get('author_username') or 'anonymous'
            key = f"{author}:{doc_key}"
            corr_id = corr.get('id', 0)
            if key not in latest_by_user_doc or corr_id > latest_by_user_doc[key].get('id', 0):
                latest_by_user_doc[key] = corr

        for corr in latest_by_user_doc.values():
            shelfmark = corr.get('shelfmark') or tr('Unknown')
            page_number = corr.get('page_number')
            # Display shelfmark with page number if available
            display_shelfmark = shelfmark
            if page_number:
                display_shelfmark = f"{shelfmark}:{page_number}"
            author = corr.get('author_username') or tr('Anonymous')
            status = corr.get('status') or 'pending'
            status_icon = {'pending': '🔵', 'approved': '✅', 'rejected': '❌', 'draft': '📝'}.get(status, '⚪')
            item = QListWidgetItem(f"{status_icon} {display_shelfmark}\n   {tr('by')} {author}")
            item.setData(Qt.ItemDataRole.UserRole, {
                'id': corr.get('id'),
                'shelfmark': shelfmark,
                'system_id': corr.get('system_id'),
                'page_number': page_number
            })
            self.all_corrections_list.addItem(item)

    def _refresh_comments_panel(self, use_cache_first=True, skip_api_calls=False):
        """Refresh the comments list panels (My Comments + All Comments)."""
        logger.debug("_refresh_comments_panel started")
        self.my_comments_list.clear()
        self.all_comments_list.clear()

        # My Comments (only if logged in)
        if self.corrections_client.is_logged_in():
            cached_my = self.corrections_client.get_cached_data('my_comments') if use_cache_first else None
            if cached_my:
                logger.debug("Using cached my comments: %s items", len(cached_my))
                self._populate_comments_list(cached_my, self.my_comments_list)

            if not skip_api_calls:
                try:
                    comments, total = self.corrections_client.get_my_comments(page_size=20)
                    logger.debug("Got %s my comments, total=%s", len(comments), total)
                    cache_data = [{'id': c.id, 'document_id': c.document_id, 'content': c.content, 'author_username': c.author_username, 'page_number': c.page_number} for c in comments]
                    self.corrections_client.set_cached_data('my_comments', cache_data)
                    self.my_comments_list.clear()
                    self._populate_comments_list(cache_data, self.my_comments_list)
                except Exception as e:
                    logger.debug("Error fetching my comments: %s", e)
                    if not cached_my:
                        item = QListWidgetItem(f"⚠️ {tr('Error')}: {str(e)[:30]}")
                        self.my_comments_list.addItem(item)
            elif not cached_my:
                item = QListWidgetItem(f"ℹ️ {tr('Offline - no cached data available')}")
                self.my_comments_list.addItem(item)
        else:
            item = QListWidgetItem(f"ℹ️ {tr('Login to see your comments')}")
            self.my_comments_list.addItem(item)

        # All Comments (no login required)
        cached_all = self.corrections_client.get_cached_data('all_comments') if use_cache_first else None
        if cached_all:
            self._populate_comments_list(cached_all, self.all_comments_list, show_author=True)

        if not skip_api_calls:
            try:
                comments, total = self.corrections_client.get_all_comments(page_size=20)
                logger.debug("Got %s all comments, total=%s", len(comments), total)
                cache_data = [{'id': c.id, 'document_id': c.document_id, 'content': c.content, 'author_username': c.author_username, 'page_number': c.page_number} for c in comments]
                self.corrections_client.set_cached_data('all_comments', cache_data)
                self.all_comments_list.clear()
                self._populate_comments_list(cache_data, self.all_comments_list, show_author=True)
            except Exception as e:
                logger.debug("Error fetching all comments: %s", e)
                if not cached_all:
                    item = QListWidgetItem(f"⚠️ {tr('Error')}: {str(e)[:30]}")
                    self.all_comments_list.addItem(item)
        elif not cached_all:
            item = QListWidgetItem(f"ℹ️ {tr('Offline - no cached data available')}")
            self.all_comments_list.addItem(item)

        logger.debug("Comments panel refresh completed")

    def _populate_comments_list(self, comments_data, target_list, show_author=False):
        """Populate comments list from data."""
        if not comments_data:
            item = QListWidgetItem(f"ℹ️ {tr('No comments yet')}")
            target_list.addItem(item)
            return
        for comment in comments_data:
            sys_id = comment.get('document_id')  # This is actually a system ID
            page_number = comment.get('page_number')
            content = comment.get('content') or ''
            text = (content[:50] + '...') if len(content) > 50 else content
            author = comment.get('author_username') or tr('Anonymous')

            # Get shelfmark and title from sys_id
            shelfmark = sys_id or tr('Unknown')
            title_preview = ''
            if sys_id and self.meta_mgr:
                shelf, title = self.meta_mgr.get_meta_for_id(sys_id)
                if shelf:
                    shelfmark = shelf
                if title:
                    words = title.split()[:4]
                    title_preview = ' '.join(words)
                    if len(title.split()) > 4:
                        title_preview += '...'

            # Display shelfmark with page number if available
            display_shelfmark = shelfmark
            if page_number:
                display_shelfmark = f"{shelfmark}:{page_number}"

            display_text = f"💬 {display_shelfmark}"
            if title_preview:
                display_text += f" - {title_preview}"
            if show_author:
                display_text += f"\n   by {author}"
            display_text += f"\n   {text}"

            item = QListWidgetItem(display_text)
            item.setData(Qt.ItemDataRole.UserRole, {'id': comment.get('id'), 'document_id': sys_id, 'page_number': page_number})
            target_list.addItem(item)

    def _discoveries_context_menu(self, pos):
        """Show context menu for discoveries list."""
        item = self.discoveries_list.itemAt(pos)
        if not item:
            return

        disc_data = item.data(Qt.ItemDataRole.UserRole)
        if not disc_data:
            return

        menu = QMenu(self)
        discovery_id = disc_data.get('id')

        # View action
        view_action = menu.addAction("👁 " + tr("View Details"))
        view_action.triggered.connect(lambda: self._on_discovery_clicked(item))

        # Get full discovery data for author check
        discovery = self.corrections_client.get_discovery(discovery_id)
        if not discovery:
            menu.exec(self.discoveries_list.mapToGlobal(pos))
            return

        current_user = self.corrections_client.current_user
        is_author = current_user and current_user.id == discovery.author_id
        is_admin = current_user and current_user.role == 'admin'

        menu.addSeparator()

        # Edit action (author or admin)
        if is_author or is_admin:
            edit_action = menu.addAction("✏️ " + tr("Edit"))
            edit_action.triggered.connect(lambda: self._edit_discovery_from_list(discovery_id))

            delete_action = menu.addAction("🗑️ " + tr("Delete"))
            delete_action.triggered.connect(lambda: self._delete_discovery_from_list(discovery_id))

        # Admin actions
        if is_admin:
            menu.addSeparator()
            pin_text = tr("Unpin") if discovery.is_pinned else tr("Pin")
            pin_action = menu.addAction("📌 " + pin_text)
            pin_action.triggered.connect(lambda: self._toggle_pin_discovery(discovery_id, not discovery.is_pinned))

            hide_text = tr("Unhide") if discovery.is_hidden else tr("Hide")
            hide_action = menu.addAction("👁 " + hide_text)
            hide_action.triggered.connect(lambda: self._toggle_hide_discovery(discovery_id, discovery.is_hidden))

        menu.exec(self.discoveries_list.mapToGlobal(pos))

    def _edit_discovery_from_list(self, discovery_id):
        """Open edit dialog for discovery from context menu."""
        dialog = DiscoveryDetailDialog(self, self.corrections_client, discovery_id)
        dialog.edit_discovery()
        self._refresh_discoveries_panel(use_cache_first=False)

    def _delete_discovery_from_list(self, discovery_id):
        """Delete discovery from context menu."""
        reply = QMessageBox.question(
            self,
            tr("Delete Discovery"),
            tr("Are you sure you want to delete this discovery?"),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            success, msg = self.corrections_client.delete_discovery(discovery_id)
            if success:
                QMessageBox.information(self, tr("Success"), tr("Discovery deleted"))
                self._refresh_discoveries_panel(use_cache_first=False)
            else:
                QMessageBox.warning(self, tr("Error"), msg)

    def _toggle_pin_discovery(self, discovery_id, pin):
        """Toggle pin status from context menu."""
        success, msg = self.corrections_client.pin_discovery(discovery_id, pin)
        if success:
            self._refresh_discoveries_panel(use_cache_first=False)
        else:
            QMessageBox.warning(self, tr("Error"), msg)

    def _toggle_hide_discovery(self, discovery_id, is_hidden):
        """Toggle hide status from context menu."""
        if is_hidden:
            success, msg = self.corrections_client.unhide_discovery(discovery_id)
        else:
            success, msg = self.corrections_client.hide_discovery(discovery_id)
        if success:
            self._refresh_discoveries_panel(use_cache_first=False)
        else:
            QMessageBox.warning(self, tr("Error"), msg)

    def _corrections_context_menu(self, pos, list_widget):
        """Show context menu for corrections list."""
        item = list_widget.itemAt(pos)
        if not item:
            return

        corr_data = item.data(Qt.ItemDataRole.UserRole)
        if not corr_data:
            return

        menu = QMenu(self)

        # View Document action
        view_doc_action = menu.addAction("📄 " + tr("View Document"))
        view_doc_action.triggered.connect(lambda: self._on_correction_clicked(item))

        # View Correction Details
        if corr_data.get('id'):
            view_details_action = menu.addAction("🔍 " + tr("View Correction Details"))
            view_details_action.triggered.connect(lambda: self._show_correction_details(corr_data.get('id')))

        menu.exec(list_widget.mapToGlobal(pos))

    def _show_correction_details(self, correction_id):
        """Show correction details dialog."""
        from corrections_ui import CorrectionDetailDialog
        # Fetch the correction object first
        correction = self.corrections_client.get_correction(correction_id)
        if correction:
            # Fetch the original V0.8 text using searcher.get_browse_page
            original_v08_text = None
            try:
                doc_id = correction.document_id or correction.system_id
                page_num = correction.page_number or 1
                if doc_id and hasattr(self, 'searcher') and self.searcher:
                    page_data = self.searcher.get_browse_page(doc_id, p_num=page_num)
                    if page_data:
                        original_v08_text = page_data.get('text', '')
                        logger.debug("Got V0.8 text from searcher: %s chars", len(original_v08_text) if original_v08_text else 0)
            except Exception as e:
                logger.debug("Error fetching V0.8 text: %s", e)

            dialog = CorrectionDetailDialog(self, self.corrections_client, correction, original_v08_text)
            dialog.exec()
        else:
            QMessageBox.warning(self, tr("Error"), tr("Could not load correction details"))

    def _comments_context_menu(self, pos, list_widget):
        """Show context menu for comments list."""
        item = list_widget.itemAt(pos)
        if not item:
            return

        comment_data = item.data(Qt.ItemDataRole.UserRole)
        if not comment_data:
            return

        menu = QMenu(self)

        # View Document action
        view_doc_action = menu.addAction("📄 " + tr("View Document"))
        view_doc_action.triggered.connect(lambda: self._on_comment_clicked(item))

        menu.exec(list_widget.mapToGlobal(pos))

    def _on_discovery_clicked(self, item):
        """Handle discovery item double-click."""
        disc = item.data(Qt.ItemDataRole.UserRole)
        if disc:
            dialog = DiscoveryDetailDialog(self, self.corrections_client, disc.get('id'))
            dialog.exec()

    def _on_correction_clicked(self, item):
        """Handle correction item double-click - open ResultDialog."""
        logger.debug("_on_correction_clicked called")
        try:
            corr_data = item.data(Qt.ItemDataRole.UserRole)
            logger.debug("corr_data=%s", corr_data)
            if corr_data:
                sys_id = corr_data.get('system_id')
                shelfmark = corr_data.get('shelfmark')
                page_num = corr_data.get('page_number') or 1
                logger.debug("sys_id=%s, shelfmark=%s, page_num=%s", sys_id, shelfmark, page_num)
                self._open_document_result_dialog(shelfmark=shelfmark, sys_id=sys_id, page_num=page_num)
        except Exception as e:
            logger.debug("Error in _on_correction_clicked: %s", e)
            import traceback
            traceback.print_exc()

    def _on_comment_clicked(self, item):
        """Handle comment item double-click - open ResultDialog."""
        comment = item.data(Qt.ItemDataRole.UserRole)
        if comment:
            doc_id = comment.get('document_id')
            page_num = comment.get('page_number') or 1
            if doc_id:
                # doc_id is actually a system ID, not a shelfmark
                self._open_document_result_dialog(sys_id=doc_id, page_num=page_num)

    # ========== Joins Panel Methods ==========

    def _refresh_joins_panel(self, use_cache_first=True, skip_api_calls=False):
        """Refresh the joins list panels (My Joins + All Joins)."""
        logger.debug("_refresh_joins_panel started")
        self.my_joins_list.clear()
        self.all_joins_list.clear()

        # My Joins (only if logged in)
        if self.corrections_client.is_logged_in():
            cached_my = self.corrections_client.get_cached_data('my_joins') if use_cache_first else None
            if cached_my:
                logger.debug("Using cached my joins: %s items", len(cached_my))
                self._populate_joins_list(cached_my, self.my_joins_list)

            if not skip_api_calls:
                try:
                    joins, total = self.corrections_client.get_my_joins(limit=20)
                    logger.debug("Got %s my joins, total=%s", len(joins), total)
                    cache_data = [{
                        'id': j.id, 'fragment_a': j.fragment_a, 'fragment_b': j.fragment_b,
                        'document_id_a': j.document_id_a, 'document_id_b': j.document_id_b,
                        'relationship_type': j.relationship_type, 'notes': j.notes,
                        'created_by_username': j.created_by_username
                    } for j in joins]
                    self.corrections_client.set_cached_data('my_joins', cache_data)
                    self.my_joins_list.clear()
                    self._populate_joins_list(cache_data, self.my_joins_list)
                except Exception as e:
                    logger.debug("Error fetching my joins: %s", e)
                    if not cached_my:
                        item = QListWidgetItem(f"⚠️ {tr('Error')}: {str(e)[:30]}")
                        self.my_joins_list.addItem(item)
            elif not cached_my:
                item = QListWidgetItem(f"ℹ️ {tr('Offline - no cached data available')}")
                self.my_joins_list.addItem(item)
        else:
            item = QListWidgetItem(f"ℹ️ {tr('Login to see your joins')}")
            self.my_joins_list.addItem(item)

        # All Joins (user-created only)
        cached_all = self.corrections_client.get_cached_data('all_joins') if use_cache_first else None
        if cached_all:
            self._populate_joins_list(cached_all, self.all_joins_list, show_author=True)

        if not skip_api_calls:
            try:
                joins, total = self.corrections_client.search_joins(source='user', limit=20)
                logger.debug("Got %s all joins, total=%s", len(joins), total)
                cache_data = [{
                    'id': j.id, 'fragment_a': j.fragment_a, 'fragment_b': j.fragment_b,
                    'document_id_a': j.document_id_a, 'document_id_b': j.document_id_b,
                    'relationship_type': j.relationship_type, 'notes': j.notes,
                    'created_by_username': j.created_by_username
                } for j in joins]
                self.corrections_client.set_cached_data('all_joins', cache_data)
                self.all_joins_list.clear()
                self._populate_joins_list(cache_data, self.all_joins_list, show_author=True)
            except Exception as e:
                logger.debug("Error fetching all joins: %s", e)
                if not cached_all:
                    item = QListWidgetItem(f"⚠️ {tr('Error')}: {str(e)[:30]}")
                    self.all_joins_list.addItem(item)
        elif not cached_all:
            item = QListWidgetItem(f"ℹ️ {tr('Offline - no cached data available')}")
            self.all_joins_list.addItem(item)

        logger.debug("Joins panel refresh completed")

    def _populate_joins_list(self, joins_data, target_list, show_author=False):
        """Populate joins list from data."""
        if not joins_data:
            item = QListWidgetItem(f"ℹ️ {tr('No joins yet')}")
            target_list.addItem(item)
            return

        rel_labels = {
            'physical_join': tr('Physical join'),
            'same_composition': tr('Same composition')
        }

        for join in joins_data:
            frag_a = join.get('fragment_a', '')
            frag_b = join.get('fragment_b', '')
            rel_type = join.get('relationship_type', '')
            author = join.get('created_by_username') or ''

            rel_display = rel_labels.get(rel_type, '')

            display_text = f"🔗 {frag_a} ↔ {frag_b}"
            if rel_display:
                display_text += f"\n   {rel_display}"
            if show_author and author:
                display_text += f"\n   by {author}"

            item = QListWidgetItem(display_text)
            item.setData(Qt.ItemDataRole.UserRole, join)
            target_list.addItem(item)

    def _joins_context_menu(self, pos, list_widget):
        """Show context menu for joins list."""
        item = list_widget.itemAt(pos)
        if not item:
            return

        join_data = item.data(Qt.ItemDataRole.UserRole)
        if not join_data:
            return

        menu = QMenu(self)

        # Open Fragment A
        frag_a = join_data.get('fragment_a', '')
        doc_id_a = join_data.get('document_id_a')
        if frag_a:
            action_a = menu.addAction(f"📄 {tr('Open')} {frag_a}")
            action_a.triggered.connect(lambda: self._open_join_fragment(frag_a, doc_id_a))

        # Open Fragment B
        frag_b = join_data.get('fragment_b', '')
        doc_id_b = join_data.get('document_id_b')
        if frag_b:
            action_b = menu.addAction(f"📄 {tr('Open')} {frag_b}")
            action_b.triggered.connect(lambda: self._open_join_fragment(frag_b, doc_id_b))

        menu.addSeparator()

        # Copy shelfmarks
        copy_action = menu.addAction("📋 " + tr("Copy shelfmarks"))
        copy_action.triggered.connect(lambda: self._copy_join_shelfmarks(frag_a, frag_b))

        # Delete (for own joins only or admin)
        is_my_list = (list_widget == self.my_joins_list)
        is_admin = (self.corrections_client.is_logged_in() and
                   self.corrections_client.current_user and
                   self.corrections_client.current_user.role == 'admin')

        if is_my_list or is_admin:
            menu.addSeparator()
            delete_action = menu.addAction("🗑️ " + tr("Delete join"))
            delete_action.triggered.connect(lambda: self._delete_join_from_list(join_data.get('id'), is_my_list))

        menu.exec(list_widget.mapToGlobal(pos))

    def _on_join_clicked(self, item):
        """Handle join item double-click - open Fragment A."""
        join_data = item.data(Qt.ItemDataRole.UserRole)
        if join_data:
            frag_a = join_data.get('fragment_a', '')
            doc_id_a = join_data.get('document_id_a')
            self._open_join_fragment(frag_a, doc_id_a)

    def _open_join_fragment(self, shelfmark, doc_id=None):
        """Open a fragment from a join - navigate to browse tab."""
        if doc_id:
            self._open_document_result_dialog(sys_id=doc_id)
        elif shelfmark:
            # Use shelfmark for browse
            self.browse_shelf_input.setText(shelfmark)
            self._set_last_browse_field("shelf")
            self.browse_load()
            self.tabs.setCurrentWidget(self.browse_tab)

    def _copy_join_shelfmarks(self, frag_a, frag_b):
        """Copy join shelfmarks to clipboard."""
        from PyQt6.QtWidgets import QApplication
        clipboard = QApplication.clipboard()
        clipboard.setText(f"{frag_a} ↔ {frag_b}")

    def _delete_join_from_list(self, join_id, is_my_list):
        """Delete a join from the community panel."""
        if not join_id:
            return

        reply = QMessageBox.question(
            self, tr("Confirm Delete"),
            tr("Are you sure you want to delete this join?"),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            success, msg = self.corrections_client.delete_join(join_id)
            if success:
                QMessageBox.information(self, tr("Success"), tr("Join deleted"))
                self._refresh_joins_panel(use_cache_first=False)
            else:
                QMessageBox.critical(self, tr("Error"), tr("Failed to delete: {}").format(msg))

    def _show_joins_feed_dialog(self):
        """Show the full joins feed dialog."""
        from corrections_ui import JoinsFeedDialog

        def browse_shelfmark(shelfmark):
            self.browse_shelf_input.setText(shelfmark)
            self._set_last_browse_field("shelf")
            self.browse_load()
            self.tabs.setCurrentWidget(self.browse_tab)

        dialog = JoinsFeedDialog(self, self.corrections_client, on_browse=browse_shelfmark)
        dialog.exec()

    def create_settings_tab(self):
        panel = QWidget(); layout = QVBoxLayout()

        settings_header = QHBoxLayout()
        settings_header.addStretch()
        btn_help = QPushButton("?")
        btn_help.setFixedWidth(30)
        btn_help.setStyleSheet("background-color: #f39c12; color: white; font-weight: bold; border-radius: 15px;")
        btn_help.clicked.connect(lambda: self.open_help_center(anchor="settings"))
        settings_header.addWidget(btn_help)
        layout.addLayout(settings_header)
        
        gb_data = QGroupBox(tr("Data & Index"))
        dl = QVBoxLayout()
        btn_dl = QPushButton(tr("Download Transcriptions (Zenodo)")); btn_dl.clicked.connect(lambda: QDesktopServices.openUrl(QUrl("https://doi.org/10.5281/zenodo.17734473")))
        dl.addWidget(btn_dl)
        self.btn_build_index = QPushButton(tr("Build / Rebuild Index")); self.btn_build_index.clicked.connect(self.run_indexing)
        self.btn_build_index.setEnabled(False)
        dl.addWidget(self.btn_build_index)
        self.index_progress = QProgressBar(); dl.addWidget(self.index_progress)
        gb_data.setLayout(dl); layout.addWidget(gb_data)

        # Application / Updates
        gb_app = QGroupBox(tr("Application"))
        app_layout = QHBoxLayout()

        self.lbl_version = QLabel(f"Version: {APP_VERSION}")
        self.btn_check_updates = QPushButton(tr("Check for Updates"))
        self.btn_check_updates.clicked.connect(self.check_updates_manual)

        app_layout.addWidget(self.lbl_version)
        app_layout.addStretch()
        app_layout.addWidget(self.btn_check_updates)
        gb_app.setLayout(app_layout)
        layout.addWidget(gb_app)
        
        gb_about = QGroupBox(tr("About"))
        abl = QVBoxLayout()
        about_html_en = """
        <style>
            h3 { margin-bottom: 0px; margin-top: 10px; }
            p { margin-top: 5px; margin-bottom: 5px; line-height: 1.4; }
            a { color: #2980b9; text-decoration: none; }
        </style>
        <div style='font-family: Arial; font-size: 13px;'>
            <div style='text-align:center;'>
                <h2 style='margin-bottom:5px;'>Genizah Search Pro {APP_VERSION}</h2>
                <p style='color: #7f8c8d;'>Developed by Hillel Gershuni (<a href='mailto:gershuni@gmail.com'>gershuni@gmail.com</a>)</p>
            </div>
            <hr>

            <h3>Dedicated to the memory of our beloved teacher, Prof. Menachem Kahana z"l</h3>
            
            <h3>Credits</h3>
            <p>This tool was developed with the coding assistance of <b>Gemini 3.0</b> and <b>GPT 5.1</b>. My thanks to Avi Shmidman, Elisha Rosenzweig, Ephraim Meiri, Elazar Gershuni, Itai Kagan, Elnatan Chen and Adiel Breuer for their advice and support.</p>

            <h3>Data Source & Acknowledgments</h3>
            <p>This software is built on the transcription dataset produced by the <b>MiDRASH Project</b>. I am grateful to the project leaders – Daniel Stoekl Ben Ezra, Marina Rustow, Nachum Dershowitz, Avi Shmidman, and Judith Olszowy-Schlanger – and to Tsafra Siew and Yitzchak Gila from the National Library of Israel. Many thanks also to the rest of the project team: Luigi Bambaci, Benjamin Kiessling, Hayim Lapin, Nurit Ezer, Elena Lolli, Berat Kurar Barakat, Sharva Gogawale, Moshe Lavee, Vered Raziel Kretzmer, and Daria Vasyutinsky Shapira.</p>
            <p>Making such a complex and valuable dataset freely available to the public is a significant step for Open Science, and I deeply appreciate their generosity in allowing everyone to access these texts.</p>
            <h3>License</h3> 
            
            <p>The underlying dataset is licensed under the Creative Commons Attribution 4.0 International (<a href='https://creativecommons.org/licenses/by/4.0/'>CC BY 4.0</a>) license</p>

            <h3>Citation</h3>
            <p>If you use these results in your research, please cite the creators of the dataset: Stoekl Ben Ezra, Daniel, Luigi Bambaci, Benjamin Kiessling, Hayim Lapin, Nurit Ezer, Elena Lolli, Marina Rustow, et al. MiDRASH Automatic Transcriptions. Data set. Zenodo, 2025. <a href='https://doi.org/10.5281/zenodo.17734473'>https://doi.org/10.5281/zenodo.17734473</a>. You can also mention you used this program: Genizah Search Pro by Hillel Gershuni.</p>
        </div>
        """
        
        about_txt = tr("ABOUT_HTML") if CURRENT_LANG == 'he' else about_html_en.replace("{APP_VERSION}", APP_VERSION)
        txt_about = QTextBrowser()
        txt_about.setHtml(about_txt)
        txt_about.setOpenExternalLinks(True)
        abl.addWidget(txt_about)

        # Citation Row
        cit_row = QHBoxLayout()
        cit_row.addWidget(QLabel(tr("Citation:")))

        citation_str = "Stoekl Ben Ezra, Daniel, Luigi Bambaci, Benjamin Kiessling, Hayim Lapin, Nurit Ezer, Elena Lolli, Marina Rustow, et al. MiDRASH Automatic Transcriptions. Data set. Zenodo, 2025. https://doi.org/10.5281/zenodo.17734473."

        self.txt_citation = QLineEdit(citation_str)
        self.txt_citation.setReadOnly(True)
        self.txt_citation.setCursorPosition(0)
        cit_row.addWidget(self.txt_citation)

        btn_copy = QPushButton(tr("Copy"))
        btn_copy.setToolTip(tr("Copy Citation"))
        btn_copy.setFixedSize(60, 24) # Small
        btn_copy.clicked.connect(self.copy_citation)
        cit_row.addWidget(btn_copy)

        abl.addLayout(cit_row)

        gb_about.setLayout(abl); layout.addWidget(gb_about)

        # Data Sources section
        gb_data = QGroupBox(tr("Data Sources"))
        data_layout = QVBoxLayout()

        data_html = "<table style='font-size: 12px; border-collapse: collapse;'>"
        data_html += "<tr><th style='text-align:left; padding: 3px 10px 3px 0;'>Source</th>"
        data_html += "<th style='text-align:left; padding: 3px 10px 3px 0;'>Version</th>"
        data_html += "<th style='text-align:left; padding: 3px 10px 3px 0;'>Status</th></tr>"

        # PGP
        try:
            from shared.document_service import get_pgp_service
            pgp_svc = get_pgp_service()
            pgp_ver = pgp_svc.get_version() if pgp_svc.is_available() else None
        except Exception:
            pgp_ver = None
        pgp_status = f"v{pgp_ver}" if pgp_ver else "Not installed"
        data_html += f"<tr><td style='padding: 3px 10px 3px 0;'>PGP Documents</td><td>{pgp_status}</td>"
        data_html += f"<td>{'&#10003;' if pgp_ver else '&#8212;'}</td></tr>"

        # FJMS
        try:
            from shared.fjms_service import get_fjms_service
            fjms_svc = get_fjms_service()
            fjms_ver = fjms_svc.get_version() if fjms_svc.is_available() else None
        except Exception:
            fjms_ver = None
        fjms_status = f"v{fjms_ver}" if fjms_ver else "Not installed"
        data_html += f"<tr><td style='padding: 3px 10px 3px 0;'>FJMS Catalog</td><td>{fjms_status}</td>"
        data_html += f"<td>{'&#10003;' if fjms_ver else '&#8212;'}</td></tr>"

        # NLI
        try:
            from shared.nli_crossref_service import get_nli_crossref_service
            nli_svc = get_nli_crossref_service()
            nli_ver = nli_svc.get_version() if nli_svc.is_available() else None
        except Exception:
            nli_ver = None
        nli_status = f"v{nli_ver}" if nli_ver else "Not installed"
        data_html += f"<tr><td style='padding: 3px 10px 3px 0;'>NLI Crossref</td><td>{nli_status}</td>"
        data_html += f"<td>{'&#10003;' if nli_ver else '&#8212;'}</td></tr>"

        data_html += "</table>"

        txt_data = QTextBrowser()
        txt_data.setHtml(data_html)
        txt_data.setMaximumHeight(100)
        data_layout.addWidget(txt_data)
        gb_data.setLayout(data_layout)
        layout.addWidget(gb_data)

        panel.setLayout(layout)
        return panel

    def copy_citation(self):
        citation = "Stoekl Ben Ezra, D., Bambaci, L., Kiessling, B., Lapin, H., Ezer, N., Lolli, E., Rustow, M., Dershowitz, N., Kurar Barakat, B., Gogawale, S., Shmidman, A., Lavee, M., Siew, T., Raziel Kretzmer, V., Vasyutinsky Shapira, D., Olszowy-Schlanger, J., & Gila, Y. (2025). MiDRASH Automatic Transcriptions. Zenodo. https://doi.org/10.5281/zenodo.17734473"
        QApplication.clipboard().setText(citation)
        QMessageBox.information(self, tr("Copied"), tr("Citation copied to clipboard!"))

    # --- HELP TEXTS ---
    def open_help_center(self, anchor=None):
        """Open the bundled Help.html with optional anchor scrolling and fallback content."""
        help_path = Config.HELP_FILE
        dlg = HelpDialog(
            self,
            tr("Genizah Help"),
            source_path=help_path,
            anchor=anchor,
            fallback_html=self._build_help_fallback_html(),
            lang=CURRENT_LANG,
        )
        dlg.exec()

    def get_search_help_text(self):
        if CURRENT_LANG == 'he': return tr("SEARCH_HELP_HTML")
        return """<h3>Search Modes</h3><ul><li><b>Exact:</b> Only finds exact matches.</li><li><b>Variants (?):</b> Basic OCR errors.</li><li><b>Extended (??):</b> More variants.</li><li><b>Maximum (???):</b> Aggressive swapping (Use caution).</li><li><b>Fuzzy (~):</b> Levenshtein distance (1-2 typos).</li><li><b>Regex:</b> Advanced patterns.</li><li><b>Title:</b> Search in composition titles (metadata).</li><li><b>Shelfmark:</b> Search for shelfmarks (metadata).</li><li><b>Responsa (R):</b> Search syntax inspired by the Bar-Ilan Responsa Project, with prefix/suffix expansion, wildcards, spelling variants, and proximity gaps. Use the Query Builder for visual construction.</li></ul><hr><b>Gap:</b> Max distance between words (irrelevant for Title/Shelfmark)."""

    def get_comp_help_text(self):
        if CURRENT_LANG == 'he': return tr("COMP_HELP_HTML")
        return """<h3>Composition Search</h3><p>Finds parallels between a source text and the Genizah.</p><ul><li><b>Chunk:</b> Words per search block (5-7 recommended).</li><li><b>Max Freq:</b> Filter out common phrases.</li><li><b>Filter >:</b> Group results if a title appears frequently (move to Appendix).</li></ul>"""

    def get_browse_help_text(self):
        if CURRENT_LANG == 'he': return tr("BROWSE_HELP_HTML")
        return """<h3>Browse Manuscripts</h3><ul><li><b>System ID:</b> Enter an ID to load a manuscript.</li><li><b>View All:</b> Switch to continuous view of the full text.</li><li><b>Save:</b> Export the manuscript text to a file.</li></ul>"""

    def get_settings_help_text(self):
        if CURRENT_LANG == 'he': return tr("SETTINGS_HELP_HTML")
        return """<h3>Settings & Index</h3><ul><li><b>Build/Rebuild Index:</b> Required on first run or after corpus updates.</li><li><b>About:</b> View version, credits, and citation details.</li></ul>"""

    def _build_help_fallback_html(self):
        sections = [
            ("search", tr("Search Help")),
            self.get_search_help_text(),
            ("composition", tr("Composition Help")),
            self.get_comp_help_text(),
            ("browse", tr("Browse Help")),
            self.get_browse_help_text(),
            ("settings", tr("Settings Help")),
            self.get_settings_help_text(),
        ]
        html_parts = ["<div style='font-family: Arial; font-size: 13px;'>"]
        for i in range(0, len(sections), 2):
            anchor, heading = sections[i]
            body = sections[i + 1]
            html_parts.append(f"<h2 id='{anchor}'>{heading}</h2>")
            html_parts.append(body)
            html_parts.append("<hr>")
        html_parts.append("</div>")
        return "".join(html_parts)

    def _sanitize_filename(self, text, fallback):
        clean = re.sub(r"[^\w\u0590-\u05FF\s-]", "", text or "")
        clean = re.sub(r"\s+", "_", clean).strip("_")
        return clean or fallback

    def _get_default_save_folder(self):
        """Get the default folder for saving reports. Checks last used location first."""
        # Check if user has a saved last location
        cfg = load_app_config()
        last_folder = cfg.get('last_save_folder')
        if last_folder and os.path.isdir(last_folder):
            return last_folder

        # Default: My Documents\Genizah Search Pro
        # Use Windows Shell API to get the correct Documents folder
        documents_folder = None
        try:
            import ctypes.wintypes
            CSIDL_PERSONAL = 5  # My Documents
            buf = ctypes.create_unicode_buffer(ctypes.wintypes.MAX_PATH)
            ctypes.windll.shell32.SHGetFolderPathW(None, CSIDL_PERSONAL, None, 0, buf)
            if buf.value:
                documents_folder = buf.value
        except Exception:
            pass

        # Fallback: try common locations
        if not documents_folder or not os.path.isdir(documents_folder):
            for folder_name in ["Documents", "My Documents"]:
                candidate = os.path.join(os.path.expanduser("~"), folder_name)
                if os.path.isdir(candidate):
                    documents_folder = candidate
                    break

        if not documents_folder or not os.path.isdir(documents_folder):
            documents_folder = os.path.expanduser("~")

        default_folder = os.path.join(documents_folder, "GenizahSearchPro", "Reports")
        try:
            os.makedirs(default_folder, exist_ok=True)
        except Exception:
            # Fallback to old reports dir if we can't create the new one
            return Config.REPORTS_DIR

        return default_folder

    def _get_unique_filepath(self, filepath):
        """If file exists, add (1), (2), etc. until we find a unique name."""
        if not os.path.exists(filepath):
            return filepath

        base, ext = os.path.splitext(filepath)
        counter = 1
        while True:
            new_path = f"{base} ({counter}){ext}"
            if not os.path.exists(new_path):
                return new_path
            counter += 1

    def _save_last_folder(self, filepath):
        """Remember the folder where user saved a file."""
        folder = os.path.dirname(filepath)
        if folder and os.path.isdir(folder):
            save_app_config({'last_save_folder': folder})

    def _default_report_path(self, hint, fallback):
        filename = self._sanitize_filename(hint, fallback)

        # Lab Mode: save to Lab Dir
        if getattr(self, 'btn_lab_mode_toggle', None) and self.btn_lab_mode_toggle.isChecked():
            base_dir = Config.REPORTS_DIR
        else:
            # Use smart default folder (last used or My Documents\Genizah Search Pro)
            base_dir = self._get_default_save_folder()

        filepath = os.path.join(base_dir, f"{filename}.txt")
        return self._get_unique_filepath(filepath)

    def _get_credit_header(self):
        english_text = (
            "Generated by Genizah Search Pro\n"
            "Data Source: MiDRASH Automatic Transcriptions (Stoekl Ben Ezra et al., 2025)\n"
            "Dataset available at: https://doi.org/10.5281/zenodo.17734473\n"
            "================================================================================\n"
        )

        final_text = english_text
        if CURRENT_LANG == 'he':
            final_text = tr("REPORT_CREDIT_TXT")
            
        return final_text + "\n"

    def _get_lab_config_block(self):
        if getattr(self, 'btn_lab_mode_toggle', None) and self.btn_lab_mode_toggle.isChecked() and self.lab_engine:
            settings_dump = json.dumps({
                'custom_variants': self.lab_engine.settings.custom_variants,
                'candidate_limit': self.lab_engine.settings.candidate_limit,
                'min_should_match': self.lab_engine.settings.min_should_match,
                'gap_penalty': self.lab_engine.settings.gap_penalty,
                'ignore_matres': self.lab_engine.settings.ignore_matres,
                'phonetic_expansion': self.lab_engine.settings.phonetic_expansion,
            }, indent=2, ensure_ascii=False)
            return f"\n[LAB MODE CONFIGURATION]\n{settings_dump}\n================================================================================\n"
        return ""

    # --- LOGIC ---
    def open_search_settings(self):
        """Open the Search Settings dialog for variant configuration."""
        if not self.lab_engine:
            return
        d = SearchSettingsDialog(self, self.lab_engine.settings)
        if d.exec():
            # Refresh preset/slider visibility based on new setting
            use_slider = getattr(self.lab_engine.settings, 'variant_use_slider', False)
            if hasattr(self, 'variant_presets_widget'):
                self.variant_presets_widget.setVisible(not use_slider)
            if hasattr(self, 'variant_slider_widget'):
                self.variant_slider_widget.setVisible(use_slider)

    def _on_search_mode_changed(self, index):
        """Show/hide variant controls and swap query/tag input based on selected mode."""
        is_exact = (index == 0)
        is_variants = (index == 1)
        is_responsa = (index == self.MODE_RESPONSA)  # 2
        is_pgp_tags = (index == self.MODE_PGP_TAGS)  # 7

        if hasattr(self, 'variant_controls_container'):
            self.variant_controls_container.setVisible(is_variants and not is_pgp_tags)
        if is_variants:
            self._update_variant_count_preview()

        # Show/hide Responsa sub-options based on mode
        if hasattr(self, 'responsa_sub_row'):
            self.responsa_sub_row.setVisible(is_responsa)

        # Dismiss mode combo glow when user selects Responsa
        if is_responsa and self._mode_glow_active:
            self._stop_mode_glow()
            save_app_config({'hint_responsa_seen': True})

        # PGP Tags mode: hide row1, show tag combo, hide search params
        if hasattr(self, 'tag_search_combo'):
            self.search_row1_container.setVisible(not is_pgp_tags)
            self.tag_search_combo.setVisible(is_pgp_tags)
            self.search_params_container.setVisible(not is_pgp_tags)

    def _pulse_mode_glow(self):
        """Toggle glow border on mode combo for feature discovery."""
        if self._mode_glow_on:
            self.mode_combo.setStyleSheet("QComboBox { border: 2px solid #10b981; border-radius: 4px; }")
        else:
            self.mode_combo.setStyleSheet("")
        self._mode_glow_on = not self._mode_glow_on

    def _stop_mode_glow(self):
        """Remove mode combo glow and item highlight."""
        self._mode_glow_active = False
        if hasattr(self, '_mode_glow_timer'):
            self._mode_glow_timer.stop()
        self.mode_combo.setStyleSheet("")
        self.mode_combo.setToolTip(tr("Responsa-Project style grammatical expansion for Hebrew search"))
        # Clear the green background and sparkle from the Responsa item
        from PyQt6.QtGui import QBrush
        self.mode_combo.model().item(self.MODE_RESPONSA).setBackground(QBrush())
        self.mode_combo.setItemText(self.MODE_RESPONSA, tr("Responsa (R)"))

    def _pulse_tabular_glow(self):
        """Toggle glow border on tabular button for feature discovery."""
        if self._tabular_glow_on:
            self.btn_query_builder.setStyleSheet("font-size: 11px; padding: 2px 8px; border: 2px solid #10b981; border-radius: 4px;")
        else:
            self.btn_query_builder.setStyleSheet("font-size: 11px; padding: 2px 8px;")
        self._tabular_glow_on = not self._tabular_glow_on

    def _stop_tabular_glow(self):
        """Remove tabular button glow."""
        self._tabular_glow_active = False
        if hasattr(self, '_tabular_glow_timer'):
            self._tabular_glow_timer.stop()
        self.btn_query_builder.setStyleSheet("font-size: 11px; padding: 2px 8px;")
        self.btn_query_builder.setToolTip(tr("Open the tabular query builder"))

    def _on_comp_mode_changed(self, index):
        """Show/hide variant slider for composition based on selected mode."""
        # Index 1 = Variants mode
        is_variants = (index == 1)
        if hasattr(self, 'comp_variant_slider_container'):
            self.comp_variant_slider_container.setVisible(is_variants)

    def _on_boundary_mode_changed(self, index):
        """Update UI based on boundary mode selection."""
        mode = self.boundary_mode_combo.currentData() if hasattr(self, 'boundary_mode_combo') else 'full'

        # Show/hide advanced button for non-full modes
        if hasattr(self, 'btn_boundary_advanced'):
            self.btn_boundary_advanced.setVisible(mode in ('boundary', 'combined'))

        # Update tooltip based on mode
        tooltips = {
            'full': tr("Search all text chunks regardless of paragraph breaks"),
            'boundary': tr("Show only matches where the matching text spans a paragraph break in your source"),
            'combined': tr("Search everything, but rank cross-paragraph matches higher")
        }
        if hasattr(self, 'boundary_mode_combo'):
            self.boundary_mode_combo.setToolTip(tooltips.get(mode, ''))

        # Update stats
        self._update_boundary_stats()

        # Save to settings
        if hasattr(self, 'lab_engine') and self.lab_engine:
            self.lab_engine.settings.boundary_mode = mode
            self.lab_engine.settings.save()

    def _on_boundary_delimiter_changed(self, index):
        """Save delimiter setting and update stats when delimiter changes."""
        delimiter = self.boundary_delimiter_combo.currentData() if hasattr(self, 'boundary_delimiter_combo') else '\n'

        # Save to settings
        if hasattr(self, 'lab_engine') and self.lab_engine:
            self.lab_engine.settings.boundary_delimiter = delimiter
            self.lab_engine.settings.save()

        # Update stats
        self._update_boundary_stats()

    def _update_boundary_stats(self):
        """Update the boundary statistics label based on current text and settings."""
        if not hasattr(self, 'boundary_stats_label') or not hasattr(self, 'comp_text_area'):
            return

        text = self.comp_text_area.toPlainText().strip()
        if not text:
            self.boundary_stats_label.setText("")
            return

        try:
            from genizah_core import get_boundary_stats

            chunk_size = self.spin_chunk.value() if hasattr(self, 'spin_chunk') else 5
            delimiter = self.boundary_delimiter_combo.currentData() if hasattr(self, 'boundary_delimiter_combo') else '\n'
            if hasattr(self, 'lab_engine') and self.lab_engine:
                min_distance = self.lab_engine.settings.min_delimiter_distance
            else:
                min_distance = getattr(self, '_min_delimiter_distance_temp', 3)

            stats = get_boundary_stats(text, delimiter, chunk_size, min_distance)

            if stats['boundary_count'] > 0:
                self.boundary_stats_label.setText(
                    tr("{} boundaries detected, {} chunks will cross them").format(
                        stats['boundary_count'], stats['crossing_chunk_count']
                    )
                )
                self.boundary_stats_label.setStyleSheet("color: #5dade2; font-size: 11px;")
            else:
                mode = self.boundary_mode_combo.currentData() if hasattr(self, 'boundary_mode_combo') else 'full'
                if mode in ('boundary', 'combined'):
                    self.boundary_stats_label.setText(tr("No paragraph breaks detected in text!"))
                    self.boundary_stats_label.setStyleSheet("color: #e74c3c; font-size: 11px;")
                else:
                    self.boundary_stats_label.setText("")
        except Exception as e:
            traceback.print_exc()
            self.boundary_stats_label.setText("")

    def _open_boundary_advanced_dialog(self):
        """Open dialog for advanced boundary search settings."""
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QFormLayout, QDoubleSpinBox, QSpinBox, QDialogButtonBox

        dialog = QDialog(self)
        dialog.setWindowTitle(tr("Advanced cross-paragraph settings"))
        dialog.setMinimumWidth(350)

        layout = QVBoxLayout(dialog)
        form = QFormLayout()

        # Get current settings
        settings = self.lab_engine.settings if hasattr(self, 'lab_engine') and self.lab_engine else None

        # Cross-paragraph boost slider
        boost_spin = QDoubleSpinBox()
        boost_spin.setRange(1.0, 3.0)
        boost_spin.setSingleStep(0.1)
        boost_spin.setValue(settings.boundary_boost if settings else getattr(self, '_boundary_boost_temp', 1.5))
        boost_spin.setToolTip(tr("Score multiplier for cross-paragraph matches"))
        form.addRow(tr("Cross-paragraph boost") + ":", boost_spin)

        # Min boundary matches
        min_matches_spin = QSpinBox()
        min_matches_spin.setRange(0, 10)
        min_matches_spin.setValue(settings.min_boundary_matches if settings else getattr(self, '_min_boundary_matches_temp', 0))
        min_matches_spin.setToolTip(tr("Minimum number of cross-paragraph matches required"))
        form.addRow(tr("Min. cross-paragraph matches") + ":", min_matches_spin)

        # Min delimiter distance
        min_distance_spin = QSpinBox()
        min_distance_spin.setRange(1, 10)
        min_distance_spin.setValue(settings.min_delimiter_distance if settings else getattr(self, '_min_delimiter_distance_temp', 3))
        min_distance_spin.setToolTip(tr("Ignore separators that are too close together"))
        form.addRow(tr("Min. words between separators") + ":", min_distance_spin)

        layout.addLayout(form)

        # Buttons
        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        layout.addWidget(buttons)

        if dialog.exec() == QDialog.DialogCode.Accepted:
            if settings:
                settings.boundary_boost = boost_spin.value()
                settings.min_boundary_matches = min_matches_spin.value()
                settings.min_delimiter_distance = min_distance_spin.value()
                settings.save()
            else:
                # Store values temporarily on self for this session (won't persist)
                self._boundary_boost_temp = boost_spin.value()
                self._min_boundary_matches_temp = min_matches_spin.value()
                self._min_delimiter_distance_temp = min_distance_spin.value()
            self._update_boundary_stats()

    def _set_variant_preset(self, pairs_count):
        """Set variant level from preset button."""
        self._current_variant_preset = pairs_count

        # Update button states
        if hasattr(self, 'btn_variant_basic'):
            self.btn_variant_basic.setChecked(pairs_count == 30)
            self.btn_variant_extended.setChecked(pairs_count == 70)
            self.btn_variant_maximum.setChecked(pairs_count == 150)

        # Update variant manager
        if hasattr(self, 'var_mgr') and self.var_mgr:
            self.var_mgr.set_variant_level(pairs_count)

        # Sync slider if visible
        if hasattr(self, 'variant_slider'):
            self.variant_slider.blockSignals(True)
            self.variant_slider.setValue(pairs_count)
            self.variant_slider_label.setText(str(pairs_count))
            self.variant_slider.blockSignals(False)

        # Update preview
        self._update_variant_count_preview()

    def _get_current_variant_pairs_count(self):
        """Get the current variant pairs count (from preset or slider)."""
        use_slider = getattr(self.lab_engine.settings if hasattr(self, 'lab_engine') and self.lab_engine else None, 'variant_use_slider', False) if hasattr(self, 'lab_engine') else False
        if use_slider and hasattr(self, 'variant_slider'):
            return self.variant_slider.value()
        elif hasattr(self, '_current_variant_preset'):
            return self._current_variant_preset
        return 70  # Default

    def _sync_variant_sliders(self, value, source='search'):
        """Keep variant sliders synchronized between search and composition tabs."""
        if source == 'search' and hasattr(self, 'comp_variant_slider'):
            self.comp_variant_slider.blockSignals(True)
            self.comp_variant_slider.setValue(value)
            self.comp_variant_slider_label.setText(str(value))
            self.comp_variant_slider.blockSignals(False)
        elif source == 'comp' and hasattr(self, 'variant_slider'):
            self.variant_slider.blockSignals(True)
            self.variant_slider.setValue(value)
            self.variant_slider_label.setText(str(value))
            self.variant_slider.blockSignals(False)

    # Shortcut prefixes: longest first to avoid partial matches (??? before ??)
    _SHORTCUT_PREFIXES = [
        ('???', 'variants_maximum'),
        ('??', 'variants_extended'),
        ('R', 'responsa'),
        ('?', 'variants'),
        ('=', 'literal'),
        ('~', 'fuzzy'),
        ('/', 'Regex'),
        ('$', 'Title'),
        ('#', 'Shelfmark'),
    ]

    def _on_query_text_changed(self):
        """Handle live text changes: detect shortcut prefixes and update variant preview."""
        text = self.query_input.text()
        # Check for shortcut prefix followed by space
        for prefix, target_mode in self._SHORTCUT_PREFIXES:
            if text.startswith(prefix + ' ') or text == prefix + ' ':
                clean = text[len(prefix):].lstrip()
                # Block signals to avoid recursive textChanged
                self.query_input.blockSignals(True)
                self.query_input.setText(clean)
                self.query_input.blockSignals(False)
                # Switch mode combo
                modes = ['literal', 'variants', 'responsa', 'fuzzy', 'Regex', 'Title', 'Shelfmark']
                if target_mode in ('variants_extended', 'variants_maximum'):
                    target_mode = 'variants'
                try:
                    combo_idx = modes.index(target_mode)
                    self.mode_combo.setCurrentIndex(combo_idx)
                except ValueError:
                    pass
                break
        self._update_variant_count_preview()

    def _update_variant_count_preview(self):
        """Update the variant count label based on current query and slider value."""
        if not hasattr(self, 'variant_count_label') or not hasattr(self, 'query_input'):
            return
        if not hasattr(self, 'var_mgr') or not self.var_mgr:
            return

        query = self.query_input.text().strip()
        if not query:
            self.variant_count_label.setText("")
            return

        # Strip search prefixes
        for prefix in ['?', '??', '???', '=', '~', '/', '#', '$']:
            if query.startswith(prefix):
                query = query[len(prefix):].strip()
                break

        words = query.split()
        if not words:
            self.variant_count_label.setText("")
            return

        try:
            # Set variant level from current UI (preset or slider)
            pairs_count = self._get_current_variant_pairs_count()
            self.var_mgr.set_variant_level(pairs_count)

            # Calculate total variants for all words
            total_variants = 0
            for word in words:
                if len(word) >= 2:
                    variants = self.var_mgr.get_variants(word, 'variants', limit=500)
                    total_variants += len(variants)
                else:
                    total_variants += 1  # Single char = 1 variant (itself)

            self.variant_count_label.setText(f"≈{total_variants}")
        except Exception:
            self.variant_count_label.setText("")

    def update_lab_ui_state(self, checked):
        """Disable standard controls when Lab Mode is active."""
        # Search Tab
        if hasattr(self, 'mode_combo'): self.mode_combo.setEnabled(not checked)
        if hasattr(self, 'gap_input'): self.gap_input.setEnabled(not checked)
        if hasattr(self, 'chk_lab_deep'): self.chk_lab_deep.setEnabled(checked)

        # Composition Tab
        if hasattr(self, 'comp_mode_combo'): self.comp_mode_combo.setEnabled(not checked)
        if hasattr(self, 'spin_freq'): self.spin_freq.setEnabled(not checked)
        if hasattr(self, 'chk_lab_deep_comp'): self.chk_lab_deep_comp.setEnabled(checked)

    def on_deep_scan_toggled_search(self, checked):
        if hasattr(self, 'chk_lab_deep_comp'):
            self.chk_lab_deep_comp.blockSignals(True)
            self.chk_lab_deep_comp.setChecked(checked)
            self.chk_lab_deep_comp.blockSignals(False)

    def on_deep_scan_toggled_comp(self, checked):
        if hasattr(self, 'chk_lab_deep'):
            self.chk_lab_deep.blockSignals(True)
            self.chk_lab_deep.setChecked(checked)
            self.chk_lab_deep.blockSignals(False)

    def on_lab_mode_toggled_search(self, checked):
        # Show/Hide Panel
        if hasattr(self, 'lab_panel_search'):
            self.lab_panel_search.setVisible(checked)

        # Sync Comp Button
        if hasattr(self, 'btn_lab_mode_toggle_comp'):
            self.btn_lab_mode_toggle_comp.blockSignals(True)
            self.btn_lab_mode_toggle_comp.setChecked(checked)
            self.btn_lab_mode_toggle_comp.blockSignals(False)
            # Ensure comp panel visibility matches too
            if hasattr(self, 'lab_panel_comp'):
                self.lab_panel_comp.setVisible(checked)

        self.update_lab_ui_state(checked)

    def on_lab_mode_toggled_comp(self, checked):
        # Show/Hide Panel
        if hasattr(self, 'lab_panel_comp'):
            self.lab_panel_comp.setVisible(checked)

        # Sync Search Button
        if hasattr(self, 'btn_lab_mode_toggle'):
            self.btn_lab_mode_toggle.blockSignals(True)
            self.btn_lab_mode_toggle.setChecked(checked)
            self.btn_lab_mode_toggle.blockSignals(False)
            # Ensure search panel visibility matches too
            if hasattr(self, 'lab_panel_search'):
                self.lab_panel_search.setVisible(checked)

        self.update_lab_ui_state(checked)

    def _open_domain_filter_dialog(self):
        """Open the domain filter dialog for post-search dynamic filtering."""
        if not self._result_domain_counts:
            return

        # Count results with no domain data
        uncategorized_count = sum(
            1 for sid in self.result_row_by_sys_id.keys()
            if sid not in self._result_domain_map or not self._result_domain_map[sid]
        )

        dlg = DomainFilterDialog(
            self, result_domains=self._result_domain_counts,
            excluded_domains=self._domain_exclusions.copy(),
            uncategorized_count=uncategorized_count,
        )
        if dlg.exec() == QDialog.DialogCode.Accepted:
            self._domain_exclusions = dlg.get_excluded_domains()
            self._update_domain_filter_label()
            self._apply_domain_exclusions()

    def _update_domain_filter_label(self):
        """Update the domain filter label badge to show exclusion state."""
        if self._domain_exclusions:
            count = len(self._domain_exclusions)
            if count == 1:
                name = self._domain_display_name(next(iter(self._domain_exclusions)))
                self.lbl_domain_filter.setText(f"[-{name}]")
            else:
                self.lbl_domain_filter.setText(f"[{count} excluded]")
            self.lbl_domain_filter.setStyleSheet("color: #e74c3c; font-size: 11px;")  # Red for exclusion
            self.lbl_domain_filter.setVisible(True)
        else:
            self.lbl_domain_filter.setVisible(False)
            self.lbl_domain_filter.setStyleSheet("color: #9b59b6; font-size: 11px;")

    def _domain_display_name(self, en_name):
        """Get display name for a domain (Hebrew if UI is Hebrew, else English)."""
        if CURRENT_LANG == 'he':
            if hasattr(self, '_domain_name_map') and en_name in self._domain_name_map:
                return self._domain_name_map[en_name]
            translated = tr(en_name)
            if translated != en_name:
                return translated
        return en_name

    def _apply_domain_exclusions(self):
        """Apply domain exclusions by hiding/showing table rows."""
        hide_uncategorized = "Uncategorized" in self._domain_exclusions

        if not self._domain_exclusions:
            # No exclusions -- show all rows
            for row in range(self.results_table.rowCount()):
                self.results_table.setRowHidden(row, False)
            visible = self.results_table.rowCount()
        else:
            visible = 0
            for row in range(self.results_table.rowCount()):
                # Read sys_id directly from table cell (survives sorting)
                item = self.results_table.item(row, self.COL_SYS_ID)
                sys_id = item.text().strip() if item else None
                # Get domains for this result
                result_domains = self._result_domain_map.get(sys_id, []) if sys_id else []
                if not result_domains:
                    # No domain data -- hide if Uncategorized is excluded
                    self.results_table.setRowHidden(row, hide_uncategorized)
                    if not hide_uncategorized:
                        visible += 1
                elif all(d in self._domain_exclusions for d in result_domains):
                    # ALL domains excluded -- hide
                    self.results_table.setRowHidden(row, True)
                else:
                    # At least one domain not excluded -- show
                    self.results_table.setRowHidden(row, False)
                    visible += 1
        total = len(self.last_results) if self.last_results else 0
        if self._domain_exclusions:
            self.status_label.setText(
                tr("Showing {} of {} results (filtering {} domains)").format(visible, total, len(self._domain_exclusions))
            )
        else:
            self.status_label.setText(
                tr("Showing {} of {} results").format(
                    min(self.results_loaded, total), total
                )
            )

    def _on_domain_enrichment_loaded(self, raw_domains):
        """Handle async domain enrichment results from DomainEnrichmentWorker.

        Processes raw domain data into display-ready structures and updates
        already-loaded result rows with domain badges.
        """
        from shared.fjms_service import qualify_domain_name

        self._result_domain_map = {}
        self._domain_name_map = getattr(self, '_domain_name_map', {})
        domain_counts = {}

        for sys_id, doms in raw_domains.items():
            child_names = {d['domain'] for d in doms}
            filtered = [
                qualify_domain_name(d['domain'], d.get('parent_domain'))
                for d in doms
                if not (d.get('parent_domain') and d['parent_domain'] in child_names and d['parent_domain'] != d['domain'])
            ]
            if filtered:
                self._result_domain_map[sys_id] = filtered
                for d in filtered:
                    domain_counts[d] = domain_counts.get(d, 0) + 1
            for d in doms:
                qname = qualify_domain_name(d['domain'], d.get('parent_domain'))
                if qname != d['domain'] and d.get('domain_heb') and d.get('parent_domain_heb'):
                    self._domain_name_map[qname] = f"{d['domain_heb']} ({d['parent_domain_heb']})"
                if d.get('domain_heb') and d['domain'] not in self._domain_name_map:
                    self._domain_name_map[d['domain']] = d['domain_heb']
                if d.get('parent_domain_heb') and d.get('parent_domain') and d['parent_domain'] not in self._domain_name_map:
                    self._domain_name_map[d['parent_domain']] = d['parent_domain_heb']

        self._result_domain_counts = domain_counts
        self._has_result_domains = bool(domain_counts)
        self.btn_domain_filter.setEnabled(self._has_result_domains)

        # Update domain column cells for already-loaded rows
        for row in range(self.results_table.rowCount()):
            item = self.results_table.item(row, self.COL_SYS_ID)
            if not item:
                continue
            sid = item.text().strip()
            domain_names = self._result_domain_map.get(sid, [])
            domain_text = ", ".join(self._domain_display_name(d) for d in domain_names) if domain_names else ""
            domain_item = QTableWidgetItem(domain_text)
            if domain_names:
                domain_item.setForeground(QColor("#8e44ad"))
            self.results_table.setItem(row, self.COL_DOMAIN, domain_item)

        # Apply any remembered domain exclusions after enrichment
        if self._domain_exclusions and self._has_result_domains:
            self._apply_domain_exclusions()

    def _navigate_to_search_with_domain(self, domain_name):
        """Navigate to search tab with domain context (exclusions cleared)."""
        self._domain_exclusions = set()  # Clear exclusions when navigating from browse
        self._update_domain_filter_label()
        self.tabs.setCurrentWidget(self.search_tab)
        # Note: domain will appear in post-search filter after user runs a search

    # --- Composition Domain Filter ---

    def _collect_comp_domain_data(self, main, appx, filtered, filt_appx):
        """Collect domain data for composition results."""
        all_sys_ids = set()
        for item in main:
            sid = item.get('sys_id')
            if not sid:
                sid, _ = self.meta_mgr.parse_header_smart(item.get('raw_header', ''))
            if sid:
                all_sys_ids.add(sid)
        for items in appx.values():
            for item in items:
                sid = item.get('sys_id')
                if not sid:
                    sid, _ = self.meta_mgr.parse_header_smart(item.get('raw_header', ''))
                if sid:
                    all_sys_ids.add(sid)
        for item in filtered:
            sid = item.get('sys_id')
            if not sid:
                sid, _ = self.meta_mgr.parse_header_smart(item.get('raw_header', ''))
            if sid:
                all_sys_ids.add(sid)
        for items in filt_appx.values():
            for item in items:
                sid = item.get('sys_id')
                if not sid:
                    sid, _ = self.meta_mgr.parse_header_smart(item.get('raw_header', ''))
                if sid:
                    all_sys_ids.add(sid)

        if not all_sys_ids:
            self._comp_result_domain_map = {}
            self._comp_result_domain_counts = {}
            self._comp_has_result_domains = False
            self.btn_comp_domain_filter.setEnabled(False)
            return

        from shared.fjms_service import get_fjms_service
        fjms = get_fjms_service()
        if not fjms.is_available():
            self.btn_comp_domain_filter.setEnabled(False)
            return

        raw_domains = fjms.get_domains_for_sys_ids(list(all_sys_ids))
        self._comp_result_domain_map = {}
        if not hasattr(self, '_domain_name_map'):
            self._domain_name_map = {}
        from shared.fjms_service import qualify_domain_name
        for sys_id, doms in raw_domains.items():
            child_names = {d['domain'] for d in doms}
            filtered_doms = [qualify_domain_name(d['domain'], d.get('parent_domain')) for d in doms if not (d.get('parent_domain') and d['parent_domain'] in child_names and d['parent_domain'] != d['domain'])]
            if filtered_doms:
                self._comp_result_domain_map[sys_id] = filtered_doms
            for d in doms:
                qname = qualify_domain_name(d['domain'], d.get('parent_domain'))
                if qname != d['domain'] and d.get('domain_heb') and d.get('parent_domain_heb'):
                    self._domain_name_map[qname] = f"{d['domain_heb']} ({d['parent_domain_heb']})"
                if d.get('domain_heb') and d['domain'] not in self._domain_name_map:
                    self._domain_name_map[d['domain']] = d['domain_heb']
                if d.get('parent_domain_heb') and d.get('parent_domain') and d['parent_domain'] not in self._domain_name_map:
                    self._domain_name_map[d['parent_domain']] = d['parent_domain_heb']

        self._comp_result_domain_counts = {}
        for sys_id, domain_names in self._comp_result_domain_map.items():
            for d in domain_names:
                self._comp_result_domain_counts[d] = self._comp_result_domain_counts.get(d, 0) + 1

        self._comp_has_result_domains = bool(self._comp_result_domain_counts)
        self.btn_comp_domain_filter.setEnabled(self._comp_has_result_domains)

    def _open_comp_domain_filter_dialog(self):
        """Open the domain filter dialog for composition results."""
        if not self._comp_result_domain_counts:
            return

        # Count uncategorized items
        all_comp_sids = set()
        for item in getattr(self, 'comp_main', []):
            sid = item.get('sys_id')
            if not sid:
                sid, _ = self.meta_mgr.parse_header_smart(item.get('raw_header', ''))
            if sid:
                all_comp_sids.add(sid)
        uncategorized_count = sum(
            1 for sid in all_comp_sids
            if sid not in self._comp_result_domain_map or not self._comp_result_domain_map[sid]
        )

        dlg = DomainFilterDialog(
            self, result_domains=self._comp_result_domain_counts,
            excluded_domains=self._comp_domain_exclusions.copy(),
            uncategorized_count=uncategorized_count,
        )
        if dlg.exec() == QDialog.DialogCode.Accepted:
            self._comp_domain_exclusions = dlg.get_excluded_domains()
            self._update_comp_domain_filter_label()
            self._apply_comp_domain_exclusions()

    def _update_comp_domain_filter_label(self):
        """Update the composition domain filter label."""
        if self._comp_domain_exclusions:
            count = len(self._comp_domain_exclusions)
            if count == 1:
                name = self._domain_display_name(next(iter(self._comp_domain_exclusions)))
                self.lbl_comp_domain_filter.setText(f"[-{name}]")
            else:
                self.lbl_comp_domain_filter.setText(f"[{count} {tr('excluded')}]")
            self.lbl_comp_domain_filter.setStyleSheet("color: #e74c3c; font-size: 11px;")
            self.lbl_comp_domain_filter.setVisible(True)
        else:
            self.lbl_comp_domain_filter.setVisible(False)
            self.lbl_comp_domain_filter.setStyleSheet("color: #9b59b6; font-size: 11px;")

    def _apply_comp_domain_exclusions(self):
        """Apply domain exclusions by hiding/showing composition tree items."""
        if not self._comp_domain_exclusions:
            # Show all items
            root = self.comp_tree.invisibleRootItem()
            for i in range(root.childCount()):
                section = root.child(i)
                for j in range(section.childCount()):
                    section.child(j).setHidden(False)
            return

        hide_uncategorized = "Uncategorized" in self._comp_domain_exclusions
        root = self.comp_tree.invisibleRootItem()
        for i in range(root.childCount()):
            section = root.child(i)
            for j in range(section.childCount()):
                node = section.child(j)
                item_data = node.data(0, Qt.ItemDataRole.UserRole)
                if not item_data or not isinstance(item_data, dict):
                    continue
                sid = item_data.get('sys_id')
                if not sid:
                    sid, _ = self.meta_mgr.parse_header_smart(item_data.get('raw_header', ''))
                result_domains = self._comp_result_domain_map.get(sid, []) if sid else []
                if not result_domains:
                    node.setHidden(hide_uncategorized)
                elif all(d in self._comp_domain_exclusions for d in result_domains):
                    node.setHidden(True)
                else:
                    node.setHidden(False)

    def _open_query_builder(self):
        """Open the tabular query builder dialog."""
        # Dismiss tabular button glow on first use
        if self._tabular_glow_active:
            self._stop_tabular_glow()
            save_app_config({'hint_tabular_seen': True})
        dlg = TabularQueryBuilderDialog(self)
        # Sync search options into dialog from outer checkboxes
        dlg.chk_opt_variants.setChecked(self.chk_responsa_variants.isChecked())
        dlg.chk_opt_ja.setChecked(self.chk_responsa_ja.isChecked())
        dlg.chk_opt_flex.setChecked(self.chk_responsa_flex.isChecked())
        dlg.chk_opt_bidir.setChecked(self.chk_bidirectional.isChecked())
        if dlg.exec() == QDialog.DialogCode.Accepted:
            syntax = dlg.get_syntax()
            negated = dlg.get_negated_words()
            # Sync search options back to outer checkboxes
            self.chk_responsa_variants.setChecked(dlg.chk_opt_variants.isChecked())
            self.chk_responsa_ja.setChecked(dlg.chk_opt_ja.isChecked())
            self.chk_responsa_flex.setChecked(dlg.chk_opt_flex.isChecked())
            self.chk_bidirectional.setChecked(dlg.chk_opt_bidir.isChecked())
            if syntax.strip():
                # One-way sync: builder -> text field
                self.query_input.setText(syntax)
                # Add negated words to exclude input (append to existing)
                if negated:
                    existing_exclude = self.exclude_input.text().strip()
                    exclude_parts = existing_exclude.split() if existing_exclude else []
                    for w in negated:
                        if w not in exclude_parts:
                            exclude_parts.append(w)
                    self.exclude_input.setText(' '.join(exclude_parts))
                # Auto-trigger search
                self.start_search()

    def toggle_search(self):
        # PGP Tags mode — execute tag search instead of text search
        if self.mode_combo.currentIndex() == self.MODE_PGP_TAGS:
            self._execute_tag_search()
            return
        if not self.searcher: return
        if self.is_searching: self.stop_search()
        else: self.start_search()

    def start_search(self):
        query = self.query_input.text().strip()
        if not query: return

        # Detect query prefix (?, ??, ???, ~, /) - Delegated to Core
        # Skip prefix parsing in Responsa mode -- # is Responsa syntax, not Shelfmark
        is_responsa = (self.mode_combo.currentIndex() == self.MODE_RESPONSA)
        mode_override, clean_query = self.searcher.parse_query_syntax(query, responsa_mode=is_responsa)

        if mode_override:
            # Map mode string back to combo index
            # modes list must match the order in init_ui (Responsa at index 2)
            modes = ['literal', 'variants', 'responsa', 'fuzzy', 'Regex', 'Title', 'Shelfmark']
            # Handle 'exact' vs 'literal' naming difference if present
            core_mode = mode_override
            if core_mode == 'exact': core_mode = 'literal'
            # Map old extended/maximum to variants (slider controls intensity)
            if core_mode in ('variants_extended', 'variants_maximum'):
                core_mode = 'variants'

            try:
                combo_idx = modes.index(core_mode)
                self.mode_combo.setCurrentIndex(combo_idx)
                query = clean_query
            except ValueError:
                pass # Mode not found in UI list

        mode_idx = self.mode_combo.currentIndex()
        # Map combo index to mode string (Responsa uses 'exact' as base mode)
        if mode_idx == self.MODE_RESPONSA:
            mode = 'exact'  # Base mode; Responsa pipeline takes over via responsa_options
        else:
            modes = ['literal', 'variants', None, 'fuzzy', 'Regex', 'Title', 'Shelfmark']
            mode = modes[mode_idx] if mode_idx < len(modes) else 'literal'

        # Update variant level and max changes from UI before search
        if mode == 'variants' and self.var_mgr:
            pairs_count = self._get_current_variant_pairs_count()
            self.var_mgr.set_variant_level(pairs_count)
            # Update max_changes in settings
            if self.lab_engine and hasattr(self, 'spin_max_changes'):
                self.lab_engine.settings.variant_max_changes = self.spin_max_changes.value()
        gap = int(self.gap_input.text()) if self.gap_input.text().isdigit() else 0

        # Get Excluded Words
        exclude_text = self.exclude_input.text().strip()
        exclude_words = exclude_text.split() if exclude_text else []

        self.last_search_query = query

        self.is_searching = True; self.btn_search.setText(tr("Stop")); self.btn_search.setStyleSheet("background-color: #c0392b; color: white;")
        self.search_progress.setRange(0, 100); self.search_progress.setValue(0); self.search_progress.setVisible(True)

        # Stop any previous metadata loading to prevent race conditions
        if self.meta_loader and self.meta_loader.isRunning():
            self.meta_loader.request_cancel()
            self.meta_loader.wait()

        # Clear item references BEFORE clearing the table to avoid accessing deleted items
        self.shelfmark_items_by_sid = {}
        self.title_items_by_sid = {}

        self.results_table.setRowCount(0) 
        for b in self.export_buttons: b.setEnabled(False)
        self.result_row_by_sys_id = {}
        self.hovered_row = -1

        # Build Responsa options if Responsa mode is selected in combo
        responsa_options = None
        if self.mode_combo.currentIndex() == self.MODE_RESPONSA:
            responsa_options = {
                'responsa_mode': True,
                'variants': self.chk_responsa_variants.isChecked(),
                'ja': self.chk_responsa_ja.isChecked(),
                'flex_spacing': self.chk_responsa_flex.isChecked(),
                'bidirectional': self.chk_bidirectional.isChecked(),
                'variant_mode': 'variants' if self.chk_responsa_variants.isChecked() else 'exact',
            }

        if self.btn_lab_mode_toggle.isChecked():
            if not self.lab_engine:
                QMessageBox.warning(self, tr("Error"), tr("Lab Engine not initialized."))
                self.reset_ui()
                return

            deep = self.chk_lab_deep.isChecked()
            limit = self.lab_engine.settings.lab_scan_limit

            self.search_thread = LabSearchThread(self.lab_engine, query, mode, gap, deep_scan=deep, scan_limit=limit)
        else:
            self.search_thread = SearchThread(self.searcher, query, mode, gap, exclude_words=exclude_words, responsa_options=responsa_options)

        self.search_thread.results_signal.connect(self.on_search_finished)
        self.search_thread.progress_signal.connect(lambda c, t: (self.search_progress.setMaximum(t), self.search_progress.setValue(c)))

        if hasattr(self.search_thread, 'status_signal'):
             self.search_thread.status_signal.connect(self.status_label.setText)

        self.search_thread.error_signal.connect(self.on_error)
        self.search_thread.start()

    def stop_search(self):
        if self.search_thread.isRunning(): self.search_thread.terminate(); self.search_thread.wait()
        self.reset_ui()

    def reset_ui(self):
        self.is_searching = False; self.btn_search.setText(tr("Search")); self.btn_search.setStyleSheet("background-color: #27ae60; color: white;")
        self.search_progress.setVisible(False)

    def on_error(self, err): self.reset_ui(); QMessageBox.critical(self, tr("Error"), str(err))

    def render_asterisks_to_html(self, text):
        if not text: return ""
        t = SearchEngine.format_snippet(text, style='html_inline')
        return f"<div dir='rtl'>{t}</div>"

    def check_scroll_load(self, value):
        bar = self.results_table.verticalScrollBar()
        if bar.maximum() > 0 and value >= bar.maximum() * 0.95:
            self.load_next_batch()

    def load_next_batch(self):
        if self.results_loaded >= len(self.last_results):
            return

        # Determine range
        start_idx = self.results_loaded
        end_idx = min(start_idx + BATCH_SIZE, len(self.last_results))
        batch = self.last_results[start_idx:end_idx]

        if not batch: return

        self.results_table.setSortingEnabled(False)
        current_row = self.results_table.rowCount()
        self.results_table.setRowCount(current_row + len(batch))

        ids_to_fetch = []

        for i, res in enumerate(batch):
            row_idx = current_row + i
            meta = res['display']
            parsed = self.meta_mgr.parse_full_id_components(res.get('raw_header', ''))
            sid = parsed['sys_id'] or meta.get('id')

            # Metadata Check
            shelf, title = self.meta_mgr.get_meta_for_id(sid)
            library_code = self.meta_mgr.get_library_for_id(sid)
            needs_fetch = (shelf == "Unknown" and (not title))
            if needs_fetch: ids_to_fetch.append(sid)

            # Checkbox
            item_chk = QTableWidgetItem()
            item_chk.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled | Qt.ItemFlag.ItemIsSelectable)
            item_chk.setCheckState(Qt.CheckState.Unchecked)
            item_chk.setData(Qt.ItemDataRole.UserRole, res)
            self.results_table.setItem(row_idx, self.COL_CHECKBOX, item_chk)

            # Actions
            # Align star to the edge (Right in English, Left in Hebrew)
            is_hebrew = (CURRENT_LANG == 'he')

            # Force LTR layout to control geometric expansion manually
            align = Qt.AlignmentFlag.AlignLeft if is_hebrew else Qt.AlignmentFlag.AlignRight
            actions_widget = ActionsHoverWidget(alignment=align)
            actions_widget.setLayoutDirection(Qt.LayoutDirection.LeftToRight)

            list_btn = self._create_action_button("☆", tr("Add to List"), parent=self.results_table)
            list_btn.clicked.connect(lambda _, r=res, b=list_btn: self.search_add_row_to_list(r, b))
            list_btn.setProperty("action_role", "list_star")

            browse_btn = self._create_action_button("📖", tr("Browse manuscript"), lambda _, r=res: self.open_result_in_browse_from_table(r))
            view_btn = self._create_action_button("👁", tr("View result"), lambda _, r=res: self.show_full_text_for_result(r))

            if is_hebrew:
                # Hebrew: Star on Left (added first), then others expand to Right
                actions_widget.add_btn(list_btn, always_visible=True)
                actions_widget.add_btn(view_btn)
                actions_widget.add_btn(browse_btn)
            else:
                # English: Star on Right (added last), others expand to Left
                actions_widget.add_btn(browse_btn)
                actions_widget.add_btn(view_btn)
                actions_widget.add_btn(list_btn, always_visible=True)

            self.results_table.setCellWidget(row_idx, self.COL_ACTIONS, actions_widget)

            # System ID
            item_sid = QTableWidgetItem(sid)
            item_sid.setData(Qt.ItemDataRole.UserRole, res)
            self.results_table.setItem(row_idx, self.COL_SYS_ID, item_sid)

            # Shelf/Title/Library
            if needs_fetch:
                item_shelf = ShelfmarkTableWidgetItem(tr("Loading..."))
                item_title = QTableWidgetItem(tr("Loading..."))
            else:
                item_shelf = ShelfmarkTableWidgetItem(shelf if shelf else tr("Unknown"))
                item_title = QTableWidgetItem(title if title else "")

            self.results_table.setItem(row_idx, self.COL_SHELF, item_shelf)

            # Library column with tooltip for full name
            item_library = QTableWidgetItem(library_code if library_code else "")
            if library_code:
                full_library_name = get_library_display(library_code, short=False)
                item_library.setToolTip(full_library_name)
            self.results_table.setItem(row_idx, self.COL_LIBRARY, item_library)

            self.results_table.setItem(row_idx, self.COL_TITLE, item_title)

            # Map for updates
            self.shelfmark_items_by_sid[sid] = item_shelf
            self.title_items_by_sid[sid] = item_title
            self.result_row_by_sys_id[sid] = row_idx

            # Snippet
            html_snippet = self.render_asterisks_to_html(res.get('snippet', ''))
            lbl = QLabel(html_snippet)
            lbl.setProperty("filter_text", res.get('snippet', ''))
            lbl.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
            self.results_table.setCellWidget(row_idx, self.COL_SNIPPET, lbl)

            # Img
            self.results_table.setItem(row_idx, self.COL_IMG, QTableWidgetItem(str(meta.get('img', ''))))
            # Src
            self.results_table.setItem(row_idx, self.COL_SRC, QTableWidgetItem(str(meta.get('source', ''))))
            # PGP badge
            if sid and sid in self._pgp_transcription_sys_ids:
                pgp_item = QTableWidgetItem("PGP")
                pgp_item.setForeground(QColor("#27ae60"))
                self.results_table.setItem(row_idx, self.COL_PGP, pgp_item)
            else:
                self.results_table.setItem(row_idx, self.COL_PGP, QTableWidgetItem(""))

            # Domain column
            domain_names = self._result_domain_map.get(sid, [])
            domain_text = ", ".join(self._domain_display_name(d) for d in domain_names) if domain_names else ""
            domain_item = QTableWidgetItem(domain_text)
            if domain_names:
                domain_item.setForeground(QColor("#8e44ad"))
            self.results_table.setItem(row_idx, self.COL_DOMAIN, domain_item)

            self._update_search_row_list_indicator(row_idx, res)

        self.results_loaded = end_idx
        self.results_table.setSortingEnabled(True)
        self._apply_results_table_filters()

        # Update Status (include expanded term count for Responsa searches)
        expanded_count = getattr(self, '_responsa_expanded_count', 0)
        if expanded_count > 0:
            self.status_label.setText(tr("Showing {} of {} results (searching {} expanded terms)").format(
                self.results_loaded, len(self.last_results), expanded_count
            ))
        else:
            self.status_label.setText(tr("Showing {} of {} results").format(self.results_loaded, len(self.last_results)))

        # Trigger Metadata
        if ids_to_fetch:
            self.start_metadata_loading(ids_to_fetch)

    def on_search_finished(self, results):
        self.reset_ui()
        self.chk_search_header.blockSignals(True)
        self.chk_search_header.setChecked(False)
        self.chk_search_header.blockSignals(False)
        self.lbl_search_export.setText(tr("Export Results") + ":")

        if not results:
            self.status_label.setText(tr("No results found."))
            self.last_results = []
            for b in self.export_buttons: b.setEnabled(False)
            self.results_table.setRowCount(0)
            self.result_row_by_sys_id = {}
            self.shelfmark_items_by_sid = {}
            self.title_items_by_sid = {}
            # Disable domain filter button when no results
            self.btn_domain_filter.setEnabled(False)
            self._has_result_domains = False
            self._result_domain_counts = {}
            self._result_domain_map = {}
            return

        self.last_results = results
        self.results_loaded = 0
        self.results_table.setRowCount(0)
        self.result_row_by_sys_id = {}
        self.shelfmark_items_by_sid = {}
        self.title_items_by_sid = {}
        self._res_map_by_sid = {r['display']['id']: r for r in results}

        # Track Responsa expanded term count for status label
        self._responsa_expanded_count = results[0].get('responsa_expanded_count', 0) if results else 0

        # Display explosion guard warning if present
        if results and results[0].get('responsa_warning'):
            warning = results[0]['responsa_warning']
            self.status_label.setText(warning)
            self.status_label.setStyleSheet("color: #f39c12; font-weight: bold;")
            def _restore_status():
                self.status_label.setText(
                    tr("Showing {} of {} results").format(
                        min(self.results_loaded, len(self.last_results)),
                        len(self.last_results)
                    )
                )
                self.status_label.setStyleSheet("")
            QTimer.singleShot(5000, _restore_status)

        for b in self.export_buttons: b.setEnabled(True)

        # Hide Source column if secondary source file (V0.7) is missing or empty
        # If Config.FILE_V7 is missing/empty, it implies we only have one source (V0.8) in index
        has_multiple_sources = os.path.exists(Config.FILE_V7) and os.path.getsize(Config.FILE_V7) > 0
        self.results_table.setColumnHidden(self.COL_SRC, not has_multiple_sources)

        # Initialize domain data (will be populated asynchronously by DomainEnrichmentWorker)
        self._result_domain_map = {}
        self._domain_name_map = {}
        self._result_domain_counts = {}
        self._has_result_domains = False
        self.btn_domain_filter.setEnabled(False)

        self.load_next_batch()

        # Auto-fit columns to content (like double-clicking the column border)
        for col in (self.COL_SYS_ID, self.COL_LIBRARY, self.COL_SHELF, self.COL_IMG):
            self.results_table.resizeColumnToContents(col)

        # Launch domain enrichment worker (async -- results appear first, domains fill in later)
        from gui_threads import DomainEnrichmentWorker
        sys_ids = [r.get('display', {}).get('id') for r in results if r.get('display', {}).get('id')]
        if sys_ids:
            if self._domain_worker and self._domain_worker.isRunning():
                self._domain_worker.wait()
            self._domain_worker = DomainEnrichmentWorker(results)
            self._domain_worker.finished.connect(self._on_domain_enrichment_loaded)
            self._domain_worker.start()

        # Launch PGP badge worker to mark results with transcriptions
        sys_ids = [r.get('display', {}).get('id') for r in results if r.get('display', {}).get('id')]
        if sys_ids:
            if self._pgp_badge_worker and self._pgp_badge_worker.isRunning():
                self._pgp_badge_worker.wait()
            self._pgp_badge_worker = PGPBadgeWorker(sys_ids)
            self._pgp_badge_worker.finished.connect(self._on_pgp_badges_loaded)
            self._pgp_badge_worker.start()

    def _open_results_filter_dialog(self, column):
        if column == 1:
            self.open_list_filter_dialog()
            return

        header_item = self.results_table.horizontalHeaderItem(column)
        column_label = header_item.text() if header_item else str(column)
        current = self.results_filters.get(column, {})
        dlg = ColumnFilterDialog(
            self,
            column_label,
            current_text=current.get("text", ""),
            exclude=current.get("exclude", False),
        )
        if dlg.exec():
            text = dlg.get_text().strip()
            if text:
                self.results_filters[column] = {"text": text, "exclude": dlg.is_exclude()}
            else:
                self.results_filters.pop(column, None)
            self._update_results_filter_indicators()
            self._apply_results_table_filters()


    def _update_results_filter_indicators(self):
        for column in (self.COL_SHELF, self.COL_TITLE, self.COL_SNIPPET, self.COL_DOMAIN):
            self.chk_search_header.set_filter_active(column, column in self.results_filters)

    def _results_filter_text_for_row(self, row, column):
        if column == self.COL_SNIPPET:
            widget = self.results_table.cellWidget(row, column)
            if widget is not None:
                raw = widget.property("filter_text")
                if raw:
                    return str(raw)
                if hasattr(widget, "text"):
                    return widget.text()
        item = self.results_table.item(row, column)
        return item.text() if item else ""

    def _apply_results_table_filters(self):
        # 1. Gather rules
        list_active = self.list_filter_state.get('active', False)
        list_mode = self.list_filter_state.get('mode', 'in')
        target_lists = self.list_filter_state.get('lists', set())

        # Use cached IDs if available, or empty set (will be populated on first activation via update_cache)
        target_sys_ids = self.list_filter_state.get('cached_ids', set())

        # Domain exclusion state
        has_domain_exclusions = bool(self._domain_exclusions) and self._has_result_domains
        hide_uncategorized = "Uncategorized" in self._domain_exclusions if has_domain_exclusions else False

        if not self.results_filters and not list_active and not has_domain_exclusions:
            for row in range(self.results_table.rowCount()):
                self.results_table.setRowHidden(row, False)
            return

        for row in range(self.results_table.rowCount()):
            visible = True

            # A. Check Column Filters
            for column, rule in self.results_filters.items():
                cell_text = self._results_filter_text_for_row(row, column)
                if not self._text_matches_filter(cell_text, rule):
                    visible = False
                    break

            if not visible:
                self.results_table.setRowHidden(row, True)
                continue

            # B. Check List Filter
            if list_active:
                # Get SysID from column 2
                item = self.results_table.item(row, 2)
                sys_id = item.text().strip() if item else ""

                in_list = sys_id in target_sys_ids

                if list_mode == 'in':
                    if not in_list: visible = False
                elif list_mode == 'not_in':
                    if in_list: visible = False

            # C. Check Domain Exclusions
            if visible and has_domain_exclusions:
                item = self.results_table.item(row, self.COL_SYS_ID)
                sys_id = item.text().strip() if item else None
                result_domains = self._result_domain_map.get(sys_id, []) if sys_id else []
                if not result_domains:
                    if hide_uncategorized:
                        visible = False
                elif all(d in self._domain_exclusions for d in result_domains):
                    visible = False

            self.results_table.setRowHidden(row, not visible)

    def _on_pgp_badges_loaded(self, pgp_sys_ids):
        """Handle PGP badge worker results - update badge column for all rows."""
        self._pgp_transcription_sys_ids = pgp_sys_ids
        for row in range(self.results_table.rowCount()):
            item = self.results_table.item(row, self.COL_SYS_ID)
            if item:
                sys_id = item.text().strip()
                if sys_id in pgp_sys_ids:
                    pgp_item = QTableWidgetItem("PGP")
                    pgp_item.setForeground(QColor("#27ae60"))
                    self.results_table.setItem(row, self.COL_PGP, pgp_item)
                else:
                    self.results_table.setItem(row, self.COL_PGP, QTableWidgetItem(""))
        self._apply_results_table_filters()

    def _on_pgp_tags_loaded(self, tags):
        """Handle PGP tags worker results - populate tag dropdown with categorized Hebrew translations."""
        from pgp_tag_translations import get_categorized_tags_for_display
        self._pgp_tags = tags
        lang = CURRENT_LANG  # 'he' or 'en'
        self.tag_search_combo.blockSignals(True)
        self.tag_search_combo.clear()
        self.tag_search_combo.addItem("", "")  # empty placeholder
        categorized = get_categorized_tags_for_display(tags, lang)
        for header, display, en_tag in categorized:
            if en_tag == "":
                # Category header — non-selectable separator
                self.tag_search_combo.addItem(display)
                idx = self.tag_search_combo.count() - 1
                model = self.tag_search_combo.model()
                item = model.item(idx)
                if item:
                    item.setEnabled(False)
            else:
                self.tag_search_combo.addItem(display, en_tag)
        self.tag_search_combo.blockSignals(False)

    def _execute_tag_search(self):
        """Execute a search by PGP tag from the dropdown."""
        tag = self.tag_search_combo.currentData()
        if not tag:
            # Fallback: user may have typed text directly
            tag = self.tag_search_combo.currentText().strip()
        if not tag:
            return
        self.status_label.setText(tr("Searching tag: {}...").format(tag))
        if self._pgp_tag_search_worker and self._pgp_tag_search_worker.isRunning():
            self._pgp_tag_search_worker.wait()
        self._pgp_tag_search_worker = PGPTagSearchWorker(tag)
        self._pgp_tag_search_worker.finished.connect(self._on_tag_search_results)
        self._pgp_tag_search_worker.start()

    def _on_tag_search_results(self, tag, results):
        """Handle tag search results - display in results table."""
        if not results:
            self.status_label.setText(tr("No results for tag: {}").format(tag))
            self.last_results = []
            self.results_loaded = 0
            self.results_table.setRowCount(0)
            self.result_row_by_sys_id = {}
            self.shelfmark_items_by_sid = {}
            self.title_items_by_sid = {}
            return

        # Filter to only results whose sys_id exists in local csv_bank
        valid_results = []
        for r in results:
            sid = r.get('sys_id', '')
            if sid and self.meta_mgr and self.meta_mgr.csv_bank.get(sid):
                valid_results.append(r)

        if not valid_results:
            self.status_label.setText(tr("No local results for tag: {}").format(tag))
            return

        # Convert tag results to search result format
        formatted = []
        for r in valid_results:
            sid = r.get('sys_id', '')
            shelf, title = self.meta_mgr.get_meta_for_id(sid) if self.meta_mgr else ('Unknown', '')
            library_code = self.meta_mgr.get_library_for_id(sid) if self.meta_mgr else ''
            desc = r.get('description', '') or ''
            doc_type = r.get('document_type', '') or ''
            transcription = r.get('transcription', '') or ''
            if transcription:
                # Use PGP transcription as snippet (Hebrew text)
                clean_text = transcription.replace('\n', ' ').replace('\r', ' ').strip()
                snippet = clean_text[:150] + ('...' if len(clean_text) > 150 else '')
            elif desc:
                snippet = f"{doc_type}: {desc[:120]}..."
            else:
                snippet = doc_type
            formatted.append({
                'display': {
                    'id': sid,
                    'shelfmark': shelf,
                    'title': title or '',
                    'library': library_code or '',
                    'img': '',
                    'source': '',
                },
                'snippet': snippet,
                'raw_header': '',
            })

        # Mark all as PGP
        self._pgp_transcription_sys_ids = {r['display']['id'] for r in formatted}

        self.chk_search_header.blockSignals(True)
        self.chk_search_header.setChecked(False)
        self.chk_search_header.blockSignals(False)

        self.last_results = formatted
        self.results_loaded = 0
        self.results_table.setRowCount(0)
        self.result_row_by_sys_id = {}
        self.shelfmark_items_by_sid = {}
        self.title_items_by_sid = {}
        self._res_map_by_sid = {r['display']['id']: r for r in formatted}

        for b in self.export_buttons: b.setEnabled(True)
        self.results_table.setColumnHidden(self.COL_SRC, True)

        self.load_next_batch()
        self.status_label.setText(tr("Tag: {} - {} results").format(tag, len(formatted)))


    def _search_by_pgp_tag(self, tag):
        """Entry point for searching by PGP tag (from browse/result dialog links)."""
        self.tabs.setCurrentWidget(self.search_tab)
        # Switch to PGP Tags mode
        self.mode_combo.setCurrentIndex(self.MODE_PGP_TAGS)
        if hasattr(self, 'tag_search_combo'):
            # Find combo item by English tag stored in userData
            idx = self.tag_search_combo.findData(tag)
            if idx >= 0:
                self.tag_search_combo.setCurrentIndex(idx)
            else:
                self.tag_search_combo.setCurrentText(tag)
        self._execute_tag_search()

    def _open_comp_filter_dialog(self, column):
        header_item = self.comp_tree.headerItem()
        column_label = header_item.text(column) if header_item else str(column)
        current = self.comp_filters.get(column, {})
        dlg = ColumnFilterDialog(
            self,
            column_label,
            current_text=current.get("text", ""),
            exclude=current.get("exclude", False),
        )
        if dlg.exec():
            text = dlg.get_text().strip()
            if text:
                self.comp_filters[column] = {"text": text, "exclude": dlg.is_exclude()}
            else:
                self.comp_filters.pop(column, None)
            self._update_comp_filter_indicators()
            self._apply_comp_tree_filters()

    def _update_comp_filter_indicators(self):
        for column in (self.comp_col_library, self.comp_col_shelfmark, self.comp_col_title, self.comp_col_context, self.comp_col_ms_context):
            self.chk_comp_header.set_filter_active(column, column in self.comp_filters)

    def _apply_comp_tree_filters(self):
        root = self.comp_tree.invisibleRootItem()
        if not self.comp_filters:
            def unhide(node):
                node.setHidden(False)
                for j in range(node.childCount()):
                    unhide(node.child(j))

            for i in range(root.childCount()):
                unhide(root.child(i))
            return

        def visit(node):
            visible_any = False
            for i in range(node.childCount()):
                if visit(node.child(i)):
                    visible_any = True

            data = node.data(0, Qt.ItemDataRole.UserRole + 1)
            if data:
                matches = self._comp_data_matches_filters(node, data)
                node_visible = matches or visible_any
            else:
                node_visible = visible_any

            node.setHidden(not node_visible)
            return node_visible

        for i in range(root.childCount()):
            visit(root.child(i))

    def _comp_data_matches_filters(self, node, data):
        for column, rule in self.comp_filters.items():
            if column == 1:
                text = data.get("shelfmark", "")
                if not text:
                    text = node.text(1)
            elif column == 2:
                text = data.get("title", "")
                if not text:
                    text = node.text(2)
            elif column == self.comp_col_context:
                text = data.get("source_ctx", "")
            elif column == self.comp_col_ms_context:
                text = data.get("ms_ctx", "")
            else:
                continue
            if not self._text_matches_filter(text, rule):
                return False
        return True

    def _text_matches_filter(self, text, rule):
        needle = (rule.get("text") or "").strip()
        if not needle:
            return True
        haystack = (text or "").lower()
        contains = needle.lower() in haystack
        if rule.get("exclude"):
            return not contains
        return contains

    def start_metadata_loading(self, ids):
        if not ids:
            return
        logger.debug("start_metadata_loading: %d ids, sample=%s", len(ids), ids[:10])

        if self.meta_loader and self.meta_loader.isRunning():
            self.meta_loader.request_cancel()
            self.meta_loader.wait()

        self.meta_cached_count = len([sid for sid in ids if sid and sid in self.meta_mgr.nli_cache])
        self.meta_to_fetch_count = len([sid for sid in ids if sid and sid not in self.meta_mgr.nli_cache])
        self.meta_progress_current = 0

        # Update initial metadata from cache for all rows
        for res in self.last_results:
            sid = res['display']['id']
            shelf = res['display'].get('shelfmark', '')
            title = res['display'].get('title', '')

            # Prefer fresh cache if available
            _, _, cached_shelf, cached_title = self._get_meta_for_header(res.get('raw_header', ''))
            shelf = cached_shelf or shelf
            title = cached_title or title

            if sid in self.shelfmark_items_by_sid and shelf:
                self.shelfmark_items_by_sid[sid].setText(shelf)
                res['display']['shelfmark'] = shelf

            if sid in self.title_items_by_sid and title:
                self.title_items_by_sid[sid].setText(title)
                res['display']['title'] = title

        self._apply_results_table_filters()

        if self.meta_to_fetch_count == 0:
            self.status_label.setText(tr("Metadata already loaded for {} items.").format(self.meta_cached_count))
            return

        self.meta_loader = ShelfmarkLoaderThread(self.meta_mgr, ids)
        self.meta_loader.progress_signal.connect(self.on_meta_progress)
        self.meta_loader.finished_signal.connect(self.on_meta_finished)
        self.meta_loader.error_signal.connect(lambda err: QMessageBox.critical(self, tr("Metadata Error"), err))
        self.status_label.setText(self._format_metadata_status())
        self.meta_loader.start()

    def on_meta_progress(self, curr, total, sid):
        self.meta_progress_current = curr
        self.status_label.setText(self._format_metadata_status())

        meta = self.meta_mgr.nli_cache.get(sid, {})
        shelf = meta.get('shelfmark', 'Unknown')
        title = meta.get('title', '')

        if sid in self.shelfmark_items_by_sid:
            if sid not in self.shelfmark_items_by_sid or sid not in self.title_items_by_sid:
                logger.debug(
                    "Meta progress sid not in table maps: sid=%s in_shelf=%s in_title=%s",
                    sid,
                    sid in self.shelfmark_items_by_sid,
                    sid in self.title_items_by_sid,
                )

            try:
                self.shelfmark_items_by_sid[sid].setText(shelf)
            except RuntimeError:
                pass # Item deleted

        if sid in self.title_items_by_sid:
            try:
                self.title_items_by_sid[sid].setText(title)
            except RuntimeError:
                pass # Item deleted

        if not hasattr(self, '_res_map_by_sid'):
            self._res_map_by_sid = {r['display']['id']: r for r in self.last_results} # Lazy build or build in search_finished

        if sid in self._res_map_by_sid:
            r = self._res_map_by_sid[sid]
            r['display']['shelfmark'] = shelf
            r['display']['title'] = title

        self._apply_results_table_filters()

    def on_meta_finished(self, cancelled):
        total_loaded = self.meta_cached_count + self.meta_progress_current
        total_expected = self.meta_cached_count + self.meta_to_fetch_count
        if cancelled:
            self.status_label.setText(tr("Metadata load cancelled. Loaded {}/{}.").format(total_loaded, total_expected))
        else:
            self.status_label.setText(tr("Loaded {} items.").format(total_expected))
        self.meta_loader = None

    def _format_metadata_status(self):
        total_expected = self.meta_cached_count + self.meta_to_fetch_count
        total_loaded = self.meta_cached_count + self.meta_progress_current
        progress_part = ""
        if self.meta_to_fetch_count:
            progress_part = f" ({self.meta_progress_current}/{self.meta_to_fetch_count})"
        return tr("Metadata loaded: {}/{}").format(total_loaded, total_expected) + progress_part

    def _create_action_button(self, content, tooltip, callback=None, parent=None):
        btn = QToolButton(parent or self.results_table)
        if isinstance(content, str):
            btn.setText(content)
            f = btn.font()
            f.setPointSize(12)
            btn.setFont(f)
        else:
            btn.setIcon(content)
        btn.setToolTip(tooltip)
        btn.setAutoRaise(True)
        btn.setCursor(Qt.CursorShape.PointingHandCursor)
        btn.setFixedSize(28, 24)
        if callback:
            btn.clicked.connect(callback)
        return btn

    def _is_item_in_non_recent_list(self, sys_id, img=None, fl_id=None):
        if not self.lists_mgr or not sys_id:
            return False
        item_id = self.lists_mgr._build_item_id(sys_id, img=img, fl_id=fl_id)
        item = self.lists_mgr.get_item(item_id)
        return bool(item and item.get('lists'))

    def _set_add_to_list_button_label(self, button, in_list):
        if not button:
            return
        button.setText(_format_add_to_list_label(in_list))

    def _update_browse_add_to_list_button(self):
        if not hasattr(self, 'btn_browse_add_to_list'):
            return
        in_list = self._is_item_in_non_recent_list(
            self.current_browse_sid,
            img=self.current_browse_p,
            fl_id=self._normalize_fl_id(self.browse_fl_input.text().strip()),
        )
        self._set_add_to_list_button_label(self.btn_browse_add_to_list, in_list)

    def _update_search_row_list_indicator(self, row, res=None):
        actions_widget = self.results_table.cellWidget(row, self.COL_ACTIONS)
        if not isinstance(actions_widget, ActionsHoverWidget):
            return
        if res is None:
            item = self.results_table.item(row, self.COL_SYS_ID)
            if item:
                res = item.data(Qt.ItemDataRole.UserRole)
        if not res:
            return

        display = res.get('display', {}) if isinstance(res, dict) else {}
        sys_id = display.get('id')
        img = display.get('img')
        fl_id = self._normalize_fl_id(self._extract_fl_id(res))
        in_list = self._is_item_in_non_recent_list(sys_id, img=img, fl_id=fl_id)

        for btn in actions_widget.buttons:
            if btn.property("action_role") == "list_star":
                btn.setText(_format_list_star(in_list))
                break

    def _update_search_action_stars(self):
        if not hasattr(self, 'results_table'):
            return
        for row in range(self.results_table.rowCount()):
            self._update_search_row_list_indicator(row)

    def on_table_cell_entered(self, row, col):
        if row == self.hovered_row:
            return

        # Hide previous
        if self.hovered_row != -1:
            w = self.results_table.cellWidget(self.hovered_row, self.COL_ACTIONS)
            if isinstance(w, ActionsHoverWidget):
                w.set_buttons_visible(False)

        self.hovered_row = row

        # Show new
        w = self.results_table.cellWidget(row, self.COL_ACTIONS)
        if isinstance(w, ActionsHoverWidget):
            w.set_buttons_visible(True)

    def on_lists_table_cell_entered(self, row, col):
        if not hasattr(self, "lists_items_table"):
            return
        if row == self.lists_hovered_row:
            return

        if self.lists_hovered_row != -1:
            w = self.lists_items_table.cellWidget(self.lists_hovered_row, 6)
            if isinstance(w, ActionsHoverWidget):
                w.set_buttons_visible(False)

        self.lists_hovered_row = row
        w = self.lists_items_table.cellWidget(row, 6)
        if isinstance(w, ActionsHoverWidget):
            w.set_buttons_visible(True)

    def eventFilter(self, source, event):
        # Smart tooltips for truncated cell text in results table
        if source == self.results_table.viewport() and event.type() == QEvent.Type.ToolTip:
            pos = event.pos()
            index = self.results_table.indexAt(pos)
            if index.isValid():
                col = index.column()
                # Skip checkbox and actions columns
                if col not in (self.COL_CHECKBOX, self.COL_ACTIONS, self.COL_IMG):
                    item = self.results_table.item(index.row(), col)
                    if item and item.text():
                        rect = self.results_table.visualRect(index)
                        fm = self.results_table.fontMetrics()
                        text_width = fm.horizontalAdvance(item.text())
                        if text_width > rect.width() - 8:
                            QToolTip.showText(event.globalPos(), item.text())
                            return True
            QToolTip.hideText()
            return True
        if source == self.results_table and event.type() == QEvent.Type.Leave:
            if self.hovered_row != -1:
                w = self.results_table.cellWidget(self.hovered_row, self.COL_ACTIONS)
                if isinstance(w, ActionsHoverWidget):
                    w.set_buttons_visible(False)
                self.hovered_row = -1
        if hasattr(self, "lists_items_table") and source == self.lists_items_table and event.type() == QEvent.Type.Leave:
            if self.lists_hovered_row != -1:
                w = self.lists_items_table.cellWidget(self.lists_hovered_row, 6)
                if isinstance(w, ActionsHoverWidget):
                    w.set_buttons_visible(False)
                self.lists_hovered_row = -1
        return super().eventFilter(source, event)

    def _collect_sorted_results(self):
        sorted_results = []
        rows = self.results_table.rowCount()
        for i in range(rows):
            item = self.results_table.item(i, self.COL_SYS_ID)
            if item:
                res = item.data(Qt.ItemDataRole.UserRole)
                if res:
                    sorted_results.append(res)
        if not sorted_results:
            sorted_results = self.last_results
        return sorted_results

    def _extract_fl_id(self, res):
        if not isinstance(res, dict):
            return None
        display = res.get('display', {}) or {}
        raw_header = res.get('raw_header') or res.get('full_header', '')
        fl_id = None

        if raw_header and self.meta_mgr:
            parsed = self.meta_mgr.parse_full_id_components(raw_header)
            fl_id = parsed.get('fl_id') or fl_id

        if not fl_id:
            fl_id = display.get('fl_id')

        if not fl_id:
            img_field = display.get('img')
            if img_field:
                m = re.search(r'FL\\s*-?(\\d+)', str(img_field))
                if m:
                    fl_id = m.group(1)

        if not fl_id:
            uid_field = res.get('uid')
            if uid_field:
                m = re.search(r'FL\\s*-?(\\d+)', str(uid_field))
                if m:
                    fl_id = m.group(1)

        return fl_id

    def show_full_text(self):
        row = self.results_table.currentRow()
        if row < 0: return

        sorted_results = self._collect_sorted_results()
        if not sorted_results:
            return
        if row >= len(sorted_results):
            row = 0
        ResultDialog(self, sorted_results, row, self.meta_mgr, self.searcher).exec()

    def show_full_text_for_result(self, res):
        if not res:
            return
        sorted_results = self._collect_sorted_results()
        if not sorted_results:
            sorted_results = [res]
        target_index = 0
        for idx, candidate in enumerate(sorted_results):
            if candidate is res:
                target_index = idx
                break
            cand_id = candidate.get('display', {}).get('id')
            if cand_id and cand_id == res.get('display', {}).get('id'):
                target_index = idx
        ResultDialog(self, sorted_results, target_index, self.meta_mgr, self.searcher).exec()

    def open_result_in_browse_from_table(self, res):
        if not res:
            return
        display = res.get('display', {})
        sid = display.get('id')
        shelf = ""
        title = ""
        if sid:
            shelf, title = self.meta_mgr.get_meta_for_id(sid)
        if not shelf:
            shelf = display.get('shelfmark', '')
        if not title:
            title = display.get('title', '')
        fl_id = self._extract_fl_id(res)
        self.open_result_in_browse(res, shelfmark=shelf, title=title, fl_id=fl_id)

    def on_search_select_all_toggled(self, checked):
        """Handle Select All checkbox toggle."""
        state = Qt.CheckState.Checked if checked else Qt.CheckState.Unchecked
        self.results_table.blockSignals(True)
        for i in range(self.results_table.rowCount()):
            item = self.results_table.item(i, self.COL_CHECKBOX)
            if item:
                item.setCheckState(state)
        self.results_table.blockSignals(False)
        self._update_search_export_label()

    def on_search_result_item_changed(self, item):
        """Handle individual checkbox changes in search results."""
        if item.column() != self.COL_CHECKBOX:
            return

        # Check if all items are checked to sync "Select All"
        all_checked = True
        has_selection = False

        # Avoid full iteration if possible, but we need to check "Select All" status
        rows = self.results_table.rowCount()
        # To optimize, we just iterate. Rows are usually < 500.
        for i in range(rows):
            it = self.results_table.item(i, self.COL_CHECKBOX)
            if it:
                if it.checkState() == Qt.CheckState.Unchecked:
                    all_checked = False
                else:
                    has_selection = True

        # If no items, all_checked is trivially true but we don't want to check the box
        if rows == 0:
            all_checked = False

        self.chk_search_header.blockSignals(True)
        self.chk_search_header.setChecked(all_checked)
        self.chk_search_header.blockSignals(False)

        self._update_search_export_label(has_selection)

    def _update_search_export_label(self, has_selection=None):
        if has_selection is None:
            # Check if any selected
            has_selection = False
            for i in range(self.results_table.rowCount()):
                it = self.results_table.item(i, self.COL_CHECKBOX)
                if it and it.checkState() == Qt.CheckState.Checked:
                    has_selection = True
                    break

        if has_selection:
            self.lbl_search_export.setText(tr("Export selected results") + ":")
        else:
            self.lbl_search_export.setText(tr("Export Results") + ":")

        # Update add-to-list button state
        if hasattr(self, 'btn_add_to_list'):
            self.btn_add_to_list.setEnabled(has_selection)

    def search_add_selected_to_list(self):
        """Add selected search results to a list."""
        if not self.lists_mgr:
            return

        # Collect selected items
        items = []
        for i in range(self.results_table.rowCount()):
            chk = self.results_table.item(i, self.COL_CHECKBOX)
            if chk and chk.checkState() == Qt.CheckState.Checked:
                res = chk.data(Qt.ItemDataRole.UserRole)
                if res:
                    display = res.get('display', {})
                    sys_id = display.get('id')
                    if sys_id:
                        items.append({
                            'sys_id': sys_id,
                            'img': display.get('img'),
                            'fl_id': self._normalize_fl_id(self._extract_fl_id(res))
                        })

        if items:
            source = f"Search: {self.last_search_query[:50]}" if self.last_search_query else "Search"
            self.show_add_to_list_menu(items, source=source, anchor_widget=self.btn_add_to_list)

    def search_add_row_to_list(self, res, anchor_widget=None):
        """Add a single search result row to a list."""
        if not self.lists_mgr or not isinstance(res, dict):
            return
        display = res.get('display', {})
        sys_id = display.get('id')
        if not sys_id:
            return
        items = [{
            'sys_id': sys_id,
            'img': display.get('img'),
            'fl_id': self._normalize_fl_id(self._extract_fl_id(res)),
        }]
        source = f"Search: {self.last_search_query[:50]}" if self.last_search_query else "Search"
        self.show_add_to_list_menu(items, source=source, anchor_widget=anchor_widget or self.results_table)

    def _collect_selected_comp_pages(self):
        selected = []
        sel_main, sel_appx, sel_filt, sel_filt_appx, sel_known = self._collect_checked_comp_items_struct()

        def add_item(item):
            if not item:
                return
            if item.get('type') in ('manuscript', 'part'):
                pages = item.get('pages', [])
                if pages:
                    selected.extend(pages)
                else:
                    selected.append(item)
            else:
                selected.append(item)

        for item in sel_main:
            add_item(item)
        for group_items in sel_appx.values():
            for item in group_items:
                add_item(item)
        for item in sel_filt:
            add_item(item)
        for group_items in sel_filt_appx.values():
            for item in group_items:
                add_item(item)
        for item in sel_known:
            add_item(item)

        return selected

    def comp_add_selected_to_list(self):
        """Add selected composition results to a list."""
        if not self.lists_mgr:
            return

        pages = self._collect_selected_comp_pages()
        items = []
        for page in pages:
            raw_header = page.get('raw_header', '')
            sys_id, p_num, _, _ = self._get_meta_for_header(raw_header)
            if not sys_id:
                continue
            parsed = self.meta_mgr.parse_full_id_components(raw_header) if raw_header else {}
            fl_id = self._normalize_fl_id(parsed.get('fl_id'))
            items.append({
                'sys_id': sys_id,
                'img': p_num,
                'fl_id': fl_id,
            })

        if items:
            title = self.comp_title_input.text().strip()
            source = f"Composition: {title[:50]}" if title else "Composition Search"
            self.show_add_to_list_menu(items, source=source, anchor_widget=self.btn_comp_add_to_list)

    def open_result_in_browse(self, res, shelfmark=None, title=None, fl_id=None):
        sid = None
        if isinstance(res, dict):
            display = res.get('display')
            if isinstance(display, dict):
                sid = display.get('id')
            if not sid:
                sid = res.get('sys_id')
            if not sid:
                raw_header = res.get('raw_header') or res.get('full_header')
                if raw_header:
                    sid, _ = self.meta_mgr.parse_header_smart(raw_header)
        if not sid:
            QMessageBox.warning(self, tr("Error"), tr("No System ID found for this result."))
            return

        # Persist highlight data for browse pane
        self.browse_highlight_data = res.get('page_highlights', []) if isinstance(res, dict) else []
        self.browse_highlight_pattern = res.get('highlight_pattern') if isinstance(res, dict) else None

        derived_fl_id = fl_id or self._extract_fl_id(res)
        if shelfmark:
            # Add library prefix to shelfmark
            display_shelf = shelfmark
            library_code = self.meta_mgr.get_library_for_id(sid)
            if library_code:
                library = get_library_display(library_code, short=False)
                display_shelf = f"{library} | {shelfmark}"
            info_text = f"<b>{display_shelf}</b>"
            if title:
                info_text += f"<br>{title}"
            self.browse_info_lbl.setText(info_text)
        if derived_fl_id:
            fl_digits = re.sub(r"\\D", "", str(derived_fl_id))
            self.browse_fl_input.setText(f"FL{fl_digits}" if fl_digits else str(derived_fl_id))
            self._set_last_browse_field("fl")
        else:
            self.browse_fl_input.setText("")
            self.browse_shelf_input.clear()        # Clear stale shelfmark
            self._set_last_browse_field("sys")     # Force sys_id priority
        self.browse_sys_input.setText(sid)
        self.tabs.setCurrentWidget(self.browse_tab)
        self.browse_load()

    def send_result_to_composition(self, res, source_text=None, title=None):
        if not source_text:
            if not res.get('full_text'):
                res['full_text'] = self.searcher.get_full_text_by_id(res.get('uid', '')) or res.get('text', '')
            source_text = res.get('full_text') or res.get('text', '')
        self.comp_text_area.setPlainText(source_text)
        if title:
            self.comp_title_input.setText(title)
            
        sys_id = None
        raw_header = res.get('raw_header') or res.get('full_header', '')
        
        # 1. Try to parse strictly 99... from header
        if raw_header and self.meta_mgr:
            parsed_sid, _ = self.meta_mgr.parse_header_smart(raw_header)
            sys_id = parsed_sid
            
        # 2. Fallback to existing display ID if parsing failed
        if not sys_id:
            sys_id = res['display'].get('id')
        # ------------------------------------------------------

        if sys_id:
            entries = list(self.excluded_raw_entries)
            # Add only if not already present
            if sys_id not in entries:
                entries.append(sys_id)
                self.set_excluded_entries("\n".join(entries))
                
        self.tabs.setCurrentWidget(self.composition_tab)
        self.comp_text_area.setFocus()

    def _sanitize_for_excel(self, text):
        """Cleans text to prevent Excel XML corruption.

        Uses shared sanitization utility for consistency with Web app.
        """
        return shared_sanitize_excel(text)

    def _add_docx_highlighted_runs(self, paragraph, text):
        parts = str(text or "").split('*')
        for i, part in enumerate(parts):
            if not part:
                continue
            run = paragraph.add_run(part)
            if i % 2 == 1:
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                run.font.bold = True

    def _set_paragraph_rtl(self, paragraph):
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        ppr = paragraph._p.get_or_add_pPr()
        bidi = ppr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            ppr.append(bidi)
        bidi.set(qn("w:val"), "1")

    def _set_table_rtl(self, table):
        tbl_pr = table._tbl.tblPr
        bidi_visual = tbl_pr.find(qn("w:bidiVisual"))
        if bidi_visual is None:
            bidi_visual = OxmlElement("w:bidiVisual")
            tbl_pr.append(bidi_visual)
        bidi_visual.set(qn("w:val"), "1")

    def _set_table_width_pct(self, table, pct=100):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "pct")
        tbl_w.set(qn("w:w"), str(int(pct * 50)))

    def export_results(self, fmt='xlsx'):
        """
        Export results handling specific formats directly.
        fmt: 'xlsx', 'csv', 'txt', or 'docx'
        """
        base_path = self._default_report_path(self.last_search_query, tr("Search_Results"))
        default_path = os.path.splitext(base_path)[0] + f".{fmt}"

        filters = {'xlsx': "Excel (*.xlsx)", 'csv': "CSV (*.csv)", 'txt': "Text (*.txt)", 'docx': "Word (*.docx)"}
        selected_filter = filters.get(fmt, "All Files (*.*)")

        path, _ = QFileDialog.getSaveFileName(self, tr("Export Results"), default_path, selected_filter)
        if not path: return

        # Prepare tabular data
        headers = [tr("System ID"), tr("Library"), tr("Shelfmark"), tr("Title"), tr("Image/Page"), tr("Source"), tr("Snippet")]
        data_rows = []

        # Collect results to export (Selected or All)
        results_to_export = []

        # Check if any are selected in the table
        has_selection = False
        selected_rows_data = []

        # Iterate table to respect user selection and visual order
        rows = self.results_table.rowCount()
        for i in range(rows):
            chk_item = self.results_table.item(i, self.COL_CHECKBOX)
            if chk_item and chk_item.checkState() == Qt.CheckState.Checked:
                has_selection = True
                res = chk_item.data(Qt.ItemDataRole.UserRole)
                if res:
                    selected_rows_data.append(res)

        if has_selection:
            results_to_export = selected_rows_data
        else:
            # Fallback to last_results (original order) if nothing selected
            results_to_export = self.last_results

        for r in results_to_export:
            d = r['display']
            sid = d.get('id', '')

            # Fetch fresh metadata (Important for Lab Mode)
            shelf, title = self.meta_mgr.get_meta_for_id(sid)
            if not shelf or shelf == "Unknown":
                shelf = d.get('shelfmark', '')
            if not title:
                title = d.get('title', '')

            # Use raw_file_hl so highlight markers remain intact
            # Clean snippet: remove newlines (input is now clean text with asterisks)
            raw_hl = r.get('raw_file_hl', '')
            snippet = str(raw_hl).strip().replace('\n', ' ').replace('\r', ' ')
            snippet = re.sub(r'\s+', ' ', snippet)

            # Get library info (full name for export)
            library_code = self.meta_mgr.get_library_for_id(sid) if sid else ''
            library_name = get_library_display(library_code, short=False) if library_code else ''

            data_rows.append([
                sid,
                library_name,
                shelf,
                title,
                str(d.get('img', '')),
                d.get('source', ''),
                snippet
            ])

        credit_text = self._get_credit_header()
        def _strip_search_prefix(text):
            return re.sub(r'^(\\?\\?\\?|\\?\\?|\\?|=|~|/|\\$|#)\\s*', '', text or "")
        export_query = _strip_search_prefix(self.last_search_query)
        search_info_lines = [
            tr("Search Query") + f": {export_query}",
            tr("Search Mode") + f": {self.mode_combo.currentText()}",
            tr("Gap") + f": {self.gap_input.text()}",
            tr("Lab Mode") + f": {'On' if self.btn_lab_mode_toggle.isChecked() else 'Off'}",
        ]
        if self.btn_lab_mode_toggle.isChecked():
            search_info_lines.append(tr("Deep Scan") + f": {'On' if self.chk_lab_deep.isChecked() else 'Off'}")
        search_info_text = "\n".join(search_info_lines)

        # --- XLSX with inline highlighting ---
        if fmt == 'xlsx':
            try:
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = tr("Search Results")
                ws.sheet_view.rightToLeft = True

                # Fonts used for rich text snippets
                font_red = InlineFont(color='FF0000', b=True)
                font_normal = InlineFont(color='000000', b=False)

                # Helper to write rich text cells
                def write_rich_cell(row, col, text):
                    safe_text = self._sanitize_for_excel(text)

                    if '*' not in safe_text:
                        ws.cell(row=row, column=col, value=safe_text)
                        return

                    # Split by asterisk markers
                    parts = safe_text.split('*')
                    rich_string = CellRichText()

                    for i, part in enumerate(parts):
                        if not part:
                            continue
                        # Odd indices represent highlighted text
                        if i % 2 == 1:
                            rich_string.append(TextBlock(font_red, part))
                        else:
                            # Even indices are plain text
                            rich_string.append(TextBlock(font_normal, part))

                    ws.cell(row=row, column=col, value=rich_string)

                # Credit header
                current_row = 1
                for line in credit_text.split('\n'):
                    if not line.strip(): continue
                    cell = ws.cell(row=current_row, column=1, value=self._sanitize_for_excel(line))
                    cell.font = Font(bold=True, color="555555")
                    current_row += 1
                for line in search_info_text.split('\n'):
                    if not line.strip():
                        continue
                    cell = ws.cell(row=current_row, column=1, value=self._sanitize_for_excel(line))
                    cell.font = Font(bold=True, color="555555")
                    current_row += 1
                current_row += 1

                # Table headers
                for col_idx, header in enumerate(headers, 1):
                    cell = ws.cell(row=current_row, column=col_idx, value=header)
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
                current_row += 1

                # Data rows
                for row_data in data_rows:
                    for col_idx, val in enumerate(row_data, 1):
                        val_str = str(val)

                        # Column 7 holds the snippet (after adding Library column)
                        if col_idx == 7:
                            write_rich_cell(current_row, col_idx, val_str)
                        else:
                            # Strip markers/HTML in other columns
                            clean_val = val_str.replace('*', '')
                            ws.cell(row=current_row, column=col_idx, value=self._sanitize_for_excel(clean_val))

                    current_row += 1

                # Column widths
                ws.column_dimensions['A'].width = 15
                ws.column_dimensions['B'].width = 20
                ws.column_dimensions['C'].width = 25  # Library
                ws.column_dimensions['D'].width = 40
                ws.column_dimensions['G'].width = 80  # Wider snippet column

                wb.save(path)
                self._save_last_folder(path)
                QMessageBox.information(self, tr("Saved"), tr("Saved to {}").format(path))

            except Exception as e:
                QMessageBox.critical(self, tr("Error"), f"Failed to save XLSX:\n{str(e)}")

        # --- CSV ---
        elif fmt == 'csv':
            try:
                with open(path, 'w', encoding='utf-8-sig', newline='') as f:
                    f.write(credit_text)
                    f.write("\n" + search_info_text + "\n")
                    writer = csv.writer(f)
                    writer.writerow([])
                    writer.writerow(headers)
                    for row in data_rows:
                        # Strip highlight markers for CSV
                        clean_row = [str(val).replace('*', '') for val in row]
                        writer.writerow(clean_row)
                self._save_last_folder(path)
                QMessageBox.information(self, tr("Saved"), tr("Saved to {}").format(path))
            except Exception as e:
                QMessageBox.critical(self, tr("Error"), f"Failed to save CSV:\n{str(e)}")

        # --- DOCX ---
        elif fmt == 'docx':
            try:
                doc = Document()
                for line in credit_text.split('\n'):
                    if not line.strip():
                        continue
                    p = doc.add_paragraph(line.strip())
                    if p.runs:
                        p.runs[0].font.bold = True
                for line in search_info_text.split('\n'):
                    if not line.strip():
                        continue
                    doc.add_paragraph(line.strip())
                doc.add_paragraph("")

                headers = [tr("System ID"), tr("Library"), tr("Shelfmark"), tr("Title"), tr("Image/Page"), tr("Source"), tr("Snippet")]
                table = doc.add_table(rows=1, cols=len(headers))
                self._set_table_width_pct(table, 100)
                hdr_cells = table.rows[0].cells
                for idx, header in enumerate(headers):
                    hdr_cells[idx].text = header

                for row in data_rows:
                    row_cells = table.add_row().cells
                    for col_idx, val in enumerate(row):
                        cell = row_cells[col_idx]
                        if col_idx == 6:  # Snippet column (after adding Library)
                            cell.text = ""
                            self._add_docx_highlighted_runs(cell.paragraphs[0], val)
                        else:
                            cell.text = str(val).replace('*', '')

                if CURRENT_LANG == "he":
                    doc.styles["Normal"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    for p in doc.paragraphs:
                        self._set_paragraph_rtl(p)
                    self._set_table_rtl(table)
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                self._set_paragraph_rtl(p)

                doc.save(path)
                self._save_last_folder(path)
                QMessageBox.information(self, tr("Saved"), tr("Saved to {}").format(path))
            except Exception as e:
                QMessageBox.critical(self, tr("Error"), f"Failed to save DOCX:\n{str(e)}")

        # --- TXT ---
        else:
            try:
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(credit_text)
                    f.write("\n" + search_info_text + "\n\n")
                    for r in results_to_export:
                        # Clean snippet: remove newlines for single-line export
                        snippet = r.get('raw_file_hl', '').strip().replace('\n', ' ').replace('\r', '')
                        f.write(f"=== {r['display']['shelfmark']} | {r['display']['title']} ===\n{snippet}\n\n")
                self._save_last_folder(path)
                QMessageBox.information(self, tr("Saved"), tr("Saved to {}").format(path))
            except Exception as e:
                QMessageBox.critical(self, tr("Error"), f"Failed to save TXT:\n{str(e)}")

    def export_comp_report(self, fmt='xlsx'):
        # 1. איסוף נתונים (לוגיקה יציבה)

        # Check for Selection
        has_selection = bool(self._collect_checked_comp_page_uids())

        if has_selection:
            # Reconstruct lists from tree selection
            c_main, c_appx, c_filt, c_filt_appx, c_known = self._collect_checked_comp_items_struct()
        else:
            # Use all data
            c_main = self.comp_main
            c_appx = self.comp_appendix
            c_filt = self.comp_filtered_main
            c_filt_appx = self.comp_filtered_appendix
            c_known = self.comp_known

        all_filtered = c_filt[:]
        for v in c_filt_appx.values():
            all_filtered.extend(v)

        if not (c_main or c_appx or c_known or all_filtered):
            QMessageBox.warning(self, tr("Save"), tr("No composition data to export."))
            return

        # 2. טעינת מטא-דאטה
        all_ids = []
        def collect_ids(item_list):
            for item in item_list:
                if item.get('type') == 'part':
                    # For Parts, collect all folios
                    folios = item.get('folios', [])
                    all_ids.extend(folios)
                elif item.get('type') == 'manuscript' and item.get('sys_id'):
                    all_ids.append(item['sys_id'])
                else:
                    sid, _ = self.meta_mgr.parse_header_smart(item.get('raw_header', ''))
                    if sid: all_ids.append(sid)

        collect_ids(c_main)
        for group_items in c_appx.values(): collect_ids(group_items)
        collect_ids(c_known)
        collect_ids(all_filtered)

        unique_ids = list(set(all_ids))
        if unique_ids:
            missing = [uid for uid in unique_ids if uid not in self.meta_mgr.nli_cache]
            if missing:
                self._fetch_metadata_with_dialog(missing, title=tr("Fetching metadata before export..."))

        # 3. שמירה
        comp_title = self.comp_title_input.text().strip() or tr("Untitled Composition")
        base_path = self._default_report_path(comp_title, tr("Composition_Report"))
        default_path = os.path.splitext(base_path)[0] + f".{fmt}"
        
        filters = {'xlsx': "Excel (*.xlsx)", 'csv': "CSV (*.csv)", 'txt': "Text (*.txt)", 'docx': "Word (*.docx)"}
        selected_filter = filters.get(fmt, "All Files (*.*)")

        path, _ = QFileDialog.getSaveFileName(self, tr("Save Report"), default_path, selected_filter)
        if not path: return

        credit_text = self._get_credit_header()
        query_text = self.comp_text_area.toPlainText().strip()
        comp_settings_lines = [
            tr("Chunk: ") + f"{self.spin_chunk.value()}",
            tr("Max Freq: ") + f"{self.spin_freq.value()}",
            tr("Search Mode") + f": {self.comp_mode_combo.currentText()}",
            tr("Filter > ") + f"{self.spin_filter.value()}",
            tr("Lab Mode") + f": {'On' if self.btn_lab_mode_toggle_comp.isChecked() else 'Off'}",
        ]
        if self.btn_lab_mode_toggle_comp.isChecked():
            comp_settings_lines.append(tr("Deep Scan") + f": {'On' if self.chk_lab_deep_comp.isChecked() else 'Off'}")

        # Use shared sanitization utility for consistency with Web app
        sanitize_for_excel = shared_sanitize_excel

        def _clean_and_marker(text):
            """Prepares HTML for export: converts spans to *, removes other tags."""
            t = str(text or "")
            # 1. Convert Spans to Markers (Handle content inside)
            t = re.sub(r"<span[^>]*>(.*?)</span>", r"*\1*", t, flags=re.DOTALL)

            # 2. Remove newlines and BR
            t = t.replace("<br>", " ").replace("<br/>", " ").replace("\n", " ").replace("\r", " ")

            # 3. Remove any remaining HTML tags (div, etc)
            t = re.sub(r'<[^>]+>', '', t)

            # 4. Collapse multiple spaces
            t = re.sub(r'\s+', ' ', t)

            # 5. Merge adjacent asterisks
            t = re.sub(r'\*(\s+)\*', r'\1', t)

            return t.strip()

        # ==========================================
        #  XLSX & CSV Logic
        # ==========================================
        appendix_count = sum(len(v) for v in c_appx.values())
        filtered_total = len(c_filt) + sum(len(v) for v in c_filt_appx.values())
        known_count = len(c_known)
        total_count = len(c_main) + appendix_count + known_count + filtered_total

        if fmt in ['xlsx', 'csv', 'docx']:
            table_rows = []

            def add_rows(items, category, group_name=""):
                for ms_item in items:
                    item_type = ms_item.get('type', '')
                    if item_type == 'part':
                        # For Parts, use Part display name and Oxford title
                        part_display = ms_item.get('part_display', '')
                        oxford_title = ms_item.get('oxford_title', '')
                        sid = ms_item.get('sys_id', '')
                        shelf = f"📖 {part_display}" if part_display else sid
                        title = oxford_title or ""
                        ms_score = ms_item.get('score', 0)
                        # Get library info
                        library_code = self.meta_mgr.get_library_for_id(sid)
                        library_display = get_library_display(library_code, short=False) if library_code else ""

                        for page in ms_item.get('pages', []):
                             p_sid, p_num, p_shelf, _ = self._get_meta_for_header(page['raw_header'])

                             src_clean = _clean_and_marker(page.get('source_ctx', ''))
                             ms_clean = _clean_and_marker(page.get('text', ''))

                             # Show both Part and folio info
                             display_shelf = f"{shelf} [{p_shelf}]" if p_shelf else shelf

                             table_rows.append([
                                category,
                                group_name,
                                p_sid or sid or "",
                                library_display,
                                display_shelf or "",
                                title or "",
                                str(p_num or ""),
                                f"{ms_score} (P:{page.get('score',0)})",
                                src_clean,
                                ms_clean
                             ])
                    elif item_type == 'manuscript':
                        sid = ms_item['sys_id']
                        shelf, title = self.meta_mgr.get_meta_for_id(sid)
                        if not shelf or shelf == "Unknown":
                             shelf = self.meta_mgr.get_shelfmark_from_header(ms_item.get('raw_header', ''))
                        # Get library info
                        library_code = self.meta_mgr.get_library_for_id(sid)
                        library_display = get_library_display(library_code, short=False) if library_code else ""

                        ms_score = ms_item.get('score', 0)

                        for page in ms_item.get('pages', []):
                             _, p_num, _, _ = self._get_meta_for_header(page['raw_header'])

                             src_clean = _clean_and_marker(page.get('source_ctx', ''))
                             ms_clean = _clean_and_marker(page.get('text', ''))

                             table_rows.append([
                                category,
                                group_name,
                                sid or "",
                                library_display,
                                shelf or "",
                                title or "",
                                str(p_num or ""),
                                f"{ms_score} (P:{page.get('score',0)})",
                                src_clean,
                                ms_clean
                             ])
                    else:
                        # Fallback
                        sid, p_num, shelf, title = self._get_meta_for_header(ms_item.get('raw_header', ''))
                        # Get library info
                        library_code = self.meta_mgr.get_library_for_id(sid)
                        library_display = get_library_display(library_code, short=False) if library_code else ""
                        src_clean = _clean_and_marker(ms_item.get('source_ctx', ''))
                        ms_clean = _clean_and_marker(ms_item.get('text', ''))

                        table_rows.append([
                            category,
                            group_name,
                            sid or "",
                            library_display,
                            shelf or "",
                            title or "",
                            str(p_num or ""),
                            str(ms_item.get('score', 0)),
                            src_clean,
                            ms_clean
                        ])

            if self.chk_comp_flat.isChecked():
                all_items = self._collect_comp_items(
                    c_main, c_appx,
                    c_filt, c_filt_appx,
                    c_known
                )
                flat_items = self._sort_comp_items(all_items)
                add_rows(flat_items, tr("All Results"))
            else:
                add_rows(c_main, tr("Main Manuscripts"))
                for sig, items in sorted(c_appx.items(), key=lambda x: len(x[1]), reverse=True):
                    add_rows(items, tr("Appendix"), sig)
                add_rows(c_filt, tr("Filtered Main"))
                for sig, items in sorted(c_filt_appx.items(), key=lambda x: len(x[1]), reverse=True):
                    add_rows(items, tr("Filtered Appendix"), sig)
                add_rows(c_known, tr("Excluded Manuscripts"))

            # --- XLSX ---
            if fmt == 'xlsx':
                try:
                    wb = openpyxl.Workbook()
                    ws_report = wb.active
                    ws_report.title = tr("Report View")
                    ws_report.sheet_view.rightToLeft = True

                    ws_raw = wb.create_sheet(tr("Raw Data"))
                    ws_raw.sheet_view.rightToLeft = True
                    ws_query = wb.create_sheet(tr("Query information"))
                    ws_query.sheet_view.rightToLeft = True

                    font_red = InlineFont(color='FF0000', b=True)
                    font_normal = InlineFont(color='000000', b=False)
                    header_fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
                    header_font = Font(bold=True)

                    def write_rich_cell(ws, row, col, text):
                        safe_text = sanitize_for_excel(text)

                        if '*' not in safe_text:
                            ws.cell(row=row, column=col, value=safe_text)
                            return

                        parts = safe_text.split('*')
                        rich_string = CellRichText()

                        for i, part in enumerate(parts):
                            if not part:
                                continue
                            if i % 2 == 1:
                                rich_string.append(TextBlock(font_red, part))
                            else:
                                rich_string.append(TextBlock(font_normal, part))

                        ws.cell(row=row, column=col, value=rich_string)

                    def _write_credit_block(ws, start_row):
                        curr = start_row
                        for line in credit_text.split('\n'):
                            clean_line = line.strip()
                            if not clean_line or "====" in clean_line:
                                continue
                            c = ws.cell(row=curr, column=1, value=sanitize_for_excel(line))
                            c.font = Font(bold=True, color="555555")
                            curr += 1
                        return curr + 1

                    def _write_headers(ws, row_idx, headers):
                        for idx, h in enumerate(headers, 1):
                            c = ws.cell(row=row_idx, column=idx, value=h)
                            c.font = Font(bold=True, color="FFFFFF")
                            c.fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
                        return row_idx + 1

                    def _write_report_header(ws, row_idx, label):
                        c = ws.cell(row=row_idx, column=1, value=label)
                        c.font = header_font
                        c.fill = header_fill
                        return row_idx + 1

                    def _write_report_row(ws, row_idx, row_data):
                        for idx, val in enumerate(row_data, 1):
                            val_str = str(val)
                            if idx in (9, 10):  # Source Context and Manuscript Text columns (shifted due to Library)
                                write_rich_cell(ws, row_idx, idx, val_str)
                            else:
                                ws.cell(row=row_idx, column=idx, value=sanitize_for_excel(val_str))
                        return row_idx + 1

                    headers = [
                        tr("Category"),
                        tr("Group"),
                        tr("System ID"),
                        tr("Library"),
                        tr("Shelfmark"),
                        tr("Title"),
                        tr("Image"),
                        tr("Score"),
                        tr("Source Context"),
                        tr("Manuscript Text"),
                    ]

                    # Raw Data sheet (flat)
                    raw_row = _write_credit_block(ws_raw, 1)
                    header_row = raw_row
                    raw_row = _write_headers(ws_raw, raw_row, headers)
                    for row_data in table_rows:
                        for idx, val in enumerate(row_data, 1):
                            val_str = str(val)
                            if idx in (9, 10):  # Source Context and Manuscript Text columns
                                write_rich_cell(ws_raw, raw_row, idx, val_str)
                            else:
                                ws_raw.cell(row=raw_row, column=idx, value=sanitize_for_excel(val_str))
                        raw_row += 1
                    if table_rows:
                        ws_raw.auto_filter.ref = f"A{header_row}:J{raw_row - 1}"

                    # Query Information sheet
                    ws_query.column_dimensions['A'].width = 22
                    ws_query.column_dimensions['B'].width = 120
                    ws_query.append([tr("Search Text"), query_text])
                    ws_query.append([])
                    ws_query.append([tr("Search Settings"), ""])
                    for line in comp_settings_lines:
                        label, _, value = line.partition(":")
                        ws_query.append([label.strip(), value.strip()])
                    for row in ws_query.iter_rows(min_row=1, max_row=ws_query.max_row, min_col=1, max_col=2):
                        for cell in row:
                            cell.alignment = Alignment(wrap_text=True, vertical="top")
                            if cell.column == 1 and cell.row in (1, 3):
                                cell.font = Font(bold=True)

                    # Report View sheet (hierarchical)
                    report_row = _write_credit_block(ws_report, 1)
                    report_row = _write_headers(ws_report, report_row, headers)

                    def add_report_section(category_label, items, groups=None, outline_level=1):
                        nonlocal report_row
                        if not items and not groups:
                            return
                        section_start = report_row
                        report_row = _write_report_header(ws_report, report_row, category_label)

                        def add_items_with_group(group_name, group_items, group_level):
                            nonlocal report_row
                            if group_name:
                                group_start = report_row
                                report_row = _write_report_header(ws_report, report_row, group_name)
                            ms_ranges = []
                            for ms_item in group_items:
                                ms_item_rows = []
                                item_type = ms_item.get('type', '')
                                if item_type == 'part':
                                    part_display = ms_item.get('part_display', '')
                                    oxford_title = ms_item.get('oxford_title', '')
                                    sid = ms_item.get('sys_id', '')
                                    shelf = f"📖 {part_display}" if part_display else sid
                                    title = oxford_title or ""
                                    ms_score = ms_item.get('score', 0)
                                    # Get library info
                                    library_code = self.meta_mgr.get_library_for_id(sid)
                                    library_display = get_library_display(library_code, short=False) if library_code else ""

                                    for page in ms_item.get('pages', []):
                                        p_sid, p_num, p_shelf, _ = self._get_meta_for_header(page['raw_header'])
                                        src_clean = _clean_and_marker(page.get('source_ctx', ''))
                                        ms_clean = _clean_and_marker(page.get('text', ''))
                                        display_shelf = f"{shelf} [{p_shelf}]" if p_shelf else shelf
                                        ms_item_rows.append([
                                            "",
                                            "",
                                            p_sid or sid or "",
                                            display_shelf or "",
                                            library_display,
                                            title or "",
                                            str(p_num or ""),
                                            f"{ms_score} (P:{page.get('score',0)})",
                                            src_clean,
                                            ms_clean
                                        ])
                                elif item_type == 'manuscript':
                                    sid = ms_item['sys_id']
                                    shelf, title = self.meta_mgr.get_meta_for_id(sid)
                                    if not shelf or shelf == "Unknown":
                                        shelf = self.meta_mgr.get_shelfmark_from_header(ms_item.get('raw_header', ''))
                                    # Get library info
                                    library_code = self.meta_mgr.get_library_for_id(sid)
                                    library_display = get_library_display(library_code, short=False) if library_code else ""

                                    ms_score = ms_item.get('score', 0)

                                    for page in ms_item.get('pages', []):
                                        _, p_num, _, _ = self._get_meta_for_header(page['raw_header'])
                                        src_clean = _clean_and_marker(page.get('source_ctx', ''))
                                        ms_clean = _clean_and_marker(page.get('text', ''))
                                        ms_item_rows.append([
                                            "",
                                            "",
                                            sid or "",
                                            shelf or "",
                                            library_display,
                                            title or "",
                                            str(p_num or ""),
                                            f"{ms_score} (P:{page.get('score',0)})",
                                            src_clean,
                                            ms_clean
                                        ])
                                else:
                                    sid, p_num, shelf, title = self._get_meta_for_header(ms_item.get('raw_header', ''))
                                    # Get library info
                                    library_code = self.meta_mgr.get_library_for_id(sid)
                                    library_display = get_library_display(library_code, short=False) if library_code else ""
                                    src_clean = _clean_and_marker(ms_item.get('source_ctx', ''))
                                    ms_clean = _clean_and_marker(ms_item.get('text', ''))
                                    ms_item_rows.append([
                                        "",
                                        "",
                                        sid or "",
                                        shelf or "",
                                        library_display,
                                        title or "",
                                        str(p_num or ""),
                                        str(ms_item.get('score', 0)),
                                        src_clean,
                                        ms_clean
                                    ])

                                if ms_item_rows:
                                    ms_start = report_row
                                    report_row = _write_report_header(ws_report, report_row, f"{ms_item_rows[0][3]} | {ms_item_rows[0][5]}")
                                    for row_data in ms_item_rows:
                                        report_row = _write_report_row(ws_report, report_row, row_data)
                                    ms_end = report_row - 1
                                    if ms_end >= ms_start + 1:
                                        ws_report.row_dimensions.group(ms_start + 1, ms_end, outline_level=group_level + 1, hidden=False)
                                    ms_ranges.append((ms_start, ms_end))

                            if group_name and ms_ranges:
                                group_end = report_row - 1
                                if group_end >= group_start + 1:
                                    ws_report.row_dimensions.group(group_start + 1, group_end, outline_level=group_level, hidden=False)

                        if groups:
                            for sig, group_items in groups:
                                add_items_with_group(f"{sig} ({len(group_items)})", group_items, outline_level + 1)
                        if items:
                            add_items_with_group("", items, outline_level + 1)

                        section_end = report_row - 1
                        if section_end >= section_start + 1:
                            ws_report.row_dimensions.group(section_start + 1, section_end, outline_level=outline_level, hidden=False)

                    if self.chk_comp_flat.isChecked():
                        add_report_section(tr("All Results"), self._sort_comp_items(all_items), outline_level=1)
                    else:
                        add_report_section(tr("Main Results"), self._sort_comp_items(c_main), outline_level=1)
                        add_report_section(tr("Appendix - Grouped"), [], groups=sorted(c_appx.items(), key=lambda x: len(x[1]), reverse=True), outline_level=1)
                        add_report_section(tr("Filtered"), self._sort_comp_items(c_filt), groups=sorted(c_filt_appx.items(), key=lambda x: len(x[1]), reverse=True), outline_level=1)
                        add_report_section(tr("Excluded"), self._sort_comp_items(c_known), outline_level=1)

                    ws_report.sheet_properties.outlinePr.summaryBelow = True

                    dims = {'D': 20, 'E': 18, 'F': 30, 'I': 50, 'J': 60}  # Columns shifted for Library
                    for col, width in dims.items():
                        ws_report.column_dimensions[col].width = width
                        ws_raw.column_dimensions[col].width = width

                    if c_known:
                        ws_excluded = wb.create_sheet(tr("Excluded Manuscripts"))
                        ws_excluded.sheet_view.rightToLeft = True
                        ws_excluded.append([tr("System ID"), tr("Shelfmark"), tr("Title")])
                        for cell in ws_excluded[1]:
                            cell.font = Font(bold=True, color="FFFFFF")
                            cell.fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")

                        def _excluded_row(item):
                            item_type = item.get('type', '')
                            if item_type == 'part':
                                sid = item.get('sys_id', '')
                                shelf = item.get('part_display', '')
                                title = item.get('oxford_title', '')
                                return [sid, shelf, title]
                            if item_type == 'manuscript':
                                sid = item.get('sys_id', '')
                                shelf, title = self.meta_mgr.get_meta_for_id(sid)
                                if not shelf or shelf == "Unknown":
                                    shelf = self.meta_mgr.get_shelfmark_from_header(item.get('raw_header', ''))
                                return [sid, shelf or "", title or ""]
                            sid, _, shelf, title = self._get_meta_for_header(item.get('raw_header', ''))
                            return [sid or "", shelf or "", title or ""]

                        for item in self._sort_comp_items(c_known):
                            ws_excluded.append(_excluded_row(item))

                    wb.save(path)
                    self._save_last_folder(path)
                    QMessageBox.information(self, tr("Saved"), tr("Saved to {}").format(path))
                except Exception as e:
                    QMessageBox.critical(self, tr("Error"), f"Failed to save XLSX:\n{e}")

            # --- CSV ---
            elif fmt == 'csv':
                try:
                    headers = [
                        tr("Category"),
                        tr("Group"),
                        tr("System ID"),
                        tr("Library"),
                        tr("Shelfmark"),
                        tr("Title"),
                        tr("Image"),
                        tr("Score"),
                        tr("Source Context"),
                        tr("Manuscript Text"),
                    ]
                    with open(path, 'w', encoding='utf-8-sig', newline='') as f:
                        f.write(credit_text)
                        writer = csv.writer(f)
                        writer.writerow([])
                        writer.writerow(headers)
                        for row in table_rows:
                            clean_row = [str(val).replace('*', '') for val in row]
                            writer.writerow(clean_row)
                    self._save_last_folder(path)
                    QMessageBox.information(self, tr("Saved"), tr("Saved to {}").format(path))
                except Exception as e:
                    QMessageBox.critical(self, tr("Error"), f"Failed to save CSV:\n{e}")

            # --- DOCX ---
            elif fmt == 'docx':
                try:
                    doc = Document()
                    section = doc.sections[-1]
                    new_width, new_height = section.page_height, section.page_width
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = new_width
                    section.page_height = new_height

                    for line in credit_text.split('\n'):
                        clean_line = line.strip()
                        if not clean_line or "====" in clean_line:
                            continue
                        p = doc.add_paragraph(clean_line)
                        if p.runs:
                            p.runs[0].font.bold = True
                    doc.add_paragraph("")
                    doc.add_paragraph(tr("Report Summary"))
                    stats_lines = [
                        f"{tr('Total Manuscripts Found')}: {total_count}",
                        f"{tr('Main Manuscripts')}: {len(c_main)}",
                        f"{tr('Main Appendix (Groups)')}: {len(c_appx)}",
                        f"{tr('Filtered by Text (Manuscripts)')}: {filtered_total}",
                        f"{tr('Excluded Manuscripts')}: {known_count}",
                    ]
                    for line in stats_lines:
                        doc.add_paragraph(line)
                    doc.add_paragraph(tr("See Appendix for query information and excluded manuscripts."))
                    doc.add_paragraph("")

                    headers = [
                        tr("Category"),
                        tr("Group"),
                        tr("System ID"),
                        tr("Library"),
                        tr("Shelfmark"),
                        tr("Title"),
                        tr("Image"),
                        tr("Score"),
                        tr("Source Context"),
                        tr("Manuscript Text"),
                    ]
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.autofit = False
                    self._set_table_width_pct(table, 100)
                    hdr_cells = table.rows[0].cells
                    for idx, header in enumerate(headers):
                        hdr_cells[idx].text = header
                    available_width = section.page_width - section.left_margin - section.right_margin
                    ratios = [0.04, 0.05, 0.06, 0.08, 0.06, 0.08, 0.05, 0.04, 0.27, 0.27]  # Adjusted for Library
                    for idx, ratio in enumerate(ratios):
                        width = int(available_width * ratio)
                        for cell in table.columns[idx].cells:
                            cell.width = width

                    for row in table_rows:
                        row_cells = table.add_row().cells
                        for col_idx, val in enumerate(row):
                            cell = row_cells[col_idx]
                            if col_idx in (8, 9):  # Source Context and Manuscript Text columns
                                cell.text = ""
                                self._add_docx_highlighted_runs(cell.paragraphs[0], val)
                            else:
                                cell.text = str(val).replace('*', '')

                    doc.add_page_break()
                    doc.add_heading(tr("Appendix"), level=1)
                    doc.add_heading(tr("Query information"), level=2)
                    doc.add_paragraph(tr("Search Text") + f": {query_text}")
                    doc.add_paragraph("")
                    doc.add_paragraph(tr("Search Settings"))
                    for line in comp_settings_lines:
                        doc.add_paragraph(line)

                    doc.add_paragraph("")
                    doc.add_heading(tr("Excluded Manuscripts"), level=2)
                    excluded_table = None
                    if c_known:
                        excluded_table = doc.add_table(rows=1, cols=3)
                        excluded_table.autofit = False
                        self._set_table_width_pct(excluded_table, 100)
                        excluded_hdr = excluded_table.rows[0].cells
                        for idx, header in enumerate([tr("System ID"), tr("Shelfmark"), tr("Title")]):
                            excluded_hdr[idx].text = header
                        excluded_widths = [0.2, 0.3, 0.5]
                        for idx, ratio in enumerate(excluded_widths):
                            width = int(available_width * ratio)
                            for cell in excluded_table.columns[idx].cells:
                                cell.width = width
                        for item in self._sort_comp_items(c_known):
                            item_type = item.get('type', '')
                            if item_type == 'part':
                                sid = item.get('sys_id', '')
                                shelf = item.get('part_display', '')
                                title = item.get('oxford_title', '')
                            elif item_type == 'manuscript':
                                sid = item.get('sys_id', '')
                                shelf, title = self.meta_mgr.get_meta_for_id(sid)
                                if not shelf or shelf == "Unknown":
                                    shelf = self.meta_mgr.get_shelfmark_from_header(item.get('raw_header', ''))
                            else:
                                sid, _, shelf, title = self._get_meta_for_header(item.get('raw_header', ''))
                            row_cells = excluded_table.add_row().cells
                            row_cells[0].text = str(sid or "")
                            row_cells[1].text = str(shelf or "")
                            row_cells[2].text = str(title or "")
                    else:
                        doc.add_paragraph(tr("None"))

                    if CURRENT_LANG == "he":
                        doc.styles["Normal"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        for p in doc.paragraphs:
                            self._set_paragraph_rtl(p)
                        self._set_table_rtl(table)
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    self._set_paragraph_rtl(p)
                        if excluded_table is not None:
                            self._set_table_rtl(excluded_table)
                            for row in excluded_table.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        self._set_paragraph_rtl(p)

                    doc.save(path)
                    self._save_last_folder(path)
                    QMessageBox.information(self, tr("Saved"), tr("Saved to {}").format(path))
                except Exception as e:
                    QMessageBox.critical(self, tr("Error"), f"Failed to save DOCX:\n{e}")

        # --- TXT ---
        else:
            try:
                sep = "=" * 80
                appendix_count = sum(len(v) for v in c_appx.values())
                filtered_total = len(c_filt) + sum(len(v) for v in c_filt_appx.values())
                known_count = len(c_known)
                total_count = len(c_main) + appendix_count + known_count + filtered_total

                def _fmt_ms_entry(ms_item):
                    item_type = ms_item.get('type', '')
                    if item_type == 'part':
                        part_display = ms_item.get('part_display', '')
                        oxford_title = ms_item.get('oxford_title', '')
                        sid = ms_item.get('sys_id', '')
                        ms_block = [sep, f"📖 PART: {part_display} | {oxford_title} (ID: {sid}) | Total Score: {ms_item.get('score', 0)}", sep]

                        for page in ms_item.get('pages', []):
                             p_sid, p_num, p_shelf, _ = self._get_meta_for_header(page['raw_header'])
                             src_clean = _clean_and_marker(page.get('source_ctx', ''))
                             ms_clean = _clean_and_marker(page.get('text', ''))
                             folio_info = f" [{p_shelf}]" if p_shelf else ""
                             ms_block.append(f"\n--- Page {p_num}{folio_info} (Score: {page.get('score',0)}) ---")
                             ms_block.append(tr("Source Context") + ":\n" + src_clean)
                             ms_block.append(tr("Manuscript") + ":\n" + ms_clean)
                        return ms_block
                    elif item_type == 'manuscript':
                        sid = ms_item['sys_id']
                        shelf, title = self.meta_mgr.get_meta_for_id(sid)
                        if not shelf or shelf == "Unknown":
                            shelf = self.meta_mgr.get_shelfmark_from_header(ms_item.get('raw_header', ''))

                        ms_block = [sep, f"MANUSCRIPT: {shelf} | {title} (ID: {sid}) | Total Score: {ms_item.get('score', 0)}", sep]

                        for page in ms_item.get('pages', []):
                             _, p_num, _, _ = self._get_meta_for_header(page['raw_header'])
                             src_clean = _clean_and_marker(page.get('source_ctx', ''))
                             ms_clean = _clean_and_marker(page.get('text', ''))
                             ms_block.append(f"\n--- Page {p_num} (Score: {page.get('score',0)}) ---")
                             ms_block.append(tr("Source Context") + ":\n" + src_clean)
                             ms_block.append(tr("Manuscript") + ":\n" + ms_clean)
                        return ms_block
                    else:
                        return self._fmt_item_legacy(ms_item)

                summary_lines = [
                    sep, tr("COMPOSITION REPORT SUMMARY"), sep,
                    f"Title: {comp_title}",
                    f"{tr('Total Manuscripts Found')}: {total_count}"
                ]
                detail_lines = [sep, tr("ALL RESULTS"), sep]

                if self.chk_comp_flat.isChecked():
                    flat_items = self._sort_comp_items(
                        self._collect_comp_items(c_main, c_appx, c_filt, c_filt_appx, c_known)
                    )
                    for item in flat_items: detail_lines.extend(_fmt_ms_entry(item))
                else:
                    summary_lines.extend([
                        f"{tr('Main Manuscripts')}: {len(c_main)}",
                        f"{tr('Main Appendix (Groups)')}: {len(c_appx)}",
                        f"{tr('Filtered by Text (Manuscripts)')}: {filtered_total}",
                        f"{tr('Excluded Manuscripts')}: {known_count}"
                    ])
                    detail_lines = [sep, tr("MAIN MANUSCRIPTS"), sep]
                    for item in c_main: detail_lines.extend(_fmt_ms_entry(item))
                    if c_appx:
                        detail_lines.extend([sep, tr("MAIN APPENDIX") + " (Grouped)", sep])
                        for sig, items in sorted(c_appx.items(), key=lambda x: len(x[1]), reverse=True):
                            detail_lines.append(f"=== GROUP: {sig} ({len(items)} items) ===")
                            for item in items: detail_lines.extend(_fmt_ms_entry(item))
                    if c_filt:
                        detail_lines.extend([sep, tr("FILTERED"), sep])
                        for item in c_filt: detail_lines.extend(_fmt_ms_entry(item))
                    if c_filt_appx:
                        detail_lines.extend([sep, tr("FILTERED APPENDIX") + " (Grouped)", sep])
                        for sig, items in sorted(c_filt_appx.items(), key=lambda x: len(x[1]), reverse=True):
                            detail_lines.append(f"=== GROUP: {sig} ({len(items)} items) ===")
                            for item in items: detail_lines.extend(_fmt_ms_entry(item))
                    if c_known:
                        detail_lines.extend([sep, tr("EXCLUDED MANUSCRIPTS"), sep])
                        for item in c_known: detail_lines.extend(_fmt_ms_entry(item))

                with open(path, 'w', encoding='utf-8') as f:
                    f.write(credit_text)
                    all_lines = summary_lines + detail_lines
                    f.write("\n".join(all_lines).strip() + "\n")
                self._save_last_folder(path)
                QMessageBox.information(self, tr("Saved"), tr("Saved to {}").format(path))

            except Exception as e:
                QMessageBox.critical(self, tr("Error"), f"Failed to save TXT:\n{e}")

    # Composition & Browse
    def open_filter_dialog(self, col_idx=None):
        # If called from CheckBoxHeader, col_idx is passed
        # If called from Composition (filter text), no arg is passed (or handled differently)
        # Note: button.clicked signal passes bool (checked state), so we need to filter that out

        # Check if this is the Search Tab's results table callback (must be int but not bool)
        if isinstance(col_idx, int) and not isinstance(col_idx, bool):
            if col_idx == 1:
                self.open_list_filter_dialog()
                return

            # Standard text filter for other columns
            current_filter = self.results_filters.get(col_idx, {})
            dlg = ColumnFilterDialog(self, col_index=col_idx, current_filter=current_filter)
            if dlg.exec():
                new_filter = dlg.get_filter_data()
                if new_filter['active']:
                    self.results_filters[col_idx] = new_filter
                else:
                    self.results_filters.pop(col_idx, None)

                # Update header visual state
                header = self.results_table.horizontalHeader()
                if hasattr(header, 'set_filter_active'):
                    header.set_filter_active(col_idx, new_filter['active'])

                self._apply_results_table_filters()
            return

        # Legacy/Composition text filter
        dlg = FilterTextDialog(self, current_sources=self.filter_sources)
        # Restore enabled sources state
        if self.filter_enabled_sources:
            dlg.enabled_sources = self.filter_enabled_sources.copy()
            dlg._refresh_sources_list()
        if dlg.exec():
            self.filter_sources = dlg.get_sources()
            self.filter_enabled_sources = dlg.get_enabled_sources()

    def _get_filter_text(self):
        """Get combined filter text from enabled sources."""
        if not self.filter_sources or not self.filter_enabled_sources:
            return ""
        texts = [self.filter_sources[ref] for ref in self.filter_enabled_sources
                 if ref in self.filter_sources]
        return " ".join(texts)

    def _update_list_filter_cache(self):
        """Cache the set of system IDs for the currently selected lists to optimize filtering."""
        target_lists = self.list_filter_state.get('lists', 'all')

        target_sys_ids = set()
        if self.lists_mgr:
            lists_to_check = []
            if target_lists == "all":
                all_lists = self.lists_mgr.get_all_lists(include_recent=False)
                lists_to_check = [l['id'] for l in all_lists]
            else:
                lists_to_check = target_lists

            for lst_id in lists_to_check:
                items = self.lists_mgr.get_items_in_list(lst_id)
                for item in items:
                    sid = item.get('sys_id')
                    if sid: target_sys_ids.add(str(sid))

        self.list_filter_state['cached_ids'] = target_sys_ids

    def open_list_filter_dialog(self):
        dlg = ListFilterDialog(self, self.lists_mgr, current_state=self.list_filter_state)
        if dlg.exec():
            mode, selected_lists = dlg.get_selection()
            self.list_filter_state['mode'] = mode
            self.list_filter_state['lists'] = selected_lists
            self.list_filter_state['active'] = True # Auto-activate on OK

            self._update_list_filter_cache()

            # Update header
            header = self.results_table.horizontalHeader()
            if hasattr(header, 'set_filter_active'):
                 header.set_filter_active(1, True)
                 header.set_star_active(1, True)

            self._apply_results_table_filters()

    def toggle_list_filter(self, col_idx=1):
        if col_idx != 1: return

        # Toggle active state
        self.list_filter_state['active'] = not self.list_filter_state['active']
        is_active = self.list_filter_state['active']

        # Ensure cache is populated if activating for the first time
        if is_active and 'cached_ids' not in self.list_filter_state:
            self._update_list_filter_cache()

        # Update header
        header = self.results_table.horizontalHeader()
        if hasattr(header, 'set_filter_active'):
             header.set_star_active(1, is_active)
             header.set_filter_active(1, is_active)

        self._apply_results_table_filters()

    def load_comp_file(self):
        path, _ = QFileDialog.getOpenFileName(self, tr("Load"), "", "Text (*.txt)")
        if path:
            with open(path, 'r', encoding='utf-8') as f: self.comp_text_area.setPlainText(f.read())

    def open_exclude_dialog(self):
        dlg = ExcludeDialog(self, existing_entries=self.excluded_raw_entries)
        if dlg.exec():
            self.set_excluded_entries(dlg.get_entries_text())

    def set_excluded_entries(self, entries_text: str):
        entries = [e.strip() for e in entries_text.splitlines() if e.strip()]
        self.excluded_raw_entries = entries

        sys_ids = set()
        shelves = set()
        for e in entries:
            cleaned = re.sub(r"\s+", "", e)
            digits_only = re.sub(r"\D", "", cleaned)
            if digits_only and digits_only == cleaned:
                sys_ids.add(cleaned)
            else:
                norm = self._normalize_shelfmark(e)
                if norm:
                    shelves.add(norm)

        self.excluded_sys_ids = sys_ids
        self.excluded_shelfmarks = shelves
        self.lbl_exclude_status.setText(tr("Excluded: {}").format(len(entries)))

    def _normalize_shelfmark(self, shelfmark: str) -> str:
        """Normalize shelfmarks using the canonical function from genizah_core."""
        return normalize_shelfmark(shelfmark)

    def _ensure_shelf_map(self):
        """Build a mapping from normalized shelfmark to sys_id for quick lookups."""
        if self._shelf_to_sys is not None:
            return
        self._shelf_to_sys = {}
        if not self.meta_mgr:
            return
        # Build mapping from csv_bank
        for sys_id, meta in self.meta_mgr.csv_bank.items():
            shelf = meta.get("shelfmark")
            if shelf:
                norm = self._normalize_shelfmark(shelf)
                if norm and norm not in self._shelf_to_sys:
                    self._shelf_to_sys[norm] = sys_id
        # Also add from nli_cache
        for sys_id, meta in self.meta_mgr.nli_cache.items():
            shelf = meta.get("shelfmark")
            if shelf:
                norm = self._normalize_shelfmark(shelf)
                if norm and norm not in self._shelf_to_sys:
                    self._shelf_to_sys[norm] = sys_id

    def _get_meta_for_header(self, raw_header):
        """Return (sys_id, p_num, shelfmark, title) preferring metadata bank for shelfmarks."""
        sys_id, p_num = self.meta_mgr.parse_header_smart(raw_header)

        shelf = "Unknown"
        title = ""

        if sys_id:
            # Use the new unified lookup
            shelf, title = self.meta_mgr.get_meta_for_id(sys_id)

        # Fallback to header parsing if CSV/Cache failed
        if not shelf or shelf == "Unknown":
            shelf = self.meta_mgr.get_shelfmark_from_header(raw_header) or "Unknown"

        return sys_id, p_num, shelf, title

    def _item_matches_exclusion(self, item):
        # For Part items, check all folios in the Part
        item_type = item.get('type', '')
        if item_type == 'part':
            # Check sys_id field directly
            direct_sid = item.get('sys_id')
            if direct_sid and direct_sid in self.excluded_sys_ids:
                return True
            # Check all folios in the Part
            for folio_sid in item.get('folios', []):
                if folio_sid in self.excluded_sys_ids:
                    return True

        sys_id, _ = self.meta_mgr.parse_header_smart(item.get('raw_header', ''))
        if sys_id and sys_id in self.excluded_sys_ids:
            return True

        if sys_id and sys_id not in self.meta_mgr.nli_cache:
            self.meta_mgr.fetch_nli_data(sys_id)

        _, _, shelf, _ = self._get_meta_for_header(item.get('raw_header', ''))
        norm_shelf = self._normalize_shelfmark(shelf)
        if norm_shelf and norm_shelf in self.excluded_shelfmarks:
            return True
        return False

    def _apply_manual_exclusions(self, main, appx):
        if not (self.excluded_sys_ids or self.excluded_shelfmarks):
            return main, appx, []

        known = []
        filtered_main = []
        for item in main:
            if self._item_matches_exclusion(item):
                known.append(item)
            else:
                filtered_main.append(item)

        filtered_appx = {}
        for key, items in appx.items():
            kept = []
            for item in items:
                if self._item_matches_exclusion(item):
                    known.append(item)
                else:
                    kept.append(item)
            if kept:
                filtered_appx[key] = kept

        return filtered_main, filtered_appx, known

    def toggle_composition(self):
        if self.is_comp_running:
            if getattr(self, 'group_thread', None) and self.group_thread.isRunning():
                # Disconnect signals to prevent race conditions during stop
                try: self.group_thread.finished_signal.disconnect()
                except: pass
                try: self.group_thread.error_signal.disconnect()
                except: pass

                self.group_thread.requestInterruption()
                self.group_thread.wait()

                QMessageBox.information(self, tr("Stopped"), tr("Grouping stopped. Showing ungrouped results."))
                # Pass explicit empty dicts for other arguments to avoid crashes
                self.display_comp_results(self.comp_raw_items or [], {}, {}, self.comp_raw_filtered or [], {}, {})
            elif getattr(self, 'comp_thread', None) and self.comp_thread.isRunning():
                self.comp_thread.terminate()
                self.comp_thread.wait()
            self.is_comp_running = False
            self.reset_comp_ui()
        else:
            self.run_composition()
        
    def reset_comp_ui(self):
        self.is_comp_running = False; self.btn_comp_run.setText(tr("Analyze Composition"))
        self.btn_comp_run.setStyleSheet("background-color: #2980b9; color: white;")
        self.comp_progress.setVisible(False)

    def run_composition(self, custom_text=None):
        """
        Main entry point for Composition Search.
        FINAL VERIFIED VERSION.
        """
        txt = (custom_text if custom_text is not None else self.comp_text_area.toPlainText()).strip()
        if not txt:
            QMessageBox.warning(self, tr("Error"), tr("Please enter text to search."))
            return

        # Auto-fill title with first 4 words if empty
        if not self.comp_title_input.text().strip():
            words = txt.split()[:4]
            auto_title = " ".join(words)
            if len(txt.split()) > 4:
                auto_title += "..."
            self.comp_title_input.setText(auto_title)

        self.is_comp_running = True
        self.btn_comp_run.setText(tr("Stop"))
        self.btn_comp_run.setStyleSheet("background-color: #c0392b; color: white;")
        
        if hasattr(self, 'btn_comp_recursive'):
            self.btn_comp_recursive.setEnabled(False)
            
        self.comp_progress.setVisible(True)
        self.comp_progress.setRange(0, 0)
        self.comp_progress.setValue(0)
        self.comp_tree.clear()
        self.comp_progress.setFormat(tr("Scanning chunks..."))
        
        # איפוס נתונים
        self.comp_raw_items = []
        self.comp_filtered = []
        self.comp_known = []
        
        for b in self.comp_export_buttons: b.setEnabled(False)

        chunk_size = self.spin_chunk.value()
        
        # מיפוי מצב חיפוש (simplified to single Variants mode with slider)
        available_modes = ['literal', 'variants', 'fuzzy']
        idx = self.comp_mode_combo.currentIndex()
        if 0 <= idx < len(available_modes):
            mode = available_modes[idx]
        else:
            mode = 'variants'

        # Update variant level from slider before search
        if mode == 'variants' and hasattr(self, 'comp_variant_slider') and self.var_mgr:
            self.var_mgr.set_variant_level(self.comp_variant_slider.value())

        excluded_ids = self.excluded_raw_entries

        # Get boundary search parameters from UI
        boundary_mode = self.boundary_mode_combo.currentData() if hasattr(self, 'boundary_mode_combo') else 'full'
        boundary_delimiter = self.boundary_delimiter_combo.currentData() if hasattr(self, 'boundary_delimiter_combo') else '\n'

        # Get advanced boundary settings from LabSettings or temporary values
        if hasattr(self, 'lab_engine') and self.lab_engine:
            boundary_boost = self.lab_engine.settings.boundary_boost
            min_boundary_matches = self.lab_engine.settings.min_boundary_matches
            min_delimiter_distance = self.lab_engine.settings.min_delimiter_distance
        else:
            # Use temporary values if set, otherwise defaults
            boundary_boost = getattr(self, '_boundary_boost_temp', 1.5)
            min_boundary_matches = getattr(self, '_min_boundary_matches_temp', 0)
            min_delimiter_distance = getattr(self, '_min_delimiter_distance_temp', 3)

        # 1. נתיב מעבדה (LAB MODE)
        if self.btn_lab_mode_toggle_comp.isChecked():
            if not self.lab_engine:
                QMessageBox.warning(self, tr("Error"), tr("Lab Engine not initialized."))
                self.reset_comp_ui()
                return

            # Robustness: Pass resolved System IDs if available, to catch items excluded by shelfmark
            # where the user didn't explicitly type the ID.
            final_excluded_ids = excluded_ids
            if self.excluded_sys_ids:
                final_excluded_ids = list(self.excluded_sys_ids)

            deep = self.chk_lab_deep_comp.isChecked()
            limit = self.lab_engine.settings.lab_scan_limit

            self.comp_thread = LabCompositionThread(
                self.lab_engine,
                txt,
                mode,
                chunk_size=chunk_size,
                excluded_ids=final_excluded_ids,
                filter_text=self._get_filter_text(),
                deep_scan=deep,
                scan_limit=limit,
                boundary_mode=boundary_mode,
                boundary_delimiter=boundary_delimiter,
                boundary_boost=boundary_boost,
                min_boundary_matches=min_boundary_matches,
                min_delimiter_distance=min_delimiter_distance
            )
            self.comp_thread.scan_finished_signal.connect(self.on_comp_scan_finished)

        # 2. נתיב רגיל (STANDARD MODE)
        else:
            if not self.searcher:
                QMessageBox.warning(self, tr("Error"), tr("Search engine not loaded."))
                self.reset_comp_ui()
                return

            # --- התיקון הקריטי כאן: הסרת progress_callback ---
            self.comp_thread = CompositionThread(
                self.searcher,
                txt,
                chunk=chunk_size,
                freq=self.spin_freq.value(),
                mode=mode,
                filter_text=self._get_filter_text(),
                threshold=self.spin_filter.value(),
                boundary_mode=boundary_mode,
                boundary_delimiter=boundary_delimiter,
                boundary_boost=boundary_boost,
                min_boundary_matches=min_boundary_matches,
                min_delimiter_distance=min_delimiter_distance
            )
            if hasattr(self.comp_thread, 'scan_finished_signal'):
                 self.comp_thread.scan_finished_signal.connect(self.on_comp_scan_finished)
            else:
                 self.comp_thread.finished_signal.connect(self.on_comp_search_finished)

        self.comp_thread.progress_signal.connect(self.on_comp_progress)
        
        if hasattr(self.comp_thread, 'status_signal'):
             self.comp_thread.status_signal.connect(self.on_comp_status_update)
             
        self.comp_thread.error_signal.connect(self.on_comp_error)
        
        self.comp_thread.start()

    def on_comp_display_mode_changed(self, _checked):
        if self.is_comp_running:
            return
        if not self._has_comp_results():
            return
        if not self.chk_comp_flat.isChecked() and not self.comp_has_grouped_results:
            if self.comp_raw_items or self.comp_raw_filtered:
                self.start_grouping(self.comp_raw_items or [], self.comp_raw_filtered or [])
                return
        if self.comp_grouped_main or self.comp_grouped_filtered_main:
            self.display_comp_results(
                self.comp_grouped_main,
                self.comp_grouped_appendix,
                self.comp_grouped_summary,
                self.comp_grouped_filtered_main,
                self.comp_grouped_filtered_appendix,
                self.comp_grouped_filtered_summary,
            )

    def run_recursive_composition(self):
        if self.is_comp_running:
            return
        base_text = self.comp_text_area.toPlainText().strip()
        if not base_text:
            return

        if not self._has_comp_results():
            self.pending_recursive_search = True
            self.run_composition()
            return

        selected_uids = self._collect_checked_comp_page_uids()
        if selected_uids:
            source_uids = selected_uids
        else:
            source_uids = self._collect_all_comp_page_uids()
            if not source_uids:
                QMessageBox.information(self, tr("No Results"), tr("No composition matches found."))
                return

        extra_texts = []
        for uid in source_uids:
            full_text = self.searcher.get_full_text_by_id(uid)
            if full_text:
                extra_texts.append(full_text)

        if not extra_texts:
            QMessageBox.information(self, tr("No Data"), tr("Could not load full text for selected results."))
            return

        combined_text = base_text + "\n\n" + "\n\n".join(extra_texts)
        self.run_composition(custom_text=combined_text)

    def on_comp_status_update(self, status):
        self.comp_progress.setFormat(status)
        if status.startswith("Scanning batch") or status.startswith("Scanning items"):
            self.comp_progress.setRange(0, 0)

    def on_comp_progress(self, curr, total):
        if total:
            self.comp_progress.setRange(0, total)
        else:
            self.comp_progress.setRange(0, 0)
        self.comp_progress.setValue(curr)

    def on_comp_error(self, err):
        """Handle errors during composition search."""
        self.reset_comp_ui()
        QMessageBox.critical(self, tr("Error"), str(err))
        
    def on_comp_scan_finished(self, result_obj):
        self.is_comp_running = False
        self.reset_comp_ui()

        known_raw = []
        if isinstance(result_obj, dict):
            items = result_obj.get('main', [])
            filtered_items = result_obj.get('filtered', [])
            known_raw = result_obj.get('known', []) 
        else:
            items = result_obj or []
            filtered_items = []

        manuscripts = self.searcher.group_pages_by_manuscript(items)
        filtered_manuscripts = self.searcher.group_pages_by_manuscript(filtered_items)
        known_manuscripts = self.searcher.group_pages_by_manuscript(known_raw)

        self.comp_raw_items = manuscripts
        self.comp_raw_filtered = filtered_manuscripts
        
        self.comp_known = known_manuscripts 

        if not manuscripts and not filtered_manuscripts and not known_manuscripts:
            QMessageBox.information(self, tr("No Results"), tr("No composition matches found."))
            self.pending_recursive_search = False
            return

        if self.chk_comp_flat.isChecked():
            self.comp_has_grouped_results = False
            self.comp_grouped_main = manuscripts
            self.comp_grouped_appendix = {}
            self.comp_grouped_summary = {}
            self.comp_grouped_filtered_main = filtered_manuscripts
            self.comp_grouped_filtered_appendix = {}
            self.comp_grouped_filtered_summary = {}
            self.display_comp_results(manuscripts, {}, {}, filtered_manuscripts, {}, {})
            return

        self.start_grouping(manuscripts, filtered_manuscripts)

    def start_grouping(self, items, filtered_items=None):
        self.is_comp_running = True
        self.comp_has_grouped_results = False
        self.btn_comp_run.setText(tr("Stop"))
        self.btn_comp_run.setStyleSheet("background-color: #c0392b; color: white;")
        self.comp_progress.setVisible(True)
        total_items = (len(items) if items else 0) + (len(filtered_items) if filtered_items else 0)
        self.comp_progress.setRange(0, total_items)
        self.comp_progress.setValue(0)
        self.comp_progress.setFormat(tr("Grouping compositions..."))

        self.group_thread = GroupingThread(
            self.searcher, items, self.spin_filter.value(), filtered_items=filtered_items
        )
        self.group_thread.progress_signal.connect(self.on_comp_progress)
        self.group_thread.status_signal.connect(lambda s: self.comp_progress.setFormat(s))
        self.group_thread.finished_signal.connect(self.on_comp_finished)
        self.group_thread.error_signal.connect(self.on_grouping_error)
        self.group_thread.start()

    def on_grouping_error(self, err):
        QMessageBox.critical(self, tr("Grouping Error"), err)
        # Fallback to ungrouped display
        self.comp_has_grouped_results = False
        self.comp_grouped_main = self.comp_raw_items or []
        self.comp_grouped_appendix = {}
        self.comp_grouped_summary = {}
        self.comp_grouped_filtered_main = self.comp_raw_filtered or []
        self.comp_grouped_filtered_appendix = {}
        self.comp_grouped_filtered_summary = {}
        self.display_comp_results(self.comp_raw_items or [], {}, {}, self.comp_raw_filtered or [], {}, {})

    def on_comp_finished(self, main_res, main_appx, main_summ, filt_res, filt_appx, filt_summ):
        self.comp_has_grouped_results = True
        
        final_main, final_appx, manual_known = self._apply_manual_exclusions(main_res, main_appx)
        
        self.comp_grouped_main = final_main or []
        self.comp_grouped_appendix = final_appx or {}
        self.comp_grouped_summary = main_summ or {} 
        
        if manual_known:
            if not self.comp_known:
                self.comp_known = []
            # Deduplicate by sys_id/part_id to avoid duplicate entries
            existing_ids = set()
            for item in self.comp_known:
                if item.get('type') == 'part':
                    existing_ids.add(f"PART:{item.get('part_id')}")
                else:
                    existing_ids.add(item.get('sys_id'))
            for item in manual_known:
                if item.get('type') == 'part':
                    key = f"PART:{item.get('part_id')}"
                else:
                    key = item.get('sys_id')
                if key and key not in existing_ids:
                    self.comp_known.append(item)
                    existing_ids.add(key)
            
        self.comp_grouped_filtered_main = filt_res or []
        self.comp_grouped_filtered_appendix = filt_appx or {}
        self.comp_grouped_filtered_summary = filt_summ or {}
        
        self.display_comp_results(
            self.comp_grouped_main, 
            self.comp_grouped_appendix, 
            self.comp_grouped_summary, 
            filt_res, 
            filt_appx, 
            filt_summ
        )

    def _collect_comp_items(self, main_res, main_appx, filt_res, filt_appx, known):
        all_items = []
        all_items.extend(main_res or [])
        for group_items in (main_appx or {}).values():
            all_items.extend(group_items)
        all_items.extend(filt_res or [])
        for group_items in (filt_appx or {}).values():
            all_items.extend(group_items)
        all_items.extend(known or [])
        return all_items

    def on_comp_header_clicked(self, section):
        if section not in (0, 1, 2, 3):
            return
        mode_map = {0: "score", 1: "shelfmark", 2: "title", 3: "system_id"}
        new_mode = mode_map.get(section, "score")
        if new_mode == self.comp_sort_mode:
            self.comp_sort_reverse = not self.comp_sort_reverse
        else:
            self.comp_sort_mode = new_mode
            self.comp_sort_reverse = new_mode == "score"
        order = Qt.SortOrder.DescendingOrder if self.comp_sort_reverse else Qt.SortOrder.AscendingOrder
        header = self.comp_tree.header()
        header.setSortIndicator(section, order)
        if self.is_comp_running or not self._has_comp_results():
            return
        if not self.chk_comp_flat.isChecked() and not self.comp_has_grouped_results:
            if self.comp_raw_items or self.comp_raw_filtered:
                self.start_grouping(self.comp_raw_items or [], self.comp_raw_filtered or [])
                return
        if self.comp_grouped_main or self.comp_grouped_filtered_main:
            self.display_comp_results(
                self.comp_grouped_main,
                self.comp_grouped_appendix,
                self.comp_grouped_summary,
                self.comp_grouped_filtered_main,
                self.comp_grouped_filtered_appendix,
                self.comp_grouped_filtered_summary,
            )

    def _current_comp_sort_mode(self):
        return self.comp_sort_mode or "score"

    def _get_comp_item_meta(self, item):
        sid = None
        shelf = None
        title = None
        item_type = item.get('type', '')
        if item_type == 'part':
            sid = item.get('sys_id')
            shelf = item.get('part_display', '')
            title = item.get('oxford_title', '')
        elif item_type == 'manuscript' and item.get('sys_id'):
            sid = item['sys_id']
            shelf, title = self.meta_mgr.get_meta_for_id(sid)
            if not shelf or shelf == "Unknown":
                shelf = self.meta_mgr.get_shelfmark_from_header(item.get('raw_header', ''))
        else:
            sid, _, shelf, title = self._get_meta_for_header(item.get('raw_header', ''))
        return sid or "", shelf or "", title or ""

    def _comp_sort_key(self, item, mode=None):
        sort_mode = mode or self._current_comp_sort_mode()
        if sort_mode == "score":
            return item.get('score', 0)

        sid, shelf, title = self._get_comp_item_meta(item)

        if sort_mode == "title":
            return title.casefold()
        if sort_mode == "system_id":
            return sid.casefold()

        shelf_key = natural_sort_key(shelf or sid)
        sid_key = natural_sort_key(sid)
        return (shelf_key, sid_key)

    def _sort_comp_items(self, items, mode=None):
        sort_mode = mode or self._current_comp_sort_mode()
        reverse = self.comp_sort_reverse if sort_mode == self.comp_sort_mode else sort_mode == "score"
        return sorted(items, key=lambda item: self._comp_sort_key(item, sort_mode), reverse=reverse)

    def _build_comp_preview_label(self, text_content):
        if not text_content:
            return QLabel("")
        flat = text_content.replace("\n", " ").replace("\r", " ").strip()
        return HiddenScrollArea(flat)

    def _set_comp_tree_text(self, node, column, text):
        node.setText(column, text)
        self._update_comp_tree_tooltip(node, column)

    def _format_score_with_boundary(self, item):
        """Format score string with boundary indicator if applicable."""
        score = int(item.get('score', 0))
        has_boundary = item.get('has_boundary_matches', False)
        boundary_count = item.get('boundary_match_count', 0)

        if has_boundary and boundary_count > 0:
            # Show "🔗 score" for boundary matches
            return f"🔗 {score}"
        return str(score)

    def _get_boundary_tooltip(self, item):
        """Get tooltip text for boundary match indicator."""
        has_boundary = item.get('has_boundary_matches', False)
        boundary_count = item.get('boundary_match_count', 0)
        boundary_quality = item.get('boundary_quality', 0)

        if has_boundary and boundary_count > 0:
            return tr("{} cross-paragraph matches").format(boundary_count)
        return ""

    def _process_snippet_queue(self):
        if not hasattr(self, 'snippet_queue') or not self.snippet_queue:
            return

        # Process a batch (e.g., 20)
        count = 0
        limit = 20

        while self.snippet_queue and count < limit:
            node = self.snippet_queue.pop(0)
            try:
                # Check if node is still valid (C++ object might be deleted)
                if node.treeWidget():
                    self._apply_comp_node_previews(node)
            except RuntimeError:
                pass # Node deleted
            count += 1

        if self.snippet_queue:
            QTimer.singleShot(10, self._process_snippet_queue)

    def _update_comp_tree_tooltip(self, node, column):
        text = node.text(column)
        if not text:
            node.setToolTip(column, "")
            return

        # Use the tree's font metrics for accuracy
        fm = self.comp_tree.fontMetrics()
        width = self.comp_tree.columnWidth(column)

        # For Column 0, we must subtract space for checkboxes and indentation
        if column == 0:
            # Approx 20px for checkbox + tree indentation level
            level = 0
            temp = node
            while temp.parent():
                level += 1
                temp = temp.parent()
            
            # Subtract indentation (default is 20 per level) and checkbox width
            width -= (self.comp_tree.indentation() * level) + 30

        elided = fm.elidedText(text, Qt.TextElideMode.ElideRight, width - 8)
        
        # If the text is elided (contains '...'), show the full text in tooltip
        node.setToolTip(column, text if elided != text else "")

    def _refresh_comp_tree_tooltips(self):
        root = self.comp_tree.invisibleRootItem()

        def visit(node):
            for col in (0, 1, 2, 3):
                self._update_comp_tree_tooltip(node, col)
            for i in range(node.childCount()):
                visit(node.child(i))

        for i in range(root.childCount()):
            visit(root.child(i))

    def _apply_comp_node_previews(self, node):
        data = node.data(0, Qt.ItemDataRole.UserRole + 1)
        if not data:
            return
            
        src_txt = data.get("source_ctx", "")
        ms_txt = data.get("ms_ctx", "")
        anchor = data.get("anchor")

        src_widget = HiddenScrollArea(src_txt.replace("\n", " "))
        self.comp_tree.setItemWidget(node, self.comp_col_context, src_widget)
        
        ms_widget = HiddenScrollArea(ms_txt.replace("\n", " "), anchor_text=anchor)
        self.comp_tree.setItemWidget(node, self.comp_col_ms_context, ms_widget)

    def _clear_comp_node_previews(self, node):
        if not node.data(0, Qt.ItemDataRole.UserRole + 1):
            return
        self.comp_tree.setItemWidget(node, self.comp_col_context, QLabel(""))
        self.comp_tree.setItemWidget(node, self.comp_col_ms_context, QLabel(""))

    def _set_comp_node_previews(self, node, source_text, ms_text, highlight_pattern=None, defer_widgets=False):
        if highlight_pattern and source_text:
            try:
                # Apply highlighting to Source Text if pattern exists
                regex = re.compile(highlight_pattern, re.IGNORECASE)
                # Only apply if not already highlighted (simple check)
                if '*' not in source_text:
                    source_text = regex.sub(r'*\g<0>*', source_text)
            except Exception:
                pass

        match = re.search(r'\*(.*?)\*', source_text or "")
        anchor = match.group(1) if match else None

        node.setData(
            0,
            Qt.ItemDataRole.UserRole + 1,
            {
                "source_ctx": source_text or "",
                "ms_ctx": ms_text or "",
                "anchor": anchor
            },
        )
        # Defer widget creation during batch operations to prevent freeze
        if not defer_widgets:
            self._apply_comp_node_previews(node)
        else:
            if not hasattr(self, 'snippet_queue'):
                self.snippet_queue = []
            self.snippet_queue.append(node)
            if len(self.snippet_queue) == 1:
                 QTimer.singleShot(10, self._process_snippet_queue)

    def display_comp_results(self, main_res, main_appx, main_summ, filt_res, filt_appx, filt_summ):
        # 1. איפוס וניקוי
        self.is_comp_running = False
        self.btn_comp_run.setText(tr("Analyze Composition"))
        self.btn_comp_run.setStyleSheet("background-color: #2980b9; color: white;")
        self.comp_progress.setVisible(False)
        for b in self.comp_export_buttons: b.setEnabled(True)

        # Reset Select All Checkbox
        self.chk_comp_header.blockSignals(True)
        self.chk_comp_header.setChecked(False)
        self.chk_comp_header.blockSignals(False)
        self._update_comp_export_label()

        if getattr(self, 'group_thread', None):
            self.group_thread.wait()
        self.group_thread = None

        self.comp_raw_items = main_res
        self.comp_raw_filtered = filt_res

        clean_main, clean_appx, known_main = self._apply_manual_exclusions(main_res, main_appx)
        clean_filt, clean_filt_appx, known_filt = self._apply_manual_exclusions(filt_res, filt_appx)

        if not hasattr(self, 'comp_known'): self.comp_known = []

        # Deduplicate by sys_id/part_id to avoid duplicate entries
        new_known = known_main + known_filt
        if new_known:
            existing_ids = set()
            for item in self.comp_known:
                if item.get('type') == 'part':
                    existing_ids.add(f"PART:{item.get('part_id')}")
                else:
                    existing_ids.add(item.get('sys_id'))
            for item in new_known:
                if item.get('type') == 'part':
                    key = f"PART:{item.get('part_id')}"
                else:
                    key = item.get('sys_id')
                if key and key not in existing_ids:
                    self.comp_known.append(item)
                    existing_ids.add(key)

        self.comp_main = clean_main
        self.comp_appendix = clean_appx
        self.comp_summary = main_summ
        self.comp_filtered_main = clean_filt
        self.comp_filtered_appendix = clean_filt_appx
        self.comp_filtered_summary = filt_summ
        
        self.comp_grouped_main = clean_main
        self.comp_grouped_appendix = clean_appx
        self.comp_grouped_summary = main_summ
        self.comp_grouped_filtered_main = clean_filt
        self.comp_grouped_filtered_appendix = clean_filt_appx
        self.comp_grouped_filtered_summary = filt_summ

        # Display Limit Logic
        full_main_count = len(clean_main)
        visible_main = clean_main

        msg_color = "black"
        if len(visible_main) < full_main_count:
            status_msg = tr("Showing top {} of {} results. (Export for full list)").format(len(visible_main), full_main_count)
            msg_color = "#e67e22" # Orange
        else:
            status_msg = tr("Found {} results.").format(full_main_count)

        if hasattr(self, 'lbl_comp_status'):
            self.lbl_comp_status.setText(status_msg)
            self.lbl_comp_status.setStyleSheet(f"color: {msg_color}; font-weight: bold;")

        ids_to_fetch = set()
        def _collect_id(item):
            sid = item.get('sys_id')
            if not sid:
                sid, _ = self.meta_mgr.parse_header_smart(item.get('raw_header', ''))
            if sid and sid not in self.meta_mgr.nli_cache:
                 ids_to_fetch.add(sid)

        # Collect domain data for composition results
        self._collect_comp_domain_data(clean_main, clean_appx, clean_filt, clean_filt_appx)

        # Block itemChanged signal during tree population to prevent O(n²) updates
        self.comp_tree_updating = True
        self.comp_tree.setUpdatesEnabled(False)
        self.comp_tree.clear()

        def make_checkable(node):
            node.setFlags(node.flags() | Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled | Qt.ItemFlag.ItemIsSelectable)
            node.setCheckState(0, Qt.CheckState.Unchecked)

        def add_manuscript_node(parent, ms_item):
            item_type = ms_item.get('type', '')

            def set_boundary_tooltip(node, item):
                """Set tooltip for boundary indicator if applicable."""
                boundary_tip = self._get_boundary_tooltip(item)
                if boundary_tip:
                    node.setToolTip(0, boundary_tip)

            def get_library_info(sid):
                """Get library code and full name for display."""
                library_code = self.meta_mgr.get_library_for_id(sid) if sid else ''
                library_full = get_library_display(library_code, short=False) if library_code else ''
                return library_code, library_full

            if item_type == 'part':
                # Part node - show Part display name with 📖 icon
                part_display = ms_item.get('part_display', '')
                oxford_title = ms_item.get('oxford_title', '')
                sid = ms_item.get('sys_id', '')
                shelf = f"📖 {part_display}" if part_display else sid
                t = oxford_title or ""
                library_code, library_full = get_library_info(sid)

                ms_node = QTreeWidgetItem(parent)
                self._set_comp_tree_text(ms_node, 0, self._format_score_with_boundary(ms_item))
                set_boundary_tooltip(ms_node, ms_item)
                self._set_comp_tree_text(ms_node, self.comp_col_shelfmark, shelf)
                self._set_comp_tree_text(ms_node, self.comp_col_library, library_code)
                if library_full:
                    ms_node.setToolTip(self.comp_col_library, library_full)
                self._set_comp_tree_text(ms_node, self.comp_col_title, t)
                self._set_comp_tree_text(ms_node, self.comp_col_sysid, ms_item.get('part_id', ''))
                make_checkable(ms_node)
                ms_node.setData(0, Qt.ItemDataRole.UserRole, ms_item)

                pages = ms_item.get('pages', [])
                folios = ms_item.get('folios', [])
                if len(pages) == 1:
                    p_item = pages[0]
                    p_sid, p_num, p_shelf, _ = self._get_meta_for_header(p_item['raw_header'])
                    folio_info = f" [{p_shelf}]" if p_shelf else ""
                    self._set_comp_tree_text(ms_node, self.comp_col_shelfmark, f"{shelf} ({tr('Image')} {p_num}{folio_info})")
                    self._set_comp_node_previews(ms_node, p_item.get('source_ctx', ''), p_item.get('text', ''), p_item.get('highlight_pattern'), defer_widgets=True)
                else:
                    if pages:
                        p0 = pages[0]
                        _, p0_num, _, _ = self._get_meta_for_header(p0['raw_header'])
                        folio_count = f", {len(folios)} folios" if len(folios) > 1 else ""
                        self._set_comp_tree_text(ms_node, self.comp_col_shelfmark, f"{shelf} ({len(pages)} matches{folio_count})")
                        self._set_comp_node_previews(ms_node, p0.get('source_ctx', ''), p0.get('text', ''), p0.get('highlight_pattern'), defer_widgets=True)

                    for p_item in pages:
                        p_sid, p_num, p_shelf, _ = self._get_meta_for_header(p_item['raw_header'])
                        folio_info = f" [{p_shelf}]" if p_shelf else ""
                        page_node = QTreeWidgetItem(ms_node)
                        self._set_comp_tree_text(page_node, 0, self._format_score_with_boundary(p_item))
                        set_boundary_tooltip(page_node, p_item)
                        self._set_comp_tree_text(page_node, self.comp_col_shelfmark, f"{tr('Image')} {p_num}{folio_info}")
                        self._set_comp_tree_text(page_node, self.comp_col_library, "")
                        self._set_comp_tree_text(page_node, self.comp_col_title, "")
                        self._set_comp_tree_text(page_node, self.comp_col_sysid, p_sid or "")
                        make_checkable(page_node)
                        page_node.setData(0, Qt.ItemDataRole.UserRole, p_item)
                        self._set_comp_node_previews(page_node, p_item.get('source_ctx', ''), p_item.get('text', ''), p_item.get('highlight_pattern'), defer_widgets=True)
            elif item_type == 'manuscript':
                sid = ms_item['sys_id']
                shelf, t = self.meta_mgr.get_meta_for_id(sid)
                if not shelf or shelf == "Unknown":
                    header_shelf = self.meta_mgr.get_shelfmark_from_header(ms_item.get('raw_header', ''))
                    if header_shelf: shelf = header_shelf
                library_code, library_full = get_library_info(sid)

                ms_node = QTreeWidgetItem(parent)
                self._set_comp_tree_text(ms_node, 0, self._format_score_with_boundary(ms_item))
                set_boundary_tooltip(ms_node, ms_item)
                self._set_comp_tree_text(ms_node, self.comp_col_shelfmark, shelf or tr("Unknown Shelfmark"))
                self._set_comp_tree_text(ms_node, self.comp_col_library, library_code)
                if library_full:
                    ms_node.setToolTip(self.comp_col_library, library_full)
                self._set_comp_tree_text(ms_node, self.comp_col_title, t or "")
                self._set_comp_tree_text(ms_node, self.comp_col_sysid, sid)
                make_checkable(ms_node)
                ms_node.setData(0, Qt.ItemDataRole.UserRole, ms_item)

                pages = ms_item.get('pages', [])
                if len(pages) == 1:
                    p_item = pages[0]
                    _, p_num, _, _ = self._get_meta_for_header(p_item['raw_header'])
                    self._set_comp_tree_text(ms_node, self.comp_col_shelfmark, f"{shelf or tr('Unknown Shelfmark')} ({tr('Image')} {p_num})")
                    self._set_comp_node_previews(ms_node, p_item.get('source_ctx', ''), p_item.get('text', ''), p_item.get('highlight_pattern'), defer_widgets=True)
                else:
                    if pages:
                        p0 = pages[0]
                        _, p0_num, _, _ = self._get_meta_for_header(p0['raw_header'])
                        self._set_comp_tree_text(ms_node, self.comp_col_shelfmark, f"{shelf or tr('Unknown Shelfmark')} ({tr('Image')} {p0_num}...)")
                        self._set_comp_node_previews(ms_node, p0.get('source_ctx', ''), p0.get('text', ''), p0.get('highlight_pattern'), defer_widgets=True)

                    for p_item in pages:
                        _, p_num, _, _ = self._get_meta_for_header(p_item['raw_header'])
                        page_node = QTreeWidgetItem(ms_node)
                        self._set_comp_tree_text(page_node, 0, self._format_score_with_boundary(p_item))
                        set_boundary_tooltip(page_node, p_item)
                        self._set_comp_tree_text(page_node, self.comp_col_shelfmark, f"{tr('Image')} {p_num}")
                        self._set_comp_tree_text(page_node, self.comp_col_library, "")
                        self._set_comp_tree_text(page_node, self.comp_col_title, "")
                        self._set_comp_tree_text(page_node, self.comp_col_sysid, "")
                        make_checkable(page_node)
                        page_node.setData(0, Qt.ItemDataRole.UserRole, p_item)
                        self._set_comp_node_previews(page_node, p_item.get('source_ctx', ''), p_item.get('text', ''), p_item.get('highlight_pattern'), defer_widgets=True)
            else:
                # Fallback
                sid, _, shelf, title = self._get_meta_for_header(ms_item.get('raw_header', ''))
                library_code, library_full = get_library_info(sid)
                node = QTreeWidgetItem(parent)
                self._set_comp_tree_text(node, 0, self._format_score_with_boundary(ms_item))
                set_boundary_tooltip(node, ms_item)
                self._set_comp_tree_text(node, self.comp_col_shelfmark, shelf)
                self._set_comp_tree_text(node, self.comp_col_library, library_code)
                if library_full:
                    node.setToolTip(self.comp_col_library, library_full)
                self._set_comp_tree_text(node, self.comp_col_title, title)
                self._set_comp_tree_text(node, self.comp_col_sysid, sid)
                make_checkable(node)
                node.setData(0, Qt.ItemDataRole.UserRole, ms_item)
                self._set_comp_node_previews(node, ms_item.get('source_ctx', ''), ms_item.get('text', ''), ms_item.get('highlight_pattern'), defer_widgets=True)

            _collect_id(ms_item)


        if self.chk_comp_flat.isChecked():
            all_flat = self._collect_comp_items(
                clean_main, clean_appx, clean_filt, clean_filt_appx, self.comp_known
            )
            sorted_flat = self._sort_comp_items(all_flat)
            visible_flat = sorted_flat

            root = QTreeWidgetItem(self.comp_tree, [tr("All Results ({})").format(len(visible_flat))])
            root.setExpanded(True)
            make_checkable(root)

            # Use batched loading for flat view too
            if visible_flat:
                self.comp_tree.setUpdatesEnabled(True)
                self._update_recursive_button_state()
                self._start_batched_tree_load(root, visible_flat)
                if ids_to_fetch:
                    self.start_metadata_loading(list(ids_to_fetch))
                return  # Early return - batched loading handles the rest

        else:
            # 1. Main Results - Using batched loading for performance
            sorted_main = self._sort_comp_items(clean_main)
            visible_sorted_main = sorted_main
            root_main = None

            if visible_sorted_main:
                root_main = QTreeWidgetItem(self.comp_tree, [tr("Main Results ({})").format(len(visible_sorted_main))])
                root_main.setData(0, Qt.ItemDataRole.UserRole + 100, "ROOT_MAIN")
                root_main.setExpanded(True)
                make_checkable(root_main)
                # Store items for batched loading later
                root_main.setData(0, Qt.ItemDataRole.UserRole + 202, visible_sorted_main)

            # 2. Appendix - Using Virtual Children for performance
            if clean_appx:
                total_appx = sum(len(v) for v in clean_appx.values())
                root_appx = QTreeWidgetItem(self.comp_tree, [tr("Appendix - Grouped ({})").format(total_appx)])
                root_appx.setData(0, Qt.ItemDataRole.UserRole + 100, "ROOT_APPX")
                root_appx.setExpanded(True)
                make_checkable(root_appx)

                sorted_groups = sorted(clean_appx.items(), key=lambda x: len(x[1]), reverse=True)
                for sig, items in sorted_groups:
                    group_node = QTreeWidgetItem(root_appx, ["", "", f"{sig} ({len(items)})", ""])
                    make_checkable(group_node)
                    # Virtual children: store items data, add placeholder
                    group_node.setData(0, Qt.ItemDataRole.UserRole + 200, items)  # Store items for lazy load
                    placeholder = QTreeWidgetItem(group_node, [tr("Loading...")])
                    placeholder.setData(0, Qt.ItemDataRole.UserRole + 201, "PLACEHOLDER")

            # 3. Filtered 
            total_filt = len(clean_filt) + sum(len(v) for v in clean_filt_appx.values())
            if total_filt > 0:
                root_filt = QTreeWidgetItem(self.comp_tree, [tr("Filtered ({})").format(total_filt)])
                root_filt.setData(0, Qt.ItemDataRole.UserRole + 100, "ROOT_FILT")
                root_filt.setExpanded(True)
                make_checkable(root_filt)
                
                # Filtered Main
                for item in self._sort_comp_items(clean_filt):
                    add_manuscript_node(root_filt, item)
                
                # Filtered Appendix - Using Virtual Children for performance
                for sig, items in sorted(clean_filt_appx.items(), key=lambda x: len(x[1]), reverse=True):
                    g_node = QTreeWidgetItem(root_filt, ["", "", f"{sig} ({len(items)})", ""])
                    make_checkable(g_node)
                    # Virtual children: store items data, add placeholder
                    g_node.setData(0, Qt.ItemDataRole.UserRole + 200, items)
                    placeholder = QTreeWidgetItem(g_node, [tr("Loading...")])
                    placeholder.setData(0, Qt.ItemDataRole.UserRole + 201, "PLACEHOLDER")

            # 4. Excluded
            if self.comp_known:
                root_known = QTreeWidgetItem(self.comp_tree, [tr("Excluded ({})").format(len(self.comp_known))])
                root_known.setData(0, Qt.ItemDataRole.UserRole + 100, "ROOT_KNOWN")
                root_known.setForeground(0, Qt.GlobalColor.darkGray)
                make_checkable(root_known)
                
                for item in self._sort_comp_items(self.comp_known):
                    add_manuscript_node(root_known, item)

            # Start batched loading for Main Results (deferred to prevent freeze)
            if root_main and visible_sorted_main:
                self.comp_tree.setUpdatesEnabled(True)
                self._update_recursive_button_state()
                # Start batched loading - this will call filters when done
                self._start_batched_tree_load(root_main, visible_sorted_main)
                if ids_to_fetch:
                    self.start_metadata_loading(list(ids_to_fetch))
                return  # Early return - batched loading handles the rest

        self.comp_tree.setUpdatesEnabled(True)
        self.comp_tree_updating = False  # Re-enable itemChanged signal
        self._update_comp_filter_indicators()
        self._apply_comp_tree_filters()
        self._update_recursive_button_state()

        if ids_to_fetch:
            self.start_metadata_loading(list(ids_to_fetch))

    def _make_node_checkable(self, node):
        """Make a tree node checkable."""
        node.setFlags(node.flags() | Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled | Qt.ItemFlag.ItemIsSelectable)
        node.setCheckState(0, Qt.CheckState.Unchecked)

    def _add_manuscript_node(self, parent, ms_item, defer_widgets=True):
        """Add a manuscript/part node to the tree. Used for lazy/batched loading."""
        def get_library_info(sid):
            library_code = self.meta_mgr.get_library_for_id(sid) if sid else ''
            library_full = get_library_display(library_code, short=False) if library_code else ''
            return library_code, library_full

        item_type = ms_item.get('type', '')
        if item_type == 'part':
            part_display = ms_item.get('part_display', '')
            oxford_title = ms_item.get('oxford_title', '')
            sid = ms_item.get('sys_id', '')
            shelf = f"📖 {part_display}" if part_display else sid
            t = oxford_title or ""
            library_code, library_full = get_library_info(sid)

            ms_node = QTreeWidgetItem(parent)
            self._set_comp_tree_text(ms_node, 0, str(int(ms_item.get('score', 0))))
            self._set_comp_tree_text(ms_node, self.comp_col_shelfmark, shelf)
            self._set_comp_tree_text(ms_node, self.comp_col_library, library_code)
            if library_full:
                ms_node.setToolTip(self.comp_col_library, library_full)
            self._set_comp_tree_text(ms_node, self.comp_col_title, t)
            self._set_comp_tree_text(ms_node, self.comp_col_sysid, ms_item.get('part_id', ''))
            self._make_node_checkable(ms_node)
            ms_node.setData(0, Qt.ItemDataRole.UserRole, ms_item)

            pages = ms_item.get('pages', [])
            folios = ms_item.get('folios', [])
            if len(pages) == 1:
                p_item = pages[0]
                p_sid, p_num, p_shelf, _ = self._get_meta_for_header(p_item['raw_header'])
                folio_info = f" [{p_shelf}]" if p_shelf else ""
                self._set_comp_tree_text(ms_node, self.comp_col_shelfmark, f"{shelf} ({tr('Image')} {p_num}{folio_info})")
                self._set_comp_node_previews(ms_node, p_item.get('source_ctx', ''), p_item.get('text', ''), p_item.get('highlight_pattern'), defer_widgets=defer_widgets)
            else:
                if pages:
                    p0 = pages[0]
                    _, p0_num, _, _ = self._get_meta_for_header(p0['raw_header'])
                    folio_count = f", {len(folios)} folios" if len(folios) > 1 else ""
                    self._set_comp_tree_text(ms_node, self.comp_col_shelfmark, f"{shelf} ({len(pages)} matches{folio_count})")
                    self._set_comp_node_previews(ms_node, p0.get('source_ctx', ''), p0.get('text', ''), p0.get('highlight_pattern'), defer_widgets=defer_widgets)

                for p_item in pages:
                    p_sid, p_num, p_shelf, _ = self._get_meta_for_header(p_item['raw_header'])
                    folio_info = f" [{p_shelf}]" if p_shelf else ""
                    page_node = QTreeWidgetItem(ms_node)
                    self._set_comp_tree_text(page_node, 0, str(int(p_item.get('score', 0))))
                    self._set_comp_tree_text(page_node, self.comp_col_shelfmark, f"{tr('Image')} {p_num}{folio_info}")
                    self._set_comp_tree_text(page_node, self.comp_col_library, "")
                    self._set_comp_tree_text(page_node, self.comp_col_title, "")
                    self._set_comp_tree_text(page_node, self.comp_col_sysid, p_sid or "")
                    self._make_node_checkable(page_node)
                    page_node.setData(0, Qt.ItemDataRole.UserRole, p_item)
                    self._set_comp_node_previews(page_node, p_item.get('source_ctx', ''), p_item.get('text', ''), p_item.get('highlight_pattern'), defer_widgets=defer_widgets)
        elif item_type == 'manuscript':
            sid = ms_item['sys_id']
            shelf, t = self.meta_mgr.get_meta_for_id(sid)
            if not shelf or shelf == "Unknown":
                header_shelf = self.meta_mgr.get_shelfmark_from_header(ms_item.get('raw_header', ''))
                if header_shelf: shelf = header_shelf
            library_code, library_full = get_library_info(sid)

            ms_node = QTreeWidgetItem(parent)
            self._set_comp_tree_text(ms_node, 0, str(int(ms_item.get('score', 0))))
            self._set_comp_tree_text(ms_node, self.comp_col_shelfmark, shelf or tr("Unknown Shelfmark"))
            self._set_comp_tree_text(ms_node, self.comp_col_library, library_code)
            if library_full:
                ms_node.setToolTip(self.comp_col_library, library_full)
            self._set_comp_tree_text(ms_node, self.comp_col_title, t or "")
            self._set_comp_tree_text(ms_node, self.comp_col_sysid, sid)
            self._make_node_checkable(ms_node)
            ms_node.setData(0, Qt.ItemDataRole.UserRole, ms_item)

            pages = ms_item.get('pages', [])
            if len(pages) == 1:
                p_item = pages[0]
                _, p_num, _, _ = self._get_meta_for_header(p_item['raw_header'])
                self._set_comp_tree_text(ms_node, self.comp_col_shelfmark, f"{shelf or tr('Unknown Shelfmark')} ({tr('Image')} {p_num})")
                self._set_comp_node_previews(ms_node, p_item.get('source_ctx', ''), p_item.get('text', ''), p_item.get('highlight_pattern'), defer_widgets=defer_widgets)
            else:
                if pages:
                    p0 = pages[0]
                    _, p0_num, _, _ = self._get_meta_for_header(p0['raw_header'])
                    self._set_comp_tree_text(ms_node, self.comp_col_shelfmark, f"{shelf or tr('Unknown Shelfmark')} ({tr('Image')} {p0_num}...)")
                    self._set_comp_node_previews(ms_node, p0.get('source_ctx', ''), p0.get('text', ''), p0.get('highlight_pattern'), defer_widgets=defer_widgets)

                for p_item in pages:
                    _, p_num, _, _ = self._get_meta_for_header(p_item['raw_header'])
                    page_node = QTreeWidgetItem(ms_node)
                    self._set_comp_tree_text(page_node, 0, str(int(p_item.get('score', 0))))
                    self._set_comp_tree_text(page_node, self.comp_col_shelfmark, f"{tr('Image')} {p_num}")
                    self._set_comp_tree_text(page_node, self.comp_col_library, "")
                    self._set_comp_tree_text(page_node, self.comp_col_title, "")
                    self._set_comp_tree_text(page_node, self.comp_col_sysid, "")
                    self._make_node_checkable(page_node)
                    page_node.setData(0, Qt.ItemDataRole.UserRole, p_item)
                    self._set_comp_node_previews(page_node, p_item.get('source_ctx', ''), p_item.get('text', ''), p_item.get('highlight_pattern'), defer_widgets=defer_widgets)
        else:
            # Fallback
            sid, _, shelf, title = self._get_meta_for_header(ms_item.get('raw_header', ''))
            library_code, library_full = get_library_info(sid)
            node = QTreeWidgetItem(parent)
            self._set_comp_tree_text(node, 0, str(int(ms_item.get('score', 0))))
            self._set_comp_tree_text(node, self.comp_col_shelfmark, shelf)
            self._set_comp_tree_text(node, self.comp_col_library, library_code)
            if library_full:
                node.setToolTip(self.comp_col_library, library_full)
            self._set_comp_tree_text(node, self.comp_col_title, title)
            self._set_comp_tree_text(node, self.comp_col_sysid, sid)
            self._make_node_checkable(node)
            node.setData(0, Qt.ItemDataRole.UserRole, ms_item)
            self._set_comp_node_previews(node, ms_item.get('source_ctx', ''), ms_item.get('text', ''), ms_item.get('highlight_pattern'), defer_widgets=defer_widgets)

    def _add_single_node_to_tree(self, parent, ms_item):
        """Dedicated helper to add one row to the tree."""
        sid = ms_item.get('sys_id')
        if not sid:
            sid, _ = self.meta_mgr.parse_header_smart(ms_item.get('raw_header', ''))

        shelf, t = self.meta_mgr.get_meta_for_id(sid)
        display_shelf = shelf if shelf and shelf != "Unknown" else (sid if sid else "Loading...")

        # Get library info
        library_code = self.meta_mgr.get_library_for_id(sid)
        library_display = get_library_display(library_code, short=False) if library_code else ""

        node = QTreeWidgetItem(parent)
        self._set_comp_tree_text(node, 0, str(int(ms_item.get('score', 0)))) # עיגול הציון
        self._set_comp_tree_text(node, self.comp_col_shelfmark, display_shelf)
        self._set_comp_tree_text(node, self.comp_col_library, library_display)
        self._set_comp_tree_text(node, self.comp_col_title, t or "")
        self._set_comp_tree_text(node, self.comp_col_sysid, sid)

        node.setFlags(node.flags() | Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled | Qt.ItemFlag.ItemIsSelectable)
        node.setCheckState(0, Qt.CheckState.Unchecked)
        node.setData(0, Qt.ItemDataRole.UserRole, ms_item)

        pages = ms_item.get('pages', [])

        if len(pages) == 1:
            p_item = pages[0]
            p_num = "Img"
            if 'raw_header' in p_item:
                _, p_num_extracted, _, _ = self._get_meta_for_header(p_item['raw_header'])
                if p_num_extracted: p_num = p_num_extracted

            self._set_comp_tree_text(node, self.comp_col_shelfmark, f"{display_shelf} (Img {p_num})")
            self._set_comp_node_previews(node, p_item.get('source_ctx', ''), p_item.get('text', ''), p_item.get('highlight_pattern'))

        elif len(pages) > 1:
             self._set_comp_tree_text(node, self.comp_col_shelfmark, f"{display_shelf} ({len(pages)} matches)")

             if pages:
                 first_p = pages[0]
                 self._set_comp_node_previews(node, first_p.get('source_ctx', ''), first_p.get('text', ''), first_p.get('highlight_pattern'))
             # ---------------------------------------------------------

             for p_item in pages:
                child = QTreeWidgetItem(node)
                self._set_comp_tree_text(child, 0, str(int(p_item.get('score', 0))))

                p_num_child = "?"
                if 'raw_header' in p_item:
                    _, p_val, _, _ = self._get_meta_for_header(p_item['raw_header'])
                    if p_val: p_num_child = p_val

                child.setText(1, f"Img {p_num_child}")

                child.setData(0, Qt.ItemDataRole.UserRole, p_item)
                self._set_comp_node_previews(child, p_item.get('source_ctx', ''), p_item.get('text', ''), p_item.get('highlight_pattern'))

    # ========== Batched Tree Loading for Performance ==========
    def _start_batched_tree_load(self, parent_node, items, batch_size=50):
        """Start loading items into tree in batches to prevent UI freeze."""
        self._batch_queue = list(items)  # Copy to avoid mutation issues
        self._batch_parent = parent_node
        self._batch_size = batch_size
        self._batch_index = 0
        # Start first batch immediately
        self._process_tree_batch()

    def _process_tree_batch(self):
        """Process one batch of items and schedule next batch."""
        if not hasattr(self, '_batch_queue') or self._batch_index >= len(self._batch_queue):
            # Done loading - cleanup
            self._batch_queue = None
            self._batch_parent = None
            self.comp_tree.setUpdatesEnabled(True)
            self.comp_tree_updating = False  # Re-enable itemChanged signal
            self._update_comp_filter_indicators()
            self._apply_comp_tree_filters()
            return

        # Process batch
        self.comp_tree.setUpdatesEnabled(False)
        end_index = min(self._batch_index + self._batch_size, len(self._batch_queue))

        for i in range(self._batch_index, end_index):
            self._add_manuscript_node(self._batch_parent, self._batch_queue[i])

        self._batch_index = end_index
        self.comp_tree.setUpdatesEnabled(True)

        # Schedule next batch with small delay to let UI breathe
        if self._batch_index < len(self._batch_queue):
            QTimer.singleShot(0, self._process_tree_batch)

    def _trigger_lazy_metadata_fetch(self):
        """Starts background fetching for items that are currently displayed but missing data."""
        missing_ids = set()
        for item in self.batch_items:
            sid = item.get('sys_id')
            if not sid:
                sid, _ = self.meta_mgr.parse_header_smart(item.get('raw_header'))
            
            if sid and sid not in self.meta_mgr.nli_cache:
                 missing_ids.add(sid)
        
        if missing_ids:
            self.start_metadata_loading(list(missing_ids))
    
    def on_comp_tree_item_changed(self, item, column):
        if self.comp_tree_updating or column != 0:
            return

        self.comp_tree_updating = True
        state = item.checkState(0)
        if item.childCount() > 0 and state in (Qt.CheckState.Checked, Qt.CheckState.Unchecked):
            for i in range(item.childCount()):
                item.child(i).setCheckState(0, state)

        self._sync_parent_check_state(item)

        # Sync "Select All" checkbox state
        all_checked = True
        root = self.comp_tree.invisibleRootItem()
        if root.childCount() == 0:
            all_checked = False
        else:
            for i in range(root.childCount()):
                if root.child(i).checkState(0) == Qt.CheckState.Unchecked:
                    all_checked = False
                    break

        self.chk_comp_header.blockSignals(True)
        self.chk_comp_header.setChecked(all_checked)
        self.chk_comp_header.blockSignals(False)

        self.comp_tree_updating = False
        self._update_recursive_button_state()
        self._update_comp_export_label()

    def on_comp_header_toggled(self, checked):
        """Toggle all root items in the composition tree."""
        state = Qt.CheckState.Checked if checked else Qt.CheckState.Unchecked

        self.comp_tree.blockSignals(True)
        root = self.comp_tree.invisibleRootItem()
        for i in range(root.childCount()):
            item = root.child(i)
            item.setCheckState(0, state)
            self._set_check_state_recursive(item, state)
        self.comp_tree.blockSignals(False)

        self._update_comp_export_label()
        self._update_recursive_button_state()

    def _set_check_state_recursive(self, item, state):
        for i in range(item.childCount()):
            child = item.child(i)
            child.setCheckState(0, state)
            self._set_check_state_recursive(child, state)

    def _update_comp_export_label(self):
        has_selection = bool(self._collect_checked_comp_page_uids())
        if has_selection:
            self.lbl_comp_export.setText(tr("Export selected results"))
        else:
            self.lbl_comp_export.setText(tr("Save Report"))
        if hasattr(self, 'btn_comp_add_to_list'):
            self.btn_comp_add_to_list.setEnabled(has_selection)

    def _collect_checked_comp_items_struct(self):
        """
        Collect checked items maintaining the structure (Main, Appendix, etc.)
        Returns: (main, appendix, filtered_main, filtered_appendix, known)
        Only items that are CHECKED (or have checked descendants) are returned.
        """
        sel_main = []
        sel_appx = {}
        sel_filt = []
        sel_filt_appx = {}
        sel_known = []

        # Helper to collect checked children from a node
        def collect_from_node(node):
            collected = []
            # If node is a leaf (Manuscript or Page)
            data = node.data(0, Qt.ItemDataRole.UserRole)
            if data:
                if node.checkState(0) == Qt.CheckState.Checked:
                    collected.append(data)
                return collected

            # If node is a group container (no data)
            for k in range(node.childCount()):
                child = node.child(k)
                child_data = child.data(0, Qt.ItemDataRole.UserRole)
                if not child_data:
                    # Recursive group (e.g. Appendix group)
                    res = collect_from_node(child)
                    collected.extend(res)
                    continue

                # It is a manuscript item
                if child.checkState(0) == Qt.CheckState.Unchecked:
                    continue

                # If fully checked or partially checked
                if child.checkState(0) == Qt.CheckState.Checked:
                    # Full manuscript selected
                    collected.append(child_data)
                elif child.checkState(0) == Qt.CheckState.PartiallyChecked:
                    # Some pages selected
                    # Clone item
                    import copy
                    new_item = copy.copy(child_data)
                    new_item['pages'] = []

                    # Find checked pages
                    for p_idx in range(child.childCount()):
                        p_node = child.child(p_idx)
                        if p_node.checkState(0) == Qt.CheckState.Checked:
                            p_data = p_node.data(0, Qt.ItemDataRole.UserRole)
                            new_item['pages'].append(p_data)

                    if new_item['pages']:
                            collected.append(new_item)
            return collected

        root = self.comp_tree.invisibleRootItem()

        # If Flat Mode:
        if self.chk_comp_flat.isChecked():
            # Root 0 is "All Results"
            if root.childCount() > 0:
                sel_main = collect_from_node(root.child(0))
            return sel_main, {}, [], {}, []

        # Use node data (UserRole+100) to identify categories
        for i in range(root.childCount()):
            node = root.child(i)
            node_type = node.data(0, Qt.ItemDataRole.UserRole + 100)

            items = collect_from_node(node)
            if not items: continue

            if node_type == "ROOT_MAIN":
                 sel_main.extend(items)
            elif node_type == "ROOT_APPX":
                 # Custom traversal for Appendix to preserve grouping structure
                 for k in range(node.childCount()):
                     group_node = node.child(k)
                     group_sig_full = group_node.text(2)
                     group_sig = group_sig_full.rpartition(' (')[0] # Remove count

                     group_items = collect_from_node(group_node)
                     if group_items:
                         sel_appx[group_sig] = group_items

            elif node_type == "ROOT_FILT":
                 # Iterate children to separate Main/Appendix in filtered
                 for k in range(node.childCount()):
                     child = node.child(k)
                     c_data = child.data(0, Qt.ItemDataRole.UserRole)
                     if c_data:
                         # Direct item -> Filtered Main
                         if child.checkState(0) == Qt.CheckState.Checked:
                             sel_filt.append(c_data)
                         elif child.checkState(0) == Qt.CheckState.PartiallyChecked:
                             # Reuse logic from helper
                             import copy
                             new_item = copy.copy(c_data)
                             new_item['pages'] = []
                             for p_idx in range(child.childCount()):
                                 p_node = child.child(p_idx)
                                 if p_node.checkState(0) == Qt.CheckState.Checked:
                                     new_item['pages'].append(p_node.data(0, Qt.ItemDataRole.UserRole))
                             if new_item['pages']: sel_filt.append(new_item)
                     else:
                         # Group node -> Filtered Appendix
                         g_sig = child.text(2).rpartition(' (')[0]
                         g_items = collect_from_node(child)
                         if g_items:
                             sel_filt_appx[g_sig] = g_items

            elif node_type == "ROOT_KNOWN":
                 sel_known.extend(items)

        return sel_main, sel_appx, sel_filt, sel_filt_appx, sel_known

    def on_comp_tree_item_expanded(self, item):
        # Check for virtual children (lazy loading)
        virtual_items = item.data(0, Qt.ItemDataRole.UserRole + 200)
        if virtual_items is not None:
            # Remove placeholder child
            while item.childCount() > 0:
                child = item.child(0)
                if child.data(0, Qt.ItemDataRole.UserRole + 201) == "PLACEHOLDER":
                    item.removeChild(child)
                else:
                    break

            # Populate real children - block itemChanged signal
            self.comp_tree_updating = True
            self.comp_tree.setUpdatesEnabled(False)
            sorted_items = self._sort_comp_items(virtual_items)
            for ms_item in sorted_items:
                self._add_manuscript_node(item, ms_item)
            self.comp_tree.setUpdatesEnabled(True)
            self.comp_tree_updating = False

            # Clear virtual data to prevent re-population
            item.setData(0, Qt.ItemDataRole.UserRole + 200, None)

        if item.childCount() > 0:
            self._clear_comp_node_previews(item)

    def on_comp_tree_item_collapsed(self, item):
        if item.childCount() > 0:
            self._apply_comp_node_previews(item)

    def _sync_parent_check_state(self, item):
        parent = item.parent()
        if not parent:
            return

        states = []
        for i in range(parent.childCount()):
            states.append(parent.child(i).checkState(0))

        if all(s == Qt.CheckState.Checked for s in states):
            parent.setCheckState(0, Qt.CheckState.Checked)
        elif all(s == Qt.CheckState.Unchecked for s in states):
            parent.setCheckState(0, Qt.CheckState.Unchecked)
        else:
            parent.setCheckState(0, Qt.CheckState.PartiallyChecked)

        self._sync_parent_check_state(parent)

    def _collect_checked_comp_page_uids(self):
        uids = set()

        def visit(node):
            data = node.data(0, Qt.ItemDataRole.UserRole)
            if data:
                if data.get('type') in ('manuscript', 'part'):
                    if node.childCount() > 0:
                        for i in range(node.childCount()):
                            visit(node.child(i))
                    else:
                        if node.checkState(0) == Qt.CheckState.Checked:
                            pages = data.get('pages', [])
                            if pages and pages[0].get('uid'):
                                uids.add(pages[0]['uid'])
                else:
                    if node.checkState(0) == Qt.CheckState.Checked:
                        uid = data.get('uid')
                        if uid:
                            uids.add(uid)
            else:
                for i in range(node.childCount()):
                    visit(node.child(i))

        root = self.comp_tree.invisibleRootItem()
        for i in range(root.childCount()):
            visit(root.child(i))

        return sorted(uids)

    def _collect_all_comp_page_uids(self):
        uids = set()

        def add_from_item(item):
            if item.get('type') in ('manuscript', 'part'):
                for page in item.get('pages', []):
                    uid = page.get('uid')
                    if uid:
                        uids.add(uid)
            else:
                uid = item.get('uid')
                if uid:
                    uids.add(uid)

        for item in self.comp_main:
            add_from_item(item)
        for group_items in self.comp_appendix.values():
            for item in group_items:
                add_from_item(item)
        for item in self.comp_filtered_main:
            add_from_item(item)
        for group_items in self.comp_filtered_appendix.values():
            for item in group_items:
                add_from_item(item)
        for item in self.comp_known:
            add_from_item(item)

        return sorted(uids)

    def _update_recursive_button_state(self):
        if self.is_comp_running:
            self.btn_comp_recursive.setEnabled(False)
            return
        checked = self._collect_checked_comp_page_uids()
        if checked:
            self.btn_comp_recursive.setText(tr("Recursive Search in Results"))
        else:
            self.btn_comp_recursive.setText(tr("Full Recursive Search"))
        self.btn_comp_recursive.setEnabled(True)

    def _has_comp_results(self):
        if self.comp_main or self.comp_appendix or self.comp_known:
            return True
        if self.comp_filtered_main or self.comp_filtered_appendix:
            return True
        return False

    def show_comp_detail(self, item, col):
        # 1. Validate Click
        data = item.data(0, Qt.ItemDataRole.UserRole)
        if not data: return # It's a structural node, ignore
        
        # 2. Flatten the Tree to create a navigation list (PAGES ONLY)
        flat_list = []
        clicked_index = -1
        
        # If user clicked a Manuscript/Part Node (top level), check if it's single page or multi
        target_item = item
        if data.get('type') in ('manuscript', 'part'):
            if item.childCount() > 0:
                # Multi-page: Auto-select first child
                target_item = item.child(0)
            else:
                # Single-page: The manuscript node IS the target
                pass

        # Helper to process a page node or a single-page manuscript/part
        def process_page_data(node_data, node_ref):
            # If it's a manuscript/part node (single page), extract the single page data
            if node_data.get('type') in ('manuscript', 'part'):
                pages = node_data.get('pages', [])
                if len(pages) == 1:
                    node_data = pages[0]
                else:
                    return # Should not happen for leaf traversal

            sid, p, shelf, title = self._get_meta_for_header(node_data['raw_header'])

            ready_data = {
                'uid': node_data['uid'],
                'raw_header': node_data['raw_header'],
                'text': node_data['text'], # Snippet
                'full_text': None, # Will be fetched by Dialog on load
                'source_ctx': node_data.get('source_ctx', ''),
                'highlight_pattern': node_data.get('highlight_pattern'),
                'display': {
                    'shelfmark': shelf,
                    'title': title,
                    'img': p,
                    'source': node_data.get('src_lbl', 'Source')
                }
            }
            flat_list.append(ready_data)

            if node_ref is target_item:
                nonlocal clicked_index
                clicked_index = len(flat_list) - 1

        # Traverse Tree Logic for Manuscript Grouping
        root = self.comp_tree.invisibleRootItem()
        for i in range(root.childCount()):
            category_node = root.child(i) # "Main", "Appendix", etc.

            # Recurse into category
            for j in range(category_node.childCount()):
                sub_node = category_node.child(j)

                if sub_node.childCount() > 0:
                    d = sub_node.data(0, Qt.ItemDataRole.UserRole)
                    if d and d.get('type') in ('manuscript', 'part'):
                        # It is a Manuscript/Part with multiple pages
                        for k in range(sub_node.childCount()):
                            page_node = sub_node.child(k)
                            process_page_data(page_node.data(0, Qt.ItemDataRole.UserRole), page_node)
                    else:
                        # It is an Appendix Group
                        for k in range(sub_node.childCount()):
                            ms_node = sub_node.child(k)
                            # Check if multi-page or single-page
                            if ms_node.childCount() > 0:
                                for m in range(ms_node.childCount()):
                                    page_node = ms_node.child(m)
                                    process_page_data(page_node.data(0, Qt.ItemDataRole.UserRole), page_node)
                            else:
                                # Single page manuscript/part in Appendix
                                process_page_data(ms_node.data(0, Qt.ItemDataRole.UserRole), ms_node)
                else:
                    # Leaf Manuscript/Part (Single Page) in Main
                    d = sub_node.data(0, Qt.ItemDataRole.UserRole)
                    if d and d.get('type') in ('manuscript', 'part'):
                        process_page_data(d, sub_node)

        if clicked_index == -1: return

        # 3. Open Dialog with List
        ResultDialog(self, flat_list, clicked_index, self.meta_mgr, self.searcher).exec()

    def _refresh_comp_tree_metadata(self):

        def update_node(node):
            node_data = node.data(0, Qt.ItemDataRole.UserRole)
            if not node_data:
                return

            sys_id, _, shelf, title = self._get_meta_for_header(node_data.get('raw_header', ''))

            node.setText(1, shelf)
            node.setText(2, title)
            node.setText(3, sys_id or '')

        root = self.comp_tree.invisibleRootItem()
        for i in range(root.childCount()):
            group = root.child(i)
            for j in range(group.childCount()):
                child = group.child(j)
                if child.childCount() > 0:
                    for k in range(child.childCount()):
                        update_node(child.child(k))
                else:
                    update_node(child)

    def _fmt_item_legacy(self, item):
        # Fallback for old page style if needed
        sid, p_num, shelf, title = self._get_meta_for_header(item.get('raw_header', ''))

        def clean(t):
            t = str(t or "")
            t = re.sub(r'<span[^>]*>', '*', t).replace('</span>', '*')
            t = t.replace("<br>", " ").replace("\n", " ").replace("\r", "")
            t = re.sub(r'<[^>]+>', '', t)
            t = re.sub(r'\s+', ' ', t)
            t = re.sub(r'\*(\s+)\*', r'\1', t)
            return t.strip()

        return [
            "=" * 80,
            f"{shelf or sid} | {title or 'Untitled'} | Img: {p_num} | Version: {item.get('src_lbl','')} | ID: {item.get('uid', sid)} (Score: {item.get('score', 0)})",
            tr("Source Context") + ":", clean(item.get('source_ctx', '')), "",
            tr("Manuscript") + ":", clean(item.get('text', '')), ""
        ]

    def _format_comp_entry(self, item):
        sys_id, page, shelfmark, title = self._resolve_meta_labels(item['raw_header'])

        header = f"{shelfmark} | {title} (System ID: {sys_id}, Img: {page or 'N/A'})"
        source_ctx = (item.get('source_ctx', '') or '').strip() or "[No source excerpt available]"
        ms_ctx = (item.get('text', '') or '').strip() or "[No manuscript excerpt available]"
        src_label = item.get('src_lbl', 'Source')

        return "\n".join([
            header,
            f"Source [{sys_id} | {src_label}]:",
            source_ctx,
            f"MS [{sys_id} | Img {page or 'N/A'}]:",
            ms_ctx
        ])

    def _fetch_metadata_with_dialog(self, system_ids, title="Loading metadata..."):

        to_fetch = [sid for sid in system_ids if sid and sid not in self.meta_mgr.nli_cache]
        if not to_fetch:
            return False

        dialog = QProgressDialog(tr("Loading shelfmarks and titles..."), tr("Cancel"), 0, len(to_fetch), self)
        dialog.setWindowTitle(title) 
        dialog.setWindowModality(Qt.WindowModality.ApplicationModal)
        dialog.setAutoClose(False)
        dialog.setAutoReset(False)
        dialog.setMinimumDuration(0)

        loop = QEventLoop(self)
        cancelled = False

        worker = ShelfmarkLoaderThread(self.meta_mgr, to_fetch)

        def on_progress(curr, total, sid):
            dialog.setMaximum(total)
            dialog.setValue(curr)
            dialog.setLabelText(f"Loaded {curr}/{total} (ID: {sid})")

        def on_finished(was_cancelled):
            nonlocal cancelled
            cancelled = was_cancelled
            dialog.reset()
            loop.quit()
            if was_cancelled:
                QMessageBox.information(self, "Metadata", tr("Loading metadata was cancelled."))

        def on_error(err):
            QMessageBox.critical(self, tr("Metadata Error"), err)
            dialog.reset()
            loop.quit()

        def handle_cancel():
            worker.request_cancel()

        dialog.canceled.connect(handle_cancel)
        worker.progress_signal.connect(on_progress)
        worker.finished_signal.connect(on_finished)
        worker.error_signal.connect(on_error)

        worker.start()
        dialog.show()
        loop.exec()
        worker.wait()

        return cancelled

    def _resolve_meta_labels(self, raw_header):
        sid, page, shelf, title = self._get_meta_for_header(raw_header)
        sys_id = sid or "Unknown System ID"

        if sid and sid not in self.meta_mgr.nli_cache:
            meta = self.meta_mgr.fetch_nli_data(sid)
            title = title or (meta.get('title') if meta else None)

        shelf_lbl = shelf or f"[Shelfmark missing for {sys_id}]"
        title_lbl = title or "[Title missing]"

        return sys_id, page, shelf_lbl, title_lbl

    def _update_part_state_for_sid(self, sid):
        """Refresh Part context (Neubauer) for the given system ID."""
        part_id = self.meta_mgr.get_part_for_folio(sid) if self.meta_mgr else None
        if part_id:
            folios = self.meta_mgr.get_folios_for_part(part_id) or []
            self.current_browse_part_id = part_id
            self.current_browse_part_folios = folios
            self.current_browse_part_folio_idx = folios.index(sid) if sid in folios else 0
        else:
            self.current_browse_part_id = None
            self.current_browse_part_folios = []
            self.current_browse_part_folio_idx = 0
        return part_id

    def browse_load(self):
        if not self.searcher: return
        self.browse_highlight_data = []
        self.browse_highlight_pattern = None
        sid = self.browse_sys_input.text().strip()
        shelf_query = self.browse_shelf_input.text().strip()
        fl_id = self.browse_fl_input.text().strip()
        if not sid and not fl_id and not shelf_query: return

        # Reset UI (skip when reading desk is active to preserve stacked view)
        if not self.browse_reading_desk_active:
            self.browse_text.setText(tr("Loading metadata..."))
        if not self.browse_reading_desk_active:
            self.browse_viewer.load_images({})  # Clear viewer

        page_data = None

        # Check if this is a Part identifier (Neubauer)
        part_id, is_part = self.meta_mgr.parse_part_identifier(shelf_query)
        if is_part:
            self._browse_load_part(part_id)
            return

        # Determine priority based on last edited field (default: shelfmark > system ID > FL)
        priority = []
        if self.last_browse_field == "fl" and fl_id:
            priority.append("fl")
        elif self.last_browse_field == "shelf" and shelf_query:
            priority.append("shelf")
        elif self.last_browse_field == "sys" and sid:
            priority.append("sys")

        if shelf_query and "shelf" not in priority:
            priority.append("shelf")
        if sid and "sys" not in priority:
            priority.append("sys")
        if fl_id and "fl" not in priority:
            priority.append("fl")

        def format_option(opt, idx):
            base = opt['shelfmark']
            title = (opt.get('title') or "").strip()
            if title:
                base = f"{base} | {title}"
            label = f"{idx + 1}. {base}"
            if len(label) > 60:
                label = label[:57] + "..."
            return label

        for field in priority:
            if field == "fl":
                if not fl_id:
                    continue
                pd = self.searcher.get_browse_page_by_fl(fl_id, sid or None)
                if pd:
                    page_data = pd
                    sid = pd.get('sys_id', sid)
                    self.browse_sys_input.setText(sid or "")
                    self.browse_fl_input.setText(pd.get('fl_id', fl_id))
                    break
                # If no other identifiers exist, stop and warn
                if not sid and not shelf_query:
                    QMessageBox.warning(self, tr("Error"), tr("FL not found."))
                    return
                continue

            if field == "shelf":
                if not shelf_query:
                    continue
                shelf_res = self.meta_mgr.resolve_system_by_shelfmark(shelf_query)
                if shelf_res['sys_id']:
                    sid = shelf_res['sys_id']
                    if shelf_res['selected_shelfmark']:
                        self.browse_shelf_input.setText(shelf_res['selected_shelfmark'])
                    self.browse_sys_input.setText(sid or "")
                    break
                elif shelf_res['options']:
                    options = shelf_res['options']
                    if len(options) == 1:
                        opt = options[0]
                        sid = opt['sys_id']
                        self.browse_shelf_input.setText(opt['shelfmark'])
                        self.browse_sys_input.setText(sid or "")
                        break
                    display_options = [format_option(opt, idx) for idx, opt in enumerate(options)]
                    choice, ok = QInputDialog.getItem(
                        self, tr("Shelfmark"), tr("Multiple shelfmarks found. Select one:"), display_options, 0, False
                    )
                    if not ok:
                        return
                    if choice in display_options:
                        idx = display_options.index(choice)
                        opt = options[idx]
                        sid = opt['sys_id']
                        self.browse_shelf_input.setText(opt['shelfmark'])
                        self.browse_sys_input.setText(sid or "")
                        break
                # Shelfmark was the chosen path; stop if not resolved
                QMessageBox.warning(self, tr("Error"), tr("Shelfmark not found."))
                return

            if field == "sys":
                if sid:
                    break

        if not sid:
            msg = tr("FL not found.")
            if shelf_query:
                msg = tr("Shelfmark not found.")
            QMessageBox.warning(self, tr("Error"), msg)
            return

        # If this folio belongs to an Oxford Part, load the Part context directly
        part_for_sid = self.meta_mgr.get_part_for_folio(sid)
        if part_for_sid:
            self._browse_load_part(part_for_sid, target_folio=sid)
            return

        self.current_browse_sid = sid
        self.current_browse_p = page_data['p_num'] if page_data else None
        self.current_browse_internal_idx = None

        # Add to Recently Viewed
        if self.lists_mgr:
            fl_id = self._normalize_fl_id(page_data.get('fl_id') if page_data else self.browse_fl_input.text().strip())
            img = page_data.get('p_num') if page_data else None
            self.lists_mgr.add_to_recent(sid, fl_id=fl_id, img=img)

        # Disable controls until loaded
        self.btn_b_catalog.setEnabled(False)
        self.btn_b_save.setEnabled(False)
        self.btn_b_toggle_img.setEnabled(False)

        # Trigger Unified Enrichment (Meta + Images)
        # Disconnect old worker to prevent stale signals and GC crash
        if hasattr(self, 'enrich_browse_worker') and self.enrich_browse_worker is not None:
            try:
                self.enrich_browse_worker.finished_signal.disconnect(self.on_browse_enriched_loaded)
            except (TypeError, RuntimeError):
                pass
        self.enrich_browse_worker = EnrichMetadataThread(self.meta_mgr, sid)
        self.enrich_browse_worker.finished_signal.connect(self.on_browse_enriched_loaded)
        self.enrich_browse_worker.start()

        # Try to render page text immediately if possible
        if self.browse_reading_desk_active:
            # Reading desk is active: instead of rendering single-page view,
            # add the newly resolved manuscript to the reading desk.
            sid = self.current_browse_sid
            shelfmark, _ = self.meta_mgr.get_meta_for_id(sid)
            if not shelfmark or shelfmark == "Unknown":
                shelfmark = sid
            self._browse_rd_add_entry(sid, shelfmark)
        elif page_data:
            self.browse_render_page(page_data)
        elif self.current_browse_sid:
            self.browse_load_page()

    def _browse_load_part(self, part_id, from_end=False, target_folio=None):
        """
        Load a Codicological Part (Neubauer) for browsing.

        Args:
            part_id: The Part ID to load
            from_end: If True, start at the last folio (for backwards navigation)
            target_folio: If provided, position at this specific folio within the Part
        """
        if not self.searcher or not self.meta_mgr:
            return

        # Get all folios in this Part
        folios = self.meta_mgr.get_folios_for_part(part_id)
        if not folios:
            QMessageBox.warning(self, tr("Error"), tr("No folios found for this Part."))
            return

        # Get Part metadata from Oxford
        part_meta = self.meta_mgr.get_part_metadata(part_id)

        # Store Part browsing state
        self.current_browse_part_id = part_id
        self.current_browse_part_folios = folios

        # Determine starting folio index
        if target_folio and target_folio in folios:
            folio_idx = folios.index(target_folio)
        elif from_end:
            folio_idx = len(folios) - 1
        else:
            folio_idx = 0

        self.current_browse_part_folio_idx = folio_idx

        # Load the selected folio
        target_sid = folios[folio_idx]
        self.current_browse_sid = target_sid
        self.current_browse_p = None
        self.current_browse_internal_idx = None

        # Update UI fields
        self.browse_sys_input.setText(target_sid)
        # Use the display format: "heb. d. 29 part 2"
        display_name = self.meta_mgr.codico_mgr.get_part_display_name(part_id)
        self.browse_shelf_input.setText(display_name)

        # Display Part metadata in info label
        part_title = part_meta.get('title', '') if part_meta else ''
        part_contents = part_meta.get('contents', '') if part_meta else ''

        # Get our CSV title if available (for Hebrew)
        shelf, csv_title = self.meta_mgr.get_meta_for_id(target_sid)

        # Add library prefix
        library_code = self.meta_mgr.get_library_for_id(target_sid)
        if library_code:
            library = get_library_display(library_code, short=False)
            shelf = f"{library} | {shelf}" if shelf else library

        # Format folio range info
        folio_range = part_meta.get('folio_range', []) if part_meta else []

        # Build info label with Part info integrated into shelfmark
        # Extract part number from part_id (e.g., "MS. Heb. b. 10/43" -> "43")
        part_num = part_id.split('/')[-1] if '/' in part_id else ''

        # Build combined: "MS heb. b. 10/79 (part 43: fols. 79-82)"
        shelf_with_part = f"{shelf or target_sid}"
        if part_num:
            shelf_with_part += f" (part {part_num}"
            if len(folio_range) == 2:
                if folio_range[0] == folio_range[1]:
                    shelf_with_part += f": fol. {folio_range[0]}"
                else:
                    shelf_with_part += f": fols. {folio_range[0]}–{folio_range[1]}"
            shelf_with_part += ")"

        info_text = f"<b>{shelf_with_part}</b>"
        tooltip_parts = []
        if part_title:
            truncated, full = _truncate_title(part_title)
            info_text += f"<br/><span style='font-size: 11px;'>{truncated}</span>"
            if full: tooltip_parts.append(full)
        if csv_title and csv_title != part_title:
            truncated, full = _truncate_title(csv_title)
            info_text += f"<br/>{truncated}"
            if full: tooltip_parts.append(full)

        self.browse_info_lbl.setText(info_text)
        self.browse_info_lbl.setToolTip('\n'.join(tooltip_parts) if tooltip_parts else '')

        # Disable controls until loaded
        self.btn_b_catalog.setEnabled(False)
        self.btn_b_save.setEnabled(False)
        self.btn_b_toggle_img.setEnabled(False)

        # Load images from Part directly (Oxford images)
        part_images = self.meta_mgr.get_part_images(part_id)
        part_meta = self.meta_mgr.get_part_metadata(part_id)
        if part_images:
            # Convert Part images to format expected by viewer (include thumb_url)
            images_ext = [{
                'label': img.get('label', ''),
                'url': img.get('full_url', ''),
                'thumb_url': img.get('thumb_url', ''),
                'folio_num': img.get('folio_num')
            } for img in part_images]
            # Get target folio number first
            shelf_for_folio, _ = self.meta_mgr.get_meta_for_id(target_sid)
            folio_num = _get_folio_number_from_shelfmark(shelf_for_folio)
            image_idx = _get_initial_image_index({'images_ext': images_ext}, folio_num)
            # Load images with target_folio for dynamic generation if needed
            self.browse_viewer.load_images({
                'images_nli': [],
                'images_ext': images_ext,
                'oxford_part_id': part_id,
                'oxford_part_metadata': part_meta,  # Include for dynamic image generation
                'attribution': "From the collections of the Bodleian Libraries, Oxford",
            }, initial_idx=image_idx, target_folio=folio_num)

        # Load text for the selected folio
        self.browse_load_page()

        # Enable controls
        self.btn_b_catalog.setEnabled(True)
        self.btn_b_save.setEnabled(True)
        self.btn_b_toggle_img.setEnabled(True)
        self.btn_find_parallels.setEnabled(True)
        self.btn_browse_add_to_list.setEnabled(True)

        # Trigger metadata enrichment for Oxford Part manuscript (Phase 33 gap closure)
        # ALWAYS start thread unconditionally -- enrich_metadata handles cache internally
        # and builds on top of basic CSV metadata with NLI crossref and FJMS data.
        target_sid_for_enrich = self.current_browse_sid
        if hasattr(self, 'enrich_browse_worker') and self.enrich_browse_worker is not None:
            try:
                self.enrich_browse_worker.finished_signal.disconnect(self.on_browse_enriched_loaded)
            except (TypeError, RuntimeError):
                pass
        self.enrich_browse_worker = EnrichMetadataThread(self.meta_mgr, target_sid_for_enrich)
        self.enrich_browse_worker.finished_signal.connect(self.on_browse_enriched_loaded)
        self.enrich_browse_worker.start()

    def browse_navigate(self, d):
        if not self.current_browse_sid: return

        idx_arg = self.current_browse_internal_idx
        
        p_arg = None
        if self.current_browse_p is not None:
            try: p_arg = int(self.current_browse_p)
            except: p_arg = 0

        page_data = self.searcher.get_browse_page(
            self.current_browse_sid, 
            p_num=p_arg, 
            next_prev=d,
            absolute_index=idx_arg,
            allow_cross=True
        )

        if not page_data:
            QMessageBox.warning(self, tr("Nav"), tr("Not found or end."))
            return

        new_sid = page_data.get('sys_id', self.current_browse_sid)
        is_new_manuscript = new_sid != self.current_browse_sid
        if is_new_manuscript:
            self.current_browse_sid = new_sid
            self.browse_sys_input.setText(new_sid)
            shelf, _ = self.meta_mgr.get_meta_for_id(new_sid)
            if shelf and shelf != "Unknown":
                self.browse_shelf_input.setText(shelf)
            self._set_last_browse_field("sys")
            # Refresh Part context
            self._update_part_state_for_sid(new_sid)
            # Render page FIRST so text is ready before community status stores it
            self.browse_render_page(page_data)
            # Flag: page already rendered, skip browse_load_page in enriched callback
            self._browse_nav_rendered = True
            # Trigger metadata enrichment for new manuscript
            # ALWAYS start thread unconditionally -- enrich_metadata handles cache internally
            if hasattr(self, 'enrich_browse_worker') and self.enrich_browse_worker is not None:
                try:
                    self.enrich_browse_worker.finished_signal.disconnect(self.on_browse_enriched_loaded)
                except (TypeError, RuntimeError):
                    pass
            self.enrich_browse_worker = EnrichMetadataThread(self.meta_mgr, new_sid)
            self.enrich_browse_worker.finished_signal.connect(self.on_browse_enriched_loaded)
            self.enrich_browse_worker.start()
        else:
            self.browse_render_page(page_data)
            # Re-fetch PGP for same-manuscript page navigation (new manuscript handled by enriched_loaded)
            self._browse_refresh_pgp_for_page()
    
    def browse_render_page(self, pd):
        if pd.get('sys_id') and pd.get('sys_id') != self.current_browse_sid:
            self.current_browse_sid = pd['sys_id']

        self.current_browse_p = pd['p_num']
        
        if 'internal_index' in pd:
            self.current_browse_internal_idx = pd['internal_index']
        else:
            self.current_browse_internal_idx = pd.get('current_idx', 1) - 1

        # If we are rendering a page, we are not in View All mode
        if self.btn_b_all.isChecked():
            self.btn_b_all.blockSignals(True)
            self.btn_b_all.setChecked(False)
            self.btn_b_all.blockSignals(False)
            self.btn_b_toggle_img.setEnabled(True)
            self.browse_viewer.setVisible(self.btn_b_toggle_img.isChecked())

        # Enable core buttons immediately when content is available
        self.btn_find_parallels.setEnabled(True)
        self.btn_browse_add_to_list.setEnabled(True)
        self.btn_b_save.setEnabled(True)
        self.btn_b_all.setEnabled(True)
        if self.current_browse_sid:
            self.btn_b_catalog.setEnabled(True)

        # Apply highlights (manual spans first, then pattern)
        page_text = pd['text']
        page_text = self._apply_browse_highlights(page_text, pd.get('uid'))
        if self.browse_highlight_pattern:
            try:
                regex = re.compile(self.browse_highlight_pattern, re.IGNORECASE)
                page_text = regex.sub(r'*\g<0>*', page_text)
            except Exception:
                pass
        browse_html_text = page_text.replace('\n', '<br>')
        self.browse_text.setHtml(f"<div dir='rtl'>{browse_html_text}</div>")
        apply_find_highlight(self.browse_text, self.browse_find_input.text().strip())
        
        full_header = pd.get('full_header', '')
        _, _, shelf, title = self._get_meta_for_header(full_header)

        # Create display versions: full name for info label, short code for input field
        display_shelf = shelf  # For info label (full library name)
        input_shelf = shelf    # For input field (short library code)
        library_code = None
        if shelf and self.current_browse_sid:
            library_code = self.meta_mgr.get_library_for_id(self.current_browse_sid)
            if library_code:
                library_full = get_library_display(library_code, short=False)
                display_shelf = f"{library_full} | {shelf}"
                input_shelf = f"{library_code} {shelf}"

        # Add Oxford Part info if available - integrated into shelfmark
        part_id = self.current_browse_part_id
        if not part_id and self.current_browse_sid:
            part_id = self.meta_mgr.get_part_for_folio(self.current_browse_sid)

        if part_id:
            part_meta = self.meta_mgr.get_part_metadata(part_id)
            folio_range = part_meta.get('folio_range', []) if part_meta else []
            oxford_title = part_meta.get('title', '') if part_meta else ''

            # Extract part number from part_id (e.g., "MS. Heb. b. 10/43" -> "43")
            part_num = part_id.split('/')[-1] if '/' in part_id else ''

            # Build combined shelfmark: "MS heb. b. 10/79 (part 43: fols. 79-82)"
            shelf_with_part = f"{display_shelf}"
            if part_num:
                shelf_with_part += f" (part {part_num}"
                if len(folio_range) == 2:
                    if folio_range[0] == folio_range[1]:
                        shelf_with_part += f": fol. {folio_range[0]}"
                    else:
                        shelf_with_part += f": fols. {folio_range[0]}–{folio_range[1]}"
                shelf_with_part += ")"

            info_text = f"<b>{shelf_with_part}</b>"
            tooltip_parts = []
            if oxford_title:
                truncated, full = _truncate_title(oxford_title)
                info_text += f"<br/><span style='font-size: 11px;'>{truncated}</span>"
                if full: tooltip_parts.append(full)
            if title and title != oxford_title:
                truncated, full = _truncate_title(title)
                info_text += f"<br/>{truncated}"
                if full: tooltip_parts.append(full)
        else:
            info_text = f"<b>{display_shelf}</b>"
            tooltip_parts = []
            if title:
                truncated, full = _truncate_title(title)
                info_text += f"<br/>{truncated}"
                if full: tooltip_parts.append(full)

        # Append enrichment metadata from nli_cache (survives re-render)
        cached_meta = self.meta_mgr.nli_cache.get(self.current_browse_sid, {})
        catalog_entry = cached_meta.get('catalog_entry')
        if catalog_entry:
            info_text += f" | {catalog_entry}"
        self.browse_info_lbl.setText(info_text)
        self.browse_info_lbl.setToolTip('\n'.join(tooltip_parts) if tooltip_parts else '')
        if input_shelf:
            self.browse_shelf_input.setText(input_shelf)

        # Update Combo with folio labels when available (Phase 31)
        total = pd['total_pages']
        curr_idx = pd['current_idx'] # 1-based index

        # Get folio images from enriched metadata (if available)
        # Only use crossref folio labels when the image count matches the
        # page count from the search index.  When they differ the labels
        # would map to the wrong pages (e.g. crossref starts at leaf 4
        # while search-index pages start at 1).
        folio_images = self._browse_folio_images
        has_folio_labels = (
            bool(folio_images)
            and len(folio_images) > 0
            and len(folio_images) == total
        )

        self.combo_browse_page.blockSignals(True)
        # Always repopulate — stale labels persist when consecutive MSs have same page count
        self.combo_browse_page.clear()
        if has_folio_labels:
            # Use folio labels from crossref service
            for i, img in enumerate(folio_images):
                label = img.get('folio_label', str(i + 1))
                self.combo_browse_page.addItem(label)
        else:
            items = [str(i) for i in range(1, total + 1)]
            self.combo_browse_page.addItems(items)

        if 0 < curr_idx <= total:
            self.combo_browse_page.setCurrentIndex(curr_idx - 1)
        self.combo_browse_page.blockSignals(False)

        # Update folio label display (Phase 31)
        folio_label_text = ''
        if has_folio_labels and 0 < curr_idx <= len(folio_images):
            folio_label_text = folio_images[curr_idx - 1].get('folio_label', '')
        if folio_label_text:
            self.lbl_browse_folio.setText(f"{tr('Folio')} {folio_label_text}")
            self.lbl_browse_folio.setVisible(True)
        else:
            self.lbl_browse_folio.setVisible(False)

        # Update page count label (Phase 31)
        effective_count = len(folio_images) if has_folio_labels else total
        if effective_count > 0:
            self.lbl_browse_page_count.setText(f"{tr('of')} {effective_count} {tr('pages')}")
            self.lbl_browse_page_count.setVisible(True)
        else:
            self.lbl_browse_page_count.setVisible(False)

        self.btn_b_prev.setEnabled(pd['current_idx'] > 1)
        self.btn_b_next.setEnabled(pd['current_idx'] < pd['total_pages'])
        self._update_browse_add_to_list_button()

        parsed = self.meta_mgr.parse_full_id_components(full_header)
        if parsed.get('fl_id'):
            self.browse_fl_input.setText(f"FL{parsed['fl_id']}")
        else:
            self.browse_fl_input.setText("")
            
        # This tells the large image viewer on the right to jump to the correct index
        if hasattr(self, 'browse_viewer') and self.browse_viewer.isVisible():
            meta = self.meta_mgr.nli_cache.get(self.current_browse_sid, {})
            shelfmark, _ = self.meta_mgr.get_meta_for_id(self.current_browse_sid)
            folio_num = _get_folio_number_from_shelfmark(shelfmark)
            side_offset = 1 if (self.current_browse_internal_idx or 0) % 2 == 1 else 0

            # Check if folio needs dynamic images (missing from viewer but within Oxford folio_range)
            viewer_images = getattr(self.browse_viewer, 'images_ext', [])
            folio_in_viewer = any(img.get('folio_num') == folio_num for img in viewer_images) if folio_num and viewer_images else False

            if not folio_in_viewer and folio_num is not None and meta:
                # Check if this is Oxford with folio_range that includes this folio
                oxford_part_meta = meta.get('oxford_part_metadata', {})
                folio_range = oxford_part_meta.get('folio_range', [])
                if meta.get('oxford_part_id') and len(folio_range) >= 2 and folio_range[0] <= folio_num <= folio_range[1]:
                    # Need dynamic images - call load_images
                    idx = _get_folio_image_index(meta, folio_num, side_offset=side_offset)
                    self.browse_viewer.load_images(meta, idx, target_folio=folio_num)
                elif viewer_images:
                    idx = _get_folio_image_index({'images_ext': viewer_images}, folio_num if folio_num is not None else self.current_browse_p, side_offset=side_offset)
                    self.browse_viewer.set_page(idx)
            elif viewer_images:
                idx = _get_folio_image_index({'images_ext': viewer_images}, folio_num if folio_num is not None else self.current_browse_p, side_offset=side_offset)
                self.browse_viewer.set_page(idx)
        # -------------------------------

        if self.current_browse_sid in self.meta_mgr.nli_cache:
            self.fetch_browse_thumbnail(self.current_browse_sid)
        else:
            self.browse_thumb.setText(tr("Loading Meta..."))
            def worker():
                self.meta_mgr.fetch_nli_data(self.current_browse_sid)
                self.browse_thumb_resolved.emit(self.current_browse_sid, "") 
            threading.Thread(target=worker, daemon=True).start()
        
    def browse_open_catalog(self):
        if self.current_browse_sid:
            QDesktopServices.openUrl(QUrl(f"https://www.nli.org.il/he/discover/manuscripts/hebrew-manuscripts/itempage?vid=KTIV&scope=KTIV&docId=PNX_MANUSCRIPTS{self.current_browse_sid}"))

    def _on_browse_thumb_resolved(self, sid, _unused_url):
        if sid != self.current_browse_sid:
            return
            
        self.fetch_browse_thumbnail(sid)

    def start_browse_download(self, sid, thumb_url):
        if sid != self.current_browse_sid:
            return

        self.browse_thumb_url = thumb_url
        self.cancel_browse_image_thread()

        if not thumb_url:
            self.on_browse_img_failed()
            return

        self.browse_img_thread = ImageLoaderThread(thumb_url)
        self.browse_img_thread.image_loaded.connect(self.on_browse_img_loaded)
        self.browse_img_thread.load_failed.connect(self.on_browse_img_failed)
        self.browse_img_thread.start()

    def on_browse_img_loaded(self, image):
        pix = QPixmap.fromImage(image)
        # Scale carefully to avoid distortion
        if not pix.isNull():
            scaled = pix.scaled(self.browse_thumb.size(), Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
            self.browse_thumb.setPixmap(scaled)
            self.browse_thumb.setText("")
        else:
            self.on_browse_img_failed()

    def on_browse_img_failed(self):
        self.browse_thumb.setPixmap(QPixmap())
        self.browse_thumb.setText("No Preview")

    def cancel_browse_image_thread(self):
        if getattr(self, 'browse_img_thread', None) and self.browse_img_thread.isRunning():
            self.browse_img_thread.cancel()
            # Use short timeout to avoid blocking UI - thread will finish in background
            self.browse_img_thread.wait(500)

    def fetch_browse_thumbnail(self, sys_id, meta=None):
        self.cancel_browse_image_thread()
        self.browse_thumb.setText("Loading...")
        self.browse_thumb.setPixmap(QPixmap())

        # Load from cache if not provided
        meta = meta or self.meta_mgr.nli_cache.get(sys_id)
        
        # In genizah_core, we now guarantee that 'thumb_url' comes from 907 $d if available
        thumb_url = meta.get('thumb_url') if meta else None

        if thumb_url:
            self.start_browse_download(sys_id, thumb_url)
        else:
            # If metadata exists but no thumb_url, it means no image at all
            if meta:
                self.browse_thumb.setText(tr("No Image"))
            else:
                self.browse_thumb.setText(tr("Waiting..."))
    
    def check_updates_auto(self):
        """Run update checker silently at startup."""
        self.update_thread = UpdateCheckerThread(APP_VERSION, is_manual=False)
        self.update_thread.finished_signal.connect(self.on_update_result)
        self.update_thread.start()

        # Also check for sidecar data updates
        self.sidecar_update_thread = SidecarUpdateThread()
        self.sidecar_update_thread.update_available.connect(self._on_sidecar_updates)
        self.sidecar_update_thread.start()

    def check_updates_manual(self):
        """Run update checker with UI feedback."""
        self.btn_check_updates.setEnabled(False)
        self.btn_check_updates.setText(tr("Checking..."))

        self.update_thread = UpdateCheckerThread(APP_VERSION, is_manual=True)
        self.update_thread.finished_signal.connect(self.on_update_result)
        self.update_thread.error_signal.connect(self.on_update_error)
        self.update_thread.start()

    def on_update_result(self, found, version, html_url, installer_url, is_manual):
        # Reset manual button state
        if is_manual:
            self.btn_check_updates.setEnabled(True)
            self.btn_check_updates.setText(tr("Check for Updates"))

        if found:
            # Check if dismissed previously (only for auto check)
            if not is_manual:
                cfg = load_app_config()
                last_dismissed = cfg.get('last_dismissed_version')
                if last_dismissed == version:
                    return # Silent return if already dismissed

            # Show Feedback
            if is_manual:
                # Dialog for manual - offer in-app update
                msg = tr("A new version is available: {}").format(version)
                reply = QMessageBox.question(
                    self, tr("Update Available"),
                    msg + "\n\n" + tr("Would you like to update now?"),
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                )
                if reply == QMessageBox.StandardButton.Yes:
                    self.start_in_app_update(version, html_url, installer_url)
            else:
                # Notification Bar for auto
                self.update_bar.show_update(version, html_url, installer_url)

        else:
            if is_manual:
                QMessageBox.information(self, tr("Up to date"), tr("You are using the latest version."))

    def on_update_error(self, err, is_manual):
        if is_manual:
            self.btn_check_updates.setEnabled(True)
            self.btn_check_updates.setText(tr("Check for Updates"))
            QMessageBox.warning(self, tr("Update Error"), err)
        else:
            logger.warning(f"Auto-update check failed: {err}")

    def _on_sidecar_updates(self, updates):
        """Handle sidecar update availability notification."""
        if not updates:
            return

        # Build human-readable message
        lines = []
        total_mb = 0
        for u in updates:
            lines.append(f"  - {u['name']}: {u['current']} \u2192 {u['available']} ({u['size_mb']} MB)")
            total_mb += u['size_mb']

        msg = tr("New research data available:") + "\n\n"
        msg += "\n".join(lines)
        msg += f"\n\n{tr('Total download:')} {total_mb} MB"
        msg += "\n\n" + tr("Download now?")

        reply = QMessageBox.question(
            self, tr("Data Update Available"), msg,
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            self._start_sidecar_download(updates)

    def _start_sidecar_download(self, updates):
        """Download sidecar updates sequentially."""
        import os
        # Use LOCALAPPDATA for updated sidecars (bundled location is read-only)
        data_dir = os.path.join(
            os.environ.get('LOCALAPPDATA', os.path.expanduser('~')),
            'GenizahSearchPro', 'data'
        )
        self._sidecar_download_queue = list(updates)
        self._sidecar_data_dir = data_dir
        self._download_next_sidecar()

    def _download_next_sidecar(self):
        """Download the next sidecar in the queue."""
        import os
        if not self._sidecar_download_queue:
            # All downloads complete -- reset services to pick up new files
            from shared.document_service import reset_pgp_service
            from shared.fjms_service import reset_fjms_service
            from shared.nli_crossref_service import reset_nli_crossref_service
            reset_pgp_service()
            reset_fjms_service()
            reset_nli_crossref_service()
            QMessageBox.information(
                self, tr("Update Complete"),
                tr("Research data has been updated. New data will be used immediately.")
            )
            return

        update = self._sidecar_download_queue.pop(0)
        target = os.path.join(self._sidecar_data_dir, update['subdir'], update['name'])
        self._current_sidecar_download = SidecarDownloadThread(update['url'], target, update['name'])
        self._current_sidecar_download.finished_signal.connect(self._on_sidecar_download_finished)
        self._current_sidecar_download.start()

    def _on_sidecar_download_finished(self, success, result, sidecar_name):
        """Handle completion of a single sidecar download."""
        if success:
            logger.info(f"Sidecar updated: {sidecar_name} -> {result}")
        else:
            logger.warning(f"Sidecar download failed: {sidecar_name}: {result}")
        self._download_next_sidecar()

    def on_update_dismissed(self, version):
        """Save dismissed version to config."""
        save_app_config({'last_dismissed_version': version})

    def on_whats_new_dismissed(self):
        """Save that user has seen What's New for this version."""
        save_app_config({'whats_new_seen': APP_VERSION})

    def show_whats_new_dialog(self):
        """Show detailed What's New dialog."""
        dlg = WhatsNewDialog(self)
        dlg.exec()
        self.on_whats_new_dismissed()
        self.whats_new_bar.hide()

    def start_in_app_update(self, version: str, html_url: str, installer_url: str):
        """Start the in-app update process with progress dialog."""
        # Hide the notification bar if visible
        self.update_bar.hide()

        # Create and show the update progress dialog
        # Store reference to prevent garbage collection
        self.update_dialog = UpdateProgressDialog(self, version, installer_url, html_url)
        self.update_dialog.show()

        # Start the download after dialog is shown
        QTimer.singleShot(100, self.update_dialog.start_download)

    def run_indexing(self):
        # 1. Pre-check: Does the input file exist?
        if not os.path.exists(Config.FILE_V8):
            msg = tr("The transcriptions file ('Transcriptions.txt') was not found in the application folder.") + "\n\n" + \
                  tr("Would you like to locate it manually?")
            
            reply = QMessageBox.question(
                self, 
                tr("File Not Found"), 
                msg, 
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )

            if reply == QMessageBox.StandardButton.Yes:
                # Open File Dialog to let user pick the file
                path, _ = QFileDialog.getOpenFileName(
                    self, 
                    tr("Select Transcriptions File"), 
                    "", 
                    "Text Files (*.txt);;All Files (*)"
                )
                
                if path:
                    # Update Config dynamically for this session
                    Config.FILE_V8 = path
                else:
                    # User cancelled the file dialog
                    return
            else:
                # User clicked No
                return

        # 2. Standard Indexing Confirmation
        if not self.indexer: return
        
        if QMessageBox.question(self, tr("Index"), tr("Start indexing?"), QMessageBox.StandardButton.Yes|QMessageBox.StandardButton.No) == QMessageBox.StandardButton.Yes:
            self.index_progress.setRange(0, 1)
            self.index_progress.setValue(0)
            self.index_progress.setFormat(tr("Indexing... %p%"))
            
            self.ithread = IndexerThread(self.meta_mgr)
            self.ithread.progress_signal.connect(self.on_index_progress)
            self.ithread.finished_signal.connect(self.on_index_finished)
            self.ithread.error_signal.connect(self.on_index_error)
            self.ithread.start()

    def on_index_progress(self, current, total):
        self.index_progress.setRange(0, max(total, 1))
        self.index_progress.setValue(current)
        self.index_progress.setFormat(f"{current}/{total} lines")

    def on_index_finished(self, total_docs):
        self.index_progress.setValue(self.index_progress.maximum())
        self.index_progress.setFormat(tr("Indexing complete"))
        self.searcher.reload_index()
        QMessageBox.information(self, tr("Done"), tr("Indexing complete. Documents indexed: {}").format(total_docs))

    def on_index_error(self, err):
        self.index_progress.setFormat(tr("Indexing failed"))

        # Check for file locking error (caused by file managers, antivirus, etc.)
        err_lower = str(err).lower()
        if "index writer" in err_lower or "io:error" in err_lower or "worker thread" in err_lower:
            msg = tr("Indexing failed - another program is blocking file access.") + "\n\n" + \
                  tr("Solution:") + "\n" + \
                  tr("• Close file manager programs (Total Commander, Directory Opus, etc.)") + "\n" + \
                  tr("• Close Windows Explorer windows open to the application folder") + "\n" + \
                  tr("• Try again")
            QMessageBox.critical(self, tr("Indexing Error"), msg)
        else:
            QMessageBox.critical(self, tr("Indexing Error"), str(err))

    def closeEvent(self, event):
        # Ensure worker threads are stopped before the window is destroyed
        try:
            if getattr(self, 'meta_loader', None) and self.meta_loader.isRunning():
                self.meta_loader.request_cancel()
                self.meta_loader.wait()

            if getattr(self, 'search_thread', None) and self.search_thread.isRunning():
                self.search_thread.requestInterruption()
                self.search_thread.wait(2000)
                if self.search_thread.isRunning():
                    self.search_thread.terminate()
                    self.search_thread.wait()

            if getattr(self, 'comp_thread', None) and self.comp_thread.isRunning():
                self.comp_thread.requestInterruption()
                self.comp_thread.wait(2000)
                if self.comp_thread.isRunning():
                    self.comp_thread.terminate()
                    self.comp_thread.wait()

            if getattr(self, 'group_thread', None) and self.group_thread.isRunning():
                self.group_thread.requestInterruption()
                self.group_thread.wait(2000)
                if self.group_thread.isRunning():
                    self.group_thread.terminate()
                    self.group_thread.wait()

            # Stop browse tab viewer image threads
            if getattr(self, 'browse_viewer', None):
                self.browse_viewer.stop_threads()
        finally:
            super().closeEvent(event)

    def _add_single_comp_node(self, parent, ms_item):
        """Adds a node to the composition tree with parent/child logic."""
        item_type = ms_item.get('type', 'manuscript')
        is_part = item_type == 'part'

        sid = ms_item.get('sys_id')
        if not sid:
            sid, _ = self.meta_mgr.parse_header_smart(ms_item.get('raw_header', ''))

        # For Parts, use Part display name; for regular manuscripts, use shelfmark
        if is_part:
            part_display = ms_item.get('part_display', '')
            oxford_title = ms_item.get('oxford_title', '')
            shelf, t = self.meta_mgr.get_meta_for_id(sid)
            # Combine Part display with title
            display_shelf = f"📖 {part_display}" if part_display else (shelf or sid or "Loading...")
            # Prefer Oxford title, fallback to CSV title
            display_title = oxford_title if oxford_title else (t or "")
        else:
            shelf, t = self.meta_mgr.get_meta_for_id(sid)
            display_shelf = shelf if shelf and shelf != "Unknown" else (sid if sid else "Loading...")
            display_title = t or ""

        # Get library info
        library_code = self.meta_mgr.get_library_for_id(sid)
        library_display = get_library_display(library_code, short=False) if library_code else ""

        pages = ms_item.get('pages', [])
        best_snippet = ""
        best_ctx = ""
        if pages:
            best_snippet = pages[0].get('text', '')
            best_ctx = pages[0].get('source_ctx', '')

        node = QTreeWidgetItem(parent)
        self._set_comp_tree_text(node, 0, str(int(ms_item.get('score', 0))))
        self._set_comp_tree_text(node, self.comp_col_shelfmark, display_shelf)
        self._set_comp_tree_text(node, self.comp_col_library, library_display)
        self._set_comp_tree_text(node, self.comp_col_title, display_title)
        self._set_comp_tree_text(node, self.comp_col_sysid, ms_item.get('part_id', '') if is_part else sid)

        node.setFlags(node.flags() | Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled | Qt.ItemFlag.ItemIsSelectable)
        node.setCheckState(0, Qt.CheckState.Unchecked)
        node.setData(0, Qt.ItemDataRole.UserRole, ms_item)

        node.setData(1, Qt.ItemDataRole.UserRole, (best_ctx, best_snippet))

        pattern = pages[0].get('highlight_pattern') if pages else None
        self._set_comp_node_previews(node, best_ctx, best_snippet, pattern)

        # For Parts with multiple folios, show folio count; otherwise show match count
        if is_part:
            folios = ms_item.get('folios', [])
            if len(folios) > 1:
                node.setText(1, f"{display_shelf} ({len(folios)} folios, {len(pages)} matches)")
            elif len(pages) > 1:
                node.setText(1, f"{display_shelf} ({len(pages)} matches)")

        if len(pages) > 1:
            if not is_part:
                node.setText(1, f"{display_shelf} ({len(pages)} matches)")
            for p_item in pages:
                p_num_str = "Page Match"
                raw_h = p_item.get('raw_header', '')
                match = re.search(r'(?i)Img\s*(\d+)', raw_h)
                if match:
                    p_num_str = f"Image {match.group(1)}"
                else:
                    _, p_num_ex, _, _ = self._get_meta_for_header(raw_h)
                    if p_num_ex: p_num_str = f"Image {p_num_ex}"

                # For Parts, also show which folio this match is from
                if is_part:
                    p_sid, _ = self.meta_mgr.parse_header_smart(raw_h)
                    if p_sid:
                        p_shelf, _ = self.meta_mgr.get_meta_for_id(p_sid)
                        if p_shelf and p_shelf != "Unknown":
                            p_num_str = f"{p_shelf} - {p_num_str}"

                child = QTreeWidgetItem(node)
                self._set_comp_tree_text(child, 0, str(int(p_item.get('score', 0))))
                child.setText(1, p_num_str)
                child.setData(0, Qt.ItemDataRole.UserRole, p_item)
                self._set_comp_node_previews(child, p_item.get('source_ctx', ''), p_item.get('text', ''), p_item.get('highlight_pattern'))

        elif len(pages) == 1:
            p_item = pages[0]
            p_suffix = ""
            match = re.search(r'(?i)Img\s*(\d+)', p_item.get('raw_header', ''))
            if match: p_suffix = f" (Img {match.group(1)})"
            node.setText(1, f"{display_shelf}{p_suffix}")

    def _on_comp_item_expanded(self, item):
        if item.childCount() > 0:
            self.comp_tree.setItemWidget(item, self.comp_col_context, None) 
            self.comp_tree.setItemWidget(item, self.comp_col_ms_context, None) 

    def _on_comp_item_collapsed(self, item):
        if item.childCount() > 0:
            stored_data = item.data(1, Qt.ItemDataRole.UserRole)
            if stored_data:
                ctx, snippet = stored_data

                # Try to retrieve highlight pattern from the main item data (Role 0)
                item_data = item.data(0, Qt.ItemDataRole.UserRole)
                pattern = None
                if item_data:
                    if item_data.get('type') in ('manuscript', 'part'):
                        pages = item_data.get('pages', [])
                        if pages:
                            pattern = pages[0].get('highlight_pattern')
                    else:
                        pattern = item_data.get('highlight_pattern')

                self._set_comp_node_previews(item, ctx, snippet, pattern)

    def on_comp_item_double_clicked(self, item, column):
        """
        Smart navigation that restores full context (Next/Prev, Source Text).
        It rebuilds the full list of results from the tree but jumps to the specific clicked item.
        """
        data = item.data(0, Qt.ItemDataRole.UserRole)
        if not data: return 

        flat_list = []
        target_index = -1
        
        clicked_node = item
        
        if data.get('type') in ('manuscript', 'part') and item.childCount() > 0:
            clicked_node = item.child(0)

        def collect_node_data(node):
            node_data = node.data(0, Qt.ItemDataRole.UserRole)
            if not node_data: return

            if node_data.get('type') in ('manuscript', 'part') and node.childCount() > 0:
                for i in range(node.childCount()):
                    collect_node_data(node.child(i))
                return

            raw_h = node_data.get('raw_header', '')
            sid, p_num, shelf, title = self._get_meta_for_header(raw_h)
            
            hl_pattern = node_data.get('highlight_pattern')
            
            ready_item = {
                'uid': node_data.get('uid', sid),
                'raw_header': raw_h,
                'text': node_data.get('text', ''),
                'full_text': None, 
                'source_ctx': node_data.get('source_ctx', ''),
                'highlight_pattern': hl_pattern,
                'display': {
                    'id': sid,
                    'shelfmark': shelf,
                    'title': title,
                    'img': p_num,
                    'source': node_data.get('src_lbl', 'Genizah Lab')
                }
            }
            
            flat_list.append(ready_item)
            
            if node is clicked_node:
                nonlocal target_index
                target_index = len(flat_list) - 1

        root = self.comp_tree.invisibleRootItem()
        for i in range(root.childCount()):
            category = root.child(i)
            if category.data(0, Qt.ItemDataRole.UserRole):
                 collect_node_data(category)
            
            for j in range(category.childCount()):
                sub = category.child(j)
                collect_node_data(sub)

        if not flat_list: return
        
        if target_index == -1: target_index = 0

        try:
            dlg = ResultDialog(self, flat_list, target_index, self.meta_mgr, self.searcher)
            dlg.exec()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to open viewer: {e}")

    def navigate_manuscript(self, direction):
        """Navigate to prev/next manuscript by file order, crossing Part boundaries."""
        current = self.current_browse_sid
        if not current:
            return

        # If browsing within a Part, navigate within the Part first
        if self.current_browse_part_id and self.current_browse_part_folios:
            new_idx = self.current_browse_part_folio_idx + direction
            if 0 <= new_idx < len(self.current_browse_part_folios):
                # Move within Part
                self.current_browse_part_folio_idx = new_idx
                new_sid = self.current_browse_part_folios[new_idx]
                self.current_browse_sid = new_sid
                self.browse_sys_input.setText(new_sid)
                shelf, _ = self.meta_mgr.get_meta_for_id(new_sid)
                if shelf and shelf != "Unknown":
                    self.browse_shelf_input.setText(shelf)
                self._set_last_browse_field("sys")
                self.browse_load_page()
                # Update images to show current folio's pages
                self._update_part_image_for_folio(new_idx)
                return
            else:
                # Crossed Part boundary - move to adjacent Part
                current_part = self.current_browse_part_id
                adjacent_part = self.meta_mgr.codico_mgr.get_adjacent_part(current_part, direction)
                if adjacent_part:
                    self._browse_load_part(adjacent_part, from_end=(direction < 0))
                    return
                else:
                    QMessageBox.information(self, tr("Nav"), tr("End of Parts in this volume."))
                    return

        # Standard file order navigation
        new_sid = self.searcher.get_adjacent_sys_id_by_file_order(current, direction)
        if new_sid:
            # Check if new folio belongs to a Part
            new_part = self.meta_mgr.get_part_for_folio(new_sid)
            if new_part:
                # Load the Part but position at this folio
                self._browse_load_part(new_part, target_folio=new_sid)
            else:
                # No Part - just load the folio
                self.current_browse_part_id = None
                self.current_browse_part_folios = []
                self.current_browse_part_folio_idx = 0
                self.browse_sys_input.setText(new_sid)
                shelf, _ = self.meta_mgr.get_meta_for_id(new_sid)
                if shelf and shelf != "Unknown":
                    self.browse_shelf_input.setText(shelf)
                self._set_last_browse_field("sys")
                self.browse_load()
        else:
            QMessageBox.information(self, tr("Nav"), tr("End of file list."))

    def _update_part_image_for_folio(self, folio_idx):
        """Update image viewer to show the current folio's images within a Part."""
        if not self.current_browse_part_id:
            return
        meta = self.meta_mgr.nli_cache.get(self.current_browse_sid, {})
        shelfmark, _ = self.meta_mgr.get_meta_for_id(self.current_browse_sid)
        folio_num = _get_folio_number_from_shelfmark(shelfmark)
        side_offset = 1 if (self.current_browse_internal_idx or 0) % 2 == 1 else 0

        # Check if folio needs dynamic images (missing from viewer but within Oxford folio_range)
        viewer_images = getattr(self.browse_viewer, 'images_ext', [])
        folio_in_viewer = any(img.get('folio_num') == folio_num for img in viewer_images) if folio_num and viewer_images else False

        if not folio_in_viewer and folio_num is not None and meta:
            oxford_part_meta = meta.get('oxford_part_metadata', {})
            folio_range = oxford_part_meta.get('folio_range', [])
            if meta.get('oxford_part_id') and len(folio_range) >= 2 and folio_range[0] <= folio_num <= folio_range[1]:
                image_idx = _get_folio_image_index(meta, folio_num, side_offset=side_offset)
                self.browse_viewer.load_images(meta, image_idx, target_folio=folio_num)
                return

        if viewer_images and self.browse_viewer.active_list:
            image_idx = _get_folio_image_index({'images_ext': viewer_images}, folio_num, side_offset=side_offset)
            self.browse_viewer.set_page(image_idx)
    
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)

if __name__ == "__main__":
    try:
        import ctypes
        if hasattr(ctypes, 'windll'):
            myappid = f'genizah.search.pro.{APP_VERSION}'
            ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
    except (ImportError, AttributeError):
        pass

    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    
    icon_path = resource_path("icon.ico")
    
    if os.path.exists(icon_path):
        app_icon = QIcon(icon_path)
        app.setWindowIcon(app_icon)
    
    window = GenizahGUI()
    window.showMaximized()
    sys.exit(app.exec())
