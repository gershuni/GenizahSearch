# -*- coding: utf-8 -*-
"""Phase 103 (D-10/D-11) — write_docx_result_block unit tests.

Tests pin the module-level DOCX per-result block writer in shared/docx_export.py:
- Block layout: heading / metadata / matched-text paragraph / URL line / separator
- LOCAL vs Genizah result dicts
- chunk_locator verbatim (D-02, no double-prefix)
- p_num fallback synthesized as 'p. N'
- Bold-red highlight runs for *-marked terms
- RTL applied on lang='he'
- Graceful on missing filepath
"""
import pytest

# Skip entire module if python-docx is not installed.
docx = pytest.importorskip('docx')


def _texts(doc):
    """Return a flat list of paragraph texts for the document."""
    return [p.text for p in doc.paragraphs]


def _make_local_result(chunk_locator=None, p_num=None, raw_file_hl='a *b* c'):
    return {
        'display': {'source': 'LOCAL', 'shelfmark': 'f.pdf', 'id': '970012345601234567'},
        'sys_id': '970012345601234567',
        'raw_file_hl': raw_file_hl,
        'chunk_locator': chunk_locator,
        'p_num': p_num,
    }


def _make_genizah_result():
    return {
        'display': {
            'source': '',
            'shelfmark': 'T-S 12.1',
            'title': 'Letter',
            'library_code': 'CUL',
            'id': '990001234567890',
            'img': '5',
        },
        'sys_id': '990001234567890',
        'raw_file_hl': 'a *b* c',
        'chunk_locator': None,
        'p_num': '5',
    }


# ---------------------------------------------------------------------------
# LOCAL result tests
# ---------------------------------------------------------------------------


def test_local_block_fields():
    """LOCAL result with filepath → doc contains filename, parent folder, filepath, 'LOCAL'."""
    from docx import Document
    from shared.docx_export import write_docx_result_block

    doc = Document()
    write_docx_result_block(doc, _make_local_result(chunk_locator='p. 3'), filepath='/x/y/f.pdf')
    texts = _texts(doc)
    combined = ' '.join(texts)
    assert 'f.pdf' in combined
    assert 'y' in combined          # parent folder
    assert '/x/y/f.pdf' in combined
    assert 'LOCAL' in combined


def test_local_block_page_locator_verbatim():
    """(REVIEWS LOW) chunk_locator='p. 3' must appear as-is — no 'page p. 3' double-prefix."""
    from docx import Document
    from shared.docx_export import write_docx_result_block

    doc = Document()
    write_docx_result_block(doc, _make_local_result(chunk_locator='p. 3'), filepath='/x/y/f.pdf')
    texts = _texts(doc)
    combined = ' '.join(texts)
    assert 'p. 3' in combined, "chunk_locator should appear verbatim in metadata line"
    assert 'page p. 3' not in combined, "Must not double-prefix chunk_locator (D-02)"


def test_local_block_page_pnum_fallback():
    """No chunk_locator but p_num='3' → metadata line contains 'p. 3' (synthesized)."""
    from docx import Document
    from shared.docx_export import write_docx_result_block

    doc = Document()
    write_docx_result_block(doc, _make_local_result(p_num='3'), filepath='/x/y/f.pdf')
    texts = _texts(doc)
    combined = ' '.join(texts)
    assert 'p. 3' in combined, "p_num fallback must be rendered as 'p. <n>'"


# ---------------------------------------------------------------------------
# Genizah result tests
# ---------------------------------------------------------------------------


def test_genizah_block_fields():
    """Genizah result → doc contains shelfmark, title, and genizahsearch.com URL."""
    from docx import Document
    from shared.docx_export import write_docx_result_block

    doc = Document()
    write_docx_result_block(doc, _make_genizah_result())
    texts = _texts(doc)
    combined = ' '.join(texts)
    assert 'T-S 12.1' in combined
    assert 'Letter' in combined
    assert 'genizahsearch.com' in combined


# ---------------------------------------------------------------------------
# Highlight runs test
# ---------------------------------------------------------------------------


def test_matched_text_highlight_runs():
    """Body paragraph must contain at least one bold+red run for the *-wrapped term."""
    from docx import Document
    from docx.shared import RGBColor
    from shared.docx_export import write_docx_result_block

    doc = Document()
    write_docx_result_block(doc, _make_local_result(raw_file_hl='plain *highlight* end'),
                            filepath='/x/y/f.pdf')
    bold_red_runs = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.bold and run.font.color.rgb == RGBColor(0xFF, 0x00, 0x00):
                bold_red_runs.append(run.text)
    assert bold_red_runs, "Expected at least one bold red run for *-marked text"
    assert 'highlight' in bold_red_runs[0]


# ---------------------------------------------------------------------------
# Two-block separator test
# ---------------------------------------------------------------------------


def test_two_blocks_separator():
    """Two calls produce 2 separator paragraphs ('_' * 40)."""
    from docx import Document
    from shared.docx_export import write_docx_result_block

    doc = Document()
    write_docx_result_block(doc, _make_local_result(), filepath='/x/y/f.pdf')
    write_docx_result_block(doc, _make_local_result(), filepath='/x/y/f.pdf')
    sep_count = sum(1 for p in doc.paragraphs if p.text == '_' * 40)
    assert sep_count == 2, f"Expected 2 separator paragraphs, got {sep_count}"


# ---------------------------------------------------------------------------
# Missing filepath graceful
# ---------------------------------------------------------------------------


def test_missing_filepath_no_error():
    """LOCAL result with filepath='' must not raise and must produce a block."""
    from docx import Document
    from shared.docx_export import write_docx_result_block

    doc = Document()
    write_docx_result_block(doc, _make_local_result(chunk_locator='p. 1'), filepath='')
    # Should produce paragraphs with LOCAL in the combined text
    combined = ' '.join(_texts(doc))
    assert 'LOCAL' in combined


# ---------------------------------------------------------------------------
# RTL on lang='he'
# ---------------------------------------------------------------------------


def test_rtl_applied_he():
    """With lang='he', at least one block paragraph must have w:bidi set in pPr."""
    from docx import Document
    from docx.oxml.ns import qn
    from shared.docx_export import write_docx_result_block

    doc = Document()
    write_docx_result_block(doc, _make_local_result(chunk_locator='p. 1'),
                            filepath='/x/y/f.pdf', lang='he')
    bidi_found = any(
        p._p.get_or_add_pPr().find(qn('w:bidi')) is not None
        for p in doc.paragraphs
    )
    assert bidi_found, "Expected at least one paragraph with w:bidi set for lang='he'"
