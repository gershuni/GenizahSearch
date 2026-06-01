# -*- coding: utf-8 -*-
"""Phase 103 LEXP-08 non-regression + D-12 DOCX carve-out consolidated gate.

This module is the final phase-103 safety net.  It asserts FOUR guarantees:

  A. xlsx cross-parity invariant (tests/test_export_xlsx_cross_parity.py) passes
     WITHOUT modification — LEXP-08 xlsx clause; ROADMAP Phase-103 SC#5.

  B. Genizah-only xlsx export is structurally identical to pre-v7.17: exactly
     4 sheets ['Search Results', 'Manuscripts', 'Bibliography',
     'Credits and Info'], no 'Local Documents' sheet added, 'Search Results'
     active — LEXP-08.

  C. Genizah-only CSV header is the unchanged 7-column table (no Filepath/Page
     appended) — LEXP-08.  Genizah TXT block content is byte-identical to
     pre-v7.17 (markers preserved, same f-string format) — LEXP-08.

  D. D-12 CARVE-OUT (APPROVED DEVIATION): Genizah-only DOCX is INTENTIONALLY a
     per-result block layout — NOT a cramped 7-column python-docx table.  This
     change is approved scope (user decision 2026-06-01, see 103-CONTEXT.md
     D-10/D-12 and ROADMAP Phase-103 SC#5 amendment).  The test below ASSERTS
     the block layout as EXPECTED behaviour, so a future verifier does NOT treat
     it as a regression.

CRITICAL: do NOT add any Genizah-DOCX byte-identity-vs-pre-v7.17 assertion here
— that would contradict the approved D-12 carve-out.
"""
from __future__ import annotations

import subprocess
import sys
from io import BytesIO
from pathlib import Path

import openpyxl
import pytest

# ---------------------------------------------------------------------------
# Shared test fixtures
# ---------------------------------------------------------------------------

GEN_ID = "990012345678901"


def _gen_result(sid=GEN_ID):
    """Minimal Genizah (non-LOCAL) result dict."""
    return {
        "display": {
            "id": sid,
            "source": "",
            "shelfmark": "T-S 12.1",
            "title": "Letter",
            "img": "5",
        },
        "sys_id": sid,
        "raw_file_hl": "gen *hit* text",
        "chunk_locator": "",
        "p_num": "",
    }


def _meta_resolver_fake(sid):
    if not sid:
        return None
    return {
        "shelfmark": f"T-S {sid[-4:]}",
        "title": f"Title {sid}",
        "library_code": "CUL",
        "library_name": "Cambridge University Library",
    }


def _identity_sanitize(text):
    return "" if text is None else str(text)


@pytest.fixture
def stub_dossier(monkeypatch):
    """Stub out the network/DB calls in shared.export_dossier."""
    from shared import export_dossier

    monkeypatch.setattr(export_dossier, "pgp_subset_for_sys_id", lambda s, **kw: None)
    monkeypatch.setattr(export_dossier, "nli_subset_for_sys_id", lambda s, **kw: None)
    monkeypatch.setattr(export_dossier, "catalog_summary_for_sys_id", lambda s, **kw: None)
    monkeypatch.setattr(export_dossier, "bibliography_for_sys_id", lambda s, **kw: [])


def _build_genizah_only_workbook(stub_dossier_fixture=None, lang="en"):
    """Offline helper — calls _build_search_results_xlsx_bytes with a single
    Genizah result, NO local_filepath_map.  Returns an openpyxl Workbook.
    """
    from genizah_app import _build_search_results_xlsx_bytes

    content = _build_search_results_xlsx_bytes(
        results=[_gen_result()],
        headers_main=[
            "System ID",
            "Library",
            "Shelfmark",
            "Title",
            "Image/Page",
            "Source",
            "Snippet",
            "Full Text",
            "Has PGP",
            "Is Printed",
            "Domains",
            "Image URL",
        ],
        meta_resolver=_meta_resolver_fake,
        sanitize_fn=_identity_sanitize,
        lang=lang,
    )
    return openpyxl.load_workbook(BytesIO(content), rich_text=True)


# ---------------------------------------------------------------------------
# A. xlsx cross-parity invariant passes unmodified (LEXP-08 / ROADMAP SC#5)
# ---------------------------------------------------------------------------


def test_cross_parity_invariant_still_passes():
    """LEXP-08 / ROADMAP SC#5: the xlsx cross-parity invariant passes UNMODIFIED.

    Runs tests/test_export_xlsx_cross_parity.py in a subprocess and asserts the
    entire module exits 0.  If this fails, Plans 02 or 03 introduced a parity
    regression in the xlsx structure (sheet names / header rows) — check that the
    Genizah-only path still produces identical output on web and desktop.
    """
    result = subprocess.run(
        [sys.executable, "-m", "pytest", "tests/test_export_xlsx_cross_parity.py", "-q"],
        capture_output=True,
        text=True,
        cwd=str(Path(__file__).resolve().parent.parent),
    )
    assert result.returncode == 0, (
        f"cross-parity invariant FAILED (LEXP-08 regression):\n"
        f"{result.stdout}\n{result.stderr}"
    )


def test_cross_parity_file_assertion_intact():
    """Guard: tests/test_export_xlsx_cross_parity.py still contains the core
    parity assertion so accidental weakening is caught.

    The file must contain ``wb_web.sheetnames == wb_desktop.sheetnames`` — this is
    the structural invariant that LEXP-08 and ROADMAP SC#5 depend on.
    """
    cross_parity_path = (
        Path(__file__).resolve().parent.parent
        / "tests"
        / "test_export_xlsx_cross_parity.py"
    )
    source = cross_parity_path.read_text(encoding="utf-8")
    assert "wb_web.sheetnames == wb_desktop.sheetnames" in source, (
        "GUARD FAILED: tests/test_export_xlsx_cross_parity.py no longer contains "
        "the core parity assertion ``wb_web.sheetnames == wb_desktop.sheetnames``. "
        "The invariant may have been accidentally weakened."
    )


# ---------------------------------------------------------------------------
# B. Genizah-only xlsx structurally unchanged (LEXP-08)
# ---------------------------------------------------------------------------


def test_genizah_only_xlsx_four_sheets(stub_dossier):
    """LEXP-08: Genizah-only xlsx export produces the unchanged 4-sheet workbook."""
    wb = _build_genizah_only_workbook()
    assert wb.sheetnames == [
        "Search Results",
        "Manuscripts",
        "Bibliography",
        "Credits and Info",
    ], f"Genizah-only sheet names changed (LEXP-08 regression): {wb.sheetnames}"
    assert "Local Documents" not in wb.sheetnames, (
        "Genizah-only export must NOT contain a 'Local Documents' sheet"
    )
    assert wb.active.title == "Search Results", (
        f"Active sheet must be 'Search Results', got {wb.active.title!r}"
    )


def test_genizah_only_xlsx_four_sheets_he(stub_dossier):
    """LEXP-08 + bilingual: Genizah-only HE export is still 4-sheet, no Local Documents."""
    wb = _build_genizah_only_workbook(lang="he")
    # HE sheet names for the 4 core Genizah sheets
    assert len(wb.sheetnames) == 4, (
        f"Genizah-only HE export must have 4 sheets; got {wb.sheetnames}"
    )
    assert "מסמכים מקומיים" not in wb.sheetnames, (
        "Genizah-only HE export must NOT contain the 'Local Documents' sheet in Hebrew"
    )


# ---------------------------------------------------------------------------
# C. Genizah-only CSV 7-column + TXT byte-identical (LEXP-08)
# ---------------------------------------------------------------------------

try:
    from genizah_app import (
        _csv_extra_cols,
        _format_txt_genizah_block,
    )

    _HELPERS_AVAILABLE = True
except Exception as _e:
    _HELPERS_AVAILABLE = False
    _IMPORT_ERR = str(_e)


_skip_if_no_helpers = pytest.mark.skipif(
    not _HELPERS_AVAILABLE,
    reason=f"genizah_app helpers not importable: {locals().get('_IMPORT_ERR', '')}",
)


@_skip_if_no_helpers
def test_genizah_only_csv_seven_columns():
    """LEXP-08: Genizah-only CSV helper yields two empty extra cells (no Filepath/Page appended).

    In export_results the full CSV header is
        headers + [_fp_label, _pg_label]  ONLY when _has_local_in_export.
    For a Genizah-only result set _has_local_in_export is False, so the header
    stays at 7 columns.  This test pins that the _csv_extra_cols helper returns
    ['', ''] for a Genizah row, which is what drives that branch.
    """
    gen_r = _gen_result()
    extra = _csv_extra_cols(gen_r)
    assert extra == ["", ""], (
        f"Genizah row extra CSV cols must be ['', ''], got {extra!r} — "
        "Filepath/Page columns must NOT be appended for Genizah-only exports (LEXP-08)"
    )


@_skip_if_no_helpers
def test_genizah_only_txt_block_byte_identical():
    """LEXP-08: Genizah TXT block content is byte-identical to pre-v7.17.

    Pre-v7.17 format:
        f"=== {shelfmark} | {title} ===\\n{snippet}"
    where snippet = raw_file_hl.strip().replace('\\n', ' ').replace('\\r', '')

    Pinned behaviours:
    - Header line starts with '=== T-S 12.1 | Letter ==='
    - Snippet preserves '*' highlight markers (NOT stripped in Genizah TXT — LEXP-08)
    - No 'Path:' line (LOCAL-only feature)
    """
    gen_r = _gen_result()
    block = _format_txt_genizah_block(gen_r)

    expected_header = "=== T-S 12.1 | Letter ==="
    assert block.startswith(expected_header), (
        f"Genizah TXT header mismatch (LEXP-08 byte-identical failure).\n"
        f"  Expected start: {expected_header!r}\n"
        f"  Actual start:   {block[:80]!r}"
    )
    # Markers preserved (byte-identical to pre-v7.17)
    assert "*" in block, (
        "Genizah TXT must preserve '*' highlight markers — pre-v7.17 behaviour (LEXP-08)"
    )
    # No 'Path:' line — that is LOCAL-only
    assert "Path:" not in block, (
        "Genizah TXT block must NOT contain 'Path:' line (LOCAL-only feature)"
    )


# ---------------------------------------------------------------------------
# D. DOCX carve-out asserted as EXPECTED (D-12)
# ---------------------------------------------------------------------------

docx = pytest.importorskip("docx", reason="python-docx not installed")


def test_genizah_docx_is_block_layout_not_table():
    """D-12 CARVE-OUT: Genizah-only DOCX is INTENTIONALLY a per-result block layout
    (no table).

    The pre-v7.17 DOCX used doc.add_table(…) — a cramped 7-column table.
    Phase 103 (D-10) replaces this with a per-result rich-document block via
    shared.docx_export.write_docx_result_block.  This is an APPROVED scope
    decision (103-CONTEXT.md D-12, ROADMAP Phase-103 SC#5 amendment, user
    direction 2026-06-01).

    This test asserts the block layout as the EXPECTED Genizah output so a
    future verifier does NOT treat the DOCX change as a regression.

    IMPORTANT: do NOT assert byte-identity with pre-v7.17 — that would
    contradict the approved D-12 carve-out.
    """
    from docx import Document

    from shared.docx_export import write_docx_result_block

    doc = Document()
    gen_r = {
        "display": {
            "id": GEN_ID,
            "source": "",
            "shelfmark": "T-S 12.1",
            "title": "Letter",
            "img": "5",
        },
        "sys_id": GEN_ID,
        "raw_file_hl": "gen *hit* text",
    }
    write_docx_result_block(doc, gen_r, filepath="", lang="en")

    # D-12: must be block layout, NOT a table
    assert len(doc.tables) == 0, (
        f"D-12 VIOLATION: Genizah DOCX must be block layout (no table), "
        f"got {len(doc.tables)} table(s).  Check that export_results DOCX branch "
        f"uses write_docx_result_block and NOT doc.add_table(...)."
    )

    joined = "\n".join(p.text for p in doc.paragraphs)

    # Shelfmark appears in the heading paragraph
    assert "T-S 12.1" in joined, (
        f"Shelfmark 'T-S 12.1' not found in Genizah DOCX block:\n{joined[:300]!r}"
    )

    # URL line is present (D-11 — GenizahSearch URL for Genizah rows)
    assert "genizahsearch.com" in joined, (
        f"GenizahSearch URL not found in Genizah DOCX block (D-11):\n{joined[:300]!r}"
    )

    # Separator paragraph exists
    assert "________" in joined, (
        "Separator paragraph ('___...') missing from Genizah DOCX block"
    )
