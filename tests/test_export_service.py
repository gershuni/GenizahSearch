# -*- coding: utf-8 -*-
"""
Comprehensive tests for the unified export service.

Tests cover:
1. Text processing utilities
2. Excel export functionality
3. Word export functionality
4. All export types: search results, lists, parallels, browse
5. RTL handling for Hebrew text
6. Error handling for empty/invalid data
"""

import pytest
import io
import openpyxl
from docx import Document
from unittest.mock import MagicMock

# Import the export service module
from web.export_service import (
    ExportService,
    get_export_service,
    sanitize_text_for_excel,
    clean_text_single_line,
    remove_highlight_markers,
    make_safe_filename,
    contains_any_term,
    extract_search_terms,
    add_hebrew_paragraph,
    add_highlighted_hebrew_paragraph,
    create_excel_workbook,
    style_excel_header,
    get_cell_alignment,
    save_workbook_to_bytes,
    save_document_to_bytes,
    CREDITS_TEXT,
)


class TestTextProcessingUtilities:
    """Test text processing helper functions."""

    def test_sanitize_text_for_excel_removes_control_chars(self):
        """Control characters should be removed from Excel text."""
        text = "Hello\x00World\x0bTest\x1f"
        result = sanitize_text_for_excel(text)
        assert "\x00" not in result
        assert "\x0b" not in result
        assert "\x1f" not in result
        assert "Hello" in result
        assert "World" in result

    def test_sanitize_text_for_excel_keeps_hebrew(self):
        """Hebrew characters should be preserved."""
        text = "שלום עולם"
        result = sanitize_text_for_excel(text)
        assert result == text

    def test_sanitize_text_for_excel_empty_input(self):
        """Empty input should return empty string."""
        assert sanitize_text_for_excel("") == ""
        assert sanitize_text_for_excel(None) == ""

    def test_clean_text_single_line_removes_newlines(self):
        """Line breaks should be replaced with spaces."""
        text = "Line 1\nLine 2\r\nLine 3\rLine 4"
        result = clean_text_single_line(text)
        assert "\n" not in result
        assert "\r" not in result
        assert "Line 1" in result
        assert "Line 4" in result

    def test_clean_text_single_line_collapses_spaces(self):
        """Multiple spaces should be collapsed to single space."""
        text = "Hello    World"
        result = clean_text_single_line(text)
        assert "  " not in result

    def test_clean_text_single_line_empty_input(self):
        """Empty input should return empty string."""
        assert clean_text_single_line("") == ""
        assert clean_text_single_line(None) == ""

    def test_remove_highlight_markers(self):
        """Asterisk markers should be removed."""
        text = "This *is* a *test*"
        result = remove_highlight_markers(text)
        assert "*" not in result
        assert "This is a test" == result

    def test_remove_highlight_markers_empty(self):
        """Empty input should return empty string."""
        assert remove_highlight_markers("") == ""
        assert remove_highlight_markers(None) == ""

    def test_make_safe_filename_removes_illegal_chars(self):
        """Illegal filesystem characters should be removed."""
        query = 'test:file/name*"query?'
        result = make_safe_filename(query)
        assert ":" not in result
        assert "/" not in result
        assert "*" not in result
        assert '"' not in result
        assert "?" not in result

    def test_make_safe_filename_limits_length(self):
        """Filename should be limited to max_length."""
        query = "a" * 100
        result = make_safe_filename(query, max_length=50)
        assert len(result) <= 50

    def test_make_safe_filename_default_for_empty(self):
        """Empty query should return default."""
        result = make_safe_filename("", default="fallback")
        assert result == "fallback"

    def test_make_safe_filename_hebrew(self):
        """Hebrew query should work correctly."""
        query = "שלום עולם"
        result = make_safe_filename(query)
        assert result == "שלום_עולם"

    def test_contains_any_term_positive(self):
        """Should return True when term is found."""
        assert contains_any_term("Hello World", ["world"]) == True
        assert contains_any_term("Hello World", ["HELLO"]) == True

    def test_contains_any_term_negative(self):
        """Should return False when term is not found."""
        assert contains_any_term("Hello World", ["foo"]) == False

    def test_contains_any_term_empty(self):
        """Should return False for empty inputs."""
        assert contains_any_term("", ["test"]) == False
        assert contains_any_term("test", []) == False

    def test_extract_search_terms_filters_operators(self):
        """Should filter out search operators."""
        query = "hello =exact ?fuzzy ~similar /regex $var #tag world"
        result = extract_search_terms(query)
        assert "hello" in result
        assert "world" in result
        assert "=exact" not in result
        assert "?fuzzy" not in result

    def test_extract_search_terms_empty(self):
        """Empty query should return empty list."""
        assert extract_search_terms("") == []
        assert extract_search_terms(None) == []


class TestExcelUtilities:
    """Test Excel-specific utilities."""

    def test_create_excel_workbook(self):
        """Should create workbook with correct settings."""
        wb, ws = create_excel_workbook("Test Sheet", rtl_sheet=True)
        assert ws.title == "Test Sheet"
        assert ws.sheet_view.rightToLeft == True

    def test_create_excel_workbook_truncates_long_name(self):
        """Sheet name should be truncated to 31 chars."""
        wb, ws = create_excel_workbook("A" * 50)
        assert len(ws.title) <= 31

    def test_style_excel_header(self):
        """Headers should be styled correctly."""
        wb, ws = create_excel_workbook()
        headers = ["Col1", "Col2", "Col3"]
        style_excel_header(ws, headers)

        # Check header row values
        assert ws.cell(row=1, column=1).value == "Col1"
        assert ws.cell(row=1, column=2).value == "Col2"
        assert ws.cell(row=1, column=3).value == "Col3"

        # Check styling applied
        assert ws.cell(row=1, column=1).font.bold == True

    def test_get_cell_alignment_rtl(self):
        """RTL alignment should have correct settings."""
        align = get_cell_alignment('rtl')
        assert align.horizontal == "right"
        assert align.readingOrder == 2

    def test_get_cell_alignment_ltr(self):
        """LTR alignment should have correct settings."""
        align = get_cell_alignment('ltr')
        assert align.horizontal == "left"

    def test_get_cell_alignment_center(self):
        """Center alignment should have correct settings."""
        align = get_cell_alignment('center')
        assert align.horizontal == "center"

    def test_save_workbook_to_bytes(self):
        """Should return valid bytes."""
        wb, ws = create_excel_workbook()
        ws.append(["test", "data"])
        content = save_workbook_to_bytes(wb)
        assert isinstance(content, bytes)
        assert len(content) > 0

        # Verify it's a valid Excel file
        stream = io.BytesIO(content)
        loaded_wb = openpyxl.load_workbook(stream)
        assert loaded_wb.active.cell(row=1, column=1).value == "test"


class TestWordUtilities:
    """Test Word-specific utilities."""

    def test_save_document_to_bytes(self):
        """Should return valid bytes."""
        doc = Document()
        doc.add_paragraph("Test content")
        content = save_document_to_bytes(doc)
        assert isinstance(content, bytes)
        assert len(content) > 0

    def test_add_hebrew_paragraph(self):
        """Should add paragraph with RTL formatting."""
        doc = Document()
        add_hebrew_paragraph(doc, "שלום עולם", bold=True)
        assert len(doc.paragraphs) == 1
        # Check bold was applied
        assert doc.paragraphs[0].runs[0].bold == True

    def test_add_highlighted_hebrew_paragraph(self):
        """Should add paragraph with highlighted terms."""
        doc = Document()
        add_highlighted_hebrew_paragraph(doc, "This *is* highlighted")
        assert len(doc.paragraphs) == 1


class TestExportService:
    """Test the main ExportService class."""

    @pytest.fixture
    def mock_meta_mgr(self):
        """Create a mock MetadataManager."""
        mgr = MagicMock()
        mgr.get_meta_for_id.return_value = ("T-S 12.345", "Test Title")
        return mgr

    @pytest.fixture
    def export_service(self, mock_meta_mgr):
        """Create an ExportService with mock dependencies."""
        return ExportService(meta_mgr=mock_meta_mgr)

    @pytest.fixture
    def sample_search_results(self):
        """Sample search results for testing."""
        return [
            {
                'display': {
                    'shelfmark': 'T-S 12.345',
                    'title': 'כתב יד עברי',
                    'id': '9912345678901234',
                },
                'snippet': 'This is a *highlighted* snippet',
                'full_text': 'Full text content here',
                'sort_score': 0.95,
            },
            {
                'display': {
                    'shelfmark': 'T-S 12.346',
                    'title': 'מסמך נוסף',
                    'id': '9912345678901235',
                },
                'snippet': 'Another snippet',
                'full_text': 'More text',
                'sort_score': 0.85,
            },
        ]

    @pytest.fixture
    def sample_list_items(self):
        """Sample list items for testing."""
        return [
            {
                'sys_id': '9912345678901234',
                'fl_id': 'FL123',
                'note': 'Test note',
                'added': '2025-01-29',
            },
            {
                'sys_id': '9912345678901235',
                'fl_id': 'FL124',
                'note': '',
                'added': '2025-01-29',
            },
        ]

    @pytest.fixture
    def sample_parallels_results(self):
        """Sample parallels results for testing."""
        return [
            {
                'raw_header': 'header_9912345678901234_page1',
                'score': 85,
                'source_ctx': 'Source *context* text',
                'text': 'Manuscript match text',
            },
        ]

    @pytest.fixture
    def sample_browse_data(self):
        """Sample browse data for testing."""
        return {
            'shelfmark': 'T-S 12.345',
            'title': 'Test manuscript',
            'sys_id': '9912345678901234',
            'text': 'Hebrew manuscript text here',
            'p_num': 1,
        }

    def test_export_search_results_excel(self, export_service, sample_search_results):
        """Should export search results to valid Excel file."""
        content, filename = export_service.export_search_results_excel(
            sample_search_results, "test query"
        )

        assert isinstance(content, bytes)
        assert filename.endswith('.xlsx')
        assert "test_query" in filename

        # Verify Excel content
        stream = io.BytesIO(content)
        wb = openpyxl.load_workbook(stream)
        ws = wb.active

        # Check headers
        assert ws.cell(row=1, column=1).value == "Shelfmark"
        # Check data row
        assert ws.cell(row=2, column=1).value == "T-S 12.345"

    def test_export_search_results_excel_rehydrates_compact_full_text(self, export_service, monkeypatch):
        """Compact export-state rows should still produce full-text Excel."""
        from web.state import state as runtime_state

        searcher = MagicMock()
        searcher.get_full_text_by_id.return_value = "Rehydrated full text"
        monkeypatch.setattr(runtime_state, 'searcher', searcher, raising=False)

        compact_results = [{
            'uid': 'uid-1',
            'display': {
                'shelfmark': 'T-S 12.345',
                'title': 'title',
                'id': '9912345678901234',
            },
            'snippet': 'snippet',
            'full_text_excerpt': 'short excerpt',
            'sort_score': 0.95,
        }]

        content, _filename = export_service.export_search_results_excel(compact_results, "query")

        wb = openpyxl.load_workbook(io.BytesIO(content))
        ws = wb.active
        assert ws.cell(row=2, column=7).value == "Rehydrated full text"
        searcher.get_full_text_by_id.assert_called_once_with('uid-1')

    def test_excel_export_rehydrates_display_from_uid(self, export_service):
        """SEED-002: when a compacted (uid-only) row is exported, the helper
        ``_resolve_result_display`` should rehydrate shelfmark/title/library
        from ``meta_mgr.get_meta_for_id`` via the parsed sys_id."""
        export_service.meta_mgr.get_meta_for_id.return_value = ('T-S 12.345', 'Test Title')
        export_service.meta_mgr.get_library_for_id.return_value = 'CUL'
        export_service.meta_mgr.parse_full_id_components.return_value = {
            'sys_id': '9912345678901234',
            'ie_id': 'IE1',
            'p_num': '1',
            'fl_id': 'FL1',
        }

        compact_results = [{
            'uid': '9912345678901234_IE1_P1_FL1',
            'sort_score': 0.95,
            'snippet': 'snippet',
            'match_terms': [],
        }]
        content, _filename = export_service.export_search_results_excel(compact_results, "query")

        wb = openpyxl.load_workbook(io.BytesIO(content))
        ws = wb.active
        # Column 1 = Shelfmark, 3 = Title, 4 = System ID
        assert ws.cell(row=2, column=1).value == 'T-S 12.345'
        assert ws.cell(row=2, column=3).value == 'Test Title'
        assert ws.cell(row=2, column=4).value == '9912345678901234'
        # Verify get_meta_for_id was called with the extracted sys_id.
        export_service.meta_mgr.get_meta_for_id.assert_any_call('9912345678901234')

    def test_excel_export_graceful_degradation_on_unknown_uid(self, export_service):
        """SEED-002: when the uid can't parse to a sys_id and no raw_header is
        present, the export must fall back to 'Unknown' (not crash, not produce
        a MagicMock-coerced string)."""
        # Critical: explicitly configure ALL three mock returns. Without this,
        # MagicMock returns MagicMock objects that coerce to "<MagicMock ...>"
        # strings in the cell — a false-positive failure mode.
        export_service.meta_mgr.parse_full_id_components.return_value = {'sys_id': None}
        export_service.meta_mgr.get_meta_for_id.return_value = ('Unknown', '')
        export_service.meta_mgr.get_library_for_id.return_value = ''

        compact_results = [{
            'uid': 'malformed-no-sys-id',
            'sort_score': 0.0,
            'snippet': 's',
            'match_terms': [],
        }]
        content, _filename = export_service.export_search_results_excel(compact_results, "query")

        wb = openpyxl.load_workbook(io.BytesIO(content))
        ws = wb.active
        # Must be the literal string 'Unknown', not a MagicMock repr.
        assert ws.cell(row=2, column=1).value == 'Unknown'

    def test_excel_output_equivalent_legacy_vs_compacted(self, export_service):
        """SEED-002: a legacy row (with display dict) and a compacted row
        (uid only) that resolve to the same sys_id must produce identical
        Excel cell values for Shelfmark, Library code, and Title."""
        export_service.meta_mgr.get_meta_for_id.return_value = ('T-S 99.1', 'Same Title')
        export_service.meta_mgr.get_library_for_id.return_value = 'CUL'
        export_service.meta_mgr.parse_full_id_components.return_value = {
            'sys_id': '9912345678901111',
            'ie_id': 'IE1',
            'p_num': '1',
            'fl_id': 'FL1',
        }

        legacy_row = {
            'display': {
                'shelfmark': 'T-S 99.1',
                'title': 'Same Title',
                'library_code': 'CUL',
                'id': '9912345678901111',
            },
            'uid': '9912345678901111_IE1_P1_FL1',
            'sort_score': 0.5,
            'snippet': 's',
            'match_terms': [],
        }
        compact_row = {
            'uid': '9912345678901111_IE1_P1_FL1',
            'sort_score': 0.5,
            'snippet': 's',
            'match_terms': [],
        }

        legacy_content, _ = export_service.export_search_results_excel([legacy_row], "q")
        compact_content, _ = export_service.export_search_results_excel([compact_row], "q")

        legacy_wb = openpyxl.load_workbook(io.BytesIO(legacy_content))
        compact_wb = openpyxl.load_workbook(io.BytesIO(compact_content))
        legacy_ws = legacy_wb.active
        compact_ws = compact_wb.active

        # Columns 1 (Shelfmark), 2 (Library), 3 (Title) must match.
        for col in (1, 2, 3):
            assert legacy_ws.cell(row=2, column=col).value == compact_ws.cell(row=2, column=col).value, \
                f"column {col} differs: legacy={legacy_ws.cell(row=2, column=col).value!r}, " \
                f"compact={compact_ws.cell(row=2, column=col).value!r}"

    def test_export_search_results_excel_empty_raises(self, export_service):
        """Should raise ValueError for empty results."""
        with pytest.raises(ValueError, match="No results to export"):
            export_service.export_search_results_excel([], "query")

    def test_export_search_results_word(self, export_service, sample_search_results):
        """Should export search results to valid Word file."""
        content, filename = export_service.export_search_results_word(
            sample_search_results, "test query"
        )

        assert isinstance(content, bytes)
        assert filename.endswith('.docx')
        assert "test_query" in filename

    def test_export_search_results_word_empty_raises(self, export_service):
        """Should raise ValueError for empty results."""
        with pytest.raises(ValueError, match="No results to export"):
            export_service.export_search_results_word([], "query")

    def test_export_list_excel(self, export_service, sample_list_items):
        """Should export list to valid Excel file."""
        content, filename = export_service.export_list_excel(
            "list_123", "My List", sample_list_items
        )

        assert isinstance(content, bytes)
        assert "My_List" in filename
        assert filename.endswith('.xlsx')

        # Verify Excel content
        stream = io.BytesIO(content)
        wb = openpyxl.load_workbook(stream)
        ws = wb.active

        # Check headers
        assert ws.cell(row=1, column=1).value == "#"
        assert ws.cell(row=1, column=2).value == "Shelfmark"

    def test_export_list_excel_empty_raises(self, export_service):
        """Should raise ValueError for empty list."""
        with pytest.raises(ValueError, match="List is empty"):
            export_service.export_list_excel("list_123", "Test", [])

    def test_export_parallels_excel(self, export_service, sample_parallels_results):
        """Should export parallels to valid Excel file."""
        content, filename = export_service.export_parallels_excel(
            sample_parallels_results, []
        )

        assert isinstance(content, bytes)
        assert filename == "parallels_results.xlsx"

    def test_export_parallels_excel_with_filtered(self, export_service, sample_parallels_results):
        """Should include filtered results in export."""
        filtered = [{'raw_header': 'filtered_991234', 'score': 50, 'source_ctx': '', 'text': ''}]
        content, filename = export_service.export_parallels_excel(
            sample_parallels_results, filtered
        )

        assert isinstance(content, bytes)

    def test_export_parallels_excel_empty_raises(self, export_service):
        """Should raise ValueError for empty parallels."""
        with pytest.raises(ValueError, match="No parallels results"):
            export_service.export_parallels_excel([], [])

    def test_export_parallels_word(self, export_service, sample_parallels_results):
        """Should export parallels to valid Word file."""
        content, filename = export_service.export_parallels_word(
            sample_parallels_results, []
        )

        assert isinstance(content, bytes)
        assert filename == "parallels_results.docx"

    def test_export_browse_word(self, export_service, sample_browse_data):
        """Should export browse data to valid Word file."""
        content, filename = export_service.export_browse_word(sample_browse_data)

        assert isinstance(content, bytes)
        assert filename == "T-S_12345.docx"

    def test_export_browse_word_with_pages(self, export_service):
        """Should handle multi-page browse data."""
        browse_data = {
            'shelfmark': 'T-S 12.345',
            'title': 'Test',
            'sys_id': '991234',
            'view_all': True,
            'pages': [
                {'p_num': 1, 'text': 'Page 1 text'},
                {'p_num': 2, 'text': 'Page 2 text'},
            ],
        }
        content, filename = export_service.export_browse_word(browse_data)
        assert isinstance(content, bytes)

    def test_export_browse_word_empty_raises(self, export_service):
        """Should raise ValueError for empty browse data."""
        with pytest.raises(ValueError, match="No browse data"):
            export_service.export_browse_word({})

    def test_get_metadata_with_valid_id(self, export_service):
        """Should return metadata for valid ID."""
        shelfmark, title = export_service.get_metadata("9912345678901234")
        assert shelfmark == "T-S 12.345"
        assert title == "Test Title"

    def test_get_metadata_without_meta_mgr(self):
        """Should return defaults without MetadataManager."""
        svc = ExportService(meta_mgr=None)
        shelfmark, title = svc.get_metadata("9912345678901234")
        assert shelfmark == "Unknown"
        assert title == ""

    def test_get_metadata_empty_id(self, export_service):
        """Should return defaults for empty ID."""
        shelfmark, title = export_service.get_metadata("")
        assert shelfmark == "Unknown"


class TestGetExportService:
    """Test the singleton getter function."""

    def test_returns_export_service(self):
        """Should return ExportService instance."""
        svc = get_export_service()
        assert isinstance(svc, ExportService)

    def test_caches_instance(self):
        """Should return same instance for same meta_mgr."""
        svc1 = get_export_service()
        svc2 = get_export_service()
        assert svc1 is svc2

    def test_updates_meta_mgr(self):
        """Should update instance when meta_mgr changes."""
        mock_mgr = MagicMock()
        svc = get_export_service(mock_mgr)
        assert svc.meta_mgr is mock_mgr


class TestCreditsText:
    """Test that credits are properly defined."""

    def test_credits_has_required_entries(self):
        """Credits should have all required lines."""
        assert len(CREDITS_TEXT) >= 4
        assert any("Dicta Genizah Search" in line for line in CREDITS_TEXT)
        assert any("MiDRASH" in line for line in CREDITS_TEXT)
        assert any("doi.org" in line for line in CREDITS_TEXT)


class TestEdgeCases:
    """Test edge cases and special scenarios."""

    def test_hebrew_text_in_excel(self):
        """Hebrew text should be properly exported to Excel."""
        svc = ExportService()
        results = [{
            'display': {
                'shelfmark': 'ת-ס 12.345',
                'title': 'כותרת בעברית',
                'id': '991234',
            },
            'snippet': 'טקסט *מודגש* בעברית',
            'full_text': 'טקסט מלא',
            'sort_score': 0.9,
        }]

        content, _ = svc.export_search_results_excel(results, "שאילתה")

        stream = io.BytesIO(content)
        wb = openpyxl.load_workbook(stream)
        ws = wb.active

        # Hebrew shelfmark should be preserved
        assert ws.cell(row=2, column=1).value == 'ת-ס 12.345'

    def test_very_long_text_in_excel(self):
        """Very long text should be truncated."""
        svc = ExportService()
        long_text = "A" * 50000  # Exceeds Excel cell limit
        results = [{
            'display': {'shelfmark': 'Test', 'title': '', 'id': '1'},
            'snippet': '',
            'full_text': long_text,
            'sort_score': 0.5,
        }]

        content, _ = svc.export_search_results_excel(results, "")

        stream = io.BytesIO(content)
        wb = openpyxl.load_workbook(stream)
        ws = wb.active

        # Full text should be truncated
        assert len(ws.cell(row=2, column=7).value) <= 32000

    def test_special_characters_in_filename(self):
        """Special characters in query should create safe filename."""
        svc = ExportService()
        results = [{
            'display': {'shelfmark': 'Test', 'title': '', 'id': '1'},
            'snippet': '',
            'full_text': '',
            'sort_score': 0.5,
        }]

        _, filename = svc.export_search_results_excel(results, 'query:with/special*chars')
        assert ":" not in filename
        assert "/" not in filename
        assert "*" not in filename
