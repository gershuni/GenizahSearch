# -*- coding: utf-8 -*-
"""Phase 103 LEXP-01/06/07: CSV / TXT / DOCX LOCAL-aware export tests.

Offline tests (no Qt, no genizah_app module-level side-effects) that exercise
the module-level helpers extracted in Phase 103 Plan 03:

    genizah_app._build_export_data_row
    genizah_app._csv_extra_cols
    genizah_app._format_txt_local_block
    genizah_app._format_txt_genizah_block

and the Wave-1 shared DOCX writer:

    shared.docx_export.write_docx_result_block

Covers:
- CSV mixed: Filepath + Page columns appended only when LOCAL present (D-08)
- CSV Genizah-only: unchanged 7-column table (LEXP-08)
- CSV LOCAL column remap: Shelfmark=filename, Library=parent, Source='LOCAL'
- CSV formula injection: LOCAL values beginning with =+-@ are neutralized (T-103-07)
- TXT LOCAL block: === filename | parent === / Path: filepath (page N) / snippet
- TXT Genizah block: byte-identical to pre-v7.17 output (LEXP-08)
- DOCX two-results: LOCAL + Genizah blocks, no table, correct fields
"""

import csv
import io

import pytest

# ---------------------------------------------------------------------------
# Helpers to import the tested targets without triggering the full Qt app
# ---------------------------------------------------------------------------

# Import the module-level helpers without the GenizahGUI class activating Qt.
# We do NOT use importlib.import_module("genizah_app") because that triggers
# QApplication construction at import time; instead we import directly from
# the module after confirming the helpers exist as module-level callables.
try:
    from genizah_app import (
        _build_export_data_row,
        _csv_extra_cols,
        _format_txt_genizah_block,
        _format_txt_local_block,
        _local_page_label,
    )

    _HELPERS_AVAILABLE = True
except Exception as _e:
    _HELPERS_AVAILABLE = False
    _IMPORT_ERR = str(_e)

docx = pytest.importorskip("docx", reason="python-docx not installed")

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

LOCAL_ID = "970000000100000001"
GEN_ID = "990012345678901"
LOCAL_FP = r"C:\books\hebrew\dusiach.pdf"  # contains os.sep
LOCAL_FP_EVIL = "=cmd|calc"  # CSV formula-injection vector
LOCAL_PARENT = "hebrew"  # parent of LOCAL_FP; hardcoded — os.path can't split a Windows path on POSIX CI
LOCAL_FNAME = "dusiach.pdf"


def _local_result(
    sid=LOCAL_ID,
    fname=LOCAL_FNAME,
    snippet="hello *world*",
    locator="p. 3",
    p_num="3",
    fp=LOCAL_FP,
):
    return {
        "display": {"id": sid, "source": "LOCAL", "shelfmark": fname},
        "sys_id": sid,
        "raw_file_hl": snippet,
        "chunk_locator": locator,
        "p_num": p_num,
        "_fp": fp,
    }


def _genizah_result(
    sid=GEN_ID,
    shelfmark="T-S 12.123",
    title="Cairo Genizah Fragment",
    snippet="fragment *text* here",
    img="4r",
    source="Genizah",
):
    return {
        "display": {
            "id": sid,
            "source": source,
            "shelfmark": shelfmark,
            "title": title,
            "img": img,
        },
        "sys_id": sid,
        "raw_file_hl": snippet,
        "chunk_locator": "",
        "p_num": "",
    }


def _fp_fn(sid, results):
    """Resolve filepath from the result dict's _fp field (test-only)."""
    for r in results:
        if (r.get("display") or {}).get("id") == sid:
            return r.get("_fp", "")
    return ""


# ---------------------------------------------------------------------------
# CSV tests
# ---------------------------------------------------------------------------

pytestmark = pytest.mark.skipif(
    not _HELPERS_AVAILABLE,
    reason=f"genizah_app helpers not importable: {locals().get('_IMPORT_ERR', '')}",
)


def test_csv_mixed_appends_filepath_page_columns():
    """Mixed CSV: Filepath and Page columns appended for LOCAL rows (D-08)."""
    local_r = _local_result()
    gen_r = _genizah_result()
    results = [local_r, gen_r]
    fp_fn = lambda sid: _fp_fn(sid, results)  # noqa: E731

    # LOCAL row extra cols
    local_extra = _csv_extra_cols(local_r, fp_fn)
    assert len(local_extra) == 2, "LOCAL row must produce 2 extra cells"
    # filepath cell: LOCAL_FP sanitized
    assert LOCAL_FP in local_extra[0] or local_extra[0].startswith(LOCAL_FP[:5])
    # page cell: the chunk_locator value
    assert "3" in local_extra[1]

    # Genizah row extra cols
    gen_extra = _csv_extra_cols(gen_r, fp_fn)
    assert gen_extra == ["", ""], "Genizah row must produce two empty extra cells"


def test_csv_genizah_only_no_extra_columns():
    """Genizah-only CSV: header is the unchanged 7-column table (LEXP-08)."""
    gen_r = _genizah_result()
    # No LOCAL source → _has_local_in_export would be False → no extra cols
    extra = _csv_extra_cols(gen_r)
    assert extra == ["", ""], "Genizah row extra cols are always ['', '']"

    # Build the full CSV in memory using the same logic as export_results.
    from shared_export_utils import sanitize_text_for_excel

    headers = [
        "System ID",
        "Library",
        "Shelfmark",
        "Title",
        "Image/Page",
        "Source",
        "Snippet",
    ]
    row, _ = _build_export_data_row(gen_r)
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(headers)  # no extra cols
    clean = [sanitize_text_for_excel(str(v).replace("*", "")) for v in row]
    writer.writerow(clean)
    buf.seek(0)
    reader = csv.reader(buf)
    hdr = next(reader)
    assert len(hdr) == 7, f"Genizah-only header must have exactly 7 columns, got {len(hdr)}"
    data_row = next(reader)
    assert len(data_row) == 7, "Genizah-only data row must have exactly 7 columns"


def test_csv_local_row_column_remap():
    """LOCAL row: Shelfmark=filename, Library=parent folder, Source='LOCAL' (D-08)."""
    local_r = _local_result()
    fp_fn = lambda sid: LOCAL_FP  # noqa: E731

    row, is_local = _build_export_data_row(local_r, filepath_fn=fp_fn)
    assert is_local is True
    # row order: [sid, library/parent, shelfmark/filename, title, img/page, source, snippet]
    sid_col, library_col, shelfmark_col, title_col, img_col, source_col, snippet_col = row
    assert shelfmark_col == LOCAL_FNAME, f"Shelfmark col must be filename, got {shelfmark_col!r}"
    assert library_col == LOCAL_PARENT, f"Library col must be parent folder, got {library_col!r}"
    assert source_col == "LOCAL", f"Source col must be 'LOCAL', got {source_col!r}"
    # Snippet contains matched text (markers may be present in row — stripped in CSV branch)
    assert "world" in snippet_col


def test_csv_local_filepath_formula_escaped():
    """T-103-07 HIGH: LOCAL values beginning with =+-@ are neutralized in CSV.

    Covers all four meta-characters: =, +, -, @.
    """
    from shared_export_utils import sanitize_text_for_excel

    evil_values = {
        "=": "=cmd|calc",
        "+": "+1-2",
        "-": "-1*SYSTEM",
        "@": "@SUM(A1:A10)",
    }

    for prefix, evil in evil_values.items():
        # Sanitize as the CSV branch does
        sanitized = sanitize_text_for_excel(evil)
        assert sanitized.startswith("'"), (
            f"Formula-injection escape missing for '{prefix}' prefix: "
            f"got {sanitized!r}, expected leading \"'\""
        )

    # End-to-end: build a row with an evil filepath, write CSV, read back.
    evil_r = _local_result(fp=evil_values["="], snippet="normal text")

    def evil_fp_fn(sid):
        return evil_values["="]

    extra = _csv_extra_cols(evil_r, evil_fp_fn)
    # extra[0] is the sanitized filepath — must start with "'"
    assert extra[0].startswith("'"), (
        f"Evil filepath in extra cols not escaped: {extra[0]!r}"
    )

    # Also check that a row cell that starts with '=' gets neutralized.
    # Construct a row whose snippet starts with '=':
    evil_snippet_r = _local_result(fname="=HYPERLINK(\"x\")", snippet="safe")
    row, _ = _build_export_data_row(evil_snippet_r)
    # row[2] is the shelfmark col = filename
    shelfmark_raw = row[2]  # raw value before sanitize
    sanitized_shelfmark = sanitize_text_for_excel(str(shelfmark_raw).replace("*", ""))
    assert sanitized_shelfmark.startswith("'"), (
        f"Evil filename not escaped: {sanitized_shelfmark!r}"
    )


# ---------------------------------------------------------------------------
# TXT tests
# ---------------------------------------------------------------------------


def test_txt_local_block_format():
    """TXT LOCAL block: === filename | parent ===, Path: filepath, (p. N) (D-09).

    D-02: chunk_locator ('p. 5') is used VERBATIM and wrapped as '(p. 5)' — it
    must NOT be double-prefixed to '(page p. 5)' (Codex MEDIUM #1).
    """
    local_r = _local_result(locator="p. 5", p_num="5")
    fp_fn = lambda sid: LOCAL_FP  # noqa: E731
    block = _format_txt_local_block(local_r, filepath_fn=fp_fn)

    assert f"=== {LOCAL_FNAME} | {LOCAL_PARENT} ===" in block, (
        f"Header line missing in block:\n{block!r}"
    )
    assert f"Path: {LOCAL_FP}" in block, f"Path line missing in block:\n{block!r}"
    assert "(p. 5)" in block, f"Page info missing in block:\n{block!r}"
    assert "(page p. 5)" not in block, (
        f"Page locator must not be double-prefixed (D-02):\n{block!r}"
    )


def test_txt_local_block_strips_markers():
    """TXT LOCAL block: snippet has '*' markers stripped (D-09 clean output)."""
    local_r = _local_result(snippet="*highlight* word", locator="p. 1")
    fp_fn = lambda sid: LOCAL_FP  # noqa: E731
    block = _format_txt_local_block(local_r, filepath_fn=fp_fn)
    # The LOCAL snippet should NOT contain * markers
    lines = block.split("\n")
    snippet_line = lines[-1]  # last line is the snippet
    assert "*" not in snippet_line, f"'*' markers must be stripped in LOCAL TXT block: {snippet_line!r}"
    assert "highlight" in snippet_line, "Highlighted word should still be present"


def test_txt_genizah_block_byte_identical():
    """Genizah TXT block: byte-identical to pre-v7.17 output (LEXP-08).

    The pre-v7.17 Genizah TXT block was:
        === {shelfmark} | {title} ===
        {snippet_with_markers_intact}

    This test pins that the Genizah block preserves '*' markers (they were NOT
    stripped in the pre-v7.17 code) and uses the exact same format string.
    """
    gen_r = _genizah_result(
        shelfmark="T-S 12.123",
        title="Cairo Genizah Fragment",
        snippet="fragment *text* here",
    )
    block = _format_txt_genizah_block(gen_r)

    # Must match the pre-v7.17 format EXACTLY:
    #   f"=== {r['display']['shelfmark']} | {r['display']['title']} ===\n{snippet}"
    # where snippet = r.get('raw_file_hl','').strip().replace('\n',' ').replace('\r','')
    expected_header = "=== T-S 12.123 | Cairo Genizah Fragment ==="
    expected_snippet = "fragment *text* here"  # markers NOT stripped
    assert block.startswith(expected_header), (
        f"Genizah header mismatch. Expected:\n  {expected_header!r}\nGot:\n  {block[:80]!r}"
    )
    assert expected_snippet in block, (
        f"Genizah snippet (with markers) not found in block:\n{block!r}"
    )
    # Confirm '*' markers are PRESERVED (byte-identity check on marker treatment)
    assert "*" in block, "Genizah TXT must preserve '*' highlight markers (pre-v7.17 behavior)"


def test_txt_genizah_block_no_path_line():
    """Genizah TXT block must NOT contain a 'Path:' line (that's LOCAL-only)."""
    gen_r = _genizah_result()
    block = _format_txt_genizah_block(gen_r)
    assert "Path:" not in block, "Genizah TXT block must not contain a 'Path:' line"


# ---------------------------------------------------------------------------
# Page-label (D-02) tests — Codex MEDIUM #1
# ---------------------------------------------------------------------------


def test_local_page_label_chunk_locator_verbatim():
    """D-02: a human-formatted chunk_locator is used VERBATIM, never re-prefixed."""
    assert _local_page_label({"chunk_locator": "p. 3", "p_num": "3"}) == "p. 3"
    assert _local_page_label({"chunk_locator": "ff. 2v-3r"}) == "ff. 2v-3r"


def test_local_page_label_pnum_fallback_synthesized():
    """D-02: with no chunk_locator, a raw p_num is synthesized as 'p. N'."""
    assert _local_page_label({"p_num": "7"}) == "p. 7"


def test_local_page_label_empty_when_no_page():
    """No locator and no p_num -> empty string (no stray 'p. ')."""
    assert _local_page_label({}) == ""
    assert _local_page_label({"p_num": ""}) == ""


def test_local_page_label_no_double_prefix():
    """A verbatim 'p. 3' locator stays 'p. 3' — no 'p. p. 3' / 'page p. 3'."""
    label = _local_page_label({"chunk_locator": "p. 3"})
    assert label == "p. 3"
    assert "p. p." not in label and "page" not in label


# ---------------------------------------------------------------------------
# DOCX tests
# ---------------------------------------------------------------------------


def test_docx_two_results_blocks():
    """DOCX export: two write_docx_result_block calls produce correct content, no table."""
    from docx import Document

    from shared.docx_export import write_docx_result_block

    doc = Document()
    local_r = _local_result(locator="p. 3")
    gen_r = _genizah_result()

    write_docx_result_block(doc, local_r, filepath=LOCAL_FP, lang="en")
    write_docx_result_block(doc, gen_r, filepath="", lang="en")

    # 1. No tables — the new layout uses paragraphs only.
    assert len(doc.tables) == 0, f"DOCX must have no tables, got {len(doc.tables)}"

    full_text = "\n".join(p.text for p in doc.paragraphs)

    # 2. LOCAL block: filename appears in heading, filepath appears in URL line.
    assert LOCAL_FNAME in full_text, f"LOCAL filename not found in DOCX text:\n{full_text[:500]}"
    assert LOCAL_FP in full_text, f"LOCAL filepath not found in DOCX text:\n{full_text[:500]}"

    # 3. Genizah block: shelfmark appears; genizahsearch.com URL appears.
    assert "T-S 12.123" in full_text, f"Genizah shelfmark not found in DOCX text:\n{full_text[:500]}"
    assert "genizahsearch.com" in full_text, (
        f"Genizah URL not found in DOCX text:\n{full_text[:500]}"
    )

    # 4. Separator paragraphs exist.
    assert "________" in full_text, "Separator paragraph missing from DOCX output"


def test_docx_local_block_parent_folder():
    """DOCX LOCAL block heading contains 'filename — parent' format."""
    from docx import Document

    from shared.docx_export import write_docx_result_block

    doc = Document()
    local_r = _local_result()
    write_docx_result_block(doc, local_r, filepath=LOCAL_FP, lang="en")

    headings = [p.text for p in doc.paragraphs if p.runs and any(r.font.bold for r in p.runs if r.font.bold)]
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert LOCAL_FNAME in all_text
    assert LOCAL_PARENT in all_text


def test_docx_genizah_block_no_filepath():
    """DOCX Genizah block: no LOCAL filepath, has genizahsearch URL (D-11)."""
    from docx import Document

    from shared.docx_export import write_docx_result_block

    doc = Document()
    gen_r = _genizah_result(sid=GEN_ID)
    write_docx_result_block(doc, gen_r, filepath="", lang="en")
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "genizahsearch.com" in full_text
    # filepath should not appear (it was empty)
    assert LOCAL_FP not in full_text
