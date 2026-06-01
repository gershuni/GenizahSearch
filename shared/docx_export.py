# -*- coding: utf-8 -*-
"""Phase 103 (D-10/D-11): reusable per-result DOCX block writer.

Module-level (no GenizahGUI / Qt dependency) so both genizah_app.py::export_results
and Phase 104's export_comp_report can render one result as a "research handout"
block instead of a cramped table row.
"""
from __future__ import annotations

import logging

logger = logging.getLogger(__name__)

_GENIZAH_URL_BASE = "https://genizahsearch.com"


def _add_highlighted_runs(paragraph, text):
    """Split ``text`` on ``*`` markers; odd-index parts become bold red."""
    from docx.shared import RGBColor

    parts = str(text or "").split("*")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.bold = True


def _set_paragraph_rtl(paragraph):
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    bidi.set(qn("w:val"), "1")


def _genizah_url_for(sys_id):
    if not sys_id:
        return ""
    return f"{_GENIZAH_URL_BASE}/?sys_id={sys_id}"


def _page_label(result_dict):
    """D-02: use ``chunk_locator`` VERBATIM (already e.g. 'p. 3'); only the raw
    ``p_num`` fallback gets a synthesized 'p. ' label. Never double-prefix."""
    locator = result_dict.get("chunk_locator")
    if locator:
        return str(locator)
    p_num = result_dict.get("p_num")
    if p_num:
        return f"p. {p_num}"
    return ""


def write_docx_result_block(doc, result_dict, filepath: str = "", lang: str = "en", full_text: str = "") -> None:
    """Write one per-result block to a python-docx Document (Phase 103 D-10/D-11).

    Block layout (both Genizah and LOCAL):
      - Heading paragraph (bold): ``{Shelfmark} — {Title}`` (Genizah) or
        ``{Filename} — {Parent folder}`` (LOCAL)
      - Metadata line: ``{Library} · {Image/Page} · {Source}`` (Genizah) or
        ``{full filepath} · {page locator} · LOCAL`` (LOCAL)
      - Matched-text paragraph with bold-red ``*``-highlights
      - URL line: GenizahSearch URL (Genizah) or full filepath (LOCAL) (D-11)
      - A thin separator paragraph

    ``filepath`` is the pre-resolved LOCAL filepath (caller resolves via the
    batch-primed cache; never looked up here). ``lang='he'`` applies RTL.
    The page segment uses ``chunk_locator`` VERBATIM (D-02) — no double-prefix.
    Designed to be reusable by Phase 104 export_comp_report.
    """
    import os

    d = result_dict.get("display") or {}
    is_local = d.get("source") == "LOCAL"
    raw_hl = result_dict.get("raw_file_hl", "") or result_dict.get("snippet", "") or ""
    # EXPUX-04: expand to the fuller matched passage (capped + highlighted) when
    # full_text is available; otherwise keep the ±60-char snippet. Expand BEFORE
    # newline-collapsing so the matched-term recovery sees the original markers.
    if full_text:
        from shared_export_utils import build_expanded_context
        raw_hl = build_expanded_context(full_text, raw_hl)
    # Strip XML-illegal control chars (NUL / form-feed etc. from page text) so
    # python-docx/lxml can't abort the export with "All strings must be XML
    # compatible …". Covers the expanded-context AND the ±60 fallback paths.
    from shared_export_utils import strip_xml_illegal_chars
    raw_hl = strip_xml_illegal_chars(str(raw_hl).replace("\n", " ").replace("\r", " "))
    # D-02 page label (chunk_locator verbatim; p_num fallback only).
    page = _page_label(result_dict)

    paras = []
    if is_local:
        sid = d.get("id") or result_dict.get("sys_id") or ""
        filename = d.get("shelfmark") or sid
        parent = os.path.basename(os.path.dirname(filepath)) if filepath else ""
        heading_text = f"{filename} — {parent}" if parent else f"{filename}"
        # page is already 'p. 3' (D-02) — use verbatim, do NOT wrap in 'page {..}'.
        meta_text = " · ".join([p for p in [filepath, page, "LOCAL"] if p])
        url_text = filepath or ""
    else:
        shelfmark = d.get("shelfmark", "") or ""
        title = d.get("title", "") or ""
        library = d.get("library_code", "") or d.get("library", "") or ""
        img = str(d.get("img", "") or page or "")
        source = d.get("source", "") or "Genizah"
        heading_text = f"{shelfmark} — {title}" if title else f"{shelfmark}"
        meta_text = " · ".join([p for p in [library, img, source] if p])
        url_text = _genizah_url_for(d.get("id") or result_dict.get("sys_id") or "")

    # Sanitize heading/meta/url too — filenames/titles are low-risk but a NUL
    # anywhere aborts the whole doc, so be defensive across every added string.
    h = doc.add_paragraph(strip_xml_illegal_chars(heading_text))
    if h.runs:
        h.runs[0].font.bold = True
    paras.append(h)
    if meta_text:
        paras.append(doc.add_paragraph(strip_xml_illegal_chars(meta_text)))
    body = doc.add_paragraph()
    _add_highlighted_runs(body, raw_hl)  # raw_hl already sanitized above
    paras.append(body)
    if url_text:
        paras.append(doc.add_paragraph(strip_xml_illegal_chars(url_text)))
    sep = doc.add_paragraph("_" * 40)
    paras.append(sep)

    if lang == "he":
        for p in paras:
            try:
                _set_paragraph_rtl(p)
            except Exception:
                logger.debug("RTL apply failed on a DOCX block paragraph", exc_info=True)
