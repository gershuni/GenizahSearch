# -*- coding: utf-8 -*-
"""Phase 95 LOCAL indexer (REQ-1, REQ-3, REQ-4, REQ-5, D-01..D-44).

Qt-free. Owns: PyMuPDF/python-docx/TXT extraction, LOCAL Tantivy side-index
+ LOCAL LAB side-index builders, SQLite cache (folders / processed_files /
local_pages / local_files), incremental re-extract, delete-by-uid, two-phase
commit with crash recovery.

Critical divergences from main Tantivy index (CONTEXT D-08 / D-13 / D-20 / D-21
+ RESEARCH Pitfall #2):
  - unique_id field uses tokenizer_name="raw" (main index omits this kwarg).
    Required because LOCAL deletes-by-term on rescan; main never deletes.
  - One Tantivy doc per PDF page (D-03), 20-paragraph chunk for DOCX (D-04).
  - LOCAL hits never write to shared corpus files (Seewald's prototype patched
    the shared corpus - we don't; SPEC Constraint #6).

Per CONTEXT D-02: _fix_rtl_line, _fix_rtl_page, _join_fragmented_lines are
PORTED VERBATIM from seewald_addition/genizah_make_index.py:67-105 but NEVER
invoked at runtime. They exist as a regression-prevention contract for the
future fallback path (pdfplumber/pypdf). Tests in tests/test_local_indexer.py
exercise the helpers in isolation.

HIGH-3 review fix: _delete_file uses crash-safe THREE-STEP protocol:
  1. Mark pending_delete=1 in SQLite
  2. Tantivy delete + commit
  3. SQLite final DELETE
Startup recovery calls _recover_pending_deletes() to replay steps 2-3 for
any rows left with pending_delete=1 from a previous crash.

HIGH-4 review fix (option b LOCKED): page text for LAB rebuild is read from
main LOCAL Tantivy content stored field by UID - NOT from SQLite (this plan's
schema stores metadata only).

MEDIUM-2 review fix: TXT extraction tries utf-8-sig with errors=strict; on
UnicodeDecodeError, attempts cp1255 (legacy Windows Hebrew); on both failures,
the file is marked extraction_status='encoding_error'.
"""
from __future__ import annotations

import datetime
import errno
import gc
import hashlib
import json
import logging
import os
import re
import sqlite3
import threading
import time
import unicodedata
from typing import Callable, Iterator, Optional

# Phase 97 F-02 — neutralize XML-bomb / XXE primitives before openpyxl loads.
# defuse_stdlib() patches xml.etree.ElementTree and friends. Best-effort:
# if defusedxml is somehow missing in a dev install, log a warning and proceed.
try:
    import defusedxml
    defusedxml.defuse_stdlib()
except ImportError:
    import logging as _logging
    _logging.getLogger(__name__).warning(
        "Phase 97: defusedxml not installed; XML-bomb defense degraded"
    )

import fitz  # PyMuPDF - D-01
from docx import Document as _DocxDoc  # python-docx - D-04
import tantivy

from shared.local_sys_id import (
    is_local_sys_id,  # noqa: F401
    generate_local_sys_id,
    _canonical_filepath,
)

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
_SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".txt", ".html", ".xlsx", ".csv"}
# Phase 97 D-NEW-4 — error statuses that keep a row even for unsupported extensions
_ERROR_STATUSES_KEPT = {
    "oversized", "error", "encoding_error",
    "changed_during_index", "zip_bomb_suspected",
    "unreachable", "timeout",
}
_DOCX_CHUNK_PARAGRAPHS = 20            # D-04
_SCANNED_PDF_CHAR_THRESHOLD = 50       # D-05
_EMPTY_PAGE_CHAR_THRESHOLD = 10        # D-06
_COMMIT_BATCH_SIZE = 25                # D-21 — Phase 95 dead-code per D-02 preservation;
                                       # Phase 95 D-21 — superseded by Phase 97 _CommitTriggers below.
                                       # Kept as dead-code constant per D-02 preservation.
_MAX_COLLISION_RETRIES = 4             # D-19
_UNIQUE_ID_PREFIX = "LOCAL"            # D-34
_DEFAULT_TXT_ENCODING = "utf-8-sig"    # D-07 starting policy

# ---------------------------------------------------------------------------
# Phase 97 C-02 / C-05 constants + commit trigger + zip-bomb helpers
# ---------------------------------------------------------------------------
# Phase 97 C-02 — heap-sampling branch DROPPED per RESEARCH Issue #1
# (tantivy-py 0.25.1 has no writer.get_memory_usage()); the heap_size= arg
# to index.writer() is a memory ceiling, not a commit trigger.
# TODO(tantivy >= 0.26): add a heap-size trigger when get_memory_usage() exists.

# Phase 97 C-05 size/zip-bomb constants
_MAX_FILE_SIZE = 100 * 1024 * 1024            # raw 100 MB hard skip
_MAX_UNCOMPRESSED_BYTES = 500 * 1024 * 1024   # zip-container 500 MB
_MAX_CELLS_PER_SHEET = 100_000                # 100K cells (consumed by Wave C XLSX extractor)
_MAX_CHARS_PER_CHUNK = 1_000_000              # 1 MB per chunk (consumed by Wave C extractors)
_XLSX_ROW_WINDOW = 500                        # F-02: rows per Tantivy chunk in XLSX extractor
_CSV_ROW_WINDOW = 200                         # F-03: rows per Tantivy chunk in CSV extractor
_CSV_ENCODINGS = ("utf-8-sig", "cp1255", "utf-16-le")  # F-05: encoding chain for CSV


class XlsxZipBombSuspected(Exception):
    """Phase 97 C-05 — raised by _check_zip_bomb or XLSX extractor when limits exceeded."""


class _PhaseAwareETA:
    """Phase 97 U-01 — separate EWMA smoothing for 4 sub-phases.

    compose_overall_eta returns sum() of remaining-phase ETAs because phases
    are sequential (walk -> extract -> commit -> rebuild). Codex MEDIUM #5
    resolved: test and code agree on sum composition.
    """

    PHASES = ("walking", "extracting", "committing", "rebuilding_lab")
    ALPHA = 0.2  # EWMA weight for new samples

    def __init__(self) -> None:
        import time as _time
        self._bytes_per_sec: dict = {p: 0.0 for p in self.PHASES}
        self._remaining_bytes: dict = {p: 0 for p in self.PHASES}
        self._last_sample: dict = {p: _time.monotonic() for p in self.PHASES}

    def record(self, phase: str, bytes_done: int) -> None:
        """Feed a new bytes_done observation into the EWMA for the given phase."""
        import time as _time
        now = _time.monotonic()
        dt = max(now - self._last_sample[phase], 1e-6)
        rate = bytes_done / dt
        self._bytes_per_sec[phase] = (
            self.ALPHA * rate + (1 - self.ALPHA) * self._bytes_per_sec[phase]
        )
        self._last_sample[phase] = now

    def set_remaining(self, phase: str, bytes_remaining: int) -> None:
        """Set the estimated remaining bytes for a phase."""
        self._remaining_bytes[phase] = bytes_remaining

    def phase_eta_seconds(self, phase: str) -> float:
        """Return ETA in seconds for a single phase."""
        rate = self._bytes_per_sec[phase]
        if rate <= 0:
            return float("inf") if self._remaining_bytes[phase] > 0 else 0.0
        return self._remaining_bytes[phase] / rate

    def compose_overall_eta(self) -> float:
        """Sequential composition — sum of remaining phase ETAs (Codex MEDIUM #5 fix).

        Returns the sum of per-phase ETAs for all phases that still have
        remaining bytes. Returns 0.0 when no bytes remain.
        """
        etas = [
            self.phase_eta_seconds(p)
            for p in self.PHASES
            if self._remaining_bytes[p] > 0
        ]
        return sum(etas) if etas else 0.0


class _CommitTriggers:
    """Phase 97 C-02 — byte/count/time commit triggers.

    Fires when ANY of the following is reached:
      - Source bytes in this batch >= 200 MB (BYTES_THRESHOLD)
      - Files processed in this batch >= 100 (FILES_THRESHOLD)
      - Seconds elapsed since last commit >= 60 (SECONDS_THRESHOLD)

    Heap-sampling path DROPPED per RESEARCH Issue #1: tantivy-py 0.25.1 has
    no writer.get_memory_usage() method. The heap_size= arg to index.writer()
    is a memory ceiling, not a commit trigger.
    TODO(tantivy >= 0.26): add heap trigger when get_memory_usage() exists.
    """

    BYTES_THRESHOLD = 200 * 1024 * 1024  # 200 MB
    FILES_THRESHOLD = 100
    SECONDS_THRESHOLD = 60.0

    def __init__(self) -> None:
        self._batch_bytes = 0
        self._batch_files = 0
        self._batch_start = time.monotonic()

    def record_file(self, source_size: int) -> None:
        """Record that one file of source_size bytes was processed."""
        self._batch_bytes += int(source_size or 0)
        self._batch_files += 1

    def should_commit(self) -> bool:
        """Return True if any commit threshold has been crossed."""
        return (
            self._batch_bytes >= self.BYTES_THRESHOLD
            or self._batch_files >= self.FILES_THRESHOLD
            or (time.monotonic() - self._batch_start) >= self.SECONDS_THRESHOLD
        )

    def reset(self) -> None:
        """Reset all counters after a commit."""
        self._batch_bytes = 0
        self._batch_files = 0
        self._batch_start = time.monotonic()


def _check_zip_bomb(filepath: str) -> str | None:
    """Phase 97 C-05 — pre-check zip-container files for uncompressed-size bombs.

    Iterates zipfile.ZipInfo.file_size records and sums uncompressed sizes.
    Returns a reason string if the total exceeds _MAX_UNCOMPRESSED_BYTES (500 MB),
    or if the file is not a valid zip container. Returns None if safe.

    Apply ONLY to .docx / .xlsx — PDF/TXT/CSV/HTML use raw-size cap only.
    This check runs BEFORE handing the file to openpyxl/python-docx, so a
    maliciously crafted zip claiming 600 MB uncompressed is rejected in < 1 ms
    without allocating any decompression memory (Codex MEDIUM #1).
    """
    import zipfile as _zipfile
    try:
        with _zipfile.ZipFile(filepath, "r") as zf:
            total_uncompressed = sum(info.file_size for info in zf.infolist())
            if total_uncompressed > _MAX_UNCOMPRESSED_BYTES:
                return (
                    f"uncompressed size {total_uncompressed} "
                    f"exceeds limit {_MAX_UNCOMPRESSED_BYTES}"
                )
    except _zipfile.BadZipFile:
        return "not a valid zip-container file"
    return None


# ---------------------------------------------------------------------------
# EncodingError (MEDIUM-2 review fix)
# ---------------------------------------------------------------------------
class EncodingError(Exception):
    """Raised by extract_txt when both utf-8-sig and cp1255 fail (MEDIUM-2)."""
    pass


# ---------------------------------------------------------------------------
# RTL helpers - DEAD CODE per D-02.
# Ported VERBATIM from seewald_addition/genizah_make_index.py:67-105.
# These helpers are NEVER invoked in v1 runtime. They exist as a
# regression-prevention contract for a future pdfplumber/pypdf fallback path.
# ---------------------------------------------------------------------------

def _rtl_ratio(text: str) -> float:  # pragma: no cover
    """Fraction of RTL chars among alpha chars. DEAD CODE per D-02."""
    alpha = [c for c in text if c.isalpha()]
    if not alpha:
        return 0.0
    rtl = sum(1 for c in alpha if unicodedata.bidirectional(c) in ("R", "AL", "AN"))
    return rtl / len(alpha)


def _fix_rtl_line(line: str) -> str:  # pragma: no cover
    """Reverse a pdfplumber mirror-reversed RTL line. DEAD CODE per D-02."""
    s = line.strip()
    if not s or _rtl_ratio(s) <= 0.4:
        return line
    lead = len(line) - len(line.lstrip())
    tail = len(line) - len(line.rstrip())
    core = s[::-1]
    return line[:lead] + core + (line[len(line) - tail:] if tail else "")


def _fix_rtl_page(text: str) -> str:  # pragma: no cover
    """Apply per-line RTL fix + re-glue punctuation. DEAD CODE per D-02."""
    if not text:
        return text
    lines = [_fix_rtl_line(ln) for ln in text.splitlines()]
    result = "\n".join(lines)
    result = re.sub(r"(\w)\s+([,.])", r"\1\2", result)
    result = re.sub(r"([,.])\s+(\w)", r"\1 \2", result)
    return result


def _join_fragmented_lines(text: str) -> str:  # pragma: no cover
    """Join pages where each word is on its own line. DEAD CODE per D-02."""
    lines = text.splitlines()
    non_empty = [line for line in lines if line.strip()]
    if len(non_empty) < 4:
        return text
    single = sum(1 for line in non_empty if len(line.split()) <= 1)
    if single / len(non_empty) < 0.60:
        return text
    paragraphs, current = [], []
    for line in lines:
        s = line.strip()
        if s:
            current.append(s)
        elif current:
            paragraphs.append(" ".join(current))
            current = []
    if current:
        paragraphs.append(" ".join(current))
    return "\n\n".join(paragraphs)


# ---------------------------------------------------------------------------
# Phase 96 D-F4: detect pathological one-word-per-line PDF extraction
# ---------------------------------------------------------------------------
# This is the LIVE detection used by extract_pdf_pages (NOT dead code, unlike
# its dead-code cousin _join_fragmented_lines above). We use 0.70 (vs the
# dead-code's 0.60) per RESEARCH §2 — more conservative threshold preserves
# documents with legitimate one-word paragraphs (chapter numbers, table cells).

_SINGLE_WORD_RATIO_THRESHOLD = 0.70  # Phase 96 D-F4, RESEARCH §2
_SINGLE_WORD_MIN_SAMPLE = 5          # below this, do not trigger detection


def _detect_single_word_per_line(text: str) -> bool:
    """Phase 96 D-F4: return True if `text` looks like pathological
    one-word-per-line output from PyMuPDF's `get_text("blocks")` mode.

    Heuristic:
      - split on newlines, keep non-empty lines after .strip()
      - if fewer than 5 non-empty lines: return False (small-sample guard;
        could be a title page or table cell)
      - compute single_word_ratio = (# lines with <= 1 word) / (# non-empty lines)
      - return True iff ratio >= 0.70

    Examples:
        >>> _detect_single_word_per_line("")
        False
        >>> _detect_single_word_per_line("one\\ntwo\\nthree\\nfour\\n")
        False
        >>> _detect_single_word_per_line("one\\ntwo\\nthree\\nfour\\nfive\\nsix\\n")
        True
        >>> _detect_single_word_per_line("the quick brown fox\\njumps over the\\nlazy dog under\\nthe bright morning\\nsun rises\\nslowly today")
        False
    """
    if not text:
        return False
    lines = text.splitlines()
    non_empty = [ln for ln in lines if ln.strip()]
    if len(non_empty) < _SINGLE_WORD_MIN_SAMPLE:
        return False
    single = sum(1 for ln in non_empty if len(ln.split()) <= 1)
    ratio = single / len(non_empty)
    return ratio >= _SINGLE_WORD_RATIO_THRESHOLD


# ---------------------------------------------------------------------------
# Schema builders
# ---------------------------------------------------------------------------

def build_local_schema() -> tantivy.Schema:
    """LOCAL Tantivy side-index schema (REQ-3).

    CRITICAL DIVERGENCE: unique_id uses tokenizer_name="raw" (main index at
    genizah_core.py:5125 omits this kwarg). Required for delete_documents()
    to work on rescan - RESEARCH Pitfall #2 / tantivy-py #297.

    Without tokenizer_name="raw", writer.delete_documents("unique_id", uid)
    SILENTLY does nothing on rescans and doubles the page-row count on every
    modify. Schema-divergence comment is required per PATTERNS.md.
    """
    builder = tantivy.SchemaBuilder()
    # CRITICAL: tokenizer_name="raw" - main index at genizah_core.py:5125 omits this
    # and is rebuilt from scratch so the bug stays latent. For LOCAL where incremental
    # delete IS the central operation, raw is mandatory. tantivy-py issue #297.
    builder.add_text_field("unique_id", stored=True, tokenizer_name="raw")
    builder.add_text_field("content", stored=True, tokenizer_name="whitespace")
    builder.add_text_field("content_head", stored=False, tokenizer_name="whitespace")
    builder.add_text_field("content_tail", stored=False, tokenizer_name="whitespace")
    builder.add_text_field("line_starts", stored=False, tokenizer_name="whitespace")
    builder.add_text_field("line_ends", stored=False, tokenizer_name="whitespace")
    builder.add_text_field("source", stored=True)
    builder.add_text_field("full_header", stored=True)
    builder.add_text_field("shelfmark", stored=True)
    builder.add_text_field("scope", stored=True)
    builder.add_text_field("boundaries", stored=True)
    # Phase 97 LD-1 — two new fields defined once here; Waves E/F populate, do NOT add more.
    # U-02: raw tokenizer REQUIRED so writer.delete_documents("scan_run_id", run_id)
    # matches; default tokenizer would tokenize the UUID and never hit a Term lookup.
    builder.add_text_field("scan_run_id", stored=True, tokenizer_name="raw")
    # D-NEW-5: human-readable per-chunk display string; stored only, no search.
    builder.add_text_field("chunk_locator", stored=True)
    return builder.build()


# ---------------------------------------------------------------------------
# Phase 97 LD-1: schema marker helpers (detect schema drift, trigger rebuild)
# ---------------------------------------------------------------------------

def _compute_schema_marker(schema_fn=None) -> str:
    """Return a short hex hash of build_local_schema source.

    When build_local_schema() changes (new fields added), the marker changes
    and triggers atomic rebuild so old on-disk indexes with wrong field sets
    are replaced rather than silently opened with a mismatched schema.

    Uses inspect.getsource for simplicity — acceptable because the schema
    function body is stable between minor commits and any meaningful field
    change bumps the marker as intended.
    """
    import hashlib
    import inspect
    fn = schema_fn if schema_fn is not None else build_local_schema
    src = inspect.getsource(fn)
    return hashlib.sha256(src.encode("utf-8")).hexdigest()[:16]


def _read_schema_marker(index_dir: str) -> str | None:
    """Read the .schema_version marker from index_dir; returns None if absent."""
    p = os.path.join(index_dir, ".schema_version")
    if not os.path.isfile(p):
        return None
    try:
        with open(p, "r", encoding="utf-8") as f:
            return f.read().strip() or None
    except OSError:
        return None


def _write_schema_marker(index_dir: str, marker: str) -> None:
    """Write .schema_version marker file alongside the index dir."""
    os.makedirs(index_dir, exist_ok=True)
    with open(os.path.join(index_dir, ".schema_version"), "w", encoding="utf-8") as f:
        f.write(marker)


def build_local_lab_schema() -> tantivy.Schema:
    """LOCAL LAB side-index schema (D-09).

    Same raw-tokenizer divergence on unique_id as the main LOCAL schema.
    fingerprint_dyn computed at build time using current LAB dynamic_rank_map.
    """
    builder = tantivy.SchemaBuilder()
    # CRITICAL: tokenizer_name="raw" - same divergence as build_local_schema().
    builder.add_text_field("unique_id", stored=True, tokenizer_name="raw")  # Pitfall #2
    builder.add_text_field("text_normalized", stored=True, tokenizer_name="simple")
    builder.add_text_field("text_ngram", stored=False, tokenizer_name="whitespace")
    # NOTE: main LAB uses self.LAB_FINGERPRINT_FIELD; we mirror as "fingerprint".
    builder.add_text_field("fingerprint", stored=False, tokenizer_name="simple")
    builder.add_text_field("fingerprint_dyn", stored=False, tokenizer_name="simple")
    builder.add_text_field("full_header", stored=True)
    builder.add_text_field("shelfmark", stored=True)
    builder.add_text_field("source", stored=True)
    builder.add_text_field("content", stored=True, tokenizer_name="simple")
    return builder.build()


# ---------------------------------------------------------------------------
# Phase 97 R-03: zstd cached_text compression helpers
# ---------------------------------------------------------------------------

_ZSTD_LEVEL = 3  # Phase 97 R-03 — balance per RESEARCH benchmark (3-5 typical)


def compress_cached_text(text: str) -> tuple[bytes, int]:
    """Compress text with zstd for storage in local_pages.cached_text.

    Returns (zstd_compressed_bytes, uncompressed_len_in_bytes).
    zstandard must be importable (pinned in requirements.txt as of Phase 97).
    """
    import zstandard
    payload = text.encode("utf-8")
    compressed = zstandard.ZstdCompressor(level=_ZSTD_LEVEL).compress(payload)
    return compressed, len(payload)


def decompress_cached_text(blob: bytes) -> str:
    """Decompress zstd-compressed bytes back to text string."""
    import zstandard
    return zstandard.ZstdDecompressor().decompress(blob).decode("utf-8")


# ---------------------------------------------------------------------------
# SQLite schema init (D-35)
# ---------------------------------------------------------------------------

def init_sqlite(db_path: str) -> sqlite3.Connection:
    """Create tables per D-35 (Phase 95 baseline) + Phase 97 columns and return an open connection.

    Tables:
      - folders: registered source folders with status tracking (+ Phase 97 counter cols)
      - processed_files: narrow mtime-cache; status for two-phase commit (+ scan_run_id, mtime_ns)
      - local_pages: per-page UID tracking + Phase 97 cached_text + chunk_locator
      - local_files: rich metadata for Browse panel + per-file status panel
      - scan_runs: Phase 97 R-01 lifecycle table (replaces _pending_cleanup sentinel)
      - pending_dir_cleanup: Phase 97 R-02 GC table for .old-<ts> rebuild dirs

    Phase 97 LD-2: Fresh installs get the FULL Phase 97 schema directly (user_version=2).
    Pre-existing Phase 95 DBs at user_version=0 are migrated by local_indexer_migrations.run()
    called in LocalIndexer.__init__ immediately after init_sqlite().
    """
    conn = sqlite3.connect(db_path)
    conn.execute("PRAGMA journal_mode=WAL")  # Pitfall #6 mitigation
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS folders (
            folder_id        INTEGER PRIMARY KEY AUTOINCREMENT,
            path             TEXT    UNIQUE NOT NULL,
            added_at         REAL,
            last_scanned_at  REAL,
            status           TEXT    NOT NULL DEFAULT 'active',
            indexed_count    INTEGER NOT NULL DEFAULT 0,
            error_count      INTEGER NOT NULL DEFAULT 0,
            pending_count    INTEGER NOT NULL DEFAULT 0,
            oversized_count  INTEGER NOT NULL DEFAULT 0,
            last_aggregate_at REAL
        );
        CREATE TABLE IF NOT EXISTS processed_files (
            filepath     TEXT    PRIMARY KEY,
            mtime        REAL,
            size         INTEGER,
            sys_id       TEXT,
            status       TEXT    NOT NULL DEFAULT 'committed',
            scan_run_id  TEXT,
            mtime_ns     INTEGER
        );
        CREATE TABLE IF NOT EXISTS local_pages (
            sys_id                      TEXT    NOT NULL,
            uid                         TEXT    NOT NULL,
            page_num                    INTEGER NOT NULL,
            cached_text                 BLOB,
            cached_text_codec           TEXT    NOT NULL DEFAULT 'zstd',
            cached_text_uncompressed_len INTEGER,
            extraction_format_version   INTEGER NOT NULL DEFAULT 1,
            chunk_locator               TEXT,
            PRIMARY KEY (sys_id, page_num)
        );
        CREATE TABLE IF NOT EXISTS local_files (
            file_id           INTEGER PRIMARY KEY AUTOINCREMENT,
            sys_id            TEXT    NOT NULL UNIQUE,
            filepath          TEXT    NOT NULL,
            folder_id         INTEGER NOT NULL REFERENCES folders(folder_id),
            display_title     TEXT,
            original_filename TEXT    NOT NULL,
            file_extension    TEXT    NOT NULL,
            page_count        INTEGER NOT NULL DEFAULT 0,
            file_size_bytes   INTEGER NOT NULL,
            extraction_status TEXT    NOT NULL,
            last_indexed_at   REAL    NOT NULL,
            sha256_full       TEXT,
            error_msg         TEXT,
            pending_delete    INTEGER NOT NULL DEFAULT 0
        );
        CREATE TABLE IF NOT EXISTS scan_runs (
            scan_run_id  TEXT PRIMARY KEY,
            started_at   REAL NOT NULL,
            ended_at     REAL,
            status       TEXT NOT NULL CHECK (status IN ('running', 'completed', 'canceled', 'discarded'))
        );
        CREATE TABLE IF NOT EXISTS pending_dir_cleanup (
            path        TEXT PRIMARY KEY,
            kind        TEXT NOT NULL,
            created_at  REAL NOT NULL DEFAULT (strftime('%s','now'))
        );
    """)
    conn.commit()
    # Phase 97 LD-2: stamp target user_version on fresh (empty) DBs so they
    # skip the migration ladder on first open.
    # Pre-existing Phase 95 DBs already have data so this guard yields False
    # for them — LocalIndexer.__init__ will call migrations.run(conn) for those.
    existing_files = conn.execute("SELECT COUNT(*) FROM processed_files").fetchone()[0]
    if existing_files == 0:
        conn.execute("PRAGMA user_version = 2")
        conn.commit()
    # else: leave user_version at 0; LocalIndexer.__init__ calls migrations.run(conn).
    return conn


# ---------------------------------------------------------------------------
# Folder overlap detection (D-17 + D-42)
# ---------------------------------------------------------------------------

def check_folder_overlap(
    candidate: str,
    existing_paths: list[str],
) -> Optional[str]:
    """Return the conflicting existing path if candidate overlaps, else None.

    Overlap = same path, ancestor, or descendant.
    Uses _canonical_filepath + os.path.normcase + os.path.commonpath per D-17.
    """
    cand = _canonical_filepath(candidate)
    for existing in existing_paths:
        exist = _canonical_filepath(existing)
        if cand == exist:
            return existing
        try:
            common = os.path.normcase(os.path.commonpath([cand, exist]))
        except ValueError:
            # Different drives on Windows
            continue
        if common == cand or common == exist:
            return existing
    return None


# ---------------------------------------------------------------------------
# UID + full_header builders (D-34)
# ---------------------------------------------------------------------------

def _make_uid(sys_id: str, page_num: int) -> str:
    """Construct the Tantivy unique_id for a LOCAL page doc (D-34)."""
    return f"{_UNIQUE_ID_PREFIX}_{sys_id}_P{page_num}"


def _make_full_header(sys_id: str, page_num: int, file_id: int) -> str:
    """Construct the full_header stored field value for a LOCAL page doc (D-34)."""
    return f"{sys_id}_LOCAL_P{page_num}_F{file_id:04d}"


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------

def extract_pdf_pages(
    filepath: str,
) -> Iterator[tuple[int, str, str]]:
    """Extract text page-by-page using PyMuPDF (D-01 / D-03).

    Yields (page_num, text, title) - D-03 one-doc-per-page model.

    D-06: pages with < 10 chars after strip are skipped silently.
    D-05: caller must check total chars across all yielded pages; if < 50,
          file gets status='no_text_layer'.

    D-02: RTL helpers are NOT invoked in v1 (dead code).

    Phase 96 D-F4: each page is first extracted via get_text("blocks") (the
    Phase 95 default which preserves paragraph structure). If the result
    trips _detect_single_word_per_line (>= 70% single-word lines across >= 5
    non-empty lines — the pathological case where the PDF was laid out via
    per-word Tj operators at distinct positions), we re-extract that page via
    get_text("text", sort=True). The sort=True flag is load-bearing per
    PyMuPDF docs — it requests spatial sort (top-left to bottom-right) which
    is the documented remedy for non-sequential content-stream order.
    The fallback is per-page, not per-document: a document with mostly-good
    pages and one pathological page recovers only the bad page.
    """
    doc = fitz.open(filepath)
    try:
        title = (doc.metadata or {}).get("title") or os.path.basename(filepath)
        for page_num, page in enumerate(doc, start=1):
            # Primary: blocks mode preserves paragraph structure
            blocks = page.get_text("blocks")
            text_parts = [b[4].strip() for b in blocks if b[6] == 0 and b[4].strip()]
            text = "\n\n".join(text_parts)

            # Phase 96 D-F4: detect pathological one-word-per-line output
            # and fall back to get_text("text", sort=True) for THIS PAGE only.
            if _detect_single_word_per_line(text):
                try:
                    fallback_text = page.get_text("text", sort=True)
                    if fallback_text and fallback_text.strip():
                        text = fallback_text
                except Exception:
                    # If the fallback itself errors, keep the blocks output —
                    # one-word-per-line is still better than no text at all.
                    pass

            if len(text.strip()) < _EMPTY_PAGE_CHAR_THRESHOLD:
                continue  # D-06: skip empty pages (post-fallback)
            yield page_num, text, title
    finally:
        doc.close()


def extract_docx_pages(
    filepath: str,
) -> Iterator[tuple[int, str, str]]:
    """Extract text in 20-paragraph chunks from a DOCX file (D-04).

    Yields (chunk_num, text, title).
    D-04: fixed 20-paragraph windows (NOT Seewald's contains_page_break heuristic).
    """
    doc = _DocxDoc(filepath)
    basename = os.path.basename(filepath)
    try:
        title = doc.core_properties.title or basename
    except Exception:
        title = basename

    paragraphs = [p.text for p in doc.paragraphs]
    chunk_num = 0
    for start in range(0, max(len(paragraphs), 1), _DOCX_CHUNK_PARAGRAPHS):
        chunk_num += 1
        chunk_paras = paragraphs[start : start + _DOCX_CHUNK_PARAGRAPHS]
        text = "\n".join(p for p in chunk_paras if p.strip())
        if text.strip():
            yield chunk_num, text, title


def extract_txt(filepath: str) -> Iterator[tuple[int, str, str]]:
    """TXT extraction (D-07 + MEDIUM-2 review fix: strict utf-8-sig + cp1255 fallback;
    encoding_error on both failures - no silent replacement-character indexing).

    Encoding policy:
      1. Try utf-8-sig with errors='strict'.
      2. On UnicodeDecodeError, try cp1255 with errors='strict' (legacy Windows Hebrew).
      3. On second UnicodeDecodeError, raise EncodingError - caller (LocalIndexer._index_one_file)
         catches and sets local_files.extraction_status = 'encoding_error'. No docs emitted.

    DO NOT use errors='replace'. The replacement character (U+FFFD) silently corrupts
    the index and breaks search for the affected files (MEDIUM-2 review fix).
    """
    try:
        with open(filepath, "r", encoding="utf-8-sig", errors="strict") as f:
            text = f.read()
    except UnicodeDecodeError as utf8_err:
        try:
            with open(filepath, "r", encoding="cp1255", errors="strict") as f:
                text = f.read()
            logger.info(
                "extract_txt fallback to cp1255 for %s (utf-8-sig failed: %s)",
                filepath,
                utf8_err,
            )
        except UnicodeDecodeError as cp1255_err:
            raise EncodingError(
                f"Cannot decode {filepath} as utf-8-sig or cp1255: "
                f"utf-8 error: {utf8_err}; cp1255 error: {cp1255_err}"
            ) from cp1255_err
    yield (1, text, os.path.basename(filepath))


# ---------------------------------------------------------------------------
# Phase 97 F-01: HTML encoding detection helper
# ---------------------------------------------------------------------------

def _detect_html_encoding(raw_bytes: bytes) -> str:
    """F-01 encoding chain: <meta charset> -> utf-8 byte-sniff -> cp1255 fallback.

    Reads up to 1 KB of ASCII chars to find a <meta charset> declaration.
    If not found, tries strict UTF-8 decode on the first 4 KB.
    Falls back to cp1255 (legacy Windows Hebrew) if UTF-8 also fails.
    """
    head = raw_bytes[:1024].decode("ascii", errors="ignore").lower()
    m = re.search(r'<meta[^>]+charset\s*=\s*["\']?([\w-]+)', head)
    if m:
        return m.group(1).strip()
    try:
        raw_bytes[:4096].decode("utf-8", errors="strict")
        return "utf-8"
    except UnicodeDecodeError:
        pass
    return "cp1255"


# ---------------------------------------------------------------------------
# Phase 97 F-01: HTML extractor (lxml.html, NOT BeautifulSoup)
# ---------------------------------------------------------------------------

def extract_html_pages(filepath: str) -> Iterator[tuple[int, str, str, str, bool]]:
    """Phase 97 F-01. Yields (chunk_num, text, title, chunk_locator, is_rtl).
    F-06: NO _fix_rtl_* calls — lxml.html returns logical-order strings already.

    Chunking strategy:
      - Semantic: chunk at h1/h2 boundaries when len(headings) >= 3 AND
        avg paragraphs-per-heading >= 5.
      - Fallback: 20-paragraph windows when sparse.

    Encoding chain: <meta charset> -> utf-8 byte-sniff -> cp1255.
    is_rtl: True when <html dir="rtl"> or <body dir="rtl">.
    """
    import lxml.html

    with open(filepath, "rb") as f:
        raw = f.read()

    encoding = _detect_html_encoding(raw)
    try:
        text = raw.decode(encoding, errors="replace")
    except LookupError:
        text = raw.decode("cp1255", errors="replace")

    tree = lxml.html.fromstring(text)

    # Strip script and style content before chunking
    for tag in tree.iter("script", "style"):
        parent = tag.getparent()
        if parent is not None:
            parent.remove(tag)

    # Extract title from <title> element
    title_elem = tree.find(".//title")
    doc_title = (
        (title_elem.text_content() or os.path.basename(filepath)).strip()
        if title_elem is not None
        else os.path.basename(filepath)
    )

    # Detect RTL from html or body dir attribute
    html_dir = tree.get("dir", "") or ""
    body_elem = tree.find(".//body")
    body_dir = (body_elem.get("dir", "") or "") if body_elem is not None else ""
    is_rtl = html_dir.lower() == "rtl" or body_dir.lower() == "rtl"

    headings = list(tree.iter("h1", "h2"))
    paragraphs = list(tree.iter("p"))
    avg_inter = (len(paragraphs) / max(len(headings), 1)) if headings else 0
    use_semantic = len(headings) >= 3 and avg_inter >= 5

    if use_semantic:
        for chunk_num, h in enumerate(headings, start=1):
            heading_text = (h.text_content() or "").strip()
            buf = [heading_text] if heading_text else []
            sib = h.getnext()
            while sib is not None and sib.tag not in ("h1", "h2"):
                content = sib.text_content() or ""
                if content.strip():
                    buf.append(content)
                sib = sib.getnext()
            chunk_text = "\n".join(s.strip() for s in buf if s.strip())
            if not chunk_text:
                continue
            if len(chunk_text) > _MAX_CHARS_PER_CHUNK:
                chunk_text = chunk_text[:_MAX_CHARS_PER_CHUNK]
            locator_heading = heading_text[:40] if heading_text else f"section {chunk_num}"
            yield chunk_num, chunk_text, doc_title, f"§ {locator_heading}", is_rtl
    else:
        chunk_num = 0
        for start in range(0, max(len(paragraphs), 1), 20):
            chunk_paras = paragraphs[start: start + 20]
            texts = [
                (p.text_content() or "").strip()
                for p in chunk_paras
                if (p.text_content() or "").strip()
            ]
            chunk_text = "\n".join(texts)
            if not chunk_text:
                continue
            if len(chunk_text) > _MAX_CHARS_PER_CHUNK:
                chunk_text = chunk_text[:_MAX_CHARS_PER_CHUNK]
            chunk_num += 1
            end = min(start + 20, len(paragraphs))
            yield chunk_num, chunk_text, doc_title, f"¶ {start + 1}-{end}", is_rtl


# ---------------------------------------------------------------------------
# Phase 97 F-02: XLSX extractor (openpyxl streaming + zip-bomb pre-check)
# ---------------------------------------------------------------------------

def extract_xlsx_pages(filepath: str) -> Iterator[tuple[int, str, str, str, bool]]:
    """Phase 97 F-02. Yields (chunk_num, text, title, chunk_locator, is_rtl).
    F-06: NO _fix_rtl_* calls — openpyxl returns logical-order strings already.

    Pre-check: _check_zip_bomb() rejects bombs before openpyxl opens the file.
    Chunking: one Tantivy doc per (sheet, _XLSX_ROW_WINDOW) window.
    is_rtl: True when sheetView.rightToLeft is set on the sheet.
    F-04: uniform row text as 'cell1 | cell2 | cell3 | ...' (no header assumption).
    """
    from openpyxl import load_workbook

    bomb_reason = _check_zip_bomb(filepath)
    if bomb_reason:
        raise XlsxZipBombSuspected(bomb_reason)

    title = os.path.basename(filepath)

    # Pre-read RTL metadata from a non-streaming open (read_only=True worksheets
    # lack sheet_view in openpyxl; metadata-only open is fast — no cell data loaded).
    rtl_map: dict[str, bool] = {}
    try:
        wb_meta = load_workbook(filepath, read_only=False, data_only=True)
        try:
            for sn in wb_meta.sheetnames:
                ws_meta = wb_meta[sn]
                sv = getattr(ws_meta, "sheet_view", None)
                rtl_map[sn] = bool(getattr(sv, "rightToLeft", False)) if sv is not None else False
        finally:
            wb_meta.close()
    except Exception:
        pass  # If metadata read fails, is_rtl stays False for all sheets

    wb = load_workbook(filepath, read_only=True, data_only=True)
    try:
        chunk_num = 0
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            is_rtl = rtl_map.get(sheet_name, False)
            rows_in_window: list[str] = []
            cells_seen = 0
            window_start_row = 1
            for row_num, row in enumerate(ws.iter_rows(values_only=True), start=1):
                cell_strs = [str(c) if c is not None else "" for c in row]
                cells_seen += len(cell_strs)
                if cells_seen > _MAX_CELLS_PER_SHEET:
                    raise XlsxZipBombSuspected(
                        f"cells_seen {cells_seen} exceeds _MAX_CELLS_PER_SHEET "
                        f"{_MAX_CELLS_PER_SHEET} in sheet {sheet_name!r}"
                    )
                line = " | ".join(cell_strs)
                if line.strip():
                    rows_in_window.append(line)
                if (row_num - window_start_row + 1) >= _XLSX_ROW_WINDOW:
                    chunk_num += 1
                    chunk_text = "\n".join(rows_in_window)
                    if len(chunk_text) > _MAX_CHARS_PER_CHUNK:
                        chunk_text = chunk_text[:_MAX_CHARS_PER_CHUNK]
                    locator = f"{sheet_name}!R{window_start_row}:R{row_num}"
                    yield chunk_num, chunk_text, title, locator, is_rtl
                    rows_in_window = []
                    window_start_row = row_num + 1
            # Flush trailing partial window
            if rows_in_window:
                chunk_num += 1
                last_row = window_start_row + len(rows_in_window) - 1
                chunk_text = "\n".join(rows_in_window)
                if len(chunk_text) > _MAX_CHARS_PER_CHUNK:
                    chunk_text = chunk_text[:_MAX_CHARS_PER_CHUNK]
                locator = f"{sheet_name}!R{window_start_row}:R{last_row}"
                yield chunk_num, chunk_text, title, locator, is_rtl
    finally:
        wb.close()


# ---------------------------------------------------------------------------
# Phase 97 F-03 + F-05: CSV extractor (csv.Sniffer + encoding chain)
# ---------------------------------------------------------------------------

def extract_csv_pages(filepath: str) -> Iterator[tuple[int, str, str, str]]:
    """Phase 97 F-03 + F-05. Yields (chunk_num, text, title, chunk_locator).
    F-06: NO _fix_rtl_* calls — csv.reader returns logical-order strings already.

    Encoding chain: utf-8-sig -> cp1255 -> utf-16-le.
    Delimiter detection: csv.Sniffer over first 4 KB; fallback to csv.excel.
    Chunking: one Tantivy doc per _CSV_ROW_WINDOW (200) rows.
    F-04: uniform row text as 'cell1 | cell2 | cell3 | ...' (no header assumption).
    Total decode failure raises EncodingError.
    """
    import csv as _csv

    title = os.path.basename(filepath)
    chosen_encoding = None
    sample_text = None

    for enc in _CSV_ENCODINGS:
        try:
            with open(filepath, "r", encoding=enc, newline="") as f:
                sample_text = f.read(4096)
            chosen_encoding = enc
            break
        except UnicodeDecodeError:
            continue

    if chosen_encoding is None:
        raise EncodingError(
            f"CSV decode failed across all encodings {_CSV_ENCODINGS}: {filepath}"
        )

    try:
        dialect = _csv.Sniffer().sniff(sample_text, delimiters=",;\t")
    except _csv.Error:
        dialect = _csv.excel  # fallback to comma-delimited

    rows_in_window: list[str] = []
    window_start = 1
    chunk_num = 0
    row_num = 0  # track last row number for trailing-flush locator

    with open(filepath, "r", encoding=chosen_encoding, newline="") as f:
        reader = _csv.reader(f, dialect=dialect)
        for row_num, row in enumerate(reader, start=1):
            cell_strs = [str(c) if c is not None else "" for c in row]
            line = " | ".join(cell_strs)
            if line.strip():
                rows_in_window.append(line)
            if (row_num - window_start + 1) >= _CSV_ROW_WINDOW:
                chunk_num += 1
                chunk_text = "\n".join(rows_in_window)
                yield chunk_num, chunk_text, title, f"rows {window_start}-{row_num}"
                rows_in_window = []
                window_start = row_num + 1

    # Flush trailing partial window — use row_num (last row seen) for the locator
    if rows_in_window:
        chunk_num += 1
        chunk_text = "\n".join(rows_in_window)
        yield chunk_num, chunk_text, title, f"rows {window_start}-{row_num}"


# ---------------------------------------------------------------------------
# Phase 97 D-NEW-2 — folder reachability helper (module-level, called from scan_all)
# ---------------------------------------------------------------------------

def _check_folder_reachable(folder_path: str, max_retries: int = 3) -> tuple[bool, str]:
    """Phase 97 D-NEW-2: errno-discriminated reachability check with retry/backoff.

    Returns (reachable: bool, status_label: str).
    status_label is one of {'active', 'unreachable', 'timeout'}.

    ENOENT / EACCES -> non-retryable 'unreachable' (no sleep).
    ETIMEDOUT / EAGAIN -> retry up to max_retries times with 2.0s backoff,
                          then 'timeout' if all attempts fail.
    Any other OSError is re-raised unchanged.
    """
    for attempt in range(max_retries):
        try:
            if os.path.isdir(folder_path):
                return True, "active"
            # os.path.isdir returned False without raising — folder does not exist
            return False, "unreachable"
        except OSError as exc:
            if exc.errno in (errno.ETIMEDOUT, errno.EAGAIN):
                # Transient network timeout — retry with 2s backoff
                if attempt < max_retries - 1:
                    time.sleep(2.0)
                    continue
                return False, "timeout"
            if exc.errno in (errno.ENOENT, errno.EACCES):
                # Non-retryable (missing or permission denied)
                return False, "unreachable"
            raise
    return False, "timeout"


# ---------------------------------------------------------------------------
# LocalIndexer class
# ---------------------------------------------------------------------------

class LocalIndexer:
    """Qt-free LOCAL indexer engine (REQ-1, REQ-3, REQ-4, REQ-5).

    Owns:
      - Tantivy LOCAL side-index (build_local_schema)
      - SQLite metadata sidecar (folders / processed_files / local_pages / local_files)
      - Per-format extraction (PDF / DOCX / TXT)
      - Incremental re-extract on mtime change
      - Crash-safe delete via pending_delete two-phase protocol (HIGH-3)
      - Startup crash recovery (_recover_pending_deletes + pending inserts)
      - Folder overlap detection
      - Cooperative cancellation hooks

    Constructor:
        index_dir: Directory for LOCAL Tantivy side-index files.
        lab_index_dir: Directory for LOCAL LAB side-index files (may be unused by Plan 03).
        db_path: Path to SQLite sidecar (local_index.sqlite3 or any path).
        progress_cb: Optional callable(current_index, total, filename) for per-file progress.
        file_finished_cb: Optional callable(filename, status, pages, error_msg).
    """

    def __init__(
        self,
        index_dir: str,
        lab_index_dir: str,
        db_path: str,
        progress_cb: Optional[Callable] = None,
        file_finished_cb: Optional[Callable] = None,
    ) -> None:
        self._index_dir = index_dir
        self._lab_index_dir = lab_index_dir
        self._db_path = db_path
        self._progress_cb = progress_cb
        self._file_finished_cb = file_finished_cb

        # Thread-local SQLite connections (D-threading-fix).
        #
        # Python's sqlite3 module ties each connection to the thread that
        # created it (check_same_thread=True by default). LocalIndexer is
        # constructed on the main/GUI thread but scan_all() runs inside
        # LocalIndexerWorker(QThread) — a different thread. Using a single
        # shared connection therefore raises:
        #   sqlite3.ProgrammingError: SQLite objects created in a thread can
        #   only be used in that same thread.
        #
        # Fix: each thread gets its own connection opened lazily via the
        # _conn property (see below). Do NOT use check_same_thread=False —
        # that bypasses Python's safety guard and risks corruption when
        # multiple threads write concurrently. Per-thread connections are
        # safe here because the QMutex in MyLibraryTab (D-25) ensures only
        # one thread is writing at any given time.
        self._thread_local = threading.local()

        # Eagerly open a connection on the constructing thread so that
        # startup work (init_sqlite + _recover_pending_deletes) runs
        # immediately without requiring callers to prime the thread-local.
        _conn = init_sqlite(db_path)
        _conn.row_factory = sqlite3.Row
        self._thread_local._conn = _conn

        # Phase 97 LD-2: run idempotent migration ladder (0->1->2).
        # Called here, before any Tantivy work, so the schema additions
        # (cached_text, scan_run_id, chunk_locator, etc.) are present for
        # _write_page_doc in the same session.
        from shared.local_indexer_migrations import run as _run_migrations
        _run_migrations(self._thread_local._conn)  # Phase 97 D-NEW-1 ladder

        # Phase 97 LD-3 (Wave E placeholder): scan_run_id defaults to None
        # until _begin_scan_run sets it (Task 4).
        self._current_scan_run_id: str | None = None

        # Tantivy LOCAL side-index
        schema = build_local_schema()
        os.makedirs(index_dir, exist_ok=True)

        # R97.2-A: pre-initialize _writer = None so the post-rebuild gate
        # `if self._writer is None:` below can check it. rebuild_main_index_atomic
        # may populate self._writer via _reopen_internal_writer_index() — in
        # that case the gate skips the retry loop. Without this attribute
        # initialization, the gate raises AttributeError on first __init__.
        self._writer = None

        # Determine if this is a fresh dir (no meta.json = normal first-run),
        # a schema-mismatch (marker file present but hash changed), or a
        # genuine corruption (open raises on an existing index).
        _meta_exists = os.path.isfile(os.path.join(index_dir, "meta.json"))
        _schema_mismatch = False
        if _meta_exists:
            expected_marker = _compute_schema_marker(build_local_schema)
            actual_marker = _read_schema_marker(index_dir)
            if actual_marker != expected_marker:
                _schema_mismatch = True

        if not _meta_exists:
            # Fresh directory — create index from scratch (normal first-run)
            self._index = tantivy.Index(schema, path=index_dir)
            _write_schema_marker(index_dir, _compute_schema_marker(build_local_schema))
        else:
            # Existing index — open and check for corruption/schema mismatch
            _needs_rebuild = _schema_mismatch
            _open_exc_captured = None
            if not _needs_rebuild:
                try:
                    self._index = tantivy.Index.open(index_dir)
                except Exception as _open_exc:
                    logger.warning(
                        "LOCAL Tantivy index open failed: %r — attempting atomic rebuild",
                        _open_exc,
                    )
                    _needs_rebuild = True
                    _open_exc_captured = _open_exc

            if _needs_rebuild:
                if _schema_mismatch:
                    logger.warning(
                        "LOCAL schema marker mismatch — attempting atomic rebuild"
                    )
                import uuid as _uuid
                _recovery_run_id = _uuid.uuid4().hex
                try:
                    # Initialize _index to a fresh empty one so rebuild_main_index_atomic
                    # can check committed rows via _conn (which is ready at this point)
                    self._index = tantivy.Index(schema, path=index_dir)
                    self.rebuild_main_index_atomic(
                        _recovery_run_id,
                        close_searcher_cb=lambda: None,
                        reload_searcher_cb=lambda: None,
                    )
                    # R97.2-A: NO redundant tantivy.Index(...) reopen here.
                    # rebuild_main_index_atomic already calls
                    # _reopen_internal_writer_index() at :2799 which sets BOTH
                    # self._index and self._writer. A second reopen would leak
                    # the writer lock (LockBusy in the 2026-05-26 cascade).
                except Exception as _rebuild_exc:
                    logger.error("LOCAL atomic rebuild failed: %r", _rebuild_exc)
                    raise RuntimeError(
                        "LOCAL index corrupt and rebuild from cached_text failed. "
                        "Use 'Reset My Library' in Advanced settings to delete the cache and rescan."
                    ) from _rebuild_exc
        # Bug-1 fix: retry writer acquisition with backoff.
        # On Windows, after a language-switch restart the previous process may
        # still hold the Tantivy writer lock for a brief moment (the old process
        # is still exiting when the new process calls __init__).  Three retries
        # with exponential back-off (250 ms / 1 s / 2 s) cover the typical
        # process-exit window without blocking the UI for more than ~3 s.
        #
        # R97.2-A: gate the retry loop on `self._writer is None` so it skips
        # when rebuild_main_index_atomic already populated self._writer via
        # _reopen_internal_writer_index(). Without the gate this loop would
        # try to acquire a second writer on top of the rebuild's writer ->
        # LockBusy (the literal LockBusy in the 2026-05-26 cascade).
        if self._writer is None:
            _writer_retries = 3
            _writer_delays = [0.25, 1.0, 2.0]
            for _attempt in range(_writer_retries + 1):
                try:
                    self._writer = self._index.writer(heap_size=15_000_000)
                    break
                except Exception as _exc:  # noqa: BLE001
                    if _attempt < _writer_retries:
                        logger.warning(
                            "LocalIndexer: writer lock attempt %d/%d failed (%s); "
                            "retrying in %.1fs",
                            _attempt + 1, _writer_retries, _exc,
                            _writer_delays[_attempt],
                        )
                        time.sleep(_writer_delays[_attempt])
                    else:
                        raise

        # Batch tracking for two-phase commit (D-21)
        self._pending_filepaths: list[str] = []

        # Phase 97 C-02 — byte/count/time commit trigger (supersedes fixed 25-file batch)
        self._commit_triggers = _CommitTriggers()

        # HIGH-3 review fix: run pending-delete recovery at init
        recovered = self._recover_pending_deletes()
        if recovered:
            logger.info("LocalIndexer init: recovered %d pending deletes", recovered)

    # ------------------------------------------------------------------
    # Thread-local connection property (D-threading-fix)
    # ------------------------------------------------------------------

    @property
    def _conn(self) -> sqlite3.Connection:
        """Return the SQLite connection for the current thread.

        Creates a new connection (via init_sqlite) the first time any thread
        accesses this property. The constructing thread pre-populates its slot
        in __init__ so startup work (init_sqlite tables + _recover_pending_deletes)
        runs immediately.

        Each thread owns its own sqlite3.Connection. This is safe for concurrent
        reads from the main/UI thread while the QMutex-serialised worker thread
        performs writes, because SQLite WAL mode allows one writer + many readers
        without blocking.
        """
        conn = getattr(self._thread_local, "_conn", None)
        if conn is None:
            conn = init_sqlite(self._db_path)
            conn.row_factory = sqlite3.Row
            self._thread_local._conn = conn
        return conn

    # ------------------------------------------------------------------
    # Folder management
    # ------------------------------------------------------------------

    def add_folder(self, path: str) -> bool:
        """Normalize path, check for overlap with existing folders, and INSERT.

        Returns True if added, False if already present or overlapping.
        Raises ValueError with an explanatory message on overlap.
        """
        canonical = _canonical_filepath(path)
        existing = self._get_folder_paths()
        conflict = check_folder_overlap(canonical, existing)
        if conflict is not None:
            raise ValueError(
                f"Folder '{path}' overlaps with existing registered folder '{conflict}'"
            )
        now = time.time()
        try:
            self._conn.execute(
                "INSERT INTO folders (path, added_at, status) VALUES (?, ?, 'active')",
                (canonical, now),
            )
            self._conn.commit()
            return True
        except sqlite3.IntegrityError:
            # Already exists (UNIQUE constraint)
            return False

    def remove_folder(self, folder_path: str) -> int:
        """Synchronous delete: remove all files in folder, then delete the folder row.

        Returns the number of files removed.
        """
        canonical = _canonical_filepath(folder_path)
        row = self._conn.execute(
            "SELECT folder_id FROM folders WHERE path = ?", (canonical,)
        ).fetchone()
        if row is None:
            return 0
        folder_id = row["folder_id"]
        # Get all sys_ids + filepaths for this folder
        files = self._conn.execute(
            "SELECT sys_id, filepath FROM local_files WHERE folder_id = ?",
            (folder_id,),
        ).fetchall()
        count = 0
        for f in files:
            self._delete_file(f["sys_id"], f["filepath"])
            count += 1
        # Commit Tantivy
        try:
            self._writer.commit()
        except Exception as exc:
            logger.warning("remove_folder: Tantivy commit failed: %s", exc)
        # Delete folder row
        self._conn.execute("DELETE FROM folders WHERE folder_id = ?", (folder_id,))
        self._conn.commit()
        return count

    def get_filepath(self, sys_id: str) -> Optional[str]:
        """Return the canonical filepath for a LOCAL sys_id, or None if not found.

        Used by the Browse panel 'Open file' button (D-28) to obtain the source
        file path so os.startfile() can launch the OS default application.
        """
        row = self._conn.execute(
            "SELECT filepath FROM local_files WHERE sys_id = ?", (sys_id,)
        ).fetchone()
        return row["filepath"] if row else None

    def list_all_filepaths(self) -> list:
        """Phase 96 D-F1: enumerate every on-disk filepath in the index.

        Used by MyLibraryTab._on_worker_finished to prune stale opt-outs after
        a rescan completes. Returns canonical filepaths from the local_files
        SQLite table (preferred over direct _conn access from the tab).

        Per 96-08-WIRING-NOTES.md §LocalIndexer public API: Option (A) —
        recommended public method over direct _conn access (Option B).
        """
        cur = self._conn.execute("SELECT filepath FROM local_files")
        return [row[0] for row in cur.fetchall()]

    def get_file_status_for_folder(self, folder_path: str) -> dict:
        """Phase 96 fix-2: return prior scan status for all files in folder_path.

        Called by _UnifiedFileTreeWidget.populate_for_folder so that Pages and
        Status columns are populated immediately on tab open (not only after a
        new scan). Without this, the tree opens with all status columns empty,
        which made the previous QSplitter+status_table look the same.

        Returns a dict mapping canonical filepath -> {'pages': int, 'status': str}
        for every file that has been indexed for this folder. Files not yet
        indexed (pending) are absent from the dict.
        """
        from shared.local_sys_id import _canonical_filepath
        canonical = _canonical_filepath(folder_path)
        row = self._conn.execute(
            "SELECT folder_id FROM folders WHERE path = ?", (canonical,)
        ).fetchone()
        if row is None:
            return {}
        folder_id = row["folder_id"]
        rows = self._conn.execute(
            "SELECT filepath, page_count, extraction_status "
            "FROM local_files WHERE folder_id = ? AND pending_delete = 0",
            (folder_id,),
        ).fetchall()
        return {
            r["filepath"]: {"pages": r["page_count"], "status": r["extraction_status"]}
            for r in rows
        }

    def list_folders(self) -> list[dict]:
        """Return all registered folders ordered by added_at (D-15 / D-16).

        Each dict has keys: folder_id, path, added_at, last_scanned_at, status.
        Used by MyLibraryTab to populate the folder list widget and by
        prescan_count_all for the W8 aggregate ceiling check.
        """
        rows = self._conn.execute(
            "SELECT folder_id, path, added_at, last_scanned_at, status "
            "FROM folders ORDER BY added_at"
        ).fetchall()
        return [
            {
                "folder_id": row["folder_id"],
                "path": row["path"],
                "added_at": row["added_at"],
                "last_scanned_at": row["last_scanned_at"],
                "status": row["status"],
            }
            for row in rows
        ]

    def prescan_count_all(self) -> tuple[int, int]:
        """W8: aggregate prescan across all registered, available folders.

        Iterates every registered folder; folders with status in
        ('unavailable', 'unreachable', 'timeout') OR whose path is not a
        directory (D-40) are excluded from the sum.
        Used by MyLibraryTab Refresh per D-16 multi-folder support.
        The aggregate is the input to the D-26/D-41 ceiling-check dialog —
        thresholds apply to the AGGREGATE, not per-folder.

        Returns: (total_file_count, total_bytes)
        """
        total_files = 0
        total_bytes = 0
        for folder in self.list_folders():
            path = folder["path"]
            # Phase 97 D-NEW-2 / LD-9: skip-set includes Phase 95 'unavailable' + new 'unreachable'/'timeout'.
            if folder["status"] in ("unavailable", "unreachable", "timeout"):
                continue
            if not os.path.isdir(path):
                continue  # treat as unavailable
            f, b = self.prescan_count(path)
            total_files += f
            total_bytes += b
        return total_files, total_bytes

    def _get_folder_paths(self) -> list[str]:
        rows = self._conn.execute("SELECT path FROM folders").fetchall()
        return [r["path"] for r in rows]

    # ------------------------------------------------------------------
    # Scan
    # ------------------------------------------------------------------

    def scan_all(self, cancel_check: Callable[[], bool] = lambda: False) -> dict:
        """Scan all registered folders; return summary dict.

        Returns: {"indexed": N, "skipped": N, "errors": N, "cancelled": bool}

        Phase 97 LD-6: inserts a scan_runs row with status='running' at start;
        updates to 'completed' at clean exit or 'canceled' on cancellation.
        Exceptions (crashes) leave the row at 'running' so start_recovery_probe()
        detects it on next startup.
        """
        # Phase 97 LD-6: begin scan_run lifecycle tracking.
        run_id = self._begin_scan_run()
        try:
            result = self._scan_all_impl(cancel_check)
        except Exception:
            # Leave status='running' so recovery probe fires on next startup.
            raise
        if result.get("cancelled"):
            self._end_scan_run(run_id, "canceled")
        else:
            self._end_scan_run(run_id, "completed")
        return result

    def _scan_all_impl(self, cancel_check: Callable[[], bool] = lambda: False) -> dict:
        """Internal implementation of scan_all (split out for scan_run lifecycle wrapper)."""
        result = {"indexed": 0, "skipped": 0, "errors": 0, "cancelled": False}
        # Phase 97 D-NEW-3: per-scan retry counter + re-queue list (allocated fresh per scan_run).
        # _scan_run_retries[filepath] = number of times this file was marked 'changed_during_index'
        # during this scan_run. Max 3 retries; 4th attempt gives up.
        self._scan_run_retries: dict[str, int] = {}
        self._re_queue: list[str] = []
        folders = self._conn.execute(
            "SELECT folder_id, path, status FROM folders"
        ).fetchall()

        for folder in folders:
            if cancel_check():
                result["cancelled"] = True
                break

            folder_path = folder["path"]
            folder_id = folder["folder_id"]

            # Phase 97 D-NEW-2 — errno-discriminated reachability check + retry.
            # D-40 extended: ENOENT/EACCES -> 'unreachable' (non-retryable);
            # ETIMEDOUT/EAGAIN -> 'timeout' (3-retry × 2s backoff).
            # Phase 97 D-NEW-2 / LD-9: skip-set includes Phase 95 'unavailable' + new 'unreachable'/'timeout'.
            # Pre-check: if already marked 'unreachable'/'timeout', skip _check_folder_reachable
            # on auto-rescan (user must manually Refresh to retry unreachable folders).
            if folder["status"] in ("unavailable", "unreachable", "timeout"):
                # Already known unreachable — skip without re-probing on auto-rescan.
                # User can trigger a manual Refresh to retry.
                continue
            reachable, status_label = _check_folder_reachable(folder_path)
            if not reachable:
                self._conn.execute(
                    "UPDATE folders SET status = ? WHERE folder_id = ?",
                    (status_label, folder_id),
                )
                self._conn.commit()
                logger.info(
                    "Phase 97 D-NEW-2: folder '%s' status=%s — preserving existing rows",
                    folder_path, status_label,
                )
                continue

            # Mark folder as active + update last_scanned_at
            self._conn.execute(
                "UPDATE folders SET status = 'active', last_scanned_at = ? WHERE folder_id = ?",
                (time.time(), folder_id),
            )
            self._conn.commit()

            # Build set of current files on disk
            disk_files: dict[str, int] = {}  # canonical_path -> file_size
            for filepath, file_size in self._iterate_supported_files(
                folder_path, cancel_check
            ):
                disk_files[filepath] = file_size

            if cancel_check():
                result["cancelled"] = True
                break

            # Get cached files for this folder (D-NEW-8: include mtime_ns for exact comparison)
            cached = {
                row["filepath"]: row
                for row in self._conn.execute(
                    "SELECT pf.filepath, pf.mtime, pf.mtime_ns, pf.size, pf.sys_id, pf.status "
                    "FROM processed_files pf "
                    "JOIN local_files lf ON pf.sys_id = lf.sys_id "
                    "WHERE lf.folder_id = ?",
                    (folder_id,),
                ).fetchall()
            }

            # Detect deleted files (in cache but not on disk)
            for cached_path, cached_row in list(cached.items()):
                if cached_path not in disk_files:
                    if cancel_check():
                        result["cancelled"] = True
                        break
                    sys_id = cached_row["sys_id"]
                    if sys_id:
                        self._delete_file(sys_id, cached_path)

            if result["cancelled"]:
                break

            # Index new / modified files
            total_files = len(disk_files)
            for idx, (filepath, file_size) in enumerate(disk_files.items()):
                if cancel_check():
                    result["cancelled"] = True
                    break

                cached_row = cached.get(filepath)
                try:
                    stat = os.stat(filepath)
                    mtime = stat.st_mtime
                    mtime_ns = stat.st_mtime_ns   # D-NEW-8: nanosecond precision
                    fsize = stat.st_size
                except OSError as exc:
                    logger.warning("scan_all: stat failed for %s: %s", filepath, exc)
                    result["errors"] += 1
                    continue

                # D-NEW-8: cache-hit comparison uses integer mtime_ns (exact) instead of
                # float mtime with 0.01s tolerance. Legacy Phase 95 rows have mtime_ns=NULL;
                # (cached_row["mtime_ns"] or 0) evaluates to 0 for NULL, which is never
                # equal to any real mtime_ns (which is always > 0 for real files since epoch).
                # This correctly forces a cache miss for all Phase 95 legacy rows on first
                # Phase 97 scan, allowing mtime_ns to be populated.
                if (
                    cached_row is not None
                    and cached_row["status"] == "committed"
                    and (cached_row["mtime_ns"] or 0) == mtime_ns
                    and (cached_row["size"] or 0) == fsize
                ):
                    # Phase 97 U-02 (RESEARCH Issue #4 LOCK): scan_run_id NOT set on cache-hit skip.
                    # The cached_row keeps its prior scan_run_id (possibly NULL for legacy rows or
                    # some prior committed UUID). Setting it here would cause discard_run(current_run)
                    # to wrongly delete previously committed rows that this run merely visited.
                    # Pinned by tests/test_scan_run_id.py::test_no_run_id_on_skipped.
                    result["skipped"] += 1
                    continue

                # Fire progress callback only for files actually being indexed/re-indexed
                # (not for cache hits) — matches D-23 "per file being processed" semantics
                if self._progress_cb:
                    self._progress_cb(idx, total_files, os.path.basename(filepath))

                # Phase 97 C-05 raw-size hard skip (LD-9 — write to BOTH tables)
                if fsize > _MAX_FILE_SIZE:
                    now = time.time()
                    sys_id_for_status = generate_local_sys_id(_canonical_filepath(filepath))
                    self._conn.execute(
                        "INSERT OR REPLACE INTO processed_files "
                        "(filepath, mtime, mtime_ns, size, sys_id, status, scan_run_id) "
                        "VALUES (?, ?, ?, ?, ?, 'oversized', ?)",
                        (filepath, mtime, mtime_ns, fsize, sys_id_for_status,
                         self._current_scan_run_id),
                    )
                    # LD-9: also write local_files.extraction_status so the UI tree shows it.
                    self._upsert_local_files_status(
                        sys_id_for_status, filepath, folder_id, "oversized", fsize, now
                    )
                    self._conn.commit()
                    if self._file_finished_cb:
                        try:
                            cb_path = _canonical_filepath(filepath)
                        except Exception:
                            cb_path = filepath
                        self._file_finished_cb(cb_path, "oversized", 0, "")
                    result["oversized"] = result.get("oversized", 0) + 1
                    # Phase 97 C-02 — record for commit trigger even on skip
                    self._commit_triggers.record_file(fsize)
                    if self._commit_triggers.should_commit():
                        self._commit_batch()
                        self._commit_triggers.reset()
                    continue

                # Phase 97 C-05 zip-bomb defense — only for zip-container formats
                ext_lower = os.path.splitext(filepath)[1].lower()
                if ext_lower in (".docx", ".xlsx"):
                    bomb_reason = _check_zip_bomb(filepath)
                    if bomb_reason:
                        now = time.time()
                        sys_id_for_bomb = generate_local_sys_id(_canonical_filepath(filepath))
                        self._conn.execute(
                            "INSERT OR REPLACE INTO processed_files "
                            "(filepath, mtime, mtime_ns, size, sys_id, status, scan_run_id) "
                            "VALUES (?, ?, ?, ?, ?, 'zip_bomb_suspected', ?)",
                            (filepath, mtime, mtime_ns, fsize, sys_id_for_bomb,
                             self._current_scan_run_id),
                        )
                        # LD-9: also write local_files.extraction_status
                        self._upsert_local_files_status(
                            sys_id_for_bomb, filepath, folder_id,
                            "zip_bomb_suspected", fsize, now, error_msg=bomb_reason,
                        )
                        self._conn.commit()
                        logger.warning(
                            "Phase 97 C-05: skipping %s — %s", filepath, bomb_reason
                        )
                        if self._file_finished_cb:
                            try:
                                cb_path = _canonical_filepath(filepath)
                            except Exception:
                                cb_path = filepath
                            self._file_finished_cb(cb_path, "zip_bomb_suspected", 0, bomb_reason)
                        result["zip_bomb_suspected"] = result.get("zip_bomb_suspected", 0) + 1
                        self._commit_triggers.record_file(fsize)
                        if self._commit_triggers.should_commit():
                            self._commit_batch()
                            self._commit_triggers.reset()
                        continue

                # Need to index (new or modified)
                if cached_row is not None and cached_row["sys_id"]:
                    # Modified: delete old docs first (D-36)
                    self._delete_file(cached_row["sys_id"], filepath)

                # Phase 97 D-NEW-3: record pre-extraction stat for TOCTOU detection.
                # Take a fresh stat IMMEDIATELY before extraction to minimise the window.
                _pre_mtime_ns = mtime_ns
                _pre_size = fsize
                try:
                    _pre_stat = os.stat(filepath)
                    _pre_mtime_ns = _pre_stat.st_mtime_ns
                    _pre_size = _pre_stat.st_size
                except OSError:
                    pass  # use values from earlier stat if re-stat fails

                status, pages = self._index_one_file(filepath, folder_id, cancel_check)

                # Phase 97 D-NEW-3: post-extraction stat — detect file changed during indexing.
                _changed_during_index = False
                try:
                    post_stat = os.stat(filepath)
                    if (post_stat.st_mtime_ns != _pre_mtime_ns or post_stat.st_size != _pre_size):
                        # File changed during indexing — re-queue (max 3 retries per scan_run)
                        retries = self._scan_run_retries.get(filepath, 0)
                        if retries < 3:
                            self._scan_run_retries[filepath] = retries + 1
                            self._re_queue.append(filepath)
                            self._conn.execute(
                                "INSERT OR REPLACE INTO processed_files "
                                "(filepath, mtime, mtime_ns, size, sys_id, status, scan_run_id) "
                                "VALUES (?, ?, ?, ?, "
                                "(SELECT sys_id FROM processed_files WHERE filepath = ?), "
                                "'changed_during_index', ?)",
                                (filepath, mtime, mtime_ns, fsize, filepath,
                                 self._current_scan_run_id),
                            )
                            self._conn.commit()
                            logger.info(
                                "Phase 97 D-NEW-3: %s changed during index (retry %d/3)",
                                filepath, retries + 1,
                            )
                            _changed_during_index = True
                        else:
                            logger.warning(
                                "Phase 97 D-NEW-3: %s changed during index 3× — giving up",
                                filepath,
                            )
                except OSError:
                    pass  # post-stat failure is non-fatal — proceed with indexed status

                if _changed_during_index:
                    # Remove from pending so _commit_batch won't overwrite 'changed_during_index'
                    # with 'committed'. The Tantivy doc for this file may be partial, but it
                    # will be re-indexed on the next scan_run (re-queued above).
                    try:
                        canonical_fp = _canonical_filepath(filepath)
                        if canonical_fp in self._pending_filepaths:
                            self._pending_filepaths.remove(canonical_fp)
                        if filepath in self._pending_filepaths:
                            self._pending_filepaths.remove(filepath)
                    except (ValueError, AttributeError):
                        pass
                    result["errors"] += 1
                    continue  # skip _commit_batch + _file_finished_cb for this file

                if status in ("ok", "no_text_layer", "encoding_error", "unsupported"):
                    result["indexed"] += 1
                else:
                    result["errors"] += 1

                if self._file_finished_cb:
                    # Phase 96 fix-7 (Codex P1.2): emit canonical filepath so
                    # update_file_status can do an exact dict lookup — basename
                    # alone was ambiguous when two folders contain same filename.
                    # [Rule 1 fix] Removed local `from shared.local_sys_id import
                    # _canonical_filepath` — it's a module-level import (line 71).
                    # The local import caused UnboundLocalError at the earlier
                    # usage on the oversized/zip-bomb branch (Python hoists local
                    # variable assignment to the function scope top).
                    try:
                        cb_path = _canonical_filepath(filepath)
                    except Exception:
                        cb_path = filepath
                    self._file_finished_cb(cb_path, status, pages, "")

                # Phase 97 C-02 — byte/count/time trigger (supersedes Phase 95 D-21 fixed
                # 25-file batch). _COMMIT_BATCH_SIZE is kept as dead-code constant per D-02.
                self._commit_triggers.record_file(fsize)
                if self._commit_triggers.should_commit():
                    self._commit_batch()
                    self._commit_triggers.reset()

        # Final commit
        if not result["cancelled"] and self._pending_filepaths:
            self._commit_batch()

        return result

    def prescan_count(
        self,
        folder_path: str,
        cancel_check: Optional[Callable[[], bool]] = None,
    ) -> tuple[int, int]:
        """Fast count of supported files + total bytes for D-26 ceiling dialog.

        Phase 97 C-03: accepts optional cancel_check callable so PrescanWorker
        can abort the os.walk mid-scan when the user presses Cancel.
        Returns (-1, -1) when cancelled.
        """
        file_count = 0
        total_bytes = 0
        try:
            for dirpath, _dirs, files in os.walk(folder_path, followlinks=False):
                if cancel_check is not None and cancel_check():
                    return -1, -1
                try:
                    for fname in files:
                        ext = os.path.splitext(fname)[1].lower()
                        if ext in _SUPPORTED_EXTENSIONS:
                            fpath = os.path.join(dirpath, fname)
                            try:
                                total_bytes += os.path.getsize(fpath)
                                file_count += 1
                            except OSError:
                                pass
                except OSError as exc:
                    logger.warning("prescan_count: OSError in %s: %s", dirpath, exc)
        except OSError as exc:
            logger.warning("prescan_count: OSError walking %s: %s", folder_path, exc)
        return file_count, total_bytes

    def estimate_index_size(self) -> int:
        """Phase 97 C-06 — sum sizes of all files under index dir.

        Used by MyLibraryTab._update_disk_indicator() to compute the current
        LOCAL Tantivy index footprint for the merge-headroom warning.
        Recursively sums Tantivy segment files + SQLite sidecar.
        """
        total = 0
        try:
            for root, _dirs, files in os.walk(self._index_dir):
                for name in files:
                    try:
                        total += os.path.getsize(os.path.join(root, name))
                    except OSError:
                        continue
        except OSError:
            pass
        return total

    # ------------------------------------------------------------------
    # Startup recovery
    # ------------------------------------------------------------------

    def startup_recovery(self) -> dict:
        """Two-pass startup recovery (D-21 + HIGH-3 review fix).

        Pass A (HIGH-3): recover any crash-interrupted deletes (_recover_pending_deletes).
        Pass B (D-21): re-extract files with status='pending'.

        Returns: {'pending_deletes_recovered': N, 'pending_inserts_recovered': M}
        """
        # Pass A: HIGH-3 - finish any interrupted deletes
        deletes_recovered = self._recover_pending_deletes()

        # Pass B: D-21 - re-extract pending inserts
        pending_rows = self._conn.execute(
            "SELECT filepath, sys_id FROM processed_files WHERE status = 'pending'"
        ).fetchall()
        inserts_recovered = 0
        for row in pending_rows:
            filepath = row["filepath"]
            sys_id = row["sys_id"]
            if not os.path.exists(filepath):
                # File gone - clean up
                if sys_id:
                    self._delete_file(sys_id, filepath)
                else:
                    self._conn.execute(
                        "DELETE FROM processed_files WHERE filepath = ?", (filepath,)
                    )
                    self._conn.commit()
                continue
            # Delete stale data and re-extract
            if sys_id:
                self._delete_file(sys_id, filepath)
            # Find folder_id for this file
            folder_row = self._conn.execute(
                "SELECT f.folder_id FROM folders f WHERE ? LIKE f.path || '%'",
                (filepath,),
            ).fetchone()
            folder_id = folder_row["folder_id"] if folder_row else 1
            self._index_one_file(filepath, folder_id, lambda: False)
            inserts_recovered += 1

        if self._pending_filepaths:
            self._commit_batch()

        return {
            "pending_deletes_recovered": deletes_recovered,
            "pending_inserts_recovered": inserts_recovered,
        }

    def _recover_pending_deletes(self) -> int:
        """HIGH-3 review fix: replay Tantivy delete + SQLite cleanup for pending rows.

        Called at __init__ AND from startup_recovery() Pass A.
        Idempotent: re-deleting already-deleted Tantivy UIDs is a no-op.

        Returns: count of completed deletes.
        """
        pending_rows = self._conn.execute(
            "SELECT sys_id FROM local_files WHERE pending_delete = 1"
        ).fetchall()
        n = 0
        for row in pending_rows:
            sys_id = row["sys_id"]
            # Step 2: Tantivy delete-by-uid loop + commit (idempotent)
            uid_rows = self._conn.execute(
                "SELECT uid FROM local_pages WHERE sys_id = ?", (sys_id,)
            ).fetchall()
            for uid_row in uid_rows:
                uid = uid_row["uid"]
                self._writer.delete_documents("unique_id", uid)
            try:
                self._writer.commit()
            except Exception as exc:
                logger.warning(
                    "_recover_pending_deletes: Tantivy commit failed for sys_id=%s: %s",
                    sys_id, exc,
                )
                continue
            # Step 3: SQLite final cleanup
            self._conn.execute("DELETE FROM local_pages WHERE sys_id = ?", (sys_id,))
            self._conn.execute("DELETE FROM local_files WHERE sys_id = ?", (sys_id,))
            self._conn.execute(
                "DELETE FROM processed_files WHERE sys_id = ?", (sys_id,)
            )
            self._conn.commit()
            n += 1
        logger.info("HIGH-3 recovery: completed %d pending deletes", n)
        return n

    # ------------------------------------------------------------------
    # Internal indexing helpers
    # ------------------------------------------------------------------

    def _iterate_supported_files(
        self,
        folder_path: str,
        cancel_check: Callable[[], bool],
    ) -> Iterator[tuple[str, int]]:
        """Yield (canonical_filepath, file_size) for ALL files in folder.

        Note: unsupported extensions are included so _index_one_file can record
        them with extraction_status='unsupported'. The extension check happens
        inside _index_one_file, not here.
        """
        try:
            for dirpath, _dirs, files in os.walk(folder_path, followlinks=False):
                if cancel_check():
                    return
                try:
                    for fname in files:
                        # Phase 97.1: per-file cancel check (was per-directory only).
                        # For directories with hundreds of files, the prior
                        # outer-only check meant cancel latency was one full
                        # directory of file processing. See debug session
                        # `.planning/debug/phase-97-freeze-winerror-3.md`.
                        if cancel_check():
                            return
                        fpath = os.path.join(dirpath, fname)
                        canonical = _canonical_filepath(fpath)
                        try:
                            fsize = os.path.getsize(canonical)
                        except OSError:
                            fsize = 0
                        yield canonical, fsize
                except OSError as exc:
                    logger.warning(
                        "_iterate_supported_files: OSError in %s: %s", dirpath, exc
                    )
        except OSError as exc:
            logger.warning(
                "_iterate_supported_files: OSError walking %s: %s", folder_path, exc
            )

    def _index_one_file(
        self,
        filepath: str,
        folder_id: int,
        cancel_check: Callable[[], bool],
    ) -> tuple[str, int]:
        """Index a single file: extract pages, write Tantivy docs, mark pending.

        Returns (extraction_status, page_count).
        """
        canonical = _canonical_filepath(filepath)
        ext = os.path.splitext(filepath)[1].lower()

        # Generate sys_id with collision retry
        sys_id = None
        for slot in range(_MAX_COLLISION_RETRIES):
            candidate = generate_local_sys_id(canonical, slot=slot)
            # Check for collision (another file already owns this sys_id)
            existing = self._conn.execute(
                "SELECT filepath FROM local_files WHERE sys_id = ?", (candidate,)
            ).fetchone()
            if existing is None or existing["filepath"] == canonical:
                sys_id = candidate
                break
            logger.warning(
                "_index_one_file: sys_id collision at slot %d for %s, retrying",
                slot, canonical,
            )
        if sys_id is None:
            logger.error(
                "_index_one_file: exhausted %d collision slots for %s",
                _MAX_COLLISION_RETRIES, canonical,
            )
            return ("error", 0)

        try:
            _stat = os.stat(canonical)
            file_size = _stat.st_size
            mtime = _stat.st_mtime
            mtime_ns = _stat.st_mtime_ns   # D-NEW-8: nanosecond precision
        except OSError as exc:
            logger.warning("_index_one_file: cannot stat %s: %s", canonical, exc)
            return ("error", 0)

        basename = os.path.basename(canonical)

        # Mark as pending BEFORE extraction (two-phase commit step 1)
        # D-NEW-8: write mtime_ns alongside mtime so cache-hit check can use int comparison.
        self._conn.execute(
            """
            INSERT OR REPLACE INTO processed_files
                (filepath, mtime, mtime_ns, size, sys_id, status, scan_run_id)
            VALUES (?, ?, ?, ?, ?, 'pending', ?)
            """,
            (canonical, mtime, mtime_ns, file_size, sys_id, self._current_scan_run_id),
        )
        self._conn.commit()

        # WR-01 FIX: pre-insert local_files row with status='pending' BEFORE
        # extraction so _write_page_doc can read the real file_id and bake
        # it into the per-page full_header (D-34 unique F-suffix per file).
        # The row gets UPDATEd to its final state by _finish_file once
        # extraction completes. Use INSERT OR IGNORE so re-index runs don't
        # clobber an existing row (the modified-file path already called
        # _delete_file at scan_all line ~725 to clear stale rows).
        now = time.time()
        self._conn.execute(
            """
            INSERT OR IGNORE INTO local_files (
                sys_id, filepath, folder_id, display_title,
                original_filename, file_extension, page_count,
                file_size_bytes, extraction_status, last_indexed_at,
                sha256_full, error_msg, pending_delete
            ) VALUES (?, ?, ?, ?, ?, ?, 0, ?, 'pending', ?, NULL, NULL, 0)
            """,
            (
                sys_id, canonical, folder_id, basename,
                basename, ext, file_size, now,
            ),
        )
        self._conn.commit()

        # Extract pages
        if ext not in _SUPPORTED_EXTENSIONS:
            return self._finish_file(
                sys_id, canonical, folder_id, basename, ext,
                0, file_size, "unsupported", None, mtime,
            )

        pages_written = 0
        extraction_status = "ok"
        error_msg = None
        display_title = basename

        try:
            if ext == ".pdf":
                pages_written, extraction_status, display_title = self._extract_and_write_pdf(
                    sys_id, canonical, folder_id, cancel_check
                )
            elif ext == ".docx":
                pages_written, extraction_status, display_title = self._extract_and_write_docx(
                    sys_id, canonical, folder_id, cancel_check
                )
            elif ext == ".txt":
                pages_written, extraction_status, display_title = self._extract_and_write_txt(
                    sys_id, canonical, folder_id
                )
            elif ext == ".html":
                pages_written, extraction_status, display_title = self._extract_and_write_html(
                    sys_id, canonical, folder_id, cancel_check
                )
            elif ext == ".xlsx":
                pages_written, extraction_status, display_title = self._extract_and_write_xlsx(
                    sys_id, canonical, folder_id, cancel_check
                )
            elif ext == ".csv":
                pages_written, extraction_status, display_title = self._extract_and_write_csv(
                    sys_id, canonical, folder_id, cancel_check
                )
        except XlsxZipBombSuspected as exc:
            extraction_status = "zip_bomb_suspected"
            error_msg = str(exc)
            logger.warning("_index_one_file: zip_bomb_suspected for %s: %s", canonical, exc)
        except EncodingError as exc:
            extraction_status = "encoding_error"
            error_msg = str(exc)
            logger.info("_index_one_file: encoding_error for %s: %s", canonical, exc)
        except Exception as exc:
            extraction_status = "error"
            error_msg = str(exc)
            logger.warning("_index_one_file: extraction error for %s: %s", canonical, exc)

        return self._finish_file(
            sys_id, canonical, folder_id, basename, ext,
            pages_written, file_size, extraction_status, error_msg, mtime,
            display_title=display_title,
        )

    def _finish_file(
        self,
        sys_id: str,
        filepath: str,
        folder_id: int,
        basename: str,
        ext: str,
        page_count: int,
        file_size: int,
        extraction_status: str,
        error_msg: Optional[str],
        mtime: float,
        display_title: Optional[str] = None,
    ) -> tuple[str, int]:
        """Upsert local_files row. Returns (extraction_status, page_count).

        WR-01 FIX: _index_one_file now pre-INSERTs a placeholder row before
        extraction so _write_page_doc can read the real file_id.  Therefore
        _finish_file UPDATEs the existing row instead of INSERT OR REPLACE
        (which would DELETE+INSERT and reassign file_id, breaking the
        already-written full_header).  If for some reason the pre-INSERT
        was missed, fall through to INSERT.
        """
        now = time.time()
        cursor = self._conn.execute(
            """
            UPDATE local_files
               SET filepath = ?, folder_id = ?, display_title = ?,
                   original_filename = ?, file_extension = ?, page_count = ?,
                   file_size_bytes = ?, extraction_status = ?, last_indexed_at = ?,
                   sha256_full = NULL, error_msg = ?, pending_delete = 0
             WHERE sys_id = ?
            """,
            (
                filepath, folder_id, display_title or basename,
                basename, ext, page_count,
                file_size, extraction_status, now,
                error_msg,
                sys_id,
            ),
        )
        if cursor.rowcount == 0:
            # Fallback — no pre-existing row to update (shouldn't normally
            # happen since _index_one_file always pre-INSERTs).
            self._conn.execute(
                """
                INSERT OR REPLACE INTO local_files (
                    sys_id, filepath, folder_id, display_title,
                    original_filename, file_extension, page_count,
                    file_size_bytes, extraction_status, last_indexed_at,
                    sha256_full, error_msg, pending_delete
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, NULL, ?, 0)
                """,
                (
                    sys_id, filepath, folder_id, display_title or basename,
                    basename, ext, page_count,
                    file_size, extraction_status, now,
                    error_msg,
                ),
            )
        self._conn.commit()
        # Track for batch commit
        self._pending_filepaths.append(filepath)
        return (extraction_status, page_count)

    def _upsert_local_files_status(
        self,
        sys_id: str,
        filepath: str,
        folder_id: int,
        extraction_status: str,
        file_size: int,
        now: float,
        error_msg: Optional[str] = None,
    ) -> None:
        """Phase 97 LD-9 — write extraction_status to local_files (UI source of truth).

        For statuses set BEFORE _index_one_file's pre-INSERT runs (oversized,
        zip_bomb_suspected), we UPSERT via INSERT ... ON CONFLICT(sys_id) DO UPDATE.
        This ensures the UI tree (which reads local_files.extraction_status) shows
        the correct status even when _finish_file is never reached.

        LD-9 rationale: Wave D folder counter aggregation reads local_files.extraction_status.
        Writing only processed_files.status leaves the UI blind to these states.
        """
        basename = os.path.basename(filepath)
        ext = os.path.splitext(filepath)[1].lower()
        self._conn.execute(
            """
            INSERT INTO local_files (
                sys_id, filepath, folder_id, display_title,
                original_filename, file_extension, page_count,
                file_size_bytes, extraction_status, last_indexed_at,
                sha256_full, error_msg, pending_delete
            ) VALUES (?, ?, ?, ?, ?, ?, 0, ?, ?, ?, NULL, ?, 0)
            ON CONFLICT(sys_id) DO UPDATE SET
                extraction_status = excluded.extraction_status,
                error_msg = excluded.error_msg,
                file_size_bytes = excluded.file_size_bytes,
                last_indexed_at = excluded.last_indexed_at
            """,
            (sys_id, filepath, folder_id, basename, basename, ext,
             file_size, extraction_status, now, error_msg),
        )

    def _write_page_doc(
        self,
        sys_id: str,
        page_num: int,
        text: str,
        title: str,
        folder_id: int,
        chunk_locator: str = "",
    ) -> str:
        """Write one Tantivy doc for a page and record in local_pages. Returns uid.

        WR-01 FIX: _index_one_file now pre-INSERTs the local_files row with
        status='pending' BEFORE extraction, so file_id is populated on
        first index too. The `else 0` branch is kept as a defensive
        fallback only — under normal flow file_id is always non-zero here.

        Phase 97 LD-3 (LOAD-BEARING per Codex HIGH #2): also writes
        cached_text (zstd-compressed), cached_text_codec, cached_text_uncompressed_len,
        extraction_format_version, and chunk_locator into local_pages in the same
        write. scan_run_id from self._current_scan_run_id (set by _begin_scan_run
        in Task 4; defaults to "" for Wave A compatibility).

        Existing callers (PDF/DOCX/TXT extractors) pass chunk_locator="" by default;
        Wave F D-NEW-5 will pass real locator strings.
        """
        # Get file_id from local_files (pre-inserted by _index_one_file).
        file_row = self._conn.execute(
            "SELECT file_id FROM local_files WHERE sys_id = ?", (sys_id,)
        ).fetchone()
        file_id = file_row["file_id"] if file_row else 0

        uid = _make_uid(sys_id, page_num)
        full_header = _make_full_header(sys_id, page_num, file_id)
        basename = os.path.basename(
            self._conn.execute(
                "SELECT filepath FROM processed_files WHERE sys_id = ?", (sys_id,)
            ).fetchone()["filepath"]
        )
        # Build content_head / content_tail for snippet generation
        words = text.split()
        head = " ".join(words[:50]) if words else ""
        tail = " ".join(words[-50:]) if words else ""

        # Phase 97 LD-1: Tantivy doc now includes scan_run_id + chunk_locator
        doc = tantivy.Document(
            unique_id=[uid],
            content=[text],
            content_head=[head],
            content_tail=[tail],
            line_starts=[""],
            line_ends=[""],
            source=["LOCAL"],
            full_header=[full_header],
            shelfmark=[basename],
            scope=["page"],
            boundaries=[""],
            # Phase 97 LD-1: empty string when scan_run_id is None (Wave A; Wave E populates)
            scan_run_id=[self._current_scan_run_id or ""],
            chunk_locator=[chunk_locator or ""],
        )
        self._writer.add_document(doc)

        # Phase 97 LD-3: compress text and write ALL Phase 97 columns in the SINGLE
        # canonical Tantivy write site (Codex HIGH #2 / ADVICE LD-3).
        cached_bytes, uncompressed_len = compress_cached_text(text)
        self._conn.execute(
            "INSERT OR REPLACE INTO local_pages "
            "(sys_id, uid, page_num, cached_text, cached_text_codec, "
            " cached_text_uncompressed_len, extraction_format_version, chunk_locator) "
            "VALUES (?, ?, ?, ?, 'zstd', ?, 1, ?)",
            (sys_id, uid, page_num, cached_bytes, uncompressed_len, chunk_locator or ""),
        )
        self._conn.commit()
        return uid

    def _extract_and_write_pdf(
        self,
        sys_id: str,
        filepath: str,
        folder_id: int,
        cancel_check: Callable[[], bool],
    ) -> tuple[int, str, str]:
        """Extract PDF pages and write Tantivy docs. Returns (pages_written, status, title)."""
        pages_written = 0
        total_chars = 0
        display_title = os.path.basename(filepath)

        for page_num, text, title in extract_pdf_pages(filepath):
            if cancel_check():
                # Rollback partial pages
                self._rollback_partial(sys_id)
                return (pages_written, "cancelled", display_title)
            display_title = title
            total_chars += len(text)
            # Phase 97 D-NEW-5: PDF locator — "p. N" (1-based page number).
            self._write_page_doc(
                sys_id, page_num, text, title, folder_id,
                chunk_locator=f"p. {page_num}",
            )
            pages_written += 1

        if total_chars < _SCANNED_PDF_CHAR_THRESHOLD and pages_written == 0:
            return (0, "no_text_layer", display_title)

        return (pages_written, "ok", display_title)

    def _extract_and_write_docx(
        self,
        sys_id: str,
        filepath: str,
        folder_id: int,
        cancel_check: Callable[[], bool],
    ) -> tuple[int, str, str]:
        """Extract DOCX chunks and write Tantivy docs. Returns (pages_written, status, title)."""
        pages_written = 0
        display_title = os.path.basename(filepath)

        for chunk_num, text, title in extract_docx_pages(filepath):
            if cancel_check():
                self._rollback_partial(sys_id)
                return (pages_written, "cancelled", display_title)
            display_title = title
            # Phase 97 D-NEW-5: DOCX locator — "paragraphs N-M" (1-based, _DOCX_CHUNK_PARAGRAPHS window).
            _para_start = (chunk_num - 1) * _DOCX_CHUNK_PARAGRAPHS + 1
            _para_end = chunk_num * _DOCX_CHUNK_PARAGRAPHS
            self._write_page_doc(
                sys_id, chunk_num, text, title, folder_id,
                chunk_locator=f"paragraphs {_para_start}-{_para_end}",
            )
            pages_written += 1

        return (pages_written, "ok", display_title)

    def _extract_and_write_txt(
        self,
        sys_id: str,
        filepath: str,
        folder_id: int,
    ) -> tuple[int, str, str]:
        """Extract TXT and write a single Tantivy doc. Returns (pages_written, status, title)."""
        pages_written = 0
        display_title = os.path.basename(filepath)

        for page_num, text, title in extract_txt(filepath):
            display_title = title
            self._write_page_doc(sys_id, page_num, text, title, folder_id)
            pages_written += 1

        return (pages_written, "ok", display_title)

    def _extract_and_write_html(
        self,
        sys_id: str,
        filepath: str,
        folder_id: int,
        cancel_check: Callable[[], bool],
    ) -> tuple[int, str, str]:
        """Phase 97 F-01: Extract HTML chunks and write Tantivy docs.

        Yields (chunk_num, text, title, locator, is_rtl) from extract_html_pages.
        is_rtl is logged at DEBUG level (stored in-memory only for Phase 97;
        Wave F D-NEW-5 can persist as a column if needed downstream).
        Returns (pages_written, status, title).
        """
        pages_written = 0
        display_title = os.path.basename(filepath)

        for chunk_num, text, title, locator, is_rtl in extract_html_pages(filepath):
            if cancel_check():
                self._rollback_partial(sys_id)
                return (pages_written, "cancelled", display_title)
            display_title = title
            self._write_page_doc(sys_id, chunk_num, text, title, folder_id, chunk_locator=locator)
            pages_written += 1
            logger.debug(
                "Phase 97 F-06 is_rtl=%s for sys_id=%s chunk=%d", is_rtl, sys_id, chunk_num
            )

        return (pages_written, "ok" if pages_written > 0 else "no_text_layer", display_title)

    def _extract_and_write_xlsx(
        self,
        sys_id: str,
        filepath: str,
        folder_id: int,
        cancel_check: Callable[[], bool],
    ) -> tuple[int, str, str]:
        """Phase 97 F-02: Extract XLSX chunks and write Tantivy docs.

        XlsxZipBombSuspected propagates to _index_one_file's except clause (LD-9 dual write).
        Returns (pages_written, status, title).
        """
        pages_written = 0
        display_title = os.path.basename(filepath)

        for chunk_num, text, title, locator, is_rtl in extract_xlsx_pages(filepath):
            if cancel_check():
                self._rollback_partial(sys_id)
                return (pages_written, "cancelled", display_title)
            display_title = title
            self._write_page_doc(sys_id, chunk_num, text, title, folder_id, chunk_locator=locator)
            pages_written += 1
            logger.debug(
                "Phase 97 F-06 is_rtl=%s for sys_id=%s chunk=%d", is_rtl, sys_id, chunk_num
            )

        return (pages_written, "ok" if pages_written > 0 else "no_text_layer", display_title)

    def _extract_and_write_csv(
        self,
        sys_id: str,
        filepath: str,
        folder_id: int,
        cancel_check: Callable[[], bool],
    ) -> tuple[int, str, str]:
        """Phase 97 F-03 + F-05: Extract CSV chunks and write Tantivy docs.

        EncodingError propagates to _index_one_file's except clause (LD-9 dual write).
        Returns (pages_written, status, title).
        """
        pages_written = 0
        display_title = os.path.basename(filepath)

        for chunk_num, text, title, locator in extract_csv_pages(filepath):
            if cancel_check():
                self._rollback_partial(sys_id)
                return (pages_written, "cancelled", display_title)
            display_title = title
            self._write_page_doc(sys_id, chunk_num, text, title, folder_id, chunk_locator=locator)
            pages_written += 1

        return (pages_written, "ok" if pages_written > 0 else "no_text_layer", display_title)

    def _rollback_partial(self, sys_id: str) -> None:
        """Roll back partial page docs for a file that was cancelled mid-extraction."""
        # Remove any local_pages rows added so far for this sys_id
        uid_rows = self._conn.execute(
            "SELECT uid FROM local_pages WHERE sys_id = ?", (sys_id,)
        ).fetchall()
        for uid_row in uid_rows:
            self._writer.delete_documents("unique_id", uid_row["uid"])
        try:
            self._writer.rollback()
        except Exception:
            pass
        try:
            self._writer = self._index.writer(heap_size=15_000_000)
        except Exception:
            pass
        self._conn.execute("DELETE FROM local_pages WHERE sys_id = ?", (sys_id,))
        self._conn.execute("DELETE FROM processed_files WHERE sys_id = ?", (sys_id,))
        self._conn.commit()

    # ------------------------------------------------------------------
    # Delete (HIGH-3 three-step crash-safe protocol)
    # ------------------------------------------------------------------

    def _delete_file(self, sys_id: str, filepath: str) -> None:
        """Crash-safe three-step delete (HIGH-3 review fix).

        Step 1: Mark pending_delete=1 in SQLite (durable; recovery picks this up).
        Step 2: Tantivy delete-by-uid + commit.
        Step 3: SQLite final cleanup.

        If crash between step 1 and step 2: _recover_pending_deletes() replays steps 2-3.
        If crash between step 2 and step 3: _recover_pending_deletes() re-runs step 2
                                             (idempotent) then step 3.
        """
        # Step 1: Mark pending_delete
        self._conn.execute(
            "UPDATE local_files SET pending_delete = 1 WHERE sys_id = ?", (sys_id,)
        )
        self._conn.commit()

        # Step 2: Tantivy delete-by-uid
        uid_rows = self._conn.execute(
            "SELECT uid FROM local_pages WHERE sys_id = ?", (sys_id,)
        ).fetchall()
        for uid_row in uid_rows:
            uid = uid_row["uid"]
            self._writer.delete_documents("unique_id", uid)
        try:
            self._writer.commit()
        except Exception as exc:
            logger.warning(
                "_delete_file: Tantivy commit failed for sys_id=%s: %s", sys_id, exc
            )
            return  # Step 3 will be replayed by recovery

        # Step 3: SQLite final cleanup
        self._conn.execute("DELETE FROM local_pages WHERE sys_id = ?", (sys_id,))
        self._conn.execute("DELETE FROM local_files WHERE sys_id = ?", (sys_id,))
        self._conn.execute(
            "DELETE FROM processed_files WHERE sys_id = ?", (sys_id,)
        )
        self._conn.commit()

    # ------------------------------------------------------------------
    # Two-phase commit (D-21)
    # ------------------------------------------------------------------

    @staticmethod
    def _is_windows_access_denied(exc: BaseException) -> bool:
        """Detect Windows ``os error 5`` (ERROR_ACCESS_DENIED) raised by the
        Tantivy Python binding during writer.commit().

        Tantivy surfaces the OS error inside a ValueError whose message
        contains ``"An IO error occurred"`` and ``"os error 5"``.  We match
        on the message rather than the exception type because the binding
        normalizes platform errors into ValueError.
        """
        msg = str(exc) or ""
        return (
            "os error 5" in msg
            or "Access is denied" in msg
            or "access is denied" in msg
        )

    def _commit_writer_with_retry(self) -> None:
        """Commit the Tantivy writer with retry/backoff on Windows os error 5.

        User-reported BLOCKER (Category 2): indexing large folders on Windows
        intermittently hit ``ValueError: An IO error occurred: 'Access is
        denied. (os error 5)'`` during ``self._writer.commit()``. The typical
        Windows causes are:
          (a) an antivirus / Windows Defender briefly scanning a new segment
              file just after rename;
          (b) a reader (e.g. a live SearchEngine.local_searcher) holding a
              handle the writer needs to atomically rename.

        Retry envelope: up to 3 attempts at 250 ms, 1 s, 2 s.  Only retried
        on Windows-access-denied; all other exceptions propagate immediately.
        On final failure, raise ValueError with detailed context (which dir,
        attempts, the underlying message).
        """
        import time as _time

        attempts = 0
        delays = (0.25, 1.0, 2.0)
        last_exc = None
        for delay in (0.0, *delays):
            if delay > 0:
                _time.sleep(delay)
            attempts += 1
            try:
                self._writer.commit()
                if attempts > 1:
                    logger.info(
                        "Tantivy writer.commit() succeeded on retry %d (dir=%s)",
                        attempts, self._index_dir,
                    )
                return
            except Exception as exc:  # noqa: BLE001
                last_exc = exc
                if not self._is_windows_access_denied(exc):
                    raise
                logger.warning(
                    "Tantivy writer.commit() hit Windows access-denied "
                    "(attempt %d/%d, dir=%s): %s",
                    attempts, len(delays) + 1, self._index_dir, exc,
                )
        # All retries exhausted.
        raise ValueError(
            f"Tantivy writer.commit() failed after {attempts} attempts "
            f"with Windows access-denied error on {self._index_dir!r}. "
            f"Pending files: {len(self._pending_filepaths)}. "
            f"Underlying error: {last_exc!r}. "
            "Common causes: (a) antivirus scanning new segment files, "
            "(b) another process holding a reader handle on this index. "
            "Try closing the live SearchEngine reader before commit, or "
            "exclude the LocalIndex directory from real-time antivirus scanning."
        )

    def _commit_batch(self) -> None:
        """Two-phase commit per D-21.

        Phase 1: Tantivy writer.commit() (Tantivy docs are now durable on disk).
        Phase 2: SQLite UPDATE status='committed' for all pending filepaths.

        Category-2 BLOCKER fix: wraps Tantivy commit in retry/backoff so
        a transient Windows ``os error 5`` (antivirus or reader contention)
        does not crash the entire indexing batch.

        Phase 97 R-04 + LD-8: Phase 2 wrapped in PRAGMA synchronous=FULL +
        BEGIN IMMEDIATE + ROLLBACK on failure to ensure the durability bracket
        survives power loss between the Tantivy commit and the SQLite UPDATE.
        """
        if not self._pending_filepaths:
            return
        # Phase 1: Tantivy commit (with retry for Windows access-denied races).
        self._commit_writer_with_retry()
        # Phase 97 R-04 + LD-8: durability bracket with ROLLBACK on failure.
        self._conn.execute("PRAGMA synchronous = FULL")
        try:
            self._conn.execute("BEGIN IMMEDIATE")
            placeholders = ",".join("?" * len(self._pending_filepaths))
            self._conn.execute(
                f"UPDATE processed_files SET status = 'committed' "  # noqa: S608
                f"WHERE filepath IN ({placeholders})",
                self._pending_filepaths,
            )
            # Phase 97 C-04 + LD-9: refresh folder counters from local_files
            # INSIDE this same BEGIN IMMEDIATE transaction so counter update
            # and status='committed' UPDATE are atomic.
            self._refresh_folder_counters_for(list(self._pending_filepaths))
            self._conn.execute("COMMIT")
        except Exception:
            # LD-8: ROLLBACK before restoring synchronous=NORMAL.
            try:
                self._conn.execute("ROLLBACK")
            except sqlite3.OperationalError:
                pass  # already rolled back (no active transaction)
            raise
        finally:
            self._conn.execute("PRAGMA synchronous = NORMAL")
        self._pending_filepaths.clear()

    def _refresh_folder_counters_for(self, filepaths: list) -> None:
        """Phase 97 C-04 + LD-9 — recompute aggregate counts for folders touched
        by this batch. Uses local_files.folder_id + local_files.extraction_status
        (UI source of truth per LD-9) — NOT processed_files.status + path LIKE prefix.

        Codex MEDIUM #3 fix: ``LIKE 'C:\\foo%'`` would match ``C:\\foobar``.
        The replacement query JOINs local_files ON folders.folder_id (the
        relational key) and aggregates extraction_status directly.

        Must be called INSIDE an active BEGIN IMMEDIATE transaction
        (enforced by _commit_batch caller context).
        """
        if not filepaths:
            return
        placeholders = ",".join("?" * len(filepaths))
        self._conn.execute(
            f"""
            UPDATE folders SET
                indexed_count = (
                    SELECT COUNT(*) FROM local_files
                    WHERE local_files.folder_id = folders.folder_id
                      AND local_files.extraction_status IN ('ok', 'committed')
                ),
                error_count = (
                    SELECT COUNT(*) FROM local_files
                    WHERE local_files.folder_id = folders.folder_id
                      AND local_files.extraction_status IN (
                          'error', 'encoding_error', 'changed_during_index', 'no_text_layer'
                      )
                ),
                pending_count = (
                    SELECT COUNT(*) FROM local_files
                    WHERE local_files.folder_id = folders.folder_id
                      AND local_files.extraction_status = 'pending'
                ),
                oversized_count = (
                    SELECT COUNT(*) FROM local_files
                    WHERE local_files.folder_id = folders.folder_id
                      AND local_files.extraction_status IN ('oversized', 'zip_bomb_suspected')
                ),
                last_aggregate_at = strftime('%s', 'now')
            WHERE folder_id IN (
                SELECT DISTINCT local_files.folder_id FROM local_files
                WHERE local_files.filepath IN ({placeholders})
            )
            """,  # noqa: S608
            filepaths,
        )

    # ------------------------------------------------------------------
    # Phase 97 R-02: Atomic rebuild helpers (LD-4 / LD-5)
    # ------------------------------------------------------------------

    def _retry_windows_rename(self, src: str, dst: str) -> None:
        """Rename src->dst with Windows-access-denied retry (same shape as _commit_writer_with_retry)."""
        import time as _time
        delays = (0.0, 0.25, 1.0, 2.0)
        last_exc = None
        for delay in delays:
            if delay > 0:
                _time.sleep(delay)
            try:
                os.rename(src, dst)
                return
            except OSError as exc:
                last_exc = exc
                if not self._is_windows_access_denied(exc):
                    raise
        raise OSError(
            f"Atomic rename {src!r} -> {dst!r} failed after retries: {last_exc!r}"
        )

    def _close_internal_writer_index(self) -> None:
        """Close LocalIndexer's own _writer + _index handles before os.rename (LD-5)."""
        try:
            if self._writer is not None:
                try:
                    self._writer.rollback()
                except Exception:
                    pass
                self._writer = None
        except Exception:
            self._writer = None
        self._index = None

    def _reopen_internal_writer_index(self) -> None:
        """Reopen _writer + _index after atomic swap so LocalIndexer can continue."""
        self._index = tantivy.Index(build_local_schema(), path=self._index_dir)
        self._writer = self._index.writer(heap_size=256 * 1024 * 1024)

    def _reextract_one_page_or_skip(self, filepath: str, page_num: int) -> str:
        """Best-effort re-extract a single page from source file; return "" on failure."""
        try:
            ext = os.path.splitext(filepath)[1].lower()
            if ext == ".pdf":
                for pn, txt, _t in extract_pdf_pages(filepath):
                    if pn == page_num:
                        return txt
            elif ext == ".docx":
                for pn, txt, _t in extract_docx_pages(filepath):
                    if pn == page_num:
                        return txt
            elif ext == ".txt":
                for pn, txt, _t in extract_txt(filepath):
                    if pn == page_num:
                        return txt
        except Exception:
            return ""
        return ""

    def rebuild_main_index_atomic(
        self,
        scan_run_id: str,
        close_searcher_cb: Callable[[], None],
        reload_searcher_cb: Callable[[], None],
    ) -> None:
        """Phase 97 R-02 — atomic temp-dir swap with full 7-handle closure.

        7-step protocol (RESEARCH Issue #3 + Codex HIGH #4 + ADVICE LD-4/LD-5):
          1. Build fresh index in <dir>.rebuild-<scan_run_id>/ from cached_text rows
             (fallback to source re-extract for NULL-cached_text rows).
          2. Validate via fresh_index.searcher().
          3. Close ALL 7 handles (external searcher CB + internal writer/index).
          4. os.rename(live -> <dir>.old-<ts>) with Windows retry.
          5. os.rename(rebuild -> live) with Windows retry; rollback on failure.
          6. Reload all searchers (engine + internal).
          7. Schedule .old- cleanup via pending_dir_cleanup INSERT.
        """
        import shutil
        import time as _time

        rebuild_dir = f"{self._index_dir}.rebuild-{scan_run_id}"
        old_dir = f"{self._index_dir}.old-{int(_time.time())}"

        # --- Step 1: build fresh index from cached_text ---
        if os.path.isdir(rebuild_dir):
            shutil.rmtree(rebuild_dir, ignore_errors=True)
        os.makedirs(rebuild_dir, exist_ok=True)
        fresh_schema = build_local_schema()
        fresh_index = tantivy.Index(fresh_schema, path=rebuild_dir)
        fresh_writer = fresh_index.writer(heap_size=256 * 1024 * 1024)
        docs_written = 0
        try:
            for row in self._conn.execute("""
                SELECT lp.sys_id, lp.uid, lp.page_num,
                       lp.cached_text, lp.cached_text_codec, lp.chunk_locator,
                       lf.original_filename, lf.file_id,
                       pf.scan_run_id AS pf_scan_run_id,
                       pf.filepath
                FROM local_pages lp
                INNER JOIN processed_files pf ON pf.sys_id = lp.sys_id
                INNER JOIN local_files lf ON lf.sys_id = lp.sys_id
                WHERE pf.status = 'committed'
                ORDER BY lp.sys_id, lp.page_num
            """):
                if row["cached_text"] is not None:
                    if row["cached_text_codec"] == "zstd":
                        text = decompress_cached_text(row["cached_text"])
                    else:
                        text = row["cached_text"].decode("utf-8", errors="replace")
                else:
                    # Legacy fallback: re-extract from source file (pre-Phase-97 rows)
                    fp = row["filepath"]
                    if not fp:
                        continue
                    text = self._reextract_one_page_or_skip(fp, row["page_num"])
                    if not text:
                        continue

                full_header = _make_full_header(row["sys_id"], row["page_num"], row["file_id"])
                words = text.split()
                head = " ".join(words[:50]) if words else ""
                tail = " ".join(words[-50:]) if words else ""
                doc = tantivy.Document(
                    unique_id=[row["uid"]],
                    content=[text],
                    content_head=[head],
                    content_tail=[tail],
                    line_starts=[""],
                    line_ends=[""],
                    source=["LOCAL"],
                    full_header=[full_header],
                    shelfmark=[row["original_filename"]],
                    scope=["page"],
                    boundaries=[""],
                    scan_run_id=[row["pf_scan_run_id"] or ""],
                    chunk_locator=[row["chunk_locator"] or ""],
                )
                fresh_writer.add_document(doc)
                docs_written += 1

            fresh_writer.commit()
            fresh_writer.wait_merging_threads()

            # --- Step 2: validate ---
            _validation_searcher = fresh_index.searcher()
            # R97.2-B: explicit null + gc.collect() forces immediate Rust-side
            # drop of writer + index BEFORE os.rename below. `del` is "weak" on
            # Windows because Python GC may delay the Rust drop, causing
            # os.rename to race against a still-held lock file.
            _validation_searcher = None
            fresh_writer = None
            fresh_index = None
            gc.collect()

        except Exception:
            shutil.rmtree(rebuild_dir, ignore_errors=True)
            raise

        if docs_written == 0:
            # No committed rows at all — this is an error: we refuse to swap in an
            # empty index when SQLite says there should be data (prevent silent data loss)
            # UNLESS there genuinely are no committed rows.
            committed_count = self._conn.execute(
                "SELECT COUNT(*) FROM processed_files WHERE status = 'committed'"
            ).fetchone()[0]
            if committed_count > 0:
                shutil.rmtree(rebuild_dir, ignore_errors=True)
                raise RuntimeError(
                    f"Rebuild produced 0 docs but SQLite has {committed_count} committed rows. "
                    "All cached_text is NULL and source files are missing. "
                    "Use 'Reset My Library' in Advanced settings to delete the cache and rescan."
                )

        # --- Step 3: close ALL 7 handles (LD-5) ---
        close_searcher_cb()               # closes engine: local_searcher, local_index,
                                          #   local_lab_searcher, _local_lab_index
        self._close_internal_writer_index()  # closes LocalIndexer._writer + _index

        # --- Steps 4+5: rename live -> .old, rebuild -> live ---
        try:
            self._retry_windows_rename(self._index_dir, old_dir)
        except OSError:
            # Could not move live aside — reopen handles and re-raise
            self._reopen_internal_writer_index()
            reload_searcher_cb()
            raise

        try:
            self._retry_windows_rename(rebuild_dir, self._index_dir)
        except OSError:
            # Rollback: restore old as live
            try:
                self._retry_windows_rename(old_dir, self._index_dir)
            except OSError:
                logger.error(
                    "CRITICAL: both rename steps failed during atomic rebuild; "
                    "old index at %s may be unreachable", old_dir,
                )
            self._reopen_internal_writer_index()
            reload_searcher_cb()
            raise

        # --- Step 6: reload ---
        self._reopen_internal_writer_index()
        reload_searcher_cb()

        # --- Step 7: write schema marker + schedule .old- cleanup ---
        _write_schema_marker(self._index_dir, _compute_schema_marker(build_local_schema))
        self._conn.execute(
            "INSERT OR REPLACE INTO pending_dir_cleanup (path, kind, created_at) "
            "VALUES (?, 'rebuild_old', strftime('%s','now'))",
            (old_dir,),
        )
        self._conn.commit()
        logger.info(
            "Phase 97 R-02: atomic rebuild complete (%d docs, old dir scheduled for GC: %s)",
            docs_written, old_dir,
        )

    def rebuild_lab_index_atomic(
        self,
        scan_run_id: str,
        close_cb: Callable[[], None],
        reload_cb: Callable[[], None],
    ) -> None:
        """Mirror of rebuild_main_index_atomic for LOCAL_LAB_INDEX_DIR.

        LAB index is derivable (D-09) — for Wave A this delegates to the
        caller's reload_cb which will call rebuild_lab_index() from genizah_core.
        Full implementation follows in Wave D/E.
        """
        # LAB index is always derivable from the main index; Wave A stub
        # simply closes + schedules a reload via the callback.
        close_cb()
        reload_cb()

    def clean_pending_rebuild_dirs(self) -> None:
        """Delete .old-<ts> rebuild directories recorded in pending_dir_cleanup."""
        import shutil
        rows = self._conn.execute(
            "SELECT path FROM pending_dir_cleanup WHERE kind = 'rebuild_old'"
        ).fetchall()
        for r in rows:
            try:
                shutil.rmtree(r["path"], ignore_errors=True)
                self._conn.execute(
                    "DELETE FROM pending_dir_cleanup WHERE path = ?", (r["path"],)
                )
            except Exception as exc:
                logger.warning(
                    "Phase 97: clean_pending_rebuild_dirs: remove %r failed: %r",
                    r["path"], exc,
                )
        self._conn.commit()

    # ------------------------------------------------------------------
    # Phase 97 LD-6: scan_runs lifecycle helpers
    # ------------------------------------------------------------------

    def _begin_scan_run(self) -> str:
        """Insert a scan_runs row with status='running'; return the new scan_run_id.

        Called at the start of scan_all so that an interrupted run is detectable
        via start_recovery_probe() on next startup (replaces _pending_cleanup sentinel).
        """
        import uuid as _uuid
        import time as _time
        run_id = _uuid.uuid4().hex
        self._conn.execute(
            "INSERT INTO scan_runs (scan_run_id, started_at, status) VALUES (?, ?, 'running')",
            (run_id, _time.time()),
        )
        self._conn.commit()
        self._current_scan_run_id = run_id
        logger.info("Phase 97 LD-6: scan_run begin scan_run_id=%s", run_id)
        return run_id

    def _end_scan_run(self, scan_run_id: str, status: str = "completed") -> None:
        """Update scan_runs SET status=?, ended_at=NOW() for the given scan_run_id.

        Valid status values: 'completed', 'canceled', 'discarded'.
        Clears self._current_scan_run_id if this run matches.
        """
        import time as _time
        assert status in ("completed", "canceled", "discarded"), status
        self._conn.execute(
            "UPDATE scan_runs SET ended_at = ?, status = ? WHERE scan_run_id = ?",
            (_time.time(), status, scan_run_id),
        )
        self._conn.commit()
        logger.info("Phase 97 LD-6: scan_run end scan_run_id=%s status=%s", scan_run_id, status)
        if self._current_scan_run_id == scan_run_id:
            self._current_scan_run_id = None

    def start_recovery_probe(self) -> list:
        """Phase 97 R-01: return scan_run_ids with status='running'.

        Called from MyLibraryTab.__init__ after migrations; an interrupted run
        (crash / power loss between Tantivy commit and SQLite COMMIT) leaves a
        row with status='running'. MyLibraryTab shows the 3-button recovery modal
        when this returns a non-empty list.

        Returns: list of scan_run_id strings (str, not Row objects).
        """
        rows = self._conn.execute(
            "SELECT scan_run_id FROM scan_runs WHERE status = 'running'"
        ).fetchall()
        return [r["scan_run_id"] for r in rows]

    def discard_run(self, run_id: str) -> dict:
        """Phase 97 U-02 + LD-7 — remove all docs/rows from this scan_run.

        Returns dict of counts removed per source.

        Protocol per ADVICE LD-7:
          1) rollback uncommitted writer state (writer.rollback() or fallback)
          2) writer.delete_documents("scan_run_id", run_id) + commit
          3) ONE SQLite transaction:
             DELETE local_pages / local_files / processed_files for this run
             UPDATE scan_runs SET status='discarded'
          4) refresh folder counters for affected paths
          5) reopen writer/index handles (rollback may have invalidated them)

        Pinned by tests/test_scan_run_id.py.
        """
        import sqlite3 as _sqlite3
        counts: dict = {
            "tantivy_rolled_back": False,
            "local_pages": 0,
            "local_files": 0,
            "processed_files": 0,
        }

        # --- Step 1: rollback uncommitted writer state ---
        if self._writer is not None:
            try:
                self._writer.rollback()
                counts["tantivy_rolled_back"] = True
                # rollback() invalidates the writer — null it so Step 2 opens a fresh one
                self._writer = None
            except Exception as exc:  # noqa: BLE001
                logger.warning(
                    "Phase 97 LD-7: writer.rollback() not available or failed: %r "
                    "— falling back to commit-then-delete",
                    exc,
                )
                try:
                    self._commit_writer_with_retry()
                except Exception:
                    logger.exception(
                        "Phase 97 LD-7: fallback commit failed; proceeding with delete-by-term"
                    )
                # Null writer so Step 5 reopens a fresh handle
                self._writer = None

        # --- Step 2: delete committed docs by term + commit ---
        # Open a fresh writer for the delete-by-term operation (needed whether we
        # rolled back or committed in Step 1, since both invalidate the old writer).
        try:
            _del_writer = self._index.writer(heap_size=15_000_000)
            _del_writer.delete_documents("scan_run_id", run_id)
            _del_writer.commit()
        except Exception:
            logger.exception("Phase 97 LD-7: delete_documents/commit failed for run %s", run_id)

        # --- Step 3: single SQLite transaction for all related row deletions ---
        # Collect affected filepaths first so we can refresh counters after.
        affected_rows = self._conn.execute(
            "SELECT filepath FROM processed_files WHERE scan_run_id = ?", (run_id,)
        ).fetchall()
        affected_paths = [r["filepath"] if isinstance(r, dict) else r[0] for r in affected_rows]

        try:
            self._conn.execute("BEGIN IMMEDIATE")
            cur = self._conn.execute(
                "DELETE FROM local_pages WHERE sys_id IN ("
                "  SELECT sys_id FROM processed_files WHERE scan_run_id = ?"
                ")",
                (run_id,),
            )
            counts["local_pages"] = cur.rowcount
            cur = self._conn.execute(
                "DELETE FROM local_files WHERE sys_id IN ("
                "  SELECT sys_id FROM processed_files WHERE scan_run_id = ?"
                ")",
                (run_id,),
            )
            counts["local_files"] = cur.rowcount
            cur = self._conn.execute(
                "DELETE FROM processed_files WHERE scan_run_id = ?", (run_id,)
            )
            counts["processed_files"] = cur.rowcount
            import time as _time
            self._conn.execute(
                "UPDATE scan_runs SET status = 'discarded', ended_at = ? "
                "WHERE scan_run_id = ?",
                (_time.time(), run_id),
            )
            self._conn.execute("COMMIT")
        except Exception:
            try:
                self._conn.execute("ROLLBACK")
            except _sqlite3.OperationalError:
                pass
            raise

        # --- Step 4: refresh folder counters for affected folders ---
        if affected_paths:
            try:
                self._refresh_folder_counters_for(affected_paths)
                self._conn.commit()
            except Exception:
                logger.exception("Phase 97 LD-7: _refresh_folder_counters_for failed after discard")

        # --- Step 5: reopen writer/index handles (rollback may have invalidated) ---
        try:
            self._writer = self._index.writer(heap_size=15_000_000)
        except Exception:
            logger.exception("Phase 97 LD-7: writer reopen failed after discard")

        logger.info("Phase 97 LD-7 discard_run scan_run_id=%s counts=%s", run_id, counts)
        if self._current_scan_run_id == run_id:
            self._current_scan_run_id = None
        return counts

    def keep_run(self, run_id: str) -> None:
        """Phase 97 U-02 — explicit final commit + mark scan_run completed.

        Called when the user chooses 'Keep partial' from the Cancel modal.
        """
        if self._writer is not None:
            try:
                self._commit_writer_with_retry()
            except Exception:
                logger.exception("Phase 97 U-02: keep_run commit failed for run %s", run_id)
        self._end_scan_run(run_id, "completed")
        logger.info("Phase 97 U-02 keep_run scan_run_id=%s preserved", run_id)

    # ------------------------------------------------------------------
    # LOCAL LAB side-index builder (D-09 + D-38 + HIGH-4 + W5 Option C)
    # ------------------------------------------------------------------

    def build_lab_side_index(
        self,
        lab_weights: dict,
        *,
        fingerprint_dyn_fn: Callable,   # Callable[[str, dict | None], str] — computes fingerprint_dyn
        fingerprint_static_fn: Callable, # Callable[[str], str] — computes static fingerprint
        normalize_text_fn: Callable,     # Callable[[str], str] — normalizes content
        lab_schema_version: int = 1,
        dynamic_rank_map=None,
    ) -> None:
        """Build LOCAL LAB side-index (D-09) and write .meta.json (D-38).

        W5 LOCKED — Option C: fingerprint helpers are passed as keyword-only
        callback functions from genizah_core.py's SearchEngine. This keeps
        shared/local_indexer.py free of any genizah_core import (no circular-dep
        risk). Options A and B are STRUCK per plan 95-06.

        HIGH-4 review fix (option b LOCKED): page text is sourced from the MAIN
        LOCAL Tantivy stored `content` field via _iterate_lab_source_rows(). It
        does NOT re-extract from source files (option a — fails on D-40 unavailable
        folders) and does NOT read from a SQLite `page_text` column (option c —
        this plan does not add that column).

        D-38: writes <LOCAL_LAB_INDEX_DIR>/.meta.json with:
          - weights_hash = sha256(json.dumps(lab_weights, sort_keys=True))
          - lab_schema_version
          - last_built_at (ISO 8601 UTC)

        Arguments:
            lab_weights: dict of current LAB dynamic weights (for weights_hash).
            fingerprint_dyn_fn: callable(content, dynamic_rank_map) -> str.
            fingerprint_static_fn: callable(content) -> str.
            normalize_text_fn: callable(content) -> str.
            lab_schema_version: int (bump when LAB schema changes to invalidate).
            dynamic_rank_map: dict | None (forwarded to fingerprint_dyn_fn).
        """
        os.makedirs(self._lab_index_dir, exist_ok=True)
        schema = build_local_lab_schema()
        lab_index = tantivy.Index(schema, path=self._lab_index_dir)
        # Register tokenizers needed by the LAB schema
        try:
            lab_index.register_tokenizer(
                "whitespace",
                tantivy.TextAnalyzerBuilder(tantivy.Tokenizer.whitespace()).build(),
            )
        except Exception:
            pass  # May fail on reopen — non-fatal
        try:
            lab_index.register_tokenizer(
                "simple",
                tantivy.TextAnalyzerBuilder(tantivy.Tokenizer.simple()).build(),
            )
        except Exception:
            pass

        writer = lab_index.writer(heap_size=50_000_000)

        rows_written = 0
        for sys_id, uid, page_num, file_id, content in self._iterate_lab_source_rows():
            try:
                fingerprint_dyn = fingerprint_dyn_fn(content, dynamic_rank_map)
                fingerprint_static = fingerprint_static_fn(content)
                text_normalized = normalize_text_fn(content)
                full_header = _make_full_header(sys_id, page_num, file_id)

                doc = tantivy.Document(
                    unique_id=uid,
                    content=content,
                    text_normalized=text_normalized,
                    fingerprint=fingerprint_static,
                    fingerprint_dyn=fingerprint_dyn,
                    full_header=full_header,
                    shelfmark=sys_id,
                    source="LOCAL",
                )
                writer.add_document(doc)
                rows_written += 1
            except Exception as exc:
                logger.warning(
                    "build_lab_side_index: failed to write uid %s: %r — skipping", uid, exc
                )

        writer.commit()
        writer.wait_merging_threads()

        # Write .meta.json (D-38).
        weights_hash = hashlib.sha256(
            json.dumps(lab_weights, sort_keys=True).encode("utf-8")
        ).hexdigest()
        meta = {
            "weights_hash": weights_hash,
            "lab_schema_version": lab_schema_version,
            "last_built_at": datetime.datetime.utcnow().isoformat() + "Z",
        }
        meta_path = os.path.join(self._lab_index_dir, ".meta.json")
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(meta, f, indent=2)

        logger.info(
            "LOCAL LAB side-index built: %d pages, weights_hash=%s",
            rows_written,
            weights_hash[:8],
        )

    def _iterate_lab_source_rows(self) -> Iterator[tuple[str, str, int, int, str]]:
        """HIGH-4 review fix (option b LOCKED): yield (sys_id, uid, page_num, file_id, content)
        sourced from the MAIN LOCAL Tantivy stored content field.

        The SQLite local_files table stores metadata only (no page_text column per Plan 03 +
        HIGH-4 decision). `local_pages` tracks the (sys_id, uid, page_num) mapping; the page
        TEXT lives in the main LOCAL Tantivy index at the `content` stored field.

        Generator steps:
          1. SELECT (sys_id, uid, page_num, file_id) from local_pages JOIN local_files.
          2. Open the main LOCAL Tantivy index.
          3. For each uid: search by unique_id (raw tokenizer), retrieve stored `content`.
          4. Yield the tuple.

        Failure modes:
          - main LOCAL Tantivy missing: log WARNING and yield nothing.
          - uid present in local_pages but missing from main LOCAL Tantivy: log WARNING, skip.
        """
        try:
            if not os.path.isdir(self._index_dir):
                logger.warning(
                    "HIGH-4: main LOCAL Tantivy missing at %s; LAB rebuild empty",
                    self._index_dir,
                )
                return
            main_schema = build_local_schema()
            main_index = tantivy.Index(main_schema, path=self._index_dir)
            # Register raw tokenizer so unique_id term queries work
            try:
                main_index.register_tokenizer(
                    "raw",
                    tantivy.TextAnalyzerBuilder(tantivy.Tokenizer.raw()).build(),
                )
            except Exception:
                pass  # Non-fatal if already registered
            searcher = main_index.searcher()
        except Exception as exc:
            logger.warning(
                "HIGH-4: cannot open main LOCAL Tantivy for LAB rebuild: %r", exc
            )
            return

        rows = self._conn.execute(
            "SELECT lp.sys_id, lp.uid, lp.page_num, lf.file_id "
            "FROM local_pages lp "
            "JOIN local_files lf ON lf.sys_id = lp.sys_id "
            "WHERE lf.pending_delete = 0 "
            "ORDER BY lp.sys_id, lp.page_num"
        ).fetchall()

        for row in rows:
            sys_id = row["sys_id"]
            uid = row["uid"]
            page_num = row["page_num"]
            file_id = row["file_id"]
            try:
                # Use raw term query on unique_id (raw tokenizer — Pitfall #2)
                q = main_index.parse_query(uid, ["unique_id"])
                top = searcher.search(q, limit=1).hits
                if not top:
                    logger.warning(
                        "HIGH-4: uid %s in local_pages but missing from main LOCAL Tantivy — skipping",
                        uid,
                    )
                    continue
                _, doc_addr = top[0]
                doc = searcher.doc(doc_addr)
                content = doc.get_first("content") or ""
                yield sys_id, uid, page_num, file_id, content
            except Exception as exc:
                logger.warning("HIGH-4: uid %s read failed: %r — skipping", uid, exc)
                continue

    @staticmethod
    def read_lab_meta(lab_index_dir: str) -> Optional[dict]:
        """Read .meta.json for LOCAL LAB invalidation check (D-38).

        Returns the parsed dict, or None if the file is absent or unreadable.
        """
        meta_path = os.path.join(lab_index_dir, ".meta.json")
        if not os.path.exists(meta_path):
            return None
        try:
            with open(meta_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as exc:
            logger.warning("LOCAL LAB meta read failed: %r", exc)
            return None

    # ------------------------------------------------------------------
    # Close
    # ------------------------------------------------------------------

    def close(self) -> None:
        """Flush + close Tantivy writer and SQLite connection.

        Explicitly deletes the Tantivy writer and index objects to release the
        Tantivy lockfile (LockBusy prevention for subsequent open() calls in
        the same process or a different process).
        """
        if self._pending_filepaths:
            try:
                self._commit_batch()
            except Exception as exc:
                logger.warning("LocalIndexer.close: commit_batch failed: %s", exc)
        try:
            self._writer.commit()
        except Exception:
            pass
        # Delete writer first (releases Tantivy lockfile), then index object
        try:
            del self._writer
        except Exception:
            pass
        try:
            del self._index
        except Exception:
            pass
        try:
            self._conn.close()
        except Exception:
            pass
