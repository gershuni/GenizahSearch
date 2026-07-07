# -*- coding: utf-8 -*-
"""Dump the full BH witness index: main table + tzerufim paragraphs -> TSV/JSON."""
import json
import docx

PATH = r"C:\Users\gersh\Dropbox\שמידמן\ספר ברכת המזון\סיגלה מאוחדת\מפתח כתבי היד.docx"
OUT_JSON = r"C:\Genizahsearch\same_work_spike\probe\data\bh_index_raw.json"
OUT_TSV = r"C:\Genizahsearch\same_work_spike\probe\data\bh_index_raw.tsv"

d = docx.Document(PATH)

# --- main table ---
t = d.tables[0]
rows = []
for r in t.rows[1:]:  # skip header
    cells = [c.text.strip() for c in r.cells]
    if len(cells) >= 4 and (cells[0] or cells[2]):
        rows.append({
            "siglum": cells[0],
            "library_he": cells[1],
            "shelfmark": cells[2],
            "details": cells[3],
        })

# --- tzerufim paragraphs: "SHELFMARK SIGLUM" lines under Heading-2 library headings ---
tzerufim = []
current_lib = None
in_tzerufim = False
outside_genizah = False
for p in d.paragraphs:
    txt = p.text.strip()
    if not txt:
        continue
    style = p.style.name
    if style.startswith("Heading 1"):
        in_tzerufim = "צירופים" in txt
        continue
    if style.startswith("Heading 2"):
        current_lib = txt
        outside_genizah = "מחוץ לגניזה" in txt
        continue
    if in_tzerufim and current_lib and not outside_genizah and style == "Normal":
        # line format: shelfmark tokens ... then siglum (Hebrew) at end
        tzerufim.append({"library_he": current_lib, "line": txt})

data = {"table_rows": rows, "tzerufim_lines": tzerufim}
with open(OUT_JSON, "w", encoding="utf-8") as f:
    json.dump(data, f, ensure_ascii=False, indent=1)

with open(OUT_TSV, "w", encoding="utf-8") as f:
    f.write("siglum\tlibrary_he\tshelfmark\tdetails\n")
    for r in rows:
        f.write(f"{r['siglum']}\t{r['library_he']}\t{r['shelfmark']}\t{r['details'][:100]}\n")

# summary
libs = {}
for r in rows:
    libs[r["library_he"]] = libs.get(r["library_he"], 0) + 1
print(f"table rows: {len(rows)}")
print(f"tzerufim lines: {len(tzerufim)}")
print("libraries (utf-8 escaped):")
for k, v in sorted(libs.items(), key=lambda x: -x[1]):
    print(f"  {k.encode('unicode_escape').decode()[:60]}: {v}")
