# -*- coding: utf-8 -*-
"""Peek at the Birkat Hamazon witness index docx structure."""
import docx

PATH = r"C:\Users\gersh\Dropbox\שמידמן\ספר ברכת המזון\סיגלה מאוחדת\מפתח כתבי היד.docx"
OUT = r"C:\Genizahsearch\same_work_spike\probe\data\bh_docx_peek.txt"

d = docx.Document(PATH)
lines = []
lines.append(f"paragraphs: {len(d.paragraphs)}")
lines.append(f"tables: {len(d.tables)}")
for ti, t in enumerate(d.tables):
    lines.append(f"table {ti}: rows={len(t.rows)} cols={len(t.columns)}")
    # dump first 5 rows of each table
    for r in t.rows[:5]:
        cells = [c.text.strip().replace("\n", " / ") for c in r.cells]
        lines.append("  ROW: " + " || ".join(cells))
lines.append("--- first 80 non-empty paragraphs ---")
count = 0
for p in d.paragraphs:
    txt = p.text.strip()
    if txt:
        lines.append(f"[{p.style.name}] {txt[:200]}")
        count += 1
        if count >= 80:
            break

with open(OUT, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))
print(f"wrote {OUT}, {len(lines)} lines")
