from nicegui import app
from fastapi import Response
from fastapi.responses import RedirectResponse
from web.state import state
import requests
import re
from genizah_core import Config
import io
import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.cell.rich_text import TextBlock, CellRichText
from openpyxl.cell.text import InlineFont
from openpyxl.utils import get_column_letter
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from urllib.parse import urlparse


# Hebrew font name - David is widely available on Windows/Mac and designed for Hebrew
HEBREW_FONT_NAME = "David"


def _set_paragraph_rtl(paragraph):
    """Set RTL direction and right alignment for a paragraph containing Hebrew text."""
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    bidi.set(qn("w:val"), "1")


def _set_run_rtl_font(run, font_size=None):
    """Set Hebrew-compatible font for a run and mark it as RTL."""
    run.font.name = HEBREW_FONT_NAME
    # Set complex script font (for Hebrew/Arabic)
    r = run._r
    rPr = r.get_or_add_rPr()
    # Set cs (complex script) font
    cs_font = rPr.find(qn("w:rFonts"))
    if cs_font is None:
        cs_font = OxmlElement("w:rFonts")
        rPr.insert(0, cs_font)
    cs_font.set(qn("w:cs"), HEBREW_FONT_NAME)
    cs_font.set(qn("w:ascii"), HEBREW_FONT_NAME)
    cs_font.set(qn("w:hAnsi"), HEBREW_FONT_NAME)
    # Mark as RTL text
    rtl = rPr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        rPr.append(rtl)
    if font_size:
        run.font.size = Pt(font_size)


def _apply_rtl_to_document(doc):
    """Apply RTL formatting to all paragraphs in a document."""
    # Set default paragraph style to RTL
    style = doc.styles['Normal']
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Apply RTL to all paragraphs
    for p in doc.paragraphs:
        _set_paragraph_rtl(p)
        for run in p.runs:
            _set_run_rtl_font(run)

# Import corrections API components
from backend.models.database import init_db
from backend.api.routes import auth, users, corrections, comments, documents, versions, discoveries

# Allowed domains for image proxy (prevents SSRF attacks)
ALLOWED_IMAGE_DOMAINS = [
    'rosetta.nli.org.il',
    'iiif.nli.org.il',
    'www.nli.org.il',
    'nli.org.il',
    'hebrew.bodleian.ox.ac.uk',
]

def init_api_routes():
    """Register API routes."""

    # Initialize corrections database
    try:
        init_db()
        print("Corrections database initialized")
    except Exception as e:
        print(f"Warning: Could not initialize corrections database: {e}")

    # Include corrections API routers
    app.include_router(auth.router, prefix="/api/v1", tags=["auth"])
    app.include_router(users.router, prefix="/api/v1", tags=["users"])
    app.include_router(corrections.router, prefix="/api/v1", tags=["corrections"])
    app.include_router(comments.router, prefix="/api/v1", tags=["comments"])
    app.include_router(documents.router, prefix="/api/v1", tags=["documents"])
    app.include_router(versions.router, prefix="/api/v1", tags=["versions"])
    app.include_router(discoveries.router, prefix="/api/v1", tags=["discoveries"])

    def fetch_fl_ids_from_nli(system_id: str, _cache={}, _cache_time={}) -> list:
        """Fetch ALL FL IDs from NLI IIIF manifest (contains all pages). Results are cached for 5 min."""
        import time as _time
        CACHE_TTL = 300  # 5 minutes

        # Check cache first
        if system_id in _cache:
            cache_age = _time.time() - _cache_time.get(system_id, 0)
            if cache_age < CACHE_TTL:
                return _cache[system_id]

        # Use IIIF manifest endpoint - this has ALL page images, unlike MARC which only has 1
        url = f"https://iiif.nli.org.il/IIIFv21/DOCID/PNX_MANUSCRIPTS{system_id}-1/manifest"
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        }
        try:
            resp = requests.get(url, headers=headers, timeout=15, verify=True)
            if resp.status_code == 200:
                data = resp.json()
                fl_ids = []

                # Extract FL IDs from canvas images in order
                if 'sequences' in data and data['sequences']:
                    for canvas in data['sequences'][0].get('canvases', []):
                        images = canvas.get('images', [])
                        if images:
                            resource = images[0].get('resource', {})
                            service = resource.get('service', {})
                            service_id = service.get('@id', '')
                            # Extract FL number (e.g. .../FL7734473/...)
                            fl_match = re.search(r'FL(\d+)', service_id)
                            if fl_match:
                                fl_ids.append(fl_match.group(1))

                if fl_ids:
                    # Cache successful result
                    _cache[system_id] = fl_ids
                    _cache_time[system_id] = _time.time()
                    print(f"Cached {len(fl_ids)} FL IDs for {system_id} from IIIF manifest")
                    return fl_ids
        except Exception as e:
            print(f"Failed to fetch FL IDs from IIIF manifest for {system_id}: {e}")

        # Fallback to MARC API (only has 1 FL ID typically)
        try:
            marc_url = f"https://iiif.nli.org.il/IIIFv21/marc/bib/{system_id}"
            resp = requests.get(marc_url, headers=headers, timeout=10, verify=True)
            if resp.status_code == 200:
                fl_ids = re.findall(r'FL(\d+)', resp.text)
                seen = set()
                unique_fl_ids = []
                for fl_id in fl_ids:
                    if fl_id not in seen:
                        seen.add(fl_id)
                        unique_fl_ids.append(fl_id)
                if unique_fl_ids:
                    _cache[system_id] = unique_fl_ids
                    _cache_time[system_id] = _time.time()
                    print(f"Cached {len(unique_fl_ids)} FL IDs from MARC for {system_id}")
                return unique_fl_ids
        except Exception as e:
            print(f"MARC fallback also failed for {system_id}: {e}")

        return []

    @app.get('/api/nli_image/{fl_id}')
    def nli_image(fl_id: str):
        """
        Fetch NLI image by FL ID. Tries IIIF first (for valid IDs), then Rosetta.
        """
        # Clean the FL ID
        digits = re.sub(r"\D", "", str(fl_id))
        if not digits:
            return Response(content="Invalid FL ID", status_code=400)

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Referer': 'https://www.nli.org.il/',
        }

        # Try IIIF first (works for valid FL IDs, returns real images)
        iiif_url = f"https://iiif.nli.org.il/IIIFv21/FL{digits}/full/max/0/default.jpg"
        try:
            resp = requests.get(iiif_url, headers=headers, timeout=15, verify=True)
            if resp.status_code == 200 and 'image' in resp.headers.get('Content-Type', ''):
                # Verify it's not a tiny placeholder (real images are > 5KB)
                if len(resp.content) > 5000:
                    return Response(
                        content=resp.content,
                        media_type=resp.headers.get('Content-Type', 'image/jpeg')
                    )
        except Exception as e:
            print(f"IIIF failed for FL{digits}: {e}")

        # Fallback to Rosetta - but filter out the "no image" placeholder
        rosetta_url = f"https://rosetta.nli.org.il/delivery/DeliveryManagerServlet?dps_func=thumbnail&dps_pid=FL{digits}"
        try:
            resp = requests.get(rosetta_url, headers=headers, timeout=15, verify=True)
            if resp.status_code == 200 and 'image' in resp.headers.get('Content-Type', ''):
                # The "no image" placeholder is ~1615 bytes, real images are larger
                if len(resp.content) > 2000:
                    return Response(
                        content=resp.content,
                        media_type=resp.headers.get('Content-Type', 'image/png')
                    )
        except Exception as e:
            print(f"Rosetta failed for FL{digits}: {e}")

        return Response(content="Image not found", status_code=404)

    # Image cache: (sys_id, page) -> (content, content_type, timestamp)
    _image_cache = {}
    _IMAGE_CACHE_TTL = 600  # 10 minutes

    @app.get('/api/nli_image_by_sysid/{sys_id}')
    def nli_image_by_sysid(sys_id: str, page: int = 0):
        """
        Fetch NLI image by System ID. Dynamically gets FL IDs from NLI MARC API.
        """
        import time as _time
        cache_key = (sys_id, page)

        # Check image cache first
        if cache_key in _image_cache:
            content, content_type, cached_at = _image_cache[cache_key]
            if _time.time() - cached_at < _IMAGE_CACHE_TTL:
                return Response(
                    content=content,
                    media_type=content_type,
                    headers={"Cache-Control": "public, max-age=600"}
                )

        # Fetch FL IDs from NLI (this function has its own cache)
        fl_ids = fetch_fl_ids_from_nli(sys_id)
        if not fl_ids:
            return Response(content="No FL IDs found for this document", status_code=404)

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Referer': 'https://www.nli.org.il/',
        }

        # If page index specified, try that specific FL ID first
        if 0 <= page < len(fl_ids):
            fl_id = fl_ids[page]
            iiif_url = f"https://iiif.nli.org.il/IIIFv21/FL{fl_id}/full/max/0/default.jpg"
            try:
                resp = requests.get(iiif_url, headers=headers, timeout=15, verify=True)
                if resp.status_code == 200 and 'image' in resp.headers.get('Content-Type', '') and len(resp.content) > 5000:
                    content_type = resp.headers.get('Content-Type', 'image/jpeg')
                    # Cache the image
                    _image_cache[cache_key] = (resp.content, content_type, _time.time())
                    return Response(
                        content=resp.content,
                        media_type=content_type,
                        headers={"Cache-Control": "public, max-age=600"}
                    )
            except Exception:
                pass

        # Fallback: try each FL ID until one works
        for fl_id in fl_ids:
            iiif_url = f"https://iiif.nli.org.il/IIIFv21/FL{fl_id}/full/max/0/default.jpg"
            try:
                resp = requests.get(iiif_url, headers=headers, timeout=15, verify=True)
                if resp.status_code == 200 and 'image' in resp.headers.get('Content-Type', '') and len(resp.content) > 5000:
                    content_type = resp.headers.get('Content-Type', 'image/jpeg')
                    _image_cache[cache_key] = (resp.content, content_type, _time.time())
                    return Response(
                        content=resp.content,
                        media_type=content_type,
                        headers={"Cache-Control": "public, max-age=600"}
                    )
            except Exception:
                pass

        return Response(content="Image not found", status_code=404)

    # Oxford image cache: (sys_id, page) -> (content, content_type, timestamp)
    _oxford_image_cache = {}

    def _extract_folio_number(shelfmark: str) -> int:
        """Extract folio number from Oxford shelfmark like 'MS heb. f.21/21' -> 21"""
        # Pattern: MS heb. X.YY/ZZ where ZZ is the folio number
        match = re.search(r'/(\d+)', shelfmark)
        if match:
            return int(match.group(1))
        return 0

    @app.get('/api/oxford_image/{sys_id}')
    def oxford_image(sys_id: str, page: int = 0):
        """
        Fetch Oxford image by System ID using CodicologicalManager.
        Automatically finds the correct folio image based on shelfmark.
        """
        import time as _time
        cache_key = (sys_id, page)

        # Check image cache first
        if cache_key in _oxford_image_cache:
            content, content_type, cached_at = _oxford_image_cache[cache_key]
            if _time.time() - cached_at < _IMAGE_CACHE_TTL:
                return Response(
                    content=content,
                    media_type=content_type,
                    headers={"Cache-Control": "public, max-age=600"}
                )

        if not state.meta_mgr or not state.meta_mgr.codico_mgr:
            return Response(content="Oxford manager not initialized", status_code=503)

        codico = state.meta_mgr.codico_mgr
        if not getattr(codico, '_loaded', False):
            return Response(content="Oxford database still loading", status_code=503)

        # Get the Part ID for this system ID
        part_id = codico.get_part_for_folio(sys_id)

        # Get shelfmark for folio number extraction
        shelfmark = ''
        shelfmark_tuple = state.meta_mgr.get_meta_for_id(sys_id)
        if shelfmark_tuple and shelfmark_tuple[0]:
            shelfmark = shelfmark_tuple[0]

        if not part_id and shelfmark:
            # Try to find part by shelfmark
            part_id, is_part = codico.parse_part_identifier(shelfmark)
            if not is_part:
                part_id = None

        if not part_id:
            return Response(content="No Oxford Part found for this document", status_code=404)

        # Get images for this part
        images = codico.get_part_images(part_id)
        if not images:
            return Response(content="No images available for this Part", status_code=404)

        # Extract folio number from shelfmark and find the matching image
        folio_num = _extract_folio_number(shelfmark)
        img_data = None
        img_url = ''

        if folio_num > 0:
            # Find all images matching this folio number (both 'a' and 'b' sides)
            folio_images = [img for img in images if img.get('folio_num') == folio_num]

            if folio_images:
                # Use page param to select: 0 = first (recto/a), 1 = second (verso/b)
                idx = min(page, len(folio_images) - 1)
                img_data = folio_images[idx]
            else:
                # If image not in database but folio is in range, generate URL dynamically
                metadata = codico.part_metadata.get(part_id, {})
                folio_range = metadata.get('folio_range', [])
                if len(folio_range) >= 2 and folio_range[0] <= folio_num <= folio_range[1]:
                    # Generate URL based on part_id pattern
                    # MS. Heb. f. 21/1 -> MS_HEB_f_21_18a.jpg or MS_HEB_f_21_18b.jpg
                    match = re.match(r'^MS\.?\s*Heb\.?\s*([a-z])\.?\s*(\d+)', part_id, re.IGNORECASE)
                    if match:
                        letter, volume = match.groups()
                        # page 0 = recto (a), page 1 = verso (b)
                        side = 'b' if page == 1 else 'a'
                        img_url = f"https://hebrew.bodleian.ox.ac.uk/fragments/full/MS_HEB_{letter}_{volume}_{folio_num}{side}.jpg"

        # Fallback to page index if folio not found and no dynamic URL
        if not img_data and not img_url:
            if page < 0 or page >= len(images):
                page = 0
            img_data = images[page]

        if not img_url and img_data:
            img_url = img_data.get('full_url', '')

        if not img_url:
            return Response(content="No image URL available", status_code=404)

        # Fetch the image from Oxford
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Referer': 'https://hebrew.bodleian.ox.ac.uk/',
        }

        try:
            resp = requests.get(img_url, headers=headers, timeout=30, verify=True)
            if resp.status_code == 200 and 'image' in resp.headers.get('Content-Type', ''):
                content_type = resp.headers.get('Content-Type', 'image/jpeg')
                # Cache the image
                _oxford_image_cache[cache_key] = (resp.content, content_type, _time.time())
                return Response(
                    content=resp.content,
                    media_type=content_type,
                    headers={"Cache-Control": "public, max-age=600"}
                )
            else:
                return Response(content=f"Failed to fetch image: {resp.status_code}", status_code=resp.status_code)
        except Exception as e:
            return Response(content=f"Error fetching image: {e}", status_code=500)

    @app.get('/api/oxford_image_url/{sys_id}')
    def oxford_image_url(sys_id: str, page: int = 0):
        """
        Get the direct Oxford URL for an image (no proxy).
        Returns JSON with the URL that the browser can fetch directly.
        """
        if not state.meta_mgr or not state.meta_mgr.codico_mgr:
            return {"error": "Oxford manager not initialized", "url": None}

        codico = state.meta_mgr.codico_mgr
        if not getattr(codico, '_loaded', False):
            return {"error": "Oxford database still loading", "url": None}

        # Get shelfmark for folio number extraction
        shelfmark = ''
        shelfmark_tuple = state.meta_mgr.get_meta_for_id(sys_id)
        if shelfmark_tuple and shelfmark_tuple[0]:
            shelfmark = shelfmark_tuple[0]

        part_id = codico.get_part_for_folio(sys_id)
        if not part_id and shelfmark:
            part_id, is_part = codico.parse_part_identifier(shelfmark)
            if not is_part:
                part_id = None

        if not part_id:
            return {"error": "No Oxford Part found", "url": None}

        images = codico.get_part_images(part_id)
        folio_num = _extract_folio_number(shelfmark)
        img_url = None

        if folio_num > 0:
            # Find images for this folio
            folio_images = [img for img in images if img.get('folio_num') == folio_num]
            if folio_images:
                idx = min(page, len(folio_images) - 1)
                img_url = folio_images[idx].get('full_url', '')
            else:
                # Generate dynamically if not in database
                metadata = codico.part_metadata.get(part_id, {})
                folio_range = metadata.get('folio_range', [])
                if len(folio_range) >= 2 and folio_range[0] <= folio_num <= folio_range[1]:
                    match = re.match(r'^MS\.?\s*Heb\.?\s*([a-z])\.?\s*(\d+)', part_id, re.IGNORECASE)
                    if match:
                        letter, volume = match.groups()
                        side = 'b' if page == 1 else 'a'
                        img_url = f"https://hebrew.bodleian.ox.ac.uk/fragments/full/MS_HEB_{letter}_{volume}_{folio_num}{side}.jpg"

        if not img_url and images:
            idx = min(page, len(images) - 1)
            img_url = images[idx].get('full_url', '')

        return {"url": img_url, "folio": folio_num, "page": page, "part_id": part_id}

    @app.get('/api/oxford_images/{sys_id}')
    def oxford_images_list(sys_id: str):
        """
        Get list of available Oxford images for a system ID.
        Returns JSON with image metadata.
        """
        if not state.meta_mgr or not state.meta_mgr.codico_mgr:
            return {"error": "Oxford manager not initialized", "images": []}

        codico = state.meta_mgr.codico_mgr
        if not getattr(codico, '_loaded', False):
            return {"error": "Oxford database still loading", "images": []}

        # Get the Part ID for this system ID
        part_id = codico.get_part_for_folio(sys_id)
        if not part_id:
            return {"error": "No Oxford Part found", "images": [], "part_id": None}

        # Get images for this part
        images = codico.get_part_images(part_id)

        return {
            "part_id": part_id,
            "images": [
                {
                    "index": i,
                    "label": img.get('label', ''),
                    "folio_num": img.get('folio_num'),
                    "url": f"/api/oxford_image/{sys_id}?page={i}"
                }
                for i, img in enumerate(images)
            ]
        }

    @app.get('/api/oxford_debug')
    def oxford_debug():
        """
        Debug endpoint to check Oxford mapping status.
        """
        if not state.meta_mgr:
            return {"error": "MetadataManager not initialized"}

        codico = state.meta_mgr.codico_mgr
        if not codico:
            return {"error": "CodicologicalManager not initialized"}

        csv_bank_count = len(state.meta_mgr.csv_bank)
        oxford_part_count = sum(1 for v in state.meta_mgr.csv_bank.values() if v.get('oxford_part_id'))

        return {
            "codico_loaded": getattr(codico, '_loaded', False),
            "csv_bank_entries": csv_bank_count,
            "entries_with_oxford_part_id": oxford_part_count,
            "folio_to_part_mappings": len(codico.folio_to_part),
            "part_metadata_count": len(codico.part_metadata),
            "sample_mappings": list(codico.folio_to_part.items())[:5] if codico.folio_to_part else [],
        }

    @app.get('/api/browse_debug/{sys_id}')
    def browse_debug(sys_id: str):
        """
        Debug endpoint to check browse data for a specific sys_id.
        """
        if not state.meta_mgr:
            return {"error": "MetadataManager not initialized"}
        if not state.searcher:
            return {"error": "SearchEngine not initialized"}

        # Get shelfmark from csv_bank
        csv_entry = state.meta_mgr.csv_bank.get(sys_id, {})
        shelfmark, title = state.meta_mgr.get_meta_for_id(sys_id)

        # Get browse_map data
        browse_map = state.searcher._load_browse_map() if hasattr(state.searcher, '_load_browse_map') else {}
        browse_entries = browse_map.get(sys_id, [])

        # Extract FL IDs from browse entries
        import re
        fl_ids = []
        for entry in browse_entries:
            full_header = entry.get('full_header', '')
            match = re.search(r'FL(\d+)', full_header)
            if match:
                fl_ids.append(match.group(1))

        # Check Oxford mapping
        part_id = None
        if state.meta_mgr.codico_mgr:
            part_id = state.meta_mgr.codico_mgr.get_part_for_folio(sys_id)

        return {
            "sys_id": sys_id,
            "csv_entry": csv_entry,
            "shelfmark": shelfmark,
            "title": title,
            "browse_entries_count": len(browse_entries),
            "fl_ids_found": fl_ids[:5],  # First 5
            "oxford_part_id": part_id,
            "is_oxford_shelfmark": shelfmark.lower().startswith('ms heb') or shelfmark.lower().startswith('ms. heb') if shelfmark else False,
        }

    @app.get('/api/proxy_image')
    def proxy_image(url: str):
        """
        Proxy image requests to NLI to bypass Referer checks.
        Spoofs the Referer header to look like it's coming from nli.org.il.
        """
        if not url:
            return Response(status_code=400)

        # Validate URL format and domain to prevent SSRF attacks
        try:
            parsed = urlparse(url)
            if parsed.scheme not in ('http', 'https'):
                return Response(content="Invalid URL scheme", status_code=400)
            if not parsed.netloc:
                return Response(content="Invalid URL", status_code=400)
            if parsed.netloc not in ALLOWED_IMAGE_DOMAINS:
                return Response(content="Domain not allowed", status_code=403)
        except Exception:
            return Response(content="Invalid URL format", status_code=400)

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Referer': 'https://www.nli.org.il/',
            'Accept': 'image/avif,image/webp,image/apng,image/svg+xml,image/*,*/*;q=0.8',
        }

        try:
            # Fetch the image with timeout
            resp = requests.get(url, headers=headers, timeout=15, verify=True)
            if resp.status_code == 200:
                return Response(
                    content=resp.content,
                    media_type=resp.headers.get('Content-Type', 'image/jpeg')
                )
            else:
                print(f"Proxy got status {resp.status_code} for URL: {url}")
                return Response(status_code=resp.status_code)
        except requests.Timeout:
            print(f"Proxy timeout for URL: {url}")
            return Response(content="Request timeout", status_code=504)
        except Exception as e:
            print(f"Proxy error for {url}: {e}")
            return Response(status_code=500)

    @app.get('/api/export/excel')
    def export_excel():
        if not state.last_results:
            return Response("No results to export", status_code=400)

        # Get search query for filename and highlighting
        search_query = state.current_search_query or ""
        search_terms = [t.strip() for t in search_query.split() if t.strip() and not t.startswith(('=', '?', '~', '/', '$', '#'))]

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Genizah Results"

        # Enable RTL for the sheet (Hebrew content)
        ws.sheet_view.rightToLeft = True

        headers = ["Shelfmark", "Title", "System ID", "Score", "Snippet", "Full Text"]
        ws.append(headers)

        # Style header row
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment

        # Define column widths
        column_widths = {
            'A': 25,   # Shelfmark
            'B': 35,   # Title
            'C': 18,   # System ID
            'D': 10,   # Score
            'E': 50,   # Snippet
            'F': 80,   # Full Text
        }
        for col_letter, width in column_widths.items():
            ws.column_dimensions[col_letter].width = width

        # Alignments for different column types
        # Hebrew text columns: RTL alignment (no wrap for cleaner single-line display)
        rtl_alignment = Alignment(horizontal="right", vertical="top", wrap_text=False, readingOrder=2)
        # LTR columns (System ID, Score)
        ltr_alignment = Alignment(horizontal="left", vertical="top")
        # Score column: number alignment
        score_alignment = Alignment(horizontal="center", vertical="top")
        # Yellow highlight for cells containing search terms
        highlight_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")

        def clean_text_single_line(text):
            """Replace line breaks with spaces and clean up."""
            if not text:
                return ""
            text = text.replace('\r\n', ' ').replace('\n', ' ').replace('\r', ' ')
            while '  ' in text:
                text = text.replace('  ', ' ')
            return text.strip()

        def contains_search_term(text):
            """Check if text contains any search term."""
            if not text or not search_terms:
                return False
            text_lower = text.lower()
            return any(term.lower() in text_lower for term in search_terms)

        for res in state.last_results:
            display = res.get('display', {})

            # Clean snippet - remove highlighting markers and line breaks
            snippet = res.get('snippet', '').replace('*', '')
            snippet = clean_text_single_line(snippet)

            # Clean full text - replace line breaks with spaces
            full_text = res.get('full_text', '')[:32000]  # Excel cell limit safety
            full_text = clean_text_single_line(full_text)

            row = [
                display.get('shelfmark', ''),
                display.get('title', ''),
                display.get('id', ''),
                str(res.get('sort_score', '')),
                snippet,
                full_text
            ]
            # Sanitize for illegal chars
            clean_row = []
            for cell in row:
                if isinstance(cell, str):
                    # Remove illegal chars (XML 1.0 invalid chars)
                    cell = "".join(ch for ch in cell if (0x20 <= ord(ch) <= 0xD7FF) or (0xE000 <= ord(ch) <= 0xFFFD) or ch in "\t")
                clean_row.append(cell)
            ws.append(clean_row)

            # Apply alignment to the row just added
            current_row = ws.max_row
            # Shelfmark (A) - RTL Hebrew
            ws.cell(row=current_row, column=1).alignment = rtl_alignment
            # Title (B) - RTL Hebrew
            ws.cell(row=current_row, column=2).alignment = rtl_alignment
            # System ID (C) - LTR
            ws.cell(row=current_row, column=3).alignment = ltr_alignment
            # Score (D) - centered
            ws.cell(row=current_row, column=4).alignment = score_alignment
            # Snippet (E) - RTL Hebrew + highlight if contains search term
            snippet_cell = ws.cell(row=current_row, column=5)
            snippet_cell.alignment = rtl_alignment
            if contains_search_term(snippet):
                snippet_cell.fill = highlight_fill
            # Full Text (F) - RTL Hebrew + highlight if contains search term
            fulltext_cell = ws.cell(row=current_row, column=6)
            fulltext_cell.alignment = rtl_alignment
            if contains_search_term(full_text):
                fulltext_cell.fill = highlight_fill

        # Add credits at the bottom
        ws.append([])
        ws.append([])
        credits_start_row = ws.max_row + 1
        ws.append(['Credits'])
        ws.append(['Generated by Dicta Genizah Search (Web Version)'])
        ws.append(['Data Source: MiDRASH Automatic Transcriptions (Stoekl Ben Ezra et al., 2025)'])
        ws.append(['Dataset: https://doi.org/10.5281/zenodo.17734473'])
        ws.append(['Citation: Stoekl Ben Ezra, D., Bambaci, L., Kiessling, B., Lapin, H., Ezer, N., Lolli, E., Rustow, M., Dershowitz, N., Kurar Barakat, B., Gogawale, S., Shmidman, A., Lavee, M., Siew, T., Raziel Kretzmer, V., Vasyutinsky Shapira, D., Olszowy-Schlanger, J., & Gila, Y. (2025). MiDRASH Automatic Transcriptions. Zenodo. https://doi.org/10.5281/zenodo.17734473'])

        # Style credits section
        credits_font = Font(italic=True, color="555555")
        for row_idx in range(credits_start_row, ws.max_row + 1):
            cell = ws.cell(row=row_idx, column=1)
            cell.font = credits_font

        stream = io.BytesIO()
        wb.save(stream)
        stream.seek(0)

        # Use search query for filename, sanitized for filesystem
        safe_query = re.sub(r'[\\/*?:"<>|]', '', search_query)[:50] if search_query else "genizah_results"
        safe_query = safe_query.strip() or "genizah_results"
        filename = f"{safe_query}.xlsx"

        return Response(
            content=stream.read(),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{requests.utils.quote(filename)}"}
        )

    @app.get('/api/export/word')
    def export_word():
        if not state.last_results:
            return Response("No results to export", status_code=400)

        # Get search query for filename and title
        search_query = state.current_search_query or ""

        doc = Document()
        title_text = f'Genizah Search Results: "{search_query}"' if search_query else 'Genizah Search Results'
        doc.add_heading(title_text, 0)

        def add_snippet_with_highlighting(paragraph, snippet_text):
            """Add snippet text with highlighted search terms (marked with *)."""
            # The snippet contains *highlighted* terms - parse and format them
            parts = snippet_text.split('*')
            is_highlighted = False
            for i, part in enumerate(parts):
                if not part:
                    is_highlighted = not is_highlighted
                    continue
                run = paragraph.add_run(part)
                _set_run_rtl_font(run)
                if is_highlighted:
                    # Highlight: bold + yellow background
                    run.bold = True
                    from docx.shared import RGBColor
                    from docx.oxml.ns import qn as qn_ns
                    from docx.oxml import OxmlElement
                    # Add yellow highlight
                    rPr = run._r.get_or_add_rPr()
                    highlight = OxmlElement('w:highlight')
                    highlight.set(qn_ns('w:val'), 'yellow')
                    rPr.append(highlight)
                is_highlighted = not is_highlighted

        for i, res in enumerate(state.last_results):
            display = res.get('display', {})
            shelf = display.get('shelfmark', 'Unknown')
            title = display.get('title', '')

            # Header paragraph with shelfmark and title
            p = doc.add_paragraph()
            run = p.add_run(f"{i+1}. {shelf}")
            run.bold = True
            _set_run_rtl_font(run)
            if title:
                title_run = p.add_run(f" - {title}")
                _set_run_rtl_font(title_run)
            _set_paragraph_rtl(p)

            # Snippet (Hebrew manuscript text) with highlighting
            if res.get('snippet'):
                snippet_p = doc.add_paragraph()
                add_snippet_with_highlighting(snippet_p, res['snippet'])
                _set_paragraph_rtl(snippet_p)

            # System ID (LTR)
            doc.add_paragraph(f"System ID: {display.get('id', '')}")
            doc.add_paragraph("_" * 40)

        # Add credits
        doc.add_page_break()
        doc.add_heading('Credits', 1)
        doc.add_paragraph('Generated by Dicta Genizah Search (Web Version)')
        doc.add_paragraph('Data Source: MiDRASH Automatic Transcriptions (Stoekl Ben Ezra et al., 2025)')
        doc.add_paragraph('Dataset available at: https://doi.org/10.5281/zenodo.17734473')
        doc.add_paragraph()
        citation = doc.add_paragraph('Full Citation: ')
        citation.add_run('Stoekl Ben Ezra, D., Bambaci, L., Kiessling, B., Lapin, H., Ezer, N., Lolli, E., Rustow, M., Dershowitz, N., Kurar Barakat, B., Gogawale, S., Shmidman, A., Lavee, M., Siew, T., Raziel Kretzmer, V., Vasyutinsky Shapira, D., Olszowy-Schlanger, J., & Gila, Y. (2025). MiDRASH Automatic Transcriptions. Zenodo. https://doi.org/10.5281/zenodo.17734473').italic = True

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        # Use search query for filename, sanitized for filesystem
        safe_query = re.sub(r'[\\/*?:"<>|]', '', search_query)[:50] if search_query else "genizah_results"
        safe_query = safe_query.strip() or "genizah_results"
        filename = f"{safe_query}.docx"

        return Response(
            content=stream.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{requests.utils.quote(filename)}"}
        )

    @app.get('/api/export/parallels/excel')
    def export_parallels_excel():
        """Export parallels results to Excel."""
        # Read from global state (accessible from HTTP requests without session context)
        parallels_results = state.parallels_results or []
        filtered_results = state.parallels_filtered or []

        if not parallels_results and not filtered_results:
            return Response("No parallels results to export", status_code=400)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Parallels Results"

        # Enable RTL for the sheet (Hebrew content)
        ws.sheet_view.rightToLeft = True

        headers = ["#", "Shelfmark", "Title", "Score", "Source Context", "Manuscript Match", "Filtered"]
        ws.append(headers)

        # Style header row
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment

        # Define column widths
        column_widths = {
            'A': 6,    # #
            'B': 25,   # Shelfmark
            'C': 35,   # Title
            'D': 10,   # Score
            'E': 50,   # Source Context
            'F': 60,   # Manuscript Match
            'G': 10,   # Filtered
        }
        for col_letter, width in column_widths.items():
            ws.column_dimensions[col_letter].width = width

        # Alignments for different column types
        rtl_alignment = Alignment(horizontal="right", vertical="top", wrap_text=True, readingOrder=2)
        ltr_alignment = Alignment(horizontal="left", vertical="top")
        center_alignment = Alignment(horizontal="center", vertical="top")

        def add_results_to_sheet(results, start_idx, is_filtered=False):
            for idx, item in enumerate(results, start_idx):
                score = item.get('score', 0)
                raw_header = item.get('raw_header', '')

                # Extract metadata
                sys_id = None
                shelfmark = 'Unknown'
                title = ''

                if raw_header and state.meta_mgr:
                    try:
                        import re
                        sys_match = re.search(r'(99\d{8,})', raw_header)
                        if sys_match:
                            sys_id = sys_match.group(1)
                            shelf_temp, title_temp = state.meta_mgr.get_meta_for_id(sys_id)
                            shelfmark = shelf_temp or shelfmark
                            title = title_temp or ''
                    except Exception:
                        pass

                source_ctx = item.get('source_ctx', '').replace('*', '')
                # Remove line breaks so text flows continuously in Excel cell
                source_ctx = source_ctx.replace('\n', ' ').replace('\r', ' ')
                while '  ' in source_ctx:
                    source_ctx = source_ctx.replace('  ', ' ')
                source_ctx = source_ctx.strip()

                # Use longer text context - prefer full_text if available, fall back to text
                # Remove line breaks so text flows continuously in Excel cell
                ms_text = item.get('text', '').replace('*', '').replace('\n', ' ').replace('\r', ' ')
                # Clean up multiple spaces
                while '  ' in ms_text:
                    ms_text = ms_text.replace('  ', ' ')
                ms_text = ms_text.strip()

                filtered_mark = 'Yes' if is_filtered else ''
                row = [idx, shelfmark, title, score, source_ctx, ms_text, filtered_mark]

                # Sanitize for illegal chars (remove newlines to keep text flowing)
                clean_row = []
                for cell in row:
                    if isinstance(cell, str):
                        cell = "".join(ch for ch in cell if (0x20 <= ord(ch) <= 0xD7FF) or (0xE000 <= ord(ch) <= 0xFFFD) or ch == "\t")
                    clean_row.append(cell)
                ws.append(clean_row)

                # Apply alignment to the row just added
                current_row = ws.max_row
                ws.cell(row=current_row, column=1).alignment = center_alignment  # #
                ws.cell(row=current_row, column=2).alignment = rtl_alignment     # Shelfmark
                ws.cell(row=current_row, column=3).alignment = rtl_alignment     # Title
                ws.cell(row=current_row, column=4).alignment = center_alignment  # Score
                ws.cell(row=current_row, column=5).alignment = rtl_alignment     # Source Context
                ws.cell(row=current_row, column=6).alignment = rtl_alignment     # Manuscript Match
                ws.cell(row=current_row, column=7).alignment = center_alignment  # Filtered

            return start_idx + len(results)

        # Add main results
        next_idx = add_results_to_sheet(parallels_results, 1, is_filtered=False)

        # Add filtered results
        if filtered_results:
            add_results_to_sheet(filtered_results, next_idx, is_filtered=True)

        # Add credits at the bottom
        ws.append([])
        ws.append([])
        credits_start_row = ws.max_row + 1
        ws.append(['Credits'])
        ws.append(['Generated by Dicta Genizah Search (Web Version)'])
        ws.append(['Data Source: MiDRASH Automatic Transcriptions (Stoekl Ben Ezra et al., 2025)'])
        ws.append(['Dataset: https://doi.org/10.5281/zenodo.17734473'])
        ws.append(['Citation: Stoekl Ben Ezra, D., Bambaci, L., Kiessling, B., Lapin, H., Ezer, N., Lolli, E., Rustow, M., Dershowitz, N., Kurar Barakat, B., Gogawale, S., Shmidman, A., Lavee, M., Siew, T., Raziel Kretzmer, V., Vasyutinsky Shapira, D., Olszowy-Schlanger, J., & Gila, Y. (2025). MiDRASH Automatic Transcriptions. Zenodo. https://doi.org/10.5281/zenodo.17734473'])

        # Style credits section
        credits_font = Font(italic=True, color="555555")
        for row_idx in range(credits_start_row, ws.max_row + 1):
            cell = ws.cell(row=row_idx, column=1)
            cell.font = credits_font

        stream = io.BytesIO()
        wb.save(stream)
        stream.seek(0)

        return Response(
            content=stream.read(),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=parallels_results.xlsx"}
        )

    @app.get('/api/export/parallels/word')
    def export_parallels_word():
        """Export parallels results to Word with RTL support for Hebrew content."""
        # Read from global state (accessible from HTTP requests without session context)
        parallels_results = state.parallels_results or []
        filtered_results = state.parallels_filtered or []

        if not parallels_results and not filtered_results:
            return Response("No parallels results to export", status_code=400)

        doc = Document()
        doc.add_heading('Genizah Parallels Search Results', 0)

        def add_results_to_doc(results, start_idx, section_title=None):
            if section_title:
                doc.add_heading(section_title, 1)

            for idx, item in enumerate(results, start_idx):
                score = item.get('score', 0)
                raw_header = item.get('raw_header', '')

                # Extract metadata
                shelfmark = 'Unknown'
                title = ''

                if raw_header and state.meta_mgr:
                    try:
                        import re
                        sys_match = re.search(r'(99\d{8,})', raw_header)
                        if sys_match:
                            sys_id = sys_match.group(1)
                            shelf_temp, title_temp = state.meta_mgr.get_meta_for_id(sys_id)
                            shelfmark = shelf_temp or shelfmark
                            title = title_temp or ''
                    except Exception:
                        pass

                # Header with shelfmark and score
                p = doc.add_paragraph()
                shelf_run = p.add_run(f"{idx}. {shelfmark}")
                shelf_run.bold = True
                _set_run_rtl_font(shelf_run)
                p.add_run(f" - Score: {score}")
                _set_paragraph_rtl(p)

                # Title (Hebrew)
                if title:
                    title_p = doc.add_paragraph()
                    title_run = title_p.add_run(title)
                    _set_run_rtl_font(title_run)
                    _set_paragraph_rtl(title_p)

                # Source Context label
                label_p1 = doc.add_paragraph()
                label_run1 = label_p1.add_run("Source Context:")
                label_run1.bold = True

                # Source Context text (Hebrew)
                source_p = doc.add_paragraph()
                source_run = source_p.add_run(item.get('source_ctx', '').replace('*', ''))
                _set_run_rtl_font(source_run)
                _set_paragraph_rtl(source_p)

                # Manuscript Match label
                label_p2 = doc.add_paragraph()
                label_run2 = label_p2.add_run("Manuscript Match:")
                label_run2.bold = True

                # Manuscript Match text (Hebrew)
                match_p = doc.add_paragraph()
                match_run = match_p.add_run(item.get('text', '').replace('*', ''))
                _set_run_rtl_font(match_run)
                _set_paragraph_rtl(match_p)

                doc.add_paragraph("_" * 60)

            return start_idx + len(results)

        # Add main results
        next_idx = add_results_to_doc(parallels_results, 1)

        # Add filtered results
        if filtered_results:
            add_results_to_doc(filtered_results, next_idx, "Filtered Results (found in source texts)")

        # Add credits
        doc.add_page_break()
        doc.add_heading('Credits', 1)
        doc.add_paragraph('Generated by Dicta Genizah Search (Web Version)')
        doc.add_paragraph('Data Source: MiDRASH Automatic Transcriptions (Stoekl Ben Ezra et al., 2025)')
        doc.add_paragraph('Dataset available at: https://doi.org/10.5281/zenodo.17734473')
        doc.add_paragraph()
        citation = doc.add_paragraph('Full Citation: ')
        citation.add_run('Stoekl Ben Ezra, D., Bambaci, L., Kiessling, B., Lapin, H., Ezer, N., Lolli, E., Rustow, M., Dershowitz, N., Kurar Barakat, B., Gogawale, S., Shmidman, A., Lavee, M., Siew, T., Raziel Kretzmer, V., Vasyutinsky Shapira, D., Olszowy-Schlanger, J., & Gila, Y. (2025). MiDRASH Automatic Transcriptions. Zenodo. https://doi.org/10.5281/zenodo.17734473').italic = True

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        return Response(
            content=stream.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=parallels_results.docx"}
        )

    @app.get('/api/export/browse/word')
    def export_browse_word():
        """Export current browse page to Word with RTL support for Hebrew content."""
        from nicegui import app as nicegui_app

        browse_data = nicegui_app.storage.user.get('browse_export_data')
        if not browse_data:
            return Response("No browse data to export", status_code=400)

        doc = Document()
        doc.add_heading('Genizah Manuscript', 0)

        # Manuscript info - shelfmark (Hebrew)
        if browse_data.get('shelfmark'):
            shelf_heading = doc.add_heading(browse_data['shelfmark'], 1)
            _set_paragraph_rtl(shelf_heading)
            for run in shelf_heading.runs:
                _set_run_rtl_font(run)

        # Title (Hebrew)
        if browse_data.get('title'):
            title_p = doc.add_paragraph()
            title_run = title_p.add_run(browse_data['title'])
            _set_run_rtl_font(title_run)
            _set_paragraph_rtl(title_p)

        # System ID (LTR)
        if browse_data.get('sys_id'):
            doc.add_paragraph(f"System ID: {browse_data['sys_id']}")

        doc.add_paragraph("_" * 60)

        # Text content (Hebrew manuscript text)
        if browse_data.get('view_all') and browse_data.get('pages'):
            doc.add_heading('Full Manuscript', 2)
            for page_data in browse_data['pages']:
                doc.add_heading(f"Page {page_data.get('p_num', '?')}", 3)
                if page_data.get('text'):
                    text_p = doc.add_paragraph()
                    text_run = text_p.add_run(page_data['text'])
                    _set_run_rtl_font(text_run)
                    _set_paragraph_rtl(text_p)
        elif browse_data.get('text'):
            doc.add_heading(f"Page {browse_data.get('p_num', '?')}", 2)
            text_p = doc.add_paragraph()
            text_run = text_p.add_run(browse_data['text'])
            _set_run_rtl_font(text_run)
            _set_paragraph_rtl(text_p)

        # Add credits
        doc.add_page_break()
        doc.add_heading('Credits', 1)
        doc.add_paragraph('Generated by Dicta Genizah Search (Web Version)')
        doc.add_paragraph('Data Source: MiDRASH Automatic Transcriptions (Stoekl Ben Ezra et al., 2025)')
        doc.add_paragraph('Dataset available at: https://doi.org/10.5281/zenodo.17734473')
        doc.add_paragraph()
        citation = doc.add_paragraph('Full Citation: ')
        citation.add_run('Stoekl Ben Ezra, D., Bambaci, L., Kiessling, B., Lapin, H., Ezer, N., Lolli, E., Rustow, M., Dershowitz, N., Kurar Barakat, B., Gogawale, S., Shmidman, A., Lavee, M., Siew, T., Raziel Kretzmer, V., Vasyutinsky Shapira, D., Olszowy-Schlanger, J., & Gila, Y. (2025). MiDRASH Automatic Transcriptions. Zenodo. https://doi.org/10.5281/zenodo.17734473').italic = True

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        return Response(
            content=stream.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=manuscript.docx"}
        )

    @app.get('/api/export/list/{list_id}/excel')
    def export_list_excel(list_id: str):
        """Export a specific list to Excel."""
        if not state.lists_mgr:
            return Response("Lists manager not available", status_code=400)

        list_data = state.lists_mgr.data.get('lists', {}).get(list_id)
        if not list_data:
            return Response("List not found", status_code=404)

        # Use get_items_in_list() - items are stored in data['items'] with list membership
        items = state.lists_mgr.get_items_in_list(list_id)
        if not items:
            return Response("List is empty", status_code=400)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = list_data.get('name', 'List')[:31]  # Excel sheet name limit

        # Enable RTL for the sheet (Hebrew content)
        ws.sheet_view.rightToLeft = True

        headers = ["#", "Shelfmark", "Title", "System ID", "FL ID", "Notes", "Added"]
        ws.append(headers)

        # Style header row
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment

        # Define column widths
        column_widths = {
            'A': 6,    # #
            'B': 25,   # Shelfmark
            'C': 35,   # Title
            'D': 18,   # System ID
            'E': 12,   # FL ID
            'F': 40,   # Notes
            'G': 18,   # Added
        }
        for col_letter, width in column_widths.items():
            ws.column_dimensions[col_letter].width = width

        # Alignments for different column types
        rtl_alignment = Alignment(horizontal="right", vertical="top", wrap_text=True, readingOrder=2)
        ltr_alignment = Alignment(horizontal="left", vertical="top")
        center_alignment = Alignment(horizontal="center", vertical="top")

        for idx, item in enumerate(items, 1):
            sys_id = item.get('sys_id', '')
            shelfmark = 'Unknown'
            title = ''

            if sys_id and state.meta_mgr:
                try:
                    shelf_temp, title_temp = state.meta_mgr.get_meta_for_id(sys_id)
                    shelfmark = shelf_temp or shelfmark
                    title = title_temp or ''
                except Exception:
                    pass

            # Clean notes - remove line breaks for single line display
            notes = item.get('note', '').replace('\n', ' ').replace('\r', ' ')
            while '  ' in notes:
                notes = notes.replace('  ', ' ')
            notes = notes.strip()

            row = [
                idx,
                shelfmark,
                title,
                sys_id,
                item.get('fl_id', ''),
                notes,
                item.get('added', '')
            ]

            # Sanitize for illegal chars
            clean_row = []
            for cell in row:
                if isinstance(cell, str):
                    cell = "".join(ch for ch in cell if (0x20 <= ord(ch) <= 0xD7FF) or (0xE000 <= ord(ch) <= 0xFFFD) or ch in "\t")
                clean_row.append(cell)
            ws.append(clean_row)

            # Apply alignment to the row just added
            current_row = ws.max_row
            ws.cell(row=current_row, column=1).alignment = center_alignment  # #
            ws.cell(row=current_row, column=2).alignment = rtl_alignment     # Shelfmark
            ws.cell(row=current_row, column=3).alignment = rtl_alignment     # Title
            ws.cell(row=current_row, column=4).alignment = ltr_alignment     # System ID
            ws.cell(row=current_row, column=5).alignment = ltr_alignment     # FL ID
            ws.cell(row=current_row, column=6).alignment = rtl_alignment     # Notes (could be Hebrew)
            ws.cell(row=current_row, column=7).alignment = ltr_alignment     # Added

        # Add credits at the bottom
        ws.append([])
        ws.append([])
        credits_start_row = ws.max_row + 1
        ws.append(['Credits'])
        ws.append(['Generated by Dicta Genizah Search (Web Version)'])
        ws.append(['Data Source: MiDRASH Automatic Transcriptions (Stoekl Ben Ezra et al., 2025)'])
        ws.append(['Dataset: https://doi.org/10.5281/zenodo.17734473'])
        ws.append(['Citation: Stoekl Ben Ezra, D., Bambaci, L., Kiessling, B., Lapin, H., Ezer, N., Lolli, E., Rustow, M., Dershowitz, N., Kurar Barakat, B., Gogawale, S., Shmidman, A., Lavee, M., Siew, T., Raziel Kretzmer, V., Vasyutinsky Shapira, D., Olszowy-Schlanger, J., & Gila, Y. (2025). MiDRASH Automatic Transcriptions. Zenodo. https://doi.org/10.5281/zenodo.17734473'])

        # Style credits section
        credits_font = Font(italic=True, color="555555")
        for row_idx in range(credits_start_row, ws.max_row + 1):
            cell = ws.cell(row=row_idx, column=1)
            cell.font = credits_font

        stream = io.BytesIO()
        wb.save(stream)
        stream.seek(0)

        filename = f"{list_data.get('name', 'list').replace(' ', '_')}.xlsx"
        return Response(
            content=stream.read(),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
